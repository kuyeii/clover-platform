import docx
doc = docx.Document()
doc.add_paragraph("Test with image:")
# We need an image.
import urllib.request
urllib.request.urlretrieve("https://dummyimage.com/100x100/000/fff", "dummy.png")
doc.add_picture("dummy.png")
doc.save("test_img.docx")
print("Saved with image")
