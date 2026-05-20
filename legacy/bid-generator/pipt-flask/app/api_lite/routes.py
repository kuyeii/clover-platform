# -*- coding: utf-8 -*-
"""
pipt-lite FastAPI 路由
仅暴露 NER 识别 + 脱敏 接口
"""
from __future__ import annotations

import logging
import os
import re
import json
import math
import html as _html
from pathlib import Path

# 加载项目根目录的 .env 文件
PRO_ENGINE_ROOT = Path(__file__).parent.parent.parent.parent
ROOT_ENV_PATH = PRO_ENGINE_ROOT / ".env"
try:
    from dotenv import load_dotenv
    load_dotenv(ROOT_ENV_PATH, override=False)  # override=False: 已有环境变量不覆盖
except ImportError:
    pass  # python-dotenv 未安装时不影响运行

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form, Body
from sqlalchemy.orm import Session
import yaml
import httpx
from .schemas import (
    RecognizeRequest, RecognizeResponse, EntityItem,
    DesensitizeRequest, DesensitizeResponse,
    BatchDesensitizeRequest, BatchDesensitizeResponse,
    RestoreRequest, RestoreResponse
)
from .docanalysis_protocol import (
    build_docanalysis_groups,
    build_docanalysis_node_index,
    build_docanalysis_system_prompt,
    extract_docanalysis_node_content,
    extract_docanalysis_text_output,
    load_docanalysis_framework,
    parse_bid_attachments_payload,
    parse_docanalysis_result_map,
    split_bid_attachments_tag,
)
from .engine import DesensitizeEngine
from .database import get_db, MappingRecord, ProjectRecord, SessionLocal
from .content_placeholder_resolve import find_illegal_pipt_bidder_placeholders, resolve_body_placeholders

logger = logging.getLogger(__name__)

router = APIRouter(tags=["脱敏服务"])

# SVG 图表生成当前产品态禁用：后端必须兜底，避免旧前端或脏配置继续触发图表工作流。
DIAGRAM_GENERATION_ENABLED = False

# 脱敏引擎单例
_engine: DesensitizeEngine | None = None


def get_engine() -> DesensitizeEngine:
    """获取或初始化脱敏引擎单例"""
    global _engine
    if _engine is None:
        _engine = DesensitizeEngine()
        logger.info("脱敏引擎初始化完成")
    return _engine

def load_profile_config(profile_name: str) -> dict:
    """从外部 config.yaml 加载脱敏配置"""
    try:
        config_path = Path(__file__).parent.parent.parent.parent / "config.yaml"
        if not config_path.exists():
            logger.warning(f"配置文件 {config_path} 不存在，使用硬编码默认值")
            return {}
        
        with open(config_path, "r", encoding="utf-8") as f:
            config = yaml.safe_load(f)
            
        profiles = config.get("pipt", {}).get("profiles", {})
        if profile_name in profiles:
            return profiles[profile_name]
        # 如果找不到指定 profile，尝试返回 default
        return profiles.get("default", {})
    except Exception as e:
        logger.error(f"解析 config.yaml 失败: {e}")
        return {}


def _read_root_env_value(env_var: str) -> str:
    """从项目根目录 .env 兜底读取单个变量，避免启动目录差异导致工作流误判未配置。"""
    try:
        if not ROOT_ENV_PATH.exists():
            return ""
        with open(ROOT_ENV_PATH, "r", encoding="utf-8") as f:
            for raw_line in f:
                line = raw_line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                key, value = line.split("=", 1)
                if key.strip() != env_var:
                    continue
                return value.strip().strip('"').strip("'")
    except Exception as e:
        logger.warning(f"读取根目录 .env 失败: {e}")
    return ""


def _get_workflow_key_with_source(workflow_name: str) -> tuple[str, str]:
    """
    从环境变量读取工作流 API Key，环境变量为空时兜底读取项目根目录 .env。
    环境变量命名规则: DIFY_WORKFLOW_{WORKFLOW_NAME.upper()}
    示例: requirement_extractor 对应 DIFY_WORKFLOW_REQUIREMENT_EXTRACTOR
    """
    env_var = f"DIFY_WORKFLOW_{workflow_name.upper()}"
    key = os.environ.get(env_var, "").strip()
    if key:
        return key, "process_env"
    key = _read_root_env_value(env_var)
    if key:
        return key, "root_env_file"
    return "", "missing"


def _get_workflow_key(workflow_name: str) -> str:
    key, _source = _get_workflow_key_with_source(workflow_name)
    if not key:
        env_var = f"DIFY_WORKFLOW_{workflow_name.upper()}"
        logger.warning(f"工作流 [{workflow_name}] 的 API Key 未配置，请在 .env 中设置 {env_var}")
    return key


def _resolve_content_workflow_name(generation_strategy: str = "") -> str:
    strategy = str(generation_strategy or "").strip().lower()
    if strategy == "response_special":
        return "response_content_writer"
    return "content_writer"


@router.get("/config/workflow-status", summary="工作流配置状态查询")
async def workflow_status():
    """返回工作流配置状态，并标记是否属于当前 manifest 纳管范围。"""
    workflows = [
        ("structure_generator", "DIFY_WORKFLOW_STRUCTURE_GENERATOR", "大纲生成", True, "managed"),
        ("content_writer", "DIFY_WORKFLOW_CONTENT_WRITER", "单章节内容生成", True, "managed"),
        ("content_group_writer", "DIFY_WORKFLOW_CONTENT_GROUP_WRITER", "H2分组正文生成", True, "managed"),
        ("content_rewrite", "DIFY_WORKFLOW_CONTENT_REWRITE", "单章节重生成", True, "managed"),
        ("response_content_writer", "DIFY_WORKFLOW_RESPONSE_CONTENT_WRITER", "响应情况正文生成", True, "managed"),
        ("diagram_generator", "DIFY_WORKFLOW_DIAGRAM_GENERATOR", "图表生成", True, "managed"),
        ("doc_analysis", "DIFY_WORKFLOW_DOC_ANALYSIS", "文档分析", True, "managed"),
        ("requirement_extractor", "DIFY_WORKFLOW_REQUIREMENT_EXTRACTOR", "需求提取", False, "legacy"),
        ("blueprint_generator", "DIFY_WORKFLOW_BLUEPRINT_GENERATOR", "全局策略蓝图", False, "legacy"),
        ("group_review_writer", "DIFY_WORKFLOW_GROUP_REVIEW_WRITER", "H2章节评估", False, "legacy"),
        ("attachment_generator", "DIFY_WORKFLOW_ATTACHMENT_GENERATOR", "智能附件生成", False, "legacy"),
        ("scoring_assistant", "DIFY_WORKFLOW_SCORING_ASSISTANT", "评分AI助手", False, "legacy"),
    ]
    status = {}
    for name, env_var, label, managed, lifecycle in workflows:
        key, source = _get_workflow_key_with_source(name)
        configured = bool(key)
        source_value = source
        if name == "diagram_generator" and not DIAGRAM_GENERATION_ENABLED:
            configured = False
            source_value = "disabled"
        status[name] = {
            "label": label,
            "env_var": env_var,
            "configured": configured,
            "source": source_value,
            "managed": managed,
            "lifecycle": lifecycle,
        }
    return status


@router.get("/config/analysis-framework", summary="获取预设解析框架配置")
async def get_analysis_framework():
    """返回预设解析框架 JSON（含每个节点的 extractionPrompt）"""
    import json
    config_path = Path(__file__).parent.parent.parent / "config" / "analysis_framework.json"
    if not config_path.exists():
        raise HTTPException(status_code=404, detail="analysis_framework.json 配置文件不存在")
    with open(config_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data


@router.post("/recognize", response_model=RecognizeResponse, summary="NER 识别")
async def recognize(request: RecognizeRequest):
    """
    对输入文本进行命名实体识别（NER），返回识别到的敏感实体。
    不修改原文，仅返回识别结果。
    """
    try:
        engine = get_engine()
        entities = engine.recognize(request.text, request.target_entities)
        return RecognizeResponse(
            entities=entities,
            entity_count=len(entities),
        )
    except Exception as e:
        logger.error(f"NER 识别失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"识别失败: {str(e)}")


@router.post("/desensitize", response_model=DesensitizeResponse, summary="文本脱敏")
async def desensitize(request: DesensitizeRequest, db: Session = Depends(get_db)):
    """
    对输入文本进行 NER 识别并脱敏处理。
    占位符通过 EntityRegistry 进行全局规范化，同一实体跨文件始终占位符相同。
    """
    try:
        engine = get_engine()

        # 应用 Profile 配置
        target_entities = request.target_entities
        method = request.method
        if target_entities is None or method is None:
            profile_config = load_profile_config(request.profile)
            if target_entities is None:
                target_entities = profile_config.get("target_entities", ["name", "phone", "id_number", "email", "addr", "bank", "car_id", "ip", "org"])
            if method is None:
                method = profile_config.get("method", "mask")

        import asyncio
        result = await asyncio.to_thread(
            engine.desensitize,
            text=request.text,
            target_entities=target_entities,
            method=method,
            placeholder_format=request.placeholder_format,
            db_session=db,
            llm_mode=getattr(request, 'llm_mode', None),  # 覆盖 LLM 模式
        )

        # 同时保留 session 级 MappingRecord，兼容旧还原接口
        if request.session_id and result.mapping_table:
            for placeholder, original in result.mapping_table.items():
                match = re.search(r"__PIPT_([a-z_]+)_", placeholder)
                entity_type = match.group(1) if match else "unknown"
                db.add(MappingRecord(
                    session_id=request.session_id,
                    placeholder=placeholder,
                    original_text=original,
                    entity_type=entity_type
                ))
            db.commit()

        return result
    except Exception as e:
        logger.error(f"脱敏失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"脱敏失败: {str(e)}")



@router.post("/desensitize/batch", response_model=BatchDesensitizeResponse, summary="批量脱敏")
async def batch_desensitize(request: BatchDesensitizeRequest, db: Session = Depends(get_db)):
    """
    对多段文本批量执行脱敏处理。
    共享同一个映射表计数器，确保占位符全局唯一。
    """
    try:
        engine = get_engine()
        
        # 应用 Profile 配置
        target_entities = request.target_entities
        method = request.method
        if target_entities is None or method is None:
            profile_config = load_profile_config(request.profile)
            if target_entities is None:
                target_entities = profile_config.get("target_entities", ["name", "phone", "id_number", "email", "addr", "bank", "car_id", "ip", "org"])
            if method is None:
                method = profile_config.get("method", "mask")
                
        results = []
        total_count = 0

        global_mappings = {}

        import asyncio
        for text in request.texts:
            result = await asyncio.to_thread(
                engine.desensitize,
                text=text,
                target_entities=target_entities,
                method=method,
                placeholder_format=request.placeholder_format,
                llm_mode=getattr(request, 'llm_mode', None),
            )
            total_count += result.entity_count
            results.append(result)
            
            # 合并映射表
            global_mappings.update(result.mapping_table)

        # 批量入库
        if request.session_id and global_mappings:
            for placeholder, original in global_mappings.items():
                match = re.search(r"__PIPT_([a-z_]+)_", placeholder)
                entity_type = match.group(1) if match else "unknown"
                record = MappingRecord(
                    session_id=request.session_id,
                    placeholder=placeholder,
                    original_text=original,
                    entity_type=entity_type
                )
                db.add(record)
            db.commit()

        return BatchDesensitizeResponse(
            results=results,
            total_entity_count=total_count,
        )
    except Exception as e:
        db.rollback()
        logger.error(f"批量脱敏失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"批量脱敏失败: {str(e)}")


@router.get("/entities", summary="获取支持的实体类型")
async def get_supported_entities():
    """返回系统支持的所有可识别实体类型"""
    from app.extension.celery_task.pipt_task.assets.constant import IDENTIFY_INFO_TO_CHINESE
    return {
        "entities": IDENTIFY_INFO_TO_CHINESE,
        "description": "key 为实体标识符，value 为中文名称",
    }


@router.post("/restore", response_model=RestoreResponse, summary="文本还原")
async def restore_text(request: RestoreRequest, db: Session = Depends(get_db)):
    """
    将占位符还原为明文。
    优先查全局 EntityRegistry（无需 session_id）；
    回退查 MappingRecord（兼容旧 session 级数据）。
    """
    import re as _re
    from app.api_lite.database import EntityRegistry, FernetEncryptor

    # 正则提取所有占位符
    placeholders = list(set(_re.findall(r'\{\{__PIPT_[a-z_]+_\d+__\}\}', request.text)))
    if not placeholders:
        return RestoreResponse(restored_text=request.text, restored_count=0)

    enc = FernetEncryptor.get()
    mapping: dict[str, str] = {}

    # 优先查全局注册表
    global_rows = db.query(EntityRegistry).filter(
        EntityRegistry.placeholder.in_(placeholders)
    ).all()
    for row in global_rows:
        mapping[row.placeholder] = enc.decrypt(row.original_text_enc)

    # 未命中的占位符回退查 MappingRecord（兼容旧数据）
    missing = [p for p in placeholders if p not in mapping]
    if missing and request.session_id:
        old_rows = db.query(MappingRecord).filter(
            MappingRecord.session_id == request.session_id,
            MappingRecord.placeholder.in_(missing)
        ).all()
        for row in old_rows:
            mapping[row.placeholder] = row.original_text

    restored_text = request.text
    count = 0
    for p, orig in mapping.items():
        if p in restored_text:
            restored_text = restored_text.replace(p, orig)
            count += 1

    logger.info(f"复原完成: {count} 处占位符（全局 {len(global_rows)} + session 回退 {count - len(global_rows)}）")
    return RestoreResponse(restored_text=restored_text, restored_count=count)


from .schemas import TemplateConfigResponse, UpdateTemplateRequest, UpdateConfigRequest

# 获取项目根目录 (根据实际层级向上推)
PRO_ENGINE_ROOT = Path(__file__).parent.parent.parent.parent

@router.get("/config/template", response_model=TemplateConfigResponse, summary="获取系统配置和大纲模板")
async def get_template_and_config(template_name: str = ""):
    """获取主要配置与指定名称的大纲模板，同时返回所有可用的模板列表"""
    try:
        config_path = PRO_ENGINE_ROOT / "config.yaml"
        templates_dir = PRO_ENGINE_ROOT / "data/templates/structures"
        
        # 获取所有模板名称
        available_templates: list[str] = []
        if templates_dir.exists():
            available_templates = sorted(f.name for f in templates_dir.glob("*.yaml") if f.is_file())

        template_name = str(template_name or "").strip()
        if template_name and ("/" in template_name or "\\" in template_name):
            raise HTTPException(status_code=400, detail="Invalid template name")
        if template_name and template_name not in available_templates:
            raise HTTPException(status_code=404, detail=f"Template not found: {template_name}")
        if not template_name and available_templates:
            template_name = available_templates[0]
            
        template_path = templates_dir / template_name if template_name else None
        
        with open(config_path, "r", encoding="utf-8") as f:
            config_dict = yaml.safe_load(f) or {}
            
        template_dict = {}
        if template_path and template_path.is_file():
            with open(template_path, "r", encoding="utf-8") as f:
                template_dict = yaml.safe_load(f) or {}
            
        return TemplateConfigResponse(
            config_dict=config_dict,
            template_dict=template_dict,
            available_templates=available_templates,
            current_template=template_name
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to read config: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/config/template", summary="删除大纲模板")
async def delete_template(template_name: str):
    """删除指定的 yaml 模板"""
    try:
        if template_name in ["standard.yaml", ""]:
            raise ValueError(f"Cannot delete pre-configured template: {template_name}")
            
        if "/" in template_name or "\\" in template_name:
            raise ValueError("Invalid template name")
            
        template_path = PRO_ENGINE_ROOT / "data/templates/structures" / template_name
        if template_path.exists():
            template_path.unlink()
            
        return {"status": "success", "message": f"Template {template_name} deleted successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.put("/config/template", summary="更新大纲模板")
async def update_template(request: UpdateTemplateRequest):
    """保存指定的 yaml 模板"""
    try:
        template_name = request.template_name or ""
        if not template_name:
            raise ValueError("Template name cannot be empty")
        # 防止目录穿越
        if "/" in template_name or "\\" in template_name:
            raise ValueError("Invalid template name")
            
        template_path = PRO_ENGINE_ROOT / "data/templates/structures" / template_name
        template_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(template_path, "w", encoding="utf-8") as f:
            yaml.dump(request.template_dict, f, allow_unicode=True, sort_keys=False)
        return {"status": "success", "message": f"Template {template_name} updated successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.put("/config/global", summary="更新系统配置")
async def update_config(request: UpdateConfigRequest):
    """保存 config.yaml"""
    try:
        config_path = PRO_ENGINE_ROOT / "config.yaml"
        with open(config_path, "w", encoding="utf-8") as f:
            yaml.dump(request.config_dict, f, allow_unicode=True, sort_keys=False)
        return {"status": "success", "message": "Config updated successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

from .schemas import GenerateStructureRequest, GenerateStructureResponse

@router.post("/config/template/generate", response_model=GenerateStructureResponse, summary="动态生成专属标书架构")
async def generate_template_architecture(request: GenerateStructureRequest):
    """调用外部系统（如 Dify）通过临时脚本动态生成项目的 YAML 结构配置"""
    try:
        import sys, subprocess, os
        from tempfile import NamedTemporaryFile
        cwd_dify = PRO_ENGINE_ROOT / "dify-bridge"
        
        prompt_str = f"你是一个资深售前解决方案架构师。请针对当前项目【{request.project_name}】，结合蓝图与需求，生成一份专属的标书结构目录YAML配置。产出格式必须符合系统标准的 blocks 数组结构，只输出合法的YAML。"
        
        prompt_file = cwd_dify / ".tmp_gen_struct_p.txt"
        data_file = cwd_dify / ".tmp_gen_struct_d.txt"
        out_file = cwd_dify / ".tmp_gen_struct_o.txt"
        
        with open(prompt_file, "w", encoding="utf-8") as f: 
            f.write(prompt_str + "\n\n" + request.blueprint)
        with open(data_file, "w", encoding="utf-8") as f: 
            f.write(request.structured_data)
            
        runner_dify = f'''
import sys, asyncio
sys.path.insert(0, "{cwd_dify}")
from src.config import DifyConfig
from src.workflow import WorkflowManager
def main():
    config = DifyConfig.from_yaml("{PRO_ENGINE_ROOT / 'config.yaml'}")
    manager = WorkflowManager(config)
    with open("{prompt_file}", "r", encoding="utf-8") as f: prompt_str = f.read()
    with open("{data_file}", "r", encoding="utf-8") as f: data_str = f.read()
    res = asyncio.run(manager.run_bid_generation(
        system_prompt=prompt_str, 
        structured_data=data_str, 
        knowledge_query="{request.project_name} 目录架构搭建",
        requires_search="false"
    ))
    with open("{out_file}", "w", encoding="utf-8") as f: f.write(res)
if __name__ == "__main__": main()
'''
        with NamedTemporaryFile("w", suffix=".py", delete=False, encoding="utf-8") as f2:
            f2.write(runner_dify)
            temp_path_dify = f2.name
            
        subprocess.run([sys.executable, temp_path_dify], cwd=cwd_dify, check=True)
        
        with open(out_file, "r", encoding="utf-8") as f:
            res_md = f.read().strip()
            
        # extract yaml
        match = re.search(r'```(?:yaml)?(.*?)```', res_md, re.DOTALL)
        yaml_str = match.group(1).strip() if match else res_md
        
        structure_dict = yaml.safe_load(yaml_str)
        if not isinstance(structure_dict, dict) or "blocks" not in structure_dict:
            # fallback wrapper
            structure_dict = {
                "name": f"{request.project_name}专属架构",
                "id": f"dynamic_struct_01",
                "blocks": structure_dict if isinstance(structure_dict, list) else []
            }
            
        for p in [temp_path_dify, prompt_file, data_file, out_file]:
            if os.path.exists(p): os.remove(p)
            
        return GenerateStructureResponse(structure_dict=structure_dict)
    except Exception as e:
        logger.error(f"Generate structure failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

from .schemas import ExtractRequirementsResponse, ExtractRequirementItem, GenerateOutlineRequest, GenerateOutlineResponse, OutlineSection
import io, json
import os
import requests
import base64
import tempfile
import uuid
import shutil
from pathlib import Path
import re

def _tag_image_with_vlm(image_path: str) -> str:
    """调用局域网内的 VLM 模型为提取出的图表/图片打标，支持失败重试"""
    try:
        vlm_api_url = os.environ.get("VLM_API_URL", "http://localhost:8000/v1/chat/completions")
        vlm_model = os.environ.get("VLM_MODEL", "qwen-vl-chat")
        # 本函数使用 OpenAI 兼容多模态 JSON（image_url），须指向 /v1/chat/completions，而非 Ollama 原生 /api/chat
        if "/api/chat" in vlm_api_url and "/v1/" not in vlm_api_url:
            logger.warning(
                "VLM_API_URL 指向 /api/chat，与当前 OpenAI 格式多模态请求不兼容，"
                "请改为例如 http://<host>:11435/v1/chat/completions"
            )
        # 可通过 .env 调整的参数
        vlm_timeout = int(os.environ.get("VLM_TIMEOUT", "120"))
        vlm_max_tokens = int(os.environ.get("VLM_MAX_TOKENS", "200"))
        max_retries = 2

        with open(image_path, "rb") as f:
            b64_img = base64.b64encode(f.read()).decode('utf-8')
            
        payload = {
            "model": vlm_model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "你是一个专业的解决方案架构解读助手。这是从标书/架构设计文档中截取的一张图片。\n\n请按以下准则提取这张图的核心信息，供下游的其他大语言模型作为事实依据参考：\n1. 【判断类型】首先指出这是一张什么图（如：系统架构图、网络拓扑图、业务流程图、数据表格、或其他配图）。\n2. 【拆解核心事实】如果是架构图/流程图，请依次列出它包含了哪些层级（如前端、后端、数据层），或者数据流向的关键节点。\n3. 【绝对禁止敏感词】如果图中出现了具体的人名、手机号、邮箱、真实的服务器 IP 地址、或者真实的甲方/乙方机构名称，请务必用“某人”、“某机构”、“[IP地址]”等泛指词替换，绝对不可将这类实体词照抄输出。\n4. 【精简干练】不要描述颜色等无关紧要的外观，直击技术或业务要点；整体字数控制在200字以内，力求让无法看到图片的外部AI通过你的描述就能脑补出图中的全貌。"},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64_img}"}}
                    ]
                }
            ],
            "max_tokens": 150
        }
        # 排队及多模态长文本推理较慢，放宽至 120s
        resp = requests.post(vlm_api_url, json=payload, timeout=120)
        if resp.status_code == 200:
            res_data = resp.json()
            content = res_data.get("choices", [{}])[0].get("message", {}).get("content", "")
            if content:
                return content.strip()
        return "无法提取图片描述"
    except Exception as e:
        logger.warning(f"VLM 打标失败 ({image_path}): {e}")
        return "图片描述提取失败"

# ───────── 通用工具函数 ─────────


# ─── PDF 分页文本索引（溯源用） ─────────────────────────────────────────
def _extract_pdf_pages_text(content_bytes: bytes) -> list[dict]:
    """
    利用 PyMuPDF 按页提取 PDF 文本，返回每页的文本内容和页码。
    用于后续将 Dify 提取的需求溯源到原文所在的页码。
    返回: [{"page": 0, "text": "..."}, ...]
    """
    pages_text = []
    try:
        import pymupdf
        doc = pymupdf.open(stream=content_bytes, filetype="pdf")
        for page_idx, page in enumerate(doc):
            text = page.get_text("text") or ""
            pages_text.append({"page": page_idx, "text": text})
        doc.close()
    except ImportError:
        logger.warning("PyMuPDF 未安装，无法进行分页文本索引")
    except Exception as e:
        logger.warning(f"PDF 分页文本提取异常: {e}")
    return pages_text


def _match_source_pages(excerpt: str, pages_text: list[dict], min_match_len: int = 8) -> list[dict]:
    """
    将 Dify 返回的原文摘录模糊匹配到 PDF 的具体页码。
    策略：截取摘录的前 min_match_len~50 个字符作为搜索关键词，
    在每页文本中进行子串搜索，返回所有命中页面的 SourcePageRef。
    """
    if not excerpt or not pages_text:
        return []

    # 清理摘录中的空白字符用于匹配
    clean_excerpt = excerpt.replace("\n", "").replace("\r", "").replace(" ", "").strip()
    if len(clean_excerpt) < min_match_len:
        return []

    # 取前50个字符作为搜索关键词（平衡精确度和容错）
    search_key = clean_excerpt[:50]

    matched_pages = []
    for page_info in pages_text:
        page_text_clean = page_info["text"].replace("\n", "").replace("\r", "").replace(" ", "")
        if search_key in page_text_clean:
            matched_pages.append({
                "page": page_info["page"],
                "excerpt": excerpt[:200],  # 截断过长的摘录
            })

    # 如果精确匹配失败，用更短的关键词再试一次
    if not matched_pages and len(clean_excerpt) >= min_match_len:
        short_key = clean_excerpt[:min_match_len]
        for page_info in pages_text:
            page_text_clean = page_info["text"].replace("\n", "").replace("\r", "").replace(" ", "")
            if short_key in page_text_clean:
                matched_pages.append({
                    "page": page_info["page"],
                    "excerpt": excerpt[:200],
                })

    return matched_pages


def _cache_pdf_file(project_id: str, content_bytes: bytes) -> str:
    """
    将上传的 PDF 文件缓存到 data/pdf_cache/{project_id}.pdf，
    返回可通过 API 访问的相对 URL。
    """
    cache_dir = PRO_ENGINE_ROOT / "data" / "pdf_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = cache_dir / f"{project_id}.pdf"
    pdf_path.write_bytes(content_bytes)
    logger.info(f"PDF 已缓存: {pdf_path}")
    return f"/api/projects/pdf/{project_id}"


def _convert_to_pdf_and_cache(project_id: str, content_bytes: bytes, filename: str) -> str:
    """
    将 DOC / DOCX 文件转换为 PDF 并缓存，返回预览 URL。
    转换策略（按优先级）：
      1. python-docx2pdf（Windows/macOS 依赖 Word，Linux 大概率失败）
      2. LibreOffice headless（Linux 首选，需要 libreoffice 已安装）
    DOCX 文件会先做排版预处理（两端/分散对齐 → 左对齐），改善 LibreOffice 渲染质量。
    失败时静默返回空字符串，不阻断主流程。
    """
    cache_dir = PRO_ENGINE_ROOT / "data" / "pdf_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    pdf_path = cache_dir / f"{project_id}.pdf"

    # 先将原始文件写入临时目录
    ext = (filename or "").lower().rsplit(".", 1)[-1]
    with tempfile.TemporaryDirectory() as tmp_dir:
        src_path = os.path.join(tmp_dir, f"source.{ext}")
        with open(src_path, "wb") as f:
            f.write(content_bytes)

        # ── DOCX 排版预处理：两端/分散对齐 → 左对齐 ────────────────────
        if ext == "docx":
            try:
                src_path = _preprocess_docx_alignment(src_path)
                logger.info("DOCX 排版预处理完成（两端/分散对齐 → 左对齐）")
            except Exception as _e:
                logger.warning(f"DOCX 排版预处理失败，使用原文件: {_e}")

        # ── 方案①：docx2pdf（仅适用于有 Word 的环境）──────────────
        try:
            import docx2pdf
            docx2pdf.convert(src_path, str(pdf_path))
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                logger.info(f"docx2pdf 转换成功: {pdf_path}")
                return f"/api/projects/pdf/{project_id}"
        except Exception as _e:
            logger.debug(f"docx2pdf 不可用: {_e}")

        # ── 方案②：LibreOffice headless（Linux 首选）────────────────
        try:
            import subprocess
            env = os.environ.copy()
            env["SAL_USE_VCLPLUGIN"] = "svp"
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--norestore",
                    "--nofirststartwizard",
                    "--convert-to", "pdf:writer_pdf_Export",
                    "--outdir", tmp_dir,
                    src_path,
                ],
                capture_output=True, timeout=120, env=env
            )
            # LibreOffice 输出文件名为 source.pdf
            lo_pdf = os.path.join(tmp_dir, "source.pdf")
            if result.returncode == 0 and os.path.exists(lo_pdf):
                import shutil
                shutil.copy(lo_pdf, str(pdf_path))
                logger.info(f"LibreOffice 转换成功: {pdf_path}")
                return f"/api/projects/pdf/{project_id}"
            else:
                logger.warning(f"LibreOffice 转换失败: {result.stderr.decode('utf-8', errors='replace')[:200]}")
        except FileNotFoundError:
            logger.warning("LibreOffice 未安装，DOCX→PDF 转换不可用")
        except Exception as _e:
            logger.warning(f"LibreOffice 转换异常: {_e}")

    logger.warning(f"所有 DOCX→PDF 转换方案均失败，project_id={project_id}")
    return ""


def _preprocess_docx_alignment(src_path: str) -> str:
    """
    预处理 DOCX 对齐属性，修复 LibreOffice 渲染时字间距异常。
    修复范围：
    - document.xml 中段落直接对齐值
    - styles.xml 中段落样式默认对齐值（解决“同文档部分段落仍异常”）
    对齐策略：
    - distribute / distributeLetter / both / justify / kashida 等 → left
    原地覆写文件，返回路径。
    """
    import os as _os
    import zipfile as _zipfile
    import tempfile as _tmp
    import shutil as _shutil
    import xml.etree.ElementTree as _ET
    from docx import Document

    WPC_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    BAD_ALIGNMENTS = {
        "distribute", "distributeLetter", "distributeAllLines", "thaiDistribute",
        "both", "justify", "lowKashida", "mediumKashida", "highKashida",
    }
    _ET.register_namespace("w", WPC_NS)
    w = {"w": WPC_NS}

    total_fixed_doc = 0
    doc = Document(src_path)

    def _fix_paragraphs(paragraphs):
        nonlocal total_fixed_doc
        for para in paragraphs:
            pPr = para._element.find(f"{{{WPC_NS}}}pPr")
            if pPr is None:
                continue
            jc = pPr.find(f"{{{WPC_NS}}}jc")
            if jc is None:
                continue
            val = str(jc.get(f"{{{WPC_NS}}}val", "") or "")
            if val in BAD_ALIGNMENTS:
                jc.set(f"{{{WPC_NS}}}val", "left")
                total_fixed_doc += 1

    _fix_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _fix_paragraphs(cell.paragraphs)
    doc.save(src_path)

    total_fixed_xml = 0
    with _zipfile.ZipFile(src_path, "r") as zin:
        entries = {info.filename: zin.read(info.filename) for info in zin.infolist()}
        infos = {info.filename: info for info in zin.infolist()}

    def _rewrite_jc(xml_bytes: bytes) -> bytes:
        nonlocal total_fixed_xml
        try:
            root = _ET.fromstring(xml_bytes)
        except Exception:
            return xml_bytes
        changed = False
        for jc in root.findall(".//w:jc", w):
            val = str(jc.get(f"{{{WPC_NS}}}val", "") or "")
            if val in BAD_ALIGNMENTS:
                jc.set(f"{{{WPC_NS}}}val", "left")
                total_fixed_xml += 1
                changed = True
        if not changed:
            return xml_bytes
        return _ET.tostring(root, encoding="utf-8", xml_declaration=True)

    for name in ("word/document.xml", "word/styles.xml"):
        if name in entries:
            entries[name] = _rewrite_jc(entries[name])

    fd, tmp_zip = _tmp.mkstemp(prefix="docx_align_fix_", suffix=".docx")
    _os.close(fd)
    try:
        with _zipfile.ZipFile(tmp_zip, "w") as zout:
            for name, data in entries.items():
                info = infos.get(name)
                if info is not None:
                    zout.writestr(info, data)
                else:
                    zout.writestr(name, data)
        _shutil.move(tmp_zip, src_path)
    finally:
        if _os.path.exists(tmp_zip):
            _os.remove(tmp_zip)

    if total_fixed_doc > 0 or total_fixed_xml > 0:
        logger.info(
            "DOCX 对齐预处理完成：段落直接属性修复 %s 处，XML样式修复 %s 处",
            total_fixed_doc,
            total_fixed_xml,
        )
    return src_path



def _extract_docx_with_tables(doc) -> str:
    """
    按文档 body 元素顺序提取 DOCX 全文（段落 + 表格），保持原始排列。
    表格转为 Markdown 格式（| col1 | col2 |），便于 LLM 理解结构化内容。
    """
    from lxml import etree

    WPC_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    PARA_TAG = f"{{{WPC_NS}}}p"
    TBL_TAG = f"{{{WPC_NS}}}tbl"
    ROW_TAG = f"{{{WPC_NS}}}tr"
    CELL_TAG = f"{{{WPC_NS}}}tc"

    def _cell_text(cell_elem) -> str:
        """提取单元格内所有段落文本，多段以空格连接"""
        texts = []
        for p in cell_elem.findall(f".//{PARA_TAG}"):
            t = "".join(node.text or "" for node in p.iter(f"{{{WPC_NS}}}t"))
            if t.strip():
                texts.append(t.strip())
        return " ".join(texts)

    def _table_to_markdown(tbl_elem) -> str:
        """将 w:tbl 元素转为 Markdown 表格"""
        rows = tbl_elem.findall(f".//{ROW_TAG}")
        if not rows:
            return ""
        md_rows = []
        for row in rows:
            cells = row.findall(f".//{CELL_TAG}")
            md_rows.append("| " + " | ".join(_cell_text(c) for c in cells) + " |")
        # 第一行后插入分隔线
        if len(md_rows) > 1:
            col_count = md_rows[0].count("|") - 1
            separator = "| " + " | ".join(["---"] * max(col_count, 1)) + " |"
            md_rows.insert(1, separator)
        return "\n".join(md_rows)

    # 遍历 body 直接子元素，按原始顺序
    body = doc.element.body
    parts = []
    for child in body:
        tag = child.tag
        if tag == PARA_TAG:
            t = "".join(node.text or "" for node in child.iter(f"{{{WPC_NS}}}t"))
            if t.strip():
                parts.append(t)
        elif tag == TBL_TAG:
            md = _table_to_markdown(child)
            if md:
                parts.append(f"\n{md}\n")

    result = "\n".join(parts)
    logger.info(f"DOCX 全文提取完成: {len(result)} 字符（含表格 Markdown）")
    return result


# ── 投标文件模块：段落定位符缓存（模块级 in-memory）──────────────────────────
# key: project_id → {"doc": docx.Document对象, "locator_map": {"P0045": 45, ...}, "doc_blocks": [...]}
_locator_cache: dict = {}


def _persist_doc_blocks_snapshot(project_id: str, doc_blocks: list[dict]) -> None:
    """
    将 doc_blocks 快照写入 projects.data，便于服务重启后至少恢复块索引列表。
    注意：仅恢复索引时无法执行真实 DOCX 切片。
    """
    if not project_id or not isinstance(doc_blocks, list):
        return
    db = SessionLocal()
    try:
        record = db.query(ProjectRecord).filter(ProjectRecord.id == project_id).first()
        if not record:
            return
        import json as _json
        data = _json.loads(record.data) if isinstance(record.data, str) else (record.data or {})
        if not isinstance(data, dict):
            data = {}
        data["__doc_blocks_cache"] = doc_blocks
        record.data = _json.dumps(data, ensure_ascii=False)
        db.commit()
    except Exception as e:
        logger.warning(f"[{project_id}] doc_blocks 快照持久化失败: {e}")
        db.rollback()
    finally:
        db.close()


def _load_doc_blocks_snapshot(project_id: str) -> list[dict]:
    """从 projects.data 中读取 doc_blocks 快照。"""
    if not project_id:
        return []
    db = SessionLocal()
    try:
        record = db.query(ProjectRecord).filter(ProjectRecord.id == project_id).first()
        if not record:
            return []
        import json as _json
        data = _json.loads(record.data) if isinstance(record.data, str) else (record.data or {})
        if not isinstance(data, dict):
            return []
        blocks = data.get("__doc_blocks_cache")
        return blocks if isinstance(blocks, list) else []
    except Exception as e:
        logger.warning(f"[{project_id}] 读取 doc_blocks 快照失败: {e}")
        return []
    finally:
        db.close()


def _build_locator_map_from_blocks(doc_blocks: list[dict]) -> dict:
    locator_map: dict[str, int] = {}
    for block in doc_blocks or []:
        if not isinstance(block, dict):
            continue
        locator = str(block.get("locator") or "").strip().upper()
        body_idx = block.get("body_idx")
        if not locator:
            continue
        try:
            locator_map[locator] = int(body_idx)
        except Exception:
            continue
    return locator_map


def _persist_docx_for_locators(project_id: str, content_bytes: bytes) -> None:
    """将原始 DOCX 持久化到磁盘，便于服务重启后恢复定位符缓存。"""
    if not project_id:
        return
    try:
        cache_dir = PRO_ENGINE_ROOT / "data" / "docx_cache"
        cache_dir.mkdir(parents=True, exist_ok=True)
        (cache_dir / f"{project_id}.docx").write_bytes(content_bytes)
    except Exception as e:
        logger.warning(f"[{project_id}] 持久化 DOCX 失败，定位缓存可能无法跨重启恢复: {e}")


def _persist_raw_document(project_id: str, raw_document: str) -> None:
    """持久化脱敏后的文本，供重提取与节点重算按 project_id 读取。"""
    if not project_id:
        return
    cache_dir = PRO_ENGINE_ROOT / "data" / "raw_doc_cache"
    cache_dir.mkdir(parents=True, exist_ok=True)
    (cache_dir / f"{project_id}.txt").write_text(raw_document or "", encoding="utf-8")


def _load_raw_document(project_id: str) -> str:
    """读取项目级缓存原文。"""
    if not project_id:
        return ""
    raw_path = PRO_ENGINE_ROOT / "data" / "raw_doc_cache" / f"{project_id}.txt"
    if not raw_path.exists():
        return ""
    try:
        return raw_path.read_text(encoding="utf-8")
    except Exception as e:
        logger.warning(f"[{project_id}] 读取 raw_document 缓存失败: {e}")
        return ""


def _restore_locator_cache_from_disk(project_id: str) -> bool:
    """当内存缓存缺失时，尝试从 data/docx_cache 重建定位符缓存。"""
    if not project_id:
        return False
    docx_path = PRO_ENGINE_ROOT / "data" / "docx_cache" / f"{project_id}.docx"
    if not docx_path.exists():
        # 兼容兜底：若原始 DOCX 不在磁盘，尝试用项目快照恢复块索引
        db_blocks = _load_doc_blocks_snapshot(project_id)
        if db_blocks:
            _locator_cache[project_id] = {
                "doc": None,
                "locator_map": _build_locator_map_from_blocks(db_blocks),
                "doc_blocks": db_blocks,
                "snapshot_only": True,
            }
            logger.info(f"[{project_id}] 已从项目快照恢复 doc_blocks: {len(db_blocks)} 个")
            return True
        return False
    try:
        import docx as _docx_mod
        _loc_doc = _docx_mod.Document(str(docx_path))
        _loc_text, _loc_map, _blocks = _extract_docx_with_locators(_loc_doc)
        _locator_cache[project_id] = {"doc": _loc_doc, "locator_map": _loc_map, "doc_blocks": _blocks}
        _persist_doc_blocks_snapshot(project_id, _blocks)
        logger.info(f"[{project_id}] 从磁盘恢复定位符缓存成功: {len(_loc_map)} 个定位点")
        return True
    except Exception as e:
        logger.warning(f"[{project_id}] 从磁盘恢复定位符缓存失败: {e}")
        return False


def _extract_docx_with_locators(doc) -> tuple:
    """
    带段落定位符变体：在每段/每表格前注入 [Pxxx] 前缀，
    供解析工作流 LLM 识别投标文件附件的起止位置。

    返回:
        text: str  — 带 [Pxxx] 定位符的完整文本
        locator_map: dict — {"P0045": 45}（定位符 → body element index）
        doc_blocks: list[dict] — 块级有序结构（block_id/locator/body_idx/text/type）
    """
    WPC_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    PARA_TAG = f"{{{WPC_NS}}}p"
    TBL_TAG  = f"{{{WPC_NS}}}tbl"
    ROW_TAG  = f"{{{WPC_NS}}}tr"
    CELL_TAG = f"{{{WPC_NS}}}tc"
    BR_TAG   = f"{{{WPC_NS}}}br"

    def _cell_text(cell_elem) -> str:
        texts = []
        for p in cell_elem.findall(f".//{PARA_TAG}"):
            t = "".join(node.text or "" for node in p.iter(f"{{{WPC_NS}}}t"))
            if t.strip():
                texts.append(t.strip())
        return " ".join(texts)

    def _table_to_markdown(tbl_elem) -> str:
        rows = tbl_elem.findall(f".//{ROW_TAG}")
        if not rows:
            return ""
        md_rows = []
        for row in rows:
            cells = row.findall(f".//{CELL_TAG}")
            md_rows.append("| " + " | ".join(_cell_text(c) for c in cells) + " |")
        if len(md_rows) > 1:
            col_count = md_rows[0].count("|") - 1
            separator = "| " + " | ".join(["---"] * max(col_count, 1)) + " |"
            md_rows.insert(1, separator)
        return "\n".join(md_rows)

    def _has_page_break(para_elem) -> bool:
        """检测段落内是否有显式分页符"""
        for br in para_elem.iter(BR_TAG):
            br_type = br.get(f"{{{WPC_NS}}}type", "")
            if br_type == "page":
                return True
        return False

    body = doc.element.body
    parts = []
    locator_map: dict = {}
    doc_blocks = []
    loc_idx = 0  # 全局定位符计数

    for body_idx, child in enumerate(body):
        tag = child.tag

        if tag == PARA_TAG:
            # 检查是否有显式分页符，先插入标记
            if _has_page_break(child):
                parts.append("---PAGE_BREAK---")

            t = "".join(node.text or "" for node in child.iter(f"{{{WPC_NS}}}t"))
            if t.strip():
                loc_key = f"P{loc_idx:04d}"
                locator_map[loc_key] = body_idx
                parts.append(f"[{loc_key}] {t}")
                doc_blocks.append({
                    "block_id": f"B{loc_idx:06d}",
                    "locator": loc_key,
                    "body_idx": body_idx,
                    "type": "paragraph",
                    "text": t,
                })
                loc_idx += 1

        elif tag == TBL_TAG:
            md = _table_to_markdown(child)
            if md:
                loc_key = f"P{loc_idx:04d}"
                locator_map[loc_key] = body_idx
                parts.append(f"[{loc_key}] \n{md}\n")
                doc_blocks.append({
                    "block_id": f"B{loc_idx:06d}",
                    "locator": loc_key,
                    "body_idx": body_idx,
                    "type": "table",
                    "text": md,
                })
                loc_idx += 1

    text = "\n".join(parts)
    logger.info(f"DOCX 带定位符提取完成: {len(text)} 字符，{loc_idx} 个定位点")
    return text, locator_map, doc_blocks


def _body_elements_to_html(doc, start_body_idx: int, end_body_idx: int) -> str:
    """
    将 docx body[start_body_idx:end_body_idx+1] 转为 HTML，
    保留段落样式（标题级别）和表格结构。并且提取图片入库。
    """
    WPC_NS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    PARA_TAG = f"{{{WPC_NS}}}p"
    TBL_TAG  = f"{{{WPC_NS}}}tbl"
    ROW_TAG  = f"{{{WPC_NS}}}tr"
    CELL_TAG = f"{{{WPC_NS}}}tc"
    STYLE_TAG = f"{{{WPC_NS}}}pStyle"
    B_TAG    = f"{{{WPC_NS}}}b"
    I_TAG    = f"{{{WPC_NS}}}i"
    U_TAG    = f"{{{WPC_NS}}}u"
    R_TAG    = f"{{{WPC_NS}}}r"
    RPR_TAG  = f"{{{WPC_NS}}}rPr"
    T_TAG    = f"{{{WPC_NS}}}t"
    PPR_TAG  = f"{{{WPC_NS}}}pPr"
    DRAWING_TAG = f"{{{WPC_NS}}}drawing"

    A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
    BLIP_TAG = f"{{{A_NS}}}blip"
    REL_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    EMBED_ATTR = f"{{{REL_NS}}}embed"


    HEADING_MAP = {
        "heading1": "h1", "heading2": "h2", "heading3": "h3",
        "heading4": "h4", "标题1": "h1", "标题2": "h2", "标题3": "h3",
    }

    def _run_to_html(run_elem) -> str:
        """将单个 run 转为内联 HTML（支持加粗/斜体/下划线）"""
        t = "".join(node.text or "" for node in run_elem.iter(T_TAG))
        if not t:
            return ""
        rpr = run_elem.find(RPR_TAG)
        if rpr is not None:
            def _is_on(tag) -> bool:
                """检测 w:b / w:i / w:u 是否激活（排除 val=false/0）"""
                elem = rpr.find(tag)
                if elem is None:
                    return False
                val = elem.get(f"{{{WPC_NS}}}val", "true").lower()
                return val not in ("false", "0", "none", "off")
            if _is_on(U_TAG):
                t = f"<u>{t}</u>"
            if _is_on(I_TAG):
                t = f"<em>{t}</em>"
            if _is_on(B_TAG):
                t = f"<strong>{t}</strong>"
        return t

    def _para_to_html(para_elem) -> str:
        """将段落转为 HTML，识别标题样式"""
        ppr = para_elem.find(PPR_TAG)
        tag = "p"
        if ppr is not None:
            style_elem = ppr.find(f".//{STYLE_TAG}")
            if style_elem is not None:
                style_val = style_elem.get(f"{{{WPC_NS}}}val", "").lower()
                tag = HEADING_MAP.get(style_val, "p")

        content = "".join(_run_to_html(r) for r in para_elem.findall(R_TAG))

        # 尝试提取段落内的图片 (drawing / blip)
        import hashlib, uuid, os
        from app.api_lite.database import SessionLocal, ImageRegistry
        
        drawings = para_elem.findall(f".//{DRAWING_TAG}")
        for dw in drawings:
            blips = dw.findall(f".//{BLIP_TAG}")
            for blip in blips:
                embed_id = blip.get(EMBED_ATTR)
                if embed_id and embed_id in doc.part.rels:
                    rel = doc.part.rels[embed_id]
                    if "image" in rel.reltype:
                        blob = rel.target_part.blob
                        img_hash = hashlib.md5(blob).hexdigest()
                        img_name = f"{img_hash}_{rel.target_part.partname.split('/')[-1]}"
                        
                        permanent_img_dir = PRO_ENGINE_ROOT / "data" / "extracted_images"
                        permanent_img_dir.mkdir(parents=True, exist_ok=True)
                        dest_img = permanent_img_dir / img_name
                        if not dest_img.exists():
                            with open(dest_img, "wb") as f:
                                f.write(blob)
                        
                        placeholder = f"__PRO_IMG_{img_hash}__"
                        # 存入数据库
                        db = SessionLocal()
                        try:
                            if not db.query(ImageRegistry).filter_by(image_hash=img_hash).first():
                                row = ImageRegistry(
                                    image_hash=img_hash,
                                    project_id=None, # 解析独立文档无 project_id 时传入 None
                                    abs_path=str(dest_img),
                                    preview_url=f"/api/extracted-images/{img_name}",
                                    placeholder=placeholder,
                                    vlm_caption=None
                                )
                                db.add(row)
                                db.commit()
                        except Exception as e:
                            db.rollback()
                            logger.error(f"ImageRegistry 写入失败: {e}")
                        finally:
                            db.close()
                            
                        # 在 HTML 内容末尾追加占位符，外部大模型会原样带入并被 RAG 知识库切片
                        content += f" <br/>![本地配图]({placeholder})<br/> "
                        
        content = content.strip()
        return f"<{tag}>{content}</{tag}>" if content else ""

    def _table_to_html(tbl_elem) -> str:
        """将 w:tbl 转为 HTML table"""
        rows = tbl_elem.findall(f".//{ROW_TAG}")
        html_rows = []
        for i, row in enumerate(rows):
            cells = row.findall(f".//{CELL_TAG}")
            cell_tag = "th" if i == 0 else "td"
            cells_html = "".join(
                f"<{cell_tag}>{''.join(_para_to_html(p) for p in c.findall(f'.//{PARA_TAG}'))}</{cell_tag}>"
                for c in cells
            )
            html_rows.append(f"<tr>{cells_html}</tr>")
        return f"<table style=\"border-collapse:collapse;width:100%\">" + "".join(html_rows) + "</table>"

    body = doc.element.body
    body_children = list(body)
    html_parts = []

    for child in body_children[start_body_idx: end_body_idx + 1]:
        tag = child.tag
        if tag == PARA_TAG:
            h = _para_to_html(child)
            if h:
                html_parts.append(h)
        elif tag == TBL_TAG:
            h = _table_to_html(child)
            if h:
                html_parts.append(h)

    return "\n".join(html_parts)


def _doc_blocks_slice_to_html(doc_blocks: list[dict], start_body_idx: int, end_body_idx: int) -> str:
    """
    当原始 DOCX 对象不可用时，基于 doc_blocks 快照退化构建 HTML。
    该模式保留段落顺序与基础文本，不保证 rich-text 样式和图片。
    """
    if not isinstance(doc_blocks, list):
        return ""
    lo, hi = (start_body_idx, end_body_idx) if start_body_idx <= end_body_idx else (end_body_idx, start_body_idx)
    selected: list[dict] = []
    for block in doc_blocks:
        if not isinstance(block, dict):
            continue
        try:
            body_idx = int(block.get("body_idx"))
        except Exception:
            continue
        if lo <= body_idx <= hi:
            selected.append(block)
    selected.sort(key=lambda b: int(b.get("body_idx", 0)))

    html_parts: list[str] = []
    for block in selected:
        text = str(block.get("text") or "").strip()
        if not text:
            continue
        escaped = _html.escape(text).replace("\n", "<br/>")
        if str(block.get("type") or "").lower() == "table":
            html_parts.append(f"<pre>{escaped}</pre>")
        else:
            html_parts.append(f"<p>{escaped}</p>")
    return "\n".join(html_parts)


def _slice_docx_bytes_by_body_range(docx_bytes: bytes, start_body_idx: int, end_body_idx: int) -> bytes:
    """
    从原始 DOCX 包中按 body 索引切片，仅保留指定区间内元素。
    保留原包中的样式/编号/关系资源，提升只读预览保真度。
    """
    import copy
    import io
    import zipfile
    import xml.etree.ElementTree as ET

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    BODY_TAG = f"{{{W_NS}}}body"
    SECTPR_TAG = f"{{{W_NS}}}sectPr"
    P_TAG = f"{{{W_NS}}}p"
    T_TAG = f"{{{W_NS}}}t"
    R_TAG = f"{{{W_NS}}}r"
    BR_TAG = f"{{{W_NS}}}br"
    LAST_RENDERED_PAGE_BREAK_TAG = f"{{{W_NS}}}lastRenderedPageBreak"
    PPR_TAG = f"{{{W_NS}}}pPr"
    PAGE_BREAK_BEFORE_TAG = f"{{{W_NS}}}pageBreakBefore"

    in_buf = io.BytesIO(docx_bytes)
    out_buf = io.BytesIO()

    with zipfile.ZipFile(in_buf, "r") as zin:
        entries: dict[str, bytes] = {}
        for info in zin.infolist():
            entries[info.filename] = zin.read(info.filename)

    doc_xml = entries.get("word/document.xml")
    if not doc_xml:
        raise ValueError("DOCX 包缺少 word/document.xml")

    root = ET.fromstring(doc_xml)
    body = root.find(f".//{BODY_TAG}")
    if body is None:
        raise ValueError("document.xml 缺少 w:body")

    original_children = list(body)
    kept_children = []
    sect_pr = None

    def _paragraph_plain_text(elem: ET.Element) -> str:
        texts = []
        for t in elem.iter(T_TAG):
            txt = t.text or ""
            if txt.strip():
                texts.append(txt.strip())
        return "".join(texts).strip()

    def _is_effectively_empty_paragraph(elem: ET.Element) -> bool:
        if elem.tag != P_TAG:
            return False
        return _paragraph_plain_text(elem) == ""

    def _trim_leading_trailing_empty_paragraphs(children: list[ET.Element]) -> list[ET.Element]:
        lo = 0
        hi = len(children) - 1
        while lo <= hi and _is_effectively_empty_paragraph(children[lo]):
            lo += 1
        while hi >= lo and _is_effectively_empty_paragraph(children[hi]):
            hi -= 1
        return children[lo: hi + 1] if lo <= hi else children

    def _remove_first_paragraph_page_break_controls(children: list[ET.Element]) -> None:
        """
        切片首段若带有“章节分页控制”，在独立预览中会出现“第一页空白、第二页才有正文”。
        仅对首个段落做温和清洗：移除 pageBreakBefore 和 run 内 page break 标记。
        """
        if not children:
            return
        first = children[0]
        if first.tag != P_TAG:
            return

        ppr = first.find(PPR_TAG)
        if ppr is not None:
            for node in list(ppr):
                if node.tag == PAGE_BREAK_BEFORE_TAG:
                    ppr.remove(node)
        else:
            ppr = ET.Element(PPR_TAG)
            first.insert(0, ppr)

        # 即使分页来自样式定义（style），这里也显式覆盖为 false，避免切片首屏空白页。
        pb_override = ET.Element(PAGE_BREAK_BEFORE_TAG)
        pb_override.set(f"{{{W_NS}}}val", "0")
        ppr.insert(0, pb_override)

        for run in list(first.findall(R_TAG)):
            for node in list(run):
                if node.tag == LAST_RENDERED_PAGE_BREAK_TAG:
                    run.remove(node)
                    continue
                if node.tag == BR_TAG:
                    br_type = (node.attrib.get(f"{{{W_NS}}}type") or "").lower()
                    if br_type == "page":
                        run.remove(node)
            if len(list(run)) == 0:
                first.remove(run)

    for idx, child in enumerate(original_children):
        if child.tag == SECTPR_TAG:
            sect_pr = copy.deepcopy(child)
            continue
        if start_body_idx <= idx <= end_body_idx:
            kept_children.append(copy.deepcopy(child))

    if not kept_children:
        raise ValueError("切片范围内无可用文档块")

    kept_children = _trim_leading_trailing_empty_paragraphs(kept_children)
    _remove_first_paragraph_page_break_controls(kept_children)

    for child in list(body):
        body.remove(child)
    for child in kept_children:
        body.append(child)
    if sect_pr is not None:
        body.append(sect_pr)

    entries["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)

    with zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, content in entries.items():
            zout.writestr(name, content)
    return out_buf.getvalue()


def _extract_raw_text_with_images(filename: str, content_bytes: bytes, use_vision_parsing: bool = False) -> tuple[str, dict]:
    """
    根据文件扩展名提取纯文本内容与图片映射表，支持 PDF / DOCX / TXT / MD
    返回: (纯文本内容, image_map图片映射字典)
    """
    ext = (filename or "").lower().split(".")[-1]
    image_map = {}

    if ext == "pdf":
        if use_vision_parsing:
            logger.info("PyMuPDF4LLM 视觉增强模式已开启，尝试解析 PDF 并提取图片...")
            try:
                import fitz
                import pymupdf4llm
                
                with tempfile.TemporaryDirectory() as temp_dir:
                    pdf_path = os.path.join(temp_dir, "temp.pdf")
                    with open(pdf_path, "wb") as f:
                        f.write(content_bytes)
                    
                    images_dir = os.path.join(temp_dir, "images")
                    os.makedirs(images_dir, exist_ok=True)
                    
                    # 使用 pymupdf4llm 导出 Markdown 和图片
                    md_text = pymupdf4llm.to_markdown(pdf_path, write_images=True, image_path=images_dir)
                    
                    # 将临时提取出的图片存放至服务端永久/长期临时目录
                    permanent_img_dir = PRO_ENGINE_ROOT / "data" / "extracted_images"
                    permanent_img_dir.mkdir(parents=True, exist_ok=True)

                    import concurrent.futures
                    import hashlib

                    # 找出文档中所有的图片引用
                    matches = list(re.finditer(r'!\[(.*?)\]\((.*?)\)', md_text))
                    
                    def _process_image(match):
                        img_rel_path = match.group(2)
                        source_img = Path(temp_dir) / img_rel_path

                        # 过滤: 文件不存在，或图片尺寸小于阈值（默认 5KB，可通过 VLM_MIN_IMAGE_SIZE_KB 调整）
                        _min_size = int(os.environ.get("VLM_MIN_IMAGE_SIZE_KB", "5")) * 1024
                        if not source_img.exists() or source_img.stat().st_size < _min_size:
                            return match.group(0), "", None

                        with open(source_img, "rb") as bf:
                            img_hash = hashlib.md5(bf.read()).hexdigest()

                        new_img_name = f"{img_hash}_{source_img.name}"
                        dest_img = permanent_img_dir / new_img_name
                        if not dest_img.exists():
                            shutil.copy(source_img, dest_img)

                        placeholder = f"__PRO_IMG_{img_hash}__"

                        # 仅安全调用部署在内网无出网风险的本地 VLM 打标
                        tag_desc = _tag_image_with_vlm(str(dest_img))
                        if not tag_desc or "无法提取" in tag_desc or len(tag_desc) < 4:
                            tag_desc = "本地配图"

                        # 注入 PostgreSQL 图片注册表保护池
                        from app.api_lite.database import SessionLocal, ImageRegistry
                        db = SessionLocal()
                        try:
                            if not db.query(ImageRegistry).filter_by(image_hash=img_hash).first():
                                row = ImageRegistry(
                                    image_hash=img_hash,
                                    project_id=None,
                                    abs_path=str(dest_img),
                                    preview_url=f"/api/extracted-images/{new_img_name}",
                                    placeholder=placeholder,
                                    vlm_caption=tag_desc,
                                    is_reference_only=1
                                )
                                db.add(row)
                                db.commit()
                        except Exception as e:
                            db.rollback()
                            logger.error(f"ImageRegistry 写入失败 (PDF): {e}")
                        finally:
                            db.close()

                        # 向外仅暴露安全的脱壳标记符及图注
                        img_info = {
                            "abs_path": str(dest_img),
                            "preview_url": f"/api/extracted-images/{new_img_name}",
                            "description": tag_desc,
                        }
                        return match.group(0), f"![{tag_desc}]({placeholder})", (placeholder, img_info)

                    # 使用线程池并发请求本地局域网大模型（并发数通过 VLM_CONCURRENT 配置，默认2）
                    _vlm_concurrent = int(os.environ.get("VLM_CONCURRENT", "2"))
                    new_md_text = md_text
                    with concurrent.futures.ThreadPoolExecutor(max_workers=_vlm_concurrent) as executor:
                        results = list(executor.map(_process_image, matches))
                        
                    for orig, repl, map_tuple in results:
                        if repl:
                            new_md_text = new_md_text.replace(orig, repl)
                        if map_tuple:
                            image_map[map_tuple[0]] = map_tuple[1]

                    return new_md_text, image_map

            except ImportError:
                logger.warning("服务器缺少 pymupdf4llm 运行环境，降级为基础纯文本解析...")
            except Exception as e:
                logger.warning(f"PyMuPDF4LLM 解析异常，正在降级处理: {e}", exc_info=True)
                
        # 默认快速纯文本解析（降级路径）
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content_bytes)) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages), image_map
        except ImportError:
            pass
        try:
            import pymupdf  # PyMuPDF (fitz)
            doc = pymupdf.open(stream=content_bytes, filetype="pdf")
            return "\n".join(page.get_text() for page in doc), image_map
        except ImportError:
            pass
        # 终极兜底：返回原始 bytes 中可打印部分
        logger.warning("未安装 pdfplumber/pymupdf，PDF 文本提取降级为字节解码")
        return content_bytes.decode("latin-1", errors="replace"), image_map

    if ext == "docx":
        # DOCX 是 ZIP 格式，python-docx 原生支持（段落 + 表格）
        try:
            import docx as _docx
            doc = _docx.Document(io.BytesIO(content_bytes))
            return _extract_docx_with_tables(doc), image_map
        except Exception:
            pass
        try:
            import zipfile, re as _re
            with zipfile.ZipFile(io.BytesIO(content_bytes)) as z:
                if "word/document.xml" in z.namelist():
                    xml = z.read("word/document.xml").decode("utf-8", errors="replace")
                    return _re.sub(r"<[^>]+>", " ", xml), image_map
        except Exception:
            pass
        logger.warning("DOCX 提取失败，降级为字节解码")
        return content_bytes.decode("utf-8", errors="replace"), image_map

    if ext == "doc":
        # 先用魔数检测真实格式
        # PK -> ZIP = 实际是 DOCX，OLE2 -> 真正的旧版 Word 二进制
        if content_bytes[:4] == b'PK\x03\x04':
            logger.info(".doc 文件实为 ZIP (DOCX)，改用 python-docx 解析")
            try:
                import docx as _docx
                doc = _docx.Document(io.BytesIO(content_bytes))
                return _extract_docx_with_tables(doc), image_map
            except Exception as _e:
                logger.warning(f"伪装成 .doc 的 DOCX 解析失败: {_e}")


        # OLE2 magic bytes = 真正的旧版 Word 二进制
        if content_bytes[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1':
            try:
                import olefile, re as _re
                with olefile.OleFileIO(io.BytesIO(content_bytes)) as ole:
                    if ole.exists("WordDocument"):
                        raw = ole.openstream("WordDocument").read()
                        text = raw.decode("utf-16-le", errors="replace")
                        text = _re.sub(r'[\x00-\x08\x0b-\x1f\x7f]', ' ', text)
                        text = _re.sub(r' {3,}', '\n', text).strip()
                        if len(text) > 50:
                            logger.info(f".doc OLE2 解析完成，提取 {len(text)} 字符")
                            return text, image_map
            except Exception as _e:
                logger.warning(f".doc OLE2 解析失败: {_e}")

        logger.warning("无法解析旧版 .doc 文件，请将文件另存为 .docx 后重新上传")
        return "[无法解析旧版 .doc 文件，请将文件另存为 .docx 后重新上传]", image_map

    # 文本/Markdown 等直接解码
    try:
        return content_bytes.decode("utf-8"), image_map
    except UnicodeDecodeError:
        return content_bytes.decode("latin-1", errors="replace"), image_map


async def _call_dify_workflow(api_key: str, inputs: dict, max_retries: int = 2) -> dict:
    """通用 Dify 工作流调用（blocking 模式），自动重试最多 max_retries 次。
    Dify 地址从 DIFY_API_URL 环境变量读取。
    """
    dify_base = os.environ.get("DIFY_API_URL", "http://localhost/v1").rstrip("/")
    dify_url = f"{dify_base}/workflows/run"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "inputs": inputs,
        "response_mode": "blocking",
        "user": "pro-engine-backend",
    }
    import asyncio as _asyncio
    last_err: Exception = RuntimeError("未知错误")
    for attempt in range(max_retries):
        try:
            # 延长超时至 1800 秒 (30分钟)，以应对极长文档或复杂 Workflow
            async with httpx.AsyncClient(timeout=1800) as client:
                resp = await client.post(dify_url, headers=headers, json=payload)
                resp.raise_for_status()
                return resp.json()
        except Exception as e:
            last_err = e
            if attempt < max_retries - 1:
                wait = 5 * (attempt + 1)
                logger.warning(f"[Dify blocking] 第 {attempt + 1} 次调用失败，{wait}s 后重试: {e}")
                await _asyncio.sleep(wait)
    raise last_err



def _parse_dify_outputs(dify_res: dict) -> dict:
    """
    解析 Dify workflow blocking 响应，兼容 structured_output / result / text 几种输出变量名
    """
    import re
    outputs = dify_res.get("data", {}).get("outputs", {})
    raw = (
        outputs.get("structured_output")
        or outputs.get("result")
        or outputs.get("text")
        or outputs
    )
    if isinstance(raw, list):
        # 兼容直接返回数组结构（如 outline 列表）
        return {"outline": raw}
    if isinstance(raw, str):
        import logging as _logging
        _log = _logging.getLogger(__name__)

        def _try_json_loads(text: str):
            try:
                return json.loads(text)
            except Exception:
                return None

        def _extract_fenced_json(text: str) -> str:
            m = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", text, flags=re.IGNORECASE)
            return (m.group(1) if m else text).strip()

        def _extract_balanced_candidates(text: str) -> list[str]:
            """从任意文本中提取平衡的 JSON 对象/数组候选（含嵌套，忽略字符串中的括号）。"""
            out: list[str] = []
            n = len(text)
            for start in range(n):
                ch = text[start]
                if ch not in "{[":
                    continue
                stack = [ch]
                in_str = False
                esc = False
                for end in range(start + 1, n):
                    c = text[end]
                    if in_str:
                        if esc:
                            esc = False
                        elif c == "\\":
                            esc = True
                        elif c == "\"":
                            in_str = False
                        continue
                    if c == "\"":
                        in_str = True
                        continue
                    if c in "{[":
                        stack.append(c)
                        continue
                    if c in "}]":
                        if not stack:
                            break
                        top = stack[-1]
                        if (top == "{" and c == "}") or (top == "[" and c == "]"):
                            stack.pop()
                            if not stack:
                                out.append(text[start:end + 1])
                                break
                        else:
                            break
            return out

        _log.info(f"[_parse_dify_outputs] raw str len={len(raw)}, last200={raw[-200:]!r}")
        cleaned = _extract_fenced_json(raw.strip())
        parsed = _try_json_loads(cleaned)
        if parsed is None:
            # 优先选择含 outline/sections 的 dict；其次接受 list；最后兜底任意 dict
            parsed_dict = None
            parsed_list = None
            fallback_dict = None
            for cand in _extract_balanced_candidates(cleaned):
                obj = _try_json_loads(cand)
                if isinstance(obj, dict):
                    if any(k in obj for k in ("outline", "sections", "items", "data")):
                        parsed_dict = obj
                        break
                    if fallback_dict is None:
                        fallback_dict = obj
                elif isinstance(obj, list) and parsed_list is None:
                    parsed_list = obj
            parsed = parsed_dict if parsed_dict is not None else (parsed_list if parsed_list is not None else (fallback_dict or {}))
        if isinstance(parsed, dict):
            _log.info(f"[_parse_dify_outputs] parsed JSON keys={list(parsed.keys())}")
        elif isinstance(parsed, list):
            _log.info(f"[_parse_dify_outputs] parsed JSON list len={len(parsed)}")
        raw = parsed
    if isinstance(raw, list):
        return {"outline": raw}
    return raw if isinstance(raw, dict) else {}

# ─────────────────────────────────


@router.post("/projects/extract", response_model=ExtractRequirementsResponse, summary="提取招标文件需求")
async def extract_project_requirements(
    file: UploadFile = File(...),
    project_name: str = Form(default=""),
    project_id: str = Form(default=""),             # 项目 ID，用于 PDF 缓存
    enable_desensitize: bool = Form(default=True),
    desensitize_profile: str = Form(default="tender"),
    use_vision_parsing: bool = Form(default=False)
):
    """
    接收前端上传的招标文件（PDF/DOCX/TXT/MD），提取纯文本与图片字典，
    先经脱敏处理，再调用 Dify 需求提取工作流。
    同时缓存 PDF 原文件并建立分页文本索引用于溯源。
    """
    try:
        content_bytes = await file.read()
        ext = (file.filename or "").lower().split(".")[-1]

        # PDF 缓存 + 分页文本索引（用于后续溯源匹配）
        pdf_url = ""
        pages_text = []
        cache_id = project_id or uuid.uuid4().hex[:12]

        if ext == "pdf":
            # PDF 直接缓存
            pdf_url = _cache_pdf_file(cache_id, content_bytes)
            pages_text = _extract_pdf_pages_text(content_bytes)
        elif ext in ("docx", "doc"):
            # DOCX/DOC 异步触发转 PDF，失败不阻断主流程
            try:
                pdf_url = _convert_to_pdf_and_cache(cache_id, content_bytes, file.filename or "")
            except Exception as _conv_err:
                logger.warning(f"DOCX→PDF 转换异常（已忽略）: {_conv_err}")

        raw_document, raw_image_map = _extract_raw_text_with_images(
            file.filename or "", content_bytes, use_vision_parsing=use_vision_parsing
        )

        # ── 投标文件定位符缓存（DOCX 专属）─────────────────────────────────
        # 生成带 [Pxxx] 定位符的文本，存入缓存，同时作为发给 Dify 的文本基础。
        # 定位符不参与脱敏（[Pxxx] 格式不会被 PIPT 识别），LLM 可见以定位附件范围。
        _loc_text_for_dify = raw_document  # 默认回退
        if ext in ("docx", "doc") and cache_id:
            try:
                import docx as _docx_mod
                _loc_doc = _docx_mod.Document(io.BytesIO(content_bytes))
                _loc_text, _loc_map, _blocks = _extract_docx_with_locators(_loc_doc)
                _locator_cache[cache_id] = {"doc": _loc_doc, "locator_map": _loc_map, "doc_blocks": _blocks}
                _persist_doc_blocks_snapshot(cache_id, _blocks)
                if ext == "docx":
                    _persist_docx_for_locators(cache_id, content_bytes)
                _loc_text_for_dify = _loc_text  # 带定位符文本发给 Dify
                logger.info(f"[{cache_id}] 定位符缓存写入: {len(_loc_map)} 个定位点")
            except Exception as _loc_err:
                logger.warning(f"[{cache_id}] 定位符缓存写入失败（已忽略）: {_loc_err}")

        text_for_dify = _loc_text_for_dify  # DOCX 时含 [Pxxx] 定位符；PDF/TXT 时为原文

        # .doc 文件无法解析时提前返回 400，避免将占位文本发给 Dify
        if raw_document.startswith("["):
            raise HTTPException(
                status_code=400,
                detail="旧版 .doc 文件无法自动解析，请将文件另存为 .docx 后重新上传。"
            )

        if enable_desensitize:
            try:
                engine = get_engine()
                # 从 config.yaml 读 tender profile 配置
                profile_config = load_profile_config(desensitize_profile)
                target_entities = profile_config.get(
                    "target_entities", ["name", "phone", "email", "id_number"]
                )
                method = profile_config.get("method", "mask")

                import asyncio
                desen_result = await asyncio.to_thread(
                    engine.desensitize,
                    text=raw_document[:300000],
                    target_entities=target_entities,
                    method=method,
                    llm_mode=os.environ.get('PIPT_LLM_MODE_EXTRACT', 'verify_only'),
                )
                text_for_dify = desen_result.desensitized_text
                mapping_table = getattr(desen_result, "mapping_table", {}) or {}
                entity_count = getattr(desen_result, "entity_count", 0) or 0
                logger.info(
                    f"招标文件脱敏完成: profile={desensitize_profile}, "
                    f"识别实体 {entity_count} 处"
                )
            except Exception as desen_err:
                logger.warning(f"脱敏处理失败，使用原文继续: {desen_err}")
                text_for_dify = raw_document[:300000]
                mapping_table = {}
                entity_count = 0
        else:
            mapping_table = {}
            entity_count = 0

        dify_key = _get_workflow_key("requirement_extractor")
        if not dify_key:
            raise HTTPException(status_code=500, detail="需求提取工作流 API Key 未配置，请在 config.yaml 中填写")
        dify_res = await _call_dify_workflow(dify_key, {
            "raw_document": text_for_dify,
            "project_name": project_name or (file.filename or "").rsplit(".", 1)[0],
        })

        structured_data = _parse_dify_outputs(dify_res)

        bid_type = structured_data.get("bid_type", "tech")
        project_summary = structured_data.get("project_summary", "")
        requirements_raw = structured_data.get("requirements", [])

        reqs = []
        for r in requirements_raw:
            if not r.get("content"):
                continue
            # 获取 Dify 返回的原文摘录
            source_excerpt = r.get("source_excerpt", "")
            # 溯源匹配：将摘录匹配到具体页码
            source_pages_data = []
            if source_excerpt and pages_text:
                source_pages_data = _match_source_pages(source_excerpt, pages_text)
            reqs.append(
                ExtractRequirementItem(
                    type=r.get("type", "tech"),
                    content=r.get("content", ""),
                    points=r.get("points"),
                    source_excerpt=source_excerpt,
                    source_pages=source_pages_data,
                )
            )

        # analysis_report: 框架化解析结果（当前由 Dify 返回，若暂未配置则返回空列表）
        analysis_report_raw = structured_data.get("analysis_report", [])

        _persist_raw_document(cache_id, text_for_dify[:300000])

        return ExtractRequirementsResponse(
            bid_type=bid_type,
            project_summary=project_summary,
            requirements=reqs,
            analysis_report=analysis_report_raw,
            mapping_table=mapping_table,
            entity_count=entity_count,
            image_map=raw_image_map,
            required_attachments=structured_data.get("required_attachments", []),
            scoring_table_template=structured_data.get("scoring_table_template", []),
            raw_document=text_for_dify,
            pdf_url=pdf_url,
            # AI 智能评估字段：由 Requirement_Extractor 工作流输出，未配置时为 None
            expected_word_count=structured_data.get("expected_word_count") or None,
            expected_chapter_count=structured_data.get("expected_chapter_count") or None,
        )

    except Exception as e:
        logger.error(f"Failed to extract requirements: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ─── SSE 版解析端点：实时推送各阶段进度 ──────────────────────────────
@router.post("/projects/extract-stream", summary="SSE 实时解析招标文件")
async def extract_project_requirements_stream(
    file: UploadFile = File(...),
    project_name: str = Form(default=""),
    project_id: str = Form(default=""),
    enable_desensitize: bool = Form(default=True),
    desensitize_profile: str = Form(default="tender"),
    use_vision_parsing: bool = Form(default=False)
):
    """
    与 /projects/extract 相同逻辑，但以 SSE 逐阶段推送进度。

    SSE 事件：
    - progress: { step, label, percent }  各阶段进度
    - result:   完整的 ExtractRequirementsResponse JSON
    - error:    { message }
    """
    import json as _json

    # 提前读取文件内容（在 SSE generator 外部，因为 UploadFile 不能跨异步边界重复使用）
    content_bytes = await file.read()
    filename = file.filename or ""
    ext = filename.lower().split(".")[-1]

    async def sse_generator():
        try:
            # ── 阶段 1：文档解析（约 10-15%） ──
            yield f"event: progress\ndata: {_json.dumps({'step': 0, 'label': '解析文档结构', 'percent': 5}, ensure_ascii=False)}\n\n"

            pdf_url = ""
            pages_text = []
            cache_id = project_id or uuid.uuid4().hex[:12]

            if ext == "pdf":
                pdf_url = _cache_pdf_file(cache_id, content_bytes)
                pages_text = _extract_pdf_pages_text(content_bytes)
            elif ext in ("docx", "doc"):
                try:
                    pdf_url = _convert_to_pdf_and_cache(cache_id, content_bytes, filename)
                except Exception as _conv_err:
                    logger.warning(f"DOCX→PDF 转换异常（已忽略）: {_conv_err}")

            raw_document, raw_image_map = _extract_raw_text_with_images(
                filename, content_bytes, use_vision_parsing=use_vision_parsing
            )

            # ── 投标文件定位符缓存（DOCX 专属）─────────────────────────────────
            # 生成带 [Pxxx] 定位符的文本，存入缓存，同时作为发给 Dify 的文本基础。
            # 定位符不参与脱敏（[Pxxx] 格式不会被 PIPT 识别），LLM 可见以定位附件范围。
            _loc_text_for_dify = raw_document
            if ext in ("docx", "doc") and cache_id:
                try:
                    import docx as _docx_mod
                    _loc_doc = _docx_mod.Document(io.BytesIO(content_bytes))
                    _loc_text, _loc_map, _blocks = _extract_docx_with_locators(_loc_doc)
                    _locator_cache[cache_id] = {"doc": _loc_doc, "locator_map": _loc_map, "doc_blocks": _blocks}
                    _persist_doc_blocks_snapshot(cache_id, _blocks)
                    if ext == "docx":
                        _persist_docx_for_locators(cache_id, content_bytes)
                    _loc_text_for_dify = _loc_text  # 带定位符文本发给 Dify/PIPT
                    logger.info(f"[SSE {cache_id}] 定位符缓存写入: {len(_loc_map)} 个定位点")
                except Exception as _loc_err:
                    logger.warning(f"[SSE {cache_id}] 定位符缓存写入失败（已忽略）: {_loc_err}")

            if raw_document.startswith("["):
                yield f"event: error\ndata: {_json.dumps({'message': '旧版 .doc 文件无法自动解析，请将文件另存为 .docx 后重新上传。'}, ensure_ascii=False)}\n\n"
                return

            yield f"event: progress\ndata: {_json.dumps({'step': 0, 'label': '文档结构解析完成', 'percent': 15}, ensure_ascii=False)}\n\n"

            # ── 阶段 2：PIPT 脱敏（约 15-50%） ──
            text_for_dify = _loc_text_for_dify  # 默认使用带定位符文本（DOCX）或原文（PDF）
            mapping_table = {}
            entity_count = 0

            if enable_desensitize:
                yield f"event: progress\ndata: {_json.dumps({'step': 1, 'label': '隐私脱敏处理中', 'percent': 20}, ensure_ascii=False)}\n\n"
                try:
                    engine = get_engine()
                    profile_config = load_profile_config(desensitize_profile)
                    target_entities = profile_config.get(
                        "target_entities", ["name", "phone", "email", "id_number"]
                    )
                    method = profile_config.get("method", "mask")

                    import asyncio
                    desen_result = await asyncio.to_thread(
                        engine.desensitize,
                        text=raw_document[:300000],
                        target_entities=target_entities,
                        method=method,
                        llm_mode=os.environ.get('PIPT_LLM_MODE_EXTRACT', 'verify_only'),
                    )
                    text_for_dify = desen_result.desensitized_text
                    mapping_table = getattr(desen_result, "mapping_table", {}) or {}
                    entity_count = getattr(desen_result, "entity_count", 0) or 0

                    yield f"event: progress\ndata: {_json.dumps({'step': 1, 'label': f'脱敏完成，识别 {entity_count} 处实体', 'percent': 50}, ensure_ascii=False)}\n\n"
                except Exception as desen_err:
                    logger.warning(f"脱敏处理失败，使用原文继续: {desen_err}")
                    text_for_dify = raw_document[:300000]
                    yield f"event: progress\ndata: {_json.dumps({'step': 1, 'label': '脱敏跳过（使用原文）', 'percent': 50}, ensure_ascii=False)}\n\n"
            else:
                yield f"event: progress\ndata: {_json.dumps({'step': 1, 'label': '跳过脱敏', 'percent': 50}, ensure_ascii=False)}\n\n"

            # ── 阶段 3：跳过 Dify 需求提取（已移至解析报告阶段） ──
            yield f"event: progress\ndata: {_json.dumps({'step': 2, 'label': '预处理完成', 'percent': 100}, ensure_ascii=False)}\n\n"

            _persist_raw_document(cache_id, text_for_dify[:300000])

            result_data = {
                "bid_type": "tech",
                "project_summary": "",
                "requirements": [],
                "analysis_report": [],
                "mapping_table": mapping_table,
                "entity_count": entity_count,
                "image_map": raw_image_map,
                "required_attachments": [],
                "scoring_table_template": [],
                "raw_document": text_for_dify,
                "pdf_url": pdf_url,
                # AI 智能评估字段（此端点仅做文本预处理，无 Dify 调用，默认 null）
                "expected_word_count": None,
                "expected_chapter_count": None,
            }
            yield f"event: result\ndata: {_json.dumps(result_data, ensure_ascii=False)}\n\n"

        except Exception as e:
            logger.error(f"SSE extract 异常: {e}", exc_info=True)
            yield f"event: error\ndata: {_json.dumps({'message': str(e)}, ensure_ascii=False)}\n\n"

    from fastapi.responses import StreamingResponse as _SSEResp
    return _SSEResp(
        sse_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )




# ─── PDF 缓存文件服务 ──────────────────────────────────────────────────
from fastapi.responses import Response as _RawResponse

@router.get("/projects/pdf/{project_id}", summary="获取已缓存的 PDF 文件")
async def get_cached_pdf(project_id: str):
    """
    提供已缓存 PDF 文件的下载/在线访问，供前端 PDF 预览面板使用。
    文件路径: data/pdf_cache/{project_id}.pdf
    """
    # 安全过滤：仅允许字母数字和连字符
    import re as _re
    if not _re.match(r'^[a-zA-Z0-9_-]+$', project_id):
        raise HTTPException(status_code=400, detail="无效的 project_id 格式")

    pdf_path = PRO_ENGINE_ROOT / "data" / "pdf_cache" / f"{project_id}.pdf"
    if not pdf_path.exists():
        raise HTTPException(status_code=404, detail="PDF 文件未找到，请先上传招标文件")
    
    content = pdf_path.read_bytes()
    return _RawResponse(
        content=content,
        media_type="application/pdf",
        headers={
            "Content-Disposition": f"inline; filename={project_id}.pdf",
            "Cache-Control": "public, max-age=3600",
        }
    )


@router.get("/extracted-images/by-hash/{image_hash}", summary="凭MD5哈希获取提取的图片")
async def get_extracted_image_by_hash(image_hash: str, db: Session = Depends(get_db)):
    """前端直接渲染 __PRO_IMG_{hash}__ 占位符用的接口"""
    import re as _re
    if not _re.match(r'^[a-fA-F0-9]+$', image_hash):
        raise HTTPException(status_code=400, detail="无效的散列格式")
        
    from app.api_lite.database import ImageRegistry
    row = db.query(ImageRegistry).filter_by(image_hash=image_hash).first()
    if not row:
        raise HTTPException(status_code=404, detail="图片实体不存在")
        
    img_path = Path(row.abs_path)
    if not img_path.exists():
        raise HTTPException(status_code=404, detail="图片物理文件已丢失")
        
    suffix = img_path.suffix.lower()
    content_type = "image/png"
    if suffix in (".jpg", ".jpeg"): content_type = "image/jpeg"
    elif suffix == ".webp": content_type = "image/webp"
    
    from fastapi.responses import FileResponse
    return FileResponse(str(img_path), media_type=content_type)


@router.get("/extracted-images/{filename}", summary="获取从招标文件中提取的图片")
async def get_extracted_image(filename: str):
    """
    前端预览图片路由：将 data/extracted_images/{filename} 暴露给浏览器。
    image_map 中的 preview_url 指向此端点，避免浏览器无法访问服务器绝对路径。
    安全措施：仅允许合法文件名（字母数字/下划线/连字符/点），防止路径穿越。
    """
    import re as _re
    if not _re.match(r'^[a-zA-Z0-9_.()-]+$', filename):
        raise HTTPException(status_code=400, detail="无效的文件名格式")

    img_path = PRO_ENGINE_ROOT / "data" / "extracted_images" / filename
    if not img_path.exists():
        raise HTTPException(status_code=404, detail="图片文件不存在")

    # 根据扩展名推断 MIME 类型
    suffix = img_path.suffix.lower()
    mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                ".gif": "image/gif", ".webp": "image/webp", ".bmp": "image/bmp"}
    media_type = mime_map.get(suffix, "application/octet-stream")

    return _RawResponse(
        content=img_path.read_bytes(),
        media_type=media_type,
        headers={"Cache-Control": "public, max-age=86400"},
    )


@router.post("/projects/upload-pdf", summary="独立上传 PDF 文件到缓存")
async def upload_pdf(
    file: UploadFile = File(...),
    project_id: str = Form(...),
):
    """
    独立的 PDF 上传端点，用于仅上传 PDF 到缓存而不触发需求提取。
    适用于用户已有提取结果但想重新加载 PDF 预览的场景。
    """
    content_bytes = await file.read()
    ext = (file.filename or "").lower().split(".")[-1]
    if ext != "pdf":
        raise HTTPException(status_code=400, detail="仅支持 PDF 格式文件")
    
    pdf_url = _cache_pdf_file(project_id, content_bytes)
    return {"pdf_url": pdf_url, "message": "PDF 已缓存"}


# ─── 投标文件附件提取端点 ────────────────────────────────────────────────────

@router.post("/bid-attachment/extract", summary="按定位符提取 DOCX 附件内容")
async def extract_bid_attachment(
    body: dict = Body(...),
):
    """
    根据解析报告中的 start_locator/end_locator，从缓存的 DOCX 中
    切取对应段落范围，转为 HTML 返回给前端编辑器。

    Body: {"project_id": str, "start_locator": str, "end_locator": str, "attachment_name": str}
    """
    import re as _re
    def _normalize_locator_token(raw: str) -> str:
        """
        将模型输出的定位符归一化为 Pxxxx 形式。
        兼容样式: [P0012] / p12 / P12~P20 / 起始P0012
        """
        s = (raw or "").strip().upper()
        if not s:
            return ""
        m = _re.search(r"P\s*0*(\d+)", s)
        if not m:
            return ""
        return f"P{int(m.group(1)):04d}"

    project_id     = (body.get("project_id") or "").strip()
    start_locator  = _normalize_locator_token(body.get("start_locator") or "")
    end_locator    = _normalize_locator_token(body.get("end_locator") or "")
    attachment_name = (body.get("attachment_name") or "").strip()

    if not project_id:
        raise HTTPException(status_code=400, detail="project_id 不能为空")
    if not start_locator or not end_locator:
        raise HTTPException(status_code=400, detail="start_locator 和 end_locator 不能为空")

    cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        _restore_locator_cache_from_disk(project_id)
        cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        raise HTTPException(
            status_code=404,
            detail=f"项目 [{project_id}] 的 DOCX 定位符缓存不存在（服务重启后可能丢失），请重新上传 DOCX 或重新解析一次"
        )

    doc = cache_entry.get("doc")
    locator_map = cache_entry.get("locator_map", {})
    doc_blocks = cache_entry.get("doc_blocks", [])

    resolved_start_locator = start_locator
    resolved_end_locator = end_locator

    start_idx = locator_map.get(resolved_start_locator)
    end_idx   = locator_map.get(resolved_end_locator)

    if start_idx is None:
        raise HTTPException(status_code=404, detail=f"定位符 {start_locator} 未找到")
    if end_idx is None:
        raise HTTPException(status_code=404, detail=f"定位符 {end_locator} 未找到")
    if start_idx > end_idx:
        start_idx, end_idx = end_idx, start_idx

    try:
        if doc is not None:
            html = _body_elements_to_html(doc, start_idx, end_idx)
        else:
            html = _doc_blocks_slice_to_html(doc_blocks, start_idx, end_idx)
            if not html:
                raise HTTPException(
                    status_code=409,
                    detail="当前仅恢复了文档块索引快照，且快照正文为空；请上传原始 DOCX 执行“重建定位”后重试"
                )
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        logger.error(f"[{project_id}] 附件内容提取失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"附件内容提取失败: {e}")

    logger.info(
        f"[{project_id}] 附件提取: {attachment_name} "
        f"{start_locator} → {end_locator}  body[{start_idx}:{end_idx}]"
    )
    return {
        "html": html,
        "attachment_name": attachment_name,
        "start_locator": start_locator,
        "end_locator": end_locator,
        "resolved_start_locator": resolved_start_locator,
        "resolved_end_locator": resolved_end_locator,
        "paragraph_count": end_idx - start_idx + 1,
        "snapshot_only": doc is None,
    }


@router.get("/bid-attachment/test-locators", summary="调试：查看项目定位符映射")
async def test_locators(project_id: str):
    """
    开发调试用：查看指定项目的定位符映射和带定位符文本片段预览。
    """
    cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        raise HTTPException(status_code=404, detail=f"项目 [{project_id}] 缓存不存在")

    locator_map = cache_entry["locator_map"]
    doc         = cache_entry["doc"]
    if doc is None:
        raise HTTPException(status_code=409, detail=f"项目 [{project_id}] 当前仅有块索引快照，缺少 DOCX 正文对象")

    # 取前 20 个定位符预览
    preview = []
    WPC_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    body_children = list(doc.element.body)
    for key in list(locator_map.keys())[:20]:
        idx = locator_map[key]
        elem = body_children[idx] if idx < len(body_children) else None
        snippet = ""
        if elem is not None:
            snippet = "".join(n.text or "" for n in elem.iter(f"{{{WPC_NS}}}t"))[:60]
        preview.append({"locator": key, "body_idx": idx, "snippet": snippet})

    return {
        "project_id": project_id,
        "total_locators": len(locator_map),
        "preview": preview,
    }


@router.get("/projects/{project_id}/doc-blocks", summary="获取文档块级索引")
async def get_doc_blocks(project_id: str):
    cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        _restore_locator_cache_from_disk(project_id)
        cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        db_blocks = _load_doc_blocks_snapshot(project_id)
        if db_blocks:
            return {
                "project_id": project_id,
                "blocks": db_blocks,
                "total_blocks": len(db_blocks),
                "snapshot_only": True,
            }
        raise HTTPException(status_code=404, detail=f"项目 [{project_id}] 的文档块缓存不存在")
    return {
        "project_id": project_id,
        "blocks": cache_entry.get("doc_blocks", []),
        "total_blocks": len(cache_entry.get("doc_blocks", [])),
        "snapshot_only": bool(cache_entry.get("snapshot_only")),
    }


@router.post("/projects/{project_id}/rebuild-locator", summary="重建 DOCX 定位缓存")
async def rebuild_doc_locator(project_id: str, file: UploadFile = File(...)):
    """
    仅重建 DOCX 定位缓存，不触发 Dify 解析流程。
    用于旧项目/重启后 docx_cache 缺失导致的块级锚点不可用问题修复。
    """
    filename = (file.filename or "").lower()
    if not filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="仅支持上传 DOCX 文件")
    content = await file.read()
    if not content:
        raise HTTPException(status_code=400, detail="上传文件为空")

    try:
        import io
        import docx as _docx_mod

        doc = _docx_mod.Document(io.BytesIO(content))
        _text, locator_map, doc_blocks = _extract_docx_with_locators(doc)
        _locator_cache[project_id] = {
            "doc": doc,
            "locator_map": locator_map,
            "doc_blocks": doc_blocks,
        }
        _persist_docx_for_locators(project_id, content)
        _persist_doc_blocks_snapshot(project_id, doc_blocks)
        logger.info(f"[{project_id}] 重建定位缓存成功: {len(doc_blocks)} 个文档块")
        return {
            "project_id": project_id,
            "blocks": len(doc_blocks),
            "locators": len(locator_map),
        }
    except Exception as e:
        logger.error(f"[{project_id}] 重建定位缓存失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"重建定位缓存失败: {e}")


@router.get("/projects/{project_id}/source-docx", summary="获取项目原始 DOCX")
async def get_project_source_docx(project_id: str):
    """
    返回项目原始 DOCX，供前端做全量分页预览。
    """
    docx_path = PRO_ENGINE_ROOT / "data" / "docx_cache" / f"{project_id}.docx"
    if not docx_path.exists():
        raise HTTPException(status_code=404, detail="项目原始 DOCX 不存在，请重新上传 DOCX 或执行重建定位")
    try:
        content = docx_path.read_bytes()
    except Exception as e:
        logger.error(f"[{project_id}] 读取原始 DOCX 失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"读取原始 DOCX 失败: {e}")

    filename = f"{project_id}.docx"
    return _RawResponse(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'inline; filename="{filename}"'},
    )


@router.post("/bid-attachment/extract-by-block", summary="按 block_id 提取 DOCX 附件内容")
async def extract_bid_attachment_by_block(body: dict = Body(...)):
    project_id = (body.get("project_id") or "").strip()
    start_block_id = (body.get("start_block_id") or "").strip()
    end_block_id = (body.get("end_block_id") or "").strip()
    attachment_name = (body.get("attachment_name") or "").strip()
    if not project_id or not start_block_id or not end_block_id:
        raise HTTPException(status_code=400, detail="project_id/start_block_id/end_block_id 不能为空")

    cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        _restore_locator_cache_from_disk(project_id)
        cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        raise HTTPException(status_code=404, detail=f"项目 [{project_id}] 的 DOCX 缓存不存在，请重新上传 DOCX 或重新解析一次")

    doc = cache_entry.get("doc")
    blocks = cache_entry.get("doc_blocks", [])
    block_map = {b.get("block_id"): b for b in blocks if isinstance(b, dict)}
    start_block = block_map.get(start_block_id)
    end_block = block_map.get(end_block_id)
    if not start_block:
        raise HTTPException(status_code=404, detail=f"block_id {start_block_id} 未找到")
    if not end_block:
        raise HTTPException(status_code=404, detail=f"block_id {end_block_id} 未找到")

    start_idx = int(start_block.get("body_idx", 0))
    end_idx = int(end_block.get("body_idx", 0))
    if start_idx > end_idx:
        start_idx, end_idx = end_idx, start_idx
        start_block_id, end_block_id = end_block_id, start_block_id

    try:
        if doc is not None:
            html = _body_elements_to_html(doc, start_idx, end_idx)
        else:
            html = _doc_blocks_slice_to_html(blocks, start_idx, end_idx)
            if not html:
                raise HTTPException(
                    status_code=409,
                    detail="当前仅恢复了文档块索引快照，且快照正文为空；请上传原始 DOCX 执行“重建定位”后重试"
                )
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        logger.error(f"[{project_id}] block 附件内容提取失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"附件内容提取失败: {e}")

    return {
        "html": html,
        "attachment_name": attachment_name,
        "start_block_id": start_block_id,
        "end_block_id": end_block_id,
        "paragraph_count": end_idx - start_idx + 1,
        "snapshot_only": doc is None,
    }


@router.post("/bid-attachment/extract-by-block-docx", summary="按 block_id 切片并返回 DOCX（二进制）")
async def extract_bid_attachment_docx_by_block(body: dict = Body(...)):
    project_id = (body.get("project_id") or "").strip()
    start_block_id = (body.get("start_block_id") or "").strip()
    end_block_id = (body.get("end_block_id") or "").strip()
    attachment_name = (body.get("attachment_name") or "").strip() or "slice"
    if not project_id or not start_block_id or not end_block_id:
        raise HTTPException(status_code=400, detail="project_id/start_block_id/end_block_id 不能为空")

    cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        _restore_locator_cache_from_disk(project_id)
        cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        raise HTTPException(status_code=404, detail=f"项目 [{project_id}] 的 DOCX 缓存不存在，请重新上传 DOCX 或重新解析一次")

    blocks = cache_entry.get("doc_blocks", [])
    block_map = {b.get("block_id"): b for b in blocks if isinstance(b, dict)}
    start_block = block_map.get(start_block_id)
    end_block = block_map.get(end_block_id)
    if not start_block:
        raise HTTPException(status_code=404, detail=f"block_id {start_block_id} 未找到")
    if not end_block:
        raise HTTPException(status_code=404, detail=f"block_id {end_block_id} 未找到")

    start_idx = int(start_block.get("body_idx", 0))
    end_idx = int(end_block.get("body_idx", 0))
    if start_idx > end_idx:
        start_idx, end_idx = end_idx, start_idx
        start_block_id, end_block_id = end_block_id, start_block_id

    docx_path = PRO_ENGINE_ROOT / "data" / "docx_cache" / f"{project_id}.docx"
    source_bytes = b""
    if docx_path.exists():
        source_bytes = docx_path.read_bytes()
    else:
        doc_obj = cache_entry.get("doc")
        if doc_obj is not None:
            import io
            buf = io.BytesIO()
            doc_obj.save(buf)
            source_bytes = buf.getvalue()

    if not source_bytes:
        raise HTTPException(status_code=409, detail="原始 DOCX 不可用，无法生成保格式切片；请上传原始 DOCX 重建定位缓存")

    try:
        sliced_bytes = _slice_docx_bytes_by_body_range(source_bytes, start_idx, end_idx)
    except Exception as e:
        logger.error(f"[{project_id}] DOCX 切片失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"DOCX 切片失败: {e}")

    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", attachment_name).strip("_") or "slice"
    filename = f"{safe_name}_{start_block_id}_{end_block_id}.docx"
    return _RawResponse(
        content=sliced_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'inline; filename="{filename}"',
            "X-Start-Block-Id": start_block_id,
            "X-End-Block-Id": end_block_id,
        },
    )


@router.delete("/projects/{project_id}/caches", summary="删除项目缓存数据")
async def delete_project_caches(project_id: str):
    """
    清理项目相关的后端缓存（PDF 文件等）。
    项目主数据存储在前端 localStorage，此端点仅做后端资源清理。
    """
    import re as _re
    if not _re.match(r'^[a-zA-Z0-9_-]+$', project_id):
        raise HTTPException(status_code=400, detail="无效的 project_id 格式")

    cleaned = []

    # 清理 PDF 缓存
    pdf_path = PRO_ENGINE_ROOT / "data" / "pdf_cache" / f"{project_id}.pdf"
    if pdf_path.exists():
        pdf_path.unlink()
        cleaned.append("pdf_cache")
        logger.info(f"已清理 PDF 缓存: {pdf_path}")

    # 清理投标文件定位符缓存
    if project_id in _locator_cache:
        del _locator_cache[project_id]
        cleaned.append("locator_cache")
        logger.info(f"已清理定位符缓存: project_id={project_id}")

    raw_path = PRO_ENGINE_ROOT / "data" / "raw_doc_cache" / f"{project_id}.txt"
    if raw_path.exists():
        raw_path.unlink()
        cleaned.append("raw_doc_cache")
        logger.info(f"已清理原文缓存: {raw_path}")

    return {"project_id": project_id, "cleaned": cleaned, "message": f"已清理 {len(cleaned)} 项资源"}


@router.post("/projects/re-extract", response_model=ExtractRequirementsResponse, summary="基于缓存原文重新提取需求")
async def re_extract_project_requirements(request: ReExtractRequirementsRequest):
    """
    接收前端传递的已缓存脱敏文本，直接调用 Dify 提取，跳过文件解析与本地脱敏过程
    """
    try:
        dify_key = _get_workflow_key("requirement_extractor")
        if not dify_key:
            raise HTTPException(status_code=500, detail="需求提取工作流 API Key 未配置，请在 config.yaml 中填写")
            
        raw_document = (request.raw_document or "").strip()
        if not raw_document:
            raw_document = _load_raw_document(request.project_id)
        if not raw_document:
            raise HTTPException(status_code=404, detail="未找到项目缓存原文，请先重新上传并解析文档")

        dify_res = await _call_dify_workflow(dify_key, {
            "raw_document": raw_document,
            "project_name": request.project_name,
        })

        structured_data = _parse_dify_outputs(dify_res)

        bid_type = structured_data.get("bid_type", "tech")
        project_summary = structured_data.get("project_summary", "")
        requirements_raw = structured_data.get("requirements", [])

        reqs = [
            ExtractRequirementItem(
                type=r.get("type", "tech"),
                content=r.get("content", ""),
                points=r.get("points"),
            )
            for r in requirements_raw
            if r.get("content")
        ]

        return ExtractRequirementsResponse(
            bid_type=bid_type,
            project_summary=project_summary,
            requirements=reqs,
            mapping_table={},  # 重提直接继承之前的即可，后端这里返回空不覆盖
            entity_count=0,
            image_map={},
            required_attachments=structured_data.get("required_attachments", []),
            scoring_table_template=structured_data.get("scoring_table_template", []),
            raw_document=raw_document,
        )

    except Exception as e:
        logger.error(f"Failed to re-extract requirements: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/projects/generate-outline", response_model=GenerateOutlineResponse, summary="AI 生成标书大纲")
def _safe_json_loads(raw: str):
    try:
        return json.loads(raw)
    except Exception:
        return None


def _normalize_heading_key(title: str) -> str:
    return re.sub(r"\s+", "", str(title or "")).strip().lower()


def _parse_structure_heading_seed_json(raw: str) -> list[dict]:
    """解析前端传入的技术部分 heading 种子，仅保留二级标题。"""
    def _as_bool(v) -> bool:
        if isinstance(v, bool):
            return v
        if v is None:
            return False
        s = str(v).strip().lower()
        if s in {"true", "1", "yes", "y", "是"}:
            return True
        if s in {"false", "0", "no", "n", "否", ""}:
            return False
        return False

    parsed = _safe_json_loads((raw or "").strip())
    if parsed is None:
        return []

    if isinstance(parsed, dict):
        candidates = (
            parsed.get("headings")
            or parsed.get("technical_sections")
            or (parsed.get("bid_structure") or {}).get("technical_sections")
            or parsed.get("sections")
            or []
        )
    elif isinstance(parsed, list):
        candidates = parsed
    else:
        candidates = []

    seeds: list[dict] = []
    for idx, item in enumerate(candidates):
        if isinstance(item, str):
            title = item.strip()
            item = {}
        elif isinstance(item, dict):
            title = str(item.get("title", "") or "").strip()
        else:
            continue

        if not title:
            continue

        level = int(item.get("level") or item.get("headingLevel") or 2)
        if level != 2:
            continue

        raw_keywords = item.get("keywords") or []
        keywords = [str(k).strip() for k in raw_keywords if str(k).strip()] if isinstance(raw_keywords, list) else []
        seeds.append({
            "id": str(item.get("id") or f"tech_heading_{idx + 1}"),
            "title": title,
            "headingLevel": 2,
            "wordCount": int(item.get("wordCount") or item.get("word_count") or 0),
            "writingHint": str(item.get("writingHint") or item.get("writing_hint") or "").strip(),
            "keywords": keywords,
            "relatedAnalysisIds": item.get("relatedAnalysisIds") or item.get("related_analysis_ids") or [],
            "score_tag": str(item.get("score_tag") or item.get("scoreTag") or "").strip(),
            "score_item_id": str(item.get("score_item_id") or item.get("scoreItemId") or "").strip(),
            "max_score": int(item.get("max_score") or item.get("maxScore") or 0),
            "criteria": str(item.get("criteria") or "").strip(),
            "related_target_ids": item.get("related_target_ids") or item.get("relatedTargetIds") or [],
            "priority_weight": float(item.get("priority_weight") or item.get("priorityWeight") or 0.0),
            "generation_strategy": str(item.get("generation_strategy") or item.get("generationStrategy") or "general").strip(),
            "response_candidate": _as_bool(item.get("response_candidate", item.get("responseCandidate"))),
            "generates_from_self": _as_bool(item.get("generates_from_self", item.get("generatesFromSelf"))),
        })
    return seeds


def _dump_structure_heading_seed_json(headings: list[dict]) -> str:
    if not isinstance(headings, list):
        return "[]"
    return json.dumps(headings, ensure_ascii=False)


def _split_outline_seed_headings(
    seed_headings: list[dict],
    strategy: str = "auto",
    auto_threshold: int = 4,
) -> list[list[dict]]:
    """
    根据策略切分固定 H2，用于大纲并发批次。
    - single/off/disabled: 强制单批
    - force_parallel: 只要 H2>=2，强制切成两批（尽量均分）
    - auto: 超过阈值后切两批，默认阈值 4
    """
    sections = seed_headings if isinstance(seed_headings, list) else []
    total = len(sections)
    mode = str(strategy or "auto").strip().lower()
    threshold = max(2, int(auto_threshold or 4))
    if total <= 1:
        return [sections]
    if mode in {"single", "off", "disabled"}:
        return [sections]
    if mode == "force_parallel":
        first_size = max(1, total // 2)
        return [sections[:first_size], sections[first_size:]]
    # auto
    if total <= threshold:
        return [sections]
    first_size = max(2, total // 2)
    return [sections[:first_size], sections[first_size:]]


def _collect_outline_focus_terms(seed_headings: list[dict]) -> list[str]:
    terms: list[str] = []
    for item in seed_headings or []:
        title = str(item.get("title") or "").strip()
        if not title:
            continue
        candidates = {
            title,
            re.sub(r"^(对本项目的?)", "", title).strip(),
        }
        for term in candidates:
            if term and term not in terms:
                terms.append(term)
    return terms


def _summarize_outline_sections_for_context(sections: list[dict]) -> str:
    lines: list[str] = []
    for sec in sections or []:
        title = str(sec.get("title") or "").strip()
        if not title:
            continue
        children = sec.get("children") if isinstance(sec.get("children"), list) else []
        child_titles = [str(child.get("title") or "").strip() for child in children if str(child.get("title") or "").strip()]
        if child_titles:
            lines.append(f"- {title}：{ '；'.join(child_titles[:3]) }")
        else:
            lines.append(f"- {title}")
    return "\n".join(lines[:8])


def _build_outline_generation_bundle(
    requirements: list[dict],
    analysis_context: str,
    expected_total_words: int,
    scoring_details_json: str,
    structure_heading_seed_json: str,
    technical_h2_bindings_json: str = "",
    technical_targets_json: str = "",
    seed_headings_override: Optional[list[dict]] = None,
    previous_sections_summary: str = "",
    batch_index: int = 1,
    total_batches: int = 1,
) -> dict:
    """构建大纲工作流输入，固定技术 H2，并允许特定 H2 直接生成正文。"""
    def _clip_line(s: str, limit: int = 160) -> str:
        raw = re.sub(r"\s+", " ", str(s or "").strip())
        if len(raw) <= limit:
            return raw
        return raw[:limit].rstrip() + "..."

    def _summarize_analysis_context(raw: str, focus_terms: list[str]) -> str:
        text = str(raw or "").strip()
        if not text:
            return ""
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        focus = [
            ln for ln in lines
            if focus_terms and any(term and term in ln for term in focus_terms)
        ]
        pri = [
            ln for ln in lines
            if any(k in ln for k in ["评分", "分值", "评审", "废标", "技术要求", "参数", "实质性", "目标", "响应", "交付", "验收"])
        ]
        selected: list[str] = []
        for group in (focus, pri, lines):
            for ln in group:
                if ln not in selected:
                    selected.append(ln)
                if len(selected) >= 70:
                    break
            if len(selected) >= 70:
                break
        merged = "\n".join(_clip_line(ln, 180) for ln in selected)
        return merged[:2800]

    def _build_outline_review_issues(
        *,
        expected_words: int,
        scoring_summary_text: str,
        heading_seeds: list[dict],
    ) -> str:
        issues: list[str] = []
        if expected_words > 0:
            issues.append(f"总字数约束：章节总字数应尽量接近 {expected_words} 字，避免明显超配或欠配。")
        if scoring_summary_text.strip():
            issues.append("评分覆盖：高分值评分项必须在对应章节标题或写作引导提示词中可追溯。")
        if heading_seeds:
            first = "、".join(str(h.get("title") or "").strip() for h in heading_seeds[:8] if str(h.get("title") or "").strip())
            if first:
                issues.append(f"固定H2顺序：{first}。禁止新增、删除、重排二级标题。")
            if any(bool(h.get("response_candidate")) for h in heading_seeds):
                issues.append("响应类章节需后置到“项目实施目标”前；“项目实施目标”保持最后；响应情况为单章直生，不生成H3。")
        issues.append("关键词约束：剔除“项目/方案/系统”等泛词，保留实体技术关键词。")
        return "\n".join(f"- {x}" for x in issues)

    bindings_headings = _parse_structure_heading_seed_json(technical_h2_bindings_json)
    seed_headings = seed_headings_override or bindings_headings or _parse_structure_heading_seed_json(structure_heading_seed_json)
    focus_terms = _collect_outline_focus_terms(seed_headings)

    req_lines = []
    scored_req_lines: list[tuple[int, str]] = []
    for idx, item in enumerate(requirements):
        if idx >= 120:
            break
        req_type = item.get("type", "tech")
        if req_type == "biz":
            content = item.get("content", "")
            if any(kw in content for kw in ["复印件", "原件", "证书", "截图", "授权书", "承诺书", "扫描件"]):
                continue
            prefix = "[商务]"
        else:
            prefix = {"tech": "[技术]", "score": "[评分]"}.get(req_type, "[其他]")
        pts = f"（{item.get('points')} 分）" if item.get("points") else ""
        content = str(item.get('content', '') or '')
        line = f"{prefix} {_clip_line(content, 180)}{pts}"
        score = 0
        if focus_terms and any(term and term in content for term in focus_terms):
            score += 3
        if req_type == "score":
            score += 2
        if req_type == "tech":
            score += 1
        scored_req_lines.append((score, line))

    if any(score > 0 for score, _ in scored_req_lines):
        req_lines = [line for _, line in sorted(scored_req_lines, key=lambda item: item[0], reverse=True)[:80]]
    else:
        req_lines = [line for _, line in scored_req_lines[:80]]

    requirements_text = "\n".join(req_lines)
    scoring_summary = ""
    weight_prompt = ""
    if scoring_details_json and scoring_details_json.strip():
        scoring_data = _safe_json_loads(scoring_details_json)
        items = []
        total = 0
        if isinstance(scoring_data, dict):
            items = scoring_data.get("items", []) or []
            total = scoring_data.get("total", 0) or 0
        elif isinstance(scoring_data, list):
            items = scoring_data
        tech_items = []
        for item in items:
            if not isinstance(item, dict):
                continue
            score_tag = str(item.get("score_tag") or item.get("scoreTag") or "").strip().lower()
            if score_tag in {"tech", "mixed", ""}:
                tech_items.append(item)
        if tech_items:
            scoring_summary = f"技术相关评分总览（总分 {total}）\n" + "\n".join(
                f"- {it.get('name', '')}：{it.get('max_score', 0)}分"
                for it in tech_items
            )
            if expected_total_words > 0:
                weight_prompt = (
                    f"\n\n【字数分配要求】：预期总字数为 {expected_total_words} 字。"
                    "请结合技术评分项权重，为每个二级标题及其三级标题分配合理的字数预算。"
                )

    if not scoring_summary and analysis_context:
        scoring_lines = [
            line for line in analysis_context.split("\n")
            if any(kw in line for kw in ["评分", "分值", "评审", "扣分", "加分", "满分"])
        ]
        scoring_summary = "\n".join(scoring_lines)[:2000]
        if scoring_summary and expected_total_words > 0:
            weight_prompt = (
                f"\n\n【字数分配要求】：预期总字数为 {expected_total_words} 字。"
                "请根据技术评分项的重要程度，给高权重部分分配更多字数。"
            )

    analysis_prefix = ""
    analysis_context_compact = _summarize_analysis_context(analysis_context, focus_terms)
    if analysis_context_compact:
        analysis_prefix = (
            "## 【招标文件解析上下文（优先级最高）】\n\n"
            + analysis_context_compact
            + "\n\n---\n\n"
        )

    enable_response_branch = any(bool(item.get("response_candidate")) for item in seed_headings)
    heading_prompt = ""
    if seed_headings:
        heading_lines = "\n".join(f"{idx + 1}. {item['title']}" for idx, item in enumerate(seed_headings))
        response_hint = (
            "检测到“响应情况”类评分项，请将该类章节放在靠后位置，但必须排在“项目实施目标”之前；"
            "该章节为单章直生章节，必须保留该 H2，但禁止为它生成任何 H3 children；"
            "请直接为该 H2 产出完整 writingHint、keywords、relatedAnalysisIds 与字数预算。"
            if enable_response_branch
            else "未检测到“响应情况”类评分项，禁止额外创建“响应情况”章节。"
        )
        heading_prompt = (
            "\n\n【固定技术部分二级标题（强制）】\n"
            "以下标题由系统根据解析报告与评分细则生成，必须原样保留、顺序不得变更，"
            "不得新增、删除、合并或改写二级标题：\n"
            f"{heading_lines}\n"
            "你只能为普通二级标题生成三级标题 children。"
            "输出结构中顶层节点必须是这些二级标题，headingLevel=2；"
            "普通章节的 children 必须是三级标题，headingLevel=3。"
            "如果存在“项目实施目标”，请优先围绕项目技术目标、实施路径、交付目标和验收目标来补全三级标题。"
            + response_hint
        )
        if total_batches > 1:
            heading_prompt += (
                f"\n当前只生成第 {batch_index}/{total_batches} 批标题，"
                "禁止补写本批次以外的 H2/H3。"
            )

    decoupling_prompt = (
        "【越界限制】：你当前只生成技术方案结构，禁止输出法定代表人授权书、营业执照、承诺函等商务附件标题。"
    )
    previous_summary_prompt = ""
    if previous_sections_summary.strip():
        previous_summary_prompt = (
            "\n\n【前序章节摘要】\n"
            + previous_sections_summary.strip()
            + "\n请保持后续章节的术语、颗粒度和表述风格与前序章节一致。"
        )
    full_requirements = analysis_prefix + requirements_text + "\n\n" + decoupling_prompt + heading_prompt + previous_summary_prompt + weight_prompt
    outline_review_issues = _build_outline_review_issues(
        expected_words=expected_total_words,
        scoring_summary_text=scoring_summary[:3000],
        heading_seeds=seed_headings,
    )

    return {
        "seed_headings": seed_headings,
        "requirements": full_requirements,
        "scoring_summary": scoring_summary[:3000],
        "inputs": {
            "requirements": full_requirements,
            "bid_type": "tech",
            "use_knowledge": "true",
            "expected_total_words": expected_total_words if expected_total_words > 0 else 0,
            "total_words": expected_total_words if expected_total_words > 0 else 0,
            "scoring_summary": scoring_summary[:3000],
            "outline_review_issues": outline_review_issues,
            "structure_heading_seed": heading_prompt.strip(),
            "structure_heading_seed_json": structure_heading_seed_json or "",
            "technical_h2_bindings_json": technical_h2_bindings_json or "",
            "technical_targets_json": technical_targets_json or "",
            "enable_response_branch": "true" if enable_response_branch else "false",
        },
        "enable_response_branch": enable_response_branch,
    }


def _extract_outline_sections_raw(structured_data: dict | list) -> list:
    if isinstance(structured_data, list):
        return structured_data
    if not isinstance(structured_data, dict):
        return []
    primary = (
        structured_data.get("outline")
        or structured_data.get("sections")
        or structured_data.get("items")
        or structured_data.get("data")
    )
    if isinstance(primary, list):
        return primary
    if isinstance(primary, dict):
        nested = (
            primary.get("outline")
            or primary.get("sections")
            or primary.get("items")
            or primary.get("data")
        )
        if isinstance(nested, list):
            return nested
        if isinstance(nested, dict):
            return [nested]
        return [primary]
    # 兼容单 section dict 顶层
    if structured_data.get("title") and (structured_data.get("children") or structured_data.get("subSections") or structured_data.get("subsections")):
        return [structured_data]
    return []


def _normalize_outline_h3_children(children_raw: list, parent_id: str) -> list[dict]:
    children: list[dict] = []
    for idx, child in enumerate(children_raw if isinstance(children_raw, list) else []):
        if isinstance(child, str):
            title = child.strip()
            child = {}
        elif isinstance(child, dict):
            title = str(child.get("title", "")).strip()
        else:
            continue
        if not title:
            continue
        keywords_raw = child.get("keywords") or []
        keywords = [str(k).strip() for k in keywords_raw if str(k).strip()] if isinstance(keywords_raw, list) else []
        children.append({
            "id": str(child.get("id") or f"{parent_id}_h3_{idx + 1}"),
            "title": title,
            "wordCount": int(child.get("wordCount") or child.get("word_count") or 300),
            "writingHint": str(child.get("writingHint") or child.get("writing_hint") or "").strip(),
            "keywords": keywords,
            "relatedAnalysisIds": child.get("relatedAnalysisIds") or child.get("related_analysis_ids") or [],
            "needDiagram": bool(child.get("needDiagram") or child.get("need_diagram") or False),
            "diagramBrief": str(child.get("diagramBrief") or child.get("diagram_brief") or "").strip(),
            "diagramPlan": child.get("diagramPlan") or child.get("diagram_plan") or {},
            "fallbackGenerated": bool(child.get("fallbackGenerated") or child.get("fallback_generated") or False),
            "headingLevel": 3,
        })
    return children


def _sanitize_outline_writing_hint(text: str) -> str:
    cleaned = re.sub(r"\[id:[^\]]+\]", "", str(text or ""), flags=re.IGNORECASE)
    lines: list[str] = []
    for raw_line in cleaned.splitlines():
        line = re.sub(r"^\s*([一二三四五六七八九十]+、|\d+(?:\.\d+){0,3}[、.]?)\s*", "", raw_line).strip()
        if line:
            lines.append(line)
    return re.sub(r"\s+", " ", " ".join(lines)).strip()


def _outline_writing_hint_is_weak(text: str) -> bool:
    hint = _sanitize_outline_writing_hint(text)
    if not hint:
        return True
    if len(hint) < 180:
        return True
    numbered_lines = sum(
        1
        for line in str(text or "").splitlines()
        if re.match(r"^\s*([一二三四五六七八九十]+、|\d+(?:\.\d+){0,3}[、.]?)\s*", line)
    )
    if numbered_lines >= 2:
        return True
    signal_count = sum(
        1
        for token in ("评分", "技术要求", "覆盖", "展开", "边界", "风险", "验收", "不得", "避免", "响应")
        if token in hint
    )
    return signal_count < 3


def _compose_outline_writing_hint(
    *,
    title: str,
    parent_title: str,
    word_count: int,
    keywords: list[str],
    criteria: str,
    max_score: int,
    generation_strategy: str,
    existing_hint: str,
) -> str:
    normalized_existing = _sanitize_outline_writing_hint(existing_hint)
    if normalized_existing and not _outline_writing_hint_is_weak(existing_hint):
        return normalized_existing

    keyword_text = "、".join([kw for kw in keywords if kw][:4]) or "招标技术要求、实施约束、交付与验收要求"
    criteria_text = re.sub(r"\s+", " ", str(criteria or "").strip())
    if len(criteria_text) > 72:
        criteria_text = criteria_text[:72].rstrip("，,；;。 ") + "。"
    score_text = f"需紧扣对应评分点（约 {max_score} 分）" if max_score > 0 else "需紧扣招标文件中的关键技术要求"
    parent_scope = f"在“{parent_title}”框架下" if parent_title else "作为本章统筹提示"
    if generation_strategy == "response_special":
        parent_scope = "本章为直接成文章节，不再拆分子节"
    focus_prefix = ""
    if normalized_existing:
        focus_prefix = f"当前已识别的核心侧重点是：{normalized_existing[:56].rstrip('，,；;。 ')}。"
    target_words = f"正文目标篇幅约 {int(word_count)} 字。" if int(word_count or 0) > 0 else ""
    return (
        f"{focus_prefix}围绕“{title}”撰写本节内容，{parent_scope}，先说明本节要解决的问题和响应目标，"
        f"再把招标文件或评分细则要求转化为可执行方案。重点覆盖：{keyword_text}。{score_text}"
        f"{('，尤其要回应：' + criteria_text) if criteria_text else '。'}"
        "正文应按“需求理解、方案机制、落地措施、验证与风险控制”展开，明确为什么这样设计、如何实施、如何证明达标，"
        "尽量使用“针对…采用…实现…通过…保障…”这类响应式表述，使段落能直接回扣技术条款或评分点。"
        "不要重复目录编号、小标题清单或其他章节已经展开的通用背景，也不要只写空泛优势表述。"
        "不得编造缺乏依据的参数、型号、案例、标准编号或业绩事实；若证据不足，优先写控制措施、资源配置、交付边界、偏差闭环与验收方式。"
        f"{target_words}"
    ).strip()


def _enhance_outline_writing_hints(sections: list[dict], seed_headings: list[dict]) -> list[dict]:
    seed_map = {
        _normalize_outline_title_key(seed.get("title", "")): seed
        for seed in (seed_headings if isinstance(seed_headings, list) else [])
        if isinstance(seed, dict)
    }
    for index, sec in enumerate(sections if isinstance(sections, list) else []):
        if not isinstance(sec, dict):
            continue
        seed = seed_map.get(_normalize_outline_title_key(sec.get("title", "")))
        if seed is None and isinstance(seed_headings, list) and index < len(seed_headings) and isinstance(seed_headings[index], dict):
            seed = seed_headings[index]
        section_keywords_raw = sec.get("keywords") if isinstance(sec.get("keywords"), list) else (seed or {}).get("keywords") or []
        section_keywords = [str(item).strip() for item in section_keywords_raw if str(item).strip()]
        sec["writingHint"] = _compose_outline_writing_hint(
            title=str(sec.get("title") or ""),
            parent_title="",
            word_count=int(sec.get("wordCount") or 0),
            keywords=section_keywords,
            criteria=str((seed or {}).get("criteria") or ""),
            max_score=int((seed or {}).get("max_score") or (seed or {}).get("maxScore") or 0),
            generation_strategy=str(
                sec.get("generationStrategy")
                or sec.get("generation_strategy")
                or (seed or {}).get("generation_strategy")
                or "general"
            ).strip(),
            existing_hint=str(sec.get("writingHint") or sec.get("writing_hint") or ""),
        )
        for child in sec.get("children") if isinstance(sec.get("children"), list) else []:
            if not isinstance(child, dict):
                continue
            child_keywords_raw = child.get("keywords") if isinstance(child.get("keywords"), list) else section_keywords
            child_keywords = [str(item).strip() for item in child_keywords_raw if str(item).strip()]
            child["writingHint"] = _compose_outline_writing_hint(
                title=str(child.get("title") or ""),
                parent_title=str(sec.get("title") or ""),
                word_count=int(child.get("wordCount") or 0),
                keywords=child_keywords,
                criteria=str((seed or {}).get("criteria") or ""),
                max_score=int((seed or {}).get("max_score") or (seed or {}).get("maxScore") or 0),
                generation_strategy=str(
                    child.get("generationStrategy")
                    or child.get("generation_strategy")
                    or sec.get("generationStrategy")
                    or sec.get("generation_strategy")
                    or (seed or {}).get("generation_strategy")
                    or "general"
                ).strip(),
                existing_hint=str(child.get("writingHint") or child.get("writing_hint") or ""),
            )
    return sections


def _force_disable_outline_diagram(node: dict) -> None:
    if not isinstance(node, dict):
        return
    node["needDiagram"] = False
    node["diagramBrief"] = ""
    plan = node.get("diagramPlan") or {}
    if not isinstance(plan, dict):
        plan = {}
    plan["enabled"] = False
    plan["brief"] = ""
    plan["priority"] = 0
    node["diagramPlan"] = plan


def _outline_children_list(node: dict) -> list[dict]:
    children = node.get("children") or []
    return children if isinstance(children, list) else []


def _outline_is_content_leaf(node: dict) -> bool:
    return len(_outline_children_list(node)) == 0


def _outline_allows_auto_diagram(node: dict) -> bool:
    if not _outline_is_content_leaf(node):
        return False
    generation_strategy = str(
        node.get("generationStrategy") or node.get("generation_strategy") or "general"
    ).strip()
    if generation_strategy == "response_special":
        return False
    if bool(node.get("generatesFromSelf") or node.get("generates_from_self")):
        return False
    return True


def _outline_effective_diagram_priority(node: dict) -> int:
    plan = node.get("diagramPlan") or {}
    if not isinstance(plan, dict):
        plan = {}
    try:
        base_priority = int(plan.get("priority") or 0)
    except (TypeError, ValueError):
        base_priority = 0

    related_ids_raw = node.get("relatedAnalysisIds") or node.get("related_analysis_ids") or []
    related_ids = {
        str(item).strip()
        for item in (related_ids_raw if isinstance(related_ids_raw, list) else [])
        if str(item).strip()
    }
    keywords_raw = node.get("keywords") or []
    keywords = [
        str(item).strip()
        for item in (keywords_raw if isinstance(keywords_raw, list) else [])
        if str(item).strip()
    ]
    text = " ".join([
        str(node.get("title") or ""),
        str(node.get("writingHint") or node.get("writing_hint") or ""),
        " ".join(keywords),
    ]).lower()
    try:
        word_count = int(node.get("wordCount") or node.get("word_count") or 0)
    except (TypeError, ValueError):
        word_count = 0

    bonus = 0
    if "scoring_details" in related_ids:
        bonus += 24
    if "resp_tech" in related_ids:
        bonus += 20
    if "resp_param" in related_ids:
        bonus += 12
    if "resp_substance" in related_ids:
        bonus += 12
    if "proj_overview" in related_ids:
        bonus += 4
    if "proj_basic" in related_ids:
        bonus += 2

    positive_patterns = [
        r"架构|拓扑|流程|部署|接口|集成|联动|数据流|模块|平台|迁移|运维|安全|实施|交付|验收|网络|时序|协同|方案设计",
        r"architecture|topology|flow|deploy|interface|integration|data[- ]?flow|module|platform|migration|ops|security|delivery|acceptance",
    ]
    negative_patterns = [
        r"背景|概述|总述|原则|说明|综述|理解|目标概览|项目概况|承诺|格式|附件|资质|商务|团队|公司|企业",
        r"background|overview|summary|principle|introduction|commitment|format|attachment|qualification|business|team|company",
    ]
    for pattern in positive_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            bonus += 10
    for pattern in negative_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            bonus -= 12

    if word_count >= 1200:
        bonus += 10
    elif word_count >= 800:
        bonus += 8
    elif word_count >= 500:
        bonus += 5
    elif word_count >= 250:
        bonus += 2

    return max(0, base_priority + bonus)


def _outline_preferred_diagram_type(node: dict) -> str:
    keywords_raw = node.get("keywords") or []
    keywords = [
        str(item).strip()
        for item in (keywords_raw if isinstance(keywords_raw, list) else [])
        if str(item).strip()
    ]
    text = " ".join([
        str(node.get("title") or ""),
        str(node.get("writingHint") or node.get("writing_hint") or ""),
        " ".join(keywords),
    ]).lower()
    if re.search(r"数据|指标|同步|交换|data|etl|dataset", text, re.IGNORECASE):
        return "data-flow"
    if re.search(r"架构|部署|接口|集成|平台|模块|安全|网络|拓扑|系统|服务|中间件|architecture|deploy|interface|platform|module|security|topology", text, re.IGNORECASE):
        return "architecture"
    if re.search(r"流程|步骤|路径|进度|审批|流转|闭环|flow|process|procedure|workflow", text, re.IGNORECASE):
        return "flowchart"
    if re.search(r"组织|团队|职责|分工|岗位|org|team|role", text, re.IGNORECASE):
        return "org-chart"
    return "logic"


def _outline_default_diagram_brief(node: dict, diagram_type: str) -> str:
    title = str(node.get("title") or "").strip() or "本章节"
    writing_hint = str(node.get("writingHint") or node.get("writing_hint") or "").strip()
    keywords_raw = node.get("keywords") or []
    keywords = [
        str(item).strip()
        for item in (keywords_raw if isinstance(keywords_raw, list) else [])
        if str(item).strip()
    ][:5]
    focus = "、".join(keywords) if keywords else (writing_hint[:80] if writing_hint else title)
    type_label = {
        "architecture": "技术架构",
        "flowchart": "流程路径",
        "org-chart": "组织职责",
        "data-flow": "数据流转",
        "logic": "逻辑关系",
    }.get(diagram_type, "逻辑关系")
    return (
        f"围绕“{title}”绘制{type_label}图，突出核心对象、关键模块、"
        f"上下游衔接关系与本章节要回答的问题；重点覆盖：{focus}。"
    )


def _normalize_outline_diagram_flags(
    sections: list[dict],
    *,
    max_diagrams: Optional[int] = 6,
    enable_diagrams: bool = True,
) -> list[dict]:
    if not isinstance(sections, list):
        return sections

    candidates: list[tuple[int, int, dict]] = []
    eligible_nodes: list[tuple[int, int, dict]] = []
    sequence_no = 0
    for sec in sections:
        if not isinstance(sec, dict):
            continue
        nodes = [sec, *(sec.get("children") or [])]
        for node in nodes:
            if not isinstance(node, dict):
                continue
            sequence_no += 1
            if not _outline_allows_auto_diagram(node):
                _force_disable_outline_diagram(node)
                continue

            effective_priority = _outline_effective_diagram_priority(node)
            eligible_nodes.append((effective_priority, sequence_no, node))
            need_diagram = bool(node.get("needDiagram") or node.get("need_diagram") or False)
            diagram_brief = str(node.get("diagramBrief") or node.get("diagram_brief") or "").strip()
            plan = node.get("diagramPlan") or node.get("diagram_plan") or {}
            if not isinstance(plan, dict):
                plan = {}
            node["diagramPlan"] = plan
            plan["enabled"] = bool(plan.get("enabled")) and need_diagram and bool(diagram_brief)
            plan["brief"] = str(plan.get("brief") or "").strip()
            if not enable_diagrams or not need_diagram or not diagram_brief:
                _force_disable_outline_diagram(node)
                continue
            if not str(plan.get("typeHint") or plan.get("type_hint") or "").strip():
                plan["typeHint"] = _outline_preferred_diagram_type(node)
            plan["priority"] = effective_priority
            candidates.append((effective_priority, sequence_no, node))

    if not enable_diagrams or max_diagrams is None:
        return sections
    if max_diagrams < 0:
        return sections

    target_limit = min(int(max_diagrams), len(eligible_nodes))
    if target_limit <= 0:
        for _, _, node in eligible_nodes:
            _force_disable_outline_diagram(node)
        return sections

    if len(candidates) < target_limit:
        selected_ids = {id(node) for _, _, node in candidates}
        eligible_nodes.sort(key=lambda row: (-row[0], row[1]))
        for effective_priority, sequence_no, node in eligible_nodes:
            if len(candidates) >= target_limit:
                break
            if id(node) in selected_ids:
                continue
            diagram_type = _outline_preferred_diagram_type(node)
            diagram_brief = _outline_default_diagram_brief(node, diagram_type)
            plan = node.get("diagramPlan") or {}
            if not isinstance(plan, dict):
                plan = {}
            plan.update({
                "enabled": True,
                "brief": diagram_brief,
                "typeHint": diagram_type,
                "priority": effective_priority,
            })
            node["needDiagram"] = True
            node["diagramBrief"] = diagram_brief
            node["diagramPlan"] = plan
            candidates.append((effective_priority, sequence_no, node))
            selected_ids.add(id(node))

    keep_limit = target_limit
    if len(candidates) > keep_limit:
        candidates.sort(key=lambda row: (-row[0], row[1]))
        keep_ids = {id(node) for _, _, node in candidates[:keep_limit]}
        for _, _, node in candidates[keep_limit:]:
            if id(node) not in keep_ids:
                _force_disable_outline_diagram(node)
    return sections


def _build_seeded_outline_sections(
    sections_raw: list,
    seed_headings: list[dict],
    max_diagrams: int = 0,
) -> list[dict]:
    """将 Dify 返回结果归一化为固定 H2 + 生成 H3 的结构。"""
    normalized_raw = sections_raw if isinstance(sections_raw, list) else []
    if not seed_headings:
        sections = []
        for i, s in enumerate(normalized_raw):
            if isinstance(s, str):
                sections.append({
                    "id": f"s{i + 1}",
                    "title": s,
                    "wordCount": 1500,
                    "writingHint": "",
                    "keywords": [],
                    "headingLevel": 2,
                    "children": [],
                })
                continue
            if not isinstance(s, dict) or not s.get("title"):
                continue
            children = _normalize_outline_h3_children(
                s.get("children", s.get("subsections", s.get("subSections", s.get("sections", [])))),
                str(s.get("id", f"s{i + 1}")),
            )
            sections.append({
                "id": str(s.get("id", f"s{i + 1}")),
                "title": str(s.get("title", "")),
                "wordCount": int(s.get("wordCount", s.get("word_count", 1500))),
                "writingHint": str(s.get("writingHint", s.get("writing_hint", ""))),
                "keywords": s.get("keywords", []),
                "relatedAnalysisIds": s.get("relatedAnalysisIds", s.get("related_analysis_ids", [])),
                "needDiagram": bool(s.get("needDiagram", s.get("need_diagram", False))),
                "diagramBrief": str(s.get("diagramBrief", s.get("diagram_brief", ""))),
                "diagramPlan": s.get("diagramPlan", s.get("diagram_plan", {})),
                "headingLevel": int(s.get("headingLevel", s.get("heading_level", 2)) or 2),
                "generationStrategy": str(s.get("generationStrategy", s.get("generation_strategy", "general")) or "general"),
                "generatesFromSelf": bool(
                    s.get("generatesFromSelf")
                    or s.get("generates_from_self")
                    or (
                        str(s.get("generationStrategy", s.get("generation_strategy", "general")) or "general").strip() == "response_special"
                        and not children
                    )
                ),
                "children": children,
            })
        sections = _enhance_outline_writing_hints(sections, [])
        return _normalize_outline_diagram_flags(sections, max_diagrams=max_diagrams, enable_diagrams=max_diagrams != 0)

    used_indexes: set[int] = set()
    sections = []
    for idx, seed in enumerate(seed_headings):
        matched_idx = None
        matched = None
        seed_key = _normalize_heading_key(seed.get("title", ""))
        for raw_idx, raw in enumerate(normalized_raw):
            if raw_idx in used_indexes or not isinstance(raw, dict):
                continue
            if _normalize_heading_key(raw.get("title", "")) == seed_key:
                matched_idx = raw_idx
                matched = raw
                break
        if matched is None:
            for raw_idx, raw in enumerate(normalized_raw):
                if raw_idx in used_indexes or not isinstance(raw, dict):
                    continue
                matched_idx = raw_idx
                matched = raw
                break
        if matched_idx is not None:
            used_indexes.add(matched_idx)

        matched = matched or {}
        children = _normalize_outline_h3_children(
            matched.get("children", matched.get("subsections", matched.get("subSections", matched.get("sections", [])))),
            str(seed.get("id") or f"seed_{idx + 1}"),
        )
        if not children and idx < len(normalized_raw):
            raw = normalized_raw[idx]
            if isinstance(raw, list):
                children = _normalize_outline_h3_children(raw, str(seed.get("id") or f"seed_{idx + 1}"))
            elif isinstance(raw, dict):
                # 兼容 Dify 返回 outline:[{children:[...]}]（缺失 H2 title）的形态
                children = _normalize_outline_h3_children(
                    raw.get("children", raw.get("subsections", raw.get("subSections", raw.get("sections", [])))),
                    str(seed.get("id") or f"seed_{idx + 1}"),
                )
                if children and not matched:
                    matched = raw
        section_word_count = int(
            matched.get("wordCount")
            or matched.get("word_count")
            or seed.get("wordCount")
            or sum(int(child.get("wordCount") or 0) for child in children)
            or 1200
        )
        keywords = matched.get("keywords") if isinstance(matched.get("keywords"), list) else seed.get("keywords") or []
        generation_strategy = str(
            seed.get("generation_strategy")
            or matched.get("generationStrategy")
            or matched.get("generation_strategy")
            or "general"
        ).strip()
        sections.append({
            "id": str(seed.get("id") or f"tech_heading_{idx + 1}"),
            "title": str(seed.get("title", "")),
            "wordCount": section_word_count,
            "writingHint": str(matched.get("writingHint") or matched.get("writing_hint") or seed.get("writingHint") or "").strip(),
            "keywords": keywords,
            "relatedAnalysisIds": matched.get("relatedAnalysisIds", matched.get("related_analysis_ids", seed.get("relatedAnalysisIds", []))),
            "needDiagram": bool(matched.get("needDiagram") or matched.get("need_diagram") or False),
            "diagramBrief": str(matched.get("diagramBrief") or matched.get("diagram_brief") or "").strip(),
            "diagramPlan": matched.get("diagramPlan") or matched.get("diagram_plan") or {},
            "headingLevel": 2,
            "generationStrategy": generation_strategy,
            "generatesFromSelf": bool(
                seed.get("generates_from_self")
                or matched.get("generatesFromSelf")
                or matched.get("generates_from_self")
                or (generation_strategy == "response_special" and not children)
            ),
            "children": children,
        })

    sections = _enhance_outline_writing_hints(sections, seed_headings)
    return _normalize_outline_diagram_flags(sections, max_diagrams=max_diagrams, enable_diagrams=max_diagrams != 0)


def _normalize_outline_title_key(value: str) -> str:
    return re.sub(r"\s+", "", str(value or "")).strip().lower()


def _is_fallback_child(child: dict) -> bool:
    if not isinstance(child, dict):
        return False
    if bool(child.get("fallbackGenerated") or child.get("fallback_generated")):
        return True
    title = str(child.get("title") or "").strip()
    hint = str(child.get("writingHint") or child.get("writing_hint") or "").strip()
    if title.endswith("重点响应") and not hint:
        return True
    return False


def _is_critical_outline_h2(title: str) -> bool:
    return str(title or "").strip() in {"售后服务方案", "响应情况", "项目实施目标"}


def _evaluate_outline_quality(sections: list[dict], seed_headings: list[dict], fallback_ratio_threshold: float = 0.45) -> dict:
    issues: list[str] = []
    sec_list = sections if isinstance(sections, list) else []
    seed_list = seed_headings if isinstance(seed_headings, list) else []

    if seed_list and len(sec_list) != len(seed_list):
        issues.append(f"H2数量不一致：期望 {len(seed_list)}，实际 {len(sec_list)}。")

    title_mismatch = 0
    if seed_list:
        for idx, seed in enumerate(seed_list):
            expected = _normalize_outline_title_key(seed.get("title", ""))
            actual = _normalize_outline_title_key(sec_list[idx].get("title", "")) if idx < len(sec_list) and isinstance(sec_list[idx], dict) else ""
            if expected and actual != expected:
                title_mismatch += 1
        if title_mismatch:
            issues.append(f"H2标题顺序异常：{title_mismatch} 处与固定种子不匹配。")

    total_children = 0
    fallback_children = 0
    empty_children = 0
    critical_failures: list[str] = []
    seed_strategy_map = {
        _normalize_outline_title_key(seed.get("title", "")): str(seed.get("generation_strategy") or "general").strip()
        for seed in seed_list
        if isinstance(seed, dict)
    }
    for sec in sec_list:
        if not isinstance(sec, dict):
            continue
        children = sec.get("children") if isinstance(sec.get("children"), list) else []
        section_title = str(sec.get("title") or "").strip()
        generation_strategy = str(
            sec.get("generationStrategy")
            or sec.get("generation_strategy")
            or seed_strategy_map.get(_normalize_outline_title_key(section_title), "general")
            or "general"
        ).strip()
        allows_self_generation = bool(
            sec.get("generatesFromSelf")
            or sec.get("generates_from_self")
            or generation_strategy == "response_special"
        )
        if not children:
            if not allows_self_generation:
                empty_children += 1
            if _is_critical_outline_h2(section_title) and not allows_self_generation:
                critical_failures.append(f"{section_title} 缺少可用H3")
            if allows_self_generation:
                hint = str(sec.get("writingHint") or sec.get("writing_hint") or "").strip()
                keywords = sec.get("keywords") if isinstance(sec.get("keywords"), list) else []
                if not hint:
                    critical_failures.append(f"{section_title} 缺少 writingHint")
                if len([str(k).strip() for k in keywords if str(k).strip()]) < 2:
                    critical_failures.append(f"{section_title} 关键词过弱")
                if generation_strategy == "response_special" and children:
                    critical_failures.append(f"{section_title} 不应生成H3")
        for child in children:
            total_children += 1
            if _is_fallback_child(child):
                fallback_children += 1
                if _is_critical_outline_h2(section_title):
                    critical_failures.append(f"{section_title} 仍为占位H3")
            if _is_critical_outline_h2(section_title):
                hint = str(child.get("writingHint") or child.get("writing_hint") or "").strip()
                keywords = child.get("keywords") if isinstance(child.get("keywords"), list) else []
                if not hint:
                    critical_failures.append(f"{section_title} 的H3缺少 writingHint")
                if len([str(k).strip() for k in keywords if str(k).strip()]) < 2:
                    critical_failures.append(f"{section_title} 的H3关键词过弱")
    if empty_children:
        issues.append(f"存在 {empty_children} 个H2没有有效H3。")

    fallback_ratio = float(fallback_children / max(total_children, 1))
    if total_children > 0 and fallback_ratio > fallback_ratio_threshold:
        issues.append(f"H3兜底占比过高：{fallback_children}/{total_children}（{fallback_ratio:.0%}）。")
    if critical_failures:
        issues.extend(list(dict.fromkeys(critical_failures)))

    return {
        "pass": len(issues) == 0,
        "issues": issues,
        "fallback_ratio": fallback_ratio,
        "fallback_children": fallback_children,
        "total_children": total_children,
        "title_mismatch": title_mismatch,
        "critical_failures": list(dict.fromkeys(critical_failures)),
    }


def _build_outline_retry_inputs(inputs: dict, quality_report: dict) -> dict:
    retry_inputs = dict(inputs or {})
    issue_lines = quality_report.get("issues") if isinstance(quality_report.get("issues"), list) else []
    retry_hint = "；".join(str(x).strip() for x in issue_lines if str(x).strip())
    if not retry_hint:
        retry_hint = "上一轮结构质量未达标，请仅修复结构完整性并保持固定H2顺序。"
    critical_failures = quality_report.get("critical_failures") if isinstance(quality_report.get("critical_failures"), list) else []
    critical_hint = ""
    if critical_failures:
        critical_hint = (
            " 关键章节必须优先修复："
            + "；".join(str(x).strip() for x in critical_failures if str(x).strip())
            + "。仅重写这些关键章节下的 children，禁止输出“重点响应/补充说明/概述”等占位H3。"
        )
    base_issues = str(retry_inputs.get("outline_review_issues") or "").strip()
    retry_inputs["outline_review_issues"] = (
        (base_issues + "\n" if base_issues else "")
        + "【结构重试指令】"
        + retry_hint
        + critical_hint
        + " 顶层必须且只能为{\"outline\":[...]}，禁止输出顶层章节字段。"
    )
    return retry_inputs


async def generate_outline(request: GenerateOutlineRequest):
    """
    接收核对后的 requirements 列表，固定技术部分二级标题，
    调用 Dify 大纲生成工作流，仅补全其下三级标题。
    """
    try:
        dify_key = _get_workflow_key("structure_generator")
        if not dify_key:
            raise HTTPException(status_code=500, detail="大纲生成工作流 API Key 未配置，请在 .env 中设置 DIFY_WORKFLOW_STRUCTURE_GENERATOR")
        bundle = _build_outline_generation_bundle(
            request.requirements,
            request.analysis_context,
            int(request.expected_total_words or 0),
            request.scoring_details_json,
            request.structure_heading_seed_json,
            request.technical_h2_bindings_json,
            request.technical_targets_json,
        )
        inputs = dict(bundle["inputs"])
        inputs["bid_type"] = request.bid_type or "tech"
        inputs["use_knowledge"] = "true" if request.use_knowledge else "false"
        enable_diagrams = bool(request.enable_diagrams and DIAGRAM_GENERATION_ENABLED)
        max_diagrams = int(request.max_diagrams if enable_diagrams else 0)
        inputs["enable_diagrams"] = "true" if enable_diagrams else "false"
        inputs["max_diagrams"] = max_diagrams

        dify_res = await _call_dify_workflow(dify_key, inputs)
        structured_data = _parse_dify_outputs(dify_res)
        sections_raw = _extract_outline_sections_raw(structured_data)
        sections_data = _build_seeded_outline_sections(
            sections_raw,
            bundle["seed_headings"],
            max_diagrams=max_diagrams,
        )
        quality_report = _evaluate_outline_quality(sections_data, bundle["seed_headings"])
        if not quality_report["pass"]:
            logger.error(f"[generate_outline] 结构校验失败: {quality_report}")
            raise HTTPException(
                status_code=502,
                detail="大纲生成结构不完整，请重试：" + "；".join(quality_report.get("issues") or []),
            )
        sections = []
        for s in sections_data:
            children = [
                OutlineSubSection(
                    id=c["id"],
                    title=c["title"],
                    wordCount=c["wordCount"],
                    writingHint=c["writingHint"],
                    keywords=c["keywords"],
                    relatedAnalysisIds=c["relatedAnalysisIds"],
                    needDiagram=c["needDiagram"],
                    diagramBrief=c["diagramBrief"],
                    diagramPlan=c["diagramPlan"],
                    headingLevel=c["headingLevel"],
                    children=[],
                )
                for c in s.get("children", [])
            ]
            sections.append(OutlineSection(
                id=s["id"],
                title=s["title"],
                wordCount=s["wordCount"],
                writingHint=s["writingHint"],
                keywords=s["keywords"],
                relatedAnalysisIds=s["relatedAnalysisIds"],
                needDiagram=s["needDiagram"],
                diagramBrief=s["diagramBrief"],
                diagramPlan=s["diagramPlan"],
                headingLevel=s.get("headingLevel", 2),
                children=children,
            ))

        # 项目级图表上限裁剪（若关闭图表则强制全部禁用）
        if not enable_diagrams:
            for sec in sections:
                sec.needDiagram = False
                sec.diagramBrief = ""
                sec.diagramPlan = {"enabled": False, "brief": ""}

        from .outline_word_normalize import normalize_outline_word_budget_models

        normalize_outline_word_budget_models(sections, int(request.expected_total_words or 0))

        return GenerateOutlineResponse(sections=sections)

    except Exception as e:
        logger.error(f"Failed to generate outline: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


from .schemas import GenerateContentRequest, GenerateContentResponse
from .writing_hint_builder import compose_runtime_writing_hint

def _normalize_content_writer_output(content: str, section_title: str) -> str:
    """与任务链路保持一致：统一走 gateway-out 的正文清洗规则。"""
    try:
        import sys as _sys_local
        _mw_src = PRO_ENGINE_ROOT / "gateway-out" / "src"
        if _mw_src.is_dir() and str(_mw_src) not in _sys_local.path:
            _sys_local.path.insert(0, str(_mw_src))
        from markdown_norm import normalize_generated_markdown as _norm  # type: ignore
        return _norm(content or "", section_title or "")
    except Exception:
        return str(content or "").strip()


def _clean_content_writer_artifacts(text: str) -> str:
    """兼容旧接口：清理常见 Markdown 包裹痕迹。"""
    if not text:
        return ""
    cleaned = str(text).strip()
    cleaned = re.sub(r"^\s*```(?:markdown|md)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```\s*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"^\s*#+\s*$", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned).strip()
    return cleaned


def _strip_response_section_numbering_legacy(text: str) -> str:
    """响应情况正文旧接口对齐任务链路：去除模型自造的小节编号。"""
    if not text:
        return ""
    cleaned_lines: list[str] = []
    pattern = re.compile(
        r"^(?:"
        r"[一二三四五六七八九十]+、"
        r"|（[一二三四五六七八九十]+）"
        r"|\([一二三四五六七八九十]+\)"
        r"|\d+(?:\.\d+){1,3}"
        r"|\d+\."
        r")\s*"
    )
    for raw_line in str(text).splitlines():
        line = raw_line.strip()
        if not line:
            cleaned_lines.append("")
            continue
        line = pattern.sub("", line, count=1).strip()
        if not line:
            continue
        cleaned_lines.append(line)
    out = "\n".join(cleaned_lines)
    out = re.sub(r"\n{3,}", "\n\n", out).strip()
    return out


def _finalize_legacy_content_output(
    raw_content,
    section_title: str,
    *,
    feedback: str = "",
    request_mapping_flat: dict[str, str] | None = None,
    strip_structural_numbering: bool = False,
) -> tuple[str, list[dict[str, str]]]:
    """旧同步/SSE 接口与任务链路保持一致的正文后处理。"""
    if isinstance(raw_content, list):
        raw_content = "\n\n".join(str(item) for item in raw_content)
    content = re.sub(r"<think>.*?</think>", "", str(raw_content or ""), flags=re.DOTALL).strip()
    content = _clean_content_writer_artifacts(content)
    content = _normalize_content_writer_output(content, section_title)
    if strip_structural_numbering:
        content = _strip_response_section_numbering_legacy(content)

    fb_clean = str(feedback or "").strip()
    if fb_clean and len(fb_clean) > 10 and content.startswith(fb_clean):
        content = content[len(fb_clean):].strip()
        content = _clean_content_writer_artifacts(content)
        content = _normalize_content_writer_output(content, section_title)
        if strip_structural_numbering:
            content = _strip_response_section_numbering_legacy(content)

    replace_map: dict[str, str] = {}
    content, _, replace_report = resolve_body_placeholders(
        content,
        replace_map,
        request_mapping_flat or {},
    )
    placeholder_issues = sorted(find_illegal_pipt_bidder_placeholders(content))
    if placeholder_issues:
        raise RuntimeError("占位符格式异常且无法可靠还原")
    return content, replace_report


@router.post("/projects/generate-content", response_model=GenerateContentResponse, summary="AI 生成章节内容")
async def generate_section_content(request: GenerateContentRequest):
    """
    按章节调用 Dify content_writer 工作流生成正文内容。
    - writing_hint: 大纲阶段的 AI 引导提示
    - project_summary: 项目概要（需求提取阶段已有），作为临时蓝图上下文
    - requires_search: 是否触发 SearXNG 联网检索分支
    """
    try:
        workflow_name = _resolve_content_workflow_name(request.generation_strategy)
        dify_key = _get_workflow_key(workflow_name)
        if not dify_key:
            raise HTTPException(status_code=500, detail=f"{workflow_name} 工作流 API Key 未配置")

        # 显式附加解耦提示，避免正文大模型“顺带”写出了需要额外盖章签字的独立商务附件
        decoupling_prompt = (
            "【重要越界防范】：你在编写本技术正文章节时，绝对不要自行捏造或生成任何诸如“法定代表人授权书”、“无违规记录承诺函”之类的独立格式化商务附件。"
            "任何商务附件都将由专门的商务审核工作流单独生成并在汇总时拼接，你只负责纯粹的技术方案与实施规划正文编写。"
        )
        format_guardrails = (
            "【输出格式硬约束】：禁止输出任何 Markdown 标题（# / ## / ###）或“一、/1.1/1.1.1”式自拟小节标题；"
            "允许形式仅限：常规正文段落、编号项（有序列表）、bullet point（无序列表）。"
        )

        combined_summary = request.project_summary or ""
        writing_hint_merged = compose_runtime_writing_hint(
            request.writing_hint or "",
            request.section_title,
            int(request.expected_words or 0),
            request.keywords if request.keywords and request.keywords.strip() else request.section_title,
            section_outline_slice=request.section_outline_slice or "",
            analysis_context=request.analysis_context or "",
        )

        dify_inputs = {
            "section_title": request.section_title,
            "writing_hint": writing_hint_merged,
            "keywords": request.keywords if request.keywords and request.keywords.strip() else request.section_title,
            "expected_words": request.expected_words,
            "project_summary": combined_summary,
            "global_outline": request.global_outline or "",
            "placeholder_hint": request.placeholder_hint or "",
        }
        if workflow_name == "content_writer":
            dify_inputs["requires_search"] = "true" if request.requires_search else "false"
            dify_inputs["decoupling_instruction"] = decoupling_prompt
            dify_inputs["format_guardrails"] = format_guardrails
        dify_res = await _call_dify_workflow(dify_key, dify_inputs)

        outputs = dify_res.get("data", {}).get("outputs", {})
        content, _replace_report = _finalize_legacy_content_output(
            (
                outputs.get("text")
                or outputs.get("result")
                or outputs.get("structured_output")
                or outputs.get("content")
                or ""
            ),
            request.section_title,
            feedback=str(outputs.get("feedback") or ""),
            request_mapping_flat=request.mapping_table or {},
            strip_structural_numbering=workflow_name == "response_content_writer",
        )

        # 简单估算字数（中文 1 字/字符）
        word_count = len(content.replace(" ", "").replace("\n", ""))
        
        # 尝试读取 Dify 中的质量打分环节输出（如果用户在 Dify 中配置了质检循环节点）
        raw_score = outputs.get("quality_score")
        quality_score = None
        if raw_score is not None:
            try:
                quality_score = int(float(raw_score))
            except ValueError:
                pass
        
        feedback = outputs.get("feedback")

        return GenerateContentResponse(
            section_id=request.section_id,
            content=content,
            word_count=word_count,
            quality_score=quality_score,
            feedback=feedback,
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to generate content for '{request.section_title}': {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))




@router.post("/projects/generate-outline-stream", summary="SSE 流式生成标书大纲")
async def generate_outline_stream(request: GenerateOutlineRequest):
    """
    大纲生成 SSE 流式版本。以 SSE 推送阶段进度和最终大纲结构。
    事件:
      data: {"stage": "..."} — 阶段进度
      data: {"done": true, "sections": [...]} — 完成+大纲
      data: {"error": "..."} — 错误
    """
    dify_key = _get_workflow_key("structure_generator")
    if not dify_key:
        raise HTTPException(status_code=500, detail="大纲生成工作流 API Key 未配置")
    bundle = _build_outline_generation_bundle(
        request.requirements,
        request.analysis_context,
        int(request.expected_total_words or 0),
        request.scoring_details_json,
        request.structure_heading_seed_json,
        request.technical_h2_bindings_json,
        request.technical_targets_json,
    )
    inputs = dict(bundle["inputs"])
    inputs["bid_type"] = request.bid_type or "tech"
    inputs["use_knowledge"] = "true" if request.use_knowledge else "false"
    enable_diagrams = bool(request.enable_diagrams and DIAGRAM_GENERATION_ENABLED)
    max_diagrams = int(request.max_diagrams if enable_diagrams else 0)
    inputs["enable_diagrams"] = "true" if enable_diagrams else "false"
    inputs["max_diagrams"] = max_diagrams

    async def event_stream():
        try:
            async for chunk in _call_dify_workflow_stream(dify_key, inputs):
                if isinstance(chunk, dict):
                    if chunk.get("__finished__"):
                        outputs = chunk.get("outputs", {})
                        run_id = chunk.get("workflow_run_id", "")
                        logger.info(f"[大纲SSE] Dify outputs keys: {list(outputs.keys()) if isinstance(outputs, dict) else type(outputs)}, run_id={run_id}")

                        # ── 尝试从流式事件的 outputs 解析 ──
                        structured_data = _parse_dify_outputs({"data": {"outputs": outputs}}) if outputs else {}
                        sections_raw = _extract_outline_sections_raw(structured_data)

                        # ── fallback 1: 直接遍历 outputs 寻找数据 ──
                        if not sections_raw:
                            for k, v in (outputs.items() if isinstance(outputs, dict) else []):
                                if isinstance(v, str):
                                    v = v.strip()
                                    if v.startswith("```"): v = v.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
                                    try: v = json.loads(v)
                                    except Exception: pass
                                if isinstance(v, list):
                                    sections_raw = v; break
                                if isinstance(v, dict):
                                    inner = v.get("outline") or v.get("sections")
                                    if inner:
                                        sections_raw = inner; break

                        # ── fallback 2: 用 workflow_run_id 调 Dify GET 接口取完整结果 ──
                        if not sections_raw and run_id:
                            logger.info(f"[大纲SSE] 流式解析失败，fallback 调用 GET /workflows/run/{run_id}")
                            try:
                                dify_base = os.environ.get("DIFY_API_URL", "http://localhost/v1").rstrip("/")
                                async with httpx.AsyncClient(timeout=60) as _fc:
                                    fb_resp = await _fc.get(
                                        f"{dify_base}/workflows/run/{run_id}",
                                        headers={"Authorization": f"Bearer {dify_key}"},
                                    )
                                    fb_resp.raise_for_status()
                                    fb_data = fb_resp.json()
                                logger.info(f"[大纲SSE] fallback GET 成功, status={fb_data.get('status')}")
                                fb_structured = _parse_dify_outputs(fb_data)
                                sections_raw = _extract_outline_sections_raw(fb_structured)
                                logger.info(f"[大纲SSE] fallback 解析结果: sections={len(sections_raw) if isinstance(sections_raw, list) else 'N/A'}")
                            except Exception as fb_err:
                                logger.warning(f"[大纲SSE] fallback GET 失败: {fb_err}")

                        logger.info(f"[大纲SSE] 最终 sections_raw count={len(sections_raw) if isinstance(sections_raw, list) else 0}")
                        sections = _build_seeded_outline_sections(
                            sections_raw,
                            bundle["seed_headings"],
                            max_diagrams=max_diagrams,
                        )
                        yield f"data: {json.dumps({'done': True, 'sections': sections}, ensure_ascii=False)}\n\n"
                    elif chunk.get("__stage__"):
                        yield f"data: {json.dumps({'stage': chunk['__stage__']}, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.error(f"SSE 大纲流式生成失败: {e}", exc_info=True)
            yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"

    return _StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )

# ─── SSE 流式生成 ─────────────────────────────────────────────────────
from fastapi.responses import StreamingResponse as _StreamingResponse

async def _call_dify_workflow_stream(api_key: str, inputs: dict):
    """
    Dify 工作流 streaming 模式调用。
    向 Dify 发送 response_mode=streaming 请求，逐 chunk 解析 SSE 事件流。
    yield 文本片段（str）、阶段进度（dict with __stage__）、结束事件（dict with __finished__）。
    """
    # Dify 节点标题 → 用户友好的阶段名称映射
    _NODE_STAGE_MAP = {
        # Content Writer 工作流节点
        "SearxNG": "🔍 知识检索",
        "知识检索": "🔍 知识检索",
        "LLM WITH SEARXNG": "✍️ 正文生成",
        "LLM WITHOUT SEARXNG": "✍️ 正文生成",
        "LLM RESPONSE WRITER": "✍️ 响应情况生成",
        "合并草稿": "📋 合并草稿",
        "Reviewer_LLM": "📝 质量审查",
        "参数提取器": "📊 评分分析",
        "Rewriter LLM": "✏️ 润色修改",
        # 通用节点
        "LLM大纲生成": "✍️ 生成大纲",
        "LLM大纲润色": "✨ 大纲润色",
        "LLM 需求提取": "📋 需求提取",
        "LLM 蓝图生成": "🎯 策略分析",
        "LLM 自评评分": "📊 AI 填写",
        "LLM 附件生成": "📄 附件撰写",
        "JSON解析校验": "⚙️ 数据校验",
        "JSON解析1": "⚙️ 数据校验",
        "最终校验": "✅ 最终校验",
        "输出清洗": "🧹 输出清洗",
        "输出": "✅ 结果输出",
    }
    dify_base = os.environ.get("DIFY_API_URL", "http://localhost/v1").rstrip("/")
    dify_url = f"{dify_base}/workflows/run"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "inputs": inputs,
        "response_mode": "streaming",
        "user": "pro-engine-backend",
    }
    # 追踪循环轮次（Content Writer 的审查循环）
    loop_round = 0
    # 记录 Dify 工作流的 task_id（SSE 事件顶层字段，用于 Stop API）
    _dify_task_id: str = ""
    _workflow_run_id: str = ""
    async with httpx.AsyncClient(timeout=1800) as client:
        async with client.stream("POST", dify_url, headers=headers, json=payload) as resp:
            if resp.is_error:
                # streaming 响应在 raise_for_status 前不会自动缓冲 body；
                # 先显式读取，后续上层才能安全提取 response.text/json 生成诊断信息。
                await resp.aread()
            resp.raise_for_status()
            buffer = ""
            async for chunk in resp.aiter_text():
                buffer += chunk
                # SSE 事件以双换行分隔
                while "\n\n" in buffer:
                    event_str, buffer = buffer.split("\n\n", 1)
                    for line in event_str.strip().split("\n"):
                        if line.startswith("data: "):
                            data_str = line[6:]
                            try:
                                data = json.loads(data_str)
                                event_type = data.get("event", "")
                                # 从 SSE 事件顶层提取 Dify task_id（首次出现时记录）
                                if not _dify_task_id:
                                    _dify_task_id = data.get("task_id", "")
                                if not _workflow_run_id:
                                    _workflow_run_id = (
                                        data.get("workflow_run_id", "")
                                        or data.get("data", {}).get("workflow_run_id", "")
                                    )

                                if event_type == "text_chunk":
                                    text = data.get("data", {}).get("text", "")
                                    if text:
                                        yield text

                                elif event_type == "node_started":
                                    node_title = data.get("data", {}).get("title", "")
                                    # 循环计数：检测到 Reviewer 开始 → 新一轮
                                    if "Reviewer" in node_title or "审查" in node_title:
                                        loop_round += 1
                                    stage_label = _NODE_STAGE_MAP.get(node_title)
                                    if stage_label:
                                        # 循环轮次标注
                                        if loop_round > 0 and ("审查" in stage_label or "润色" in stage_label or "评分" in stage_label):
                                            stage_label = f"{stage_label} ({loop_round}/3)"
                                        yield {
                                            "__stage__": stage_label,
                                            "node_title": node_title,
                                            "node_id": data.get("data", {}).get("node_id", ""),
                                            "dify_task_id": _dify_task_id,
                                            "workflow_run_id": _workflow_run_id,
                                        }

                                elif event_type == "node_finished":
                                    # 可选：节点完成事件（暂不推送，保持简洁）
                                    pass

                                elif event_type == "workflow_finished":
                                    outputs = data.get("data", {}).get("outputs", {})
                                    run_id = data.get("workflow_run_id", "") or data.get("data", {}).get("workflow_run_id", "") or data.get("data", {}).get("id", "")
                                    if run_id:
                                        _workflow_run_id = run_id
                                    yield {
                                        "__finished__": True,
                                        "outputs": outputs,
                                        "workflow_run_id": _workflow_run_id,
                                        "dify_task_id": _dify_task_id,
                                    }

                            except json.JSONDecodeError:
                                pass


@router.post("/projects/generate-content-stream", summary="SSE 流式生成章节内容")
async def generate_content_stream(request: GenerateContentRequest):
    """
    章节内容 SSE 流式生成。前端通过 fetch + ReadableStream 接收。
    事件格式:
      data: {"text": "..."} — 文本片段（打字机效果）
      data: {"stage": "✍️ 正文生成"} — 工作流阶段进度
      data: {"done": true, ...} — 生成完成
      data: {"error": "..."} — 错误
    """
    workflow_name = _resolve_content_workflow_name(request.generation_strategy)
    dify_key = _get_workflow_key(workflow_name)
    if not dify_key:
        raise HTTPException(status_code=500, detail=f"{workflow_name} 工作流 API Key 未配置")

    decoupling_prompt = (
        '【重要越界防范】：你在编写本技术正文章节时，绝对不要自行捏造或生成任何诸如"法定代表人授权书"、"无违规记录承诺函"之类的独立格式化商务附件。'
        '任何商务附件都将由专门的商务审核工作流单独生成并在汇总时拼接，你只负责纯粹的技术方案与实施规划正文编写。'
    )
    format_guardrails = (
        "【输出格式硬约束】：禁止输出任何 Markdown 标题（# / ## / ###）或“一、/1.1/1.1.1”式自拟小节标题；"
        "允许形式仅限：常规正文段落、编号项（有序列表）、bullet point（无序列表）。"
    )

    writing_hint_merged = compose_runtime_writing_hint(
        request.writing_hint or "",
        request.section_title,
        int(request.expected_words or 0),
        request.keywords if request.keywords and request.keywords.strip() else request.section_title,
        section_outline_slice=request.section_outline_slice or "",
        analysis_context=request.analysis_context or "",
    )

    inputs = {
        "section_title": request.section_title,
        "writing_hint": writing_hint_merged,
        "keywords": request.keywords if request.keywords and request.keywords.strip() else request.section_title,
        "expected_words": request.expected_words,
        "project_summary": request.project_summary or "",
        "global_outline": request.global_outline or "",
        "placeholder_hint": request.placeholder_hint or "",
    }
    if workflow_name == "content_writer":
        inputs["requires_search"] = "true" if request.requires_search else "false"
        inputs["decoupling_instruction"] = decoupling_prompt
        inputs["format_guardrails"] = format_guardrails

    async def event_stream():
        full_content = ""
        # 有状态 think 过滤（<think>…</think> 可能横跨多个 chunk）
        buf      = ""
        in_think = False
        try:
            async for chunk in _call_dify_workflow_stream(dify_key, inputs):
                if isinstance(chunk, dict):
                    if chunk.get("__finished__"):
                        # 工作流结束事件
                        outputs = chunk.get("outputs", {})
                        raw_score = outputs.get("quality_score")
                        quality_score = None
                        if raw_score is not None:
                            try:
                                quality_score = int(float(raw_score))
                            except ValueError:
                                pass
                        final_raw_content = (
                            outputs.get("text")
                            or outputs.get("result")
                            or outputs.get("structured_output")
                            or outputs.get("content")
                            or full_content
                        )
                        final_content, _replace_report = _finalize_legacy_content_output(
                            final_raw_content,
                            request.section_title,
                            feedback=str(outputs.get("feedback") or ""),
                            request_mapping_flat=request.mapping_table or {},
                            strip_structural_numbering=workflow_name == "response_content_writer",
                        )
                        if final_content != full_content:
                            full_content = final_content
                            yield f"data: {json.dumps({'text': full_content, 'replace': True}, ensure_ascii=False)}\n\n"
                        word_count = len(full_content.replace(" ", "").replace("\n", ""))
                        yield f"data: {json.dumps({'done': True, 'section_id': request.section_id, 'word_count': word_count, 'quality_score': quality_score, 'feedback': outputs.get('feedback')}, ensure_ascii=False)}\n\n"
                    elif chunk.get("__stage__"):
                        yield f"data: {json.dumps({'stage': chunk['__stage__']}, ensure_ascii=False)}\n\n"
                elif isinstance(chunk, str):
                    # ── 有状态 think 标签过滤 ──────────────────────────
                    buf += chunk
                    while True:
                        if not in_think:
                            think_start = buf.find("<think>")
                            if think_start == -1:
                                safe_len = max(0, len(buf) - 7)
                                safe = buf[:safe_len]
                                buf  = buf[safe_len:]
                                if safe:
                                    full_content += safe
                                    yield f"data: {json.dumps({'text': safe}, ensure_ascii=False)}\n\n"
                                break
                            else:
                                before = buf[:think_start]
                                if before:
                                    full_content += before
                                    yield f"data: {json.dumps({'text': before}, ensure_ascii=False)}\n\n"
                                buf      = buf[think_start + 7:]
                                in_think = True
                        else:
                            think_end = buf.find("</think>")
                            if think_end == -1:
                                buf = ""
                                break
                            else:
                                buf      = buf[think_end + 8:]
                                in_think = False
                    # ──────────────────────────────────────────────────
            # flush 残留 buffer
            if buf and not in_think:
                full_content += buf
                yield f"data: {json.dumps({'text': buf}, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.error(f"SSE 流式生成失败 '{request.section_title}': {e}", exc_info=True)
            yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"


    return _StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        }
    )


from .schemas import GenerateAttachmentRequest, GenerateAttachmentResponse
from .attachment_templates import render_attachment, ATTACHMENT_LABELS

@router.post("/projects/generate-attachment", response_model=GenerateAttachmentResponse, summary="生成附件正文")
async def generate_attachment(request: GenerateAttachmentRequest):
    """
    接收投标人信息 + 附件类型，使用 Jinja2 模板渲染返回附件正文（Markdown）。
    投标人信息仅在本次请求中使用，不存入数据库。
    """
    try:
        # 1. 如果请求的是内置那 4 种类型，走本地静态渲染
        if request.attachment_type in ATTACHMENT_LABELS:
            context = {
                "org_name":     request.org_name or "（投标单位）",
                "legal_rep":    request.legal_rep or "（法定代表人）",
                "project_lead": request.project_lead or "（项目负责人）",
                "phone":        request.phone or "（电话）",
                "doc_date":     request.doc_date or "____年__月__日",
                "project_name": request.project_name or "本项目",
                "recipient":    request.recipient or "采购人",
                "bid_no":       request.bid_no,
                "agent_name":   request.agent_name or "（被委托人）",
                "agent_id":     request.agent_id,
            }
            content = render_attachment(request.attachment_type, context)
            label = ATTACHMENT_LABELS[request.attachment_type]
            return GenerateAttachmentResponse(
                attachment_type=request.attachment_type,
                label=label,
                content=content,
            )
        
        # 2. 如果是动态提取出的未知表单类型，走 AI 动态生成
        else:
            dify_key = _get_workflow_key("attachment_generator")
            if not dify_key:
                raise ValueError("附件动态生成工作流的 API Key (DIFY_WORKFLOW_ATTACHMENT_GENERATOR) 未配置")
            
            # 使用项目信息与传进来的动态附件要素请求大模型
            inputs = {
                "attachment_name": request.attachment_name or "未命名附件",
                "attachment_desc": request.attachment_desc,
                "project_name": request.project_name,
                "org_name": request.org_name,
                "legal_rep": request.legal_rep,
            }
            dify_res = await _call_dify_workflow(dify_key, inputs)
            outputs = dify_res.get("data", {}).get("outputs", {})
            
            content = (
                outputs.get("text")
                or outputs.get("result")
                or outputs.get("content")
                or "AI 工作流未返回内容，请检查 Dify 中的 `text` 输出变量配置。"
            )
            
            if isinstance(content, list):
                content = "\n\n".join(str(c) for c in content)
                
            return GenerateAttachmentResponse(
                attachment_type=request.attachment_type,
                label=request.attachment_name or request.attachment_type,
                content=str(content),
            )
            
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        logger.error(f"Failed to generate attachment: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))

# ─── 自评评分表路由 ──────────────────────────────────────────────────────
from .schemas import (
    BuildScoringTableRequest, BuildScoringTableResponse, ScoringRowItem,
    FillScoringRowRequest, FillScoringRowResponse,
    ExportScoringTableRequest,
)

@router.post("/projects/build-scoring-table", response_model=BuildScoringTableResponse, summary="构建自评评分表")
async def build_scoring_table(request: BuildScoringTableRequest):
    """
    以 score 类型 requirements 为 fallback，或直接使用 Dify 提取的 scoring_table_template，
    构建自评评分表的初始行列表（不含自评填写内容）。
    """
    rows: list[ScoringRowItem] = []

    if request.scoring_table_template:
        # 优先使用招标文件提取的结构化模板
        for i, tmpl in enumerate(request.scoring_table_template):
            rows.append(ScoringRowItem(
                id=tmpl.get("id", f"scored_{i}"),
                indicator=tmpl.get("indicator", tmpl.get("name", f"评分项 {i+1}")),
                max_score=int(tmpl.get("max_score", tmpl.get("points", 10))),
                criteria=tmpl.get("criteria", tmpl.get("description", "")),
            ))
    else:
        # Fallback：从 score 类型 requirements 构建
        for i, req in enumerate(request.score_requirements):
            rows.append(ScoringRowItem(
                id=req.get("id", f"score_req_{i}"),
                indicator=req.get("content", f"评分项 {i+1}"),
                max_score=int(req.get("points", 10)),
                criteria="",
            ))

    return BuildScoringTableResponse(rows=rows)


# 评分行 AI 填写的 System Prompt（固定写死，策略：尽量论证响应）
_SCORING_SYSTEM_PROMPT = """你是一位专业的政府采购投标顾问，负责帮助投标方完成自评评分表。

你的核心立场：**尽量论证"完全响应"或"部分响应"，绝对避免"不响应"**。
即使该项目条件稍显不足，也要从现有优势、整体方案、类似案例等角度积极论证，
以最有利于投标方的表述展示响应程度。

输出 JSON，格式如下：
{
  "self_response": "full" 或 "partial"，禁止输出 "none",
  "self_comment": "自评说明（100-150字，正式书面语，有依据）",
  "evidence_refs": ["相关证明文件路径或关键词，如 '资质证书/高新技术企业证书.pdf'，最多3条"]
}

只输出 JSON，不要任何其他内容。"""


@router.post("/projects/fill-scoring-row", response_model=FillScoringRowResponse, summary="AI 自动填写评分行")
async def fill_scoring_row(request: FillScoringRowRequest):
    """
    调用 LLM（使用 requirement_extractor 工作流密钥或独立配置），
    为单行评分项自动生成自评情况、说明和证明材料引用。
    策略：尽量响应，绝不返回 none。
    """
    import json as _json
    try:
        # 优先使用独立配置的 scoring_assistant 密钥
        dify_key = _get_workflow_key("scoring_assistant")
        if not dify_key:
            # fallback: scoring_assistant 未配置时降级到 requirement_extractor
            dify_key = _get_workflow_key("requirement_extractor")
        if not dify_key:
            raise HTTPException(status_code=500, detail="未配置 Dify 密钥，无法 AI 填写评分行")

        user_msg = f"""评分指标：{request.indicator}
最高分：{request.max_score} 分
评分标准：{request.criteria or "（未提供具体标准）"}
项目概要：{request.project_summary or "（未提供）"}
其他需求上下文：{request.requirements_context or "（未提供）"}

请按要求输出 JSON。"""

        dify_res = await _call_dify_workflow(dify_key, {
            "raw_document": user_msg,
            "_system_override": _SCORING_SYSTEM_PROMPT,  # Dify 侧如支持则注入
        })
        outputs = dify_res.get("data", {}).get("outputs", {})

        # ── 优先：新 DSL Code 节点直接输出了拆分字段 ──
        if outputs.get("self_response"):
            self_response = outputs["self_response"]
            if self_response not in ("full", "partial"):
                self_response = "partial"
            evidence_raw = outputs.get("evidence_refs", "[]")
            if isinstance(evidence_raw, str):
                try:
                    evidence_list = _json.loads(evidence_raw)
                except Exception:
                    evidence_list = []
            elif isinstance(evidence_raw, list):
                evidence_list = evidence_raw
            else:
                evidence_list = []
            return FillScoringRowResponse(
                row_id=request.row_id,
                self_response=self_response,
                self_comment=outputs.get("self_comment", "我方具备相关能力，能够响应本评分项要求。"),
                evidence_refs=evidence_list[:3],
            )

        # ── 降级：旧 DSL 输出纯文本 JSON ──
        raw_text = (
            outputs.get("text") or outputs.get("result") or outputs.get("content") or ""
        )
        if isinstance(raw_text, list):
            raw_text = "\n".join(str(x) for x in raw_text)

        import re as _re

        # 尝试清洗 <think> 标签（推理模型常见）
        clean_text = _re.sub(r'<think>.*?</think>', '', raw_text, flags=_re.DOTALL).strip()

        # 尝试匹配 ```json ... ``` 包裹的结构
        match_md = _re.search(r'```json\s*(.*?)\s*```', clean_text, _re.DOTALL)
        if match_md:
            json_str = match_md.group(1)
        else:
            # 降级匹配第一个 { 到最后一个 }
            match_brace = _re.search(r'\{.*\}', clean_text, _re.DOTALL)
            json_str = match_brace.group(0) if match_brace else "{}"

        try:
            parsed = _json.loads(json_str) if json_str else {}
        except _json.JSONDecodeError as e:
            logger.error(f"JSON Decode Error in fill_scoring_row: {e}. Raw Text: {json_str}")
            parsed = {}

        self_response = parsed.get("self_response", "partial")
        if self_response not in ("full", "partial"):
            self_response = "partial"  # 强制不允许 none

        return FillScoringRowResponse(
            row_id=request.row_id,
            self_response=self_response,
            self_comment=parsed.get("self_comment", "我方具备相关能力，能够响应本评分项要求。"),
            evidence_refs=parsed.get("evidence_refs", []),
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to AI fill scoring row '{request.indicator}': {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/projects/export-scoring-table", summary="导出评分表为 Excel")
async def export_scoring_table(request: ExportScoringTableRequest):
    """
    将评分表导出为 Excel 文件（openpyxl）。
    返回文件流，前端直接触发下载。
    TODO: 后续 gateway-forge 阶段需整合为 Word 表格形式。
    """
    from io import BytesIO
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from fastapi.responses import StreamingResponse
    except ImportError:
        raise HTTPException(status_code=500, detail="服务器缺少 openpyxl，请先 pip install openpyxl")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "自评评分表"

    # 表头样式
    header_fill = PatternFill("solid", fgColor="1A6FA8")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    headers = ["评分指标", "最高分", "评分标准", "自评情况", "自评说明", "证明材料引用"]
    col_widths = [30, 10, 40, 12, 50, 45]
    for col_idx, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell = ws.cell(row=1, column=col_idx, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border
        ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = w
    ws.row_dimensions[1].height = 28

    # 自评情况中文映射
    response_label = {"full": "响应", "partial": "部分响应", "none": "不响应", "": "未填写"}

    for row_idx, row in enumerate(request.rows, 2):
        vals = [
            row.indicator,
            row.max_score,
            row.criteria,
            response_label.get(row.self_response, row.self_response),
            row.self_comment,
            "\n".join(row.evidence_refs),
        ]
        for col_idx, val in enumerate(vals, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = thin_border
        ws.row_dimensions[row_idx].height = 60

    # 合计行
    total_max = sum(r.max_score for r in request.rows)
    ws.append(["合计", total_max, "", "", "", ""])

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    from urllib.parse import quote as _quote_fn

    filename = f"{request.project_name}_自评评分表.xlsx"
    filename_star = _quote_fn(filename, safe="")
    cd = f'attachment; filename="scoring.xlsx"; filename*=UTF-8\'\'{filename_star}'
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": cd},
    )

# ─── 全局蓝图生成 ─────────────────────────────────────────────────────────
from .schemas import GenerateBlueprintRequest, GenerateBlueprintResponse, BlueprintData

@router.post("/projects/generate-blueprint", response_model=GenerateBlueprintResponse, summary="生成全局蓝图")
async def generate_blueprint(request: GenerateBlueprintRequest):
    """
    基于项目需求和大纲，生成全局投标策略蓝图。
    调用 ProEngine_Blueprint_Generator 工作流。
    """
    import json as _json
    try:
        dify_key = _get_workflow_key("blueprint_generator")
        if not dify_key:
            # Fallback 默认蓝图（开发兜底）
            return GenerateBlueprintResponse(blueprint=BlueprintData(
                positioning="展示深厚行业经验与技术领先性，打造高性价比方案",
                strategy="充分响应招标需求，在附加分项目上寻求突破，提供超出预期的售后保障",
                highlights=["自研核心技术的安全可靠性", "行业首创的快速交付模式", "总包一站式闭环服务"],
                writing_style="正式、专业、数据驱动"
            ))

        req_summary = "\\n".join([f"- [{r.get('type')}] {r.get('content')}" for r in request.requirements[:30]])
        outline_summary = "\\n".join([f"- {s.get('title')}" for s in request.outline])

        dify_res = await _call_dify_workflow(dify_key, {
            "bid_type": request.bid_type,
            "project_summary": request.project_summary,
            "requirements_summary": req_summary,
            "outline_summary": outline_summary
        })
        
        outputs = dify_res.get("data", {}).get("outputs", {})

        # ── 优先：新 DSL Code 节点直接输出了拆分字段 ──
        if outputs.get("positioning"):
            highlights_raw = outputs.get("highlights", "[]")
            if isinstance(highlights_raw, str):
                try:
                    highlights_list = _json.loads(highlights_raw)
                except Exception:
                    highlights_list = []
            elif isinstance(highlights_raw, list):
                highlights_list = highlights_raw
            else:
                highlights_list = []
            blueprint = BlueprintData(
                positioning=outputs.get("positioning", ""),
                strategy=outputs.get("strategy", ""),
                highlights=highlights_list,
                writing_style=outputs.get("writing_style", "正式、严谨庄重")
            )
            return GenerateBlueprintResponse(blueprint=blueprint)

        # ── 降级：旧 DSL 输出纯文本 JSON ──
        raw_text = outputs.get("text") or outputs.get("result") or outputs.get("content") or ""
        
        import re as _re
        match = _re.search(r'\{.*\}', raw_text, _re.DOTALL)
        parsed = _json.loads(match.group()) if match else {}
        
        # 兼容输出结构
        blueprint = BlueprintData(
            positioning=parsed.get("positioning", "展示高质量、高性价的专业方案"),
            strategy=parsed.get("strategy", "严格遵守所有要求，提供具有竞争力的优势方案"),
            highlights=parsed.get("highlights", ["优秀的行业业绩案例", "专业的技术服务团队", "完善的售后保障"]),
            writing_style=parsed.get("writing_style", "正式、严谨庄重")
        )
        return GenerateBlueprintResponse(blueprint=blueprint)
    except Exception as e:
        logger.error(f"Failed to generate blueprint: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ─── 最终文档组装 (Gateway-Forge) ──────────────────────────────────────────
import sys as _sys
from pydantic import BaseModel as _BaseModel

class _SectionItem(_BaseModel):
    id: str = ""
    title: str = ""
    content: str = ""

class _ForgeDocumentRequest(_BaseModel):
    project_id: str = ""
    project_name: str = "投标文件"
    sections: list[dict] = []        # [{id, title, content}]
    scoring_rows: list[dict] = []    # ScoringRow 列表
    attachments: list[dict] = []     # [{label, content}]
    mapping_table: dict = {}         # {占位符: 原文}（可选，主要为 BIDDER 信息占位符）
    bidder_info: dict = {}           # {orgName, legalRep, projectLead, phone, docDate}
    image_map: dict = {}             # {占位符: 本地路径}


def _sanitize_docx_slice_heading_semantics(docx_bytes: bytes) -> tuple[bytes, bool]:
    """
    清理切片内 Heading/TOC 语义，保留排版直观效果（字体、缩进、表格、图片）。
    返回: (处理后的 docx_bytes, 是否成功净化)
    """
    import io
    import re as _re
    import zipfile
    import xml.etree.ElementTree as ET

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    NS = {"w": W_NS}
    qn = lambda tag: f"{{{W_NS}}}{tag}"

    def _is_heading_like_style_id(style_id: str) -> bool:
        sid = (style_id or "").strip().lower()
        return bool(_re.match(r"heading[1-9]\d*$", sid) or sid.startswith("toc"))

    def _is_heading_like_style_name(name: str) -> bool:
        n = (name or "").strip().lower()
        return n.startswith("heading") or n.startswith("toc")

    in_buf = io.BytesIO(docx_bytes)
    out_buf = io.BytesIO()
    with zipfile.ZipFile(in_buf, "r") as zin:
        entries = {info.filename: zin.read(info.filename) for info in zin.infolist()}

    doc_xml = entries.get("word/document.xml")
    if not doc_xml:
        return docx_bytes, False

    heading_style_ids: set[str] = set()
    style_font_size_map: dict[str, str] = {}
    styles_xml = entries.get("word/styles.xml")
    if styles_xml:
        try:
            styles_root = ET.fromstring(styles_xml)
            for style in styles_root.findall(".//w:style", NS):
                if style.attrib.get(qn("type")) != "paragraph":
                    continue
                sid = style.attrib.get(qn("styleId"), "")
                # 记录样式层字号（半磅值字符串，如 24=12pt）
                sz_el = style.find("w:rPr/w:sz", NS)
                if sz_el is not None:
                    sz_val = sz_el.attrib.get(qn("val"), "")
                    if sz_val:
                        style_font_size_map[sid] = sz_val
                if _is_heading_like_style_id(sid):
                    heading_style_ids.add(sid)
                    continue
                name_el = style.find("w:name", NS)
                if name_el is not None and _is_heading_like_style_name(name_el.attrib.get(qn("val"), "")):
                    heading_style_ids.add(sid)
        except Exception:
            heading_style_ids = set()
            style_font_size_map = {}

    changed = False
    try:
        root = ET.fromstring(doc_xml)
        for ppr in root.findall(".//w:p/w:pPr", NS):
            p_style = ppr.find("w:pStyle", NS)
            if p_style is not None:
                style_id = p_style.attrib.get(qn("val"), "")
                sid_l = (style_id or "").strip().lower()
                if (
                    style_id in heading_style_ids
                    or _is_heading_like_style_id(style_id)
                    or sid_l.startswith("heading")
                    or sid_l.startswith("toc")
                ):
                    # 先保留原样式字号，再清理 heading/toc 语义
                    inherited_sz = style_font_size_map.get(style_id, "")
                    ppr.remove(p_style)
                    changed = True
                    if inherited_sz:
                        ppr_rpr = ppr.find("w:rPr", NS)
                        if ppr_rpr is None:
                            ppr_rpr = ET.SubElement(ppr, qn("rPr"))
                        # 字体统一到导出正文字体
                        rfonts = ppr_rpr.find("w:rFonts", NS)
                        if rfonts is None:
                            rfonts = ET.SubElement(ppr_rpr, qn("rFonts"))
                        rfonts.set(qn("ascii"), "Times New Roman")
                        rfonts.set(qn("hAnsi"), "Times New Roman")
                        rfonts.set(qn("eastAsia"), "宋体")
                        # 字号延用原样式
                        sz = ppr_rpr.find("w:sz", NS)
                        if sz is None:
                            sz = ET.SubElement(ppr_rpr, qn("sz"))
                        sz.set(qn("val"), inherited_sz)
                        sz_cs = ppr_rpr.find("w:szCs", NS)
                        if sz_cs is None:
                            sz_cs = ET.SubElement(ppr_rpr, qn("szCs"))
                        sz_cs.set(qn("val"), inherited_sz)
            outline_lvl = ppr.find("w:outlineLvl", NS)
            if outline_lvl is not None:
                ppr.remove(outline_lvl)
                changed = True
    except Exception:
        return docx_bytes, False

    if not changed:
        return docx_bytes, True

    entries["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, content in entries.items():
            zout.writestr(name, content)
    return out_buf.getvalue(), True


def _build_toc_entries(sections: list[dict]) -> list[dict]:
    entries: list[dict] = []
    for sec in sections or []:
        if not isinstance(sec, dict):
            continue
        lvl = int(sec.get("toc_level") or sec.get("heading_level") or 0)
        if lvl < 1 or lvl > 3:
            continue
        number = str(sec.get("heading_number") or "").strip()
        raw_title = str(sec.get("heading_text") or sec.get("title") or "").strip()
        title = f"{number} {raw_title}".strip() if number else raw_title
        if not title:
            continue
        entries.append({"level": lvl, "text": title})
    return entries


def _doc_has_unexpected_heading_semantics(doc, allowed_headings: set[str]) -> bool:
    def _norm(s: str) -> str:
        return re.sub(r"\s+", "", str(s or "")).strip().lower()

    for p in doc.paragraphs:
        try:
            style_name = (p.style.name if p.style else "") or ""
        except Exception:
            style_name = ""
        txt = (p.text or "").strip()
        if not txt:
            continue
        style_l = style_name.lower()
        if style_l.startswith("toc"):
            return True
        if style_l.startswith("heading"):
            if _norm(txt) not in allowed_headings and txt != "目录":
                return True
    return False


def _insert_toc_page(doc, toc_entries: list[dict], use_native_toc: bool) -> None:
    if not toc_entries:
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    anchor = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    title_p = anchor.insert_paragraph_before("目录")
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_p.runs:
        title_p.runs[0].bold = True

    if use_native_toc:
        field_p = anchor.insert_paragraph_before("")
        fld = OxmlElement("w:fldSimple")
        fld.set(_qn("w:instr"), ' TOC \\o "1-3" \\h \\z \\u ')
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = "（在 Word 中右键目录并选择“更新域”）"
        r.append(t)
        fld.append(r)
        field_p._p.append(fld)
    else:
        for row in toc_entries:
            indent = "    " * max(0, int(row.get("level", 1)) - 1)
            anchor.insert_paragraph_before(f"{indent}{row.get('text', '')}")

    split_p = anchor.insert_paragraph_before("")
    split_p.add_run().add_break(WD_BREAK.PAGE)


def _rebind_heading_numbering_for_export(doc) -> None:
    """
    为导出文档中的 Heading1/2/3 统一绑定多级编号语义，避免拼装后退化为静态文本编号。
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as _qn

    try:
        numbering = doc.part.numbering_part.numbering_definitions._numbering
    except Exception:
        return

    abstract_ids = [
        int(n.get(_qn("w:abstractNumId")))
        for n in numbering.findall(_qn("w:abstractNum"))
        if n.get(_qn("w:abstractNumId"))
    ]
    num_ids = [
        int(n.get(_qn("w:numId")))
        for n in numbering.findall(_qn("w:num"))
        if n.get(_qn("w:numId"))
    ]
    next_abstract_id = (max(abstract_ids) + 1) if abstract_ids else 2000
    next_num_id = (max(num_ids) + 1) if num_ids else 2000

    abstract = OxmlElement("w:abstractNum")
    abstract.set(_qn("w:abstractNumId"), str(next_abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(_qn("w:val"), "multilevel")
    abstract.append(multi)

    for ilvl, fmt, text in (
        (0, "chineseCounting", "%1、"),
        (1, "decimal", "%1.%2"),
        (2, "decimal", "%1.%2.%3"),
    ):
        lvl = OxmlElement("w:lvl")
        lvl.set(_qn("w:ilvl"), str(ilvl))
        start = OxmlElement("w:start")
        start.set(_qn("w:val"), "1")
        lvl.append(start)
        nfmt = OxmlElement("w:numFmt")
        nfmt.set(_qn("w:val"), fmt)
        lvl.append(nfmt)
        if ilvl >= 1:
            # 让 H2/H3 引用上级时统一按阿拉伯数字显示。
            is_lgl = OxmlElement("w:isLgl")
            lvl.append(is_lgl)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(_qn("w:val"), text)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(_qn("w:val"), "left")
        lvl.append(lvl_jc)
        abstract.append(lvl)

    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(_qn("w:numId"), str(next_num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(_qn("w:val"), str(next_abstract_id))
    num.append(abs_ref)
    numbering.append(num)

    def _heading_level_from_style(style_name: str) -> int:
        s = (style_name or "").strip().lower()
        if s.startswith("heading 1") or s.startswith("heading1") or s.startswith("标题1") or s.startswith("标题 1"):
            return 1
        if s.startswith("heading 2") or s.startswith("heading2") or s.startswith("标题2") or s.startswith("标题 2"):
            return 2
        if s.startswith("heading 3") or s.startswith("heading3") or s.startswith("标题3") or s.startswith("标题 3"):
            return 3
        return 0

    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        if not txt or txt == "目录":
            continue
        try:
            style_name = (p.style.name if p.style else "") or ""
        except Exception:
            style_name = ""
        level = _heading_level_from_style(style_name)
        if level < 1:
            continue
        ppr = p._p.get_or_add_pPr()
        old = ppr.find(_qn("w:numPr"))
        if old is not None:
            ppr.remove(old)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(_qn("w:val"), str(level - 1))
        num_id = OxmlElement("w:numId")
        num_id.set(_qn("w:val"), str(next_num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id)
        ppr.append(num_pr)


def _apply_toc_for_export(
    docx_bytes: bytes,
    sections: list[dict],
    *,
    prefer_native: bool,
    heading_sanitized: bool,
) -> bytes:
    from io import BytesIO
    from docx import Document

    def _strip_heading_prefix(text: str) -> str:
        raw = str(text or "").strip()
        if not raw:
            return ""
        # 兼容中文序号与多级阿拉伯序号前缀：一、 / 1.1 / 1.1.1
        return re.sub(r"^(([一二三四五六七八九十百千万]+、)|(\d+(?:\.\d+){1,2}))\s*", "", raw).strip()

    toc_entries = _build_toc_entries(sections)
    if not toc_entries:
        return docx_bytes

    doc = Document(BytesIO(docx_bytes))
    allowed: set[str] = set()
    for x in toc_entries:
        txt = str(x.get("text") or "").strip()
        if txt:
            allowed.add(re.sub(r"\s+", "", txt).strip().lower())
        stripped = _strip_heading_prefix(txt)
        if stripped:
            allowed.add(re.sub(r"\s+", "", stripped).strip().lower())
    use_native = bool(prefer_native and heading_sanitized and not _doc_has_unexpected_heading_semantics(doc, allowed))
    _insert_toc_page(doc, toc_entries, use_native_toc=use_native)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


def _rebind_heading_numbering_bytes(docx_bytes: bytes) -> bytes:
    from io import BytesIO
    from docx import Document

    try:
        doc = Document(BytesIO(docx_bytes))
        _rebind_heading_numbering_for_export(doc)
        out = BytesIO()
        doc.save(out)
        out.seek(0)
        return out.read()
    except Exception:
        return docx_bytes


def _resolve_project_source_docx_bytes(project_id: str, cache_entry: dict) -> bytes:
    docx_path = PRO_ENGINE_ROOT / "data" / "docx_cache" / f"{project_id}.docx"
    if docx_path.exists():
        return docx_path.read_bytes()

    doc_obj = cache_entry.get("doc")
    if doc_obj is not None:
        import io
        buf = io.BytesIO()
        doc_obj.save(buf)
        return buf.getvalue()
    return b""


def _build_docx_slice_segment(section: dict, project_id: str) -> tuple[bytes, bool]:
    start_block_id = str(section.get("start_block_id") or "").strip()
    end_block_id = str(section.get("end_block_id") or "").strip()
    if not start_block_id or not end_block_id:
        raise HTTPException(status_code=400, detail="docx_slice 段缺少 start_block_id/end_block_id")

    cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        _restore_locator_cache_from_disk(project_id)
        cache_entry = _locator_cache.get(project_id)
    if not cache_entry:
        raise HTTPException(status_code=404, detail=f"项目 [{project_id}] 的 DOCX 缓存不存在，无法按原文切片导出")

    blocks = cache_entry.get("doc_blocks", [])
    block_map = {b.get("block_id"): b for b in blocks if isinstance(b, dict)}
    start_block = block_map.get(start_block_id)
    end_block = block_map.get(end_block_id)
    if not start_block:
        raise HTTPException(status_code=404, detail=f"block_id {start_block_id} 未找到")
    if not end_block:
        raise HTTPException(status_code=404, detail=f"block_id {end_block_id} 未找到")

    start_idx = int(start_block.get("body_idx", 0))
    end_idx = int(end_block.get("body_idx", 0))
    if start_idx > end_idx:
        start_idx, end_idx = end_idx, start_idx

    source_bytes = _resolve_project_source_docx_bytes(project_id, cache_entry)
    if not source_bytes:
        raise HTTPException(status_code=409, detail="原始 DOCX 不可用，无法生成保格式切片")
    sliced = _slice_docx_bytes_by_body_range(source_bytes, start_idx, end_idx)
    sanitized, ok = _sanitize_docx_slice_heading_semantics(sliced)
    return sanitized, ok


def _build_hybrid_forge_docx(
    request: _ForgeDocumentRequest,
    full_mapping: dict,
) -> tuple[bytes, bool]:
    from io import BytesIO
    from docx import Document
    from docx.enum.text import WD_BREAK
    from src.forge import DocumentForge, _add_scoring_table, _add_attachments  # type: ignore
    try:
        from docxcompose.composer import Composer
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"缺少 docxcompose 依赖，无法执行保格式拼装: {e}",
        )

    project_id = (request.project_id or "").strip()
    if not project_id:
        raise HTTPException(status_code=400, detail="混合导出需要 project_id")

    # 每个段附带一个分页标志：True=拼接前插入分页符，False=连续拼接
    docx_segments: list[tuple[bytes, bool]] = []
    heading_sanitized = True
    forge = DocumentForge(
        mapping_table=full_mapping,
        bidder_info=request.bidder_info,
        image_map=request.image_map,
    )

    def _prepend_page_break(doc_obj: Document) -> Document:
        # 确保“每个模块独立起页”：在被拼接文档开头插入分页符
        if doc_obj.paragraphs:
            p = doc_obj.paragraphs[0].insert_paragraph_before()
        else:
            p = doc_obj.add_paragraph()
        p.add_run().add_break(WD_BREAK.PAGE)
        return doc_obj

    for sec in request.sections:
        if not isinstance(sec, dict):
            continue
        source_type = str(sec.get("source_type") or "markdown").strip().lower()
        if source_type == "docx_slice":
            if sec.get("inject_title") and str(sec.get("title") or "").strip():
                heading_seg = forge.build(
                    sections=[{
                        "id": sec.get("id", ""),
                        "title": sec.get("title", ""),
                        "heading_number": sec.get("heading_number", ""),
                        "heading_text": sec.get("heading_text", ""),
                        "bookmark_id": sec.get("bookmark_id", ""),
                        "content": "",
                        "heading_level": sec.get("heading_level", 1),
                        "title_only": True,
                    }],
                    scoring_rows=[],
                    attachments=[],
                )
                docx_segments.append((heading_seg, True))
            sliced, sanitized_ok = _build_docx_slice_segment(sec, project_id)
            heading_sanitized = heading_sanitized and bool(sanitized_ok)
            # 与 inject_title 属于同一模块，禁止在 H1 与正文切片之间额外分页
            docx_segments.append((sliced, False))
            continue
        md_seg = forge.build(
            sections=[{
                "id": sec.get("id", ""),
                "title": sec.get("title", ""),
                "heading_number": sec.get("heading_number", ""),
                "heading_text": sec.get("heading_text", ""),
                "bookmark_id": sec.get("bookmark_id", ""),
                "content": sec.get("content", ""),
                "heading_level": sec.get("heading_level", 1),
                "title_only": sec.get("title_only", False),
            }],
            scoring_rows=[],
            attachments=[],
        )
        docx_segments.append((md_seg, True))

    if not docx_segments:
        return forge.build(
            sections=request.sections,
            scoring_rows=request.scoring_rows,
            attachments=request.attachments,
        ), heading_sanitized

    master = Document(BytesIO(docx_segments[0][0]))
    composer = Composer(master)
    for seg, with_page_break in docx_segments[1:]:
        doc_obj = Document(BytesIO(seg))
        composer.append(_prepend_page_break(doc_obj) if with_page_break else doc_obj)

    merged_buf = BytesIO()
    composer.save(merged_buf)
    merged_buf.seek(0)

    final_doc = Document(merged_buf)
    if request.scoring_rows:
        _add_scoring_table(final_doc, request.scoring_rows)
    if request.attachments:
        _add_attachments(final_doc, request.attachments)
    _rebind_heading_numbering_for_export(final_doc)

    out = BytesIO()
    final_doc.save(out)
    out.seek(0)
    return out.read(), heading_sanitized


@router.post("/projects/forge-document", summary="组装生成最终标书 .docx")
async def forge_document(request: _ForgeDocumentRequest, db: Session = Depends(get_db)):
    """
    将已生成的所有元素整合为最终 .docx 文件：
    - 内部应用 EntityRegistry 全局还原 PIPT 占位符
    - request.mapping_table 用于还原 BIDDER 卵中的占位符
    - 各章节 Markdown 拼接转 Word
    - 自评评分表嵌入，附件追加
    """
    try:
        import re as _re
        from app.api_lite.database import EntityRegistry, FernetEncryptor

        # 步骤1：从 EntityRegistry 查询全局 PIPT 占位符映射
        all_content = " ".join(
            [s.get("content", "") for s in request.sections]
            + [a.get("content", "") for a in request.attachments]
        )
        pipt_placeholders = list(set(_re.findall(r'\{\{__PIPT_[a-z_]+_\d+__\}\}', all_content)))
        pipt_mapping: dict[str, str] = {}
        if pipt_placeholders:
            enc = FernetEncryptor.get()
            rows = db.query(EntityRegistry).filter(
                EntityRegistry.placeholder.in_(pipt_placeholders)
            ).all()
            pipt_mapping = {row.placeholder: enc.decrypt(row.original_text_enc) for row in rows}
            logger.info(f"forge-document: 查询到 {len(pipt_mapping)}/{len(pipt_placeholders)} 个 PIPT 占位符映射")

        # 步骤2：合并 PIPT 映射 + 前端传入的 BIDDER 映射
        full_mapping = {**pipt_mapping, **request.mapping_table}

        # 步骤2.1：从 ImageRegistry 查询涉及的图片并填充 request.image_map
        img_placeholders = list(set(_re.findall(r'__PRO_IMG_[a-f0-9]+__', all_content)))
        dynamic_image_map = {}
        if img_placeholders:
            from app.api_lite.database import ImageRegistry
            img_hashes = [ph.replace("__PRO_IMG_", "").replace("__", "") for ph in img_placeholders]
            img_rows = db.query(ImageRegistry).filter(
                ImageRegistry.image_hash.in_(img_hashes)
            ).all()
            for row in img_rows:
                dynamic_image_map[row.placeholder] = {"abs_path": row.abs_path, "preview_url": row.preview_url}
            
            logger.info(f"forge-document: 查询到 {len(dynamic_image_map)}/{len(img_placeholders)} 个物理图片实体")
            
        request.image_map = {**(request.image_map or {}), **dynamic_image_map}

        # 步骤3：生成 Word 文档
        _gateway_out_path = str(PRO_ENGINE_ROOT / "gateway-out")
        if _gateway_out_path not in _sys.path:
            _sys.path.insert(0, _gateway_out_path)

        from src.forge import DocumentForge  # type: ignore

        has_docx_slice = any(
            isinstance(sec, dict) and str(sec.get("source_type") or "").strip().lower() == "docx_slice"
            for sec in request.sections
        )
        heading_sanitized = True
        if has_docx_slice:
            docx_bytes, heading_sanitized = _build_hybrid_forge_docx(request, full_mapping)
        else:
            forge = DocumentForge(
                mapping_table=full_mapping,
                bidder_info=request.bidder_info,
                image_map=request.image_map,
            )
            docx_bytes = forge.build(
                sections=request.sections,
                scoring_rows=request.scoring_rows,
                attachments=request.attachments,
            )
        docx_bytes = _rebind_heading_numbering_bytes(docx_bytes)
        docx_bytes = _apply_toc_for_export(
            docx_bytes,
            request.sections,
            prefer_native=True,
            heading_sanitized=bool(heading_sanitized),
        )

        from fastapi.responses import Response
        from urllib.parse import quote

        safe_name = request.project_name.replace("/", "_").replace("\\", "_")
        filename = f"{safe_name}_标书文件.docx"
        # Starlette 要求 header 值为 latin-1；RFC 5987 的 filename* 值须为百分号编码的 UTF-8
        filename_star = quote(filename, safe="")
        content_disposition = (
            'attachment; filename="document.docx"; '
            f"filename*=UTF-8''{filename_star}"
        )
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": content_disposition},
        )
    except Exception as e:
        logger.error(f"forge-document 失败: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=str(e))


# ─── 知识库看板 (Knowledge Hub) ───────────────────────────────────────────
from .schemas import KnowledgeListResponse, KnowledgeSyncResponse, KnowledgeDocument
from datetime import datetime as _dt
import asyncio as _asyncio
from .task_manager import task_manager

_KB_SYNC_PROJECT = "__knowledge_sync__"
_kb_sync_guard = _asyncio.Lock()
_kb_sync_dedupe: dict[str, str] = {}


def _kb_sync_max_running() -> int:
    limits = task_manager.get_limits()
    return max(1, int(limits.get("max_kb_sync_running", 1)))


async def _acquire_kb_sync_dedupe_key(dedupe_key: str) -> Optional[str]:
    """同一同步 key 仅允许运行一个任务，若已有运行中任务则返回其 task_id。"""
    async with _kb_sync_guard:
        existing_task_id = _kb_sync_dedupe.get(dedupe_key)
        if not existing_task_id:
            return None
        existing = task_manager.get_task(existing_task_id)
        if existing and existing.status == "running":
            return existing_task_id
        _kb_sync_dedupe.pop(dedupe_key, None)
        return None


async def _set_kb_sync_dedupe_key(dedupe_key: str, task_id: str) -> None:
    async with _kb_sync_guard:
        _kb_sync_dedupe[dedupe_key] = task_id


async def _release_kb_sync_dedupe_key(dedupe_key: str, task_id: str) -> None:
    async with _kb_sync_guard:
        if _kb_sync_dedupe.get(dedupe_key) == task_id:
            _kb_sync_dedupe.pop(dedupe_key, None)


def _upsert_kb_sync_status_file(status_path: Path, patch: dict) -> None:
    """尽力更新 kb 同步状态文件，不阻断主流程。"""
    try:
        base = {}
        if status_path.exists():
            with open(status_path, "r", encoding="utf-8") as f:
                loaded = json.load(f)
                if isinstance(loaded, dict):
                    base = loaded
        base.update(patch or {})
        with open(status_path, "w", encoding="utf-8") as f:
            json.dump(base, f, ensure_ascii=False)
    except Exception:
        logger.warning("更新 kb sync 状态文件失败: %s", status_path, exc_info=True)


async def _run_kb_sync_subprocess(
    task_id: str,
    dedupe_key: str,
    cmd: list[str],
    *,
    env: Optional[dict[str, str]] = None,
    status_path: Optional[Path] = None,
) -> None:
    proc: Optional[_asyncio.subprocess.Process] = None
    try:
        task_manager.update_stage(task_id, "知识库同步启动中")
        if status_path:
            _upsert_kb_sync_status_file(
                status_path,
                {"job_id": task_id, "status": "starting", "task_id": task_id},
            )
        proc = await _asyncio.create_subprocess_exec(
            *cmd,
            cwd=str(PRO_ENGINE_ROOT),
            env=env,
            stdout=_asyncio.subprocess.PIPE,
            stderr=_asyncio.subprocess.PIPE,
        )
        task_manager.update_stage(task_id, "知识库同步执行中")
        if status_path:
            _upsert_kb_sync_status_file(
                status_path,
                {"job_id": task_id, "status": "running", "task_id": task_id, "pid": proc.pid},
            )
        stdout, stderr = await proc.communicate()
        if proc.returncode == 0:
            task_manager.set_result(task_id, {
                "return_code": 0,
                "stdout": (stdout or b"").decode("utf-8", errors="ignore")[-4000:],
            })
            if status_path and status_path.exists():
                _upsert_kb_sync_status_file(
                    status_path,
                    {"job_id": task_id, "status": "completed", "task_id": task_id},
                )
            return
        err = (stderr or b"").decode("utf-8", errors="ignore")[-4000:]
        out = (stdout or b"").decode("utf-8", errors="ignore")[-2000:]
        task_manager.set_error(task_id, f"sync_kb.py failed rc={proc.returncode}: {err or out or 'unknown error'}")
        if status_path:
            _upsert_kb_sync_status_file(
                status_path,
                {"job_id": task_id, "status": "failed", "task_id": task_id, "error": err or out or "unknown error"},
            )
    except _asyncio.CancelledError:
        if proc and proc.returncode is None:
            proc.terminate()
            try:
                await _asyncio.wait_for(proc.wait(), timeout=3)
            except Exception:
                proc.kill()
        task_manager.set_cancelled(task_id)
        if status_path:
            _upsert_kb_sync_status_file(
                status_path,
                {"job_id": task_id, "status": "cancelled", "task_id": task_id},
            )
        raise
    except Exception as e:
        logger.error("知识库同步任务异常: task_id=%s, err=%s", task_id, e, exc_info=True)
        task_manager.set_error(task_id, str(e))
        if status_path:
            _upsert_kb_sync_status_file(
                status_path,
                {"job_id": task_id, "status": "failed", "task_id": task_id, "error": str(e)},
            )
    finally:
        await _release_kb_sync_dedupe_key(dedupe_key, task_id)

@router.get("/knowledge/documents", response_model=KnowledgeListResponse, summary="获取远端 Dify 数据集文档状态")
async def get_knowledge_documents():
    """透传请求前往 Dify 查询指定 Dataset 中已解析出的 Documents 列表"""
    import os
    dify_url = os.getenv("DIFY_API_URL", "http://localhost/v1")
    dataset_id = os.getenv("DIFY_DATASET_ID", "")
    dataset_key = os.getenv("DIFY_DATASET_KEY", "")

    if not dataset_id or not dataset_key:
        return KnowledgeListResponse(
            dataset_info={"error": "DIFY_DATASET_ID or KEY not configured in backend."},
            documents=[]
        )

    headers = {"Authorization": f"Bearer {dataset_key}"}
    url = f"{dify_url}/datasets/{dataset_id}/documents"
    
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.get(url, headers=headers)
            resp.raise_for_status()
            data = resp.json()

        docs_raw = data.get("data", [])
        documents = []
        for d in docs_raw:
            # 兼容时间戳转换和 Dify 状态值：completed, indexing, error 等
            ts = d.get("created_at")
            upload_time = _dt.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M") if ts else "-"
            
            raw_status = d.get("indexing_status", "completed")
            if raw_status == "completed":
                status = "success"
            elif raw_status == "error":
                status = "failed"
            else:
                status = "indexing"
                
            # 补充估算体积转换
            size_bytes = d.get("word_count", 0) * 2  # Dify没有直接文件大小，用字数*2估个体积展示就好
            if size_bytes > 1024 * 1024:
                size_str = f"{size_bytes / (1024 * 1024):.1f} MB"
            else:
                size_str = f"{size_bytes / 1024:.1f} KB"
                
            documents.append(KnowledgeDocument(
                id=d.get("id", ""),
                name=d.get("name", ""),
                size=size_str,
                uploadTime=upload_time,
                status=status,
                chunks=d.get("tokens", d.get("word_count", 0)) // 500  # 大约按 500 tokens/chunk 算出大致数字供前端显示
            ))
            
        return KnowledgeListResponse(
            dataset_info={"status": "connected", "dataset_id": dataset_id},
            documents=documents
        )
    except Exception as e:
        logger.error(f"Failed to fetch dify kb documents: {e}")
        return KnowledgeListResponse(
            dataset_info={"status": "error", "message": str(e)},
            documents=[]
        )

@router.post("/knowledge/sync", response_model=KnowledgeSyncResponse, summary="触发知识库同步")
async def trigger_knowledge_sync():
    """触发受限并发的知识库同步后台任务，返回统一 task_id。"""
    script_path = PRO_ENGINE_ROOT / "scripts" / "sync_kb.py"
    if not script_path.exists():
        raise HTTPException(status_code=404, detail="sync_kb.py script not found.")
    try:
        task_manager.ensure_backend_ready()
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail={"code": "TASK_BACKEND_UNAVAILABLE", "message": str(e)})

    dedupe_key = "all"
    existing_task_id = await _acquire_kb_sync_dedupe_key(dedupe_key)
    if existing_task_id:
        return KnowledgeSyncResponse(
            message="同类知识库同步任务正在执行，已复用现有任务",
            status="running",
            task_id=existing_task_id,
        )

    allowed, details = await task_manager.try_acquire_task_slot(
        project_id=_KB_SYNC_PROJECT,
        task_type="knowledge_sync",
        enforce_project_limit=False,
        max_type_running=_kb_sync_max_running(),
    )
    if not allowed:
        raise HTTPException(
            status_code=429,
            detail={
                "code": "TASK_LIMIT_REACHED",
                "message": "知识库同步并发达到上限，请稍后重试",
                "limits": task_manager.get_limits(),
                "metrics": details,
            },
        )

    task_id: Optional[str] = None
    try:
        import sys
        task_id = task_manager.create_task("knowledge_sync", _KB_SYNC_PROJECT)
        await _set_kb_sync_dedupe_key(dedupe_key, task_id)
        bg = _asyncio.create_task(_run_kb_sync_subprocess(task_id, dedupe_key, [sys.executable, str(script_path)]))
        task_manager.set_async_task(task_id, bg)
        return KnowledgeSyncResponse(
            message="后台同步任务已启动，请通过 task_id 轮询状态",
            status="running",
            task_id=task_id,
        )
    except Exception as e:
        if task_id:
            await _release_kb_sync_dedupe_key(dedupe_key, task_id)
        logger.error(f"Failed to trigger sync: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/knowledge/sync/{doc_name}", response_model=KnowledgeSyncResponse, summary="触发单文件知识库同步")
async def trigger_knowledge_sync_single(doc_name: str):
    """触发受限并发的单文件知识库同步后台任务，返回统一 task_id。"""
    script_path = PRO_ENGINE_ROOT / "scripts" / "sync_kb.py"
    if not script_path.exists():
        raise HTTPException(status_code=404, detail="sync_kb.py script not found.")
    try:
        task_manager.ensure_backend_ready()
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail={"code": "TASK_BACKEND_UNAVAILABLE", "message": str(e)})

    local_prefix = doc_name
    if local_prefix.endswith(".txt"):
        local_prefix = local_prefix[:-4]
    local_prefix = (local_prefix or "").strip()
    if not local_prefix:
        raise HTTPException(status_code=400, detail="无效的 doc_name")

    dedupe_key = f"single:{local_prefix.lower()}"
    existing_task_id = await _acquire_kb_sync_dedupe_key(dedupe_key)
    if existing_task_id:
        return KnowledgeSyncResponse(
            message=f"文件 {doc_name} 已有同步任务在执行，已复用现有任务",
            status="running",
            task_id=existing_task_id,
        )

    allowed, details = await task_manager.try_acquire_task_slot(
        project_id=_KB_SYNC_PROJECT,
        task_type="knowledge_sync",
        enforce_project_limit=False,
        max_type_running=_kb_sync_max_running(),
    )
    if not allowed:
        raise HTTPException(
            status_code=429,
            detail={
                "code": "TASK_LIMIT_REACHED",
                "message": "知识库同步并发达到上限，请稍后重试",
                "limits": task_manager.get_limits(),
                "metrics": details,
            },
        )

    task_id: Optional[str] = None
    try:
        import sys
        task_id = task_manager.create_task("knowledge_sync", _KB_SYNC_PROJECT)
        await _set_kb_sync_dedupe_key(dedupe_key, task_id)
        bg = _asyncio.create_task(
            _run_kb_sync_subprocess(task_id, dedupe_key, [sys.executable, str(script_path), local_prefix])
        )
        task_manager.set_async_task(task_id, bg)
        return KnowledgeSyncResponse(
            message=f"后台单文件同步任务已启动: {doc_name}",
            status="running",
            task_id=task_id,
        )
    except Exception as e:
        if task_id:
            await _release_kb_sync_dedupe_key(dedupe_key, task_id)
        logger.error(f"Failed to trigger single sync: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ─── P6: 解析报告 SSE 生成端点 ──────────────────────────────────

def _collect_leaf_nodes(nodes: list, parent_label: str = "") -> list[dict]:
    """递归收集框架中的所有叶子节点（含 extractionPrompt 的节点）"""
    leaves = []
    for node in nodes:
        full_label = f"{parent_label} > {node['label']}" if parent_label else node['label']
        children = node.get("children", [])
        if children:
            leaves.extend(_collect_leaf_nodes(children, full_label))
        elif node.get("extractionPrompt"):
            leaves.append({
                "id": node["id"],
                "label": node["label"],
                "full_label": full_label,
                "extractionPrompt": node["extractionPrompt"],
            })
    return leaves


async def _setup_knowledge_base(doc_text: str, project_id: str) -> tuple[str, str] | None:
    """
    上传脱敏文档到固定知识库 bid_base，等待索引完成。
    上传前会清理同 project_id 的历史文档，防止残留堆积。

    Returns:
        tuple[dataset_id, document_id] 成功时，None 失败则回退全文模式
    """
    dataset_api_key = os.environ.get("DIFY_DATASET_API_KEY", "").strip()
    bid_base_id = os.environ.get("DIFY_BID_BASE_ID", "").strip()
    if not dataset_api_key or not bid_base_id:
        return None

    try:
        import sys
        sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent / "dify-bridge"))
        from src.config import DifyConfig
        from src.knowledge import KnowledgeManager

        dify_base = os.environ.get("DIFY_API_URL", "http://localhost/v1").rstrip("/")
        config = DifyConfig(
            base_url=dify_base,
            api_key=dataset_api_key,
            timeout=300,
        )
        km = KnowledgeManager(config)

        # ── 清理同 project 的旧文档，防止残留堆积 ──
        doc_prefix = f"bid_doc_{project_id}" if project_id else "bid_doc_temp"
        try:
            existing_docs = await km.list_documents(bid_base_id)
            for doc in existing_docs:
                if doc.get("name", "").startswith(doc_prefix):
                    try:
                        await km.delete_document(doc["id"], bid_base_id)
                        logger.info(f"清理旧文档: {doc['name']} (id={doc['id']})")
                    except Exception as del_err:
                        logger.warning(f"清理旧文档失败（可忽略）: {del_err}")
        except Exception as list_err:
            logger.warning(f"列出文档失败（继续上传）: {list_err}")

        # 文档名带 project_id 标记，便于识别和清理
        doc_name = f"{doc_prefix}.txt"

        # 上传脱敏文本到固定 bid_base 知识库
        result = await km.upload_text(
            name=doc_name,
            text=doc_text,
            dataset_id=bid_base_id,
        )
        doc_id = result.get("document", {}).get("id", "")

        # 等待索引完成（最多 5 分钟）
        if doc_id:
            indexed = await km.wait_for_indexing(doc_id, bid_base_id, timeout=300, interval=3)
            if indexed:
                logger.info(f"文档索引就绪: bid_base={bid_base_id}, doc_id={doc_id}, project={project_id}")
                return (bid_base_id, doc_id)
            else:
                logger.warning(f"文档索引超时，回退全文模式")
                try:
                    await km.delete_document(doc_id, bid_base_id)
                except Exception:
                    pass
                return None

        return None
    except Exception as e:
        logger.warning(f"知识库上传失败，回退全文模式: {e}")
        return None


async def _cleanup_knowledge_base(dataset_id: str, document_id: str):
    """清理：只删文档，不删知识库"""
    try:
        dataset_api_key = os.environ.get("DIFY_DATASET_API_KEY", "").strip()
        if not dataset_api_key:
            return

        import sys
        sys.path.insert(0, str(Path(__file__).parent.parent.parent.parent / "dify-bridge"))
        from src.config import DifyConfig
        from src.knowledge import KnowledgeManager

        dify_base = os.environ.get("DIFY_API_URL", "http://localhost/v1").rstrip("/")
        config = DifyConfig(base_url=dify_base, api_key=dataset_api_key, timeout=60)
        km = KnowledgeManager(config)
        await km.delete_document(document_id, dataset_id)
        logger.info(f"文档已清理: doc_id={document_id}")
    except Exception as e:
        logger.warning(f"文档清理失败（可忽略）: {e}")


@router.post("/projects/analyze", summary="解析报告 SSE 生成（分组并行全文模式）")
async def analyze_document(
    raw_document: str = Form(default=""),
    project_id: str = Form(default=""),
    selected_node_ids: str = Form(default=""),  # 逗号分隔的节点 ID，为空则提取全部
):
    """
    接收脱敏后的文档文本，分组并行调用 Dify 工作流提取，以 SSE 推送进度。

    支持全量提取（默认）和批量选择提取（传 selected_node_ids）。
    当 selected_node_ids 非空时，只提取指定的叶子节点，合并为一组调用 Dify。

    SSE 事件类型：
    - progress:      阶段进度 { phase, message }
    - node_complete: 节点完成 { node_id, label, content }
    - error:         分组错误 { group, error }
    - complete:      全部完成 { total_nodes, success_count }
    """
    import json as _json
    import asyncio as _asyncio
    import re as _re

    # 读取框架配置
    config_path = Path(__file__).parent.parent.parent / "config" / "analysis_framework.json"
    if not config_path.exists():
        raise HTTPException(status_code=404, detail="analysis_framework.json 不存在")
    system_prompt_base, all_nodes = load_docanalysis_framework(config_path)

    if not all_nodes:
        raise HTTPException(status_code=400, detail="框架中无节点")

    dify_key = _get_workflow_key("doc_analysis") or _get_workflow_key("requirement_extractor")
    if not dify_key:
        raise HTTPException(status_code=500, detail="需求提取工作流 API Key 未配置")

    # 解析批量选择的节点 ID
    selected_ids = set(
        nid.strip() for nid in selected_node_ids.split(",") if nid.strip()
    ) if selected_node_ids.strip() else None

    doc_source = (raw_document or "").strip() or _load_raw_document(project_id)
    if not doc_source:
        raise HTTPException(status_code=404, detail="未找到项目原文缓存，请先重新上传并解析文档")
    doc_text = doc_source[:300000]  # 300k 字符，覆盖绝大部分招标文件

    groups = build_docanalysis_groups(all_nodes, selected_ids)
    if not groups:
        raise HTTPException(status_code=400, detail="未找到可提取节点")

    total_nodes = sum(len(g["nodes"]) for g in groups)

    # ── 单组异步提取函数 ──
    async def extract_group(group: dict) -> list[dict]:
        """
        构建多任务 prompt，调 Dify 一次，返回 [{"node_id", "label", "content"}]
        """
        nodes = group["nodes"]
        group_label = group["group_label"]

        combined_system = build_docanalysis_system_prompt(system_prompt_base, nodes, group_label)
        raw_text = ""

        try:
            dify_res = await _call_dify_workflow(dify_key, {
                "system_prompt": combined_system,
                "raw_document": doc_text,
                "node_label": group_label,
            })

            outputs = dify_res.get("data", {}).get("outputs", {})
            raw_text = extract_docanalysis_text_output(outputs)
            raw_text, _attachments_payload = split_bid_attachments_tag(raw_text)
            result_map = parse_docanalysis_result_map(raw_text)

            results = []
            for n in nodes:
                content = extract_docanalysis_node_content(result_map, n["id"])
                if isinstance(content, (dict, list)):
                    content = _json.dumps(content, ensure_ascii=False, indent=2)
                results.append({"node_id": n["id"], "label": n["label"], "content": str(content)})
            return results

        except Exception as je:
            logger.warning(f"分组 [{group_label}] 结果解析失败，降级逐节点提取: {je}")
            if raw_text:
                return [{"node_id": nodes[0]["id"], "label": nodes[0]["label"], "content": raw_text}]
            return [{"node_id": n["id"], "label": n["label"], "content": "**提取失败，请重新生成**"} for n in nodes]

    # ── SSE Generator（并行执行所有组）──
    async def sse_generator():
        import asyncio as _asyncio
        queue: _asyncio.Queue = _asyncio.Queue()
        success_count = 0

        yield f"event: progress\ndata: {_json.dumps({'phase': 'analyzing', 'message': f'并行解析 {len(groups)} 组 / 共 {total_nodes} 个节点', 'total': total_nodes}, ensure_ascii=False)}\n\n"

        # 包装：每组完成后往 queue 里推结果
        async def run_group(group: dict, idx: int):
            yield_data = await extract_group(group)
            await queue.put(("group_done", idx, group["group_label"], yield_data))

        # 并行启动所有组
        tasks = [_asyncio.create_task(run_group(g, i)) for i, g in enumerate(groups)]

        # 消费 queue，实时 SSE 推送
        done_count = 0
        while done_count < len(groups):
            event_type, idx, group_label, results = await queue.get()
            done_count += 1

            yield f"event: progress\ndata: {_json.dumps({'phase': 'group_done', 'message': f'完成: {group_label} ({done_count}/{len(groups)})'}, ensure_ascii=False)}\n\n"

            for r in results:
                yield f"event: node_complete\ndata: {_json.dumps({'node_id': r['node_id'], 'label': r['label'], 'content': r['content']}, ensure_ascii=False)}\n\n"
                success_count += 1

        # 等待所有 task 真正结束
        await _asyncio.gather(*tasks, return_exceptions=True)

        yield f"event: complete\ndata: {_json.dumps({'total_nodes': total_nodes, 'success_count': success_count}, ensure_ascii=False)}\n\n"

    from fastapi.responses import StreamingResponse as _SSEResponse
    return _SSEResponse(
        sse_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/projects/{project_id}/analyze-node", summary="单节点重新提取（SSE）")
async def analyze_single_node(project_id: str, body: dict = Body(...)):
    """
    单独重新提取某个节点，按统一 docanalysis 协议返回最终结果。

    Body: { node_id, node_label, extraction_prompt, raw_document? }
    SSE 事件: {"type":"bid_attachments","items":[...]} | {"type":"done","node_id":"...","content":"..."} | {"type":"error","message":"..."}
    """
    import json as _json
    import re as _re
    from fastapi.responses import StreamingResponse as _SR

    node_id = body.get("node_id", "")
    node_label = body.get("node_label", "")
    extraction_prompt = body.get("extraction_prompt", "")
    raw_doc = str(body.get("raw_document", "") or "").strip()
    if not raw_doc:
        raw_doc = _load_raw_document(project_id)
    raw_doc = raw_doc[:300000]

    if not node_id or not raw_doc:
        raise HTTPException(status_code=400, detail="缺少 node_id 或项目原文缓存不存在")

    dify_key = _get_workflow_key("doc_analysis") or _get_workflow_key("requirement_extractor")
    if not dify_key:
        raise HTTPException(status_code=500, detail="工作流 API Key 未配置")

    # 读取 systemPrompt
    config_path = Path(__file__).parent.parent.parent / "config" / "analysis_framework.json"
    system_prompt_base = ""
    node_by_id: dict[str, dict] = {}
    if config_path.exists():
        system_prompt_base, all_nodes = load_docanalysis_framework(config_path)
        node_by_id = build_docanalysis_node_index(all_nodes)
    node_def = node_by_id.get(str(node_id).strip()) or {}
    node_label = str(node_def.get("label") or node_label or "").strip()
    extraction_prompt = str(node_def.get("extractionPrompt") or extraction_prompt or "").strip()
    if not node_label:
        node_label = node_id
    if not extraction_prompt:
        raise HTTPException(status_code=400, detail="缺少 extraction_prompt，且未在解析框架中找到该节点定义")

    async def event_generator():
        prompt_nodes = [{
            "id": node_id,
            "label": node_label,
            "extractionPrompt": extraction_prompt,
        }]
        combined_system = build_docanalysis_system_prompt(
            system_prompt_base,
            prompt_nodes,
            f"单节点重提取：{node_label}",
        )
        try:
            outputs = {}
            got_finished = False
            async for chunk in _call_dify_workflow_stream(dify_key, {
                "system_prompt": combined_system,
                "raw_document": raw_doc,
                "node_label": node_label,
            }):
                if isinstance(chunk, dict) and chunk.get("__finished__"):
                    got_finished = True
                    outputs = chunk.get("outputs", {}) or {}
                    break
            if not got_finished:
                raise RuntimeError("解析工作流异常中断（未收到 finished 事件）")

            raw_text = extract_docanalysis_text_output(outputs)
            content_text, attachments_payload = split_bid_attachments_tag(raw_text)
            result_map = parse_docanalysis_result_map(content_text)
            content = extract_docanalysis_node_content(result_map, node_id)
            if isinstance(content, (dict, list)):
                content = _json.dumps(content, ensure_ascii=False, indent=2)
            content = str(content).strip()

            bid_items = parse_bid_attachments_payload(attachments_payload)
            if bid_items:
                yield f"data: {_json.dumps({'type': 'bid_attachments', 'items': bid_items}, ensure_ascii=False)}\n\n"
            yield f"data: {_json.dumps({'type': 'done', 'node_id': node_id, 'content': content}, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.error(f"单节点提取 SSE 失败 [{node_id}]: {e}")
            yield f"data: {_json.dumps({'type': 'error', 'message': str(e)}, ensure_ascii=False)}\n\n"

    return _SR(event_generator(), media_type="text/event-stream")



@router.post("/projects/{project_id}/analysis-report", summary="保存解析报告到后端")
async def save_analysis_report(project_id: str, body: dict = Body(...)):
    """持久化解析报告，优先写入项目主记录，同时保留文件镜像。"""
    import json as _json
    import re as _re
    if not _re.match(r'^[a-zA-Z0-9_-]+$', project_id):
        raise HTTPException(status_code=400, detail="无效的 project_id")
    report = body.get("analysis_report", [])
    db = SessionLocal()
    try:
        row = db.query(ProjectRecord).filter_by(id=project_id).first()
        if row:
            if isinstance(row.data, str):
                try:
                    data = _json.loads(row.data)
                except Exception:
                    data = {}
            elif isinstance(row.data, dict):
                data = dict(row.data)
            else:
                data = {}
            if not isinstance(data, dict):
                data = {}
            data["analysisReport"] = report
            row.data = _json.dumps(data, ensure_ascii=False)
            db.add(row)
            db.commit()
    except Exception as e:
        db.rollback()
        logger.warning(f"[{project_id}] 保存 analysisReport 到项目记录失败: {e}")
    finally:
        db.close()
    save_dir = Path(__file__).parent.parent.parent.parent / "data" / "projects"
    save_dir.mkdir(parents=True, exist_ok=True)
    save_path = save_dir / f"{project_id}_analysis.json"
    with open(save_path, "w", encoding="utf-8") as f:
        _json.dump(report, f, ensure_ascii=False, indent=2)
    logger.info(f"解析报告已保存: {save_path}")
    return {"message": "保存成功", "path": str(save_path)}


@router.get("/projects/{project_id}/analysis-report", summary="读取解析报告")
async def get_analysis_report(project_id: str):
    """读取解析报告，优先从项目主记录读取，不存在再回退到文件镜像。"""
    import json as _json
    import re as _re
    if not _re.match(r'^[a-zA-Z0-9_-]+$', project_id):
        raise HTTPException(status_code=400, detail="无效的 project_id")
    db = SessionLocal()
    try:
        row = db.query(ProjectRecord).filter_by(id=project_id).first()
        if row and isinstance(row.data, dict):
            report = row.data.get("analysisReport") or row.data.get("analysis_report")
            if isinstance(report, list):
                return {
                    "analysis_report": report,
                    "analysis_v2": row.data.get("analysisV2") or row.data.get("analysis_v2") or {},
                }
    except Exception as e:
        logger.warning(f"[{project_id}] 读取 analysisReport 主记录失败: {e}")
    finally:
        db.close()
    save_path = Path(__file__).parent.parent.parent.parent / "data" / "projects" / f"{project_id}_analysis.json"
    if not save_path.exists():
        return {"analysis_report": [], "analysis_v2": {}}
    with open(save_path, "r", encoding="utf-8") as f:
        report = _json.load(f)
    return {"analysis_report": report, "analysis_v2": {}}


# ──────────────────────────────────────────────────────────────────────────────
# 知识库同步接口（异步后台触发 + 进度查询）
# ──────────────────────────────────────────────────────────────────────────────

import json as _json_sync
from datetime import datetime


def _get_kb_sync_status_path(job_id: str) -> Path:
    """获取同步任务状态文件路径"""
    status_dir = PRO_ENGINE_ROOT / "data" / "kb_sync_status"
    status_dir.mkdir(parents=True, exist_ok=True)
    return status_dir / f"{job_id}.json"


@router.post("/kb/sync", summary="触发知识库异步同步")
async def trigger_kb_sync(request: dict = Body(default={})):
    """
    在后台异步触发知识库同步（sync_kb.py），返回 job_id(task_id)。
    可选参数：file_prefix（仅同步指定前缀的文件）、llm_mode（覆盖 LLM 模式）
    """
    sync_script = PRO_ENGINE_ROOT / "scripts" / "sync_kb.py"
    if not sync_script.exists():
        raise HTTPException(status_code=404, detail="sync_kb.py script not found.")
    try:
        task_manager.ensure_backend_ready()
    except RuntimeError as e:
        raise HTTPException(status_code=500, detail={"code": "TASK_BACKEND_UNAVAILABLE", "message": str(e)})

    file_prefix = str((request or {}).get("file_prefix", "") or "").strip()
    llm_mode = str((request or {}).get("llm_mode", os.environ.get("PIPT_LLM_MODE", "augment")) or "augment").strip()
    dedupe_key = f"kb:{file_prefix.lower()}:{llm_mode.lower()}"

    existing_task_id = await _acquire_kb_sync_dedupe_key(dedupe_key)
    if existing_task_id:
        return {
            "job_id": existing_task_id,
            "task_id": existing_task_id,
            "status": "running",
            "message": "同配置知识库同步任务正在执行，已复用现有任务",
        }

    allowed, details = await task_manager.try_acquire_task_slot(
        project_id=_KB_SYNC_PROJECT,
        task_type="knowledge_sync",
        enforce_project_limit=False,
        max_type_running=_kb_sync_max_running(),
    )
    if not allowed:
        raise HTTPException(
            status_code=429,
            detail={
                "code": "TASK_LIMIT_REACHED",
                "message": "知识库同步并发达到上限，请稍后重试",
                "limits": task_manager.get_limits(),
                "metrics": details,
            },
        )

    task_id: Optional[str] = None
    try:
        task_id = task_manager.create_task("knowledge_sync", _KB_SYNC_PROJECT)
        await _set_kb_sync_dedupe_key(dedupe_key, task_id)

        status_path = _get_kb_sync_status_path(task_id)
        _upsert_kb_sync_status_file(
            status_path,
            {
                "job_id": task_id,
                "task_id": task_id,
                "status": "starting",
                "started_at": datetime.now().isoformat(),
                "file_prefix": file_prefix,
                "llm_mode": llm_mode,
                "total": 0,
                "processed": 0,
                "failed": 0,
                "current_file": "",
            },
        )

        env = os.environ.copy()
        env["PIPT_LLM_MODE"] = llm_mode
        env["KB_SYNC_JOB_ID"] = task_id
        env["KB_SYNC_STATUS_DIR"] = str(PRO_ENGINE_ROOT / "data" / "kb_sync_status")
        cmd = [sys.executable, str(sync_script)]
        if file_prefix:
            cmd.append(file_prefix)

        bg = _asyncio.create_task(
            _run_kb_sync_subprocess(
                task_id,
                dedupe_key,
                cmd,
                env=env,
                status_path=status_path,
            )
        )
        task_manager.set_async_task(task_id, bg)
        logger.info("知识库同步任务已启动: task_id=%s, llm_mode=%s, file_prefix=%s", task_id, llm_mode, file_prefix)
        return {
            "job_id": task_id,
            "task_id": task_id,
            "status": "running",
            "message": f"知识库同步已启动（LLM 模式: {llm_mode}），可通过 job_id 轮询进度",
        }
    except Exception as e:
        if task_id:
            await _release_kb_sync_dedupe_key(dedupe_key, task_id)
        logger.error("启动 sync_kb.py 失败: %s", e, exc_info=True)
        raise HTTPException(status_code=500, detail=f"同步任务启动失败: {e}")


@router.get("/kb/sync-status/{job_id}", summary="查询知识库同步任务进度")
async def get_kb_sync_status(job_id: str):
    """
    通过 job_id 查询知识库同步任务的当前状态和进度。
    状态值：starting / running / completed / failed
    """
    import re as _re
    if not _re.match(r'^[a-f0-9]{12}$', job_id):
        raise HTTPException(status_code=400, detail="无效的 job_id 格式")

    task = task_manager.get_task(job_id)
    if task:
        status_map = {
            "running": "running",
            "done": "completed",
            "error": "failed",
            "cancelled": "cancelled",
            "timeout": "failed",
        }
        mapped = status_map.get(task.status, "running")
        status_path = _get_kb_sync_status_path(job_id)
        if status_path.exists():
            try:
                with open(status_path, "r", encoding="utf-8") as f:
                    data = _json_sync.load(f)
                data.setdefault("job_id", job_id)
                data.setdefault("task_id", job_id)
                data["status"] = mapped
                return data
            except Exception:
                pass
        return {
            "job_id": job_id,
            "task_id": job_id,
            "status": mapped,
            "started_at": datetime.fromtimestamp(task.created_at).isoformat(),
            "total": 0,
            "processed": 0,
            "failed": 1 if mapped == "failed" else 0,
            "current_file": "",
            "error": task.error if mapped == "failed" else "",
        }

    status_path = _get_kb_sync_status_path(job_id)
    if not status_path.exists():
        raise HTTPException(status_code=404, detail=f"未找到任务 {job_id} 的状态记录")

    try:
        with open(status_path, "r", encoding="utf-8") as f:
            data = _json_sync.load(f)
        return data
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"读取状态文件失败: {e}")


@router.get("/kb/sync-jobs", summary="列出最近的知识库同步任务")
async def list_kb_sync_jobs():
    """返回最近 20 个同步任务的状态摘要（按时间倒序）"""
    status_dir = PRO_ENGINE_ROOT / "data" / "kb_sync_status"
    if not status_dir.exists():
        return {"jobs": []}

    jobs = []
    for f in sorted(status_dir.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)[:20]:
        try:
            with open(f, encoding="utf-8") as fp:
                data = _json_sync.load(fp)
                jobs.append({
                    "job_id": data.get("job_id"),
                    "status": data.get("status"),
                    "started_at": data.get("started_at"),
                    "total": data.get("total", 0),
                    "processed": data.get("processed", 0),
                    "failed": data.get("failed", 0),
                })
        except Exception:
            continue

    return {"jobs": jobs}


# ── PDF 导出解析报告 ──────────────────────────────────────────
class _ExportReportRequest(_BaseModel):
    project_name: str = "招标文件"
    nodes: list = []

@router.post("/projects/export-report", summary="导出解析报告 PDF")
async def export_report_pdf(req: _ExportReportRequest):
    """
    接收前端 analysisNodes JSON，生成带大纲书签的 PDF 并返回下载。
    请求 body: { "project_name": "...", "nodes": [...] }
    """
    import io
    import json as json_mod
    from datetime import datetime
    from fastapi.responses import StreamingResponse
    from weasyprint import HTML as WeasyprintHTML

    project_name = req.project_name
    nodes = req.nodes
    export_time = datetime.now().strftime("%Y-%m-%d %H:%M")

    # ── 节点内容 → HTML 片段 ──────────────────────────
    def _content_to_html(content: str) -> str:
        import re as re_mod
        txt = (content or "").strip()
        if not txt:
            return '<p class="empty">（未提取）</p>'

        # JSON 评分表（处理二次序列化：content 可能是 JSON 字符串套 JSON 字符串）
        if txt.startswith("{") or txt.startswith("[") or txt.startswith('"'):
            try:
                parsed = json_mod.loads(txt)
                # 处理双重序列化：Dify 有时把 JSON 对象序列化成字符串再放入外层 JSON
                if isinstance(parsed, str):
                    parsed = json_mod.loads(parsed)
                if isinstance(parsed, dict):
                    items = parsed.get("items") or []
                elif isinstance(parsed, list):
                    items = parsed
                else:
                    items = []

                if items:
                    rows = []
                    for it in items:
                        criteria_text = (it.get("criteria", "") or "").replace("\n", "<br/>")
                        rows.append(
                            f'<tr>'
                            f'<td>{it.get("name","")}</td>'
                            f'<td class="criteria">{criteria_text}</td>'
                            f'<td class="score">{it.get("max_score",0)}分</td>'
                            f'</tr>'
                        )
                    total = sum(it.get("max_score", 0) for it in items)
                    rows.append(
                        f'<tr class="total-row">'
                        f'<td colspan="2" style="text-align:right;font-weight:600;">合计</td>'
                        f'<td class="score" style="font-weight:700;color:#0369a1;">{total}分</td>'
                        f'</tr>'
                    )
                    note = parsed.get("note", "") if isinstance(parsed, dict) else ""
                    note_html = f'<p class="note">{note}</p>' if note else ""
                    return (
                        '<table class="score-table">'
                        '<tr class="total-row">'
                        '<th style="width:120px;text-align:left;">评分项</th>'
                        '<th style="text-align:left;">评分规则</th>'
                        '<th style="width:60px;text-align:right;">满分</th>'
                        '</tr>' + "".join(rows) + '</table>' + note_html
                    )
                else:
                    # items 为空：取 note 字段显示说明文字，防止原始 JSON 泄漏到报告
                    if isinstance(parsed, dict):
                        note = parsed.get("note", "")
                        return f'<p class="empty">{note if note else "招标文件未提供此项信息"}</p>'
            except Exception as e:
                import logging as _log
                _log.getLogger(__name__).warning(f"[export] scoring JSON parse failed, content prefix={repr(txt[:80])}, err={e}")

        # XML 列表 <要点>...</要点>
        list_re = re_mod.findall(r'<要点([^>]*)>([\s\S]*?)</要点>', txt)
        if list_re:
            items_html = []
            for i, (attrs, text_val) in enumerate(list_re):
                mandatory = 'mandatory' in attrs
                cls = ' class="mandatory"' if mandatory else ''
                items_html.append(
                    f'<div class="list-item"{cls}>'
                    f'<span class="list-num">{i+1}.</span>'
                    f'<span class="list-text">{text_val.strip()}</span>'
                    f'</div>'
                )
            return "".join(items_html)

        # XML 字段 <字段名>值</字段名>
        field_re = re_mod.findall(r'<([^/>\s][^>]*)>([\s\S]*?)</\1>', txt)
        fields = [(k, v.strip()) for k, v in field_re if not k.startswith("要点")]
        if fields:
            rows = []
            for label, value in fields:
                rows.append(
                    f'<tr>'
                    f'<td class="field-label">{label}</td>'
                    f'<td class="field-value">{value}</td>'
                    f'</tr>'
                )
            return '<table class="field-table">' + "".join(rows) + '</table>'

        # 纯文本
        return f'<p class="plain-text">{txt}</p>'

    # ── 递归生成 HTML ─────────────────────────────────
    def _walk_nodes(nodes_list: list, depth: int = 0) -> str:
        html_parts = []
        for node in nodes_list:
            node_id = node.get("id", "")
            label = node.get("label", "")
            content = node.get("content", "")
            children = node.get("children") or []
            has_children = len(children) > 0
            h_level = min(depth + 2, 4)
            html_parts.append(f'<h{h_level} id="{node_id}">{label}</h{h_level}>')
            if not has_children and content:
                html_parts.append(f'<div class="content-box">{_content_to_html(content)}</div>')
            elif not has_children:
                html_parts.append('<p class="empty">等待解析填充</p>')
            if has_children:
                html_parts.append(_walk_nodes(children, depth + 1))
        return "".join(html_parts)

    body_html = _walk_nodes(nodes)

    # WeasyPrint 支持全部标准 CSS，无需任何 hack
    full_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head><meta charset="utf-8"/>
<style>
@page {{ size: A4; margin: 15mm 14mm; }}
body {{
  font-family: 'SimSun', 'Noto Sans SC', 'Microsoft YaHei', 'PingFang SC', serif;
  font-size: 12px; color: #1f2937; line-height: 1.6;
}}
h1 {{ font-size: 20px; font-weight: 700; color: #0c4a6e; margin: 0 0 4px; }}
h2 {{ font-size: 16px; font-weight: 700; color: #0c4a6e; margin: 28px 0 8px;
      border-bottom: 2px solid #e0f2fe; padding-bottom: 6px; }}
h3 {{ font-size: 14px; font-weight: 700; color: #1e3a5f; margin: 20px 0 6px; }}
h4 {{ font-size: 12px; font-weight: 700; color: #0284c7; margin: 14px 0 4px; }}
h2, h3, h4 {{ page-break-after: avoid; }}
.header {{ border-bottom: 2px solid #0284c7; padding-bottom: 10px; margin-bottom: 24px; }}
.header p {{ font-size: 11px; color: #6b7280; margin: 0; }}
.content-box {{
  margin: 4px 0 16px; padding: 10px 12px;
  border: 1px solid #e5e7eb; border-radius: 6px; background: #f9fafb;
}}
.empty {{ color: #9ca3af; font-style: italic; font-size: 11px; }}
/* 列表项 */
.list-item {{
  display: flex; gap: 8px; margin-bottom: 6px; font-size: 12px; color: #374151;
}}
.list-item.mandatory {{
  border-left: 3px solid #fca5a5; padding-left: 8px;
}}
.list-num {{ color: #0284c7; font-weight: 600; flex-shrink: 0; min-width: 20px; }}
.list-text {{ flex: 1; }}
/* 评分表 */
.score-table {{
  width: 100%; border-collapse: collapse; font-size: 12px; margin: 8px 0;
}}
.score-table th, .score-table td {{
  padding: 6px 8px; border: 1px solid #d1d5db; vertical-align: top;
}}
.score-table .criteria {{ font-size: 11px; }}
.score-table .score {{ text-align: right; font-weight: 600; color: #0284c7; }}
.total-row {{ background: #f0f0f0; }}
.note {{ color: #9ca3af; font-size: 10px; margin-top: 4px; }}
/* 字段表 */
.field-table {{ width: 100%; border-collapse: collapse; }}
.field-label {{
  padding: 5px 8px; border-bottom: 1px solid #f3f4f6;
  font-size: 11px; color: #0284c7; vertical-align: top; width: 120px; font-weight: 600;
}}
.field-value {{
  padding: 5px 8px; border-bottom: 1px solid #f3f4f6;
  font-size: 12px; color: #374151; vertical-align: top;
}}
.plain-text {{ font-size: 12px; color: #374151; }}
</style>
</head>
<body>
<div class="header">
  <h1>{project_name} — 招标文件解析报告</h1>
  <p>导出时间：{export_time}</p>
</div>
{body_html}
</body>
</html>"""

    # WeasyPrint 生成 PDF
    pdf_bytes = WeasyprintHTML(string=full_html).write_pdf()

    pdf_buffer = io.BytesIO(pdf_bytes)
    filename = f"解析报告_{project_name}_{datetime.now().strftime('%Y%m%d')}.pdf"
    return StreamingResponse(
        pdf_buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{__import__('urllib.parse', fromlist=['quote']).quote(filename)}"},
    )
