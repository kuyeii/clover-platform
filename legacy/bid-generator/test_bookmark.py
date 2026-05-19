from docx import Document
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

doc = Document()
p = doc.add_paragraph()

# Create bookmark start
bm_start = OxmlElement('w:bookmarkStart')
bm_start.set(qn('w:id'), '0')
bm_start.set(qn('w:name'), 'MyBookmark123')

p._p.append(bm_start)

# Add run
run = p.add_run('This is bookmarked text.')

# Create bookmark end
bm_end = OxmlElement('w:bookmarkEnd')
bm_end.set(qn('w:id'), '0')

p._p.append(bm_end)

# Try adding a PAGEREF in the next line
p2 = doc.add_paragraph('Reference page: ')
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_pageref(paragraph, bookmark_name):
    # <w:fldSimple w:instr=" PAGEREF MyBookmark123 \h "> 
    # <w:r><w:t>3</w:t></w:r> 
    # </w:fldSimple>
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), f' PAGEREF {bookmark_name} \h ')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '0' # Placeholder for page number before F9
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)

add_pageref(p2, 'MyBookmark123')

doc.save('/home/haisuchen/pro-engine/test_bm.docx')
print("Saved test_bm.docx")
