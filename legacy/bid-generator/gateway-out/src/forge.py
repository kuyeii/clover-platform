# -*- coding: utf-8 -*-
"""
DocumentForge — 最终文档组装引擎

将已生成的所有元素整合为一份完整的、可交付的 .docx 标书文件：
  1. 全量占位符还原（BIDDER + PIPT 脱敏）
  2. 各章节正文拼接 → Markdown → DOCX
  3. 自评评分表嵌入（Word 内嵌表格）
  4. 附件追加（分节，每份独页）
"""

import hashlib
import html
import logging
import os
import re
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .restorer import PlaceholderRestorer
from .converter.md_to_docx import MarkdownToDocxConverter
from .markdown_norm import prepare_section_for_forge
from .svg_export import preprocess_svg_for_png

logger = logging.getLogger(__name__)

_DEFAULT_TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "official_template_cn.docx"
_PRO_ENGINE_ROOT = Path(__file__).resolve().parents[2]
_DIAGRAM_ARTIFACT_ROOT = Path(
    os.environ.get("DIAGRAM_ARTIFACT_DIR", str(_PRO_ENGINE_ROOT / "data" / "diagram_artifacts"))
)


# ── SVG 图表 → PNG 临时文件（供 DOCX 插图使用）───────────────────────────────

def _svg_to_temp_png(svg_text: str, title: str) -> Optional[str]:
    """
    使用 svglib 将 SVG 字符串渲染为临时 PNG 文件。
    以 SVG 内容的 MD5 命名，保证幂等：同一图表多次调用只渲染一次。

    Returns:
        str: 临时 PNG 绝对路径；处理失败时返回 None
    """
    try:
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM

        svg_proc = preprocess_svg_for_png(svg_text)
        slug = hashlib.md5(svg_proc.encode("utf-8", errors="replace")).hexdigest()[:16]
        tmp_dir = Path(tempfile.gettempdir()) / "proengine_diagrams"
        tmp_dir.mkdir(parents=True, exist_ok=True)
        tmp_svg = tmp_dir / f"diagram_{slug}.svg"
        tmp_png = tmp_dir / f"diagram_{slug}.png"

        # 已缓存则直接返回
        if tmp_png.exists() and tmp_png.stat().st_size > 0:
            return str(tmp_png)

        tmp_svg.write_text(svg_proc, encoding="utf-8")
        drawing = svg2rlg(str(tmp_svg))
        if drawing is None:
            logger.warning(f"svglib 无法解析图表 SVG（{title}），跳过插图")
            return None

        renderPM.drawToFile(drawing, str(tmp_png), fmt="PNG")
        tmp_svg.unlink(missing_ok=True)
        return str(tmp_png) if tmp_png.exists() else None
    except Exception as e:
        logger.warning(f"SVG→PNG 转换失败（{title}）: {e}")
        return None


def _safe_artifact_token(value: str, fallback: str = "default") -> str:
    token = re.sub(r"[^a-zA-Z0-9_.-]+", "_", str(value or "").strip())
    return token or fallback


def _get_html_attr(attrs_str: str, name: str) -> str:
    m = re.search(
        rf'\b{re.escape(name)}\s*=\s*(?:"([^"]*)"|\'([^\']*)\'|([^\s>]+))',
        attrs_str or "",
        flags=re.IGNORECASE,
    )
    if not m:
        return ""
    return html.unescape(next((g for g in m.groups() if g is not None), "")).strip()


def _read_diagram_artifact(project_id: str, diagram_id: str) -> str:
    """
    从后端落盘 artifact 读取 SVG。

    Args:
        project_id: 项目 ID，用于定位 data/diagram_artifacts/{project_id}
        diagram_id: 图表 artifact ID，不允许路径穿越

    Returns:
        str: SVG 文本；不存在或非法时返回空字符串
    """
    safe_diagram_id = _safe_artifact_token(diagram_id, "")
    if not safe_diagram_id or safe_diagram_id != str(diagram_id or "").strip():
        return ""

    candidates: list[Path] = []
    safe_project_id = _safe_artifact_token(project_id, "")
    if safe_project_id:
        candidates.append(_DIAGRAM_ARTIFACT_ROOT / safe_project_id / f"{safe_diagram_id}.svg")
    # 兼容历史正文缺 project_id 的情况：只按文件名在 artifact 根目录下兜底查找。
    candidates.extend(_DIAGRAM_ARTIFACT_ROOT.glob(f"*/{safe_diagram_id}.svg"))

    for path in candidates:
        try:
            resolved = path.resolve()
            try:
                resolved.relative_to(_DIAGRAM_ARTIFACT_ROOT.resolve())
            except ValueError:
                continue
            if resolved.exists() and resolved.is_file():
                svg_text = resolved.read_text(encoding="utf-8").strip()
                return svg_text if svg_text.lower().startswith("<svg") else ""
        except Exception as e:
            logger.warning("读取图表 artifact 失败（%s）: %s", safe_diagram_id, e)
    return ""


def _strip_diagrams_to_images(markdown: str, project_id: str = "") -> str:
    """
    扫描 Markdown 正文中的 <diagram> 块，将其转换为可被 md_to_docx 处理的格式：
    - 成功转 PNG → 替换为 ![标题](tmp_png路径) + 图注段落
    - 转换失败   → 替换为可读占位文字，不破坏文档结构
    """
    def _replace(m: re.Match) -> str:
        attrs_str = m.group(1)
        body = m.group(2)

        title = _get_html_attr(attrs_str, "title") or "架构图"
        diagram_id = _get_html_attr(attrs_str, "data-diagram-id")

        # 优先读取 artifact 引用，避免 DOCX 导出依赖前端回填大段 SVG。
        raw_svg = _read_diagram_artifact(project_id, diagram_id) if diagram_id else ""
        if raw_svg:
            png_path = _svg_to_temp_png(raw_svg, title)
            if png_path:
                return f"\n\n**图：{title}**\n\n![{title}]({png_path})\n\n"

        # 兼容旧正文：提取内嵌 SVG
        svg_m = re.search(r'<svg[\s\S]*?</svg>', body, re.IGNORECASE)
        if svg_m:
            # 基础校验：必须以 <svg 开头，防止脏数据
            raw_svg = svg_m.group(0).strip()
            if raw_svg.lower().startswith("<svg"):
                png_path = _svg_to_temp_png(raw_svg, title)
                if png_path:
                    return f"\n\n**图：{title}**\n\n![{title}]({png_path})\n\n"

        # 降级：图表转换失败时插入友好文字提示，不阻断导出流程
        logger.warning(f"图表 '{title}' SVG 无效、artifact 缺失或转换失败，以文字占位替代")
        return f"\n\n> 【图表：{title}（电子版文档中可查看）】\n\n"

    return re.sub(
        r'<diagram\s+([^>]*)>([\s\S]*?)</diagram>',
        _replace,
        markdown,
        flags=re.IGNORECASE,
    )


# 自评情况中文映射
_RESPONSE_LABEL = {
    "full": "完全响应",
    "partial": "部分响应",
    "none": "不响应",
    "": "未填写",
}


def _add_page_break(doc: Document) -> None:
    """在文档末尾添加分页符"""
    doc.add_page_break()


def _add_scoring_table(doc: Document, scoring_rows: list[dict]) -> None:
    """
    在文档末尾追加自评评分表（Word 内嵌表格）

    Args:
        doc: python-docx Document
        scoring_rows: ScoringRow 字典列表
    """
    if not scoring_rows:
        return

    _add_page_break(doc)

    # 评分表标题
    heading = doc.add_heading("自评评分表", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    headers = ["评分指标", "最高分", "评分标准", "自评情况", "自评说明", "证明材料引用"]
    col_widths_cm = [5.0, 1.5, 6.5, 2.0, 7.0, 5.5]  # 单位：cm，合计约 27.5cm

    table = doc.add_table(rows=1 + len(scoring_rows) + 1, cols=len(headers))
    table.style = "Table Grid"

    from docx.shared import Cm
    # 设置列宽
    for col_idx, width in enumerate(col_widths_cm):
        for row in table.rows:
            row.cells[col_idx].width = Cm(width)

    # ── 表头行 ──
    header_row = table.rows[0]
    for col_idx, h in enumerate(headers):
        cell = header_row.cells[col_idx]
        cell.text = h
        # 表头样式：加粗、深蓝底色、白色字体
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        # 背景色
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1A6FA8")
        tc_pr.append(shd)

    # ── 数据行 ──
    for row_idx, row_data in enumerate(scoring_rows, 1):
        row = table.rows[row_idx]
        evidence = row_data.get("evidenceRefs", row_data.get("evidence_refs", []))
        evidence_list = evidence if isinstance(evidence, list) else (str(evidence).split("\n") if evidence else [])

        values = [
            row_data.get("indicator", ""),
            str(row_data.get("maxScore", row_data.get("max_score", ""))),
            row_data.get("criteria", ""),
            _RESPONSE_LABEL.get(row_data.get("selfResponse", row_data.get("self_response", "")), ""),
            row_data.get("selfComment", row_data.get("self_comment", "")),
            "", # 证据列交由特定逻辑渲染
        ]
        for col_idx, val in enumerate(values):
            cell = row.cells[col_idx]
            
            if col_idx == 5:
                # 证据位置列：插入原生的 Word 书签页码引用指令 (PAGEREF)
                if not evidence_list:
                    continue
                p = cell.paragraphs[0]
                for i, e_id in enumerate(evidence_list):
                    if not str(e_id).strip():
                        continue
                    if i > 0:
                        p.add_run("\n")
                    r_pre = p.add_run("详见第 ")
                    r_pre.font.size = Pt(9)
                    
                    bm_name = "BM_" + str(e_id).replace("-", "_")[:35]
                    
                    fld = OxmlElement('w:fldSimple')
                    fld.set(qn('w:instr'), f' PAGEREF {bm_name} \\h ')
                    r_field = OxmlElement('w:r')
                    t_field = OxmlElement('w:t')
                    t_field.text = '--'  # 在 Word 更新域代码前显示的虚拟占位符
                    r_field.append(t_field)
                    fld.append(r_field)
                    p._p.append(fld)
                    
                    r_post = p.add_run(" 页")
                    r_post.font.size = Pt(9)
            else:
                cell.text = val
                cell.paragraphs[0].runs[0].font.size = Pt(9) if cell.paragraphs[0].runs else None

    # ── 合计行 ──
    total_row = table.rows[-1]
    total_score = sum(
        int(r.get("maxScore", r.get("max_score", 0)) or 0) for r in scoring_rows
    )
    total_row.cells[0].text = "合计"
    total_row.cells[0].paragraphs[0].runs[0].bold = True
    total_row.cells[1].text = str(total_score)
    total_row.cells[1].paragraphs[0].runs[0].bold = True

    logger.info(f"评分表追加完成: {len(scoring_rows)} 行")


def _add_attachments(doc: Document, attachments: list[dict]) -> None:
    """
    追加附件内容到文档末尾，每份附件独占新页

    Args:
        doc: python-docx Document
        attachments: [{label: str, content: str}] Markdown 格式内容
    """
    for att in attachments:
        label = att.get("label", "附件")
        content = att.get("content", "")
        if not content.strip():
            continue

        _add_page_break(doc)
        # 与正文之间空一行，再写“附件”二字并首行缩进两字符
        doc.add_paragraph("")
        annex_prefix = doc.add_paragraph("附件")
        annex_prefix.paragraph_format.first_line_indent = Pt(24)  # 约两字符缩进
        annex_prefix.paragraph_format.space_before = Pt(0)
        annex_prefix.paragraph_format.space_after = Pt(0)
        for r in annex_prefix.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        heading = doc.add_paragraph(str(label).strip())
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        heading.paragraph_format.first_line_indent = Pt(24)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(0)
        for r in heading.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 简单按行追加附件正文（附件一般是纯文本/简短 Markdown）
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            # 跳过 Markdown 标题（附件里通常不需要额外的标题层级）
            if line.startswith("#"):
                clean = line.lstrip("#").strip()
                p = doc.add_paragraph(clean)
                p.paragraph_format.first_line_indent = Pt(24)
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.first_line_indent = Pt(24)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(12)
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    if attachments:
        logger.info(f"附件追加完成: {len(attachments)} 份")


class DocumentForge:
    """
    最终文档组装器

    用法:
        forge = DocumentForge(mapping_table={...}, bidder_info={...})
        docx_bytes = forge.build(sections=[...], scoring_rows=[...], attachments=[...])
    """

    def __init__(
        self,
        mapping_table: Optional[dict] = None,
        bidder_info: Optional[dict] = None,
        image_map: Optional[dict] = None,
        template_path: Optional[str] = None,
        project_id: Optional[str] = None,
    ):
        """
        Args:
            mapping_table: PIPT 脱敏映射表
            bidder_info: 投标人信息（BidderInfo 结构）
            image_map: 图片占位符映射表 (如 {"{{IMG_0001}}": "/path/to/img.png"})
            template_path: Word 模板路径（可选）
            project_id: 项目 ID，用于解析后端落盘的图表 artifact
        """
        self.restorer = PlaceholderRestorer(mapping_table or {})
        self.bidder_info = bidder_info or {}
        self.image_map = image_map or {}
        self.project_id = project_id or ""
        env_template = os.environ.get("FORGE_TEMPLATE_PATH", "").strip()
        candidate = template_path or env_template or str(_DEFAULT_TEMPLATE_PATH)
        self.template_path = candidate if candidate and Path(candidate).exists() else None
        if self.template_path:
            logger.info("DocumentForge 使用模板: %s", self.template_path)
        else:
            logger.warning("DocumentForge 未找到模板，将使用代码默认样式")

    def build(
        self,
        sections: list[dict],
        scoring_rows: Optional[list[dict]] = None,
        attachments: Optional[list[dict]] = None,
    ) -> bytes:
        """
        执行完整的文档组装流程，返回 .docx 文件字节流

        Args:
            sections: 章节列表 [{id, title, content}]，content 为 Markdown
            scoring_rows: 自评评分表行列表（可选）
            attachments: 附件列表 [{label, content}]（可选）

        Returns:
            bytes: .docx 文件内容
        """
        scoring_rows = scoring_rows or []
        attachments = attachments or []

        # ── 步骤 1：拼接并还原 Markdown 正文 ──
        parts = []
        for sec in sections:
            title = sec.get("title", "")
            heading_text = str(sec.get("heading_text") or "").strip()
            heading_number = str(sec.get("heading_number") or "").strip()
            raw = sec.get("content", "").strip()
            sec_id = sec.get("id", "")
            heading_level = int(sec.get("heading_level", 1) or 1)
            heading_level = max(1, min(6, heading_level))
            title_only = bool(sec.get("title_only"))
            if not raw and not (title and title_only):
                continue
            plain_title = heading_text or title
            display_title = f"{heading_number} {plain_title}".strip() if heading_number else plain_title
            body = prepare_section_for_forge(display_title or title, raw) if raw else ""
            if raw and not body.strip():
                continue
            bookmark_id = str(sec.get("bookmark_id") or "").strip()
            if not bookmark_id and sec_id:
                bookmark_id = f"BM_{str(sec_id).replace('-', '_')[:35]}"
            bm_tag = f" {{#{bookmark_id}}}" if bookmark_id else ""
            if display_title:
                prefix = "#" * heading_level
                if body.strip():
                    parts.append(f"{prefix} {display_title}{bm_tag}\n\n{body}")
                else:
                    parts.append(f"{prefix} {display_title}{bm_tag}")
            elif body.strip():
                parts.append(body)

        full_markdown = "\n\n---\n\n".join(parts)

        # 全量占位符还原（先 BIDDER，后 PIPT）
        full_markdown = self.restorer.restore_all(full_markdown, self.bidder_info)

        # ── 步骤 1.5a：还原图片占位符为 Markdown 图片语法 ──
        # image_map value 可能是绝对路径字符串，或 {abs_path, preview_url} 字典（新格式）
        for img_ph, img_val in self.image_map.items():
            if isinstance(img_val, dict):
                img_path = img_val.get("abs_path") or img_val.get("preview_url", "")
            else:
                img_path = str(img_val)
            if img_path:
                if img_ph.startswith("__PRO_IMG_"):
                    # __PRO_IMG_ 占位符已经在 ![图注](__PRO_IMG_xxx__) 内部，直接替换本身即可
                    full_markdown = full_markdown.replace(img_ph, img_path)
                else:
                    # 兼容老的 {{IMG_xxx}} 独立占位符模式
                    full_markdown = full_markdown.replace(img_ph, f"\n\n![img]({img_path})\n\n")

        # ── 步骤 1.5b：将正文内的 <diagram> SVG 块转为临时 PNG ──
        # 必须在 md_to_docx 之前处理，否则 <diagram> 标签会被当作乱码文本写入 Word
        full_markdown = _strip_diagrams_to_images(full_markdown, self.project_id)

        # ── 步骤 2：Markdown → DOCX ──
        converter = MarkdownToDocxConverter(template_path=self.template_path)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name

        converter.convert(full_markdown, tmp_path)
        doc = Document(tmp_path)

        # 清理临时文件
        try:
            Path(tmp_path).unlink()
        except Exception:
            pass

        # ── 步骤 3：追加评分表 ──
        if scoring_rows:
            _add_scoring_table(doc, scoring_rows)

        # ── 步骤 4：追加附件 ──
        if attachments:
            _add_attachments(doc, attachments)

        # ── 步骤 5：输出为字节流 ──
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)

        logger.info(
            f"DocumentForge 完成: {len(sections)} 章节, "
            f"{len(scoring_rows)} 评分行, {len(attachments)} 附件"
        )
        return buf.read()
