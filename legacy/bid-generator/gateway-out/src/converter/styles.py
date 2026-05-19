from typing import Optional
# -*- coding: utf-8 -*-
"""
Word 文档样式管理
定义标书排版所需的标题、正文、表格等样式
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement


def _set_style_font(style_obj, zh_font: str, en_font: str = "Times New Roman", size_pt: int = 12, bold: bool = False) -> None:
    """同时设置 style 的西文与中文字体，避免中文回退导致字体混乱。"""
    style_obj.font.name = en_font
    style_obj.font.size = Pt(size_pt)
    style_obj.font.bold = bold
    try:
        style_obj._element.rPr.rFonts.set(qn("w:eastAsia"), zh_font)
        style_obj._element.rPr.rFonts.set(qn("w:ascii"), en_font)
        style_obj._element.rPr.rFonts.set(qn("w:hAnsi"), en_font)
    except Exception:
        pass


def _ensure_style(doc: Document, style_name: str):
    try:
        return doc.styles[style_name]
    except Exception:
        return None


def apply_bid_styles(doc: Document) -> Document:
    """
    为 Word 文档应用标书排版样式

    Args:
        doc: python-docx Document 对象

    Returns:
        Document: 应用样式后的文档
    """
    style = doc.styles

    # ==================== 正文样式 ====================
    normal_style = style["Normal"]
    _set_style_font(normal_style, zh_font="宋体", en_font="Times New Roman", size_pt=12, bold=False)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal_style.paragraph_format.line_spacing = Pt(28)
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进两字符

    # ==================== 标题样式 ====================
    # 一级标题：黑体，三号，左对齐（按公文正文层级）
    h1 = _ensure_style(doc, "Heading 1")
    if h1:
        _set_style_font(h1, zh_font="黑体", en_font="Times New Roman", size_pt=16, bold=True)
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        h1.paragraph_format.line_spacing = Pt(32)
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(6)
        h1.paragraph_format.first_line_indent = Cm(0)

    # 二级标题：楷体_GB2312，三号，左对齐
    h2 = _ensure_style(doc, "Heading 2")
    if h2:
        _set_style_font(h2, zh_font="楷体_GB2312", en_font="Times New Roman", size_pt=16, bold=True)
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        h2.paragraph_format.line_spacing = Pt(28)
        h2.paragraph_format.space_before = Pt(8)
        h2.paragraph_format.space_after = Pt(4)
        h2.paragraph_format.first_line_indent = Cm(0.74)

    # 三级标题：仿宋_GB2312，三号，左对齐
    h3 = _ensure_style(doc, "Heading 3")
    if h3:
        _set_style_font(h3, zh_font="仿宋_GB2312", en_font="Times New Roman", size_pt=16, bold=True)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(0, 0, 0)
        h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        h3.paragraph_format.line_spacing = Pt(28)
        h3.paragraph_format.space_before = Pt(6)
        h3.paragraph_format.space_after = Pt(3)
        h3.paragraph_format.first_line_indent = Cm(0.74)

    # 列表样式继承正文字体，避免项目符号处字体跳变
    for list_style_name in ("List Bullet", "List Number"):
        s = _ensure_style(doc, list_style_name)
        if s:
            _set_style_font(s, zh_font="宋体", en_font="Times New Roman", size_pt=12, bold=False)
            s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            s.paragraph_format.line_spacing = Pt(28)

    return doc


def set_page_margins(doc: Document) -> Document:
    """按模板规范设置页边距：上3.7 下3.5 左2.8 右2.6；页眉1.5 页脚2.5（厘米）。"""
    for section in doc.sections:
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(2.5)
    return doc


def _append_page_field(paragraph, zh_font: str = "宋体"):
    run = paragraph.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)  # 四号
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), zh_font)
    except Exception:
        pass
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def apply_page_number_footer(doc: Document) -> Document:
    """奇偶页页码：四号宋体，“— n —”；奇数页右下，偶数页左下。"""
    doc.settings.odd_and_even_pages_header_footer = True
    for section in doc.sections:
        odd_footer = section.footer
        p_odd = odd_footer.paragraphs[0] if odd_footer.paragraphs else odd_footer.add_paragraph()
        p_odd.text = ""
        p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_odd.add_run("— ")
        _append_page_field(p_odd)
        p_odd.add_run(" —")

        even_footer = section.even_page_footer
        p_even = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
        p_even.text = ""
        p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_even.add_run("— ")
        _append_page_field(p_even)
        p_even.add_run(" —")
    return doc
