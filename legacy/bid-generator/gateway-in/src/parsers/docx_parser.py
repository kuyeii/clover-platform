# -*- coding: utf-8 -*-
"""
DOCX 文件解析器
基于 python-docx 实现 Word 招标文件的文本、表格、图片提取
"""

import logging
from typing import Optional

from .base import BaseParser, ParsedDocument

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """DOCX 文件解析器"""

    def parse(self) -> ParsedDocument:
        """解析 DOCX 文件，返回结构化文档"""
        from docx import Document
        from .document_tree import DocumentNode

        doc = ParsedDocument(
            source_path=str(self.file_path),
            file_format="docx",
        )

        document = Document(str(self.file_path))

        # 提取元数据
        core_props = document.core_properties
        doc.metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
        }

        # 构建文档树（基于 Heading 样式）
        root = DocumentNode(level=0, title="Root")
        node_stack = [0]       # 层级栈
        node_refs = [root]     # 节点引用栈

        all_text_parts = []
        
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        
        # 按序遍历文档体中的段落与表格
        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                para = Paragraph(child, document)
                text = para.text.strip()
                if not text:
                    continue
                all_text_parts.append(text)

                # 检查是否为标题样式
                is_heading = False
                level = 0
                if para.style and para.style.name and para.style.name.startswith("Heading"):
                    try:
                        level = int(para.style.name.replace("Heading", "").strip())
                        is_heading = True
                    except (ValueError, AttributeError):
                        pass

                # 简单的正则启发式补助：常见的一、(一) 等标题（当无样式时）
                if not is_heading and (
                    text.startswith("第") and "章" in text[:10] or 
                    text.startswith("一、") or text.startswith("二、") or text.startswith("三、") or text.startswith("四、")
                ):
                    is_heading = True
                    level = 1

                if is_heading:
                    new_node = DocumentNode(level=level, title=text)
                    # 寻找父节点（向上出栈直到遇到比当前层级小的节点）
                    while len(node_stack) > 1 and node_stack[-1] >= level:
                        node_stack.pop()
                        node_refs.pop()
                    parent = node_refs[-1]
                    parent.children.append(new_node)
                    node_stack.append(level)
                    node_refs.append(new_node)
                else:
                    # 累加段落内容到当前活跃节点
                    node_refs[-1].content += text + "\n"
                    
            elif isinstance(child, CT_Tbl):
                table = Table(child, document)
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    table_data.append(row_data)
                
                doc.tables.append(table_data)
                
                # 把表格转换为易读的 Markdown 格式加入到正文中
                if table_data:
                    md_table = []
                    for i, r in enumerate(table_data):
                        md_table.append("| " + " | ".join(r) + " |")
                        if i == 0:
                            md_table.append("|" + "|".join(["---"] * len(r)) + "|")
                    
                    table_text = "\n" + "\n".join(md_table) + "\n"
                    all_text_parts.append(table_text)
                    node_refs[-1].content += table_text + "\n"

        doc.full_text = "\n".join(all_text_parts)
        doc.document_tree = root

        # 转为平铺 sections 以保持向后兼容
        doc.sections = [{"title": n.title, "content": n.content, "level": n.level} for n in root.children]

        # 提取图片信息
        for idx, rel in enumerate(document.part.rels.values()):
            if "image" in rel.reltype:
                doc.images.append({
                    "page": 0,  # DOCX 无明确页面概念
                    "index": idx,
                    "data": None,  # 暂不加载图片二进制
                    "format": rel.target_ref.split(".")[-1] if hasattr(rel, "target_ref") else "unknown",
                })

        logger.info(f"DOCX 解析完成: {len(doc.sections)} 个章节, {len(doc.tables)} 个表格, {len(doc.images)} 张图片")
        return doc

    def extract_text(self) -> str:
        """提取 DOCX 中的纯文本"""
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document = Document(str(self.file_path))
        parts = []
        for child in document.element.body.iterchildren():
            if isinstance(child, CT_P):
                para = Paragraph(child, document)
                if para.text.strip():
                    parts.append(para.text.strip())
            elif isinstance(child, CT_Tbl):
                table = Table(child, document)
                for row in table.rows:
                    row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    parts.append("| " + " | ".join(row_data) + " |")
        return "\n".join(parts)

    def extract_tables(self) -> list[list[list[str]]]:
        """提取 DOCX 中的所有表格"""
        from docx import Document

        document = Document(str(self.file_path))
        all_tables = []
        for table in document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            all_tables.append(table_data)
        return all_tables
