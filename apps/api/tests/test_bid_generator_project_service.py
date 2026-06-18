from __future__ import annotations

import json
import io
import asyncio
import sys
import unittest
import tempfile
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import AsyncMock, MagicMock, Mock, patch


API_ROOT = Path(__file__).resolve().parents[1]
API_ROOT_VALUE = str(API_ROOT)
if API_ROOT_VALUE not in sys.path:
    sys.path.insert(0, API_ROOT_VALUE)

from app.services import bid_generator_service as service


class BidGeneratorProjectServiceTests(unittest.TestCase):
    def test_create_project_raises_conflict_when_id_exists(self) -> None:
        conn = Mock()
        exists_result = Mock()
        exists_result.scalar_one.return_value = True
        duplicate_result = Mock()
        duplicate_result.first.return_value = {"id": "proj-1"}
        conn.execute.side_effect = [exists_result, duplicate_result]
        engine = MagicMock()
        engine.begin.return_value.__enter__.return_value = conn

        with patch.object(service, "get_engine", return_value=engine):
            with self.assertRaises(service.BidProjectConflict):
                service.create_project_payload(
                    {"id": "proj-1", "name": "项目一", "status": "uploading", "data": {"id": "proj-1"}}
                )

    def test_update_project_upserts_and_returns_project(self) -> None:
        row = _project_row(data={"name": "项目一", "nested": {"a": 1}})
        conn = Mock()
        exists_result = Mock()
        exists_result.scalar_one.return_value = True
        upsert_result = Mock()
        upsert_result.mappings.return_value.one.return_value = row
        conn.execute.side_effect = [exists_result, upsert_result]
        engine = MagicMock()
        engine.begin.return_value.__enter__.return_value = conn

        with patch.object(service, "get_engine", return_value=engine):
            payload = service.update_project_payload(
                "proj-1",
                {"name": "项目一", "status": "uploaded", "data": {"name": "项目一", "nested": {"a": 1}}},
            )

        self.assertEqual(payload["id"], "proj-1")
        self.assertEqual(payload["data"]["nested"]["a"], 1)
        params = conn.execute.call_args_list[1].args[1]
        self.assertEqual(params["project_id"], "proj-1")
        self.assertEqual(json.loads(params["data"])["nested"]["a"], 1)

    def test_patch_project_deep_merges_data_and_removes_keys(self) -> None:
        selected = _project_row(
            data={
                "name": "旧项目",
                "status": "uploaded",
                "nested": {"a": 1, "b": 2},
                "stale": True,
            }
        )
        updated = _project_row(
            name="新项目",
            status="ready",
            data={
                "name": "新项目",
                "status": "ready",
                "nested": {"a": 1, "b": 3},
            },
        )
        conn = Mock()
        exists_result = Mock()
        exists_result.scalar_one.return_value = True
        select_result = Mock()
        select_result.mappings.return_value.first.return_value = selected
        update_result = Mock()
        update_result.mappings.return_value.one.return_value = updated
        conn.execute.side_effect = [exists_result, select_result, update_result]
        engine = MagicMock()
        engine.begin.return_value.__enter__.return_value = conn

        with patch.object(service, "get_engine", return_value=engine):
            payload = service.patch_project_payload(
                "proj-1",
                {
                    "name": "新项目",
                    "status": "ready",
                    "data_patch": {"nested": {"b": 3}},
                    "remove_data_keys": ["stale"],
                },
            )

        self.assertEqual(payload["name"], "新项目")
        self.assertEqual(payload["status"], "ready")
        self.assertEqual(payload["data"]["nested"], {"a": 1, "b": 3})
        self.assertNotIn("stale", payload["data"])
        update_params = conn.execute.call_args_list[2].args[1]
        self.assertNotIn("stale", json.loads(update_params["data"]))
        self.assertEqual(json.loads(update_params["data"])["nested"], {"a": 1, "b": 3})

    def test_batch_create_projects_upserts_and_counts_created_updated(self) -> None:
        existing_result = Mock()
        existing_result.mappings.return_value.all.return_value = [{"id": "proj-existing"}]
        insert_result = Mock()
        conn = Mock()
        conn.execute.side_effect = [existing_result, insert_result, insert_result]
        engine = MagicMock()
        engine.begin.return_value.__enter__.return_value = conn

        with (
            patch.object(service, "get_engine", return_value=engine),
            patch.object(service, "_ensure_project_storage", return_value=None),
        ):
            payload = service.batch_create_projects_payload(
                [
                    {"id": "proj-existing", "name": "已有项目", "status": "editing", "data": {"id": "proj-existing"}},
                    {"id": "proj-new", "name": "新项目", "status": "uploading", "data": {"id": "proj-new"}},
                ]
            )

        self.assertEqual(payload, {"created": 1, "updated": 1})
        self.assertEqual(conn.execute.call_args_list[0].args[1]["project_ids"], ["proj-existing", "proj-new"])
        first_upsert = conn.execute.call_args_list[1].args[1]
        second_upsert = conn.execute.call_args_list[2].args[1]
        self.assertEqual(first_upsert["project_id"], "proj-existing")
        self.assertEqual(json.loads(first_upsert["data"])["id"], "proj-existing")
        self.assertEqual(second_upsert["project_id"], "proj-new")
        self.assertEqual(json.loads(second_upsert["data"])["id"], "proj-new")

    def test_save_analysis_report_updates_project_data_and_writes_mirror(self) -> None:
        selected = {"data": json.dumps({"name": "项目一", "analysisV2": {"schema_version": "v2"}}, ensure_ascii=False)}
        conn = Mock()
        exists_result = Mock()
        exists_result.scalar_one.return_value = True
        persist_exists = Mock()
        persist_exists.scalar_one.return_value = True
        select_result = Mock()
        select_result.mappings.return_value.first.return_value = selected
        update_result = Mock()
        conn.execute.side_effect = [persist_exists, select_result, update_result]
        engine = MagicMock()
        engine.begin.return_value.__enter__.return_value = conn
        mirror_path = Path("/tmp/clover-test/proj-1_analysis.json")

        with (
            patch.object(service, "_analysis_report_mirror_path", return_value=mirror_path),
            patch.object(service, "get_engine", return_value=engine),
            patch.object(service, "_ensure_project_storage", return_value=None),
            patch.object(Path, "mkdir", return_value=None),
            patch.object(Path, "open", MagicMock()),
        ):
            payload = service.save_analysis_report_payload("proj-1", {"analysis_report": [{"id": "n1"}]})

        self.assertEqual(payload["message"], "保存成功")
        update_params = conn.execute.call_args_list[2].args[1]
        self.assertEqual(json.loads(update_params["data"])["analysisReport"], [{"id": "n1"}])

    def test_get_analysis_report_prefers_project_data(self) -> None:
        with patch.object(
            service,
            "get_project_payload",
            return_value={
                "id": "proj-1",
                "name": "项目一",
                "status": "uploaded",
                "data": {
                    "analysisReport": [{"id": "n1"}],
                    "analysisV2": {"schema_version": "v2"},
                },
                "created_at": "2026-06-01T00:00:00+00:00",
                "updated_at": "2026-06-01T00:00:00+00:00",
            },
        ):
            payload = service.get_analysis_report_payload("proj-1")

        self.assertEqual(payload["analysis_report"], [{"id": "n1"}])
        self.assertEqual(payload["analysis_v2"], {"schema_version": "v2"})

    def test_get_analysis_report_returns_empty_when_project_and_mirror_missing(self) -> None:
        mirror_path = Mock()
        mirror_path.exists.return_value = False
        with (
            patch.object(service, "get_project_payload", side_effect=service.BidProjectNotFound()),
            patch.object(service, "_analysis_report_mirror_path", return_value=mirror_path),
        ):
            payload = service.get_analysis_report_payload("proj-1")

        self.assertEqual(payload, {"analysis_report": [], "analysis_v2": {}})

    def test_get_project_doc_blocks_reads_persisted_snapshot(self) -> None:
        blocks = [
            {"block_id": "B000001", "locator": "P0001", "body_idx": 0, "text": "第一段", "type": "paragraph"},
            {"block_id": "B000002", "locator": "P0002", "body_idx": 1, "text": "第二段", "type": "paragraph"},
        ]
        with patch.object(
            service,
            "get_project_payload",
            return_value={
                "id": "proj-1",
                "name": "项目一",
                "status": "uploaded",
                "data": {"__doc_blocks_cache": blocks},
                "created_at": "2026-06-01T00:00:00+00:00",
                "updated_at": "2026-06-01T00:00:00+00:00",
            },
        ):
            payload = service.get_project_doc_blocks_payload("proj-1")

        self.assertEqual(payload["project_id"], "proj-1")
        self.assertEqual(payload["blocks"], blocks)
        self.assertEqual(payload["total_blocks"], 2)
        self.assertTrue(payload["snapshot_only"])

    def test_get_project_doc_blocks_returns_404_when_snapshot_missing(self) -> None:
        with patch.object(
            service,
            "get_project_payload",
            return_value={
                "id": "proj-1",
                "name": "项目一",
                "status": "uploaded",
                "data": {},
                "created_at": "2026-06-01T00:00:00+00:00",
                "updated_at": "2026-06-01T00:00:00+00:00",
            },
        ):
            with self.assertRaises(Exception) as ctx:
                service.get_project_doc_blocks_payload("proj-1")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 404)

    def test_get_project_doc_blocks_rejects_unsafe_project_id(self) -> None:
        with self.assertRaises(Exception) as ctx:
            service.get_project_doc_blocks_payload("../proj")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_build_scoring_table_payload_prefers_scoring_template(self) -> None:
        payload = _run_async(
            service.build_scoring_table_payload(
                {
                    "scoring_table_template": [
                        {
                            "id": "row-1",
                            "indicator": "技术方案",
                            "max_score": 15,
                            "criteria": "完整响应",
                        }
                    ],
                    "score_requirements": [
                        {"id": "fallback-1", "content": "不应使用", "points": 5},
                    ],
                }
            )
        )

        self.assertEqual(
            payload,
            {
                "rows": [
                    {
                        "id": "row-1",
                        "indicator": "技术方案",
                        "max_score": 15,
                        "criteria": "完整响应",
                    }
                ]
            },
        )

    def test_build_scoring_table_payload_falls_back_to_score_requirements(self) -> None:
        payload = _run_async(
            service.build_scoring_table_payload(
                {
                    "score_requirements": [
                        {"id": "req-1", "content": "满足安全要求", "points": 8},
                        {"content": "具备实施案例"},
                    ]
                }
            )
        )

        self.assertEqual(
            payload,
            {
                "rows": [
                    {
                        "id": "req-1",
                        "indicator": "满足安全要求",
                        "max_score": 8,
                        "criteria": "",
                    },
                    {
                        "id": "score_req_1",
                        "indicator": "具备实施案例",
                        "max_score": 10,
                        "criteria": "",
                    },
                ]
            },
        )

    def test_generate_blueprint_payload_returns_default_when_key_missing(self) -> None:
        with patch.object(service, "_get_workflow_key", return_value=""):
            payload = _run_async(
                service.generate_blueprint_payload(
                    {
                        "bid_type": "tech",
                        "project_summary": "项目概况",
                        "requirements": [],
                        "outline": [],
                    }
                )
            )

        self.assertEqual(payload["blueprint"]["writing_style"], "正式、专业、数据驱动")
        self.assertEqual(len(payload["blueprint"]["highlights"]), 3)

    def test_generate_blueprint_payload_prefers_split_outputs(self) -> None:
        with (
            patch.object(service, "_get_workflow_key", return_value="app-blueprint"),
            patch.object(
                service,
                "_call_dify_workflow",
                new=AsyncMock(
                    return_value={
                        "data": {
                            "outputs": {
                                "positioning": "突出行业积累",
                                "strategy": "聚焦高分项",
                                "highlights": '["优势一","优势二"]',
                                "writing_style": "正式、严谨庄重",
                            }
                        }
                    }
                ),
            ),
        ):
            payload = _run_async(
                service.generate_blueprint_payload(
                    {
                        "bid_type": "tech",
                        "project_summary": "项目概况",
                        "requirements": [{"type": "tech", "content": "满足要求"}],
                        "outline": [{"title": "第一章"}],
                    }
                )
            )

        self.assertEqual(payload["blueprint"]["positioning"], "突出行业积累")
        self.assertEqual(payload["blueprint"]["highlights"], ["优势一", "优势二"])

    def test_generate_blueprint_payload_falls_back_to_json_text_output(self) -> None:
        with (
            patch.object(service, "_get_workflow_key", return_value="app-blueprint"),
            patch.object(
                service,
                "_call_dify_workflow",
                new=AsyncMock(
                    return_value={
                        "data": {
                            "outputs": {
                                "text": '{"positioning":"高性价比","strategy":"稳妥应标","highlights":["亮点A"],"writing_style":"正式"}'
                            }
                        }
                    }
                ),
            ),
        ):
            payload = _run_async(
                service.generate_blueprint_payload(
                    {
                        "bid_type": "biz",
                        "project_summary": "项目概况",
                        "requirements": [],
                        "outline": [],
                    }
                )
            )

        self.assertEqual(payload["blueprint"]["positioning"], "高性价比")
        self.assertEqual(payload["blueprint"]["highlights"], ["亮点A"])

    def test_fill_scoring_row_payload_prefers_split_outputs(self) -> None:
        with (
            patch.object(service, "_get_workflow_key", side_effect=["app-scoring", ""]),
            patch.object(
                service,
                "_call_dify_workflow",
                new=AsyncMock(
                    return_value={
                        "data": {
                            "outputs": {
                                "self_response": "full",
                                "self_comment": "完全满足评分项要求。",
                                "evidence_refs": '["资质证书","项目案例"]',
                            }
                        }
                    }
                ),
            ),
        ):
            payload = _run_async(
                service.fill_scoring_row_payload(
                    {
                        "row_id": "row-1",
                        "indicator": "技术能力",
                        "max_score": 10,
                        "criteria": "满足要求",
                    }
                )
            )

        self.assertEqual(payload["self_response"], "full")
        self.assertEqual(payload["evidence_refs"], ["资质证书", "项目案例"])

    def test_fill_scoring_row_payload_falls_back_to_json_text_output(self) -> None:
        with (
            patch.object(service, "_get_workflow_key", side_effect=["", "app-req"]),
            patch.object(
                service,
                "_call_dify_workflow",
                new=AsyncMock(
                    return_value={
                        "data": {
                            "outputs": {
                                "text": '```json {"self_response":"none","self_comment":"部分满足","evidence_refs":["附件A","附件B","附件C","附件D"]} ```'
                            }
                        }
                    }
                ),
            ),
        ):
            payload = _run_async(
                service.fill_scoring_row_payload(
                    {
                        "row_id": "row-2",
                        "indicator": "服务能力",
                        "max_score": 8,
                    }
                )
            )

        self.assertEqual(payload["self_response"], "partial")
        self.assertEqual(payload["evidence_refs"], ["附件A", "附件B", "附件C"])

    def test_fill_scoring_row_payload_raises_when_no_key_configured(self) -> None:
        with patch.object(service, "_get_workflow_key", return_value=""):
            with self.assertRaises(Exception) as ctx:
                _run_async(
                    service.fill_scoring_row_payload(
                        {
                            "row_id": "row-3",
                            "indicator": "实施方案",
                            "max_score": 6,
                        }
                    )
                )

        self.assertEqual(getattr(ctx.exception, "status_code", None), 500)

    def test_generate_attachment_payload_renders_builtin_template_natively(self) -> None:
        payload = _run_async(
            service.generate_attachment_payload(
                {
                    "attachment_type": "application_letter",
                    "org_name": "测试公司",
                    "legal_rep": "张三",
                    "project_name": "智慧园区项目",
                    "recipient": "采购人",
                    "doc_date": "2026年6月2日",
                }
            )
        )

        self.assertEqual(payload["label"], "投标申请书")
        self.assertIn("测试公司", payload["content"])
        self.assertIn("智慧园区项目", payload["content"])

    def test_generate_attachment_payload_uses_dify_for_dynamic_attachment(self) -> None:
        with (
            patch.object(service, "_get_workflow_key", return_value="app-attachment"),
            patch.object(
                service,
                "_call_dify_workflow",
                new=AsyncMock(return_value={"data": {"outputs": {"text": "动态附件正文"}}}),
            ),
        ):
            payload = _run_async(
                service.generate_attachment_payload(
                    {
                        "attachment_type": "custom_attachment",
                        "attachment_name": "补充附件",
                        "attachment_desc": "需要说明实施计划",
                        "project_name": "智慧园区项目",
                        "org_name": "测试公司",
                        "legal_rep": "张三",
                    }
                )
            )

        self.assertEqual(payload["label"], "补充附件")
        self.assertEqual(payload["content"], "动态附件正文")

    def test_generate_attachment_payload_raises_when_dynamic_key_missing(self) -> None:
        with patch.object(service, "_get_workflow_key", return_value=""):
            with self.assertRaises(Exception) as ctx:
                _run_async(
                    service.generate_attachment_payload(
                        {
                            "attachment_type": "custom_attachment",
                            "attachment_name": "补充附件",
                        }
                    )
                )

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_test_locators_payload_reads_persisted_snapshot_natively(self) -> None:
        blocks = [
            {"block_id": "B000001", "locator": "p0001", "body_idx": 0, "text": "第一段正文", "type": "paragraph"},
            {"block_id": "B000002", "locator": "P0002", "body_idx": 1, "text": "第二段正文", "type": "paragraph"},
            {"block_id": "B000003", "locator": "", "body_idx": 2, "text": "无定位符", "type": "paragraph"},
        ]
        with (
            patch.object(
                service,
                "get_project_payload",
                return_value={
                    "id": "proj-1",
                    "name": "项目一",
                    "status": "uploaded",
                    "data": {"__doc_blocks_cache": blocks},
                    "created_at": "2026-06-01T00:00:00+00:00",
                    "updated_at": "2026-06-01T00:00:00+00:00",
                },
            ),
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            payload = asyncio.run(service.test_locators_payload("proj-1"))

        legacy_import.assert_not_called()
        self.assertEqual(payload["project_id"], "proj-1")
        self.assertEqual(payload["total_locators"], 2)
        self.assertTrue(payload["snapshot_only"])
        self.assertEqual(
            payload["preview"],
            [
                {"locator": "P0001", "body_idx": 0, "snippet": "第一段正文"},
                {"locator": "P0002", "body_idx": 1, "snippet": "第二段正文"},
            ],
        )

    def test_get_template_config_reads_config_and_selected_yaml_template(self) -> None:
        root = Path("/tmp/clover-bid-template-test")
        config_path = root / "config.yaml"
        templates_dir = root / "templates" / "structures"
        selected_template = templates_dir / "custom.yaml"
        default_template = templates_dir / "standard.yaml"

        def fake_exists(path: Path) -> bool:
            return path == templates_dir

        def fake_iterdir(path: Path) -> list[Path]:
            self.assertEqual(path, templates_dir)
            return [selected_template, default_template]

        def fake_is_file(path: Path) -> bool:
            return path in {selected_template, default_template}

        def fake_open(path: Path, *args, **kwargs):
            import io

            if path == config_path:
                return io.StringIO("workspace:\n  data_dir: ./data\n")
            if path == selected_template:
                return io.StringIO("id: custom\nblocks:\n  - id: sec_1\n")
            raise FileNotFoundError(path)

        with (
            patch.object(service, "_bid_generator_config_path", return_value=config_path),
            patch.object(service, "_template_structures_dir", return_value=templates_dir),
            patch.object(Path, "exists", fake_exists),
            patch.object(Path, "iterdir", fake_iterdir),
            patch.object(Path, "is_file", fake_is_file),
            patch.object(Path, "open", fake_open),
        ):
            payload = service.get_template_config_payload("custom.yaml")

        self.assertEqual(payload["config_dict"], {"workspace": {"data_dir": "./data"}})
        self.assertEqual(payload["template_dict"], {"id": "custom", "blocks": [{"id": "sec_1"}]})
        self.assertEqual(payload["available_templates"], ["custom.yaml", "standard.yaml"])
        self.assertEqual(payload["current_template"], "custom.yaml")

    def test_get_template_config_uses_first_template_when_name_missing(self) -> None:
        root = Path("/tmp/clover-bid-template-test")
        config_path = root / "config.yaml"
        templates_dir = root / "templates" / "structures"
        selected_template = templates_dir / "a.yaml"

        def fake_open(path: Path, *args, **kwargs):
            import io

            if path == config_path:
                return io.StringIO("{}")
            if path == selected_template:
                return io.StringIO("id: a\n")
            raise FileNotFoundError(path)

        with (
            patch.object(service, "_bid_generator_config_path", return_value=config_path),
            patch.object(service, "_template_structures_dir", return_value=templates_dir),
            patch.object(Path, "exists", return_value=True),
            patch.object(Path, "iterdir", return_value=[selected_template]),
            patch.object(Path, "is_file", return_value=True),
            patch.object(Path, "open", fake_open),
        ):
            payload = service.get_template_config_payload()

        self.assertEqual(payload["current_template"], "a.yaml")
        self.assertEqual(payload["template_dict"], {"id": "a"})

    def test_get_template_config_rejects_unsafe_template_name(self) -> None:
        with self.assertRaises(Exception) as ctx:
            service.get_template_config_payload("../secret.yaml")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_get_template_config_returns_404_for_unknown_template(self) -> None:
        root = Path("/tmp/clover-bid-template-test")
        templates_dir = root / "templates" / "structures"
        known_template = templates_dir / "standard.yaml"
        with (
            patch.object(service, "_template_structures_dir", return_value=templates_dir),
            patch.object(Path, "exists", return_value=True),
            patch.object(Path, "iterdir", return_value=[known_template]),
            patch.object(Path, "is_file", return_value=True),
        ):
            with self.assertRaises(Exception) as ctx:
                service.get_template_config_payload("missing.yaml")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 404)

    def test_update_template_config_writes_yaml_natively(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            templates_dir = root / "templates" / "structures"
            with (
                patch.object(service, "_template_structures_dir", return_value=templates_dir),
                patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
            ):
                payload = _run_async(
                    service.update_template_config_payload(
                        {"template_name": "custom.yaml", "template_dict": {"id": "custom", "blocks": [{"id": "b1"}]}}
                    )
                )

            legacy_import.assert_not_called()
            self.assertEqual(payload, {"status": "success", "message": "Template custom.yaml updated successfully"})
            template_path = templates_dir / "custom.yaml"
            self.assertTrue(template_path.exists())
            self.assertEqual(service._read_yaml_mapping(template_path), {"id": "custom", "blocks": [{"id": "b1"}]})

    def test_delete_template_config_deletes_yaml_natively(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            templates_dir = root / "templates" / "structures"
            template_path = templates_dir / "custom.yaml"
            template_path.parent.mkdir(parents=True)
            template_path.write_text("id: custom\n", encoding="utf-8")
            with (
                patch.object(service, "_template_structures_dir", return_value=templates_dir),
                patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
            ):
                payload = _run_async(service.delete_template_config_payload("custom.yaml"))

            legacy_import.assert_not_called()
            self.assertEqual(payload, {"status": "success", "message": "Template custom.yaml deleted successfully"})
            self.assertFalse(template_path.exists())

    def test_delete_template_config_rejects_standard_template(self) -> None:
        with self.assertRaises(Exception) as ctx:
            _run_async(service.delete_template_config_payload("standard.yaml"))

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_update_global_config_writes_yaml_natively(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            config_path = root / "config.yaml"
            with (
                patch.object(service, "_bid_generator_config_path", return_value=config_path),
                patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
            ):
                payload = _run_async(
                    service.update_global_config_payload({"config_dict": {"workspace": {"data_dir": "./data"}}})
                )

            legacy_import.assert_not_called()
            self.assertEqual(payload, {"status": "success", "message": "Config updated successfully"})
            self.assertEqual(service._read_yaml_mapping(config_path), {"workspace": {"data_dir": "./data"}})

    def test_get_cached_pdf_payload_returns_pdf_metadata(self) -> None:
        pdf_path = Mock()
        pdf_path.exists.return_value = True
        pdf_path.read_bytes.return_value = b"%PDF-test"

        with patch.object(service, "_pdf_cache_path", return_value=pdf_path):
            payload = service.get_cached_pdf_payload("proj-1")

        self.assertEqual(payload.content, b"%PDF-test")
        self.assertEqual(payload.media_type, "application/pdf")
        self.assertEqual(payload.filename, "proj-1.pdf")
        self.assertTrue(payload.inline)

    def test_upload_pdf_payload_rejects_non_pdf_filename(self) -> None:
        with self.assertRaises(Exception) as ctx:
            service.upload_pdf_payload("proj-1", filename="file.docx", content=b"docx")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_upload_pdf_payload_writes_cache_and_returns_legacy_url(self) -> None:
        pdf_path = Mock()
        pdf_path.parent.mkdir.return_value = None
        pdf_path.write_bytes.return_value = None

        with patch.object(service, "_pdf_cache_path", return_value=pdf_path):
            payload = service.upload_pdf_payload("proj-1", filename="招标.pdf", content=b"%PDF")

        pdf_path.write_bytes.assert_called_once_with(b"%PDF")
        self.assertEqual(payload["pdf_url"], "/api/v1/bid-generator/api/projects/pdf/proj-1")
        self.assertEqual(payload["message"], "PDF 已缓存")

    def test_delete_project_caches_removes_native_cache_files_without_legacy_import(self) -> None:
        pdf_path = Mock()
        pdf_path.exists.return_value = True
        raw_path = Mock()
        raw_path.exists.return_value = True
        docx_path = Mock()
        docx_path.exists.return_value = True

        with (
            patch.object(service, "_pdf_cache_path", return_value=pdf_path),
            patch.object(service, "_raw_doc_cache_path", return_value=raw_path),
            patch.object(service, "_docx_cache_path", return_value=docx_path),
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            payload = service.delete_project_caches_payload("proj-1")

        legacy_import.assert_not_called()
        self.assertEqual(payload["project_id"], "proj-1")
        self.assertEqual(payload["cleaned"], ["pdf_cache", "raw_doc_cache", "docx_cache"])
        self.assertEqual(payload["message"], "已清理 3 项资源")
        pdf_path.unlink.assert_called_once_with()
        raw_path.unlink.assert_called_once_with()
        docx_path.unlink.assert_called_once_with()

    def test_delete_project_caches_returns_empty_when_nothing_exists(self) -> None:
        pdf_path = Mock()
        pdf_path.exists.return_value = False
        raw_path = Mock()
        raw_path.exists.return_value = False
        docx_path = Mock()
        docx_path.exists.return_value = False

        with (
            patch.object(service, "_pdf_cache_path", return_value=pdf_path),
            patch.object(service, "_raw_doc_cache_path", return_value=raw_path),
            patch.object(service, "_docx_cache_path", return_value=docx_path),
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            payload = service.delete_project_caches_payload("proj-1")

        legacy_import.assert_not_called()
        self.assertEqual(payload, {"project_id": "proj-1", "cleaned": [], "message": "已清理 0 项资源"})
        pdf_path.unlink.assert_not_called()
        raw_path.unlink.assert_not_called()
        docx_path.unlink.assert_not_called()

    def test_delete_project_caches_rejects_unsafe_project_id(self) -> None:
        with self.assertRaises(Exception) as ctx:
            service.delete_project_caches_payload("../proj")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_get_source_docx_payload_returns_docx_metadata(self) -> None:
        docx_path = Mock()
        docx_path.exists.return_value = True
        docx_path.read_bytes.return_value = b"docx-bytes"

        with patch.object(service, "_docx_cache_path", return_value=docx_path):
            payload = service.get_source_docx_payload("proj-1")

        self.assertEqual(payload.content, b"docx-bytes")
        self.assertEqual(payload.filename, "proj-1.docx")
        self.assertEqual(payload.media_type, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    def test_get_extracted_image_by_hash_reads_registry_path(self) -> None:
        image_path = Mock()
        image_path.exists.return_value = True
        image_path.read_bytes.return_value = b"png"
        image_path.suffix = ".png"
        image_path.name = "img.png"
        conn = Mock()
        exists_result = Mock()
        exists_result.scalar_one.return_value = True
        row_result = Mock()
        row_result.mappings.return_value.first.return_value = {"abs_path": "/tmp/img.png"}
        conn.execute.side_effect = [exists_result, row_result]
        engine = MagicMock()
        engine.begin.return_value.__enter__.return_value = conn

        with (
            patch.object(service, "get_engine", return_value=engine),
            patch.object(service, "Path", return_value=image_path),
        ):
            payload = service.get_extracted_image_by_hash_payload("ABCDEF")

        self.assertEqual(payload.content, b"png")
        self.assertEqual(payload.media_type, "image/png")
        self.assertEqual(payload.filename, "img.png")
        params = conn.execute.call_args_list[1].args[1]
        self.assertEqual(params["image_hash"], "abcdef")

    def test_get_extracted_image_payload_rejects_unsafe_filename(self) -> None:
        with self.assertRaises(Exception) as ctx:
            service.get_extracted_image_payload("../secret.png")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_get_extracted_image_payload_infers_media_type(self) -> None:
        image_path = Mock()
        image_path.exists.return_value = True
        image_path.read_bytes.return_value = b"jpg"
        image_path.suffix = ".jpg"
        image_path.name = "demo.jpg"

        with patch.object(service, "_extracted_image_path", return_value=image_path):
            payload = service.get_extracted_image_payload("demo.jpg")

        self.assertEqual(payload.content, b"jpg")
        self.assertEqual(payload.media_type, "image/jpeg")
        self.assertEqual(payload.cache_control, "public, max-age=86400")

    def test_get_diagram_artifact_svg_reads_project_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            project_dir = root / "__A"
            project_dir.mkdir()
            (project_dir / "abcdef1234567890.svg").write_text("<svg></svg>", encoding="utf-8")

            with patch.object(service, "_diagram_artifact_dir", return_value=root):
                payload = service.get_diagram_artifact_svg_payload("ABCDEF1234567890", project_id="项目 A")

        self.assertEqual(payload.content, b"<svg></svg>")
        self.assertEqual(payload.media_type, "image/svg+xml")
        self.assertEqual(payload.filename, "abcdef1234567890.svg")
        self.assertEqual(payload.cache_control, "public, max-age=86400")

    def test_get_diagram_artifact_svg_falls_back_to_mermaid_preview(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            project_dir = root / "default"
            project_dir.mkdir()
            (project_dir / "abcdef1234567890.mmd").write_text("flowchart TD\nA-->B", encoding="utf-8")

            with patch.object(service, "_diagram_artifact_dir", return_value=root):
                payload = service.get_diagram_artifact_svg_payload("abcdef1234567890")

        self.assertEqual(payload.media_type, "image/svg+xml")
        self.assertIn(b"<svg", payload.content)
        self.assertIn("A--&gt;B".encode("utf-8"), payload.content)
        self.assertEqual(payload.filename, "abcdef1234567890.svg")

    def test_get_mermaid_diagram_artifact_reads_source(self) -> None:
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            project_dir = root / "default"
            project_dir.mkdir()
            (project_dir / "abcdef1234567890.mmd").write_text("flowchart TD\nA-->B", encoding="utf-8")

            with patch.object(service, "_diagram_artifact_dir", return_value=root):
                payload = service.get_mermaid_diagram_artifact_payload("abcdef1234567890")

        self.assertEqual(payload.content, b"flowchart TD\nA-->B")
        self.assertEqual(payload.media_type, "text/plain; charset=utf-8")
        self.assertEqual(payload.filename, "abcdef1234567890.mmd")

    def test_get_diagram_artifact_rejects_unsafe_id(self) -> None:
        with self.assertRaises(Exception) as ctx:
            service.get_diagram_artifact_svg_payload("../bad")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_execute_diagram_for_section_runs_natively_without_legacy_task_routes(self) -> None:
        class FakeTaskManager:
            def __init__(self) -> None:
                self.stages: list[str] = []
                self.dify_task_ids: list[str] = []

            async def reserve_diagram_slot(self, project_id: str, max_diagrams: int) -> bool:
                self.reserved = (project_id, max_diagrams)
                return True

            async def release_diagram_slot(self, project_id: str) -> None:
                raise AssertionError("unexpected release")

            def update_stage(self, task_id: str, stage: str) -> None:
                self.stages.append(stage)

            def set_dify_task_id(self, task_id: str, dify_task_id: str) -> None:
                self.dify_task_ids.append(dify_task_id)

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status="running")

        async def fake_diagram_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"dify_task_id": "dify-diagram-1"}
            yield {"__finished__": True, "outputs": {"svg": "<svg><text>架构</text></svg>"}}

        task_manager = FakeTaskManager()
        with tempfile.TemporaryDirectory() as tmp_dir:
            root = Path(tmp_dir)
            with (
                patch.object(service, "_task_manager", return_value=task_manager),
                patch.object(service, "_call_dify_workflow_stream", new=fake_diagram_stream),
                patch.object(service, "_diagram_artifact_dir", return_value=root),
            ):
                diagrams, reserved, error = _run_async(
                    service._execute_diagram_for_section(
                        task_id="task-1",
                        project_id="proj-1",
                        diagram_key="diagram-key",
                        enable_diagrams=True,
                        need_diagram=True,
                        diagram_brief="生成系统架构图",
                        max_diagrams=1,
                        diagram_type_hint="architecture",
                        section_title="总体架构",
                        writing_hint="说明模块、接口、数据流",
                        raw_keywords="架构, 接口",
                        raw_global_outline="总体架构",
                    )
                )

            self.assertEqual(len(list((root / "proj-1").glob("*.svg"))), 1)

        self.assertFalse(hasattr(service, "_legacy_task_routes_module"))
        self.assertTrue(reserved)
        self.assertIsNone(error)
        self.assertEqual(diagrams[0]["type"], "architecture")
        self.assertIn("svg_url", diagrams[0])
        self.assertEqual(task_manager.dify_task_ids, ["dify-diagram-1"])
        self.assertIn("🎨 图表生成中", task_manager.stages)

    def test_get_task_status_payload_maps_running_legacy_task(self) -> None:
        task = SimpleNamespace(
            task_id="task-1",
            project_id="proj-1",
            status="running",
            stages=["开始", "__text__隐藏", "处理中"],
            current_stage="处理中",
            result={"done": True},
            partial_result={"section": "A"},
            partial_events=[
                {"event_id": 1, "phase": "old"},
                {"event_id": 3, "phase": "new"},
            ],
            partial_event_seq=3,
            error="ignored",
            created_at=1780296000,
            updated_at=1780296060,
        )

        with patch.object(service, "_require_task_owner", return_value=task):
            payload = service.get_task_status_payload("task-1", project_id="proj-1", after_event_id=1)

        self.assertEqual(payload["task_id"], "task-1")
        self.assertEqual(payload["status"], "running")
        self.assertEqual(payload["state"], "running")
        self.assertEqual(payload["progress"], 0)
        self.assertEqual(payload["stages"], ["开始", "处理中"])
        self.assertIsNone(payload["result"])
        self.assertEqual(payload["partial_result"], {"section": "A"})
        self.assertEqual(payload["partial_events"], [{"event_id": 3, "phase": "new"}])
        self.assertEqual(payload["last_partial_event_id"], 3)
        self.assertIsNone(payload["error"])
        self.assertTrue(payload["cancellable"])

    def test_ensure_project_slot_native_uses_task_manager_without_legacy_route(self) -> None:
        class FakeTaskManager:
            def ensure_backend_ready(self) -> None:
                return None

            def get_limits(self) -> dict[str, int]:
                return {
                    "max_global_running": 4,
                    "max_project_running": 1,
                    "max_project_content_running": 2,
                    "max_kb_sync_running": 1,
                }

            async def try_acquire_task_slot(self, *args, **kwargs):
                self.args = args
                self.kwargs = kwargs
                return True, {"reason": "ok"}

        task_manager = FakeTaskManager()
        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            _run_async(service._ensure_project_slot_native("proj-1", "content"))

        legacy_import.assert_not_called()
        self.assertEqual(task_manager.args[:2], ("proj-1", "content"))
        self.assertEqual(task_manager.kwargs["max_project_running"], 2)
        self.assertTrue(task_manager.kwargs["enforce_project_limit"])

    def test_ensure_project_slot_native_maps_limit_error(self) -> None:
        class FakeTaskManager:
            def ensure_backend_ready(self) -> None:
                return None

            def get_limits(self) -> dict[str, int]:
                return {
                    "max_global_running": 4,
                    "max_project_running": 1,
                    "max_project_content_running": 2,
                    "max_kb_sync_running": 1,
                }

            async def try_acquire_task_slot(self, *args, **kwargs):
                _ = args, kwargs
                return False, {
                    "reason": "project_limit",
                    "max_project_running": 2,
                    "running_project": 2,
                }

        with patch.object(service, "_task_manager", return_value=FakeTaskManager()):
            with self.assertRaises(Exception) as ctx:
                _run_async(service._ensure_project_slot_native("proj-1", "content"))

        self.assertEqual(getattr(ctx.exception, "status_code", None), 409)
        self.assertEqual(getattr(ctx.exception, "code", ""), "TASK_LIMIT_REACHED")
        self.assertIn("正在运行 2 个正文任务", str(ctx.exception))

    def test_get_task_status_payload_maps_done_legacy_task(self) -> None:
        task = SimpleNamespace(
            task_id="task-1",
            project_id="proj-1",
            status="done",
            stages=[],
            current_stage="完成",
            result={"done": True},
            partial_result={"ignored": True},
            partial_events=[],
            partial_event_seq=0,
            error=None,
            created_at=1780296000,
            updated_at=1780296060,
        )

        with patch.object(service, "_require_task_owner", return_value=task):
            payload = service.get_task_status_payload("task-1")

        self.assertEqual(payload["state"], "succeeded")
        self.assertEqual(payload["progress"], 100)
        self.assertEqual(payload["result"], {"done": True})
        self.assertIsNone(payload["partial_result"])
        self.assertFalse(payload["cancellable"])

    def test_start_content_task_payload_runs_natively_and_sets_partial_and_result(self) -> None:
        class FakeTaskManager:
            def __init__(self) -> None:
                self.created: list[tuple[str, str, str]] = []
                self.stages: list[str] = []
                self.partial_result = None
                self.result = None
                self.async_task = None
                self.dify_task_ids: list[str] = []
                self.status = "running"

            def create_task(self, task_type: str, project_id: str, workflow_name: str = "") -> str:
                self.created.append((task_type, project_id, workflow_name))
                return "task-content-1"

            def update_stage(self, task_id: str, stage: str) -> None:
                self.stages.append(stage)

            def set_dify_task_id(self, task_id: str, dify_task_id: str) -> None:
                self.dify_task_ids.append(dify_task_id)

            def set_partial_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.partial_result = payload

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.result = payload
                self.status = "done"

            def set_async_task(self, task_id: str, task: object) -> None:
                self.async_task = task

            def set_error(self, task_id: str, message: str) -> None:
                raise AssertionError(f"unexpected set_error: {message}")

            def set_cancelled(self, task_id: str) -> None:
                raise AssertionError("unexpected cancel")

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status=self.status, current_stage="✅ 正文已生成")

        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"dify_task_id": "dify-content-1"}
            yield {"__stage__": "✍️ 正文生成"}
            yield {
                "__finished__": True,
                "outputs": {
                    "text": "原始正文",
                    "feedback": "反馈",
                    "quality_score": "85",
                },
            }

        task_manager = FakeTaskManager()

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_validate_required_bidder_info", return_value=None),
            patch.object(service, "_ensure_project_slot_native", new=AsyncMock(return_value=None)),
            patch.object(service, "_get_workflow_key", side_effect=lambda name: "content-key" if name in {"content_writer", "diagram_generator"} else ""),
            patch.object(service, "_compose_runtime_writing_hint", return_value="合成提示"),
            patch.object(
                service,
                "_merge_bidder_pipt_context",
                return_value=(
                    {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                    "占位符提示",
                    {},
                ),
            ),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
            patch.object(
                service,
                "_finalize_legacy_content_output",
                return_value=("最终正文", [{"placeholder": "@@PIPT:v1:e000001:k11111111@@", "original": "张三"}], None),
            ),
            patch.object(service, "_normalize_referenced_images", return_value=("最终正文", [{"placeholder": "__PRO_IMG_1__"}])),
            patch.object(
                service,
                "_run_inline_content_diagram",
                new=AsyncMock(return_value=("最终正文", 0, None, None)),
            ),
            patch.object(service, "_persist_project_runtime", return_value=None),
            patch.object(service, "_persist_content_result_to_project", return_value=None),
            patch.object(service, "_sync_project_runtime_from_task", return_value=None),
        ):
            payload = _run_async(
                service.start_content_task_payload(
                    {
                        "project_id": "proj-1",
                        "section_id": "sec-1",
                        "section_title": "第一章",
                        "writing_hint": "提示",
                        "expected_words": 1200,
                        "project_summary": "项目概述",
                        "global_outline": "全局大纲",
                        "mapping_table": {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                        "bidder_info": {"companyName": "某公司"},
                    }
                )
            )
            _run_async(_await_task(task_manager.async_task))
        self.assertEqual(payload, {"task_id": "task-content-1", "section_id": "sec-1"})
        self.assertEqual(task_manager.created, [("content", "proj-1", "content_writer")])
        self.assertIn("✍️ 正文生成", task_manager.stages)
        self.assertEqual(task_manager.partial_result["phase"], "text_ready")
        self.assertEqual(task_manager.partial_result["content"], "最终正文")
        self.assertEqual(task_manager.result["done"], True)
        self.assertEqual(task_manager.result["content"], "最终正文")
        self.assertEqual(task_manager.result["quality_score"], 85)

    def test_start_content_rewrite_task_payload_runs_natively_and_sets_result(self) -> None:
        class FakeTaskManager:
            def __init__(self) -> None:
                self.created: list[tuple[str, str, str]] = []
                self.result = None
                self.async_task = None
                self.dify_task_ids: list[str] = []
                self.status = "running"

            def create_task(self, task_type: str, project_id: str, workflow_name: str = "") -> str:
                self.created.append((task_type, project_id, workflow_name))
                return "task-rewrite-1"

            def update_stage(self, task_id: str, stage: str) -> None:
                return None

            def set_dify_task_id(self, task_id: str, dify_task_id: str) -> None:
                self.dify_task_ids.append(dify_task_id)

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.result = payload
                self.status = "done"

            def set_async_task(self, task_id: str, task: object) -> None:
                self.async_task = task

            def set_error(self, task_id: str, message: str) -> None:
                raise AssertionError(f"unexpected set_error: {message}")

            def set_cancelled(self, task_id: str) -> None:
                raise AssertionError("unexpected cancel")

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status=self.status, current_stage="✅ 已完成")

        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"dify_task_id": "dify-rewrite-1"}
            yield {"__stage__": "✏️ 润色修改"}
            yield {
                "__finished__": True,
                "outputs": {
                    "text": "重写后的正文",
                    "feedback": "反馈",
                    "quality_score": "92",
                },
            }

        task_manager = FakeTaskManager()

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_validate_required_bidder_info", return_value=None),
            patch.object(service, "_ensure_project_slot_native", new=AsyncMock(return_value=None)),
            patch.object(service, "_get_workflow_key", return_value="rewrite-key"),
            patch.object(
                service,
                "_merge_bidder_pipt_context",
                return_value=(
                    {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                    "占位符提示",
                    {},
                ),
            ),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
            patch.object(
                service,
                "_finalize_legacy_content_output",
                return_value=("重写后的正文", [{"placeholder": "@@PIPT:v1:e000001:k11111111@@", "original": "张三"}], None),
            ),
            patch.object(service, "_persist_project_runtime", return_value=None),
            patch.object(service, "_persist_content_result_to_project", return_value=None),
            patch.object(service, "_sync_project_runtime_from_task", return_value=None),
        ):
            payload = _run_async(
                service.start_content_rewrite_task_payload(
                    {
                        "project_id": "proj-1",
                        "section_id": "sec-1",
                        "section_title": "第一章",
                        "current_content": "原始正文<diagram data-diagram-id=\"abc\"></diagram>",
                        "rewrite_instruction": "请强化技术细节",
                        "expected_words": 1200,
                        "mapping_table": {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                        "bidder_info": {"companyName": "某公司"},
                    }
                )
            )
            _run_async(_await_task(task_manager.async_task))
        self.assertEqual(payload, {"task_id": "task-rewrite-1", "section_id": "sec-1"})
        self.assertEqual(task_manager.created, [("content", "proj-1", "content_rewrite")])
        self.assertEqual(task_manager.result["done"], True)
        self.assertIn("重写后的正文", task_manager.result["content"])
        self.assertIn("<diagram", task_manager.result["content"])
        self.assertEqual(task_manager.result["quality_score"], 92)

    def test_start_content_group_task_payload_runs_natively_and_emits_group_child_done(self) -> None:
        class FakeTaskManager:
            def __init__(self) -> None:
                self.created: list[tuple[str, str, str]] = []
                self.result = None
                self.async_task = None
                self.partial_events: list[dict[str, object]] = []
                self.stages: list[str] = []
                self.dify_task_ids: list[str] = []
                self.status = "running"

            def create_task(self, task_type: str, project_id: str, workflow_name: str = "") -> str:
                self.created.append((task_type, project_id, workflow_name))
                return "task-group-1"

            def update_stage(self, task_id: str, stage: str) -> None:
                self.stages.append(stage)

            def set_dify_task_id(self, task_id: str, dify_task_id: str) -> None:
                self.dify_task_ids.append(dify_task_id)

            def append_partial_event(self, task_id: str, payload: dict[str, object]) -> None:
                self.partial_events.append(payload)

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.result = payload
                self.status = "done"

            def set_async_task(self, task_id: str, task: object) -> None:
                self.async_task = task

            def set_error(self, task_id: str, message: str) -> None:
                raise AssertionError(f"unexpected set_error: {message}")

            def set_cancelled(self, task_id: str) -> None:
                raise AssertionError("unexpected cancel")

            async def release_diagram_slot(self, project_id: str) -> None:
                return None

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status=self.status, current_stage="✅ 分组正文已生成")

        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"dify_task_id": "dify-group-1"}
            yield {"__stage__": "📦 H2 批量生成中"}
            yield {
                "__finished__": True,
                "outputs": {
                    "sections": [
                        {"section_id": "sec-1", "content": "第一节正文", "quality_score": "88", "feedback": "ok"},
                        {"section_id": "sec-2", "content": "第二节正文", "quality_score": "91", "feedback": "ok"},
                    ]
                },
            }

        task_manager = FakeTaskManager()

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_validate_required_bidder_info", return_value=None),
            patch.object(service, "_ensure_project_slot_native", new=AsyncMock(return_value=None)),
            patch.object(service, "_get_workflow_key", side_effect=lambda name: "group-key" if name == "content_group_writer" else ""),
            patch.object(
                service,
                "_merge_bidder_pipt_context",
                return_value=(
                    {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                    "占位符提示",
                    {},
                ),
            ),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
            patch.object(service, "_persist_project_runtime", return_value=None),
            patch.object(service, "_persist_group_content_result_to_project", return_value=None),
            patch.object(service, "_sync_project_runtime_from_task", return_value=None),
        ):
            payload = _run_async(
                service.start_content_group_task_payload(
                    {
                        "project_id": "proj-1",
                        "group_id": "group-a",
                        "group_title": "第一组",
                        "project_summary": "项目概述",
                        "global_outline": "全局大纲",
                        "mapping_table": {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                        "bidder_info": {"companyName": "某公司"},
                        "children": [
                            {
                                "section_id": "sec-1",
                                "section_title": "第一节",
                                "keywords": "关键词1",
                                "expected_words": 500,
                                "writing_hint": "提示1",
                                "analysis_context": "分析1",
                            },
                            {
                                "section_id": "sec-2",
                                "section_title": "第二节",
                                "keywords": "关键词2",
                                "expected_words": 600,
                                "writing_hint": "提示2",
                                "analysis_context": "分析2",
                            },
                        ],
                    }
                )
            )
            _run_async(_await_task(task_manager.async_task))
        self.assertEqual(payload, {"task_id": "task-group-1", "group_id": "group-a"})
        self.assertEqual(task_manager.created, [("content", "proj-1", "content_group_writer")])
        self.assertEqual(len(task_manager.partial_events), 2)
        self.assertTrue(all(event["phase"] == "group_child_done" for event in task_manager.partial_events))
        self.assertEqual(task_manager.result["group_id"], "group-a")
        self.assertEqual(task_manager.result["group_title"], "第一组")
        self.assertEqual(len(task_manager.result["sections"]), 2)
        self.assertEqual(task_manager.result["failed_sections"], [])

    def test_start_group_review_task_payload_runs_natively_and_sets_result(self) -> None:
        class FakeTaskManager:
            def __init__(self) -> None:
                self.created: list[tuple[str, str, str]] = []
                self.result = None
                self.async_task = None
                self.stages: list[str] = []
                self.dify_task_ids: list[str] = []
                self.status = "running"

            def create_task(self, task_type: str, project_id: str, workflow_name: str = "") -> str:
                self.created.append((task_type, project_id, workflow_name))
                return "task-review-1"

            def update_stage(self, task_id: str, stage: str) -> None:
                self.stages.append(stage)

            def set_dify_task_id(self, task_id: str, dify_task_id: str) -> None:
                self.dify_task_ids.append(dify_task_id)

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.result = payload
                self.status = "done"

            def set_async_task(self, task_id: str, task: object) -> None:
                self.async_task = task

            def set_error(self, task_id: str, message: str) -> None:
                raise AssertionError(f"unexpected set_error: {message}")

            def set_cancelled(self, task_id: str) -> None:
                raise AssertionError("unexpected cancel")

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status=self.status, current_stage="✅ 分组评估已完成")

        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"dify_task_id": "dify-review-1"}
            yield {"__stage__": "🧾 H2 章节评估中"}
            yield {
                "__finished__": True,
                "outputs": {
                    "result": {
                        "group_feedback": "整体内容较完整，建议补强实施细节。",
                        "quality_score": "87",
                    }
                },
            }

        task_manager = FakeTaskManager()

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_ensure_project_slot_native", new=AsyncMock(return_value=None)),
            patch.object(service, "_get_workflow_key", side_effect=lambda name: "review-key" if name == "group_review_writer" else ""),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
            patch.object(service, "_persist_project_runtime", return_value=None),
            patch.object(service, "_sync_project_runtime_from_task", return_value=None),
        ):
            payload = _run_async(
                service.start_group_review_task_payload(
                    {
                        "project_id": "proj-1",
                        "group_id": "group-a",
                        "group_title": "第一组",
                        "project_summary": "项目概述",
                        "group_outline": "分组大纲",
                        "group_analysis_context": "分组分析上下文",
                        "sections": [
                            {"section_id": "sec-1", "section_title": "第一节", "content": "第一节正文"},
                            {"section_id": "sec-2", "section_title": "第二节", "content": "第二节正文"},
                        ],
                    }
                )
            )
            _run_async(_await_task(task_manager.async_task))
        self.assertEqual(payload, {"task_id": "task-review-1", "group_id": "group-a"})
        self.assertEqual(task_manager.created, [("content", "proj-1", "group_review_writer")])
        self.assertIn("🧾 H2 章节评估中", task_manager.stages)
        self.assertEqual(task_manager.result["done"], True)
        self.assertEqual(task_manager.result["group_id"], "group-a")
        self.assertEqual(task_manager.result["group_title"], "第一组")
        self.assertEqual(task_manager.result["group_feedback"], "整体内容较完整，建议补强实施细节。")
        self.assertEqual(task_manager.result["quality_score"], 87)

    def test_start_diagram_task_payload_runs_natively_and_sets_result(self) -> None:
        class FakeTaskManager:
            def __init__(self) -> None:
                self.created: list[tuple[str, str, str]] = []
                self.result = None
                self.async_task = None
                self.stages: list[str] = []
                self.status = "running"

            def create_task(self, task_type: str, project_id: str, workflow_name: str = "") -> str:
                self.created.append((task_type, project_id, workflow_name))
                return "task-diagram-1"

            def update_stage(self, task_id: str, stage: str) -> None:
                self.stages.append(stage)

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.result = payload
                self.status = "done"

            def set_async_task(self, task_id: str, task: object) -> None:
                self.async_task = task

            def set_error(self, task_id: str, message: str) -> None:
                raise AssertionError(f"unexpected set_error: {message}")

            def set_cancelled(self, task_id: str) -> None:
                raise AssertionError("unexpected cancel")

            async def release_diagram_slot(self, project_id: str) -> None:
                return None

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status=self.status, current_stage="✅ 图表已完成")

        task_manager = FakeTaskManager()

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_ensure_project_slot_native", new=AsyncMock(return_value=None)),
            patch.object(service, "_get_workflow_key", side_effect=lambda name: "diagram-key" if name == service._get_diagram_workflow_name() else ""),
            patch.object(
                service,
                "_run_diagram_request",
                new=AsyncMock(
                    return_value={
                        "done": True,
                        "section_id": "sec-1",
                        "content": "正文<diagram data-diagram-id=\"abc\"></diagram>",
                        "word_count": 10,
                        "quality_score": 89,
                        "feedback": "ok",
                        "replace_report": [],
                        "diagrams_count": 1,
                    }
                ),
            ),
            patch.object(service, "_persist_project_runtime", return_value=None),
            patch.object(service, "_sync_project_runtime_from_task", return_value=None),
        ):
            payload = _run_async(
                service.start_diagram_task_payload(
                    {
                        "project_id": "proj-1",
                        "section_id": "sec-1",
                        "section_title": "第一节",
                        "base_content": "正文",
                        "enable_diagrams": True,
                        "need_diagram": True,
                        "diagram_brief": "补一张架构图",
                        "max_diagrams": 1,
                    }
                )
            )
            _run_async(_await_task(task_manager.async_task))
        self.assertEqual(payload, {"task_id": "task-diagram-1", "section_id": "sec-1"})
        self.assertEqual(task_manager.created, [("diagram", "proj-1", service._get_diagram_workflow_name())])
        self.assertIn("🎨 独立图表任务启动", task_manager.stages)
        self.assertEqual(task_manager.result["section_id"], "sec-1")
        self.assertEqual(task_manager.result["diagrams_count"], 1)

    def test_start_diagram_batch_task_payload_runs_natively_and_emits_partial_events(self) -> None:
        class FakeTaskManager:
            def __init__(self) -> None:
                self.created: list[tuple[str, str, str]] = []
                self.result = None
                self.async_task = None
                self.partial_events: list[dict[str, object]] = []
                self.stages: list[str] = []
                self.status = "running"

            def create_task(self, task_type: str, project_id: str, workflow_name: str = "") -> str:
                self.created.append((task_type, project_id, workflow_name))
                return "task-diagram-batch-1"

            def update_stage(self, task_id: str, stage: str) -> None:
                self.stages.append(stage)

            def append_partial_event(self, task_id: str, payload: dict[str, object]) -> None:
                self.partial_events.append(payload)

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.result = payload
                self.status = "done"

            def set_async_task(self, task_id: str, task: object) -> None:
                self.async_task = task

            def set_error(self, task_id: str, message: str) -> None:
                raise AssertionError(f"unexpected set_error: {message}")

            def set_cancelled(self, task_id: str) -> None:
                raise AssertionError("unexpected cancel")

            async def release_diagram_slot(self, project_id: str) -> None:
                return None

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status=self.status, current_stage="✅ 批量图表已完成")

        results = [
            {
                "done": True,
                "section_id": "sec-1",
                "content": "正文1<diagram data-diagram-id=\"a\"></diagram>",
                "word_count": 10,
                "quality_score": 80,
                "feedback": "ok1",
                "replace_report": [],
                "diagrams_count": 1,
            },
            {
                "done": True,
                "section_id": "sec-2",
                "content": "正文2",
                "word_count": 8,
                "quality_score": 78,
                "feedback": "ok2",
                "replace_report": [],
                "diagrams_count": 0,
                "diagram_error": {"message": "图表失败"},
            },
        ]
        task_manager = FakeTaskManager()

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_ensure_project_slot_native", new=AsyncMock(return_value=None)),
            patch.object(service, "_get_workflow_key", side_effect=lambda name: "diagram-key" if name == service._get_diagram_workflow_name() else ""),
            patch.object(service, "_run_diagram_request", new=AsyncMock(side_effect=results)),
            patch.object(service, "_persist_project_runtime", return_value=None),
            patch.object(service, "_sync_project_runtime_from_task", return_value=None),
        ):
            payload = _run_async(
                service.start_diagram_batch_task_payload(
                    {
                        "project_id": "proj-1",
                        "diagram_requests": [
                            {"project_id": "proj-1", "section_id": "sec-1", "section_title": "第一节", "base_content": "正文1"},
                            {"project_id": "proj-1", "section_id": "sec-2", "section_title": "第二节", "base_content": "正文2"},
                        ],
                        "enable_diagrams": True,
                    }
                )
            )
            _run_async(_await_task(task_manager.async_task))
        self.assertEqual(payload, {"task_id": "task-diagram-batch-1", "count": 2})
        self.assertEqual(task_manager.created, [("diagram", "proj-1", service._get_diagram_workflow_name())])
        self.assertEqual(len(task_manager.partial_events), 2)
        self.assertTrue(all(event["phase"] == "diagram_section_done" for event in task_manager.partial_events))
        self.assertEqual(len(task_manager.result["sections"]), 2)
        self.assertEqual(len(task_manager.result["failed_sections"]), 1)
        self.assertEqual(task_manager.result["diagrams_count"], 1)

    def test_require_task_owner_rejects_project_mismatch(self) -> None:
        task_manager = Mock()
        task_manager.get_task.return_value = SimpleNamespace(task_id="task-1", project_id="proj-1")

        with patch.object(service, "_task_manager", return_value=task_manager):
            with self.assertRaises(Exception) as ctx:
                service._require_task_owner("task-1", "other")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 403)

    def test_require_task_owner_returns_404_when_task_missing(self) -> None:
        task_manager = Mock()
        task_manager.get_task.return_value = None

        with patch.object(service, "_task_manager", return_value=task_manager):
            with self.assertRaises(Exception) as ctx:
                service._require_task_owner("task-1", None)

        self.assertEqual(getattr(ctx.exception, "status_code", None), 404)

    def test_cancel_task_payload_uses_native_task_manager_cancel(self) -> None:
        task = SimpleNamespace(
            task_id="task-1",
            project_id="proj-1",
            task_type="extract",
            workflow_name="",
            status="running",
            current_stage="解析中",
            created_at=1780296000,
        )
        task_manager = Mock()
        task_manager.cancel_task.return_value = True
        task_manager.get_task.return_value = SimpleNamespace(
            task_id="task-1",
            project_id="proj-1",
            task_type="extract",
            status="cancelled",
            current_stage="",
            created_at=1780296000,
        )

        with (
            patch.object(service, "_require_task_owner", return_value=task),
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_stop_dify_workflows_for_task", return_value=(False, "not_applicable")),
            patch.object(service, "_persist_project_task_runtime", return_value=None) as persist_runtime,
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            payload = _run_async(service.cancel_task_payload("task-1", project_id="proj-1"))

        legacy_import.assert_not_called()
        task_manager.cancel_task.assert_called_once_with("task-1")
        self.assertEqual(persist_runtime.call_count, 2)
        self.assertEqual(payload["cancelled"], True)
        self.assertEqual(payload["task_id"], "task-1")
        self.assertEqual(payload["remote_stop_status"], "not_applicable")
        self.assertEqual(payload["task_state"], "cancelled")

    def test_cancel_task_payload_returns_404_when_task_already_finished(self) -> None:
        task = SimpleNamespace(task_id="task-1", project_id="proj-1", task_type="outline", status="done", current_stage="")
        with patch.object(service, "_require_task_owner", return_value=task):
            with self.assertRaises(Exception) as ctx:
                _run_async(service.cancel_task_payload("task-1"))

        self.assertEqual(getattr(ctx.exception, "status_code", None), 404)
        self.assertEqual(getattr(ctx.exception, "code", ""), "RESOURCE_NOT_FOUND")

    def test_stop_dify_workflows_for_task_uses_native_http_client(self) -> None:
        calls: list[tuple[str, dict[str, str], dict[str, str]]] = []

        class FakeResponse:
            status_code = 200

        class FakeClient:
            async def post(self, url: str, *, headers: dict[str, str], json: dict[str, str]) -> FakeResponse:
                calls.append((url, headers, json))
                return FakeResponse()

        @asynccontextmanager
        async def fake_async_client(*, timeout: int):
            self.assertEqual(timeout, 10)
            yield FakeClient()

        task = SimpleNamespace(
            task_type="outline",
            workflow_name="",
            dify_task_id="dify-1",
            dify_task_ids=["dify-1", "dify-2"],
        )
        with (
            patch.dict(
                service.os.environ,
                {
                    "DIFY_API_URL": "http://dify.local/v1",
                    "DIFY_WORKFLOW_STRUCTURE_GENERATOR": "workflow-key",
                },
                clear=False,
            ),
            patch.object(service.httpx, "AsyncClient", fake_async_client),
        ):
            stopped, status = _run_async(service._stop_dify_workflows_for_task(task))

        self.assertTrue(stopped)
        self.assertEqual(status, "stopped")
        self.assertEqual(
            [call[0] for call in calls],
            [
                "http://dify.local/v1/workflows/tasks/dify-1/stop",
                "http://dify.local/v1/workflows/tasks/dify-2/stop",
            ],
        )
        self.assertEqual(calls[0][1], {"Authorization": "Bearer workflow-key"})
        self.assertEqual(calls[0][2], {"user": "pro-engine-backend"})

    def test_start_outline_task_payload_runs_natively_and_sets_partial_and_result(self) -> None:
        class FakeTaskManager:
            def __init__(self) -> None:
                self.created: list[tuple[str, str, str]] = []
                self.partial_result = None
                self.result = None
                self.async_task = None
                self.stages: list[str] = []
                self.status = "running"
                self.dify_task_ids: list[str] = []

            def create_task(self, task_type: str, project_id: str, workflow_name: str = "") -> str:
                self.created.append((task_type, project_id, workflow_name))
                return "task-outline-1"

            def update_stage(self, task_id: str, stage: str) -> None:
                self.stages.append(stage)

            def set_dify_task_id(self, task_id: str, dify_task_id: str) -> None:
                self.dify_task_ids.append(dify_task_id)

            def set_partial_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.partial_result = payload

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.result = payload
                self.status = "done"

            def set_async_task(self, task_id: str, task: object) -> None:
                self.async_task = task

            def set_error(self, task_id: str, message: str) -> None:
                raise AssertionError(f"unexpected set_error: {message}")

            def set_cancelled(self, task_id: str) -> None:
                raise AssertionError("unexpected cancel")

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status=self.status, current_stage="✅ 大纲结构已就绪")

        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"dify_task_id": "dify-outline-1"}
            yield {"__stage__": "✍️ 生成大纲", "workflow_run_id": "run-outline-1"}
            yield {
                "__finished__": True,
                "outputs": {
                    "structured_output": {
                        "outline": [
                            {
                                "title": "总体技术方案",
                                "children": [
                                    {
                                        "title": "建设目标",
                                        "wordCount": 500,
                                        "writingHint": "围绕建设目标展开",
                                        "keywords": ["建设目标"],
                                    }
                                ],
                            }
                        ]
                    }
                },
                "workflow_run_id": "run-outline-1",
            }

        task_manager = FakeTaskManager()

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_ensure_project_slot_native", new=AsyncMock(return_value=None)),
            patch.object(service, "_get_workflow_key", side_effect=lambda name: "outline-key" if name == "structure_generator" else ""),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
            patch.object(service, "_persist_project_runtime", return_value=None),
            patch.object(service, "_sync_project_runtime_from_task", return_value=None),
        ):
            payload = _run_async(
                service.start_outline_task_payload(
                    {
                        "project_id": "proj-1",
                        "bid_type": "tech",
                        "requirements": [{"type": "tech", "content": "满足建设要求"}],
                        "expected_total_words": 600,
                        "structure_heading_seed_json": '[{"id":"h2-1","title":"总体技术方案","generation_strategy":"general"}]',
                    }
                )
            )
            _run_async(_await_task(task_manager.async_task))
        self.assertEqual(payload, {"task_id": "task-outline-1"})
        self.assertEqual(task_manager.created, [("outline", "proj-1", "structure_generator")])
        self.assertEqual(task_manager.partial_result["phase"], "h3_meta_generating")
        self.assertEqual(task_manager.result["done"], True)
        self.assertEqual(task_manager.result["phase"], "outline_finalized")
        self.assertEqual(task_manager.result["sections"][0]["title"], "总体技术方案")
        self.assertEqual(task_manager.result["sections"][0]["wordCount"], 600)

    def test_start_extract_task_payload_delegates_multipart_arguments(self) -> None:
        upload = SimpleNamespace(filename="demo.docx")
        upload.read = AsyncMock(return_value=b"docx-bytes")

        class FakeTaskManager:
            def __init__(self) -> None:
                self.created: list[tuple[str, str]] = []
                self.stages: list[str] = []
                self.result = None
                self.async_task = None

            def create_task(self, task_type: str, project_id: str) -> str:
                self.created.append((task_type, project_id))
                return "task-extract-1"

            def update_stage(self, task_id: str, stage: str) -> None:
                self.stages.append(stage)

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.result = payload

            def set_async_task(self, task_id: str, task: object) -> None:
                self.async_task = task

            def set_error(self, task_id: str, message: str) -> None:
                raise AssertionError(f"unexpected set_error: {message}")

            def set_cancelled(self, task_id: str) -> None:
                raise AssertionError("unexpected cancel")

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status="done", current_stage="预处理完成")

        task_manager = FakeTaskManager()

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_ensure_project_slot_native", new=AsyncMock(return_value=None)),
            patch.object(service, "_extract_raw_text_with_images_native", return_value=("原文内容", {"img1": {"name": "示意图"}})),
            patch.object(service, "_extract_docx_with_locators_native", return_value=("定位原文", {"P1": 1}, [{"block_id": "B1"}])),
            patch.object(service, "_persist_project_doc_blocks_snapshot", return_value=None),
            patch.object(service, "_persist_docx_cache", return_value=None),
            patch.object(service, "_persist_project_runtime", return_value=None),
            patch.object(service, "_sync_project_runtime_from_task", return_value=None),
            patch.object(service, "_persist_raw_document", return_value=None),
            patch.object(service, "_convert_to_pdf_and_cache_native", return_value="/api/projects/pdf/proj-1"),
            patch.object(service, "_run_bid_pipt_preprocess", return_value={"text": "定位原文", "mapping_table": {}, "mapping_table_count": 0, "placeholder_manifest": {}, "placeholder_policy": {}}),
        ):
            payload = _run_async(
                service.start_extract_task_payload(
                    upload,
                    project_name="项目一",
                    project_id="proj-1",
                    enable_desensitize=False,
                )
            )
            _run_async(_await_task(task_manager.async_task))
        self.assertEqual(payload["task_id"], "task-extract-1")
        self.assertEqual(task_manager.created, [("extract", "proj-1")])
        self.assertIn("解析文档结构", task_manager.stages)
        self.assertIn("跳过脱敏", task_manager.stages)
        self.assertEqual(task_manager.result["raw_document"], "定位原文")
        self.assertEqual(task_manager.result["pdf_url"], "/api/projects/pdf/proj-1")
        self.assertEqual(task_manager.result["image_map"], {"img1": {"name": "示意图"}})

    def test_start_extract_task_payload_uses_unified_pipt_global_redaction(self) -> None:
        upload = SimpleNamespace(filename="demo.docx")
        upload.read = AsyncMock(return_value=b"docx-bytes")

        class FakeTaskManager:
            def __init__(self) -> None:
                self.result = None
                self.async_task = None

            def create_task(self, task_type: str, project_id: str) -> str:
                return "task-extract-2"

            def update_stage(self, task_id: str, stage: str) -> None:
                return None

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.result = payload

            def set_async_task(self, task_id: str, task: object) -> None:
                self.async_task = task

            def set_error(self, task_id: str, message: str) -> None:
                raise AssertionError(f"unexpected set_error: {message}")

            def set_cancelled(self, task_id: str) -> None:
                raise AssertionError("unexpected cancel")

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status="done", current_stage="预处理完成")

        task_manager = FakeTaskManager()

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_ensure_project_slot_native", new=AsyncMock(return_value=None)),
            patch.object(service, "_extract_raw_text_with_images_native", return_value=("第一段 张三 未替换。第二段 张三 已识别。", {})),
            patch.object(service, "_extract_docx_with_locators_native", return_value=("第一段 张三 未替换。第二段 张三 已识别。", {}, [])),
            patch.object(service, "_persist_project_doc_blocks_snapshot", return_value=None),
            patch.object(service, "_persist_docx_cache", return_value=None),
            patch.object(service, "_persist_project_runtime", return_value=None),
            patch.object(service, "_sync_project_runtime_from_task", return_value=None),
            patch.object(service, "_persist_raw_document", return_value=None),
            patch.object(service, "_convert_to_pdf_and_cache_native", return_value=""),
            patch.object(
                service,
                "_run_bid_pipt_preprocess",
                return_value={
                    "text": "第一段 @@PIPT:v1:e000001:k11111111@@ 未替换。第二段 @@PIPT:v1:e000001:k11111111@@ 已识别。",
                    "mapping_table": {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                    "mapping_table_count": 1,
                    "placeholder_manifest": {"@@PIPT:v1:e000001:k11111111@@": {"entity_type": "name"}},
                    "placeholder_policy": {},
                },
            ),
        ):
            payload = _run_async(
                service.start_extract_task_payload(
                    upload,
                    project_name="项目一",
                    project_id="proj-1",
                    enable_desensitize=True,
                )
            )
            _run_async(_await_task(task_manager.async_task))

        self.assertEqual(payload["task_id"], "task-extract-2")
        self.assertNotIn("张三", task_manager.result["raw_document"])
        self.assertEqual(task_manager.result["entity_count"], 1)
        self.assertEqual(
            task_manager.result["mapping_table"],
            {"@@PIPT:v1:e000001:k11111111@@": "张三"},
        )

    def test_start_analyze_task_payload_runs_natively_and_sets_result(self) -> None:
        class FakeTaskManager:
            def __init__(self) -> None:
                self.created: list[tuple[str, str]] = []
                self.stages: list[str] = []
                self.result = None
                self.async_task = None
                self.dify_task_ids: list[str] = []

            def create_task(self, task_type: str, project_id: str) -> str:
                self.created.append((task_type, project_id))
                return "task-analyze-1"

            def update_stage(self, task_id: str, stage: str) -> None:
                self.stages.append(stage)

            def set_dify_task_id(self, task_id: str, dify_task_id: str) -> None:
                self.dify_task_ids.append(dify_task_id)

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                self.result = payload

            def set_async_task(self, task_id: str, task: object) -> None:
                self.async_task = task

            def set_error(self, task_id: str, message: str) -> None:
                raise AssertionError(f"unexpected set_error: {message}")

            def set_cancelled(self, task_id: str) -> None:
                raise AssertionError("unexpected cancel")

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status="done", current_stage="商务与技术结构已生成")

        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"dify_task_id": "dify-task-1"}
            yield {
                "__finished__": True,
                "outputs": {
                    "text": json.dumps(
                        {
                            "proj_overview": "项目概述内容",
                            "proj_basic": "<项目名称>项目一</项目名称>",
                            "scoring_details": json.dumps(
                                {
                                    "total": 10,
                                    "items": [{"id": "score-1", "name": "技术方案", "max_score": 10, "criteria": "完全响应得10分"}],
                                },
                                ensure_ascii=False,
                            ),
                            "structure_attachments": "<要点>投标函</要点>",
                            "resp_tech": "技术目标",
                            "resp_param": "参数要求",
                            "resp_substance": "实施要求",
                        },
                        ensure_ascii=False,
                    )
                    + '<BID_ATTACHMENTS>[{"name":"投标函","start_locator":"P1","end_locator":"P2","description":""}]</BID_ATTACHMENTS>'
                },
            }

        task_manager = FakeTaskManager()

        async def fake_sleep(_: float) -> None:
            return None

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_get_workflow_key", return_value="doc-key"),
            patch.object(service, "load_docanalysis_framework", return_value=("系统提示", _analysis_framework_nodes())),
            patch.object(service, "_bid_generator_root", return_value=Path("/tmp")),
            patch.object(service.Path, "exists", return_value=True),
            patch.object(service, "_persist_raw_document", return_value=None),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
            patch.object(service, "get_project_payload", return_value={"id": "proj-1", "data": {"__doc_blocks_cache": _doc_blocks_fixture()}}),
            patch.object(service, "_persist_project_runtime", return_value=None),
            patch.object(service, "_sync_project_runtime_from_task", return_value=None),
            patch.object(service, "_ensure_task_running", return_value=None),
            patch.object(service, "_persist_project_analysis_report", return_value=None),
            patch.object(service, "patch_project_payload", return_value={"id": "proj-1"}),
            patch.object(service.asyncio, "sleep", new=fake_sleep),
        ):
            payload = _run_async(service.start_analyze_task_payload(project_id="proj-1", raw_document="原文内容"))
            _run_async(_await_task(task_manager.async_task))

        self.assertEqual(payload, {"task_id": "task-analyze-1"})
        self.assertEqual(task_manager.created, [("analyze", "proj-1")])
        self.assertIn("dify-task-1", task_manager.dify_task_ids)
        self.assertTrue(any(stage.startswith(service._BID_ATTACH_STAGE_PREFIX) for stage in task_manager.stages))
        self.assertTrue(any(stage.startswith(service._ANALYSIS_V2_STAGE_PREFIX) for stage in task_manager.stages))
        self.assertEqual(task_manager.result["done"], True)
        self.assertEqual(task_manager.result["success_count"], 7)
        self.assertEqual(task_manager.result["analysis_v2"]["schema_version"], 3)

    def test_start_analyze_task_payload_uses_cached_document_when_raw_missing(self) -> None:
        class FakeTaskManager:
            def create_task(self, task_type: str, project_id: str) -> str:
                return "task-analyze-2"

            def update_stage(self, task_id: str, stage: str) -> None:
                return None

            def set_dify_task_id(self, task_id: str, dify_task_id: str) -> None:
                return None

            def set_result(self, task_id: str, payload: dict[str, object]) -> None:
                return None

            def set_async_task(self, task_id: str, task: object) -> None:
                return None

            def get_task(self, task_id: str) -> SimpleNamespace:
                return SimpleNamespace(status="running", current_stage="")

        task_manager = FakeTaskManager()

        with (
            patch.object(service, "_task_manager", return_value=task_manager),
            patch.object(service, "_get_workflow_key", return_value="doc-key"),
            patch.object(service, "load_docanalysis_framework", return_value=("系统提示", _analysis_framework_nodes())),
            patch.object(service, "_bid_generator_root", return_value=Path("/tmp")),
            patch.object(service.Path, "exists", return_value=True),
            patch.object(service, "_load_raw_document", return_value="缓存原文"),
            patch.object(service, "_persist_project_runtime", return_value=None),
        ):
            payload = _run_async(service.start_analyze_task_payload(project_id="proj-1"))

        self.assertEqual(payload, {"task_id": "task-analyze-2"})

    def test_extract_requirements_payload_runs_natively_without_legacy_route(self) -> None:
        upload = SimpleNamespace(filename="demo.pdf")
        upload.read = AsyncMock(return_value=b"pdf-bytes")

        with (
            patch.object(
                service,
                "_prepare_extract_document",
                return_value={
                    "cache_id": "proj-1",
                    "pdf_url": "/api/projects/pdf/proj-1",
                    "pages_text": [{"page": 0, "text": "原文片段需求内容"}],
                    "raw_image_map": {"img1": {"name": "图1"}},
                    "text_for_dify": "脱敏后原文",
                    "mapping_table": {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                    "entity_count": 1,
                    "placeholder_manifest": {"@@PIPT:v1:e000001:k11111111@@": {"entity_type": "name"}},
                    "placeholder_policy": {},
                },
            ),
            patch.object(service, "_get_workflow_key", return_value="req-key"),
            patch.object(
                service,
                "_call_dify_workflow",
                new=AsyncMock(
                    return_value={
                        "data": {
                            "outputs": {
                                "text": json.dumps(
                                    {
                                        "bid_type": "tech",
                                        "project_summary": "项目概述",
                                        "requirements": [
                                            {
                                                "type": "tech",
                                                "content": "需求内容",
                                                "points": 5,
                                                "source_excerpt": "原文片段需求内容",
                                            }
                                        ],
                                        "analysis_report": [{"id": "n1"}],
                                        "required_attachments": [{"name": "投标函"}],
                                        "scoring_table_template": [{"indicator": "技术方案"}],
                                        "expected_word_count": 12000,
                                        "expected_chapter_count": 8,
                                    },
                                    ensure_ascii=False,
                                )
                            }
                        }
                    }
                ),
            ),
            patch.object(service, "_persist_extract_raw_document", return_value=None) as persist_raw,
        ):
            payload = _run_async(
                service.extract_requirements_payload(
                    upload,
                    project_name="项目一",
                    project_id="proj-1",
                )
            )
        persist_raw.assert_called_once_with("proj-1", "脱敏后原文")
        self.assertEqual(payload["bid_type"], "tech")
        self.assertEqual(payload["project_summary"], "项目概述")
        self.assertEqual(payload["requirements"][0]["content"], "需求内容")
        self.assertEqual(payload["requirements"][0]["source_pages"][0]["page"], 0)
        self.assertEqual(payload["analysis_report"], [{"id": "n1"}])
        self.assertEqual(payload["expected_word_count"], 12000)
        self.assertEqual(payload["expected_chapter_count"], 8)

    def test_extract_requirements_stream_response_runs_natively_without_legacy_route(self) -> None:
        upload = SimpleNamespace(filename="demo.docx")
        upload.read = AsyncMock(return_value=b"docx-bytes")
        request = SimpleNamespace(is_disconnected=_async_false)

        with (
            patch.object(
                service,
                "_prepare_extract_document",
                return_value={
                    "cache_id": "proj-1",
                    "pdf_url": "/api/projects/pdf/proj-1",
                    "pages_text": [],
                    "raw_image_map": {"img1": {"name": "图1"}},
                    "text_for_dify": "定位原文",
                    "mapping_table": {},
                    "entity_count": 0,
                    "placeholder_manifest": {},
                    "placeholder_policy": {},
                },
            ),
            patch.object(
                service,
                "_run_bid_pipt_preprocess",
                return_value={
                    "text": "预处理后原文",
                    "mapping_table": {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                    "mapping_table_count": 1,
                    "placeholder_manifest": {"@@PIPT:v1:e000001:k11111111@@": {"entity_type": "name"}},
                    "placeholder_policy": {},
                },
            ),
            patch.object(service, "_persist_extract_raw_document", return_value=None),
        ):
            response = _run_async(
                service.extract_requirements_stream_response(
                    upload,
                    project_name="项目一",
                    project_id="proj-1",
                    enable_desensitize=True,
                )
            )
            body = _run_async(_collect_streaming_response_text(response))

        _ = request
        self.assertEqual(response.media_type, "text/event-stream")
        self.assertIn('event: progress', body)
        self.assertIn('"label": "解析文档结构"', body)
        self.assertIn('"label": "文档结构解析完成"', body)
        self.assertIn('"label": "脱敏完成，识别 1 处实体"', body)
        self.assertIn('event: result', body)
        self.assertIn('"raw_document": "预处理后原文"', body)
        self.assertIn('"mapping_table": {"@@PIPT:v1:e000001:k11111111@@": "张三"}', body)

    def test_finalize_generated_body_runs_without_legacy_task_routes(self) -> None:
        with patch.object(service, "_normalize_generated_markdown", side_effect=lambda content, _title: content):
            payload = service._finalize_generated_body(
                "```markdown\n一、完全响应招标要求\n\n1.1 提供实施方案\n```",
                "响应情况",
                strip_structural_numbering=True,
            )

        self.assertFalse(hasattr(service, "_legacy_task_routes_module"))
        self.assertEqual(payload, "完全响应招标要求\n\n提供实施方案")

    def test_compose_runtime_writing_hint_runs_without_legacy_builder(self) -> None:
        with patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import:
            payload = service._compose_runtime_writing_hint(
                "原始写作意图\n\n【正文扩写与技术深度约束（必须遵守）】\n旧约束",
                "1.1 响应情况",
                900,
                "验收, 风险",
                section_outline_slice="[当前] 1.1 响应情况",
                analysis_context="招标要求：须完全响应",
            )

        legacy_import.assert_not_called()
        self.assertIn("【本节目录层级定位（只用于理解，不得输出）】", payload)
        self.assertIn("【章内承接与开篇导入要求】", payload)
        self.assertIn("原始写作意图", payload)
        self.assertNotIn("旧约束", payload)
        self.assertIn("招标要求：须完全响应", payload)
        self.assertIn("目标篇幅：约 900 字", payload)
        self.assertIn("关键词覆盖：验收, 风险", payload)

    def test_finalize_content_output_resolves_placeholders_natively(self) -> None:
        with (
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
            patch.object(service, "_normalize_generated_markdown", side_effect=lambda content, _title: content),
            patch("app.services.bid_content_placeholder_service.get_engine") as placeholder_engine,
        ):
            conn = Mock()
            exists_result = Mock()
            exists_result.scalar_one.return_value = False
            conn.execute.return_value = exists_result
            placeholder_engine.return_value.begin.return_value.__enter__.return_value = conn
            content, report, warning = service._finalize_legacy_content_output(
                "我方人员 {{PIPT_1}} 与 {{BIDDER_ORG}} 完全响应。",
                "响应情况",
                request_mapping_flat={
                    "{{__PIPT_name_1__}}": "张三",
                    "{{__BIDDER_ORG__}}": "测试公司",
                },
            )

        legacy_import.assert_not_called()
        self.assertEqual(content, "我方人员 张三 与 测试公司 完全响应。")
        self.assertIsNone(warning)
        self.assertEqual(
            report,
            [
                {"placeholder": "{{PIPT_1}}", "original": "张三", "status": "success"},
                {"placeholder": "{{BIDDER_ORG}}", "original": "测试公司", "status": "success"},
            ],
        )

    def test_finalize_content_output_warns_without_blocking_for_placeholder_issues(self) -> None:
        with patch.object(service, "_normalize_generated_markdown", side_effect=lambda content, _title: content):
            content, report, warning = service._finalize_legacy_content_output(
                "联系人 {{ PIPT-name-1 }}，备用 @@PIPT:v1:e000001:k11111111@@。",
                "响应情况",
                request_mapping_flat={},
            )

        self.assertEqual(content, "联系人 {{ PIPT-name-1 }}，备用 @@PIPT:v1:e000001:k11111111@@。")
        self.assertEqual(
            report,
            [{"placeholder": "@@PIPT:v1:e000001:k11111111@@", "original": "", "status": "miss"}],
        )
        self.assertEqual(warning["code"], "placeholder_restore_warning")
        self.assertTrue(warning["has_illegal_placeholder"])
        self.assertTrue(warning["has_unresolved_placeholder"])

    def test_finalize_content_output_resolves_from_core_vault_fallback(self) -> None:
        token = "@@PIPT:v1:e000001:k11111111@@"
        entity_conn = Mock()
        bid_exists = Mock()
        bid_exists.scalar_one.return_value = False
        entity_conn.execute.return_value = bid_exists

        vault_conn = Mock()
        core_exists = Mock()
        core_exists.scalar_one.return_value = True
        core_rows = Mock()
        core_rows.mappings.return_value.all.return_value = [
            {"placeholder": token, "original_text_enc": "张三"}
        ]
        vault_conn.execute.side_effect = [core_exists, core_rows]
        engine = MagicMock()
        engine.connect.return_value.__enter__.side_effect = [entity_conn, vault_conn]

        with (
            patch.object(service, "_normalize_generated_markdown", side_effect=lambda content, _title: content),
            patch("app.services.bid_content_placeholder_service._decrypt_vault_original_text", return_value="张三"),
            patch("app.services.bid_content_placeholder_service.get_engine", return_value=engine),
        ):
            content, report, warning = service._finalize_legacy_content_output(
                f"联系人 {token}。",
                "响应情况",
                request_mapping_flat={},
            )

        self.assertEqual(content, "联系人 张三。")
        self.assertIsNone(warning)
        self.assertEqual(report, [{"placeholder": token, "original": "张三", "status": "success"}])

    def test_parse_group_content_results_keeps_placeholder_warning_sections(self) -> None:
        children = [
            {"section_id": "sec-1", "section_title": "第一节"},
            {"section_id": "sec-2", "section_title": "第二节"},
        ]
        outputs = {
            "sections": [
                {"section_id": "sec-1", "content": "第一节正文"},
                {"section_id": "sec-2", "content": "第二节正文"},
            ]
        }

        def finalize_side_effect(section_title: str, _outputs: dict, _mapping: dict) -> dict:
            if section_title == "第一节":
                return {
                    "content": "第一节正文 {{ PIPT-name-1 }}",
                    "word_count": 20,
                    "replace_report": [],
                    "placeholder_warning": {"code": "placeholder_restore_warning", "message": "模型生成发生错误，请手动修改异常文本或重新生成。"},
                    "placeholder_issues": [],
                }
            return {
                "content": "第二节正文",
                "word_count": 4,
                "replace_report": [],
                "placeholder_issues": [],
            }

        with patch.object(service, "_finalize_single_content_result", side_effect=finalize_side_effect):
            parsed = service._parse_group_content_results(outputs, children, {})

        self.assertEqual([row["section_id"] for row in parsed["sections"]], ["sec-1", "sec-2"])
        self.assertEqual(parsed["sections"][0]["placeholder_warning"]["code"], "placeholder_restore_warning")
        self.assertEqual(parsed["failed_sections"], [])
        self.assertEqual(parsed["parse_error"], "")

    def test_rebuild_locator_payload_rebuilds_docx_snapshot_natively(self) -> None:
        blocks = [
            {"block_id": "B000000", "locator": "P0000", "body_idx": 0, "type": "paragraph", "text": "第一段"},
            {"block_id": "B000001", "locator": "P0001", "body_idx": 1, "type": "paragraph", "text": "第二段"},
        ]
        upload = SimpleNamespace(filename="demo.docx", read=MagicMock())
        upload.read.return_value = _async_value(b"docx-bytes")

        with (
            patch.object(service, "_extract_docx_blocks", return_value=blocks) as extract_blocks,
            patch.object(service, "_persist_docx_cache") as persist_docx,
            patch.object(service, "_persist_project_doc_blocks_snapshot") as persist_snapshot,
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            payload = _run_async(service.rebuild_locator_payload("proj-1", upload))

        legacy_import.assert_not_called()
        extract_blocks.assert_called_once_with(b"docx-bytes")
        persist_docx.assert_called_once_with("proj-1", b"docx-bytes")
        persist_snapshot.assert_called_once_with(project_id="proj-1", doc_blocks=blocks)
        self.assertEqual(payload["project_id"], "proj-1")
        self.assertEqual(payload["blocks"], 2)
        self.assertEqual(payload["locators"], 2)
        self.assertTrue(payload["snapshot_only"])

    def test_rebuild_locator_payload_rejects_non_docx_filename(self) -> None:
        upload = SimpleNamespace(filename="demo.pdf", read=MagicMock())

        with self.assertRaises(Exception) as ctx:
            _run_async(service.rebuild_locator_payload("proj-1", upload))

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_extract_docx_blocks_extracts_paragraphs_and_tables(self) -> None:
        import io
        import docx

        document = docx.Document()
        document.add_paragraph("第一段")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "列一"
        table.cell(0, 1).text = "列二"
        table.cell(1, 0).text = "值一"
        table.cell(1, 1).text = "值二"
        buffer = io.BytesIO()
        document.save(buffer)

        blocks = service._extract_docx_blocks(buffer.getvalue())

        self.assertEqual(blocks[0]["block_id"], "B000000")
        self.assertEqual(blocks[0]["locator"], "P0000")
        self.assertEqual(blocks[0]["type"], "paragraph")
        self.assertEqual(blocks[0]["text"], "第一段")
        self.assertEqual(blocks[1]["block_id"], "B000001")
        self.assertEqual(blocks[1]["locator"], "P0001")
        self.assertEqual(blocks[1]["type"], "table")
        self.assertIn("| 列一 | 列二 |", blocks[1]["text"])

    def test_extract_bid_attachment_payload_reads_snapshot_natively(self) -> None:
        blocks = [
            {"block_id": "B000000", "locator": "P0000", "body_idx": 0, "type": "paragraph", "text": "第一段<script>"},
            {"block_id": "B000001", "locator": "P0001", "body_idx": 1, "type": "table", "text": "| A | B |\n| --- | --- |"},
        ]
        with (
            patch.object(service, "get_project_doc_blocks_payload", return_value={"project_id": "proj-1", "blocks": blocks}),
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            payload = _run_async(
                service.extract_bid_attachment_payload(
                    {
                        "project_id": "proj-1",
                        "start_locator": "[p0]",
                        "end_locator": "P0001",
                        "attachment_name": "附件一",
                    }
                )
            )

        legacy_import.assert_not_called()
        self.assertEqual(payload["attachment_name"], "附件一")
        self.assertEqual(payload["start_locator"], "P0000")
        self.assertEqual(payload["end_locator"], "P0001")
        self.assertEqual(payload["resolved_start_locator"], "P0000")
        self.assertEqual(payload["resolved_end_locator"], "P0001")
        self.assertEqual(payload["paragraph_count"], 2)
        self.assertTrue(payload["snapshot_only"])
        self.assertIn("<p>第一段&lt;script&gt;</p>", payload["html"])
        self.assertIn("<pre>| A | B |<br/>| --- | --- |</pre>", payload["html"])

    def test_extract_bid_attachment_by_block_payload_reads_snapshot_natively(self) -> None:
        blocks = [
            {"block_id": "B000000", "locator": "P0000", "body_idx": 0, "type": "paragraph", "text": "第一段"},
            {"block_id": "B000001", "locator": "P0001", "body_idx": 1, "type": "paragraph", "text": "第二段"},
        ]
        with (
            patch.object(service, "get_project_doc_blocks_payload", return_value={"project_id": "proj-1", "blocks": blocks}),
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            payload = _run_async(
                service.extract_bid_attachment_by_block_payload(
                    {
                        "project_id": "proj-1",
                        "start_block_id": "B000001",
                        "end_block_id": "B000000",
                        "attachment_name": "附件一",
                    }
                )
            )

        legacy_import.assert_not_called()
        self.assertEqual(payload["attachment_name"], "附件一")
        self.assertEqual(payload["start_block_id"], "B000000")
        self.assertEqual(payload["end_block_id"], "B000001")
        self.assertEqual(payload["paragraph_count"], 2)
        self.assertTrue(payload["snapshot_only"])
        self.assertEqual(payload["html"], "<p>第一段</p>\n<p>第二段</p>")

    def test_extract_bid_attachment_by_block_payload_returns_404_for_missing_block(self) -> None:
        with patch.object(
            service,
            "get_project_doc_blocks_payload",
            return_value={"project_id": "proj-1", "blocks": [{"block_id": "B000000", "locator": "P0000", "body_idx": 0}]},
        ):
            with self.assertRaises(Exception) as ctx:
                _run_async(
                    service.extract_bid_attachment_by_block_payload(
                        {"project_id": "proj-1", "start_block_id": "B000000", "end_block_id": "B000002"}
                    )
                )

        self.assertEqual(getattr(ctx.exception, "status_code", None), 404)

    def test_extract_bid_attachment_by_block_docx_response_slices_cached_docx_natively(self) -> None:
        import io
        import docx

        document = docx.Document()
        document.add_paragraph("第一段")
        document.add_paragraph("第二段")
        document.add_paragraph("第三段")
        buffer = io.BytesIO()
        document.save(buffer)

        docx_path = Mock()
        docx_path.exists.return_value = True
        docx_path.read_bytes.return_value = buffer.getvalue()
        blocks = [
            {"block_id": "B000000", "locator": "P0000", "body_idx": 0, "type": "paragraph", "text": "第一段"},
            {"block_id": "B000001", "locator": "P0001", "body_idx": 1, "type": "paragraph", "text": "第二段"},
            {"block_id": "B000002", "locator": "P0002", "body_idx": 2, "type": "paragraph", "text": "第三段"},
        ]

        with (
            patch.object(service, "get_project_doc_blocks_payload", return_value={"project_id": "proj-1", "blocks": blocks}),
            patch.object(service, "_docx_cache_path", return_value=docx_path),
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            payload = _run_async(
                service.extract_bid_attachment_by_block_docx_response(
                    {
                        "project_id": "proj-1",
                        "start_block_id": "B000002",
                        "end_block_id": "B000001",
                        "attachment_name": "附件 一",
                    }
                )
            )

        legacy_import.assert_not_called()
        self.assertEqual(payload.filename, "slice_B000001_B000002.docx")
        self.assertEqual(payload.media_type, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertEqual(payload.headers, {"X-Start-Block-Id": "B000001", "X-End-Block-Id": "B000002"})

        sliced = docx.Document(io.BytesIO(payload.content))
        sliced_text = "\n".join(paragraph.text for paragraph in sliced.paragraphs)
        self.assertNotIn("第一段", sliced_text)
        self.assertIn("第二段", sliced_text)
        self.assertIn("第三段", sliced_text)

    def test_extract_bid_attachment_by_block_docx_response_returns_409_when_docx_missing(self) -> None:
        docx_path = Mock()
        docx_path.exists.return_value = False
        blocks = [{"block_id": "B000000", "locator": "P0000", "body_idx": 0, "type": "paragraph", "text": "第一段"}]

        with (
            patch.object(service, "get_project_doc_blocks_payload", return_value={"project_id": "proj-1", "blocks": blocks}),
            patch.object(service, "_docx_cache_path", return_value=docx_path),
        ):
            with self.assertRaises(Exception) as ctx:
                _run_async(
                    service.extract_bid_attachment_by_block_docx_response(
                        {"project_id": "proj-1", "start_block_id": "B000000", "end_block_id": "B000000"}
                    )
                )

        self.assertEqual(getattr(ctx.exception, "status_code", None), 409)

    def test_stream_task_progress_response_streams_native_sse_without_legacy_route(self) -> None:
        task = SimpleNamespace(
            task_id="task-1",
            project_id="proj-1",
            status="done",
            stages=[
                "生成大纲中",
                "__text__正文片段",
                '__node__{"id":"n1"}',
                '__task_event__{"event":"custom","value":1}',
            ],
            result={"done": True},
            created_at=1780296000,
        )
        request = SimpleNamespace(is_disconnected=_async_false)

        with (
            patch.object(service, "_require_task_owner", return_value=task),
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            response = _run_async(service.stream_task_progress_response("task-1", request, project_id="proj-1"))
            body = _run_async(_collect_streaming_response_text(response))

        legacy_import.assert_not_called()
        self.assertEqual(response.media_type, "text/event-stream")
        self.assertIn("event: stage", body)
        self.assertIn('"stage": "生成大纲中"', body)
        self.assertIn('"text": "正文片段"', body)
        self.assertIn("event: node_complete", body)
        self.assertIn("event: custom", body)
        self.assertIn("event: done", body)

    def test_stream_task_progress_response_emits_error_event_when_task_missing(self) -> None:
        request = SimpleNamespace(is_disconnected=_async_false)
        with patch.object(
            service,
            "_require_task_owner",
            side_effect=service.PlatformError(code="RESOURCE_NOT_FOUND", message="任务不存在或已过期", status_code=404),
        ):
            response = _run_async(service.stream_task_progress_response("task-1", request))
            body = _run_async(_collect_streaming_response_text(response))

        self.assertIn('"error": "任务不存在或已过期"', body)

    def test_analyze_node_response_streams_native_done_and_bid_attachments(self) -> None:
        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {
                "__finished__": True,
                "outputs": {
                    "text": '{"node-x": "节点内容"}<BID_ATTACHMENTS>[{"name":"附件一","start_locator":"P1","end_locator":"P2","description":"说明"}]</BID_ATTACHMENTS>'
                },
            }

        with (
            patch.object(service, "_get_workflow_key", side_effect=["app-doc-analysis"]),
            patch.object(service, "_load_raw_document", return_value="原文内容"),
            patch.object(
                service,
                "load_docanalysis_framework",
                return_value=(
                    "系统提示",
                    [{"id": "node-x", "label": "节点X", "extractionPrompt": "抽取节点X"}],
                ),
            ),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
        ):
            response = _run_async(service.analyze_node_response("proj-1", {"node_id": "node-x"}))
            body = _run_async(_collect_streaming_response_text(response))

        self.assertEqual(response.media_type, "text/event-stream")
        self.assertIn('"type": "bid_attachments"', body)
        self.assertIn('"type": "done"', body)
        self.assertIn('"node_id": "node-x"', body)
        self.assertIn('"content": "节点内容"', body)

    def test_analyze_node_response_streams_error_when_finished_missing(self) -> None:
        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            if False:
                yield None

        with (
            patch.object(service, "_get_workflow_key", side_effect=["app-doc-analysis"]),
            patch.object(service, "_load_raw_document", return_value="原文内容"),
            patch.object(
                service,
                "load_docanalysis_framework",
                return_value=(
                    "系统提示",
                    [{"id": "node-x", "label": "节点X", "extractionPrompt": "抽取节点X"}],
                ),
            ),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
        ):
            response = _run_async(service.analyze_node_response("proj-1", {"node_id": "node-x"}))
            body = _run_async(_collect_streaming_response_text(response))

        self.assertIn('"type": "error"', body)

    def test_analyze_document_response_streams_native_progress_node_and_complete(self) -> None:
        async def fake_group_extract(**kwargs):
            _ = kwargs
            return [
                {"node_id": "proj_overview", "label": "项目概述", "content": "概述内容"},
                {"node_id": "proj_basic", "label": "项目基本信息", "content": "基本信息"},
            ]

        with (
            patch.object(service, "_get_workflow_key", return_value="doc-key"),
            patch.object(service, "load_docanalysis_framework", return_value=("系统提示", _analysis_framework_nodes())),
            patch.object(service, "_bid_generator_root", return_value=Path("/tmp")),
            patch.object(service.Path, "exists", return_value=True),
            patch.object(service, "_load_raw_document", return_value="缓存原文"),
            patch.object(service, "_extract_docanalysis_group_results", new=AsyncMock(side_effect=fake_group_extract)),
        ):
            response = _run_async(service.analyze_document_response(project_id="proj-1"))
            body = _run_async(_collect_streaming_response_text(response))

        self.assertEqual(response.media_type, "text/event-stream")
        self.assertIn("event: progress", body)
        self.assertIn('"message": "并行解析 2 组 / 共 7 个节点"', body)
        self.assertIn("event: node_complete", body)
        self.assertIn('"node_id": "proj_overview"', body)
        self.assertIn("event: complete", body)
        self.assertIn('"success_count": 4', body)

    def test_analyze_document_response_raises_when_cached_document_missing(self) -> None:
        with (
            patch.object(service, "_get_workflow_key", return_value="doc-key"),
            patch.object(service, "load_docanalysis_framework", return_value=("系统提示", _analysis_framework_nodes())),
            patch.object(service, "_bid_generator_root", return_value=Path("/tmp")),
            patch.object(service.Path, "exists", return_value=True),
            patch.object(service, "_load_raw_document", return_value=""),
        ):
            with self.assertRaises(Exception) as ctx:
                _run_async(service.analyze_document_response(project_id="proj-1"))

    def test_re_extract_requirements_payload_runs_natively_with_cached_document(self) -> None:
        with (
            patch.object(service, "_load_raw_document", return_value="缓存原文"),
            patch.object(service, "_get_workflow_key", return_value="req-key"),
            patch.object(
                service,
                "_call_dify_workflow",
                new=AsyncMock(
                    return_value={
                        "data": {
                            "outputs": {
                                "text": json.dumps(
                                    {
                                        "bid_type": "tech",
                                        "project_summary": "项目概述",
                                        "requirements": [{"type": "tech", "content": "需求A", "points": 10}],
                                        "required_attachments": [{"name": "投标函"}],
                                        "scoring_table_template": [{"indicator": "技术方案"}],
                                        "expected_word_count": 15000,
                                    },
                                    ensure_ascii=False,
                                )
                            }
                        }
                    }
                ),
            ),
        ):
            payload = _run_async(
                service.re_extract_requirements_payload(
                    {
                        "project_id": "proj-1",
                        "project_name": "项目一",
                    }
                )
            )
        self.assertEqual(payload["bid_type"], "tech")
        self.assertEqual(payload["project_summary"], "项目概述")
        self.assertEqual(payload["requirements"][0]["content"], "需求A")
        self.assertEqual(payload["mapping_table"], {})
        self.assertEqual(payload["entity_count"], 0)
        self.assertEqual(payload["raw_document"], "缓存原文")
        self.assertEqual(payload["expected_word_count"], 15000)

    def test_re_extract_requirements_payload_raises_when_cached_document_missing(self) -> None:
        with patch.object(service, "_load_raw_document", return_value=""):
            with self.assertRaises(Exception) as ctx:
                _run_async(
                    service.re_extract_requirements_payload(
                        {
                            "project_id": "proj-1",
                            "project_name": "项目一",
                        }
                    )
                )

        self.assertEqual(getattr(ctx.exception, "status_code", None), 404)

    def test_generate_content_payload_runs_natively_without_legacy_route(self) -> None:
        with (
            patch.object(service, "_validate_required_bidder_info", return_value=None),
            patch.object(service, "_get_workflow_key", return_value="content-key"),
            patch.object(service, "_compose_runtime_writing_hint", return_value="合成提示"),
            patch.object(
                service,
                "_merge_bidder_pipt_context",
                return_value=(
                    {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                    "占位符提示",
                    {},
                ),
            ),
            patch.object(
                service,
                "_call_dify_workflow",
                new=AsyncMock(
                    return_value={
                        "data": {
                            "outputs": {
                                "text": "原始正文",
                                "quality_score": "88",
                                "feedback": "反馈意见",
                            }
                        }
                    }
                ),
            ),
            patch.object(
                service,
                "_finalize_legacy_content_output",
                return_value=("最终正文", [{"placeholder": "@@PIPT:v1:e000001:k11111111@@", "original": "张三"}], None),
            ),
            patch.object(
                service,
                "_run_inline_content_diagram",
                new=AsyncMock(return_value=("最终正文", 1, None, {"mode": "inline"})),
            ),
        ):
            payload = _run_async(
                service.generate_content_payload(
                    {
                        "project_id": "proj-1",
                        "section_id": "sec-1",
                        "section_title": "第一章",
                        "writing_hint": "提示",
                        "expected_words": 1200,
                        "project_summary": "项目概述",
                        "global_outline": "全局大纲",
                        "mapping_table": {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                        "bidder_info": {"companyName": "某公司"},
                    }
                )
            )
        self.assertEqual(payload["section_id"], "sec-1")
        self.assertEqual(payload["content"], "最终正文")
        self.assertEqual(payload["quality_score"], 88)
        self.assertEqual(payload["diagrams_count"], 1)
        self.assertEqual(payload["diagram_specs"], {"mode": "inline"})

    def test_generate_content_stream_response_runs_natively_without_legacy_route(self) -> None:
        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"__stage__": "✍️ 正文生成"}
            yield "第一段"
            yield {
                "__finished__": True,
                "outputs": {
                    "text": "完整正文",
                    "feedback": "反馈",
                    "quality_score": "90",
                },
            }

        with (
            patch.object(service, "_validate_required_bidder_info", return_value=None),
            patch.object(service, "_get_workflow_key", return_value="content-key"),
            patch.object(service, "_compose_runtime_writing_hint", return_value="合成提示"),
            patch.object(
                service,
                "_merge_bidder_pipt_context",
                return_value=(
                    {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                    "占位符提示",
                    {},
                ),
            ),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
            patch.object(
                service,
                "_finalize_legacy_content_output",
                return_value=("完整正文", [{"placeholder": "@@PIPT:v1:e000001:k11111111@@", "original": "张三"}], None),
            ),
            patch.object(
                service,
                "_run_inline_content_diagram",
                new=AsyncMock(return_value=("完整正文", 0, None, None)),
            ),
        ):
            response = _run_async(
                service.generate_content_stream_response(
                    {
                        "project_id": "proj-1",
                        "section_id": "sec-1",
                        "section_title": "第一章",
                        "writing_hint": "提示",
                        "expected_words": 1200,
                        "project_summary": "项目概述",
                        "global_outline": "全局大纲",
                        "mapping_table": {"@@PIPT:v1:e000001:k11111111@@": "张三"},
                        "bidder_info": {"companyName": "某公司"},
                    }
                )
            )
            body = _run_async(_collect_streaming_response_text(response))
        self.assertEqual(response.media_type, "text/event-stream")
        self.assertIn('"stage": "✍️ 正文生成"', body)
        self.assertIn('"text": "第一段"', body)
        self.assertIn('"done": true', body.lower())
        self.assertIn('"word_count": 4', body)

    def test_generate_outline_payload_returns_native_sections(self) -> None:
        with (
            patch.object(service, "_get_workflow_key", return_value="app-structure"),
            patch.object(
                service,
                "_call_dify_workflow",
                new=AsyncMock(
                    return_value={
                        "data": {
                            "outputs": {
                                "structured_output": {
                                    "outline": [
                                        {
                                            "title": "总体技术方案",
                                            "children": [{"title": "建设思路", "writingHint": "覆盖技术要求与实施路径", "keywords": ["建设", "实施"]}],
                                        }
                                    ]
                                }
                            }
                        }
                    }
                ),
            ),
        ):
            payload = _run_async(
                service.generate_outline_payload(
                    {
                        "bid_type": "tech",
                        "requirements": [{"type": "tech", "content": "满足建设要求"}],
                        "analysis_context": "技术要求：覆盖总体方案",
                        "structure_heading_seed_json": '[{"id":"h2-1","title":"总体技术方案","generation_strategy":"general"}]',
                    }
                )
            )

        self.assertEqual(payload["sections"][0]["title"], "总体技术方案")
        self.assertEqual(payload["sections"][0]["children"][0]["title"], "建设思路")

    def test_generate_outline_payload_raises_when_quality_check_fails(self) -> None:
        with (
            patch.object(service, "_get_workflow_key", return_value="app-structure"),
            patch.object(
                service,
                "_call_dify_workflow",
                new=AsyncMock(return_value={"data": {"outputs": {"structured_output": {"outline": []}}}}),
            ),
        ):
            with self.assertRaises(Exception) as ctx:
                _run_async(
                    service.generate_outline_payload(
                        {
                            "bid_type": "tech",
                            "requirements": [{"type": "tech", "content": "满足建设要求"}],
                            "structure_heading_seed_json": '[{"id":"h2-1","title":"总体技术方案","generation_strategy":"general"}]',
                        }
                    )
                )

        self.assertEqual(getattr(ctx.exception, "status_code", None), 502)

    def test_generate_outline_stream_response_emits_stage_and_done_events_natively(self) -> None:
        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"__stage__": "✍️ 生成大纲", "workflow_run_id": "run-1"}
            yield {
                "__finished__": True,
                "outputs": {
                    "structured_output": {
                        "outline": [
                            {
                                "title": "总体技术方案",
                                "children": [
                                    {
                                        "title": "建设目标",
                                        "wordCount": 500,
                                        "writingHint": "围绕建设目标、实施路径、交付要求与验收指标展开详细说明，确保方案结构完整且可执行。",
                                        "keywords": ["建设目标", "实施路径", "验收指标"],
                                    }
                                ],
                            }
                        ]
                    }
                },
                "workflow_run_id": "run-1",
            }

        with (
            patch.object(service, "_get_workflow_key", return_value="app-structure"),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
        ):
            response = _run_async(
                service.generate_outline_stream_response(
                    {
                        "bid_type": "tech",
                        "requirements": [{"type": "tech", "content": "满足建设要求"}],
                        "expected_total_words": 600,
                        "structure_heading_seed_json": '[{"id":"h2-1","title":"总体技术方案","generation_strategy":"general"}]',
                    }
                )
            )
            body = _run_async(_collect_streaming_response_text(response))

        self.assertEqual(response.media_type, "text/event-stream")
        self.assertIn('"stage": "✍️ 生成大纲"', body)
        self.assertIn('"done": true', body)
        self.assertIn('"title": "总体技术方案"', body)
        self.assertIn('"wordCount": 600', body)

    def test_generate_outline_stream_response_falls_back_to_workflow_run_result(self) -> None:
        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"__stage__": "✍️ 生成大纲", "workflow_run_id": "run-2"}
            yield {"__finished__": True, "outputs": {}, "workflow_run_id": "run-2"}

        fallback_payload = {
            "data": {
                "outputs": {
                    "structured_output": {
                        "outline": [
                            {
                                "title": "总体技术方案",
                                "children": [
                                    {
                                        "title": "实施方案",
                                        "wordCount": 480,
                                        "writingHint": "说明实施组织、关键路径、资源投入与质量保障，覆盖建设全过程的执行安排与控制点。",
                                        "keywords": ["实施方案", "资源投入", "质量保障"],
                                    }
                                ],
                            }
                        ]
                    }
                }
            }
        }

        with (
            patch.object(service, "_get_workflow_key", return_value="app-structure"),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
            patch.object(service, "_get_dify_workflow_run_result", new=AsyncMock(return_value=fallback_payload)) as fallback_mock,
        ):
            response = _run_async(
                service.generate_outline_stream_response(
                    {
                        "bid_type": "tech",
                        "requirements": [{"type": "tech", "content": "满足建设要求"}],
                        "expected_total_words": 700,
                        "structure_heading_seed_json": '[{"id":"h2-1","title":"总体技术方案","generation_strategy":"general"}]',
                    }
                )
            )
            body = _run_async(_collect_streaming_response_text(response))

        fallback_mock.assert_awaited_once_with("app-structure", "run-2")
        self.assertIn('"done": true', body)
        self.assertIn('"wordCount": 700', body)

    def test_generate_outline_stream_response_emits_error_when_quality_check_fails(self) -> None:
        async def fake_stream(*args, **kwargs):
            _ = args, kwargs
            yield {"__finished__": True, "outputs": {"structured_output": {"outline": []}}, "workflow_run_id": "run-3"}

        with (
            patch.object(service, "_get_workflow_key", return_value="app-structure"),
            patch.object(service, "_call_dify_workflow_stream", new=fake_stream),
            patch.object(service, "_get_dify_workflow_run_result", new=AsyncMock(return_value={})),
        ):
            response = _run_async(
                service.generate_outline_stream_response(
                    {
                        "bid_type": "tech",
                        "requirements": [{"type": "tech", "content": "满足建设要求"}],
                        "structure_heading_seed_json": '[{"id":"h2-1","title":"总体技术方案","generation_strategy":"general"}]',
                    }
                )
            )
            body = _run_async(_collect_streaming_response_text(response))

        self.assertIn('"error": "大纲生成结构不完整，请重试：', body)

    def test_export_report_response_builds_pdf_natively(self) -> None:
        self.assertFalse(hasattr(service.bid_workflow_execution_adapter, "export_report_response"))

        response = _run_async(
            service.export_report_response(
                {
                    "project_name": "项目一",
                    "nodes": [{"id": "n1", "label": "资质要求", "content": "满足招标文件要求"}],
                }
            )
        )

        self.assertEqual(response.media_type, "application/pdf")
        self.assertEqual(response.filename, "analysis-report.pdf")
        self.assertTrue(response.content.startswith(b"%PDF"))
        self.assertIn("Content-Disposition", response.headers)

    def test_forge_document_response_builds_docx_natively_without_legacy_adapter(self) -> None:
        import io
        import docx

        class FakeForge:
            def __init__(self, **kwargs):
                self.kwargs = kwargs

            def build(self, sections, scoring_rows=None, attachments=None):
                document = docx.Document()
                for section in sections:
                    title = section.get("title") or section.get("heading_text")
                    if title:
                        document.add_heading(title, level=int(section.get("heading_level") or 1))
                    content = section.get("content")
                    if content:
                        document.add_paragraph(content)
                buffer = io.BytesIO()
                document.save(buffer)
                return buffer.getvalue()

        with (
            patch.object(service, "create_document_forge", return_value=FakeForge()),
            patch.object(service, "_load_forge_pipt_mapping", return_value={}),
            patch.object(service, "_load_forge_image_map", return_value={}),
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            payload = _run_async(
                service.forge_document_response(
                    {
                        "project_id": "proj-1",
                        "project_name": "项目一",
                        "sections": [{"id": "s1", "title": "第一章", "content": "正文内容", "heading_level": 1}],
                    }
                )
            )

        legacy_import.assert_not_called()
        self.assertFalse(hasattr(service.bid_workflow_execution_adapter, "forge_document_response"))
        self.assertEqual(payload.media_type, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.assertEqual(payload.filename, "document.docx")
        document = docx.Document(io.BytesIO(payload.content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("第一章", text)
        self.assertIn("正文内容", text)

    def test_forge_document_response_preserves_docx_slice_natively(self) -> None:
        import io
        import docx

        source = docx.Document()
        source.add_paragraph("原文第一段")
        source.add_paragraph("原文第二段")
        source.add_paragraph("原文第三段")
        source_buffer = io.BytesIO()
        source.save(source_buffer)

        docx_path = Mock()
        docx_path.exists.return_value = True
        docx_path.read_bytes.return_value = source_buffer.getvalue()
        blocks = [
            {"block_id": "B000000", "locator": "P0000", "body_idx": 0, "type": "paragraph", "text": "原文第一段"},
            {"block_id": "B000001", "locator": "P0001", "body_idx": 1, "type": "paragraph", "text": "原文第二段"},
            {"block_id": "B000002", "locator": "P0002", "body_idx": 2, "type": "paragraph", "text": "原文第三段"},
        ]

        def fake_build(sections, scoring_rows=None, attachments=None):
            document = docx.Document()
            for section in sections:
                if section.get("title"):
                    document.add_heading(section["title"], level=int(section.get("heading_level") or 1))
                if section.get("content"):
                    document.add_paragraph(section["content"])
            buffer = io.BytesIO()
            document.save(buffer)
            return buffer.getvalue()

        fake_forge = SimpleNamespace(build=fake_build)
        with (
            patch.object(service, "create_document_forge", return_value=fake_forge),
            patch.object(service, "add_scoring_table_and_attachments", return_value=None),
            patch.object(service, "_load_forge_pipt_mapping", return_value={}),
            patch.object(service, "_load_forge_image_map", return_value={}),
            patch.object(service, "get_project_doc_blocks_payload", return_value={"project_id": "proj-1", "blocks": blocks}),
            patch.object(service, "_docx_cache_path", return_value=docx_path),
            patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import,
        ):
            payload = _run_async(
                service.forge_document_response(
                    {
                        "project_id": "proj-1",
                        "project_name": "项目一",
                        "sections": [
                            {"id": "intro", "title": "前言", "content": "生成正文", "heading_level": 1},
                            {
                                "id": "slice",
                                "title": "原文附件",
                                "source_type": "docx_slice",
                                "start_block_id": "B000001",
                                "end_block_id": "B000002",
                                "inject_title": True,
                                "heading_level": 1,
                            },
                        ],
                    }
                )
            )

        legacy_import.assert_not_called()
        document = docx.Document(io.BytesIO(payload.content))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        self.assertIn("前言", text)
        self.assertIn("生成正文", text)
        self.assertIn("原文附件", text)
        self.assertNotIn("原文第一段", text)
        self.assertIn("原文第二段", text)
        self.assertIn("原文第三段", text)

    def test_hybrid_forge_groups_contiguous_markdown_before_page_breaks(self) -> None:
        import docx

        source = docx.Document()
        source.add_paragraph("原文切片")
        source_buffer = io.BytesIO()
        source.save(source_buffer)

        docx_path = Mock()
        docx_path.exists.return_value = True
        docx_path.read_bytes.return_value = source_buffer.getvalue()
        blocks = [{"block_id": "B000000", "locator": "P0000", "body_idx": 0, "type": "paragraph", "text": "原文切片"}]
        build_calls: list[list[dict]] = []

        def fake_build(sections, scoring_rows=None, attachments=None):
            build_calls.append([dict(section) for section in sections])
            document = docx.Document()
            for section in sections:
                title = section.get("title") or section.get("heading_text")
                if title:
                    document.add_heading(title, level=int(section.get("heading_level") or 1))
                content = section.get("content")
                if content:
                    document.add_paragraph(content)
            buffer = io.BytesIO()
            document.save(buffer)
            return buffer.getvalue()

        fake_forge = SimpleNamespace(build=fake_build)
        with (
            patch.object(service, "create_document_forge", return_value=fake_forge),
            patch.object(service, "add_scoring_table_and_attachments", return_value=None),
            patch.object(service, "_load_forge_pipt_mapping", return_value={}),
            patch.object(service, "_load_forge_image_map", return_value={}),
            patch.object(service, "get_project_doc_blocks_payload", return_value={"project_id": "proj-1", "blocks": blocks}),
            patch.object(service, "_docx_cache_path", return_value=docx_path),
        ):
            _run_async(
                service.forge_document_response(
                    {
                        "project_id": "proj-1",
                        "project_name": "项目一",
                        "sections": [
                            {"id": "root", "title": "技术部分", "content": "", "heading_level": 1, "title_only": True},
                            {"id": "parent", "title": "实施方案", "content": "", "heading_level": 2, "title_only": True},
                            {"id": "child", "title": "总体方案", "content": "正文内容", "heading_level": 3},
                            {
                                "id": "slice",
                                "title": "原文附件",
                                "source_type": "docx_slice",
                                "start_block_id": "B000000",
                                "end_block_id": "B000000",
                                "inject_title": True,
                                "heading_level": 1,
                            },
                        ],
                    }
                )
            )

        markdown_calls = [call for call in build_calls if any(section.get("id") == "child" for section in call)]
        self.assertEqual(len(markdown_calls), 1)
        self.assertEqual([section.get("id") for section in markdown_calls[0]], ["root", "parent", "child"])

    def test_export_scoring_table_response_builds_xlsx_natively(self) -> None:
        import io
        import openpyxl

        with patch.object(service, "_ensure_legacy_imported", create=True) as legacy_import:
            payload = _run_async(
                service.export_scoring_table_response(
                    {
                        "project_name": "项目一",
                        "rows": [
                            {
                                "indicator": "技术方案",
                                "max_score": 10,
                                "criteria": "满足要求",
                                "self_response": "full",
                                "self_comment": "完全响应",
                                "evidence_refs": ["章节1", "附件A"],
                            }
                        ],
                    }
                )
            )

        legacy_import.assert_not_called()
        self.assertEqual(payload.media_type, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        self.assertEqual(payload.filename, "scoring.xlsx")
        self.assertFalse(payload.inline)
        self.assertIn("filename*=UTF-8", payload.headers["Content-Disposition"])

        workbook = openpyxl.load_workbook(io.BytesIO(payload.content))
        worksheet = workbook.active
        self.assertEqual(worksheet.title, "自评评分表")
        self.assertEqual(worksheet.cell(row=1, column=1).value, "评分指标")
        self.assertEqual(worksheet.cell(row=2, column=1).value, "技术方案")
        self.assertEqual(worksheet.cell(row=2, column=4).value, "响应")
        self.assertEqual(worksheet.cell(row=2, column=6).value, "章节1\n附件A")
        self.assertEqual(worksheet.cell(row=3, column=1).value, "合计")
        self.assertEqual(worksheet.cell(row=3, column=2).value, 10)

    def test_export_scoring_table_response_rejects_non_array_rows(self) -> None:
        with self.assertRaises(Exception) as ctx:
            _run_async(service.export_scoring_table_response({"project_name": "项目一", "rows": {}}))

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_bid_service_does_not_expose_knowledge_sync_triggers(self) -> None:
        self.assertFalse(hasattr(service, "trigger_knowledge_sync_payload"))
        self.assertFalse(hasattr(service, "trigger_single_knowledge_sync_payload"))
        self.assertFalse(hasattr(service, "trigger_kb_sync_payload"))

    def test_list_knowledge_images_filters_and_joins_preview_url(self) -> None:
        now = datetime(2026, 6, 1, tzinfo=timezone.utc)
        asset_row = {
            "image_hash": "abc123",
            "placeholder": "__PRO_IMG_abc123__",
            "source_doc": "doc-a",
            "source_page": 2,
            "caption": "图表",
            "image_type": "chart",
            "summary": "摘要",
            "tags_json": json.dumps(["图表", "流程"], ensure_ascii=False),
            "caption_status": "captioned",
            "created_at": now,
        }
        conn = Mock()
        assets_exists = Mock()
        assets_exists.scalar_one.return_value = True
        rows_result = Mock()
        rows_result.mappings.return_value.all.return_value = [asset_row]
        registry_exists = Mock()
        registry_exists.scalar_one.return_value = True
        registry_result = Mock()
        registry_result.mappings.return_value.all.return_value = [
            {"image_hash": "abc123", "preview_url": "/api/extracted-images/a.png"}
        ]
        conn.execute.side_effect = [assets_exists, rows_result, registry_exists, registry_result]
        engine = MagicMock()
        engine.begin.return_value.__enter__.return_value = conn

        with patch.object(service, "get_engine", return_value=engine):
            payload = service.list_knowledge_images_payload(source_doc="doc-a", caption_status="captioned", limit=999)

        self.assertEqual(payload["total"], 1)
        self.assertEqual(payload["items"][0]["tags"], ["图表", "流程"])
        self.assertEqual(payload["items"][0]["preview_url"], "/api/extracted-images/a.png")
        params = conn.execute.call_args_list[1].args[1]
        self.assertEqual(params["source_doc"], "doc-a")
        self.assertEqual(params["caption_status"], "captioned")
        self.assertEqual(params["limit"], 500)

    def test_update_knowledge_image_defaults_caption_status_to_manual(self) -> None:
        now = datetime(2026, 6, 1, tzinfo=timezone.utc)
        selected = {
            "image_hash": "abc123",
            "placeholder": "__PRO_IMG_abc123__",
            "source_doc": "doc-a",
            "source_page": 1,
            "caption": "",
            "image_type": "",
            "summary": "",
            "tags_json": "[]",
            "caption_status": "pending",
            "created_at": now,
        }
        updated = {
            **selected,
            "caption": "人工说明",
            "image_type": "diagram",
            "summary": "更新摘要",
            "tags_json": json.dumps(["A", "B"], ensure_ascii=False),
            "caption_status": "manual",
        }
        conn = Mock()
        assets_exists = Mock()
        assets_exists.scalar_one.return_value = True
        select_result = Mock()
        select_result.mappings.return_value.first.return_value = selected
        update_result = Mock()
        update_result.mappings.return_value.one.return_value = updated
        registry_exists = Mock()
        registry_exists.scalar_one.return_value = True
        registry_result = Mock()
        registry_result.mappings.return_value.all.return_value = [
            {"image_hash": "abc123", "preview_url": "/api/extracted-images/a.png"}
        ]
        conn.execute.side_effect = [assets_exists, select_result, update_result, registry_exists, registry_result]
        engine = MagicMock()
        engine.begin.return_value.__enter__.return_value = conn

        with patch.object(service, "get_engine", return_value=engine):
            payload = service.update_knowledge_image_payload(
                "abc123",
                {"caption": "人工说明", "image_type": "diagram", "summary": "更新摘要", "tags": ["A", "B"]},
            )

        self.assertEqual(payload["caption_status"], "manual")
        self.assertEqual(payload["tags"], ["A", "B"])
        self.assertEqual(payload["preview_url"], "/api/extracted-images/a.png")
        update_params = conn.execute.call_args_list[2].args[1]
        self.assertEqual(update_params["caption_status"], "manual")
        self.assertEqual(json.loads(update_params["tags_json"]), ["A", "B"])

    def test_update_knowledge_image_rejects_non_list_tags(self) -> None:
        now = datetime(2026, 6, 1, tzinfo=timezone.utc)
        conn = Mock()
        assets_exists = Mock()
        assets_exists.scalar_one.return_value = True
        select_result = Mock()
        select_result.mappings.return_value.first.return_value = {
            "image_hash": "abc123",
            "placeholder": "__PRO_IMG_abc123__",
            "source_doc": "doc-a",
            "source_page": 1,
            "caption": "",
            "image_type": "",
            "summary": "",
            "tags_json": "[]",
            "caption_status": "pending",
            "created_at": now,
        }
        conn.execute.side_effect = [assets_exists, select_result]
        engine = MagicMock()
        engine.begin.return_value.__enter__.return_value = conn

        with patch.object(service, "get_engine", return_value=engine):
            with self.assertRaises(Exception) as ctx:
                service.update_knowledge_image_payload("abc123", {"tags": "bad"})

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_register_knowledge_image_asset_native_persists_registry_and_asset(self) -> None:
        execute_calls = []

        class FakeResult:
            def __init__(self, scalar=None, first=None):
                self._scalar = scalar
                self._first = first

            def scalar_one(self):
                return self._scalar

            def first(self):
                return self._first

        class FakeConn:
            def execute(self, statement, params=None):
                execute_calls.append((str(statement), params or {}))
                sql = str(statement)
                if "to_regclass" in sql:
                    return FakeResult(scalar=True)
                if "SELECT 1 FROM bid_generator.image_registry" in sql:
                    return FakeResult(first=None)
                if "SELECT 1 FROM bid_generator.knowledge_image_assets" in sql:
                    return FakeResult(first=None)
                return FakeResult()

        engine = MagicMock()
        engine.begin.return_value.__enter__.return_value = FakeConn()

        with (
            tempfile.TemporaryDirectory() as tmp_dir,
            patch.object(service, "_bid_generator_legacy_root", return_value=Path(tmp_dir)),
            patch.object(service, "get_engine", return_value=engine),
            patch.object(
                service,
                "_tag_image_with_vlm_native",
                return_value='{"caption":"系统架构图","image_type":"系统架构图","summary":"架构说明","tags":["架构"]}',
            ),
        ):
            placeholder, caption, image_info = service._register_knowledge_image_asset_native(
                filename="demo.docx",
                image_bytes=b"fake-image-bytes",
                original_name="image.png",
            )

        self.assertTrue(placeholder.startswith("__PRO_IMG_"))
        self.assertEqual(caption, "系统架构图")
        self.assertEqual(image_info["description"], "系统架构图")
        self.assertIn("knowledge_block", image_info)
        insert_registry = [params for sql, params in execute_calls if "INSERT INTO bid_generator.image_registry" in sql][0]
        insert_asset = [params for sql, params in execute_calls if "INSERT INTO bid_generator.knowledge_image_assets" in sql][0]
        self.assertEqual(insert_registry["placeholder"], placeholder)
        self.assertEqual(insert_asset["source_doc"], "demo.docx")
        self.assertEqual(insert_asset["caption_status"], "captioned")

    def test_extract_docx_with_tables_native_registers_embedded_images(self) -> None:
        import base64
        import docx

        image_bytes = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAIAAACQd1PeAAAADElEQVR4nGP4z8AAAAMBAQDJ/pLvAAAAAElFTkSuQmCC"
        )
        image_buffer = io.BytesIO(image_bytes)

        document = docx.Document()
        document.add_paragraph("正文段落")
        document.add_picture(image_buffer)
        docx_buffer = io.BytesIO()
        document.save(docx_buffer)

        with patch.object(
            service,
            "_register_knowledge_image_asset_native",
            return_value=("__PRO_IMG_hash__", "配图说明", {"preview_url": "/api/extracted-images/hash.png"}),
        ) as register:
            text_value, image_map = service._extract_docx_with_tables_native(
                docx_buffer.getvalue(),
                filename="demo.docx",
                extract_images=True,
            )

        self.assertIn("正文段落", text_value)
        self.assertIn("![配图说明](__PRO_IMG_hash__)", text_value)
        self.assertEqual(image_map["__PRO_IMG_hash__"]["preview_url"], "/api/extracted-images/hash.png")
        register.assert_called_once()

    def test_list_kb_sync_jobs_reads_recent_status_files(self) -> None:
        files = [
            _fake_status_file("new.json", 30, {"job_id": "new", "status": "completed", "started_at": "2026-06-01", "total": 3, "processed": 3, "failed": 0}),
            _fake_status_file("old.json", 10, {"job_id": "old", "status": "failed", "started_at": "2026-05-31", "total": 2, "processed": 1, "failed": 1}),
            _fake_status_file("bad.json", 20, "not-json", bad_json=True),
        ]
        status_dir = Mock()
        status_dir.exists.return_value = True
        status_dir.glob.return_value = files

        with patch.object(service, "_kb_sync_status_dir", return_value=status_dir):
            payload = service.list_kb_sync_jobs_payload()

        self.assertEqual([job["job_id"] for job in payload["jobs"]], ["new", "old"])
        self.assertEqual(payload["jobs"][0]["processed"], 3)

    def test_list_kb_sync_jobs_returns_empty_when_dir_missing(self) -> None:
        status_dir = Mock()
        status_dir.exists.return_value = False

        with patch.object(service, "_kb_sync_status_dir", return_value=status_dir):
            payload = service.list_kb_sync_jobs_payload()

        self.assertEqual(payload, {"jobs": []})

    def test_get_kb_sync_status_reads_status_file_when_task_missing(self) -> None:
        path = Mock()
        path.exists.return_value = True
        manager = MagicMock()
        manager.__enter__.return_value.read.return_value = json.dumps(
            {"job_id": "abcdef123456", "status": "completed", "processed": 3},
            ensure_ascii=False,
        )
        path.open.return_value = manager

        with (
            patch.object(service, "_get_task", return_value=None),
            patch.object(service, "_kb_sync_status_path", return_value=path),
        ):
            payload = service.get_kb_sync_status_payload("abcdef123456")

        self.assertEqual(payload, {"job_id": "abcdef123456", "status": "completed", "processed": 3})

    def test_get_kb_sync_status_merges_running_task_with_status_file(self) -> None:
        path = Mock()
        path.exists.return_value = True
        manager = MagicMock()
        manager.__enter__.return_value.read.return_value = json.dumps(
            {"job_id": "abcdef123456", "status": "completed", "processed": 2},
            ensure_ascii=False,
        )
        path.open.return_value = manager
        task = SimpleNamespace(status="running", created_at=1780296000, error="")

        with (
            patch.object(service, "_get_task", return_value=task),
            patch.object(service, "_kb_sync_status_path", return_value=path),
        ):
            payload = service.get_kb_sync_status_payload("abcdef123456")

        self.assertEqual(payload["job_id"], "abcdef123456")
        self.assertEqual(payload["status"], "running")
        self.assertEqual(payload["processed"], 2)
        self.assertEqual(payload["task_id"], "abcdef123456")

    def test_get_kb_sync_status_returns_task_fallback_when_file_missing(self) -> None:
        path = Mock()
        path.exists.return_value = False
        task = SimpleNamespace(status="error", created_at=1780296000, error="sync failed")

        with (
            patch.object(service, "_get_task", return_value=task),
            patch.object(service, "_kb_sync_status_path", return_value=path),
        ):
            payload = service.get_kb_sync_status_payload("abcdef123456")

        self.assertEqual(payload["status"], "failed")
        self.assertEqual(payload["failed"], 1)
        self.assertEqual(payload["error"], "sync failed")

    def test_get_kb_sync_status_rejects_invalid_job_id(self) -> None:
        with self.assertRaises(Exception) as ctx:
            service.get_kb_sync_status_payload("bad-job")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 400)

    def test_get_kb_sync_status_returns_404_when_missing(self) -> None:
        path = Mock()
        path.exists.return_value = False

        with (
            patch.object(service, "_get_task", return_value=None),
            patch.object(service, "_kb_sync_status_path", return_value=path),
        ):
            with self.assertRaises(Exception) as ctx:
                service.get_kb_sync_status_payload("abcdef123456")

        self.assertEqual(getattr(ctx.exception, "status_code", None), 404)

    def test_get_knowledge_documents_returns_legacy_error_when_dify_env_missing(self) -> None:
        with patch.dict(service.os.environ, {"DIFY_DATASET_ID": "", "DIFY_DATASET_KEY": ""}, clear=False):
            payload = _run_async(service.get_knowledge_documents_payload())

        self.assertEqual(
            payload,
            {
                "dataset_info": {"error": "DIFY_DATASET_ID or KEY not configured in backend."},
                "documents": [],
            },
        )

    def test_get_knowledge_documents_maps_dify_documents(self) -> None:
        class FakeResponse:
            def raise_for_status(self) -> None:
                return None

            def json(self) -> dict:
                return {
                    "data": [
                        {
                            "id": "doc-1",
                            "name": "标书A",
                            "word_count": 2048,
                            "tokens": 7,
                            "created_at": 1780296000,
                            "indexing_status": "completed",
                        },
                        {
                            "id": "doc-2",
                            "name": "标书B",
                            "word_count": 1000,
                            "tokens": 0,
                            "created_at": "bad",
                            "indexing_status": "error",
                        },
                        {
                            "id": "doc-3",
                            "name": "标书C",
                            "word_count": 100,
                            "indexing_status": "waiting",
                        },
                    ]
                }

        class FakeClient:
            async def get(self, url: str, *, headers: dict[str, str]) -> FakeResponse:
                self.url = url
                self.headers = headers
                return FakeResponse()

        fake_client = FakeClient()

        @asynccontextmanager
        async def fake_async_client(*, timeout: int):
            self.assertEqual(timeout, 10)
            yield fake_client

        with (
            patch.dict(
                service.os.environ,
                {
                    "DIFY_API_URL": "http://dify.local/v1",
                    "DIFY_DATASET_ID": "dataset-1",
                    "DIFY_DATASET_KEY": "key-1",
                },
                clear=False,
            ),
            patch.object(service.httpx, "AsyncClient", fake_async_client),
        ):
            payload = _run_async(service.get_knowledge_documents_payload())

        self.assertEqual(payload["dataset_info"], {"status": "connected", "dataset_id": "dataset-1"})
        self.assertEqual(fake_client.url, "http://dify.local/v1/datasets/dataset-1/documents")
        self.assertEqual(fake_client.headers, {"Authorization": "Bearer key-1"})
        expected_upload_time = service.datetime.fromtimestamp(1780296000).strftime("%Y-%m-%d %H:%M")
        self.assertEqual(
            payload["documents"],
            [
                {
                    "id": "doc-1",
                    "name": "标书A",
                    "size": "4.0 KB",
                    "uploadTime": expected_upload_time,
                    "status": "success",
                    "chunks": 7,
                },
                {
                    "id": "doc-2",
                    "name": "标书B",
                    "size": "2.0 KB",
                    "uploadTime": "-",
                    "status": "failed",
                    "chunks": 2,
                },
                {
                    "id": "doc-3",
                    "name": "标书C",
                    "size": "0.2 KB",
                    "uploadTime": "-",
                    "status": "indexing",
                    "chunks": 0,
                },
            ],
        )

    def test_get_knowledge_documents_returns_error_payload_on_http_failure(self) -> None:
        class FakeClient:
            async def get(self, url: str, *, headers: dict[str, str]) -> object:
                raise RuntimeError("dify unavailable")

        @asynccontextmanager
        async def fake_async_client(*, timeout: int):
            yield FakeClient()

        with (
            patch.dict(
                service.os.environ,
                {
                    "DIFY_API_URL": "http://dify.local/v1",
                    "DIFY_DATASET_ID": "dataset-1",
                    "DIFY_DATASET_KEY": "key-1",
                },
                clear=False,
            ),
            patch.object(service.httpx, "AsyncClient", fake_async_client),
        ):
            payload = _run_async(service.get_knowledge_documents_payload())

        self.assertEqual(payload["dataset_info"]["status"], "error")
        self.assertIn("dify unavailable", payload["dataset_info"]["message"])
        self.assertEqual(payload["documents"], [])


def _project_row(
    *,
    project_id: str = "proj-1",
    name: str = "项目一",
    status: str = "uploaded",
    data: dict,
) -> dict:
    now = datetime(2026, 6, 1, tzinfo=timezone.utc)
    return {
        "id": project_id,
        "name": name,
        "status": status,
        "data": json.dumps(data, ensure_ascii=False),
        "created_at": now,
        "updated_at": now,
    }


def _fake_status_file(name: str, mtime: int, payload: object, *, bad_json: bool = False) -> Mock:
    path = Mock()
    path.name = name
    path.stat.return_value.st_mtime = mtime
    manager = MagicMock()
    if bad_json:
        manager.__enter__.return_value.read.return_value = "{"
    else:
        manager.__enter__.return_value.read.return_value = json.dumps(payload, ensure_ascii=False)
    path.open.return_value = manager
    return path


def _run_async(awaitable):
    import asyncio

    return asyncio.run(awaitable)


async def _async_value(value):
    return value


async def _async_false():
    return False


async def _await_task(task):
    return await task


async def _collect_streaming_response_text(response):
    chunks = []
    async for chunk in response.body_iterator:
        if isinstance(chunk, bytes):
            chunks.append(chunk.decode("utf-8"))
        else:
            chunks.append(str(chunk))
    return "".join(chunks)


def _analysis_framework_nodes():
    return [
        {
            "id": "group-a",
            "label": "分组A",
            "children": [
                {"id": "proj_overview", "label": "项目概述", "extractionPrompt": "提取项目概述"},
                {"id": "proj_basic", "label": "项目基本信息", "extractionPrompt": "提取项目基本信息"},
                {"id": "scoring_details", "label": "评分细则", "extractionPrompt": "提取评分细则"},
                {"id": "structure_attachments", "label": "附件结构", "extractionPrompt": "提取附件结构"},
            ],
        },
        {
            "id": "group-b",
            "label": "分组B",
            "children": [
                {"id": "resp_tech", "label": "技术目标", "extractionPrompt": "提取技术目标"},
                {"id": "resp_param", "label": "参数要求", "extractionPrompt": "提取参数要求"},
                {"id": "resp_substance", "label": "实施要求", "extractionPrompt": "提取实施要求"},
            ],
        },
    ]


def _doc_blocks_fixture():
    return [
        {"block_id": "B000001", "locator": "P0001", "body_idx": 1, "text": "投标函", "type": "paragraph"},
        {"block_id": "B000002", "locator": "P0002", "body_idx": 2, "text": "法定代表人身份证明", "type": "paragraph"},
    ]


if __name__ == "__main__":
    unittest.main()
