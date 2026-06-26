from __future__ import annotations

import asyncio
import hashlib
import io
import json
import logging
import time
from pathlib import Path
from typing import Any
from urllib.parse import quote
from uuid import UUID

import httpx
from fastapi import UploadFile
from fastapi.responses import Response
from sqlalchemy import text
from sqlalchemy.exc import SQLAlchemyError

from app.core.errors import PlatformError
from app.services.pipt_recognition_adapter import recognize_with_platform_recognizer
from app.services.pipt_gateway_service import DEFAULT_TARGET_ENTITIES, preprocess_payload
from app.services.rag_dify_service import (
    get_dataset_api_base_url,
    get_dataset_api_key,
    get_default_dataset_id,
    get_desensitized_dataset_id,
    get_raw_dataset_id,
)
from packages.py_common.db.session import get_engine


logger = logging.getLogger(__name__)


class RagKnowledgeError(Exception):
    def __init__(self, detail: str, *, status_code: int = 502) -> None:
        self.detail = detail
        self.status_code = status_code
        super().__init__(detail)


def _database_error(exc: Exception) -> PlatformError:
    logger.exception("RAG knowledge PostgreSQL operation failed")
    return PlatformError(
        code="DATABASE_ERROR",
        message="知识库本地资料数据库访问失败。",
        status_code=500,
        details={"module": "rag-web-search", "schema": "rag"},
    )


def _json_dumps(data: Any) -> str:
    return json.dumps(data, ensure_ascii=False, separators=(",", ":"))


def _json_value(value: Any, fallback: Any) -> Any:
    if value is None:
        return fallback
    if isinstance(value, (dict, list)):
        return value
    if isinstance(value, str):
        try:
            return json.loads(value)
        except json.JSONDecodeError:
            return fallback
    return fallback


def _hash_text(value: str) -> str:
    return hashlib.sha256(str(value or "").encode("utf-8")).hexdigest()


def _ensure_knowledge_storage() -> None:
    try:
        with get_engine().begin() as conn:
            conn.execute(text("CREATE SCHEMA IF NOT EXISTS rag"))
            conn.execute(
                text(
                    """
                    CREATE TABLE IF NOT EXISTS rag.knowledge_documents (
                      id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
                      name TEXT NOT NULL,
                      source_type TEXT NOT NULL DEFAULT 'file',
                      original_content BYTEA NULL,
                      content_text TEXT NOT NULL DEFAULT '',
                      content_hash CHAR(64) NOT NULL DEFAULT repeat('0', 64),
                      mime_type TEXT NULL,
                      file_size BIGINT NULL,
                      parse_status TEXT NOT NULL DEFAULT 'pending',
                      privacy_status TEXT NOT NULL DEFAULT 'pending',
                      has_sensitive BOOLEAN NOT NULL DEFAULT FALSE,
                      sensitive_count INTEGER NOT NULL DEFAULT 0,
                      sensitive_types JSONB NOT NULL DEFAULT '[]'::jsonb,
                      recognition_summary JSONB NOT NULL DEFAULT '{}'::jsonb,
                      sync_status TEXT NOT NULL DEFAULT 'pending',
                      dify_document_id TEXT NULL,
                      dify_batch TEXT NULL,
                      pipt_request_id TEXT NULL,
                      pipt_mapping_count INTEGER NOT NULL DEFAULT 0,
                      last_error TEXT NULL,
                      created_at TIMESTAMPTZ NOT NULL DEFAULT now(),
                      updated_at TIMESTAMPTZ NOT NULL DEFAULT now(),
                      parsed_at TIMESTAMPTZ NULL,
                      synced_at TIMESTAMPTZ NULL
                    )
                    """
                )
            )
            conn.execute(text("ALTER TABLE rag.knowledge_documents ADD COLUMN IF NOT EXISTS original_content BYTEA NULL"))
            conn.execute(text("ALTER TABLE rag.knowledge_documents ALTER COLUMN content_text SET DEFAULT ''"))
            conn.execute(text("ALTER TABLE rag.knowledge_documents ALTER COLUMN content_hash SET DEFAULT repeat('0', 64)"))
            conn.execute(text("ALTER TABLE rag.knowledge_documents ADD COLUMN IF NOT EXISTS parse_status TEXT NOT NULL DEFAULT 'pending'"))
            conn.execute(text("ALTER TABLE rag.knowledge_documents ADD COLUMN IF NOT EXISTS parsed_at TIMESTAMPTZ NULL"))
            conn.execute(text("ALTER TABLE rag.knowledge_documents ADD COLUMN IF NOT EXISTS raw_dify_document_id TEXT NULL"))
            conn.execute(text("ALTER TABLE rag.knowledge_documents ADD COLUMN IF NOT EXISTS raw_dify_batch TEXT NULL"))
            conn.execute(text("ALTER TABLE rag.knowledge_documents ADD COLUMN IF NOT EXISTS raw_sync_status TEXT NOT NULL DEFAULT 'pending'"))
            conn.execute(text("ALTER TABLE rag.knowledge_documents ADD COLUMN IF NOT EXISTS desensitized_dify_document_id TEXT NULL"))
            conn.execute(text("ALTER TABLE rag.knowledge_documents ADD COLUMN IF NOT EXISTS desensitized_dify_batch TEXT NULL"))
            conn.execute(text("ALTER TABLE rag.knowledge_documents ADD COLUMN IF NOT EXISTS desensitized_sync_status TEXT NOT NULL DEFAULT 'pending'"))
            conn.execute(
                text(
                    """
                    UPDATE rag.knowledge_documents
                    SET desensitized_dify_document_id = COALESCE(desensitized_dify_document_id, dify_document_id),
                        desensitized_dify_batch = COALESCE(desensitized_dify_batch, dify_batch),
                        desensitized_sync_status = CASE
                          WHEN dify_document_id IS NOT NULL AND sync_status = 'synced' THEN 'synced'
                          ELSE desensitized_sync_status
                        END
                    WHERE dify_document_id IS NOT NULL
                    """
                )
            )
            conn.execute(
                text(
                    """
                    CREATE INDEX IF NOT EXISTS idx_rag_knowledge_documents_updated_at
                      ON rag.knowledge_documents(updated_at DESC)
                    """
                )
            )
            conn.execute(
                text(
                    """
                    CREATE INDEX IF NOT EXISTS idx_rag_knowledge_documents_sync_status
                      ON rag.knowledge_documents(sync_status)
                    """
                )
            )
            conn.execute(
                text(
                    """
                    CREATE INDEX IF NOT EXISTS idx_rag_knowledge_documents_has_sensitive
                      ON rag.knowledge_documents(has_sensitive)
                    """
                )
            )
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc


def _safe_document_name(name: str | None, fallback: str = "知识库资料") -> str:
    value = Path(str(name or fallback)).name.strip()
    return value or fallback


def _normalize_entity(entity: Any) -> dict[str, Any]:
    if hasattr(entity, "model_dump"):
        raw = entity.model_dump()
    elif isinstance(entity, dict):
        raw = entity
    else:
        raw = {
            "text": getattr(entity, "text", ""),
            "entity_type": getattr(entity, "entity_type", "unknown"),
            "start": getattr(entity, "start", None),
            "end": getattr(entity, "end", None),
            "source": getattr(entity, "source", "unknown"),
            "confidence": getattr(entity, "confidence", 0),
            "reason": getattr(entity, "reason", ""),
        }
    return {
        "entity_type": str(raw.get("entity_type") or "unknown"),
        "start": raw.get("start"),
        "end": raw.get("end"),
        "source": str(raw.get("source") or "unknown"),
        "confidence": raw.get("confidence") or 0,
        "reason": str(raw.get("reason") or ""),
    }


def _recognize_privacy(text_value: str) -> dict[str, Any]:
    entities = recognize_with_platform_recognizer(
        text=text_value,
        target_entities=DEFAULT_TARGET_ENTITIES,
        llm_mode="verify_only",
    )
    normalized = [_normalize_entity(entity) for entity in entities]
    counts_by_type: dict[str, int] = {}
    for item in normalized:
        entity_type = str(item.get("entity_type") or "unknown")
        counts_by_type[entity_type] = counts_by_type.get(entity_type, 0) + 1
    return {
        "status": "recognized",
        "has_sensitive": bool(normalized),
        "sensitive_count": len(normalized),
        "sensitive_types": sorted(counts_by_type),
        "summary": {
            "counts_by_type": counts_by_type,
            "entities": normalized[:50],
            "truncated": len(normalized) > 50,
        },
    }


def _recognition_from_preprocess_result(result: dict[str, Any]) -> dict[str, Any]:
    manifest = result.get("placeholder_manifest")
    if not isinstance(manifest, dict):
        manifest = {}
    counts_by_type: dict[str, int] = {}
    for meta in manifest.values():
        if not isinstance(meta, dict):
            continue
        entity_type = str(meta.get("entity_type") or "unknown")
        counts_by_type[entity_type] = counts_by_type.get(entity_type, 0) + 1
    sensitive_count = int(result.get("mapping_table_count") or sum(counts_by_type.values()) or 0)
    return {
        "status": "recognized",
        "has_sensitive": sensitive_count > 0,
        "sensitive_count": sensitive_count,
        "sensitive_types": sorted(counts_by_type),
        "summary": {
            "counts_by_type": counts_by_type,
            "truncated": False,
        },
    }


def _decode_text_bytes(content: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover - 环境依赖缺失时给出业务错误
        raise RagKnowledgeError("当前环境缺少 python-docx，无法解析 DOCX。", status_code=503) from exc

    document = Document(io.BytesIO(content))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _extract_pdf_text(content: bytes) -> str:
    try:
        import fitz
    except Exception as exc:  # pragma: no cover - 环境依赖缺失时给出业务错误
        raise RagKnowledgeError("当前环境缺少 PyMuPDF，无法解析 PDF。", status_code=503) from exc

    parts: list[str] = []
    with fitz.open(stream=content, filetype="pdf") as document:
        for page in document:
            text_value = page.get_text("text").strip()
            if text_value:
                parts.append(text_value)
    return "\n\n".join(parts)


def _extract_upload_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix in {".txt", ".md", ".markdown", ".csv", ".json", ".html", ".htm"}:
        return _decode_text_bytes(content)
    if suffix == ".docx":
        return _extract_docx_text(content)
    if suffix == ".pdf":
        return _extract_pdf_text(content)
    raise RagKnowledgeError("暂不支持该资料格式，请上传 txt、md、docx 或 pdf。", status_code=415)


async def _read_upload_content(file: UploadFile) -> tuple[str, str, bytes]:
    filename = _safe_document_name(file.filename, "upload.txt")
    mime = file.content_type or "application/octet-stream"
    content = await file.read()
    if not content:
        raise RagKnowledgeError("Empty file upload", status_code=400)
    return filename, mime, content


def _row_to_local_document(row: Any) -> dict[str, Any]:
    updated_at = row.get("updated_at")
    created_at = row.get("created_at")
    parsed_at = row.get("parsed_at")
    synced_at = row.get("synced_at")
    sensitive_types = _json_value(row.get("sensitive_types"), [])
    summary = _json_value(row.get("recognition_summary"), {})
    return {
        "id": str(row["id"]),
        "name": str(row["name"]),
        "description": row.get("last_error"),
        "display_status": row.get("sync_status"),
        "indexing_status": row.get("sync_status"),
        "data_source_type": "local_original",
        "word_count": len(str(row.get("content_text") or "")),
        "tokens": None,
        "segment_count": None,
        "enabled": row.get("sync_status") == "synced",
        "created_at": int(created_at.timestamp()) if hasattr(created_at, "timestamp") else None,
        "updated_at": int(updated_at.timestamp()) if hasattr(updated_at, "timestamp") else None,
        "source_type": row.get("source_type"),
        "mime_type": row.get("mime_type"),
        "file_size": row.get("file_size"),
        "parse_status": row.get("parse_status"),
        "privacy_status": row.get("privacy_status"),
        "has_sensitive": bool(row.get("has_sensitive")),
        "sensitive_count": int(row.get("sensitive_count") or 0),
        "sensitive_types": sensitive_types if isinstance(sensitive_types, list) else [],
        "recognition_summary": summary if isinstance(summary, dict) else {},
        "sync_status": row.get("sync_status"),
        "dify_document_id": row.get("dify_document_id"),
        "raw_dify_document_id": row.get("raw_dify_document_id"),
        "raw_dify_batch": row.get("raw_dify_batch"),
        "raw_sync_status": row.get("raw_sync_status"),
        "desensitized_dify_document_id": row.get("desensitized_dify_document_id"),
        "desensitized_dify_batch": row.get("desensitized_dify_batch"),
        "desensitized_sync_status": row.get("desensitized_sync_status"),
        "pipt_request_id": row.get("pipt_request_id"),
        "pipt_mapping_count": int(row.get("pipt_mapping_count") or 0),
        "last_error": row.get("last_error"),
        "parsed_at": int(parsed_at.timestamp()) if hasattr(parsed_at, "timestamp") else None,
        "synced_at": int(synced_at.timestamp()) if hasattr(synced_at, "timestamp") else None,
    }


def _get_local_document_row(document_id: str) -> dict[str, Any] | None:
    try:
        UUID(str(document_id))
    except ValueError:
        return None
    _ensure_knowledge_storage()
    try:
        with get_engine().begin() as conn:
            row = conn.execute(
                text(
                    """
                    SELECT *
                    FROM rag.knowledge_documents
                    WHERE id = CAST(:document_id AS uuid)
                    """
                ),
                {"document_id": document_id},
            ).mappings().first()
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc
    return dict(row) if row else None


def _ensure_local_document_parsed(row: dict[str, Any]) -> dict[str, Any]:
    if str(row.get("parse_status") or "parsed") == "parsed" and str(row.get("content_text") or "").strip():
        return row

    local_id = str(row["id"])
    try:
        content_text = ""
        if row.get("source_type") == "file":
            raw_content = row.get("original_content")
            if isinstance(raw_content, memoryview):
                raw_bytes = raw_content.tobytes()
            elif isinstance(raw_content, bytes):
                raw_bytes = raw_content
            else:
                raw_bytes = b""
            content_text = _extract_upload_text(str(row.get("name") or "upload.txt"), raw_bytes).strip()
        else:
            content_text = str(row.get("content_text") or "").strip()
        if not content_text:
            raise RagKnowledgeError("资料没有可解析文本内容。", status_code=400)

        with get_engine().begin() as conn:
            updated = conn.execute(
                text(
                    """
                    UPDATE rag.knowledge_documents
                    SET content_text = :content_text,
                        content_hash = :content_hash,
                        parse_status = 'parsed',
                        parsed_at = now(),
                        last_error = NULL,
                        updated_at = now()
                    WHERE id = CAST(:document_id AS uuid)
                    RETURNING *
                    """
                ),
                {
                    "document_id": local_id,
                    "content_text": content_text,
                    "content_hash": _hash_text(content_text),
                },
            ).mappings().one()
        return dict(updated)
    except Exception as exc:
        detail = getattr(exc, "detail", None) or getattr(exc, "message", None) or str(exc) or "资料解析失败"
        try:
            with get_engine().begin() as conn:
                conn.execute(
                    text(
                        """
                        UPDATE rag.knowledge_documents
                        SET parse_status = 'failed',
                            sync_status = 'failed',
                            raw_sync_status = 'failed',
                            desensitized_sync_status = 'failed',
                            last_error = :last_error,
                            updated_at = now()
                        WHERE id = CAST(:document_id AS uuid)
                        """
                    ),
                    {"document_id": local_id, "last_error": str(detail)[:1000]},
                )
        except Exception:
            logger.exception("Failed to persist knowledge parse error")
        if isinstance(exc, (RagKnowledgeError, PlatformError)):
            raise
        raise RagKnowledgeError(str(detail), status_code=500) from exc


def _dify_headers() -> dict[str, str]:
    key = get_dataset_api_key().strip()
    if not key:
        raise RagKnowledgeError("知识库服务配置错误：未设置 DIFY_DATASET_API_KEY。", status_code=503)
    return {
        "Authorization": f"Bearer {key}",
        "Accept": "application/json",
    }


def _dataset_id() -> str:
    dataset_id = get_default_dataset_id().strip()
    if not dataset_id:
        raise RagKnowledgeError("知识库服务配置错误：未设置 DIFY_DEFAULT_DATASET_ID。", status_code=503)
    return dataset_id


def _raw_dataset_id() -> str:
    dataset_id = get_raw_dataset_id().strip()
    if not dataset_id:
        raise RagKnowledgeError("知识库服务配置错误：未设置 DIFY_RAW_DATASET_ID。", status_code=503)
    return dataset_id


def _desensitized_dataset_id() -> str:
    dataset_id = get_desensitized_dataset_id().strip()
    if not dataset_id:
        raise RagKnowledgeError("知识库服务配置错误：未设置 DIFY_DESENSITIZED_DATASET_ID。", status_code=503)
    return dataset_id


def _extract_description(raw: dict[str, Any]) -> str | None:
    direct = raw.get("description")
    if isinstance(direct, str) and direct.strip():
        return direct.strip()
    meta = raw.get("doc_metadata")
    if isinstance(meta, list):
        for item in meta:
            if not isinstance(item, dict):
                continue
            name = (item.get("name") or "").strip().lower()
            if name in ("description", "描述", "简介", "summary"):
                value = item.get("value")
                if isinstance(value, str) and value.strip():
                    return value.strip()
        for item in meta:
            if isinstance(item, dict):
                value = item.get("value")
                if isinstance(value, str) and value.strip():
                    return value.strip()
    return None


def _kb_document_config() -> dict[str, Any]:
    return {
        "indexing_technique": "high_quality",
        "process_rule": {
            "rules": {
                "segmentation": {
                    "separator": "\n##",
                    "chunk_overlap": 150,
                    "max_tokens": 1024,
                }
            },
            "mode": "automatic",
        },
    }


def _create_by_text_payload(name: str, text: str) -> dict[str, Any]:
    return {"name": name, "text": text, **_kb_document_config()}


def _upstream_error_message(response: httpx.Response, fallback: str = "Upstream error") -> str:
    try:
        payload = response.json()
        if isinstance(payload, dict):
            value = payload.get("message") or payload.get("detail") or payload.get("error")
            if value:
                return str(value)[:500]
    except Exception:
        pass
    return (response.text or fallback)[:500]


async def _wait_indexing_complete(
    client: httpx.AsyncClient,
    *,
    base: str,
    dataset_id: str,
    batch: str,
    document_id: str,
    headers: dict[str, str],
    timeout_seconds: float = 180.0,
    poll_interval: float = 1.0,
) -> tuple[str, str | None]:
    url = f"{base}/datasets/{dataset_id}/documents/{batch}/indexing-status"
    deadline = time.monotonic() + timeout_seconds
    last_status = "unknown"

    while time.monotonic() < deadline:
        try:
            response = await client.get(url, headers=headers)
        except httpx.RequestError as exc:
            return "error", f"Upstream error: {exc}"

        if response.status_code >= 400:
            return "error", _upstream_error_message(response, "Indexing status request failed")

        payload = response.json()
        rows = payload.get("data") if isinstance(payload, dict) else []
        rows = rows if isinstance(rows, list) else []
        row = next((item for item in rows if isinstance(item, dict) and item.get("id") == document_id), None)
        if row is None and rows and isinstance(rows[0], dict):
            row = rows[0]
        if not isinstance(row, dict):
            await asyncio.sleep(poll_interval)
            continue

        last_status = str(row.get("indexing_status") or "")
        if last_status == "completed":
            return "completed", None
        if last_status == "error":
            error = row.get("error")
            return "error", str(error) if error else "Indexing failed"
        await asyncio.sleep(poll_interval)

    return last_status, "Timeout waiting for indexing to complete"


async def create_text_document_and_wait(name: str, text: str, *, dataset_id: str | None = None) -> dict[str, Any]:
    base = get_dataset_api_base_url()
    target_dataset_id = dataset_id or _dataset_id()
    headers = {**_dify_headers(), "Content-Type": "application/json"}
    payload = _create_by_text_payload(name.strip(), text)
    url = f"{base}/datasets/{target_dataset_id}/document/create-by-text"

    async with httpx.AsyncClient(timeout=httpx.Timeout(300.0), trust_env=False) as client:
        try:
            response = await client.post(url, headers=headers, json=payload)
        except httpx.RequestError as exc:
            raise RagKnowledgeError(f"Upstream error: {exc}") from exc

        if response.status_code >= 400:
            raise RagKnowledgeError(_upstream_error_message(response, "Create document failed"), status_code=response.status_code)

        created = response.json()
        document = created.get("document") if isinstance(created, dict) else None
        batch = created.get("batch") if isinstance(created, dict) else None
        if not isinstance(document, dict) or not isinstance(batch, str):
            raise RagKnowledgeError("Invalid upstream response")
        document_id = document.get("id")
        if not isinstance(document_id, str):
            raise RagKnowledgeError("Missing document id in upstream response")

        final_status, error = await _wait_indexing_complete(
            client,
            base=base,
            dataset_id=target_dataset_id,
            batch=batch,
            document_id=document_id,
            headers=_dify_headers(),
        )

    if final_status == "completed":
        return {
            "ok": True,
            "document_id": document_id,
            "name": document.get("name") if isinstance(document.get("name"), str) else name,
            "batch": batch,
            "indexing_status": "completed",
        }

    raise RagKnowledgeError(error or f"Document created but indexing did not complete (status={final_status})", status_code=500)


async def create_file_document_and_wait(file: UploadFile) -> dict[str, Any]:
    return await create_local_file_document(file)


async def create_local_file_document(file: UploadFile) -> dict[str, Any]:
    filename, mime, content = await _read_upload_content(file)
    _ensure_knowledge_storage()
    try:
        with get_engine().begin() as conn:
            row = conn.execute(
                text(
                    """
                    INSERT INTO rag.knowledge_documents (
                      name, source_type, original_content, content_text, content_hash,
                      mime_type, file_size, parse_status, privacy_status, has_sensitive,
                      sensitive_count, sensitive_types, recognition_summary, sync_status,
                      raw_sync_status, desensitized_sync_status, updated_at
                    )
                    VALUES (
                      :name, 'file', :original_content, '', repeat('0', 64),
                      :mime_type, :file_size, 'pending', :privacy_status, :has_sensitive, :sensitive_count,
                      CAST(:sensitive_types AS jsonb), CAST(:recognition_summary AS jsonb),
                      'pending', 'pending', 'pending', now()
                    )
                    RETURNING *
                    """
                ),
                {
                    "name": filename,
                    "original_content": content,
                    "mime_type": mime,
                    "file_size": len(content),
                    "privacy_status": "pending",
                    "has_sensitive": False,
                    "sensitive_count": 0,
                    "sensitive_types": _json_dumps([]),
                    "recognition_summary": _json_dumps({}),
                },
            ).mappings().one()
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc

    document = _row_to_local_document(row)
    return {
        "ok": True,
        "document_id": document["id"],
        "name": document["name"],
        "batch": "",
        "indexing_status": document["sync_status"],
        "document": document,
    }


async def create_local_text_document(name: str, text_value: str) -> dict[str, Any]:
    clean_name = _safe_document_name(name, "文本资料")
    clean_text = str(text_value or "").strip()
    if not clean_text:
        raise RagKnowledgeError("text 不能为空", status_code=422)

    _ensure_knowledge_storage()
    try:
        with get_engine().begin() as conn:
            row = conn.execute(
                text(
                    """
                    INSERT INTO rag.knowledge_documents (
                      name, source_type, content_text, content_hash, parse_status,
                      privacy_status, has_sensitive, sensitive_count, sensitive_types,
                      recognition_summary, sync_status, raw_sync_status, desensitized_sync_status,
                      parsed_at, updated_at
                    )
                    VALUES (
                      :name, 'text', :content_text, :content_hash, 'parsed',
                      :privacy_status, :has_sensitive, :sensitive_count,
                      CAST(:sensitive_types AS jsonb), CAST(:recognition_summary AS jsonb),
                      'pending', 'pending', 'pending', now(), now()
                    )
                    RETURNING *
                    """
                ),
                {
                    "name": clean_name,
                    "content_text": clean_text,
                    "content_hash": _hash_text(clean_text),
                    "privacy_status": "pending",
                    "has_sensitive": False,
                    "sensitive_count": 0,
                    "sensitive_types": _json_dumps([]),
                    "recognition_summary": _json_dumps({}),
                },
            ).mappings().one()
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc

    document = _row_to_local_document(row)
    return {
        "ok": True,
        "document_id": document["id"],
        "name": document["name"],
        "batch": "",
        "indexing_status": document["sync_status"],
        "document": document,
    }


async def list_knowledge_documents() -> dict[str, Any]:
    _ensure_knowledge_storage()
    try:
        with get_engine().begin() as conn:
            rows = conn.execute(
                text(
                    """
                    SELECT *
                    FROM rag.knowledge_documents
                    ORDER BY updated_at DESC, created_at DESC
                    """
                )
            ).mappings().all()
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc

    documents = [_row_to_local_document(row) for row in rows]
    return {"documents": documents, "total": len(documents)}


async def sync_knowledge_document_to_dify(document_id: str) -> dict[str, Any]:
    row = _get_local_document_row(document_id)
    if row is None:
        raise RagKnowledgeError("本地资料不存在。", status_code=404)
    if (
        row.get("sync_status") == "synced"
        and row.get("raw_dify_document_id")
        and (row.get("desensitized_dify_document_id") or row.get("dify_document_id"))
    ):
        return {
            "ok": True,
            "skipped": True,
            "document_id": str(row["id"]),
            "raw_dify_document_id": row.get("raw_dify_document_id"),
            "desensitized_dify_document_id": row.get("desensitized_dify_document_id") or row.get("dify_document_id"),
            "name": row.get("name"),
            "batch": row.get("desensitized_dify_batch") or row.get("dify_batch") or "",
            "indexing_status": "synced",
        }
    row = await asyncio.to_thread(_ensure_local_document_parsed, row)

    local_id = str(row["id"])
    try:
        with get_engine().begin() as conn:
            conn.execute(
                text(
                    """
                    UPDATE rag.knowledge_documents
                    SET sync_status = 'syncing',
                        raw_sync_status = CASE WHEN raw_dify_document_id IS NULL THEN 'syncing' ELSE raw_sync_status END,
                        desensitized_sync_status = CASE
                          WHEN desensitized_dify_document_id IS NULL AND dify_document_id IS NULL THEN 'syncing'
                          ELSE desensitized_sync_status
                        END,
                        last_error = NULL,
                        updated_at = now()
                    WHERE id = CAST(:document_id AS uuid)
                    """
                ),
                {"document_id": local_id},
            )
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc

    try:
        source_text = str(row.get("content_text") or "")
        raw_created = None
        raw_document_id = str(row.get("raw_dify_document_id") or "")
        if not raw_document_id:
            raw_created = await create_text_document_and_wait(
                str(row.get("name") or "知识库资料"),
                source_text,
                dataset_id=_raw_dataset_id(),
            )
            raw_document_id = str(raw_created.get("document_id") or "")
            if not raw_document_id:
                raise RagKnowledgeError("原文知识库未返回文档 ID。", status_code=502)
            with get_engine().begin() as conn:
                conn.execute(
                    text(
                        """
                        UPDATE rag.knowledge_documents
                        SET raw_dify_document_id = :raw_dify_document_id,
                            raw_dify_batch = :raw_dify_batch,
                            raw_sync_status = 'synced',
                            updated_at = now()
                        WHERE id = CAST(:document_id AS uuid)
                        """
                    ),
                    {
                        "document_id": local_id,
                        "raw_dify_document_id": raw_document_id,
                        "raw_dify_batch": raw_created.get("batch") or "",
                    },
                )

        preprocess_result = await asyncio.to_thread(
            preprocess_payload,
            {
                "text": source_text,
                "module_code": "knowledge-base",
                "purpose": "knowledge_sync",
                "mode": "strong",
                "enabled": True,
                "target_entities": DEFAULT_TARGET_ENTITIES,
                "llm_mode": "verify_only",
            },
        )
        safe_text = str(preprocess_result.get("text") or "")
        if not safe_text:
            raise RagKnowledgeError("脱敏结果为空，已停止同步。", status_code=500)
        recognition = _recognition_from_preprocess_result(preprocess_result)

        upload_name = f"{Path(str(row.get('name') or '知识库资料')).stem}（脱敏）"
        created = await create_text_document_and_wait(upload_name, safe_text, dataset_id=_desensitized_dataset_id())
        dify_document_id = str(created.get("document_id") or "")
        if not dify_document_id:
            raise RagKnowledgeError("脱敏知识库未返回文档 ID。", status_code=502)

        try:
            with get_engine().begin() as conn:
                updated = conn.execute(
                    text(
                        """
                        UPDATE rag.knowledge_documents
                        SET sync_status = 'synced',
                            dify_document_id = :dify_document_id,
                            dify_batch = :dify_batch,
                            desensitized_dify_document_id = :dify_document_id,
                            desensitized_dify_batch = :dify_batch,
                            desensitized_sync_status = 'synced',
                            raw_sync_status = 'synced',
                            privacy_status = :privacy_status,
                            has_sensitive = :has_sensitive,
                            sensitive_count = :sensitive_count,
                            sensitive_types = CAST(:sensitive_types AS jsonb),
                            recognition_summary = CAST(:recognition_summary AS jsonb),
                            pipt_request_id = :pipt_request_id,
                            pipt_mapping_count = :pipt_mapping_count,
                            last_error = NULL,
                            synced_at = now(),
                            updated_at = now()
                        WHERE id = CAST(:document_id AS uuid)
                        RETURNING *
                        """
                    ),
                    {
                        "document_id": local_id,
                        "dify_document_id": dify_document_id,
                        "dify_batch": created.get("batch") or "",
                        "privacy_status": recognition["status"],
                        "has_sensitive": recognition["has_sensitive"],
                        "sensitive_count": recognition["sensitive_count"],
                        "sensitive_types": _json_dumps(recognition["sensitive_types"]),
                        "recognition_summary": _json_dumps(recognition["summary"]),
                        "pipt_request_id": preprocess_result.get("request_id") or "",
                        "pipt_mapping_count": int(preprocess_result.get("mapping_table_count") or 0),
                    },
                ).mappings().one()
        except (SQLAlchemyError, RuntimeError) as exc:
            raise _database_error(exc) from exc

        return {
            **created,
            "document_id": local_id,
            "raw_dify_document_id": raw_document_id,
            "desensitized_dify_document_id": dify_document_id,
            "dify_document_id": dify_document_id,
            "desensitized": True,
            "mapping_table_count": int(preprocess_result.get("mapping_table_count") or 0),
            "request_id": preprocess_result.get("request_id"),
            "document": _row_to_local_document(updated),
        }
    except Exception as exc:
        detail = getattr(exc, "detail", None) or getattr(exc, "message", None) or str(exc) or "脱敏同步失败"
        try:
            with get_engine().begin() as conn:
                conn.execute(
                    text(
                        """
                        UPDATE rag.knowledge_documents
                        SET sync_status = 'failed',
                            raw_sync_status = CASE WHEN raw_dify_document_id IS NULL THEN 'failed' ELSE raw_sync_status END,
                            desensitized_sync_status = CASE
                              WHEN desensitized_dify_document_id IS NULL AND dify_document_id IS NULL THEN 'failed'
                              ELSE desensitized_sync_status
                            END,
                            last_error = :last_error,
                            updated_at = now()
                        WHERE id = CAST(:document_id AS uuid)
                        """
                    ),
                    {"document_id": local_id, "last_error": str(detail)[:1000]},
                )
        except Exception:
            logger.exception("Failed to persist knowledge sync error")
        if isinstance(exc, (RagKnowledgeError, PlatformError)):
            raise
        raise RagKnowledgeError(str(detail), status_code=500) from exc


def _summarize_document_detail(raw: dict[str, Any]) -> dict[str, Any]:
    upload: dict[str, Any] | None = None
    data_source_detail = raw.get("data_source_detail_dict")
    if isinstance(data_source_detail, dict):
        upload_file = data_source_detail.get("upload_file")
        if isinstance(upload_file, dict):
            upload = {
                "name": upload_file.get("name"),
                "size": upload_file.get("size"),
                "extension": upload_file.get("extension"),
                "mime_type": upload_file.get("mime_type"),
            }
    return {
        "id": raw.get("id"),
        "name": raw.get("name"),
        "data_source_type": raw.get("data_source_type"),
        "created_from": raw.get("created_from"),
        "word_count": raw.get("word_count"),
        "tokens": raw.get("tokens"),
        "hit_count": raw.get("hit_count"),
        "indexing_status": raw.get("indexing_status"),
        "display_status": raw.get("display_status"),
        "doc_form": raw.get("doc_form"),
        "doc_language": raw.get("doc_language"),
        "segment_count": raw.get("segment_count"),
        "average_segment_length": raw.get("average_segment_length"),
        "indexing_latency": raw.get("indexing_latency"),
        "created_at": raw.get("created_at"),
        "updated_at": raw.get("updated_at"),
        "completed_at": raw.get("completed_at"),
        "doc_metadata": raw.get("doc_metadata"),
        "upload_file": upload,
        "enabled": raw.get("enabled"),
        "error": raw.get("error"),
    }


def _summarize_segment(segment: dict[str, Any]) -> dict[str, Any]:
    return {
        "id": segment.get("id"),
        "position": segment.get("position"),
        "content": segment.get("content") if isinstance(segment.get("content"), str) else "",
        "word_count": segment.get("word_count"),
        "tokens": segment.get("tokens"),
        "hit_count": segment.get("hit_count"),
        "status": segment.get("status"),
        "keywords": segment.get("keywords") if isinstance(segment.get("keywords"), list) else [],
    }


async def _get_dify_document_detail(document_id: str, *, dataset_id: str) -> dict[str, Any]:
    base = get_dataset_api_base_url()
    headers = _dify_headers()
    doc_url = f"{base}/datasets/{dataset_id}/documents/{document_id}"

    async with httpx.AsyncClient(timeout=httpx.Timeout(120.0), trust_env=False) as client:
        try:
            doc_response = await client.get(doc_url, headers=headers, params={"metadata": "all"})
        except httpx.RequestError as exc:
            raise RagKnowledgeError(f"Upstream error: {exc}") from exc

        if doc_response.status_code == 404:
            raise RagKnowledgeError("Document not found.", status_code=404)
        if doc_response.status_code >= 400:
            raise RagKnowledgeError(
                _upstream_error_message(doc_response, "Failed to load document"),
                status_code=doc_response.status_code,
            )

        document_raw = doc_response.json()
        if not isinstance(document_raw, dict):
            raise RagKnowledgeError("Invalid document response")

        segments: list[dict[str, Any]] = []
        page = 1
        limit = 100
        doc_form_from_segments: str | None = None
        while page <= 200:
            seg_url = f"{base}/datasets/{dataset_id}/documents/{document_id}/segments"
            try:
                seg_response = await client.get(seg_url, headers=headers, params={"page": page, "limit": limit})
            except httpx.RequestError as exc:
                raise RagKnowledgeError(f"Upstream error (segments): {exc}") from exc

            if seg_response.status_code >= 400:
                raise RagKnowledgeError(
                    _upstream_error_message(seg_response, "Failed to load segments"),
                    status_code=seg_response.status_code,
                )

            payload = seg_response.json()
            if isinstance(payload.get("doc_form"), str):
                doc_form_from_segments = payload["doc_form"]
            batch = payload.get("data") if isinstance(payload, dict) else []
            if isinstance(batch, list):
                segments.extend(_summarize_segment(item) for item in batch if isinstance(item, dict))
            if not bool(payload.get("has_more")):
                break
            page += 1

    segments.sort(key=lambda item: (item.get("position") is None, item.get("position") or 0))
    summary = _summarize_document_detail(document_raw)
    if doc_form_from_segments and not summary.get("doc_form"):
        summary["doc_form"] = doc_form_from_segments
    return {"document": summary, "segments": segments, "segment_total": len(segments)}


async def get_knowledge_document_detail(document_id: str) -> dict[str, Any]:
    local_row = _get_local_document_row(document_id)
    if local_row is not None:
        local_document = _row_to_local_document(local_row)
        dify_document_id = local_row.get("desensitized_dify_document_id") or local_row.get("dify_document_id")
        if dify_document_id and local_row.get("desensitized_sync_status") == "synced":
            payload = await _get_dify_document_detail(str(dify_document_id), dataset_id=_desensitized_dataset_id())
            document = payload.get("document") if isinstance(payload.get("document"), dict) else {}
            return {
                **payload,
                "document": {
                    **document,
                    "id": local_document["id"],
                    "name": local_document["name"],
                    "dify_document_id": dify_document_id,
                    "raw_dify_document_id": local_document.get("raw_dify_document_id"),
                    "desensitized_dify_document_id": local_document.get("desensitized_dify_document_id"),
                    "sync_status": local_document["sync_status"],
                    "raw_sync_status": local_document.get("raw_sync_status"),
                    "desensitized_sync_status": local_document.get("desensitized_sync_status"),
                    "privacy_status": local_document["privacy_status"],
                    "has_sensitive": local_document["has_sensitive"],
                    "sensitive_count": local_document["sensitive_count"],
                    "sensitive_types": local_document["sensitive_types"],
                    "pipt_request_id": local_document["pipt_request_id"],
                    "pipt_mapping_count": local_document["pipt_mapping_count"],
                },
            }

        return {
            "document": {
                "id": local_document["id"],
                "name": local_document["name"],
                "data_source_type": local_document["data_source_type"],
                "created_from": "local",
                "word_count": 0,
                "tokens": None,
                "hit_count": None,
                "indexing_status": local_document["sync_status"],
                "display_status": local_document["sync_status"],
                "doc_form": "text_model",
                "doc_language": "Chinese",
                "segment_count": 0,
                "average_segment_length": 0,
                "indexing_latency": None,
                "created_at": local_document["created_at"],
                "updated_at": local_document["updated_at"],
                "completed_at": local_document["synced_at"],
                "doc_metadata": {
                    "sync_status": local_document["sync_status"],
                    "privacy_status": local_document["privacy_status"],
                    "has_sensitive": local_document["has_sensitive"],
                    "sensitive_count": local_document["sensitive_count"],
                    "sensitive_types": local_document["sensitive_types"],
                    "recognition_summary": local_document["recognition_summary"],
                },
                "upload_file": {
                    "name": local_document["name"],
                    "size": local_document["file_size"],
                    "extension": Path(local_document["name"]).suffix,
                    "mime_type": local_document["mime_type"],
                },
                "enabled": local_document["enabled"],
                "error": local_document["last_error"],
                "sync_status": local_document["sync_status"],
                "privacy_status": local_document["privacy_status"],
                "has_sensitive": local_document["has_sensitive"],
                "sensitive_count": local_document["sensitive_count"],
                "sensitive_types": local_document["sensitive_types"],
                "pipt_request_id": local_document["pipt_request_id"],
                "pipt_mapping_count": local_document["pipt_mapping_count"],
            },
            "segments": [],
            "segment_total": 0,
        }
    return await _get_dify_document_detail(document_id, dataset_id=_dataset_id())


def _safe_export_filename(name: str | None, suffix: str) -> str:
    base = (name or "knowledge-document").strip()
    base = "".join(ch if ch.isalnum() or ch in ("-", "_", ".", " ") else "_" for ch in base)
    base = base.strip(" .") or "knowledge-document"
    if len(base) > 80:
        base = base[:80].rstrip(" .")
    return f"{base}.{suffix}"


def _content_disposition(filename: str) -> str:
    ascii_name = filename.encode("ascii", errors="ignore").decode("ascii") or "knowledge-document"
    return f"attachment; filename=\"{ascii_name}\"; filename*=UTF-8''{quote(filename)}"


def _document_export_markdown(payload: dict[str, Any]) -> str:
    document = payload.get("document") if isinstance(payload.get("document"), dict) else {}
    segments = payload.get("segments") if isinstance(payload.get("segments"), list) else []
    name = document.get("name") or document.get("id") or "knowledge-document"
    lines = [
        f"# {name}",
        "",
        f"- 文档 ID: {document.get('id') or ''}",
        f"- 索引状态: {document.get('indexing_status') or document.get('display_status') or ''}",
        f"- 分段数: {payload.get('segment_total') or len(segments)}",
        f"- Tokens: {document.get('tokens') or ''}",
        "",
    ]
    for index, segment in enumerate(segments, start=1):
        if not isinstance(segment, dict):
            continue
        content = segment.get("content") if isinstance(segment.get("content"), str) else ""
        lines.extend([f"## Segment {segment.get('position') or index}", "", content.strip(), ""])
    return "\n".join(lines).rstrip() + "\n"


async def download_knowledge_document(document_id: str, *, format: str = "markdown") -> Response:
    payload = await get_knowledge_document_detail(document_id)
    document = payload.get("document") if isinstance(payload.get("document"), dict) else {}
    name = document.get("name") if isinstance(document.get("name"), str) else document_id
    if format == "json":
        body = json.dumps(payload, ensure_ascii=False, indent=2) + "\n"
        filename = _safe_export_filename(name, "json")
        media_type = "application/json; charset=utf-8"
    else:
        body = _document_export_markdown(payload)
        filename = _safe_export_filename(name, "md")
        media_type = "text/markdown; charset=utf-8"

    return Response(
        content=body,
        media_type=media_type,
        headers={"Content-Disposition": _content_disposition(filename)},
    )


async def _delete_dify_document(document_id: str, *, dataset_id: str) -> None:
    base = get_dataset_api_base_url()
    headers = _dify_headers()
    url = f"{base}/datasets/{dataset_id}/documents/{document_id}"
    async with httpx.AsyncClient(timeout=60.0, trust_env=False) as client:
        try:
            response = await client.delete(url, headers=headers)
        except httpx.RequestError as exc:
            raise RagKnowledgeError(f"Upstream error: {exc}") from exc

    if response.status_code in {204, 404}:
        return
    if response.status_code == 400:
        raise RagKnowledgeError(_upstream_error_message(response, "Bad request"), status_code=400)
    raise RagKnowledgeError(_upstream_error_message(response), status_code=response.status_code)


async def delete_knowledge_document(document_id: str) -> Response:
    local_row = _get_local_document_row(document_id)
    if local_row is not None:
        for dify_document_id, dataset_id in (
            (local_row.get("raw_dify_document_id"), _raw_dataset_id()),
            (local_row.get("desensitized_dify_document_id") or local_row.get("dify_document_id"), _desensitized_dataset_id()),
        ):
            if dify_document_id:
                await _delete_dify_document(str(dify_document_id), dataset_id=dataset_id)
        try:
            with get_engine().begin() as conn:
                conn.execute(
                    text(
                        """
                        DELETE FROM rag.knowledge_documents
                        WHERE id = CAST(:document_id AS uuid)
                        """
                    ),
                    {"document_id": document_id},
                )
        except (SQLAlchemyError, RuntimeError) as exc:
            raise _database_error(exc) from exc
        return Response(status_code=204)

    await _delete_dify_document(document_id, dataset_id=_dataset_id())
    return Response(status_code=204)
