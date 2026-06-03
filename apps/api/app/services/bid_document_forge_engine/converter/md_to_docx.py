from typing import Optional
# -*- coding: utf-8 -*-
"""
Markdown → DOCX 转换器
将 LLM 生成的 Markdown 文本转换为标准排版的 Word 文档
"""

import logging
import os
import re
import tempfile
import zipfile
import xml.etree.ElementTree as ET
from io import BytesIO
from urllib.parse import unquote
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .styles import apply_bid_styles, set_page_margins, apply_page_number_footer

logger = logging.getLogger(__name__)


class MarkdownToDocxConverter:
    """
    Markdown → DOCX 转换器

    将 Markdown 格式的投标技术文件转换为标准排版的 Word 文档
    支持：标题层级、加粗/斜体、列表、表格（基础支持）
    """

    def __init__(self, template_path: Optional[str] = None):
        """
        Args:
            template_path: Word 模板路径（可选），为 None 时使用默认样式
        """
        self.template_path = template_path
        self._heading_num_id: Optional[int] = None
        self._temp_template_paths: list[Path] = []

    @staticmethod
    def _ensure_docx_package_numbering(docx_bytes: bytes) -> tuple[bytes, bool]:
        """
        确保 docx 包内存在 numbering.xml 与关联关系。
        返回: (处理后字节, 是否有改动)
        """
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
        REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
        ET.register_namespace("w", W_NS)
        ET.register_namespace("", CT_NS)

        def qn(ns: str, tag: str) -> str:
            return f"{{{ns}}}{tag}"

        in_buf = BytesIO(docx_bytes)

        with zipfile.ZipFile(in_buf, "r") as zin:
            entries = {info.filename: zin.read(info.filename) for info in zin.infolist()}

        changed = False

        if "word/numbering.xml" not in entries:
            numbering_xml = (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            entries["word/numbering.xml"] = numbering_xml.encode("utf-8")
            changed = True

        rels_key = "word/_rels/document.xml.rels"
        rels_xml = entries.get(rels_key)
        if rels_xml:
            rels_root = ET.fromstring(rels_xml)
            has_num_rel = False
            max_rid = 0
            for rel in rels_root.findall(qn(REL_NS, "Relationship")):
                rid = rel.attrib.get("Id", "")
                m = re.match(r"rId(\d+)$", rid)
                if m:
                    max_rid = max(max_rid, int(m.group(1)))
                if rel.attrib.get("Type", "").endswith("/relationships/numbering"):
                    has_num_rel = True
            if not has_num_rel:
                rel = ET.SubElement(rels_root, qn(REL_NS, "Relationship"))
                rel.set("Id", f"rId{max_rid + 1}")
                rel.set(
                    "Type",
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering",
                )
                rel.set("Target", "numbering.xml")
                entries[rels_key] = ET.tostring(rels_root, encoding="utf-8", xml_declaration=True)
                changed = True

        ct_key = "[Content_Types].xml"
        ct_xml = entries.get(ct_key)
        if ct_xml:
            ct_root = ET.fromstring(ct_xml)
            has_override = False
            for ov in ct_root.findall(qn(CT_NS, "Override")):
                if ov.attrib.get("PartName", "") == "/word/numbering.xml":
                    has_override = True
                    break
            if not has_override:
                ov = ET.SubElement(ct_root, qn(CT_NS, "Override"))
                ov.set("PartName", "/word/numbering.xml")
                ov.set(
                    "ContentType",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
                )
                entries[ct_key] = ET.tostring(ct_root, encoding="utf-8", xml_declaration=True)
                changed = True

        if not changed:
            return docx_bytes, False

        out_buf = BytesIO()
        with zipfile.ZipFile(out_buf, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for name, content in entries.items():
                zout.writestr(name, content)
        return out_buf.getvalue(), True

    def _prepare_template_path(self, template_path: Path) -> Path:
        """
        模板兜底：若模板缺失 numbering 部件，运行时生成临时补丁模板并返回其路径。
        """
        try:
            raw = template_path.read_bytes()
            patched, changed = self._ensure_docx_package_numbering(raw)
            if not changed:
                return template_path
            fd, temp_path = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            temp = Path(temp_path)
            temp.write_bytes(patched)
            self._temp_template_paths.append(temp)
            logger.info("模板缺失编号部件，已注入 numbering 运行时兜底: %s", template_path)
            return temp
        except Exception as e:
            logger.warning("模板编号兜底失败，回退原模板: %s", e)
            return template_path

    @staticmethod
    def _apply_run_font(run, zh_font: str = "宋体", en_font: str = "Times New Roman", size_pt: int = 12, bold: Optional[bool] = None):
        """对 run 强制设置中英文字体，避免 Word 自动回退。"""
        run.font.name = en_font
        run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), zh_font)
            run._element.rPr.rFonts.set(qn("w:ascii"), en_font)
            run._element.rPr.rFonts.set(qn("w:hAnsi"), en_font)
        except Exception:
            pass

    def _normalize_paragraph(self, p, heading_level: Optional[int] = None):
        """统一段落级排版，避免样式被模板/默认值污染。"""
        if heading_level is None:
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(28)
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0 if heading_level == 1 else 0.74)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(32 if heading_level == 1 else 28)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.widow_control = False
        p.paragraph_format.keep_with_next = False
        p.paragraph_format.keep_together = False
        for run in p.runs:
            if heading_level == 1:
                self._apply_run_font(run, zh_font="黑体", size_pt=16, bold=True)
            elif heading_level == 2:
                self._apply_run_font(run, zh_font="楷体_GB2312", size_pt=16, bold=True)
            elif heading_level == 3:
                self._apply_run_font(run, zh_font="仿宋_GB2312", size_pt=16, bold=True)
            else:
                self._apply_run_font(run, zh_font="宋体", size_pt=12, bold=False)

    @staticmethod
    def _split_heading_number(title: str) -> tuple[str, str]:
        text = (title or "").strip()
        if not text:
            return "", ""
        match = re.match(r"^(([一二三四五六七八九十]+、)|(\d+(?:\.\d+){1,2}))\s*(.+)$", text)
        if not match:
            return "", text
        return (match.group(1) or "").strip(), (match.group(4) or "").strip()

    @staticmethod
    def _build_heading_numbering(doc: Document) -> int:
        numbering = doc.part.numbering_part.numbering_definitions._numbering

        abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum")) if n.get(qn("w:abstractNumId"))]
        num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num")) if n.get(qn("w:numId"))]
        next_abstract_id = (max(abstract_ids) + 1) if abstract_ids else 1000
        next_num_id = (max(num_ids) + 1) if num_ids else 1000

        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(next_abstract_id))

        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "multilevel")
        abstract.append(multi)

        # H1: 一、 二、; H2: 1.1; H3: 1.1.1
        for ilvl, fmt, text, start in (
            (0, "chineseCounting", "%1、", "1"),
            (1, "decimal", "%1.%2", "1"),
            (2, "decimal", "%1.%2.%3", "1"),
        ):
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), str(ilvl))

            s = OxmlElement("w:start")
            s.set(qn("w:val"), start)
            lvl.append(s)

            nfmt = OxmlElement("w:numFmt")
            nfmt.set(qn("w:val"), fmt)
            lvl.append(nfmt)
            if ilvl >= 1:
                # 让 H2/H3 的上级编号按阿拉伯数字显示（例如 1.1 / 1.1.1），
                # 不受 H1 中文序号格式影响。
                is_lgl = OxmlElement("w:isLgl")
                lvl.append(is_lgl)

            lvl_text = OxmlElement("w:lvlText")
            lvl_text.set(qn("w:val"), text)
            lvl.append(lvl_text)

            lvl_jc = OxmlElement("w:lvlJc")
            lvl_jc.set(qn("w:val"), "left")
            lvl.append(lvl_jc)

            ppr = OxmlElement("w:pPr")
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), str(360 + ilvl * 360))
            ind.set(qn("w:hanging"), "0")
            ppr.append(ind)
            lvl.append(ppr)

            abstract.append(lvl)

        numbering.append(abstract)

        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(next_num_id))
        abs_ref = OxmlElement("w:abstractNumId")
        abs_ref.set(qn("w:val"), str(next_abstract_id))
        num.append(abs_ref)
        numbering.append(num)
        return next_num_id

    def _apply_heading_numbering(self, p, heading_level: int, doc: Document) -> None:
        if heading_level < 1 or heading_level > 3:
            return
        if self._heading_num_id is None:
            self._heading_num_id = self._build_heading_numbering(doc)
        ppr = p._p.get_or_add_pPr()
        old = ppr.find(qn("w:numPr"))
        if old is not None:
            ppr.remove(old)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), str(heading_level - 1))
        num_id = OxmlElement("w:numId")
        num_id.set(qn("w:val"), str(self._heading_num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id)
        ppr.append(num_pr)

    @staticmethod
    def _clear_document_body(doc: Document) -> None:
        """清空模板正文内容，保留样式、节设置、页眉页脚。"""
        body = doc._body._element
        # 保留 sectPr，清空其它正文块（段落/表格）
        for child in list(body):
            if child.tag.endswith("sectPr"):
                continue
            body.remove(child)

    def _apply_semantic_rules(self, p, line_text: str) -> None:
        """按公文语义对特定段落进行定向排版。"""
        t = (line_text or "").strip()
        if not t:
            return
        # 落款/署名靠右，右侧留 4 字符
        if re.match(r"^(落款[:：]?|署名[:：]?|投标单位[:：]?|法定代表人[:：]?)", t):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.right_indent = Cm(1.48)
        # 日期行：默认居中（与落款视觉对齐）
        elif re.match(r"^\d{4}年\d{1,2}月\d{1,2}日$", t):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
        # 联系方式：左对齐，首行缩进两字符
        elif re.match(r"^(联系方式[:：]?|联系人[:：]?|联系电话[:：]?)", t):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0.74)

    @staticmethod
    def _pick_image_style_name(doc: Document) -> Optional[str]:
        """优先使用模板里的图表段落样式。"""
        preferred = ("图表", "图片", "Image", "Figure")
        names = [s.name for s in doc.styles]
        for n in preferred:
            if n in names:
                return n
        return None

    @staticmethod
    def _convert_inline_to_anchor(run) -> None:
        """
        将 run 中的 inline 图片转换为浮动锚点，采用上下型环绕（业界常见避免正文挤压）。
        """
        try:
            drawing = run._r.xpath(".//w:drawing")[0]
            inline = drawing.xpath(".//wp:inline")[0]
        except Exception:
            return

        # inline -> anchor
        inline.tag = qn("wp:anchor")
        anchor = inline
        anchor.set(qn("wp:simplePos"), "0")
        anchor.set(qn("wp:relativeHeight"), "0")
        anchor.set(qn("wp:behindDoc"), "0")
        anchor.set(qn("wp:locked"), "0")
        anchor.set(qn("wp:layoutInCell"), "1")
        anchor.set(qn("wp:allowOverlap"), "1")

        # 已存在定位子节点则不重复追加
        has_pos_h = bool(anchor.xpath("./wp:positionH"))
        has_pos_v = bool(anchor.xpath("./wp:positionV"))
        has_wrap = bool(anchor.xpath("./wp:wrapTopAndBottom") or anchor.xpath("./wp:wrapSquare"))
        has_simple = bool(anchor.xpath("./wp:simplePos"))

        if not has_simple:
            simple_pos = OxmlElement("wp:simplePos")
            simple_pos.set("x", "0")
            simple_pos.set("y", "0")
            anchor.insert(0, simple_pos)

        if not has_pos_h:
            pos_h = OxmlElement("wp:positionH")
            pos_h.set(qn("wp:relativeFrom"), "column")
            align_h = OxmlElement("wp:align")
            align_h.text = "center"
            pos_h.append(align_h)
            anchor.insert(1, pos_h)

        if not has_pos_v:
            pos_v = OxmlElement("wp:positionV")
            pos_v.set(qn("wp:relativeFrom"), "paragraph")
            pos_off = OxmlElement("wp:posOffset")
            pos_off.text = "0"
            pos_v.append(pos_off)
            anchor.insert(2, pos_v)

        if not has_wrap:
            wrap = OxmlElement("wp:wrapTopAndBottom")
            anchor.insert(3, wrap)

    def convert(self, markdown_text: str, output_path: str) -> str:
        """
        将 Markdown 转换为 DOCX 并保存

        Args:
            markdown_text: Markdown 文本
            output_path: 输出 .docx 文件路径

        Returns:
            str: 输出文件路径
        """
        # 每次转换重置编号定义状态，避免跨任务污染
        self._heading_num_id = None

        # 从模板创建或新建文档
        if self.template_path and Path(self.template_path).exists():
            template = self._prepare_template_path(Path(self.template_path))
            doc = Document(str(template))
            # 使用模板样式前先清空模板示例正文
            self._clear_document_body(doc)
        else:
            doc = Document()

        # 应用标书样式
        doc = apply_bid_styles(doc)
        doc = set_page_margins(doc)
        doc = apply_page_number_footer(doc)

        # 逐行解析 Markdown 并写入文档
        lines = markdown_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # 空行跳过
            if not line.strip():
                i += 1
                continue

            # 标题
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()
                
                # 检测标题后缀是否包含隐藏书签标记，例如 {#BM_550e_8400}
                bm_match = re.search(r'\{#([_A-Za-z0-9]+)\}$', title)
                bookmark_id = None
                if bm_match:
                    bookmark_id = bm_match.group(1)
                    title = title[:bm_match.start()].strip()
                    
                h_level = min(level, 3)
                heading_number, heading_title = self._split_heading_number(title)
                p = doc.add_heading("", level=h_level)
                p.add_run(heading_title or title)
                self._normalize_paragraph(p, heading_level=h_level)
                self._apply_heading_numbering(p, h_level, doc)
                if bookmark_id:
                    self._add_bookmark(p, bookmark_id)
                    
                i += 1
                continue

            # 表格（检测表格开始）
            if "|" in line and i + 1 < len(lines) and re.match(r'^[\|\-\s:]+$', lines[i + 1]):
                table_lines = [line]
                i += 1
                while i < len(lines) and "|" in lines[i]:
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(doc, table_lines)
                continue

            # 无序列表
            list_match = re.match(r'^(\s*)[*\-+•●○▪◦·]\s*(.+)?$', line)
            if list_match:
                content = (list_match.group(2) or "").strip()
                if not content:
                    i += 1
                    continue
                content = self._process_inline(content)
                has_bullet_style = any(s.name == "List Bullet" for s in doc.styles)
                if has_bullet_style:
                    p = doc.add_paragraph(content, style="List Bullet")
                else:
                    p = doc.add_paragraph(f"· {content}")
                self._normalize_paragraph(p)
                i += 1
                continue

            # 有序列表
            ordered_match = re.match(r'^(\s*)(\d+)[.)]\s+(.+)$', line)
            if ordered_match:
                ordinal = ordered_match.group(2)
                content = ordered_match.group(3)
                content = self._process_inline(content)
                has_number_style = any(s.name == "List Number" for s in doc.styles)
                if has_number_style:
                    p = doc.add_paragraph(content, style="List Number")
                else:
                    p = doc.add_paragraph(f"{ordinal}. {content}")
                self._normalize_paragraph(p)
                i += 1
                continue

            # 普通段落
            line_str = line.strip()
            
            # 图片检测：如果整行是一个单独的图片语法 ![...](...)
            image_match = re.match(r'^!\[(.*?)\]\((.+)\)$', line_str)
            if image_match:
                img_path = unquote(image_match.group(2).strip().strip('"').strip("'"))
                try:
                    if Path(img_path).exists():
                        img_style = self._pick_image_style_name(doc)
                        p = doc.add_paragraph(style=img_style) if img_style else doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.first_line_indent = Cm(0)
                        p.paragraph_format.space_before = Pt(6)
                        p.paragraph_format.space_after = Pt(6)
                        run = p.add_run()
                        run.add_picture(img_path, width=Cm(14))
                        # 模板锚点优先；代码兜底将 inline 转为上下型环绕 anchor
                        self._convert_inline_to_anchor(run)
                    else:
                        logger.warning(f"图片路径不存在，跳过: {img_path}")
                except Exception as e:
                    logger.warning(f"插入图片失败 {img_path}: {e}")
                i += 1
                continue

            content = self._process_inline(line_str)
            p = doc.add_paragraph(content)
            self._normalize_paragraph(p)
            self._apply_semantic_rules(p, line_str)
            i += 1

        # 保存文件
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output))
        for temp in self._temp_template_paths:
            try:
                temp.unlink(missing_ok=True)
            except Exception:
                pass
        self._temp_template_paths.clear()
        logger.info(f"DOCX 生成完成: {output_path}")
        return str(output)

    def _add_table(self, doc: Document, table_lines: list[str]) -> None:
        """解析 Markdown 表格并添加到文档"""
        rows = []
        for line in table_lines:
            # 跳过分隔行
            if re.match(r'^[\|\-\s:]+$', line):
                continue
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            rows.append(cells)

        if not rows:
            return

        # 取第一行作为表头确定列数
        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    table.cell(i, j).text = self._process_inline(cell_text)
                    for paragraph in table.cell(i, j).paragraphs:
                        for run in paragraph.runs:
                            self._apply_run_font(run, zh_font="宋体", size_pt=10)

                    # 表头加粗
                    if i == 0:
                        for paragraph in table.cell(i, j).paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    @staticmethod
    def _add_bookmark(paragraph, bookmark_name: str) -> None:
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn
        import uuid
        
        bm_id = str(uuid.uuid4().int % 1000000)
        
        bm_start = OxmlElement('w:bookmarkStart')
        bm_start.set(qn('w:id'), bm_id)
        bm_start.set(qn('w:name'), bookmark_name)
        
        bm_end = OxmlElement('w:bookmarkEnd')
        bm_end.set(qn('w:id'), bm_id)
        
        paragraph._p.insert(0, bm_start)
        paragraph._p.append(bm_end)

    @staticmethod
    def _process_inline(text: str) -> str:
        """处理行内 Markdown 格式（加粗、斜体等），返回纯文本"""
        # 去除加粗标记
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        # 去除斜体标记
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        # 去除代码标记
        text = re.sub(r'`(.+?)`', r'\1', text)
        return text
