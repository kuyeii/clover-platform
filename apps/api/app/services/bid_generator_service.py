from __future__ import annotations

import asyncio
import ast
import base64
import json
import os
import logging
import re
import threading
import socket
import time
import io
import html
import copy
import hashlib
import uuid
import zipfile
import tempfile
import subprocess
import shutil
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Mapping, Optional
from urllib.parse import quote, urlparse

import httpx
import requests
import yaml
from fastapi import Request, UploadFile
from fastapi.responses import StreamingResponse
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from app.core.config import get_api_settings
from app.core.errors import PlatformError
from app.services.bid_attachment_template_service import ATTACHMENT_LABELS, render_attachment
from app.services.bid_bidder_pipt_service import (
    merge_bidder_pipt_context,
    validate_required_bidder_info,
)
from app.services.bid_content_placeholder_service import (
    find_illegal_pipt_bidder_placeholders_native,
    resolve_body_placeholders_native,
)
from app.services.bid_docanalysis_service import (
    build_docanalysis_groups,
    build_docanalysis_node_index,
    build_docanalysis_system_prompt,
    extract_docanalysis_node_content,
    extract_docanalysis_text_output,
    load_docanalysis_framework,
    parse_bid_attachments_payload,
    parse_docanalysis_result_map,
    split_bid_attachments_tag,
)
from app.services.bid_document_forge_service import (
    add_scoring_table_and_attachments,
    create_document_forge,
)
from app.services.bid_document_forge_engine.markdown_norm import normalize_generated_markdown
from app.services.bid_outline_service import (
    build_outline_generation_bundle,
    build_seeded_outline_sections,
    evaluate_outline_quality,
    extract_outline_sections_raw,
    normalize_outline_word_budget_dict,
    parse_dify_outputs,
)
from app.services import bid_workflow_execution_adapter
from app.services.bid_task_runtime_service import task_manager as native_task_manager
from app.services.pipt_gateway_service import preprocess_internal_payload
from sqlalchemy import bindparam, text
from sqlalchemy.exc import SQLAlchemyError
from packages.py_common.db.session import get_engine

logger = logging.getLogger(__name__)

_IMPORT_LOCK = threading.RLock()

WORKFLOWS: tuple[tuple[str, str, str, bool, str], ...] = (
    ("structure_generator", "DIFY_WORKFLOW_STRUCTURE_GENERATOR", "大纲生成", True, "managed"),
    ("content_writer", "DIFY_WORKFLOW_CONTENT_WRITER", "单章节内容生成", True, "managed"),
    ("content_group_writer", "DIFY_WORKFLOW_CONTENT_GROUP_WRITER", "H2分组正文生成", True, "managed"),
    ("content_rewrite", "DIFY_WORKFLOW_CONTENT_REWRITE", "单章节重生成", True, "managed"),
    ("response_content_writer", "DIFY_WORKFLOW_RESPONSE_CONTENT_WRITER", "响应情况正文生成", True, "managed"),
    ("diagram_generator", "DIFY_WORKFLOW_DIAGRAM_GENERATOR", "图表生成（SVG）", True, "managed"),
    ("diagram_generator_mermaid", "DIFY_WORKFLOW_DIAGRAM_GENERATOR_MERMAID", "图表生成（Mermaid）", True, "managed"),
    ("doc_analysis", "DIFY_WORKFLOW_DOC_ANALYSIS", "文档分析", True, "managed"),
    ("requirement_extractor", "DIFY_WORKFLOW_REQUIREMENT_EXTRACTOR", "需求提取", False, "legacy"),
    ("blueprint_generator", "DIFY_WORKFLOW_BLUEPRINT_GENERATOR", "全局策略蓝图", False, "legacy"),
    ("group_review_writer", "DIFY_WORKFLOW_GROUP_REVIEW_WRITER", "H2章节评估", False, "legacy"),
    ("attachment_generator", "DIFY_WORKFLOW_ATTACHMENT_GENERATOR", "智能附件生成", False, "legacy"),
    ("scoring_assistant", "DIFY_WORKFLOW_SCORING_ASSISTANT", "评分AI助手", False, "legacy"),
)

SUPPORTED_ENTITIES: dict[str, str] = {
    "name": "姓名",
    "phone": "电话号码",
    "id_number": "身份证号",
    "bank": "银行账户",
    "car_id": "车牌号",
    "ip": "IP地址",
    "email": "电子邮箱",
    "addr": "地址",
    "gender": "性别",
    "political_status": "政治面貌",
    "nation": "民族",
    "org": "企业/机构",
}

_BID_ATTACH_STAGE_PREFIX = "__bid_attachments__"
_ANALYSIS_V2_STAGE_PREFIX = "__analysis_v2__"
_TASK_EVENT_STAGE_PREFIX = "__task_event__"

_DIFY_NODE_STAGE_MAP: dict[str, str] = {
    "SearxNG": "🔍 知识检索",
    "知识检索": "🔍 知识检索",
    "LLM WITH SEARXNG": "✍️ 正文生成",
    "LLM WITHOUT SEARXNG": "✍️ 正文生成",
    "LLM RESPONSE WRITER": "✍️ 响应情况生成",
    "合并草稿": "📋 合并草稿",
    "Reviewer_LLM": "📝 质量审查",
    "参数提取器": "📊 评分分析",
    "Rewriter LLM": "✏️ 润色修改",
    "LLM大纲生成": "✍️ 生成大纲",
    "LLM大纲润色": "✨ 大纲润色",
    "LLM 需求提取": "📋 需求提取",
    "LLM 蓝图生成": "🎯 策略分析",
    "LLM 自评评分": "📊 AI 填写",
    "LLM 附件生成": "📄 附件撰写",
    "JSON解析校验": "⚙️ 数据校验",
    "JSON解析1": "⚙️ 数据校验",
    "最终校验": "✅ 最终校验",
    "输出清洗": "🧹 输出清洗",
    "输出": "✅ 结果输出",
}

_SCORING_SYSTEM_PROMPT = """你是一位专业的政府采购投标顾问，负责帮助投标方完成自评评分表。

你的核心立场：**尽量论证"完全响应"或"部分响应"，绝对避免"不响应"**。
即使该项目条件稍显不足，也要从现有优势、整体方案、类似案例等角度积极论证，
以最有利于投标方的表述展示响应程度。

输出 JSON，格式如下：
{
  "self_response": "full" 或 "partial"，禁止输出 "none",
  "self_comment": "自评说明（100-150字，正式书面语，有依据）",
  "evidence_refs": ["相关证明文件路径或关键词，如 '资质证书/高新技术企业证书.pdf'，最多3条"]
}

只输出 JSON，不要任何其他内容。"""


class BidProjectNotFound(Exception):
    pass


class BidProjectConflict(Exception):
    pass


@dataclass(frozen=True, slots=True)
class BidGeneratorFilePayload:
    content: bytes
    media_type: str
    filename: str
    inline: bool = True
    cache_control: str = "public, max-age=3600"
    headers: Mapping[str, str] | None = None


def _repo_root() -> Path:
    return get_api_settings().repo_root


def _bid_generator_root() -> Path:
    return _repo_root() / "legacy" / "bid-generator" / "pipt-flask"


def _bid_generator_legacy_root() -> Path:
    return _repo_root() / "legacy" / "bid-generator"


def _bid_generator_resource_root() -> Path:
    return _repo_root() / "apps" / "api" / "app" / "resources" / "bid_generator"


def _bid_generator_config_path() -> Path:
    return _bid_generator_resource_root() / "config.yaml"


def _legacy_app_package_path() -> Path:
    return _bid_generator_root() / "app"


def _read_root_env_value(env_var: str) -> str:
    for env_path in (_repo_root() / ".env", _repo_root() / "legacy" / "bid-generator" / ".env"):
        try:
            if not env_path.exists():
                continue
            with env_path.open("r", encoding="utf-8") as file:
                for raw_line in file:
                    line = raw_line.strip()
                    if not line or line.startswith("#") or "=" not in line:
                        continue
                    key, value = line.split("=", 1)
                    if key.strip() == env_var:
                        return value.strip().strip('"').strip("'")
        except OSError:
            continue
    return ""


def _get_workflow_key_source(workflow_name: str) -> tuple[bool, str]:
    env_var = f"DIFY_WORKFLOW_{workflow_name.upper()}"
    if os.environ.get(env_var, "").strip():
        return True, "process_env"
    if _read_root_env_value(env_var):
        return True, "root_env_file"
    return False, "missing"


def _diagram_generation_enabled() -> bool:
    return os.environ.get("ENABLE_DIAGRAM_GENERATION", "false").strip().lower() == "true"


def _get_diagram_generator_mode() -> str:
    mode = os.environ.get("DIAGRAM_GENERATOR_MODE", "svg").strip().lower()
    return "mermaid" if mode in {"mermaid", "mmd"} else "svg"


def _get_diagram_workflow_name() -> str:
    return "diagram_generator_mermaid" if _get_diagram_generator_mode() == "mermaid" else "diagram_generator"


def _dump_structure_heading_seed_json_local(headings: list[dict[str, Any]]) -> str:
    if not isinstance(headings, list):
        return "[]"
    return json.dumps(headings, ensure_ascii=False)


def _split_outline_seed_headings_local(
    seed_headings: list[dict[str, Any]],
    strategy: str = "auto",
    auto_threshold: int = 4,
) -> list[list[dict[str, Any]]]:
    sections = seed_headings if isinstance(seed_headings, list) else []
    total = len(sections)
    mode = str(strategy or "auto").strip().lower()
    threshold = max(2, int(auto_threshold or 4))
    if total <= 1:
        return [sections]
    if mode in {"single", "off", "disabled"}:
        return [sections]
    if mode == "force_parallel":
        first_size = max(1, total // 2)
        return [sections[:first_size], sections[first_size:]]
    if total <= threshold:
        return [sections]
    first_size = max(2, total // 2)
    return [sections[:first_size], sections[first_size:]]


def _check_dns_host(host: str, port: int = 443) -> dict[str, str | int | bool]:
    """轻量 DNS 诊断，不发起模型请求，也不暴露任何密钥。"""
    try:
        socket.getaddrinfo(host, port, type=socket.SOCK_STREAM)
    except socket.gaierror as exc:
        return {
            "host": host,
            "port": port,
            "resolvable": False,
            "status": "error",
            "message": f"DNS 解析失败: {exc}",
        }
    except OSError as exc:
        return {
            "host": host,
            "port": port,
            "resolvable": False,
            "status": "error",
            "message": f"DNS 检查失败: {exc}",
        }
    return {
        "host": host,
        "port": port,
        "resolvable": True,
        "status": "ok",
        "message": "DNS 可解析",
    }


def _model_provider_diagnostics() -> dict[str, dict[str, str | int | bool]]:
    return {
        "dashscope": _check_dns_host("dashscope.aliyuncs.com"),
    }


def _dify_api_diagnostics() -> dict[str, str | int | bool]:
    """诊断标书后端到 Dify API 的基础连通配置，只检查主机解析。"""
    raw_url = os.environ.get("DIFY_API_URL", "http://localhost/v1").strip() or "http://localhost/v1"
    parsed_url = raw_url if "://" in raw_url else f"http://{raw_url}"
    parsed = urlparse(parsed_url)
    host = parsed.hostname or ""
    try:
        port = parsed.port or (443 if parsed.scheme == "https" else 80)
    except ValueError as exc:
        return {
            "url_env": "DIFY_API_URL",
            "host": host,
            "port": "",
            "resolvable": False,
            "status": "error",
            "message": f"DIFY_API_URL 端口无效: {exc}",
        }
    if not host:
        return {
            "url_env": "DIFY_API_URL",
            "host": "",
            "port": port,
            "resolvable": False,
            "status": "error",
            "message": "DIFY_API_URL 缺少可解析的主机名",
        }
    return {
        "url_env": "DIFY_API_URL",
        **_check_dns_host(host, port),
    }


def get_health_payload() -> dict[str, str]:
    return {"status": "ok", "service": "pipt-lite"}


def get_workflow_status_payload() -> dict[str, Any]:
    status: dict[str, Any] = {}
    for name, env_var, label, managed, lifecycle in WORKFLOWS:
        configured, source = _get_workflow_key_source(name)
        source_value = source
        if name in {"diagram_generator", "diagram_generator_mermaid"} and not _diagram_generation_enabled():
            configured = False
            source_value = "disabled"
        status[name] = {
            "label": label,
            "env_var": env_var,
            "configured": configured,
            "source": source_value,
            "managed": managed,
            "lifecycle": lifecycle,
        }
    status["_diagnostics"] = {
        "label": "外部依赖诊断",
        "managed": False,
        "lifecycle": "diagnostic",
        "providers": _model_provider_diagnostics(),
        "dify_api": _dify_api_diagnostics(),
    }
    return status


def get_analysis_framework_payload() -> Any:
    config_path = _bid_generator_root() / "config" / "analysis_framework.json"
    if not config_path.exists():
        raise PlatformError(
            code="RESOURCE_NOT_FOUND",
            message="analysis_framework.json 配置文件不存在",
            status_code=404,
        )
    try:
        with config_path.open("r", encoding="utf-8") as file:
            return json.load(file)
    except json.JSONDecodeError as exc:
        raise PlatformError(
            code="BUSINESS_DIRECT_ERROR",
            message="analysis_framework.json 配置文件不是合法 JSON。",
            status_code=500,
        ) from exc


def get_template_config_payload(template_name: str = "") -> dict[str, Any]:
    """读取标书系统配置与大纲模板；入参为模板文件名，出参为 config/template。"""
    config_path = _bid_generator_config_path()
    templates_dir = _template_structures_dir()
    normalized_template_name = str(template_name or "").strip()
    if normalized_template_name and ("/" in normalized_template_name or "\\" in normalized_template_name):
        raise PlatformError(code="INVALID_REQUEST", message="Invalid template name", status_code=400)

    try:
        available_templates = _list_template_structure_names(templates_dir)
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="读取模板目录失败。", status_code=500) from exc

    if normalized_template_name and normalized_template_name not in available_templates:
        raise PlatformError(
            code="RESOURCE_NOT_FOUND",
            message=f"Template not found: {normalized_template_name}",
            status_code=404,
        )
    current_template = normalized_template_name or (available_templates[0] if available_templates else "")
    template_path = templates_dir / current_template if current_template else None

    return {
        "config_dict": _read_yaml_mapping(config_path),
        "template_dict": _read_yaml_mapping(template_path) if template_path else {},
        "available_templates": available_templates,
        "current_template": current_template,
    }


def get_supported_entities_payload() -> dict[str, Any]:
    return {
        "entities": SUPPORTED_ENTITIES,
        "description": "key 为实体标识符，value 为中文名称",
    }


def _database_error(exc: Exception) -> PlatformError:
    logger.exception("Bid-generator PostgreSQL operation failed")
    return PlatformError(
        code="DATABASE_ERROR",
        message="标书生成项目数据库访问失败。",
        status_code=500,
        details={"module": "bid-generator", "schema": "bid_generator"},
    )


def _ensure_project_storage() -> None:
    try:
        with get_engine().begin() as conn:
            exists = conn.execute(text("SELECT to_regclass('bid_generator.projects') IS NOT NULL")).scalar_one()
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc

    if not exists:
        raise PlatformError(
            code="DATABASE_ERROR",
            message="标书生成项目数据库表不存在。",
            status_code=500,
            details={"table": "bid_generator.projects"},
        )


def _json_value(value: Any) -> Any:
    if isinstance(value, str):
        try:
            return json.loads(value)
        except json.JSONDecodeError as exc:
            raise PlatformError(
                code="BUSINESS_DIRECT_ERROR",
                message="标书生成项目数据不是合法 JSON。",
                status_code=500,
            ) from exc
    return value if value is not None else {}


def _iso_value(value: Any) -> str:
    if value is None:
        return ""
    if hasattr(value, "isoformat"):
        return value.isoformat()
    return str(value)


def _project_from_row(row: Mapping[str, Any]) -> dict[str, Any]:
    return {
        "id": str(row["id"]),
        "name": str(row["name"]),
        "status": str(row["status"]),
        "data": _json_value(row.get("data")),
        "created_at": _iso_value(row.get("created_at")),
        "updated_at": _iso_value(row.get("updated_at")),
    }


def list_projects_payload() -> list[dict[str, Any]]:
    _ensure_project_storage()
    try:
        with get_engine().begin() as conn:
            rows = conn.execute(
                text(
                    """
                    SELECT id, name, status, data, created_at, updated_at
                    FROM bid_generator.projects
                    ORDER BY created_at DESC
                    """
                )
            ).mappings().all()
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc
    return [_project_from_row(row) for row in rows]


def get_project_payload(project_id: str) -> dict[str, Any]:
    _ensure_project_storage()
    try:
        with get_engine().begin() as conn:
            row = conn.execute(
                text(
                    """
                    SELECT id, name, status, data, created_at, updated_at
                    FROM bid_generator.projects
                    WHERE id = :project_id
                    """
                ),
                {"project_id": project_id},
            ).mappings().first()
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc
    if row is None:
        raise BidProjectNotFound()
    return _project_from_row(row)


def get_project_mappings_payload(project_id: str) -> dict[str, Any]:
    project = get_project_payload(project_id)
    data = project.get("data") if isinstance(project, dict) else {}
    if not isinstance(data, dict):
        mapping_table = {}
    else:
        mapping_table = data.get("mappingTable", {})
    try:
        count = len(mapping_table)
    except TypeError:
        count = 0
    return {"mappings": mapping_table, "count": count}


def get_project_doc_blocks_payload(project_id: str) -> dict[str, Any]:
    """读取项目文档块快照；入参为项目 ID，出参兼容 legacy doc-blocks。"""
    normalized_id = _ensure_safe_project_id(project_id)
    try:
        project = get_project_payload(normalized_id)
    except BidProjectNotFound as exc:
        raise PlatformError(
            code="RESOURCE_NOT_FOUND",
            message=f"项目 [{normalized_id}] 的文档块缓存不存在",
            status_code=404,
        ) from exc
    data = project.get("data") if isinstance(project, dict) else {}
    blocks = data.get("__doc_blocks_cache") if isinstance(data, dict) else []
    if not isinstance(blocks, list) or not blocks:
        raise PlatformError(
            code="RESOURCE_NOT_FOUND",
            message=f"项目 [{normalized_id}] 的文档块缓存不存在",
            status_code=404,
        )
    return {
        "project_id": normalized_id,
        "blocks": blocks,
        "total_blocks": len(blocks),
        "snapshot_only": True,
    }


def create_project_payload(payload: Mapping[str, Any]) -> dict[str, Any]:
    """创建标书项目；入参为 legacy ProjectCreate JSON，出参保持 legacy ProjectResponse 结构。"""
    _ensure_project_storage()
    project_id = _required_string(payload.get("id"), field="id")
    name = _required_string(payload.get("name"), field="name")
    status = _string_or_default(payload.get("status"), default="uploading")
    data = _dict_or_default(payload.get("data"))
    try:
        with get_engine().begin() as conn:
            exists = conn.execute(
                text("SELECT 1 FROM bid_generator.projects WHERE id = :project_id"),
                {"project_id": project_id},
            ).first()
            if exists is not None:
                raise BidProjectConflict()
            row = conn.execute(
                text(
                    """
                    INSERT INTO bid_generator.projects (id, name, status, data)
                    VALUES (:project_id, :name, :status, :data)
                    RETURNING id, name, status, data, created_at, updated_at
                    """
                ),
                {
                    "project_id": project_id,
                    "name": name,
                    "status": status,
                    "data": json.dumps(data, ensure_ascii=False),
                },
            ).mappings().one()
    except BidProjectConflict:
        raise
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc
    return _project_from_row(row)


def update_project_payload(project_id: str, payload: Mapping[str, Any]) -> dict[str, Any]:
    """更新或 upsert 标书项目；入参为项目 ID 与 ProjectUpdate JSON，出参为项目完整记录。"""
    _ensure_project_storage()
    normalized_id = _required_string(project_id, field="project_id")
    data = _dict_or_default(payload.get("data"))
    name = _optional_string(payload.get("name")) or _optional_string(data.get("name")) or normalized_id
    status = _optional_string(payload.get("status")) or "uploaded"
    data_json = json.dumps(data, ensure_ascii=False)
    try:
        with get_engine().begin() as conn:
            row = conn.execute(
                text(
                    """
                    INSERT INTO bid_generator.projects (id, name, status, data)
                    VALUES (:project_id, :name, :status, :data)
                    ON CONFLICT (id) DO UPDATE SET
                      name = COALESCE(:update_name, bid_generator.projects.name),
                      status = COALESCE(:update_status, bid_generator.projects.status),
                      data = COALESCE(:update_data, bid_generator.projects.data),
                      updated_at = :updated_at
                    RETURNING id, name, status, data, created_at, updated_at
                    """
                ),
                {
                    "project_id": normalized_id,
                    "name": name,
                    "status": status,
                    "data": data_json,
                    "update_name": _optional_string(payload.get("name")),
                    "update_status": _optional_string(payload.get("status")),
                    "update_data": data_json if "data" in payload else None,
                    "updated_at": _utc_now(),
                },
            ).mappings().one()
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc
    return _project_from_row(row)


def patch_project_payload(project_id: str, payload: Mapping[str, Any]) -> dict[str, Any]:
    """增量更新标书项目 data；入参支持 data_patch 深合并和 remove_data_keys 删除，出参为项目完整记录。"""
    _ensure_project_storage()
    normalized_id = _required_string(project_id, field="project_id")
    try:
        with get_engine().begin() as conn:
            row = conn.execute(
                text(
                    """
                    SELECT id, name, status, data, created_at, updated_at
                    FROM bid_generator.projects
                    WHERE id = :project_id
                    FOR UPDATE
                    """
                ),
                {"project_id": normalized_id},
            ).mappings().first()
            if row is None:
                raise BidProjectNotFound()

            data = _json_value(row.get("data"))
            if not isinstance(data, dict):
                data = {}
            name = str(row["name"])
            status = str(row["status"])
            payload_name = _optional_string(payload.get("name"))
            payload_status = _optional_string(payload.get("status"))
            if payload_name is not None:
                name = payload_name
                data["name"] = payload_name
            if payload_status is not None:
                status = payload_status
                data["status"] = payload_status

            data_patch = payload.get("data_patch", {})
            if data_patch is not None and not isinstance(data_patch, dict):
                raise PlatformError(
                    code="INVALID_REQUEST",
                    message="data_patch 必须是对象。",
                    status_code=400,
                )
            if data_patch:
                data = _deep_merge_dict(data, data_patch)

            remove_data_keys = payload.get("remove_data_keys", [])
            if remove_data_keys is not None and not isinstance(remove_data_keys, list):
                raise PlatformError(
                    code="INVALID_REQUEST",
                    message="remove_data_keys 必须是数组。",
                    status_code=400,
                )
            for key in remove_data_keys or []:
                if isinstance(key, str) and key:
                    data.pop(key, None)

            updated = conn.execute(
                text(
                    """
                    UPDATE bid_generator.projects
                    SET name = :name,
                        status = :status,
                        data = :data,
                        updated_at = :updated_at
                    WHERE id = :project_id
                    RETURNING id, name, status, data, created_at, updated_at
                    """
                ),
                {
                    "project_id": normalized_id,
                    "name": name,
                    "status": status,
                    "data": json.dumps(data, ensure_ascii=False),
                    "updated_at": _utc_now(),
                },
            ).mappings().one()
    except (BidProjectNotFound, PlatformError):
        raise
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc
    return _project_from_row(updated)


def delete_project_payload(project_id: str) -> None:
    """删除标书项目；入参为项目 ID，删除成功无返回，未命中抛 BidProjectNotFound。"""
    _ensure_project_storage()
    normalized_id = _required_string(project_id, field="project_id")
    try:
        with get_engine().begin() as conn:
            result = conn.execute(
                text("DELETE FROM bid_generator.projects WHERE id = :project_id"),
                {"project_id": normalized_id},
            )
            if int(result.rowcount or 0) == 0:
                raise BidProjectNotFound()
    except BidProjectNotFound:
        raise
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc


def delete_project_caches_payload(project_id: str) -> dict[str, Any]:
    """清理项目后端缓存；入参为项目 ID，出参兼容 legacy cleaned/message。"""
    normalized_id = _ensure_safe_project_id(project_id)
    cleaned: list[str] = []
    for cache_name, path in (
        ("pdf_cache", _pdf_cache_path(normalized_id)),
        ("raw_doc_cache", _raw_doc_cache_path(normalized_id)),
        ("docx_cache", _docx_cache_path(normalized_id)),
    ):
        try:
            if path.exists():
                path.unlink()
                cleaned.append(cache_name)
        except OSError as exc:
            raise PlatformError(code="BUSINESS_DIRECT_ERROR", message=f"清理 {cache_name} 失败。", status_code=500) from exc

    return {
        "project_id": normalized_id,
        "cleaned": cleaned,
        "message": f"已清理 {len(cleaned)} 项资源",
    }


def batch_create_projects_payload(projects: Any) -> dict[str, int]:
    """批量 upsert 标书项目；入参为 ProjectCreate 数组，出参为 created/updated 计数。"""
    _ensure_project_storage()
    if not isinstance(projects, list):
        raise PlatformError(code="INVALID_REQUEST", message="请求体必须是项目数组。", status_code=400)
    if not projects:
        return {"created": 0, "updated": 0}

    incoming: dict[str, dict[str, Any]] = {}
    for item in projects:
        if not isinstance(item, Mapping):
            raise PlatformError(code="INVALID_REQUEST", message="项目必须是对象。", status_code=400)
        project_id = _required_string(item.get("id"), field="id")
        data = _dict_or_default(item.get("data"))
        incoming[project_id] = {
            "id": project_id,
            "name": _required_string(item.get("name"), field="name"),
            "status": _string_or_default(item.get("status"), default="uploading"),
            "data": json.dumps(data, ensure_ascii=False),
        }

    try:
        with get_engine().begin() as conn:
            existing_rows = conn.execute(
                text("SELECT id FROM bid_generator.projects WHERE id = ANY(:project_ids)"),
                {"project_ids": list(incoming)},
            ).mappings().all()
            existing_ids = {str(row["id"]) for row in existing_rows}
            now = _utc_now()
            for item in incoming.values():
                conn.execute(
                    text(
                        """
                        INSERT INTO bid_generator.projects (id, name, status, data, created_at, updated_at)
                        VALUES (:project_id, :name, :status, :data, :created_at, :updated_at)
                        ON CONFLICT (id) DO UPDATE SET
                          name = EXCLUDED.name,
                          status = EXCLUDED.status,
                          data = EXCLUDED.data,
                          updated_at = EXCLUDED.updated_at
                        """
                    ),
                    {
                        "project_id": item["id"],
                        "name": item["name"],
                        "status": item["status"],
                        "data": item["data"],
                        "created_at": now,
                        "updated_at": now,
                    },
                )
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc
    created = len(set(incoming) - existing_ids)
    return {"created": created, "updated": len(incoming) - created}


def save_analysis_report_payload(project_id: str, payload: Mapping[str, Any]) -> dict[str, str]:
    """保存解析报告；入参为项目 ID 和 analysis_report 数组，出参兼容 legacy message/path。"""
    _ensure_safe_project_id(project_id)
    report = payload.get("analysis_report", [])
    if not isinstance(report, list):
        raise PlatformError(code="INVALID_REQUEST", message="analysis_report 必须是数组。", status_code=400)
    _persist_project_analysis_report(project_id=project_id, report=report)
    save_path = _analysis_report_mirror_path(project_id)
    try:
        save_path.parent.mkdir(parents=True, exist_ok=True)
        with save_path.open("w", encoding="utf-8") as file:
            json.dump(report, file, ensure_ascii=False, indent=2)
    except OSError as exc:
        logger.warning("[%s] 保存 analysisReport 文件镜像失败: %s", project_id, exc)
    return {"message": "保存成功", "path": str(save_path)}


def get_analysis_report_payload(project_id: str) -> dict[str, Any]:
    """读取解析报告；入参为项目 ID，出参包含 analysis_report 和 analysis_v2。"""
    _ensure_safe_project_id(project_id)
    try:
        project = get_project_payload(project_id)
    except BidProjectNotFound:
        project = {}
    data = project.get("data") if isinstance(project, dict) else {}
    if isinstance(data, dict):
        report = data.get("analysisReport") or data.get("analysis_report")
        if isinstance(report, list):
            analysis_v2 = data.get("analysisV2") or data.get("analysis_v2") or {}
            return {"analysis_report": report, "analysis_v2": analysis_v2 if isinstance(analysis_v2, dict) else {}}

    save_path = _analysis_report_mirror_path(project_id)
    if not save_path.exists():
        return {"analysis_report": [], "analysis_v2": {}}
    try:
        with save_path.open("r", encoding="utf-8") as file:
            report = json.load(file)
    except (OSError, json.JSONDecodeError) as exc:
        logger.warning("[%s] 读取 analysisReport 文件镜像失败: %s", project_id, exc)
        return {"analysis_report": [], "analysis_v2": {}}
    return {"analysis_report": report if isinstance(report, list) else [], "analysis_v2": {}}


def get_cached_pdf_payload(project_id: str) -> BidGeneratorFilePayload:
    """读取项目 PDF 缓存；入参为项目 ID，出参为 PDF 字节及响应元数据。"""
    normalized_id = _ensure_safe_project_id(project_id)
    pdf_path = _pdf_cache_path(normalized_id)
    if not pdf_path.exists():
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="PDF 文件未找到，请先上传招标文件", status_code=404)
    try:
        content = pdf_path.read_bytes()
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="读取 PDF 文件失败。", status_code=500) from exc
    return BidGeneratorFilePayload(
        content=content,
        media_type="application/pdf",
        filename=f"{normalized_id}.pdf",
        inline=True,
        cache_control="public, max-age=3600",
    )


def upload_pdf_payload(project_id: str, *, filename: str, content: bytes) -> dict[str, str]:
    """缓存项目 PDF；入参为项目 ID、文件名和文件字节，出参兼容 legacy pdf_url/message。"""
    normalized_id = _ensure_safe_project_id(project_id)
    if not str(filename or "").lower().endswith(".pdf"):
        raise PlatformError(code="INVALID_REQUEST", message="仅支持 PDF 格式文件", status_code=400)
    if not content:
        raise PlatformError(code="INVALID_REQUEST", message="上传文件为空。", status_code=400)
    pdf_path = _pdf_cache_path(normalized_id)
    try:
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        pdf_path.write_bytes(content)
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="缓存 PDF 文件失败。", status_code=500) from exc
    return {"pdf_url": f"/api/v1/bid-generator/api/projects/pdf/{normalized_id}", "message": "PDF 已缓存"}


def get_source_docx_payload(project_id: str) -> BidGeneratorFilePayload:
    """读取项目原始 DOCX；入参为项目 ID，出参为 DOCX 字节及响应元数据。"""
    normalized_id = _ensure_safe_project_id(project_id)
    docx_path = _docx_cache_path(normalized_id)
    if not docx_path.exists():
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="项目原始 DOCX 不存在，请重新上传 DOCX 或执行重建定位", status_code=404)
    try:
        content = docx_path.read_bytes()
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="读取原始 DOCX 失败。", status_code=500) from exc
    return BidGeneratorFilePayload(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{normalized_id}.docx",
        inline=True,
        cache_control="public, max-age=3600",
    )


def get_diagram_artifact_svg_payload(diagram_id: str, *, project_id: str = "") -> BidGeneratorFilePayload:
    """读取图表 SVG artifact；入参为图表 ID/项目 ID，出参为 SVG 字符串响应元数据。"""
    safe_id = _ensure_safe_diagram_artifact_id(diagram_id)
    artifact_dir = _diagram_artifact_dir()
    svg_path = _find_diagram_artifact_path(artifact_dir, safe_id, _safe_diagram_project_dir(project_id), ".svg")
    if svg_path is not None:
        return _read_text_artifact_payload(svg_path, media_type="image/svg+xml")

    mermaid_path = _find_diagram_artifact_path(artifact_dir, safe_id, _safe_diagram_project_dir(project_id), ".mmd")
    if mermaid_path is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="图表 artifact 不存在", status_code=404)
    rendered_svg_path = mermaid_path.with_suffix(".svg")
    if rendered_svg_path.exists():
        return _read_text_artifact_payload(rendered_svg_path, media_type="image/svg+xml")
    return BidGeneratorFilePayload(
        content=_mermaid_to_fallback_svg(_read_text_file(mermaid_path), title="Mermaid 数据流图").encode("utf-8"),
        media_type="image/svg+xml",
        filename=f"{safe_id}.svg",
        inline=True,
        cache_control="public, max-age=86400",
    )


def get_mermaid_diagram_artifact_payload(diagram_id: str, *, project_id: str = "") -> BidGeneratorFilePayload:
    """读取 Mermaid 图表源码 artifact；入参为图表 ID/项目 ID，出参为 text/plain 响应元数据。"""
    safe_id = _ensure_safe_diagram_artifact_id(diagram_id)
    path = _find_diagram_artifact_path(_diagram_artifact_dir(), safe_id, _safe_diagram_project_dir(project_id), ".mmd")
    if path is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="Mermaid 图表 artifact 不存在", status_code=404)
    return _read_text_artifact_payload(path, media_type="text/plain; charset=utf-8")


def get_task_status_payload(
    task_id: str,
    *,
    project_id: str | None = None,
    after_event_id: int = 0,
) -> dict[str, Any]:
    """读取后台任务状态；入参为任务 ID/项目 ID/事件游标，出参兼容 legacy 轮询结构。"""
    task = _require_task_owner(task_id, project_id)
    started_at = _utc_iso_from_timestamp(getattr(task, "created_at", None))
    updated_at = _utc_iso_from_timestamp(getattr(task, "updated_at", None))
    normalized_after_event_id = _non_negative_int(after_event_id)
    partial_events = [
        event for event in (getattr(task, "partial_events", None) or [])
        if isinstance(event, dict) and _non_negative_int(event.get("event_id")) > normalized_after_event_id
    ]
    status = str(getattr(task, "status", "") or "error")
    stages = [str(item) for item in (getattr(task, "stages", None) or []) if not str(item).startswith("__text__")]
    return {
        "task_id": task_id,
        "status": status,
        "state": _task_status_to_api_state(status),
        "progress": 100 if status == "done" else 0,
        "current_stage": str(getattr(task, "current_stage", "") or ""),
        "stages": stages,
        "result": getattr(task, "result", None) if status == "done" else None,
        "partial_result": getattr(task, "partial_result", None) if status == "running" else None,
        "partial_events": partial_events,
        "last_partial_event_id": _non_negative_int(getattr(task, "partial_event_seq", 0)),
        "error": getattr(task, "error", None) if status in {"error", "timeout"} else None,
        "cancelled": status == "cancelled",
        "timed_out": status == "timeout",
        "cancellable": status == "running",
        "started_at": started_at,
        "updated_at": updated_at,
    }


async def cancel_task_payload(task_id: str, *, project_id: str | None = None) -> dict[str, Any]:
    """取消后台任务；入参为任务 ID/项目 ID，出参兼容 legacy cancel 响应。"""
    task_id_value = _required_string(task_id, field="task_id")
    normalized_project_id = str(project_id or "").strip() or None
    task = _require_task_owner(task_id_value, normalized_project_id)
    if str(getattr(task, "status", "") or "") != "running":
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="任务不存在或已完成", status_code=404)

    _persist_project_task_runtime(
        task,
        runtime_state="cancelling",
        message=str(getattr(task, "current_stage", "") or "任务取消中"),
        cancellable=False,
    )
    dify_stopped, remote_stop_status = await _stop_dify_workflows_for_task(task)
    task_manager = _task_manager()
    ok = bool(task_manager.cancel_task(task_id_value))
    if not ok and not dify_stopped:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="任务不存在或已完成", status_code=404)

    latest_task = task_manager.get_task(task_id_value) or task
    _persist_project_task_runtime(
        latest_task,
        runtime_state=_task_status_to_api_state(str(getattr(latest_task, "status", "") or "cancelled")),
        message=str(getattr(latest_task, "current_stage", "") or ""),
        cancellable=False,
    )
    cancelled_at = datetime.now(timezone.utc).isoformat()
    return {
        "cancelled": True,
        "task_id": task_id_value,
        "dify_stopped": bool(dify_stopped),
        "remote_stop_status": remote_stop_status,
        "task_state": _task_status_to_api_state(str(getattr(latest_task, "status", "") or "cancelled")),
        "phase": str(getattr(latest_task, "current_stage", "") or ""),
        "cancelled_at": cancelled_at,
    }


async def start_outline_task_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """启动大纲后台任务；入参为大纲生成 JSON，出参为 task_id 响应。"""
    payload = _json_object_body(body)
    dify_key = _get_workflow_key("structure_generator")
    if not dify_key:
        raise PlatformError(code="TASK_START_FAILED", message="大纲生成工作流 API Key 未配置", status_code=500)

    requirements = payload.get("requirements", []) if isinstance(payload.get("requirements"), list) else []
    bid_type = str(payload.get("bid_type") or "tech")
    use_knowledge = bool(payload.get("use_knowledge", True))
    analysis_context = str(payload.get("analysis_context") or "")
    expected_total_words = _int_or_default(payload.get("expected_total_words"), default=0)
    enable_diagrams = bool(payload.get("enable_diagrams", False) and _diagram_generation_enabled())
    max_diagrams = _int_or_default(payload.get("max_diagrams"), default=0) if enable_diagrams else 0
    scoring_details_json = str(payload.get("scoring_details_json") or "")
    structure_heading_seed_json = str(payload.get("structure_heading_seed_json") or "")
    technical_h2_bindings_json = str(payload.get("technical_h2_bindings_json") or "")
    technical_targets_json = str(payload.get("technical_targets_json") or "")
    outline_batch_strategy = str(payload.get("outline_batch_strategy", "auto") or "auto").strip().lower()
    outline_auto_parallel_threshold = _int_or_default(payload.get("outline_auto_parallel_threshold"), default=4)
    project_id = str(payload.get("project_id") or "").strip()
    await _ensure_project_slot_native(project_id, "outline")

    bundle = build_outline_generation_bundle(
        requirements=requirements,
        analysis_context=analysis_context,
        expected_total_words=expected_total_words,
        scoring_details_json=scoring_details_json,
        structure_heading_seed_json=structure_heading_seed_json,
        technical_h2_bindings_json=technical_h2_bindings_json,
        technical_targets_json=technical_targets_json,
    )
    inputs = dict(bundle["inputs"])
    inputs["bid_type"] = bid_type
    inputs["use_knowledge"] = "true" if use_knowledge else "false"
    inputs["enable_diagrams"] = "true" if enable_diagrams else "false"
    inputs["max_diagrams"] = max_diagrams
    outline_batches = _split_outline_seed_headings_local(
        bundle.get("seed_headings") or [],
        strategy=outline_batch_strategy,
        auto_threshold=outline_auto_parallel_threshold,
    )

    task_manager = _task_manager()
    task_id = task_manager.create_task("outline", project_id, workflow_name="structure_generator")
    _persist_project_runtime(
        project_id,
        task_id=task_id,
        task_type="outline",
        runtime_state="running",
        message="大纲生成中",
        cancellable=True,
    )

    async def run_task() -> None:
        execution_trace: list[dict[str, Any]] = []

        async def execute_outline_batch(
            *,
            batch_seed_headings: list[dict[str, Any]],
            batch_index: int,
            total_batches: int,
            started_at: float,
        ) -> list[dict[str, Any]]:
            batch_bundle = build_outline_generation_bundle(
                requirements=requirements,
                analysis_context=analysis_context,
                expected_total_words=expected_total_words,
                scoring_details_json=scoring_details_json,
                structure_heading_seed_json=_dump_structure_heading_seed_json_local(batch_seed_headings),
                technical_h2_bindings_json=_dump_structure_heading_seed_json_local(batch_seed_headings),
                technical_targets_json=technical_targets_json,
            )
            batch_inputs = dict(batch_bundle["inputs"])
            batch_inputs["bid_type"] = bid_type
            batch_inputs["use_knowledge"] = "true" if use_knowledge else "false"
            batch_inputs["enable_diagrams"] = "true" if enable_diagrams else "false"
            batch_inputs["max_diagrams"] = max_diagrams

            execution_trace.append(
                {
                    "kind": "batch_started",
                    "batch_index": int(batch_index),
                    "total_batches": int(total_batches),
                    "h2_count": len(batch_seed_headings or []),
                    "at": datetime.utcnow().isoformat(),
                    "elapsed_sec": int(max(0, time.monotonic() - started_at)),
                }
            )
            _push_task_event(task_id, "execution_trace", execution_trace[-1])
            outputs = await _collect_workflow_outputs(
                task_id,
                dify_key,
                batch_inputs,
                _r=None,
                initial_stage=f"✍️ 第 {batch_index}/{total_batches} 批大纲生成中",
            )
            sections = _resolve_outline_sections_from_outputs(
                outputs,
                seed_headings=batch_seed_headings,
                max_diagrams=0,
            )
            quality_report = evaluate_outline_quality(sections, batch_seed_headings)
            if not quality_report.get("pass"):
                raise RuntimeError(
                    f"第 {batch_index}/{total_batches} 批大纲结构质量校验失败："
                    + "; ".join(quality_report.get("issues") or [])
                )
            execution_trace.append(
                {
                    "kind": "batch_finished",
                    "batch_index": int(batch_index),
                    "total_batches": int(total_batches),
                    "h2_count": len(batch_seed_headings or []),
                    "at": datetime.utcnow().isoformat(),
                    "elapsed_sec": int(max(0, time.monotonic() - started_at)),
                }
            )
            _push_task_event(task_id, "execution_trace", execution_trace[-1])
            return sections

        try:
            started_at = time.monotonic()
            execution_trace.append(
                {
                    "kind": "outline_task_started",
                    "strategy": outline_batch_strategy,
                    "auto_threshold": outline_auto_parallel_threshold,
                    "total_batches": len(outline_batches),
                    "seed_h2_count": len(bundle.get("seed_headings") or []),
                    "at": datetime.utcnow().isoformat(),
                    "elapsed_sec": 0,
                }
            )
            _push_task_event(task_id, "execution_trace", execution_trace[-1])
            task_manager.update_stage(task_id, "📤 模型连接中")
            _emit_outline_stage_event_local(task_id, "📤 模型连接中", elapsed_sec=0)
            task_manager.update_stage(task_id, "🧠 模型预热中")
            _emit_outline_stage_event_local(task_id, "🧠 模型预热中", elapsed_sec=int(time.monotonic() - started_at))
            _push_task_event(task_id, "control", {"response_branch": "enabled" if bundle.get("enable_response_branch") else "skipped"})

            seed_sections = _make_h2_seed_sections_local(bundle.get("seed_headings") or [])
            if seed_sections:
                _push_task_event(task_id, "h2_seed", {"sections": seed_sections})
                task_manager.set_partial_result(
                    task_id,
                    {
                        "phase": "h2_seed_ready",
                        "sections": seed_sections,
                        "completeness": {"h2_ready": True, "h3_ready": False, "meta_ready": False},
                    },
                )

            task_manager.update_stage(task_id, "✍️ 生成大纲")
            _emit_outline_stage_event_local(task_id, "✍️ 生成大纲", elapsed_sec=int(time.monotonic() - started_at))

            if len(outline_batches) > 1:
                progressive_sections = _make_h2_seed_sections_local(bundle.get("seed_headings") or [])
                sec_by_id = {str(s.get("id") or ""): s for s in progressive_sections}
                total_batches = len(outline_batches)
                completed_batches = 0
                batch_results: dict[int, list[dict[str, Any]]] = {}
                batch_start_ts: dict[int, float] = {}

                async def run_outline_batch(batch_index: int, batch_seed_headings: list[dict[str, Any]]) -> tuple[int, list[dict[str, Any]]]:
                    return (
                        batch_index,
                        await execute_outline_batch(
                            batch_seed_headings=batch_seed_headings,
                            batch_index=batch_index,
                            total_batches=total_batches,
                            started_at=started_at,
                        ),
                    )

                batch_tasks = [
                    asyncio.create_task(run_outline_batch(batch_index, batch_seed_headings))
                    for batch_index, batch_seed_headings in enumerate(outline_batches, start=1)
                ]
                for batch_index, batch_seed_headings in enumerate(outline_batches, start=1):
                    batch_start_ts[batch_index] = time.monotonic()
                    _push_task_event(
                        task_id,
                        "outline_batch",
                        {
                            "batch_index": batch_index,
                            "total_batches": total_batches,
                            "status": "started",
                            "h2_count": len(batch_seed_headings or []),
                            "label": f"第 {batch_index}/{total_batches} 批已启动",
                            "elapsed_sec": int(time.monotonic() - started_at),
                        },
                    )
                try:
                    task_manager.update_stage(task_id, f"✍️ 并发生成 {total_batches} 批大纲")
                    _emit_outline_stage_event_local(task_id, f"✍️ 并发生成 {total_batches} 批大纲", elapsed_sec=int(time.monotonic() - started_at))
                    for done in asyncio.as_completed(batch_tasks):
                        _ensure_task_running(task_id)
                        batch_index, batch_sections = await done
                        batch_results[batch_index] = batch_sections
                        completed_batches += 1
                        _push_task_event(
                            task_id,
                            "outline_batch",
                            {
                                "batch_index": completed_batches,
                                "completed_batches": completed_batches,
                                "finished_batch_index": batch_index,
                                "total_batches": total_batches,
                                "status": "finished",
                                "batch_elapsed_sec": int(max(0, time.monotonic() - batch_start_ts.get(batch_index, started_at))),
                                "label": f"第 {batch_index}/{total_batches} 批已完成",
                                "elapsed_sec": int(time.monotonic() - started_at),
                            },
                        )
                        _emit_outline_stage_event_local(
                            task_id,
                            f"✍️ 已完成 {completed_batches}/{total_batches} 批大纲",
                            elapsed_sec=int(time.monotonic() - started_at),
                            heartbeat=completed_batches < total_batches,
                        )
                        for item in batch_sections:
                            sid = str(item.get("id") or "")
                            target = sec_by_id.get(sid)
                            if not target:
                                continue
                            target["children"] = item.get("children") or []
                            target["wordCount"] = int(item.get("wordCount") or 0)
                            target["writingHint"] = str(item.get("writingHint") or "")
                            target["keywords"] = item.get("keywords") or []
                            target["needDiagram"] = bool(item.get("needDiagram") or item.get("need_diagram") or False)
                            target["diagramBrief"] = str(item.get("diagramBrief") or item.get("diagram_brief") or "")
                            target["diagramPlan"] = item.get("diagramPlan") or item.get("diagram_plan") or {}
                        _push_task_event(
                            task_id,
                            "partial_outline",
                            {
                                "sections": progressive_sections,
                                "completeness": {
                                    "h2_ready": True,
                                    "h3_ready": completed_batches == total_batches,
                                    "meta_ready": completed_batches == total_batches,
                                },
                            },
                        )
                        task_manager.set_partial_result(
                            task_id,
                            {
                                "phase": f"outline_batch_{completed_batches}",
                                "sections": progressive_sections,
                                "completeness": {
                                    "h2_ready": True,
                                    "h3_ready": completed_batches == total_batches,
                                    "meta_ready": completed_batches == total_batches,
                                },
                            },
                        )
                except Exception:
                    for pending in batch_tasks:
                        if not pending.done():
                            pending.cancel()
                    await asyncio.gather(*batch_tasks, return_exceptions=True)
                    raise

                sections = [
                    section
                    for batch_index in range(1, total_batches + 1)
                    for section in (batch_results.get(batch_index) or [])
                ]
                normalize_outline_word_budget_dict(sections, expected_total_words)
                final_quality = evaluate_outline_quality(sections, bundle["seed_headings"])
                if not final_quality.get("pass"):
                    raise RuntimeError("分批大纲结构归一化后校验失败：" + "; ".join(final_quality.get("issues") or []))
                task_manager.update_stage(task_id, "✅ 大纲结构已就绪")
                _emit_outline_stage_event_local(task_id, "✅ 大纲结构已就绪", elapsed_sec=int(time.monotonic() - started_at))
                task_manager.set_result(
                    task_id,
                    {
                        "done": True,
                        "sections": sections,
                        "phase": "outline_finalized",
                        "execution_trace": execution_trace,
                        "batch_strategy": outline_batch_strategy,
                        "total_batches": len(outline_batches),
                    },
                )
                _sync_project_runtime_from_task(task_manager.get_task(task_id))
                return

            outputs = await _collect_workflow_outputs(
                task_id,
                dify_key,
                inputs,
                _r=None,
                initial_stage="✍️ 生成大纲",
            )
            sections = _resolve_outline_sections_from_outputs(
                outputs,
                seed_headings=bundle["seed_headings"],
                max_diagrams=max_diagrams if enable_diagrams else 0,
            )
            quality_report = evaluate_outline_quality(sections, bundle["seed_headings"])
            if not quality_report.get("pass"):
                raise RuntimeError("大纲结构质量校验失败：" + "; ".join(quality_report.get("issues") or []))

            progressive_sections = _make_h2_seed_sections_local(bundle.get("seed_headings") or [])
            sec_by_id = {str(s.get("id") or ""): s for s in progressive_sections}
            h3_batches = _outline_sections_window_batches_local(sections, window_size=2)
            for i, batch in enumerate(h3_batches, start=1):
                _ensure_task_running(task_id)
                batch_payload = []
                for item in batch:
                    sid = str(item.get("id") or "")
                    target = sec_by_id.get(sid)
                    if not target:
                        continue
                    target["children"] = [{"id": c.get("id", ""), "title": c.get("title", ""), "headingLevel": 3} for c in (item.get("children") or [])]
                    batch_payload.append({"id": sid, "title": target.get("title", ""), "children": target["children"]})
                _push_task_event(task_id, "h3_batch", {"window_index": i, "total_windows": len(h3_batches), "items": batch_payload})
                _push_task_event(
                    task_id,
                    "partial_outline",
                    {
                        "sections": progressive_sections,
                        "completeness": {"h2_ready": True, "h3_ready": i == len(h3_batches), "meta_ready": False},
                    },
                )
                task_manager.set_partial_result(
                    task_id,
                    {
                        "phase": "h3_generating",
                        "sections": progressive_sections,
                        "completeness": {"h2_ready": True, "h3_ready": i == len(h3_batches), "meta_ready": False},
                    },
                )
                _push_task_event(
                    task_id,
                    "stage",
                    {
                        "code": "outline_generating",
                        "label": "✍️ 生成大纲",
                        "phase": 2,
                        "percent": min(65, 25 + int(i * 40 / max(len(h3_batches), 1))),
                        "elapsed_sec": int(time.monotonic() - started_at),
                        "heartbeat": True,
                    },
                )

            meta_batches = _outline_sections_window_batches_local(sections, window_size=2)
            for i, batch in enumerate(meta_batches, start=1):
                _ensure_task_running(task_id)
                batch_payload = []
                for item in batch:
                    sid = str(item.get("id") or "")
                    target = sec_by_id.get(sid)
                    if not target:
                        continue
                    target["wordCount"] = int(item.get("wordCount") or 0)
                    target["writingHint"] = str(item.get("writingHint") or "")
                    target["keywords"] = item.get("keywords") or []
                    target["needDiagram"] = bool(item.get("needDiagram") or item.get("need_diagram") or False)
                    target["diagramBrief"] = str(item.get("diagramBrief") or item.get("diagram_brief") or "")
                    target["diagramPlan"] = item.get("diagramPlan") or item.get("diagram_plan") or {}
                    child_map = {str(c.get("id") or ""): c for c in (target.get("children") or [])}
                    for child in item.get("children") or []:
                        cid = str(child.get("id") or "")
                        if cid and cid in child_map:
                            child_map[cid]["wordCount"] = int(child.get("wordCount") or 0)
                            child_map[cid]["writingHint"] = str(child.get("writingHint") or "")
                            child_map[cid]["keywords"] = child.get("keywords") or []
                            child_map[cid]["needDiagram"] = bool(child.get("needDiagram") or child.get("need_diagram") or False)
                            child_map[cid]["diagramBrief"] = str(child.get("diagramBrief") or child.get("diagram_brief") or "")
                            child_map[cid]["diagramPlan"] = child.get("diagramPlan") or child.get("diagram_plan") or {}
                    batch_payload.append({"id": sid, "wordCount": target.get("wordCount", 0), "keywords": target.get("keywords", [])})
                _push_task_event(task_id, "meta_batch", {"window_index": i, "total_windows": len(meta_batches), "items": batch_payload})
                _push_task_event(
                    task_id,
                    "partial_outline",
                    {
                        "sections": progressive_sections,
                        "completeness": {"h2_ready": True, "h3_ready": True, "meta_ready": i == len(meta_batches)},
                    },
                )
                task_manager.set_partial_result(
                    task_id,
                    {
                        "phase": "h3_meta_generating",
                        "sections": progressive_sections,
                        "completeness": {"h2_ready": True, "h3_ready": True, "meta_ready": i == len(meta_batches)},
                    },
                )
                _push_task_event(
                    task_id,
                    "stage",
                    {
                        "code": "outline_generating",
                        "label": "✍️ 生成大纲",
                        "phase": 3,
                        "percent": min(90, 70 + int(i * 20 / max(len(meta_batches), 1))),
                        "elapsed_sec": int(time.monotonic() - started_at),
                        "heartbeat": True,
                    },
                )

            task_manager.update_stage(task_id, "🧾 大纲归一化中")
            _push_task_event(task_id, "stage", {"code": "outline_finalized", "label": "大纲归一化中", "phase": 4, "percent": 95, "elapsed_sec": int(time.monotonic() - started_at)})
            normalize_outline_word_budget_dict(sections, expected_total_words)
            final_quality = evaluate_outline_quality(sections, bundle["seed_headings"])
            if not final_quality.get("pass"):
                raise RuntimeError("大纲结构归一化后校验失败：" + "; ".join(final_quality.get("issues") or []))
            task_manager.update_stage(task_id, "✅ 大纲结构已就绪")
            _emit_outline_stage_event_local(task_id, "✅ 大纲结构已就绪", elapsed_sec=int(time.monotonic() - started_at))
            task_manager.set_result(
                task_id,
                {
                    "done": True,
                    "sections": sections,
                    "phase": "outline_finalized",
                    "execution_trace": execution_trace,
                    "total_batches": len(outline_batches),
                },
            )
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except asyncio.CancelledError:
            await _best_effort_stop_dify_by_task_id(task_id)
            logger.info("[Task %s] 大纲生成任务被用户取消", task_id)
            task_manager.set_cancelled(task_id)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except Exception as exc:
            logger.error("[Task %s] 大纲生成后台任务失败: %s", task_id, exc, exc_info=True)
            task_manager.set_error(task_id, _format_dify_runtime_error(exc))
            _sync_project_runtime_from_task(task_manager.get_task(task_id))

    background_task = asyncio.create_task(run_task())
    task_manager.set_async_task(task_id, background_task)
    return {"task_id": task_id}


async def start_content_task_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """启动正文后台任务；入参为正文生成 JSON，出参为 legacy task_id 响应。"""
    payload = _json_object_body(body)
    try:
        _validate_required_bidder_info(payload.get("bidder_info", {}) or {})
    except Exception as exc:
        detail = str(exc)
        if detail:
            raise PlatformError(code="INVALID_REQUEST", message=detail, status_code=400) from exc
        raise

    generation_strategy = str(payload.get("generation_strategy") or "general").strip() or "general"
    workflow_name = _resolve_content_workflow_name(generation_strategy)
    strip_structural_numbering = workflow_name == "response_content_writer"
    dify_key = _get_workflow_key(workflow_name)
    if not dify_key:
        raise PlatformError(code="TASK_START_FAILED", message=f"{workflow_name} 工作流 API Key 未配置", status_code=500)

    project_id = _required_string(payload.get("project_id"), field="project_id")
    await _ensure_project_slot_native(project_id, "content")
    section_id = str(payload.get("section_id") or "")
    section_title = str(payload.get("section_title") or "")
    expected_words = _int_or_default(payload.get("expected_words"), default=500)
    keywords = str(payload.get("keywords") or "").strip() or section_title
    analysis_context = str(payload.get("analysis_context") or "")
    slice_text = str(payload.get("section_outline_slice") or "")
    writing_hint = _compose_runtime_writing_hint(
        str(payload.get("writing_hint") or ""),
        section_title,
        expected_words,
        keywords,
        section_outline_slice=slice_text,
        analysis_context=analysis_context,
    )

    inputs = {
        "section_title": section_title,
        "writing_hint": writing_hint,
        "keywords": keywords,
        "expected_words": expected_words,
        "project_summary": payload.get("project_summary", ""),
        "global_outline": payload.get("global_outline", ""),
        "placeholder_hint": str(payload.get("placeholder_hint") or ""),
    }
    if workflow_name == "content_writer":
        inputs["requires_search"] = "true" if bool(payload.get("requires_search", True)) else "false"
        inputs["image_map_hint"] = str(payload.get("image_map_hint", "") or "")

    enable_diagrams = bool(payload.get("enable_diagrams", False) and _diagram_generation_enabled())
    max_diagrams = _int_or_default(payload.get("max_diagrams"), default=0) if enable_diagrams else 0
    need_diagram = bool(payload.get("need_diagram", False) and enable_diagrams)
    diagram_brief = str(payload.get("diagram_brief", "") or "") if enable_diagrams else ""
    diagram_type_hint = str(payload.get("diagram_type_hint", "architecture") or "architecture")
    raw_keywords = str(payload.get("keywords", "") or "")
    raw_global_outline = str(payload.get("global_outline", "") or "")
    defer_diagram = bool(payload.get("defer_diagram", False))
    request_mapping_flat = _string_mapping(payload.get("mapping_table"))
    try:
        request_mapping_flat, merged_placeholder_hint, _bidder_context = _merge_bidder_pipt_context(
            mapping_table=request_mapping_flat,
            placeholder_hint=str(payload.get("placeholder_hint", "") or ""),
            bidder_info=payload.get("bidder_info", {}) or {},
        )
        inputs["placeholder_hint"] = merged_placeholder_hint
    except Exception:
        logger.warning("投标人信息 PIPT 归一化失败，正文任务使用请求原始占位符上下文", exc_info=True)

    task_manager = _task_manager()
    task_id = task_manager.create_task("content", project_id, workflow_name=workflow_name)
    _persist_project_runtime(
        project_id,
        task_id=task_id,
        task_type="content",
        runtime_state="running",
        message="正文生成中",
        cancellable=True,
    )

    async def run_task() -> None:
        try:
            task_manager.update_stage(
                task_id,
                "🧠 响应情况正文生成中" if workflow_name == "response_content_writer" else "🔍 知识检索与工作流执行中",
            )
            wants_diagram = workflow_name == "content_writer" and enable_diagrams and need_diagram and bool(diagram_brief.strip()) and max_diagrams > 0
            diagram_key = _get_workflow_key(_get_diagram_workflow_name()) if wants_diagram else ""
            can_defer_diagram = bool(wants_diagram and diagram_key)
            diagram_skip = None
            if workflow_name == "content_writer" and bool(payload.get("enable_diagrams")) and bool(payload.get("need_diagram")) and not can_defer_diagram:
                diagram_skip = _build_diagram_skip_payload(
                    workflow_name=workflow_name,
                    enable_diagrams=enable_diagrams,
                    need_diagram=need_diagram,
                    diagram_brief=diagram_brief,
                    max_diagrams=max_diagrams,
                    diagram_key=diagram_key,
                )

            outputs: dict[str, Any] = {}
            got_finished = False
            async for chunk in _call_dify_workflow_stream(dify_key, inputs):
                _ensure_task_running(task_id)
                if isinstance(chunk, dict):
                    if chunk.get("dify_task_id"):
                        task_manager.set_dify_task_id(task_id, chunk["dify_task_id"])
                    if chunk.get("__stage__"):
                        task_manager.update_stage(task_id, chunk["__stage__"])
                    elif chunk.get("__finished__"):
                        outputs = chunk.get("outputs", {}) if isinstance(chunk.get("outputs"), dict) else {}
                        got_finished = True
                        break
            if not got_finished:
                raise RuntimeError("内容工作流异常中断（未收到 finished 事件）")

            task_manager.update_stage(task_id, "📝 解析生成结果")
            content, replace_report, placeholder_warning = _finalize_legacy_content_output(
                outputs.get("text") or outputs.get("result") or outputs.get("structured_output") or "",
                section_title,
                feedback=str(outputs.get("feedback") or ""),
                request_mapping_flat=request_mapping_flat,
                strip_structural_numbering=strip_structural_numbering,
                audit_source="apps_api.task.start_content",
            )
            content, referenced_images = _normalize_referenced_images(content)
            quality_score = None
            raw_score = outputs.get("quality_score")
            if raw_score is not None:
                try:
                    quality_score = int(float(raw_score))
                except (TypeError, ValueError):
                    quality_score = None
            word_count = _count_visible_chars(content)

            diagrams_generated = 0
            diagram_error = None
            diagram_specs = _extract_content_diagram_specs(outputs)
            if can_defer_diagram and not defer_diagram:
                content, diagrams_generated, diagram_error, diagram_specs = await _run_inline_content_diagram(
                    payload=payload,
                    workflow_name=workflow_name,
                    content=content,
                    writing_hint=writing_hint,
                    outputs=outputs,
                )
                content, referenced_images = _normalize_referenced_images(content)
                word_count = _count_visible_chars(content)

            if can_defer_diagram and defer_diagram:
                task_manager.update_stage(task_id, "✅ 正文已生成（图表将在独立任务中生成）")
            elif not can_defer_diagram:
                task_manager.update_stage(task_id, "✅ 正文已生成")
            else:
                task_manager.update_stage(task_id, f"✅ 正文与图表已生成（{diagrams_generated} 张）" if diagrams_generated else "✅ 正文已生成")

            partial_payload: dict[str, Any] = {
                "partial": True,
                "phase": "diagram_ready" if diagrams_generated else "text_ready",
                "section_id": section_id,
                "content": content,
                "word_count": word_count,
                "quality_score": quality_score,
                "feedback": outputs.get("feedback") or None,
                "replace_report": replace_report,
                "placeholder_warning": placeholder_warning,
                "referenced_images": referenced_images,
                "diagrams_count": diagrams_generated,
            }
            if diagram_skip:
                partial_payload["diagram_skip"] = diagram_skip
            if diagram_error:
                partial_payload["diagram_error"] = diagram_error
            task_manager.set_partial_result(task_id, partial_payload)

            done_payload: dict[str, Any] = {
                "done": True,
                "section_id": section_id,
                "content": content,
                "word_count": word_count,
                "quality_score": quality_score,
                "feedback": outputs.get("feedback") or None,
                "replace_report": replace_report,
                "placeholder_warning": placeholder_warning,
                "referenced_images": referenced_images,
                "diagrams_count": diagrams_generated,
            }
            if diagram_skip:
                done_payload["diagram_skip"] = diagram_skip
            if diagram_error:
                done_payload["diagram_error"] = diagram_error
            if can_defer_diagram:
                if defer_diagram:
                    done_payload["diagram_deferred"] = True
                    done_payload["diagram_request"] = {
                        "section_id": section_id,
                        "section_title": section_title,
                        "base_content": content,
                        "writing_hint": writing_hint,
                        "keywords": raw_keywords,
                        "global_outline": raw_global_outline,
                        "diagram_brief": diagram_brief,
                        "diagram_type_hint": diagram_type_hint,
                        "diagram_specs": diagram_specs,
                        "quality_score": quality_score,
                        "feedback": outputs.get("feedback") or None,
                        "replace_report": replace_report,
                        "placeholder_warning": placeholder_warning,
                    }
                    if diagram_specs:
                        done_payload["diagram_specs"] = diagram_specs
                elif diagram_specs:
                    done_payload["diagram_specs"] = diagram_specs
            _persist_content_result_to_project(project_id, section_id, done_payload, status="done")
            task_manager.set_result(task_id, done_payload)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except asyncio.CancelledError:
            await _best_effort_stop_dify_by_task_id(task_id)
            logger.info("[Task %s] 内容生成任务被用户取消", task_id)
            task_manager.set_cancelled(task_id)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except Exception as exc:
            logger.error("[Task %s] 内容生成后台任务失败: %s", task_id, exc, exc_info=True)
            _persist_content_result_to_project(project_id, section_id, {}, status="error", error=_format_dify_runtime_error(exc))
            task_manager.set_error(task_id, _format_dify_runtime_error(exc))
            _sync_project_runtime_from_task(task_manager.get_task(task_id))

    background_task = asyncio.create_task(run_task())
    task_manager.set_async_task(task_id, background_task)
    return {"task_id": task_id, "section_id": section_id}


async def start_content_rewrite_task_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """启动单章节重生成任务；入参为重生成 JSON，出参为 legacy task_id 响应。"""
    payload = _json_object_body(body)
    try:
        _validate_required_bidder_info(payload.get("bidder_info", {}) or {})
    except Exception as exc:
        detail = str(exc)
        if detail:
            raise PlatformError(code="INVALID_REQUEST", message=detail, status_code=400) from exc
        raise

    dify_key = _get_workflow_key("content_rewrite")
    if not dify_key:
        raise PlatformError(code="TASK_START_FAILED", message="content_rewrite 工作流 API Key 未配置", status_code=500)

    project_id = _required_string(payload.get("project_id"), field="project_id")
    await _ensure_project_slot_native(project_id, "content")
    section_id = str(payload.get("section_id") or "").strip()
    section_title = str(payload.get("section_title") or "").strip()
    current_content = str(payload.get("current_content") or "")
    current_text, diagram_suffix = _split_diagram_blocks(current_content)
    if not current_text.strip():
        raise PlatformError(code="INVALID_REQUEST", message="current_content 不能为空", status_code=400)

    expected_words = _int_or_default(payload.get("expected_words"), default=0)
    rewrite_instruction = str(payload.get("rewrite_instruction") or "").strip()
    request_mapping_flat = _string_mapping(payload.get("mapping_table"))
    rewrite_placeholder_hint = str(payload.get("placeholder_hint") or "")
    try:
        request_mapping_flat, rewrite_placeholder_hint, _bidder_context = _merge_bidder_pipt_context(
            mapping_table=request_mapping_flat,
            placeholder_hint=rewrite_placeholder_hint,
            bidder_info=payload.get("bidder_info", {}) or {},
        )
    except Exception:
        logger.warning("投标人信息 PIPT 归一化失败，重生成任务使用请求原始占位符上下文", exc_info=True)
    strip_structural_numbering = str(payload.get("generation_strategy", "general") or "general").strip() == "response_special"

    task_manager = _task_manager()
    task_id = task_manager.create_task("content", project_id, workflow_name="content_rewrite")
    _persist_project_runtime(
        project_id,
        task_id=task_id,
        task_type="content",
        runtime_state="running",
        message=f"{section_title or section_id or '章节'} 重生成中",
        cancellable=True,
    )

    async def run_task() -> None:
        try:
            inputs = {
                "section_id": section_id,
                "section_title": section_title,
                "current_content": current_text,
                "rewrite_instruction": rewrite_instruction,
                "expected_words": expected_words,
                "project_summary": payload.get("project_summary", ""),
                "global_outline": payload.get("global_outline", ""),
                "section_outline_slice": payload.get("section_outline_slice", ""),
                "analysis_context": payload.get("analysis_context", ""),
                "placeholder_hint": rewrite_placeholder_hint,
            }
            task_manager.update_stage(task_id, f"🪄 正在重生成：{section_title or section_id or '未命名章节'}")
            outputs: dict[str, Any] = {}
            got_finished = False
            async for chunk in _call_dify_workflow_stream(dify_key, inputs):
                _ensure_task_running(task_id)
                if isinstance(chunk, dict):
                    if chunk.get("dify_task_id"):
                        task_manager.set_dify_task_id(task_id, chunk["dify_task_id"])
                    if chunk.get("__stage__"):
                        task_manager.update_stage(task_id, chunk["__stage__"])
                    elif chunk.get("__finished__"):
                        outputs = chunk.get("outputs", {}) if isinstance(chunk.get("outputs"), dict) else {}
                        got_finished = True
                        break
            if not got_finished:
                raise RuntimeError("内容工作流异常中断（未收到 finished 事件）")

            rewritten, replace_report, placeholder_warning = _finalize_legacy_content_output(
                outputs.get("text") or outputs.get("result") or outputs.get("structured_output") or "",
                section_title or section_id,
                feedback=str(outputs.get("feedback") or ""),
                request_mapping_flat=request_mapping_flat,
                strip_structural_numbering=strip_structural_numbering,
                audit_source="apps_api.task.start_content_rewrite",
            )
            if diagram_suffix:
                rewritten = f"{rewritten}\n{diagram_suffix}".strip() if rewritten else diagram_suffix
            payload_result: dict[str, Any] = {
                "done": True,
                "section_id": section_id,
                "content": rewritten,
                "word_count": _count_visible_chars(rewritten),
                "quality_score": None,
                "feedback": outputs.get("feedback") or None,
                "replace_report": replace_report,
                "placeholder_warning": placeholder_warning,
            }
            raw_score = outputs.get("quality_score")
            if raw_score is not None:
                try:
                    payload_result["quality_score"] = int(float(raw_score))
                except (TypeError, ValueError):
                    payload_result["quality_score"] = None
            _persist_content_result_to_project(project_id, section_id, payload_result, status="done")
            task_manager.set_result(task_id, payload_result)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except asyncio.CancelledError:
            await _best_effort_stop_dify_by_task_id(task_id)
            logger.info("[Task %s] 单章节重生成任务被用户取消", task_id)
            task_manager.set_cancelled(task_id)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except Exception as exc:
            logger.error("[Task %s] 单章节重生成任务失败: %s", task_id, exc, exc_info=True)
            _persist_content_result_to_project(project_id, section_id, {}, status="error", error=_format_dify_runtime_error(exc))
            task_manager.set_error(task_id, _format_dify_runtime_error(exc))
            _sync_project_runtime_from_task(task_manager.get_task(task_id))

    background_task = asyncio.create_task(run_task())
    task_manager.set_async_task(task_id, background_task)
    return {"task_id": task_id, "section_id": section_id}


async def start_content_group_task_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """启动 H2 分组正文任务；入参为分组生成 JSON，出参为 task_id 响应。"""
    payload = _json_object_body(body)
    try:
        _validate_required_bidder_info(payload.get("bidder_info", {}) or {})
    except Exception as exc:
        detail = str(exc)
        if detail:
            raise PlatformError(code="INVALID_REQUEST", message=detail, status_code=400) from exc
        raise

    dify_key = _get_workflow_key("content_group_writer")
    if not dify_key:
        raise PlatformError(code="TASK_START_FAILED", message="content_group_writer 工作流 API Key 未配置", status_code=500)

    project_id = _required_string(payload.get("project_id"), field="project_id")
    await _ensure_project_slot_native(project_id, "content")

    group_id = str(payload.get("group_id") or "").strip() or uuid.uuid4().hex[:8]
    group_title = str(payload.get("group_title") or "").strip() or "未命名分组"
    children = _build_group_writing_children(payload.get("children") or [])
    if not children:
        raise PlatformError(code="INVALID_REQUEST", message="children 不能为空", status_code=400)

    request_mapping_flat = _string_mapping(payload.get("mapping_table"))
    group_placeholder_hint = str(payload.get("placeholder_hint") or "")
    try:
        request_mapping_flat, group_placeholder_hint, _bidder_context = _merge_bidder_pipt_context(
            mapping_table=request_mapping_flat,
            placeholder_hint=group_placeholder_hint,
            bidder_info=payload.get("bidder_info", {}) or {},
        )
    except Exception:
        logger.warning("投标人信息 PIPT 归一化失败，分组正文任务使用请求原始占位符上下文", exc_info=True)

    task_manager = _task_manager()
    task_id = task_manager.create_task("content", project_id, workflow_name="content_group_writer")
    _persist_project_runtime(
        project_id,
        task_id=task_id,
        task_type="content",
        runtime_state="running",
        message=f"{group_title} 正文批量生成中",
        cancellable=True,
    )

    async def run_task() -> None:
        try:
            shared_analysis_context = _dedupe_join([child.get("analysis_context", "") for child in children], max_len=9000)
            group_outline_slice = _dedupe_join(
                [str(payload.get("global_outline", "") or "").strip()] + [str(child.get("section_outline_slice") or "") for child in children],
                max_len=2600,
            )
            group_search_query = _build_group_search_query(group_title, children)
            enable_diagrams = bool(payload.get("enable_diagrams", False) and _diagram_generation_enabled())
            max_diagrams = _int_or_default(payload.get("max_diagrams"), default=0) if enable_diagrams else 0
            diagram_key = _get_workflow_key(_get_diagram_workflow_name()) if enable_diagrams and max_diagrams > 0 else ""

            group_inputs = {
                "group_id": group_id,
                "group_title": group_title,
                "expected_total_words": sum(max(0, int(child.get("expected_words") or 0)) for child in children),
                "project_summary": payload.get("project_summary", ""),
                "global_outline": group_outline_slice,
                "placeholder_hint": group_placeholder_hint,
                "requires_search": "true" if bool(payload.get("requires_search", False)) else "false",
                "group_analysis_context": shared_analysis_context,
                "search_query": group_search_query,
                "children_json": json.dumps(
                    [
                        {
                            "section_id": child["section_id"],
                            "section_title": child["section_title"],
                            "keywords": child["keywords"],
                            "expected_words": child["expected_words"],
                            "writing_hint": child["writing_hint"],
                        }
                        for child in children
                    ],
                    ensure_ascii=False,
                ),
            }
            outputs = await _collect_workflow_outputs(
                task_id,
                dify_key,
                group_inputs,
                _r=None,
                initial_stage=f"📦 H2 批量生成中：{group_title}",
            )
            parsed = _parse_group_content_results(outputs, children, request_mapping_flat)
            results = list(parsed.get("sections") or [])
            failed_sections = list(parsed.get("failed_sections") or [])
            parse_error = str(parsed.get("parse_error") or "").strip()
            if parse_error:
                summary = _summarize_workflow_outputs(outputs)
                logger.warning("[Task %s] H2 批量正文解析存在缺失: %s; 返回摘要: %s", task_id, parse_error, summary)
                if results:
                    task_manager.update_stage(task_id, f"⚠️ 批量结果不完整，已保留成功章节（{len(results)}/{len(children)}）")
                else:
                    task_manager.update_stage(task_id, "⚠️ 批量结果无可用正文，已标记章节失败")

            repaired_sections, failed_sections = await _repair_group_failed_sections(
                task_id=task_id,
                children=children,
                failed_sections=failed_sections,
                request=payload,
                request_mapping_flat=request_mapping_flat,
                group_placeholder_hint=group_placeholder_hint,
                group_outline_slice=group_outline_slice,
            )
            if repaired_sections:
                repaired_ids = {str(row.get("section_id") or "") for row in repaired_sections}
                results = [row for row in results if str(row.get("section_id") or "") not in repaired_ids]
                results.extend(repaired_sections)
                task_manager.update_stage(task_id, f"🩹 已补生成缺失子章节（{len(repaired_sections)} 个）")

            child_map = {child["section_id"]: child for child in children}
            ordered_results = sorted(
                results,
                key=lambda row: int(child_map.get(str(row.get("section_id") or ""), {}).get("diagram_priority", 0)),
                reverse=True,
            )
            final_by_id: dict[str, dict[str, Any]] = {str(row.get("section_id") or ""): dict(row) for row in results}
            for done_count, row in enumerate(ordered_results, start=1):
                section_id = str(row.get("section_id") or "")
                child = child_map.get(section_id)
                if not child:
                    continue
                content = str(row.get("content") or "")
                child_need_diagram = bool(child.get("need_diagram"))
                child_diagram_brief = str(child.get("diagram_brief") or "").strip()
                child_wants_diagram = enable_diagrams and child_need_diagram and bool(child_diagram_brief) and max_diagrams > 0
                child_can_generate_diagram = bool(child_wants_diagram and diagram_key)
                child_should_report_diagram_skip = bool(payload.get("enable_diagrams", False)) and child_need_diagram
                child_diagram_skip = None
                if child_should_report_diagram_skip and not child_can_generate_diagram:
                    child_diagram_skip = _build_diagram_skip_payload(
                        workflow_name="content_writer",
                        enable_diagrams=enable_diagrams,
                        need_diagram=child_need_diagram,
                        diagram_brief=child_diagram_brief,
                        max_diagrams=max_diagrams,
                        diagram_key=diagram_key if child_wants_diagram else "",
                    )

                diagrams_generated: list[dict[str, Any]] = []
                if child_can_generate_diagram:
                    diagram_specs = row.get("diagram_specs") or row.get("diagram_spec")
                    diagrams_generated, diagram_slot_reserved, diagram_error = await _execute_diagram_for_section(
                        task_id=task_id,
                        project_id=project_id,
                        diagram_key=diagram_key,
                        enable_diagrams=enable_diagrams,
                        need_diagram=child_need_diagram,
                        diagram_brief=child_diagram_brief,
                        max_diagrams=max_diagrams,
                        diagram_type_hint=str(child.get("diagram_type_hint") or "architecture"),
                        section_title=str(child.get("section_title") or ""),
                        writing_hint=str(child.get("writing_hint") or ""),
                        raw_keywords=str(child.get("keywords") or ""),
                        raw_global_outline=group_outline_slice,
                        content_context=content,
                        diagram_specs=diagram_specs,
                    )
                    if not diagrams_generated and diagram_slot_reserved:
                        await task_manager.release_diagram_slot(project_id)
                    if diagram_error:
                        row["diagram_error"] = diagram_error
                    if diagrams_generated:
                        content = content + "\n" + "\n".join(_build_diagram_reference_tag(item) for item in diagrams_generated)
                        row["content"] = content
                        row["word_count"] = _count_visible_chars(content)
                if child_diagram_skip:
                    row["diagram_skip"] = child_diagram_skip
                final_by_id[section_id] = row
                task_manager.append_partial_event(
                    task_id,
                    {
                        "partial": True,
                        "phase": "group_child_done",
                        "group_id": group_id,
                        "section_id": section_id,
                        "content": row.get("content") or "",
                        "word_count": row.get("word_count") or 0,
                        "quality_score": row.get("quality_score"),
                        "feedback": row.get("feedback"),
                        "replace_report": row.get("replace_report") or [],
                        "diagrams_count": len(diagrams_generated),
                        "diagram_error": row.get("diagram_error"),
                        "diagram_skip": row.get("diagram_skip"),
                        "done_count": done_count,
                        "total_count": len(children),
                    },
                )

            rank = {child["section_id"]: idx for idx, child in enumerate(children)}
            results = sorted(final_by_id.values(), key=lambda row: rank.get(str(row.get("section_id") or ""), 9999))
            failed_sections.sort(key=lambda row: rank.get(str(row.get("section_id") or ""), 9999))

            result_payload = {
                "done": True,
                "group_id": group_id,
                "group_title": group_title,
                "sections": results,
                "sections_count": len(results),
                "failed_sections": failed_sections,
                "failed_count": len(failed_sections),
                "partial_success": bool(results) and bool(failed_sections),
            }
            _persist_group_content_result_to_project(project_id, results, failed_sections)
            task_manager.set_result(task_id, result_payload)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except asyncio.CancelledError:
            await _best_effort_stop_dify_by_task_id(task_id)
            logger.info("[Task %s] H2 批量正文任务被用户取消", task_id)
            task_manager.set_cancelled(task_id)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except Exception as exc:
            logger.error("[Task %s] H2 批量正文任务失败: %s", task_id, exc, exc_info=True)
            task_manager.set_error(task_id, _format_dify_runtime_error(exc))
            _sync_project_runtime_from_task(task_manager.get_task(task_id))

    background_task = asyncio.create_task(run_task())
    task_manager.set_async_task(task_id, background_task)
    return {"task_id": task_id, "group_id": group_id}


async def start_group_review_task_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """启动 H2 分组评估任务；入参为评估 JSON，出参为 task_id 响应。"""
    payload = _json_object_body(body)
    dify_key = _get_workflow_key("group_review_writer")
    if not dify_key:
        raise PlatformError(code="TASK_START_FAILED", message="group_review_writer 工作流 API Key 未配置", status_code=500)

    project_id = _required_string(payload.get("project_id"), field="project_id")
    await _ensure_project_slot_native(project_id, "content")
    group_id = str(payload.get("group_id") or "").strip() or uuid.uuid4().hex[:8]
    group_title = str(payload.get("group_title") or "").strip() or "未命名章节"
    sections = payload.get("sections") or []
    if not isinstance(sections, list) or not sections:
        raise PlatformError(code="INVALID_REQUEST", message="sections 不能为空", status_code=400)

    task_manager = _task_manager()
    task_id = task_manager.create_task("content", project_id, workflow_name="group_review_writer")
    _persist_project_runtime(
        project_id,
        task_id=task_id,
        task_type="content",
        runtime_state="running",
        message=f"{group_title} 评估中",
        cancellable=True,
    )

    async def run_task() -> None:
        try:
            inputs = {
                "group_id": group_id,
                "group_title": group_title,
                "project_summary": payload.get("project_summary", ""),
                "group_outline": payload.get("group_outline", ""),
                "group_analysis_context": payload.get("group_analysis_context", ""),
                "sections_json": json.dumps(sections, ensure_ascii=False),
            }
            outputs = await _collect_workflow_outputs(
                task_id,
                dify_key,
                inputs,
                _r=None,
                initial_stage=f"🧾 H2 章节评估中：{group_title}",
            )
            result_payload = _parse_group_review_result(outputs)
            result_payload.update(
                {
                    "done": True,
                    "group_id": group_id,
                    "group_title": group_title,
                }
            )
            task_manager.set_result(task_id, result_payload)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except asyncio.CancelledError:
            await _best_effort_stop_dify_by_task_id(task_id)
            logger.info("[Task %s] H2 分组评估任务被用户取消", task_id)
            task_manager.set_cancelled(task_id)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except Exception as exc:
            logger.error("[Task %s] H2 分组评估任务失败: %s", task_id, exc, exc_info=True)
            task_manager.set_error(task_id, _format_dify_runtime_error(exc))
            _sync_project_runtime_from_task(task_manager.get_task(task_id))

    background_task = asyncio.create_task(run_task())
    task_manager.set_async_task(task_id, background_task)
    return {"task_id": task_id, "group_id": group_id}


async def start_diagram_task_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """启动独立图表任务；入参为图表生成 JSON，出参为 task_id 响应。"""
    payload = _json_object_body(body)
    project_id = str(payload.get("project_id") or "").strip()
    section_id = str(payload.get("section_id") or "")
    base_content = str(payload.get("base_content") or "")

    task_manager = _task_manager()
    if not _diagram_generation_enabled():
        task_id = task_manager.create_task("diagram", project_id, workflow_name=_get_diagram_workflow_name())
        _persist_project_runtime(
            project_id,
            task_id=task_id,
            task_type="diagram",
            runtime_state="succeeded",
            message="图表生成已禁用，保留正文",
            cancellable=False,
        )
        result_payload = {
            "done": True,
            "section_id": section_id,
            "content": base_content,
            "word_count": _count_visible_chars(base_content),
            "quality_score": payload.get("quality_score"),
            "feedback": payload.get("feedback"),
            "replace_report": payload.get("replace_report", []) or [],
            "diagrams_count": 0,
        }
        task_manager.set_result(task_id, result_payload)
        _sync_project_runtime_from_task(task_manager.get_task(task_id))
        return {"task_id": task_id, "section_id": section_id}

    diagram_key = _get_workflow_key(_get_diagram_workflow_name())
    if not diagram_key:
        raise PlatformError(code="TASK_START_FAILED", message=f"{_get_diagram_workflow_name()} 工作流 API Key 未配置", status_code=500)

    await _ensure_project_slot_native(project_id, "diagram")
    enable_diagrams = bool(payload.get("enable_diagrams", False) and _diagram_generation_enabled())
    task_id = task_manager.create_task("diagram", project_id, workflow_name=_get_diagram_workflow_name())
    _persist_project_runtime(
        project_id,
        task_id=task_id,
        task_type="diagram",
        runtime_state="running",
        message="图表生成中",
        cancellable=True,
    )

    async def run_task() -> None:
        try:
            task_manager.update_stage(task_id, "🎨 独立图表任务启动")
            result_payload = await _run_diagram_request(task_id, {**payload, "enable_diagrams": enable_diagrams}, diagram_key)
            if result_payload.get("diagram_error"):
                task_manager.update_stage(task_id, "⚠️ 图表生成失败，已保留正文")
            task_manager.set_result(task_id, result_payload)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except asyncio.CancelledError:
            if project_id:
                await task_manager.release_diagram_slot(project_id)
            logger.info("[Task %s] 图表任务被用户取消", task_id)
            task_manager.set_cancelled(task_id)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except Exception as exc:
            if project_id:
                await task_manager.release_diagram_slot(project_id)
            logger.error("[Task %s] 图表后台任务失败: %s", task_id, exc, exc_info=True)
            task_manager.set_result(
                task_id,
                {
                    "done": True,
                    "section_id": section_id,
                    "content": base_content,
                    "word_count": _count_visible_chars(base_content),
                    "quality_score": payload.get("quality_score"),
                    "feedback": payload.get("feedback"),
                    "replace_report": payload.get("replace_report", []) or [],
                    "diagrams_count": 0,
                    "diagram_error": _build_diagram_error_payload(exc, str(payload.get("section_title") or "")),
                },
            )
            _sync_project_runtime_from_task(task_manager.get_task(task_id))

    background_task = asyncio.create_task(run_task())
    task_manager.set_async_task(task_id, background_task)
    return {"task_id": task_id, "section_id": section_id}


async def start_diagram_batch_task_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """启动批量图表任务；入参为批量图表 JSON，出参为 task_id 响应。"""
    payload = _json_object_body(body)
    project_id = str(payload.get("project_id") or "").strip()
    raw_requests = payload.get("diagram_requests") or payload.get("requests") or []
    if not isinstance(raw_requests, list):
        raise PlatformError(code="INVALID_REQUEST", message="diagram_requests 必须是数组", status_code=400)
    diagram_requests = [item for item in raw_requests if isinstance(item, dict)]
    if not diagram_requests:
        raise PlatformError(code="INVALID_REQUEST", message="diagram_requests 不能为空", status_code=400)

    task_manager = _task_manager()
    if not _diagram_generation_enabled():
        task_id = task_manager.create_task("diagram", project_id, workflow_name=_get_diagram_workflow_name())
        sections = []
        for item in diagram_requests:
            item_project_id = str(item.get("project_id") or project_id or "").strip()
            base_content = str(item.get("base_content") or "")
            sections.append(_build_diagram_task_result({**item, "project_id": item_project_id}, base_content, [], None))
        task_manager.set_result(
            task_id,
            {
                "done": True,
                "project_id": project_id,
                "sections": sections,
                "failed_sections": [],
                "diagrams_count": 0,
            },
        )
        _sync_project_runtime_from_task(task_manager.get_task(task_id))
        return {"task_id": task_id, "count": len(sections)}

    diagram_key = _get_workflow_key(_get_diagram_workflow_name())
    if not diagram_key:
        raise PlatformError(code="TASK_START_FAILED", message=f"{_get_diagram_workflow_name()} 工作流 API Key 未配置", status_code=500)

    await _ensure_project_slot_native(project_id, "diagram")
    task_id = task_manager.create_task("diagram", project_id, workflow_name=_get_diagram_workflow_name())
    _persist_project_runtime(
        project_id,
        task_id=task_id,
        task_type="diagram",
        runtime_state="running",
        message="批量图表生成中",
        cancellable=True,
    )

    async def run_task() -> None:
        sections: list[dict[str, Any]] = []
        failed_sections: list[dict[str, Any]] = []
        try:
            total = len(diagram_requests)
            for idx, item in enumerate(diagram_requests, start=1):
                _ensure_task_running(task_id)
                section_id = str(item.get("section_id") or "")
                section_title = str(item.get("section_title") or section_id or "未命名章节")
                task_manager.update_stage(task_id, f"🎨 图表生成中 {idx}/{total}: {section_title}")
                merged_request = {
                    **item,
                    "project_id": str(item.get("project_id") or project_id or "").strip(),
                    "enable_diagrams": bool(item.get("enable_diagrams", payload.get("enable_diagrams", True))),
                    "max_diagrams": int(item.get("max_diagrams", payload.get("max_diagrams", 0)) or 0),
                    "mapping_table": item.get("mapping_table", payload.get("mapping_table", {}) or {}),
                }
                result_payload = await _run_diagram_request(task_id, merged_request, diagram_key)
                sections.append(result_payload)
                if result_payload.get("diagram_error"):
                    failed_sections.append({"section_id": section_id, "error": result_payload["diagram_error"]})
                task_manager.append_partial_event(
                    task_id,
                    {
                        "partial": True,
                        "phase": "diagram_section_done",
                        "section_id": section_id,
                        "content": result_payload.get("content", ""),
                        "word_count": result_payload.get("word_count", 0),
                        "quality_score": result_payload.get("quality_score"),
                        "feedback": result_payload.get("feedback"),
                        "replace_report": result_payload.get("replace_report", []),
                        "diagrams_count": result_payload.get("diagrams_count", 0),
                        "diagram_error": result_payload.get("diagram_error"),
                        "done_count": idx,
                        "total_count": total,
                    },
                )
            task_manager.set_result(
                task_id,
                {
                    "done": True,
                    "project_id": project_id,
                    "sections": sections,
                    "failed_sections": failed_sections,
                    "diagrams_count": sum(int(row.get("diagrams_count") or 0) for row in sections),
                },
            )
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except asyncio.CancelledError:
            if project_id:
                await task_manager.release_diagram_slot(project_id)
            logger.info("[Task %s] 批量图表任务被用户取消", task_id)
            task_manager.set_cancelled(task_id)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except Exception as exc:
            if project_id:
                await task_manager.release_diagram_slot(project_id)
            logger.error("[Task %s] 批量图表后台任务失败: %s", task_id, exc, exc_info=True)
            task_manager.set_result(
                task_id,
                {
                    "done": True,
                    "project_id": project_id,
                    "sections": sections,
                    "failed_sections": failed_sections,
                    "diagrams_count": sum(int(row.get("diagrams_count") or 0) for row in sections),
                    "diagram_error": _build_diagram_error_payload(exc, "批量图表"),
                },
            )
            _sync_project_runtime_from_task(task_manager.get_task(task_id))

    background_task = asyncio.create_task(run_task())
    task_manager.set_async_task(task_id, background_task)
    return {"task_id": task_id, "count": len(diagram_requests)}


async def start_analyze_task_payload(
    *,
    raw_document: str = "",
    project_id: str = "",
    selected_node_ids: str = "",
) -> dict[str, Any]:
    """启动解析报告后台任务；入参为原文/项目/节点选择，出参为 legacy task_id 响应。"""
    normalized_project_id = _ensure_safe_project_id(project_id)
    task_manager = _task_manager()
    dify_key = _get_workflow_key("doc_analysis") or _get_workflow_key("requirement_extractor")
    if not dify_key:
        raise PlatformError(code="TASK_START_FAILED", message="需求提取工作流 API Key 未配置", status_code=500)

    config_path = _bid_generator_root() / "config" / "analysis_framework.json"
    if not config_path.exists():
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="analysis_framework.json 不存在", status_code=404)
    system_prompt_base, all_nodes = load_docanalysis_framework(config_path)
    if not all_nodes:
        raise PlatformError(code="INVALID_REQUEST", message="框架中无节点", status_code=400)

    normalized_selected_ids = {
        node_id.strip() for node_id in str(selected_node_ids or "").split(",") if node_id.strip()
    } or None
    document_source = str(raw_document or "").strip()
    if document_source:
        _persist_raw_document(normalized_project_id, document_source[:300000])
    else:
        document_source = _load_raw_document(normalized_project_id)
    if not document_source:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="未找到项目原文缓存，请先上传并解析文档", status_code=404)
    document_text = document_source[:300000]

    groups = build_docanalysis_groups(all_nodes, normalized_selected_ids)
    if not groups or not any(group.get("nodes") for group in groups):
        raise PlatformError(code="INVALID_REQUEST", message="未找到可提取节点，请检查解析框架配置", status_code=400)

    total_nodes = sum(len(group.get("nodes") or []) for group in groups)
    task_id = task_manager.create_task("analyze", normalized_project_id)
    _persist_project_runtime(
        normalized_project_id,
        task_id=task_id,
        task_type="analyze",
        runtime_state="running",
        message="解析报告生成中",
    )

    async def run_task() -> None:
        try:
            existing_project = get_project_payload(normalized_project_id)
            existing_data = existing_project.get("data") if isinstance(existing_project, dict) else {}
            if not isinstance(existing_data, dict):
                existing_data = {}
            existing_report = existing_data.get("analysisReport") or existing_data.get("analysis_report") or []
            results_by_id = _collect_analysis_content_map(existing_report) if isinstance(existing_report, list) else {}
            latest_bid_items = [item for item in (existing_data.get("bidAttachmentList") or []) if isinstance(item, dict)]

            task_manager.update_stage(task_id, f"开始分析，共 {total_nodes} 个节点")
            _push_task_event(task_id, "structure_stage", {"phase": "attachments_generating", "label": "附件结构生成中"})

            async def call_group_workflow(subset_nodes: list[dict], subset_label: str) -> str:
                combined_system = build_docanalysis_system_prompt(system_prompt_base, subset_nodes, subset_label)
                outputs: dict[str, Any] = {}
                got_finished = False
                async for chunk in _call_dify_workflow_stream(
                    dify_key,
                    {
                        "system_prompt": combined_system,
                        "raw_document": document_text,
                        "node_label": subset_label,
                    },
                ):
                    _ensure_task_running(task_id)
                    if not isinstance(chunk, dict):
                        continue
                    if chunk.get("dify_task_id"):
                        task_manager.set_dify_task_id(task_id, chunk["dify_task_id"])
                    if chunk.get("__finished__"):
                        got_finished = True
                        outputs = chunk.get("outputs", {}) or {}
                        break
                if not got_finished:
                    raise RuntimeError("解析工作流异常中断（未收到 finished 事件）")
                return extract_docanalysis_text_output(outputs)

            async def extract_group(group: dict[str, Any]) -> list[dict[str, str]]:
                subset_nodes = group.get("nodes") if isinstance(group.get("nodes"), list) else []
                group_label = str(group.get("group_label") or "")
                try:
                    raw_text = await call_group_workflow(subset_nodes, group_label)
                    bid_items: list[dict] = []
                    raw_text, attachments_payload = split_bid_attachments_tag(raw_text)
                    if attachments_payload:
                        bid_items = parse_bid_attachments_payload(attachments_payload)
                        if bid_items:
                            bid_items = _enrich_bid_attachments_with_doc_blocks(normalized_project_id, bid_items)
                            latest_bid_items.clear()
                            latest_bid_items.extend(bid_items)
                            task_manager.update_stage(task_id, f"{_BID_ATTACH_STAGE_PREFIX}{json.dumps(bid_items, ensure_ascii=False)}")

                    result_map = parse_docanalysis_result_map(raw_text)
                    if not bid_items and any(str(node.get("id") or "") == "structure_attachments" for node in subset_nodes):
                        fallback_names = _extract_chapter_names_from_text(str(result_map.get("structure_attachments", "") or ""))
                        fallback_items = [{"name": name, "start_locator": "", "end_locator": "", "description": ""} for name in fallback_names]
                        if fallback_items:
                            fallback_items = _enrich_bid_attachments_with_doc_blocks(normalized_project_id, fallback_items)
                            latest_bid_items.clear()
                            latest_bid_items.extend(fallback_items)
                            task_manager.update_stage(task_id, f"{_BID_ATTACH_STAGE_PREFIX}{json.dumps(fallback_items, ensure_ascii=False)}")

                    results: list[dict[str, str]] = []
                    for node in subset_nodes:
                        content = extract_docanalysis_node_content(result_map, str(node.get("id") or ""))
                        if isinstance(content, (dict, list)):
                            content = json.dumps(content, ensure_ascii=False, indent=2)
                        content_text = str(content)
                        node_id = str(node.get("id") or "")
                        results_by_id[node_id] = content_text
                        results.append({"node_id": node_id, "label": str(node.get("label") or ""), "content": content_text})
                    return results
                except Exception as exc:
                    logger.warning("[analyze task] 分组 [%s] 完整提取失败，降级逐节点重试: %s", group_label, exc)

                fallback_results: list[dict[str, str]] = []
                for node in subset_nodes:
                    node_id = str(node.get("id") or "")
                    node_label = str(node.get("label") or "")
                    task_manager.update_stage(task_id, f"正在进行节点重试: {node_label}...")
                    try:
                        single_raw = await call_group_workflow([node], f"{group_label} - 单独抽取 {node_label}")
                        bid_items: list[dict] = []
                        single_raw, attachments_payload = split_bid_attachments_tag(single_raw)
                        if attachments_payload:
                            bid_items = parse_bid_attachments_payload(attachments_payload)
                            if bid_items:
                                bid_items = _enrich_bid_attachments_with_doc_blocks(normalized_project_id, bid_items)
                                latest_bid_items.clear()
                                latest_bid_items.extend(bid_items)
                                task_manager.update_stage(task_id, f"{_BID_ATTACH_STAGE_PREFIX}{json.dumps(bid_items, ensure_ascii=False)}")
                        try:
                            single_result_map = parse_docanalysis_result_map(single_raw)
                            content = extract_docanalysis_node_content(single_result_map, node_id)
                        except Exception:
                            content = single_raw or "**提取异常**"
                            if node_id == "structure_attachments" and not bid_items:
                                fallback_names = _extract_chapter_names_from_text(str(single_raw or ""))
                                fallback_items = [{"name": name, "start_locator": "", "end_locator": "", "description": ""} for name in fallback_names]
                                if fallback_items:
                                    fallback_items = _enrich_bid_attachments_with_doc_blocks(normalized_project_id, fallback_items)
                                    latest_bid_items.clear()
                                    latest_bid_items.extend(fallback_items)
                                    task_manager.update_stage(task_id, f"{_BID_ATTACH_STAGE_PREFIX}{json.dumps(fallback_items, ensure_ascii=False)}")
                        if isinstance(content, (dict, list)):
                            content = json.dumps(content, ensure_ascii=False, indent=2)
                        content_text = str(content)
                        results_by_id[node_id] = content_text
                        fallback_results.append({"node_id": node_id, "label": node_label, "content": content_text})
                        await asyncio.sleep(0.5)
                    except Exception as single_exc:
                        logger.error("[analyze task] 节点 %s 降级提取失败: %s", node_id, single_exc)
                        fallback_results.append({"node_id": node_id, "label": node_label, "content": "**提取失败，请重新生成**"})
                return fallback_results

            done_groups = 0
            success_count = 0
            for group in groups:
                results = await extract_group(group)
                done_groups += 1
                task_manager.update_stage(task_id, f"完成: {group.get('group_label', '')} ({done_groups}/{len(groups)})")
                for result in results:
                    task_manager.update_stage(task_id, f"__node__{json.dumps(result, ensure_ascii=False)}")
                    success_count += 1

            _push_task_event(task_id, "structure_stage", {"phase": "business_generating", "label": "商务部分生成中"})
            _push_task_event(task_id, "structure_stage", {"phase": "technical_generating", "label": "技术部分生成中"})
            analysis_v2 = _build_analysis_v2(results_by_id, latest_bid_items)
            analysis_report = _inflate_analysis_tree(all_nodes, results_by_id)
            analysis_report = _inject_analysis_report_derived_nodes(analysis_report, analysis_v2)
            _persist_analysis_state(normalized_project_id, analysis_report, analysis_v2)
            if bool(analysis_v2.get("enable_response_branch")):
                task_manager.update_stage(task_id, "response_branch_enabled")
            else:
                task_manager.update_stage(task_id, "response_branch_skipped")
            task_manager.update_stage(task_id, f"{_ANALYSIS_V2_STAGE_PREFIX}{json.dumps(analysis_v2, ensure_ascii=False)}")
            _push_task_event(task_id, "structure_stage", {"phase": "structure_ready", "label": "商务与技术结构已生成"})
            task_manager.set_result(
                task_id,
                {
                    "total_nodes": total_nodes,
                    "success_count": success_count,
                    "done": True,
                    "analysis_report": analysis_report,
                    "analysis_v2": analysis_v2,
                },
            )
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except asyncio.CancelledError:
            await _best_effort_stop_dify_by_task_id(task_id)
            task_manager.set_cancelled(task_id)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except Exception as exc:
            logger.error("[analyze task] 后台任务失败: %s", exc, exc_info=True)
            task_manager.set_error(task_id, str(exc))
            _sync_project_runtime_from_task(task_manager.get_task(task_id))

    background_task = asyncio.create_task(run_task())
    task_manager.set_async_task(task_id, background_task)
    return {"task_id": task_id}


async def start_extract_task_payload(
    file: UploadFile,
    *,
    project_name: str = "",
    project_id: str = "",
    enable_desensitize: bool = True,
    desensitize_profile: str = "tender",
    use_vision_parsing: bool = False,
) -> dict[str, Any]:
    """启动文档解析后台任务；入参为上传文件和解析配置，出参保持 legacy task_id 响应。"""
    normalized_project_id = _required_string(project_id, field="project_id")
    await _ensure_project_slot_native(normalized_project_id, "extract")
    content_bytes = await file.read()
    filename = str(getattr(file, "filename", "") or "")
    suffix = Path(filename).suffix.lower()
    cache_id = normalized_project_id or uuid.uuid4().hex[:12]

    task_manager = _task_manager()
    task_id = task_manager.create_task("extract", normalized_project_id)
    _persist_project_runtime(
        normalized_project_id,
        task_id=task_id,
        task_type="extract",
        runtime_state="running",
        message="文档抽取中",
        cancellable=True,
    )

    async def run_task() -> None:
        try:
            task_manager.update_stage(task_id, "解析文档结构")
            pdf_url = ""
            if suffix == ".pdf":
                pdf_url = _cache_pdf_file_native(cache_id, content_bytes)
                _extract_pdf_pages_text_native(content_bytes)
            elif suffix in {".docx", ".doc"}:
                try:
                    pdf_url = _convert_to_pdf_and_cache_native(cache_id, content_bytes, filename)
                except Exception as exc:
                    logger.warning("DOC/DOCX 转 PDF 失败: %s", exc)

            raw_document, raw_image_map = _extract_raw_text_with_images_native(
                filename,
                content_bytes,
                use_vision_parsing=use_vision_parsing,
            )
            if str(raw_document).startswith("["):
                raise PlatformError(
                    code="TASK_START_FAILED",
                    message="旧版 .doc 文件无法自动解析，请将文件另存为 .docx 后重新上传。",
                    status_code=400,
                )

            text_for_dify = str(raw_document or "")
            if suffix in {".docx", ".doc"}:
                try:
                    loc_text, _loc_map, doc_blocks = _extract_docx_with_locators_native(content_bytes)
                    if doc_blocks:
                        _persist_project_doc_blocks_snapshot(project_id=cache_id, doc_blocks=doc_blocks)
                    if suffix == ".docx":
                        _persist_docx_cache(cache_id, content_bytes)
                    if loc_text:
                        text_for_dify = loc_text
                except Exception as exc:
                    logger.warning("[extract task] 定位符缓存写入失败: %s", exc)

            task_manager.update_stage(task_id, "文档结构解析完成")

            mapping_table: dict[str, str] = {}
            entity_count = 0
            placeholder_manifest: dict[str, Any] = {}
            placeholder_policy: dict[str, Any] = {}
            if enable_desensitize:
                task_manager.update_stage(task_id, "隐私脱敏处理中")
                try:
                    desensitized = _run_bid_pipt_preprocess(
                        text=text_for_dify[:300000],
                        project_id=normalized_project_id,
                        task_id=task_id,
                        profile_name=desensitize_profile,
                    )
                    text_for_dify = str(desensitized.get("text") or text_for_dify[:300000])
                    mapping_table = _string_mapping(desensitized.get("mapping_table"))
                    entity_count = _int_or_default(
                        desensitized.get("mapping_table_count"),
                        default=len(mapping_table),
                    )
                    placeholder_manifest = _dict_or_default(desensitized.get("placeholder_manifest"))
                    placeholder_policy = _dict_or_default(desensitized.get("placeholder_policy"))
                    task_manager.update_stage(task_id, f"脱敏完成，识别 {entity_count} 处实体")
                except Exception as exc:
                    logger.warning("脱敏失败: %s", exc)
                    text_for_dify = text_for_dify[:300000]
                    task_manager.update_stage(task_id, "脱敏跳过（使用原文）")
            else:
                task_manager.update_stage(task_id, "跳过脱敏")

            _persist_raw_document(cache_id, text_for_dify[:300000])
            task_manager.update_stage(task_id, "预处理完成")
            task_manager.set_result(
                task_id,
                {
                    "done": True,
                    "bid_type": "tech",
                    "project_summary": "",
                    "requirements": [],
                    "analysis_report": [],
                    "analysis_v2": {},
                    "mapping_table": mapping_table,
                    "placeholder_manifest": placeholder_manifest,
                    "placeholder_policy": placeholder_policy,
                    "entity_count": entity_count,
                    "image_map": raw_image_map if isinstance(raw_image_map, dict) else {},
                    "required_attachments": [],
                    "scoring_table_template": [],
                    "raw_document": text_for_dify,
                    "pdf_url": pdf_url,
                },
            )
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except asyncio.CancelledError:
            task_manager.set_cancelled(task_id)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except PlatformError as exc:
            task_manager.set_error(task_id, exc.message)
            _sync_project_runtime_from_task(task_manager.get_task(task_id))
        except Exception as exc:
            logger.error("[extract task] 后台任务失败: %s", exc, exc_info=True)
            task_manager.set_error(task_id, str(exc))
            _sync_project_runtime_from_task(task_manager.get_task(task_id))

    background_task = asyncio.create_task(run_task())
    task_manager.set_async_task(task_id, background_task)
    return {"task_id": task_id}


async def export_scoring_table_response(body: Mapping[str, Any]) -> Any:
    """导出评分表 Excel；入参为评分表 JSON，出参保持 legacy 二进制响应。"""
    payload = _json_object_body(body)
    project_name = _required_string(payload.get("project_name"), field="project_name")
    rows = payload.get("rows", [])
    if not isinstance(rows, list):
        raise PlatformError(code="INVALID_REQUEST", message="rows 必须是数组。", status_code=400)

    try:
        workbook_bytes = _build_scoring_table_xlsx(project_name=project_name, rows=rows)
    except PlatformError:
        raise
    except Exception as exc:
        raise PlatformError(code="EXPORT_FAILED", message=f"导出评分表失败: {exc}", status_code=500) from exc

    filename = f"{project_name}_自评评分表.xlsx"
    content_disposition = f'attachment; filename="scoring.xlsx"; filename*=UTF-8\'\'{quote(filename, safe="")}'
    return BidGeneratorFilePayload(
        content=workbook_bytes,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        filename="scoring.xlsx",
        inline=False,
        cache_control="no-store",
        headers={"Content-Disposition": content_disposition},
    )


async def forge_document_response(body: Mapping[str, Any]) -> Any:
    """组装导出标书 DOCX；入参为 forge JSON，出参保持 legacy 二进制响应。"""
    payload = _json_object_body(body)
    project_name = str(payload.get("project_name") or "投标文件").strip() or "投标文件"
    sections = payload.get("sections", [])
    scoring_rows = payload.get("scoring_rows", [])
    attachments = payload.get("attachments", [])
    mapping_table = payload.get("mapping_table", {})
    bidder_info = payload.get("bidder_info", {})
    image_map = payload.get("image_map", {})
    project_id = str(payload.get("project_id") or "").strip()
    if not isinstance(sections, list):
        raise PlatformError(code="INVALID_REQUEST", message="sections 必须是数组。", status_code=400)
    if not isinstance(scoring_rows, list):
        raise PlatformError(code="INVALID_REQUEST", message="scoring_rows 必须是数组。", status_code=400)
    if not isinstance(attachments, list):
        raise PlatformError(code="INVALID_REQUEST", message="attachments 必须是数组。", status_code=400)
    if not isinstance(mapping_table, Mapping):
        mapping_table = {}
    if not isinstance(bidder_info, Mapping):
        bidder_info = {}
    if not isinstance(image_map, Mapping):
        image_map = {}

    try:
        docx_bytes = _build_forge_document_docx(
            project_id=project_id,
            sections=sections,
            scoring_rows=scoring_rows,
            attachments=attachments,
            mapping_table=dict(mapping_table),
            bidder_info=dict(bidder_info),
            image_map=dict(image_map),
        )
    except PlatformError:
        raise
    except Exception as exc:
        logger.error("forge-document 失败: %s", exc, exc_info=True)
        raise PlatformError(code="FORGE_FAILED", message=str(exc), status_code=500) from exc

    safe_name = project_name.replace("/", "_").replace("\\", "_")
    filename = f"{safe_name}_标书文件.docx"
    content_disposition = f'attachment; filename="document.docx"; filename*=UTF-8\'\'{quote(filename, safe="")}'
    return BidGeneratorFilePayload(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="document.docx",
        inline=False,
        cache_control="no-store",
        headers={"Content-Disposition": content_disposition},
    )


async def export_report_response(body: Mapping[str, Any]) -> Any:
    """导出解析报告 PDF；入参为报告节点 JSON，出参保持 legacy 二进制响应。"""
    payload = _json_object_body(body)
    project_name = str(payload.get("project_name") or "招标文件").strip() or "招标文件"
    nodes = payload.get("nodes", [])
    if not isinstance(nodes, list):
        raise PlatformError(code="INVALID_REQUEST", message="nodes 必须是数组。", status_code=400)

    try:
        pdf_bytes = _build_analysis_report_pdf(project_name=project_name, nodes=nodes)
    except PlatformError:
        raise
    except Exception as exc:
        raise PlatformError(code="EXPORT_FAILED", message=f"导出解析报告失败: {exc}", status_code=500) from exc

    filename = f"解析报告_{project_name}_{datetime.now().strftime('%Y%m%d')}.pdf"
    content_disposition = f'attachment; filename="analysis-report.pdf"; filename*=UTF-8\'\'{quote(filename, safe="")}'
    return BidGeneratorFilePayload(
        content=pdf_bytes,
        media_type="application/pdf",
        filename="analysis-report.pdf",
        inline=False,
        cache_control="no-store",
        headers={"Content-Disposition": content_disposition},
    )


async def stream_task_progress_response(
    task_id: str,
    request: Request,
    *,
    project_id: str | None = None,
) -> StreamingResponse:
    """推送后台任务进度 SSE；入参为任务 ID/请求/项目 ID，出参保持 legacy 事件流协议。"""
    task_id_value = _required_string(task_id, field="task_id")
    normalized_project_id = str(project_id or "").strip() or None

    async def progress_stream() -> Any:
        try:
            task = _require_task_owner(task_id_value, normalized_project_id)
        except PlatformError as exc:
            yield _sse_data({"error": exc.message})
            return

        started_at = _task_started_datetime(task)
        sent = 0
        for stage in getattr(task, "stages", None) or []:
            if request is not None and await request.is_disconnected():
                logger.info("[BidTask %s] SSE client disconnected during replay", task_id_value)
                return
            for event in _task_stage_sse_events(task_id_value, stage, sent, started_at):
                yield event
            sent += 1

        terminal_events = _task_terminal_sse_events(task)
        if terminal_events:
            for event in terminal_events:
                yield event
            return

        while True:
            if request is not None and await request.is_disconnected():
                logger.info("[BidTask %s] SSE client disconnected during stream", task_id_value)
                return
            try:
                await asyncio.wait_for(task._event.wait(), timeout=30)
            except asyncio.TimeoutError:
                heartbeat_stage = "⏳ 仍在生成大纲，请稍候…"
                phase, percent = _outline_stage_meta_from_label(getattr(task, "current_stage", "") or heartbeat_stage)
                yield _sse_data(
                    {
                        "stage": heartbeat_stage,
                        "heartbeat": True,
                        "phase": phase,
                        "percent": percent or 50,
                        "elapsed_sec": _elapsed_seconds(started_at),
                    }
                )
                continue

            try:
                task = _require_task_owner(task_id_value, normalized_project_id)
            except PlatformError:
                return

            for index, stage in enumerate((getattr(task, "stages", None) or [])[sent:], start=sent):
                if request is not None and await request.is_disconnected():
                    logger.info("[BidTask %s] SSE client disconnected during events", task_id_value)
                    return
                for event in _task_stage_sse_events(task_id_value, stage, index, started_at):
                    yield event
            sent = len(getattr(task, "stages", None) or [])

            terminal_events = _task_terminal_sse_events(task)
            if terminal_events:
                for event in terminal_events:
                    yield event
                return

    return StreamingResponse(
        progress_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


def get_extracted_image_by_hash_payload(image_hash: str) -> BidGeneratorFilePayload:
    """按图片 hash 读取提取图片；入参为十六进制 hash，出参为图片字节及响应元数据。"""
    normalized_hash = _ensure_safe_image_hash(image_hash)
    try:
        with get_engine().begin() as conn:
            exists = conn.execute(text("SELECT to_regclass('bid_generator.image_registry') IS NOT NULL")).scalar_one()
            if not exists:
                raise PlatformError(code="RESOURCE_NOT_FOUND", message="图片实体不存在", status_code=404)
            row = conn.execute(
                text(
                    """
                    SELECT abs_path
                    FROM bid_generator.image_registry
                    WHERE image_hash = :image_hash
                    """
                ),
                {"image_hash": normalized_hash},
            ).mappings().first()
    except PlatformError:
        raise
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc
    if row is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="图片实体不存在", status_code=404)
    return _read_image_file_payload(Path(str(row["abs_path"])), fallback_name=f"{normalized_hash}.png")


def get_extracted_image_payload(filename: str) -> BidGeneratorFilePayload:
    """按文件名读取提取图片；入参为安全文件名，出参为图片字节及响应元数据。"""
    normalized_filename = _ensure_safe_image_filename(filename)
    return _read_image_file_payload(_extracted_image_path(normalized_filename), fallback_name=normalized_filename)


def list_knowledge_images_payload(
    *,
    source_doc: str = "",
    caption_status: str = "",
    limit: int = 200,
) -> dict[str, Any]:
    """查询知识库图片语义资产；入参为过滤条件，出参兼容 legacy items/total。"""
    safe_limit = max(1, min(int(limit or 200), 500))
    filters = []
    params: dict[str, Any] = {"limit": safe_limit}
    normalized_source_doc = str(source_doc or "").strip()
    normalized_caption_status = str(caption_status or "").strip()
    if normalized_source_doc:
        filters.append("source_doc = :source_doc")
        params["source_doc"] = normalized_source_doc
    if normalized_caption_status:
        filters.append("caption_status = :caption_status")
        params["caption_status"] = normalized_caption_status
    where_sql = f"WHERE {' AND '.join(filters)}" if filters else ""
    try:
        with get_engine().begin() as conn:
            _ensure_knowledge_image_tables(conn)
            rows = conn.execute(
                text(
                    f"""
                    SELECT image_hash, placeholder, source_doc, source_page, caption,
                           image_type, summary, tags_json, caption_status, created_at
                    FROM bid_generator.knowledge_image_assets
                    {where_sql}
                    ORDER BY created_at DESC
                    LIMIT :limit
                    """
                ),
                params,
            ).mappings().all()
            hashes = [str(row.get("image_hash") or "").lower() for row in rows]
            registries = _load_image_registry_preview_urls(conn, hashes)
    except PlatformError:
        raise
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc
    return {
        "items": [_knowledge_image_asset_payload(row, registries.get(str(row.get("image_hash") or "").lower(), "")) for row in rows],
        "total": len(rows),
    }


def update_knowledge_image_payload(image_hash: str, payload: Mapping[str, Any]) -> dict[str, Any]:
    """更新知识库图片语义资产；入参为 image_hash 与可编辑字段，出参为更新后的资产。"""
    normalized_hash = re.sub(r"[^a-fA-F0-9]", "", str(image_hash or "")).lower()
    if not normalized_hash:
        raise PlatformError(code="INVALID_REQUEST", message="无效的 image_hash", status_code=400)
    try:
        with get_engine().begin() as conn:
            _ensure_knowledge_image_tables(conn)
            row = conn.execute(
                text(
                    """
                    SELECT image_hash, placeholder, source_doc, source_page, caption,
                           image_type, summary, tags_json, caption_status, created_at
                    FROM bid_generator.knowledge_image_assets
                    WHERE image_hash = :image_hash
                    FOR UPDATE
                    """
                ),
                {"image_hash": normalized_hash},
            ).mappings().first()
            if row is None:
                raise PlatformError(code="RESOURCE_NOT_FOUND", message="知识库图片不存在", status_code=404)

            caption = str(row.get("caption") or "")
            image_type = str(row.get("image_type") or "")
            summary = str(row.get("summary") or "")
            tags_json = str(row.get("tags_json") or "[]")
            caption_status = str(row.get("caption_status") or "pending")

            if "caption" in payload and payload.get("caption") is not None:
                caption = str(payload.get("caption") or "").strip() or "知识库配图"
            if "image_type" in payload and payload.get("image_type") is not None:
                image_type = str(payload.get("image_type") or "").strip() or "其他"
            if "summary" in payload and payload.get("summary") is not None:
                summary = str(payload.get("summary") or "").strip()
            if "tags" in payload and payload.get("tags") is not None:
                tags_value = payload.get("tags")
                if not isinstance(tags_value, list):
                    raise PlatformError(code="INVALID_REQUEST", message="tags 必须是数组。", status_code=400)
                tags_json = json.dumps([str(item).strip() for item in tags_value if str(item).strip()], ensure_ascii=False)
            if "caption_status" in payload and payload.get("caption_status") is not None:
                caption_status = str(payload.get("caption_status") or "").strip() or "manual"
            else:
                caption_status = "manual"

            updated = conn.execute(
                text(
                    """
                    UPDATE bid_generator.knowledge_image_assets
                    SET caption = :caption,
                        image_type = :image_type,
                        summary = :summary,
                        tags_json = :tags_json,
                        caption_status = :caption_status
                    WHERE image_hash = :image_hash
                    RETURNING image_hash, placeholder, source_doc, source_page, caption,
                              image_type, summary, tags_json, caption_status, created_at
                    """
                ),
                {
                    "image_hash": normalized_hash,
                    "caption": caption,
                    "image_type": image_type,
                    "summary": summary,
                    "tags_json": tags_json,
                    "caption_status": caption_status,
                },
            ).mappings().one()
            registries = _load_image_registry_preview_urls(conn, [normalized_hash])
    except PlatformError:
        raise
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc
    return _knowledge_image_asset_payload(updated, registries.get(normalized_hash, ""))


def list_kb_sync_jobs_payload() -> dict[str, Any]:
    """列出最近知识库同步任务；出参兼容 legacy jobs 数组。"""
    status_dir = _kb_sync_status_dir()
    if not status_dir.exists():
        return {"jobs": []}
    jobs: list[dict[str, Any]] = []
    try:
        status_files = sorted(status_dir.glob("*.json"), key=lambda path: path.stat().st_mtime, reverse=True)[:20]
    except OSError:
        return {"jobs": []}
    for path in status_files:
        try:
            with path.open("r", encoding="utf-8") as file:
                data = json.load(file)
        except (OSError, json.JSONDecodeError):
            continue
        if not isinstance(data, dict):
            continue
        jobs.append(
            {
                "job_id": data.get("job_id"),
                "status": data.get("status"),
                "started_at": data.get("started_at"),
                "total": data.get("total", 0),
                "processed": data.get("processed", 0),
                "failed": data.get("failed", 0),
            }
        )
    return {"jobs": jobs}


def get_kb_sync_status_payload(job_id: str) -> dict[str, Any]:
    """查询知识库同步任务状态；入参为 job_id，出参兼容 legacy sync-status。"""
    normalized_job_id = _ensure_safe_kb_sync_job_id(job_id)
    task = _get_task(normalized_job_id)
    if task is not None:
        mapped_status = {
            "running": "running",
            "done": "completed",
            "error": "failed",
            "cancelled": "cancelled",
            "timeout": "failed",
        }.get(str(getattr(task, "status", "") or ""), "running")
        status_path = _kb_sync_status_path(normalized_job_id)
        if status_path.exists():
            try:
                data = _read_json_file(status_path)
                if isinstance(data, dict):
                    data.setdefault("job_id", normalized_job_id)
                    data.setdefault("task_id", normalized_job_id)
                    data["status"] = mapped_status
                    return data
            except PlatformError:
                pass
        return {
            "job_id": normalized_job_id,
            "task_id": normalized_job_id,
            "status": mapped_status,
            "started_at": _utc_iso_from_timestamp(getattr(task, "created_at", None)),
            "total": 0,
            "processed": 0,
            "failed": 1 if mapped_status == "failed" else 0,
            "current_file": "",
            "error": str(getattr(task, "error", "") or "") if mapped_status == "failed" else "",
        }

    status_path = _kb_sync_status_path(normalized_job_id)
    if not status_path.exists():
        raise PlatformError(code="RESOURCE_NOT_FOUND", message=f"未找到任务 {normalized_job_id} 的状态记录", status_code=404)
    data = _read_json_file(status_path)
    return data if isinstance(data, dict) else {}


async def get_knowledge_documents_payload() -> dict[str, Any]:
    """查询 Dify 知识库文档状态；出参兼容 legacy dataset_info/documents。"""
    dify_url = os.getenv("DIFY_API_URL", "http://localhost/v1")
    dataset_id = os.getenv("DIFY_DATASET_ID", "")
    dataset_key = os.getenv("DIFY_DATASET_KEY", "")
    if not dataset_id or not dataset_key:
        return {
            "dataset_info": {"error": "DIFY_DATASET_ID or KEY not configured in backend."},
            "documents": [],
        }

    headers = {"Authorization": f"Bearer {dataset_key}"}
    url = f"{dify_url}/datasets/{dataset_id}/documents"
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
            data = response.json()
    except Exception as exc:
        logger.error("Failed to fetch dify kb documents: %s", exc)
        return {"dataset_info": {"status": "error", "message": str(exc)}, "documents": []}

    document_rows = data.get("data") if isinstance(data, dict) else []
    if not isinstance(document_rows, list):
        document_rows = []
    documents = [_knowledge_document_payload(item) for item in document_rows if isinstance(item, dict)]
    return {"dataset_info": {"status": "connected", "dataset_id": dataset_id}, "documents": documents}


def _persist_project_analysis_report(*, project_id: str, report: list[Any]) -> None:
    try:
        with get_engine().begin() as conn:
            exists = conn.execute(text("SELECT to_regclass('bid_generator.projects') IS NOT NULL")).scalar_one()
            if not exists:
                return
            row = conn.execute(
                text(
                    """
                    SELECT data
                    FROM bid_generator.projects
                    WHERE id = :project_id
                    FOR UPDATE
                    """
                ),
                {"project_id": project_id},
            ).mappings().first()
            if row is None:
                return
            data = _json_value(row.get("data"))
            if not isinstance(data, dict):
                data = {}
            data["analysisReport"] = report
            conn.execute(
                text(
                    """
                    UPDATE bid_generator.projects
                    SET data = :data,
                        updated_at = :updated_at
                    WHERE id = :project_id
                    """
                ),
                {
                    "project_id": project_id,
                    "data": json.dumps(data, ensure_ascii=False),
                    "updated_at": _utc_now(),
                },
            )
    except (SQLAlchemyError, RuntimeError, PlatformError) as exc:
        logger.warning("[%s] 保存 analysisReport 到项目记录失败: %s", project_id, exc)


def _extract_chapter_names_from_text(text: str) -> list[str]:
    if not text:
        return []
    names = re.findall(r"<要点[^>]*>(.*?)</要点>", text, re.DOTALL)
    names = [name.strip() for name in names if name and name.strip()]
    if names:
        return names

    noise_prefix = re.compile(
        r"^(?:[\d一二三四五六七八九十百]+[.、．。）)]\s*|第[一二三四五六七八九十百\d]+[章节条]\s*|\([一二三四五六七八九十\d]+\)\s*)"
    )
    json_chars = set('{}[]"=:/')
    cleaned: list[str] = []
    for line in str(text or "").splitlines():
        line_value = noise_prefix.sub("", line.strip()).strip()
        if len(line_value) < 2:
            continue
        if any(char in json_chars for char in line_value):
            continue
        if line_value.startswith(("*", "#", "`", ">")):
            continue
        cleaned.append(line_value)
    return cleaned


def _normalize_chapter_name_for_match(name: str) -> str:
    normalized = str(name or "").strip().lower()
    if not normalized:
        return ""
    normalized = re.sub(r"<[^>]+>", "", normalized)
    normalized = re.sub(r"^[\s]*(?:\d+[.、)）]|\(?\d+\)?[.、]?|第[一二三四五六七八九十百\d]+[章节条])\s*", "", normalized)
    normalized = re.sub(r"[（(]\s*(?:商务|技术|资格|价格|响应|投标)\s*部分\s*[）)]", "", normalized)
    normalized = re.sub(r"""[《》“”"'`·\-\s:：，,。；;、/\\\\]""", "", normalized)
    return normalized


def _is_chapter_match(chapter_norm: str, block_norm: str) -> bool:
    if not chapter_norm or not block_norm:
        return False
    if chapter_norm in block_norm:
        return True
    if len(block_norm) >= 4 and block_norm in chapter_norm:
        return True
    return False


def _enrich_bid_attachments_with_doc_blocks(project_id: str, items: list[dict]) -> list[dict]:
    if not project_id or not isinstance(items, list) or not items:
        return items
    try:
        snapshot = get_project_doc_blocks_payload(project_id)
    except PlatformError:
        return items
    doc_blocks = snapshot.get("blocks") if isinstance(snapshot, dict) else []
    if not isinstance(doc_blocks, list) or not doc_blocks:
        return items

    normalized_blocks: list[dict[str, Any]] = []
    for block in doc_blocks:
        if not isinstance(block, dict):
            continue
        normalized_blocks.append(
            {
                "block_id": str(block.get("block_id") or ""),
                "locator": str(block.get("locator") or "").upper(),
                "body_idx": _non_negative_int(block.get("body_idx")),
                "text_norm": _normalize_chapter_name_for_match(str(block.get("text") or "")),
            }
        )
    if not normalized_blocks:
        return items

    matched_block_indices: list[tuple[int, int]] = []
    cursor = 0
    for item_index, item in enumerate(items):
        if not isinstance(item, dict):
            continue
        name_norm = _normalize_chapter_name_for_match(str(item.get("name") or ""))
        if not name_norm:
            continue
        hit_index = -1
        for block_index in range(cursor, len(normalized_blocks)):
            if _is_chapter_match(name_norm, str(normalized_blocks[block_index]["text_norm"])):
                hit_index = block_index
                break
        if hit_index < 0:
            for block_index in range(len(normalized_blocks)):
                if _is_chapter_match(name_norm, str(normalized_blocks[block_index]["text_norm"])):
                    hit_index = block_index
                    break
        if hit_index >= 0:
            matched_block_indices.append((item_index, hit_index))
            cursor = hit_index + 1

    if not matched_block_indices:
        return items

    enriched = [dict(item) if isinstance(item, dict) else item for item in items]
    for position, (item_index, block_index) in enumerate(matched_block_indices):
        start_block = normalized_blocks[block_index]
        next_block_index = matched_block_indices[position + 1][1] if position + 1 < len(matched_block_indices) else len(normalized_blocks)
        end_block = normalized_blocks[max(block_index, next_block_index - 1)]
        row = enriched[item_index]
        if not isinstance(row, dict):
            continue
        if not str(row.get("start_locator") or "").strip():
            row["start_locator"] = start_block.get("locator", "")
        if not str(row.get("end_locator") or "").strip():
            row["end_locator"] = end_block.get("locator", "")
        row["start_block_id"] = start_block.get("block_id", "")
        row["end_block_id"] = end_block.get("block_id", "")
    return enriched


def _collect_analysis_content_map(nodes: list[dict]) -> dict[str, str]:
    result: dict[str, str] = {}
    for node in nodes or []:
        if not isinstance(node, dict):
            continue
        node_id = str(node.get("id") or "").strip()
        children = node.get("children") or []
        if children:
            result.update(_collect_analysis_content_map(children))
            continue
        if node_id:
            result[node_id] = str(node.get("content") or "")
    return result


def _inflate_analysis_tree(nodes: list[dict], content_map: dict[str, str], parent_id: Optional[str] = None) -> list[dict]:
    tree: list[dict] = []
    for node in nodes or []:
        if not isinstance(node, dict):
            continue
        node_id = str(node.get("id") or "")
        children = _inflate_analysis_tree(node.get("children") or [], content_map, node_id.strip() or parent_id)
        tree.append(
            {
                "id": node_id,
                "label": str(node.get("label") or ""),
                "content": str(content_map.get(node_id, "")),
                "parent_id": parent_id,
                "children": children,
            }
        )
    return tree


def _parse_xml_items(text: str) -> list[str]:
    if not text:
        return []
    items = re.findall(r"<要点[^>]*>(.*?)</要点>", text, re.DOTALL)
    if items:
        return [re.sub(r"<[^>]+>", "", item).strip() for item in items if item and item.strip()]
    lines = [line.strip() for line in str(text or "").splitlines() if line.strip()]
    return [re.sub(r"^\[[^\]]+\]\s*", "", line).strip() for line in lines if line.strip()]


def _parse_xml_field_map(text: str) -> dict[str, str]:
    result: dict[str, str] = {}
    if not text:
        return result
    for key, value in re.findall(r"<([^>/]+)>(.*?)</\1>", str(text or ""), re.DOTALL):
        result[str(key).strip()] = re.sub(r"<[^>]+>", "", value).strip()
    return result


def _normalize_score_tag(value: str, name: str = "", criteria: str = "") -> str:
    normalized = str(value or "").strip().lower()
    if normalized in {"tech", "biz", "mixed"}:
        return normalized
    text_value = f"{name}\n{criteria}".lower()
    tech_keywords = ["技术", "方案", "架构", "实施", "服务响应", "功能", "性能", "团队", "驻场", "运维", "培训", "交付"]
    biz_keywords = ["资质", "商务", "报价", "价格", "业绩", "合同", "付款", "售后", "承诺", "证书", "企业"]
    tech_hit = sum(1 for keyword in tech_keywords if keyword in text_value)
    biz_hit = sum(1 for keyword in biz_keywords if keyword in text_value)
    if tech_hit and not biz_hit:
        return "tech"
    if biz_hit and not tech_hit:
        return "biz"
    return "mixed"


def _normalize_optional_bool(value: Any) -> Optional[bool]:
    if isinstance(value, bool):
        return value
    if value is None:
        return None
    normalized = str(value).strip().lower()
    if normalized in {"true", "1", "yes", "y", "是"}:
        return True
    if normalized in {"false", "0", "no", "n", "否", ""}:
        return False
    return None


def _parse_scoring_details(raw: str) -> dict[str, Any]:
    if not raw:
        return {"total": 0, "items": []}
    try:
        payload = json.loads(raw)
    except Exception:
        try:
            payload = ast.literal_eval(raw)
        except Exception:
            logger.warning("[analysis_v2] scoring_details 解析失败，使用空列表")
            return {"total": 0, "items": []}
    items = payload.get("items") if isinstance(payload, dict) else []
    normalized_items: list[dict[str, Any]] = []
    for index, item in enumerate(items or []):
        if not isinstance(item, dict):
            continue
        name = str(item.get("name") or "").strip()
        if not name:
            continue
        try:
            max_score = int(float(item.get("max_score", item.get("maxScore", 0)) or 0))
        except Exception:
            max_score = 0
        criteria = str(item.get("criteria") or "").strip()
        explicit_response = _normalize_optional_bool(item.get("is_response_item", item.get("isResponseItem")))
        normalized_items.append(
            {
                "id": str(item.get("id") or f"score_{index + 1}"),
                "name": name,
                "max_score": max_score,
                "criteria": criteria,
                "score_tag": _normalize_score_tag(str(item.get("score_tag") or ""), name=name, criteria=criteria),
                "is_response_item": explicit_response,
                "response_reason": str(item.get("response_reason", item.get("responseReason", "")) or "").strip(),
                "response_explicit": explicit_response is not None,
            }
        )
    total = payload.get("total", 0) if isinstance(payload, dict) else 0
    try:
        total = int(float(total or 0))
    except Exception:
        total = sum(item["max_score"] for item in normalized_items)
    if total <= 0:
        total = sum(item["max_score"] for item in normalized_items)
    return {"total": total, "items": normalized_items}


def _slugify_heading(text: str, fallback: str) -> str:
    slug = re.sub(r"[^\w\u4e00-\u9fa5]+", "_", str(text or "").strip().lower()).strip("_")
    return slug or fallback


def _make_structure_heading(
    *,
    title: str,
    level: int,
    category: str,
    source: str,
    source_node_id: str = "",
    source_title: str = "",
    score_tag: str = "",
    score_item_id: str = "",
    max_score: int = 0,
    criteria: str = "",
    criteria_excerpt: str = "",
    related_target_ids: Optional[list[str]] = None,
    priority_weight: float = 0.0,
    generation_strategy: str = "general",
    generation_mode: str = "derived",
    response_candidate: bool = False,
    generates_from_self: bool = False,
    start_block_id: str = "",
    end_block_id: str = "",
    start_locator: str = "",
    end_locator: str = "",
    anchor_confidence: float = 0.0,
) -> dict[str, Any]:
    safe_title = str(title or "").strip()
    return {
        "id": f"{category}_{_slugify_heading(safe_title, fallback=str(uuid.uuid4())[:8])}",
        "title": safe_title,
        "level": int(level),
        "category": category,
        "source": source,
        "source_node_id": source_node_id,
        "source_title": source_title or safe_title,
        "score_tag": score_tag,
        "score_item_id": score_item_id,
        "max_score": int(max_score or 0),
        "criteria": str(criteria or "").strip(),
        "criteria_excerpt": str(criteria_excerpt or "").strip(),
        "related_target_ids": list(related_target_ids or []),
        "priority_weight": float(priority_weight or 0.0),
        "generation_strategy": str(generation_strategy or "general"),
        "generation_mode": str(generation_mode or "derived"),
        "response_candidate": bool(response_candidate),
        "generates_from_self": bool(generates_from_self),
        "start_block_id": start_block_id,
        "end_block_id": end_block_id,
        "start_locator": start_locator,
        "end_locator": end_locator,
        "anchor_confidence": float(anchor_confidence),
        "editable_ops": ["rename", "delete"],
        "deleted": False,
        "children": [],
    }


def _criteria_excerpt(text: str, limit: int = 220) -> str:
    raw = str(text or "").strip()
    if len(raw) <= limit:
        return raw
    return raw[:limit].rstrip() + "..."


def _is_response_candidate_strict(name: str, criteria: str) -> bool:
    title = re.sub(r"\s+", "", str(name or "")).lower()
    crt = str(criteria or "").lower()
    strong_title_keys = [
        "响应情况",
        "响应程度",
        "符合性响应",
        "符合性偏离",
        "偏离情况",
        "偏离表",
        "逐条响应情况",
        "实质性条款响应情况",
        "技术条款响应情况",
    ]
    if any(key in title for key in strong_title_keys):
        return True
    if ("完全响应" in crt and "部分响应" in crt and ("不响应" in crt or "未响应" in crt or "偏离" in crt)) and ("得分" in crt or "得" in crt or "评分" in crt):
        return True
    return False


def _build_analysis_v2(content_map: dict[str, str], bid_items: list[dict]) -> dict[str, Any]:
    basic_info = _parse_xml_field_map(content_map.get("proj_basic", ""))
    scoring = _parse_scoring_details(content_map.get("scoring_details", ""))
    technical_target_nodes: list[dict[str, Any]] = []
    for node_id, label in [("resp_tech", "技术目标与范围"), ("resp_param", "参数与指标要求"), ("resp_substance", "实施与交付硬约束")]:
        content = str(content_map.get(node_id, "")).strip()
        if content:
            technical_target_nodes.append({"id": node_id, "label": label, "content": content})

    attachments: list[dict[str, Any]] = []
    for index, item in enumerate(bid_items or []):
        if not isinstance(item, dict):
            continue
        title = str(item.get("name") or "").strip()
        if not title:
            continue
        attachments.append(
            {
                **_make_structure_heading(
                    title=title,
                    level=1,
                    category="attachments",
                    source="llm",
                    source_node_id="structure_attachments",
                    source_title=title,
                    start_block_id=str(item.get("start_block_id") or ""),
                    end_block_id=str(item.get("end_block_id") or ""),
                    start_locator=str(item.get("start_locator") or ""),
                    end_locator=str(item.get("end_locator") or ""),
                    anchor_confidence=0.95 if item.get("start_block_id") or item.get("start_locator") else 0.35,
                ),
                "id": f"attachments_{index + 1}_{_slugify_heading(title, fallback=str(index + 1))}",
            }
        )

    technical_sections: list[dict[str, Any]] = []
    business_sections: list[dict[str, Any]] = []
    technical_target_ids = [str(node.get("id") or "").strip() for node in technical_target_nodes if str(node.get("id") or "").strip()]
    for index, item in enumerate(scoring.get("items", [])):
        score_tag = str(item.get("score_tag") or "mixed")
        if score_tag not in {"tech", "biz", "mixed"}:
            score_tag = "mixed"
        item_name = str(item.get("name") or f"评分项{index + 1}")
        item_criteria = str(item.get("criteria") or "")
        max_score = int(item.get("max_score") or 0)
        explicit_response = item.get("is_response_item")
        if score_tag == "biz":
            is_response = False
        elif item.get("response_explicit"):
            is_response = bool(explicit_response)
        else:
            is_response = _is_response_candidate_strict(item_name, item_criteria)
        heading = _make_structure_heading(
            title=item_name,
            level=2,
            category="technical" if score_tag != "biz" else "business",
            source="score_item",
            source_node_id="scoring_details",
            source_title=item_name,
            score_tag=score_tag,
            score_item_id=str(item.get("id") or ""),
            max_score=max_score,
            criteria=item_criteria,
            criteria_excerpt=_criteria_excerpt(item_criteria),
            related_target_ids=technical_target_ids if score_tag != "biz" else [],
            priority_weight=float(max_score),
            generation_strategy="response_special" if is_response else "general",
            generation_mode="derived",
            response_candidate=is_response,
            generates_from_self=is_response,
        )
        if score_tag == "biz":
            business_sections.append(heading)
        else:
            technical_sections.append(heading)

    objective_heading = _make_structure_heading(
        title="项目实施目标",
        level=2,
        category="technical",
        source="system",
        source_node_id="technical_targets",
        source_title="项目实施目标",
        score_tag="tech",
        related_target_ids=technical_target_ids,
        generation_strategy="objective_special",
        generation_mode="system",
        generates_from_self=False,
    )
    if not any(str(item.get("title") or "").strip() == "项目实施目标" for item in technical_sections):
        technical_sections.append(objective_heading)
    else:
        for section in technical_sections:
            if str(section.get("title") or "").strip() == "项目实施目标":
                section["generation_strategy"] = "objective_special"
                section["generation_mode"] = "system"

    response_candidates = [
        section
        for section in technical_sections
        if bool(section.get("response_candidate")) and str(section.get("title") or "").strip() != "项目实施目标"
    ]
    if len(response_candidates) > 1:
        preferred = [
            section
            for section in response_candidates
            if any(keyword in str(section.get("title") or "") for keyword in ["响应情况", "响应程度", "符合性偏离", "偏离情况"])
        ]
        keep = preferred[0] if preferred else response_candidates[0]
        for section in technical_sections:
            title = str(section.get("title") or "").strip()
            if title == str(keep.get("title") or "").strip():
                continue
            if bool(section.get("response_candidate")) and title != "项目实施目标":
                section["response_candidate"] = False
                section["generation_strategy"] = "general"
                section["generates_from_self"] = False

    response_sections = [section for section in technical_sections if bool(section.get("response_candidate")) and str(section.get("title") or "").strip() != "项目实施目标"]
    non_response_sections = [section for section in technical_sections if not bool(section.get("response_candidate")) and str(section.get("title") or "").strip() != "项目实施目标"]
    objective_sections = [section for section in technical_sections if str(section.get("title") or "").strip() == "项目实施目标"]
    technical_sections = non_response_sections + response_sections + objective_sections

    technical_h2_bindings = [
        {
            "h2_id": section.get("id", ""),
            "title": section.get("title", ""),
            "score_item_id": section.get("score_item_id", ""),
            "score_value": int(section.get("max_score") or 0),
            "score_criteria": section.get("criteria", ""),
            "score_tag": section.get("score_tag", ""),
            "related_target_ids": section.get("related_target_ids", []),
            "priority_weight": float(section.get("priority_weight") or 0.0),
            "generation_strategy": section.get("generation_strategy", "general"),
            "response_candidate": bool(section.get("response_candidate")),
            "generates_from_self": bool(section.get("generates_from_self")),
        }
        for section in technical_sections
        if not bool(section.get("deleted"))
    ]

    return {
        "schema_version": 3,
        "project_info": {
            "overview": str(content_map.get("proj_overview", "")).strip(),
            "basic_info": basic_info,
            "scoring_items": scoring.get("items", []),
            "scoring_total": int(scoring.get("total", 0) or 0),
        },
        "technical_targets": technical_target_nodes,
        "enable_response_branch": any(bool(section.get("response_candidate")) for section in technical_sections),
        "technical_h2_bindings": technical_h2_bindings,
        "bid_structure": {
            "attachments": attachments,
            "technical_sections": technical_sections,
            "business_sections": business_sections,
        },
    }


def _render_derived_structure_content(items: list[dict]) -> str:
    lines: list[str] = []
    for item in items or []:
        if not isinstance(item, dict):
            continue
        title = str(item.get("title") or "").strip()
        if title:
            lines.append(f"<要点>{title}</要点>")
    return "\n".join(lines)


def _inject_analysis_report_derived_nodes(report_nodes: list[dict], analysis_v2: dict) -> list[dict]:
    business_content = _render_derived_structure_content((analysis_v2.get("bid_structure") or {}).get("business_sections") or [])
    technical_content = _render_derived_structure_content((analysis_v2.get("bid_structure") or {}).get("technical_sections") or [])

    def walk(nodes: list[dict]) -> list[dict]:
        updated: list[dict] = []
        for node in nodes or []:
            if not isinstance(node, dict):
                continue
            node_copy = dict(node)
            node_id = str(node_copy.get("id") or "")
            if node_id == "structure_business":
                node_copy["content"] = business_content
            elif node_id == "structure_technical":
                node_copy["content"] = technical_content
            children = node_copy.get("children") or []
            if children:
                node_copy["children"] = walk(children)
            updated.append(node_copy)
        return updated

    return walk(report_nodes)


def _persist_analysis_state(project_id: str, analysis_report: list[dict], analysis_v2: dict) -> None:
    _persist_project_analysis_report(project_id=project_id, report=analysis_report)
    try:
        patch_project_payload(
            project_id,
            {
                "data_patch": {
                    "analysisV2": analysis_v2,
                    "bidAttachmentList": [
                        {
                            "name": item.get("title", ""),
                            "start_locator": item.get("start_locator", ""),
                            "end_locator": item.get("end_locator", ""),
                            "start_block_id": item.get("start_block_id", ""),
                            "end_block_id": item.get("end_block_id", ""),
                            "description": "",
                        }
                        for item in (analysis_v2.get("bid_structure", {}) or {}).get("attachments", [])
                        if not item.get("deleted")
                    ],
                }
            },
        )
    except Exception as exc:
        logger.warning("[%s] 持久化 analysis_v2 失败: %s", project_id, exc)


def _persist_project_doc_blocks_snapshot(*, project_id: str, doc_blocks: list[dict[str, Any]]) -> None:
    try:
        with get_engine().begin() as conn:
            exists = conn.execute(text("SELECT to_regclass('bid_generator.projects') IS NOT NULL")).scalar_one()
            if not exists:
                return
            row = conn.execute(
                text(
                    """
                    SELECT data
                    FROM bid_generator.projects
                    WHERE id = :project_id
                    FOR UPDATE
                    """
                ),
                {"project_id": project_id},
            ).mappings().first()
            if row is None:
                return
            data = _json_value(row.get("data"))
            if not isinstance(data, dict):
                data = {}
            data["__doc_blocks_cache"] = doc_blocks
            conn.execute(
                text(
                    """
                    UPDATE bid_generator.projects
                    SET data = :data,
                        updated_at = :updated_at
                    WHERE id = :project_id
                    """
                ),
                {
                    "project_id": project_id,
                    "data": json.dumps(data, ensure_ascii=False),
                    "updated_at": _utc_now(),
                },
            )
    except (SQLAlchemyError, RuntimeError, PlatformError) as exc:
        logger.warning("[%s] 保存 doc_blocks 快照到项目记录失败: %s", project_id, exc)


def _persist_docx_cache(project_id: str, content: bytes) -> None:
    docx_path = _docx_cache_path(project_id)
    try:
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        docx_path.write_bytes(content)
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="缓存 DOCX 文件失败。", status_code=500) from exc


def _extract_docx_blocks(content: bytes) -> list[dict[str, Any]]:
    try:
        import docx as docx_module
    except ImportError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="统一后端缺少 python-docx 依赖。", status_code=500) from exc

    doc = docx_module.Document(io.BytesIO(content))
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    paragraph_tag = f"{{{word_namespace}}}p"
    table_tag = f"{{{word_namespace}}}tbl"
    row_tag = f"{{{word_namespace}}}tr"
    cell_tag = f"{{{word_namespace}}}tc"

    def cell_text(cell_elem: Any) -> str:
        texts: list[str] = []
        for paragraph in cell_elem.findall(f".//{paragraph_tag}"):
            text_value = "".join(node.text or "" for node in paragraph.iter(f"{{{word_namespace}}}t"))
            if text_value.strip():
                texts.append(text_value.strip())
        return " ".join(texts)

    def table_to_markdown(table_elem: Any) -> str:
        rows = table_elem.findall(f".//{row_tag}")
        if not rows:
            return ""
        markdown_rows: list[str] = []
        for row in rows:
            cells = row.findall(f".//{cell_tag}")
            markdown_rows.append("| " + " | ".join(cell_text(cell) for cell in cells) + " |")
        if len(markdown_rows) > 1:
            column_count = markdown_rows[0].count("|") - 1
            markdown_rows.insert(1, "| " + " | ".join(["---"] * max(column_count, 1)) + " |")
        return "\n".join(markdown_rows)

    doc_blocks: list[dict[str, Any]] = []
    locator_index = 0
    for body_index, child in enumerate(doc.element.body):
        if child.tag == paragraph_tag:
            text_value = "".join(node.text or "" for node in child.iter(f"{{{word_namespace}}}t"))
            if not text_value.strip():
                continue
            block_type = "paragraph"
            block_text = text_value
        elif child.tag == table_tag:
            block_text = table_to_markdown(child)
            if not block_text:
                continue
            block_type = "table"
        else:
            continue

        locator = f"P{locator_index:04d}"
        doc_blocks.append(
            {
                "block_id": f"B{locator_index:06d}",
                "locator": locator,
                "body_idx": body_index,
                "type": block_type,
                "text": block_text,
            }
        )
        locator_index += 1
    return doc_blocks


def _normalize_locator_token(raw: str) -> str:
    normalized = str(raw or "").strip().upper()
    if not normalized:
        return ""
    match = re.search(r"P\s*0*(\d+)", normalized)
    if match is None:
        return ""
    return f"P{int(match.group(1)):04d}"


def _find_doc_block_by_locator(blocks: list[Any], locator: str) -> Mapping[str, Any] | None:
    normalized = _normalize_locator_token(locator)
    for block in blocks:
        if not isinstance(block, Mapping):
            continue
        if _normalize_locator_token(str(block.get("locator") or "")) == normalized:
            return block
    return None


def _find_doc_block_by_id(blocks: list[Any], block_id: str) -> Mapping[str, Any] | None:
    normalized = str(block_id or "").strip()
    for block in blocks:
        if isinstance(block, Mapping) and str(block.get("block_id") or "").strip() == normalized:
            return block
    return None


def _doc_blocks_slice_to_html(blocks: list[Any], start_body_idx: int, end_body_idx: int) -> str:
    lo, hi = (start_body_idx, end_body_idx) if start_body_idx <= end_body_idx else (end_body_idx, start_body_idx)
    selected: list[Mapping[str, Any]] = []
    for block in blocks:
        if not isinstance(block, Mapping):
            continue
        body_idx = _non_negative_int(block.get("body_idx"))
        if lo <= body_idx <= hi:
            selected.append(block)
    selected.sort(key=lambda item: _non_negative_int(item.get("body_idx")))

    html_parts: list[str] = []
    for block in selected:
        text_value = str(block.get("text") or "").strip()
        if not text_value:
            continue
        escaped = html.escape(text_value).replace("\n", "<br/>")
        if str(block.get("type") or "").lower() == "table":
            html_parts.append(f"<pre>{escaped}</pre>")
        else:
            html_parts.append(f"<p>{escaped}</p>")
    return "\n".join(html_parts)


def _slice_docx_bytes_by_body_range(docx_bytes: bytes, start_body_idx: int, end_body_idx: int) -> bytes:
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    body_tag = f"{{{word_namespace}}}body"
    sectpr_tag = f"{{{word_namespace}}}sectPr"
    paragraph_tag = f"{{{word_namespace}}}p"
    text_tag = f"{{{word_namespace}}}t"
    run_tag = f"{{{word_namespace}}}r"
    break_tag = f"{{{word_namespace}}}br"
    last_rendered_page_break_tag = f"{{{word_namespace}}}lastRenderedPageBreak"
    paragraph_props_tag = f"{{{word_namespace}}}pPr"
    page_break_before_tag = f"{{{word_namespace}}}pageBreakBefore"

    input_buffer = io.BytesIO(docx_bytes)
    output_buffer = io.BytesIO()
    with zipfile.ZipFile(input_buffer, "r") as zin:
        entries = {info.filename: zin.read(info.filename) for info in zin.infolist()}

    doc_xml = entries.get("word/document.xml")
    if not doc_xml:
        raise ValueError("DOCX 包缺少 word/document.xml")

    root = ET.fromstring(doc_xml)
    body = root.find(f".//{body_tag}")
    if body is None:
        raise ValueError("document.xml 缺少 w:body")

    original_children = list(body)
    kept_children: list[ET.Element] = []
    section_props: ET.Element | None = None

    def paragraph_plain_text(elem: ET.Element) -> str:
        texts = []
        for text_node in elem.iter(text_tag):
            text_value = text_node.text or ""
            if text_value.strip():
                texts.append(text_value.strip())
        return "".join(texts).strip()

    def is_effectively_empty_paragraph(elem: ET.Element) -> bool:
        return elem.tag == paragraph_tag and paragraph_plain_text(elem) == ""

    def trim_empty_paragraphs(children: list[ET.Element]) -> list[ET.Element]:
        lo = 0
        hi = len(children) - 1
        while lo <= hi and is_effectively_empty_paragraph(children[lo]):
            lo += 1
        while hi >= lo and is_effectively_empty_paragraph(children[hi]):
            hi -= 1
        return children[lo: hi + 1] if lo <= hi else children

    def remove_first_paragraph_page_break_controls(children: list[ET.Element]) -> None:
        if not children:
            return
        first = children[0]
        if first.tag != paragraph_tag:
            return
        paragraph_props = first.find(paragraph_props_tag)
        if paragraph_props is not None:
            for node in list(paragraph_props):
                if node.tag == page_break_before_tag:
                    paragraph_props.remove(node)
        else:
            paragraph_props = ET.Element(paragraph_props_tag)
            first.insert(0, paragraph_props)

        page_break_override = ET.Element(page_break_before_tag)
        page_break_override.set(f"{{{word_namespace}}}val", "0")
        paragraph_props.insert(0, page_break_override)

        for run in list(first.findall(run_tag)):
            for node in list(run):
                if node.tag == last_rendered_page_break_tag:
                    run.remove(node)
                    continue
                if node.tag == break_tag and (node.attrib.get(f"{{{word_namespace}}}type") or "").lower() == "page":
                    run.remove(node)
            if len(list(run)) == 0:
                first.remove(run)

    lo, hi = (start_body_idx, end_body_idx) if start_body_idx <= end_body_idx else (end_body_idx, start_body_idx)
    for idx, child in enumerate(original_children):
        if child.tag == sectpr_tag:
            section_props = copy.deepcopy(child)
            continue
        if lo <= idx <= hi:
            kept_children.append(copy.deepcopy(child))

    if not kept_children:
        raise ValueError("切片范围内无可用文档块")

    kept_children = trim_empty_paragraphs(kept_children)
    remove_first_paragraph_page_break_controls(kept_children)

    for child in list(body):
        body.remove(child)
    for child in kept_children:
        body.append(child)
    if section_props is not None:
        body.append(section_props)

    entries["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(output_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, content in entries.items():
            zout.writestr(name, content)
    return output_buffer.getvalue()


def _build_scoring_table_xlsx(*, project_name: str, rows: list[Any]) -> bytes:
    try:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    except ImportError as exc:
        raise PlatformError(code="EXPORT_FAILED", message="服务器缺少 openpyxl，请先 pip install openpyxl", status_code=500) from exc

    _ = project_name
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "自评评分表"

    header_fill = PatternFill("solid", fgColor="1A6FA8")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    headers = ["评分指标", "最高分", "评分标准", "自评情况", "自评说明", "证明材料引用"]
    col_widths = [30, 10, 40, 12, 50, 45]
    for col_idx, (header, width) in enumerate(zip(headers, col_widths), 1):
        cell = worksheet.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border
        worksheet.column_dimensions[worksheet.cell(row=1, column=col_idx).column_letter].width = width
    worksheet.row_dimensions[1].height = 28

    response_label = {"full": "响应", "partial": "部分响应", "none": "不响应", "": "未填写"}
    total_max = 0.0
    for row_idx, raw_row in enumerate(rows, 2):
        row = raw_row if isinstance(raw_row, Mapping) else {}
        max_score = _number_value(row.get("max_score"))
        total_max += max_score
        values = [
            str(row.get("indicator") or ""),
            max_score,
            str(row.get("criteria") or ""),
            response_label.get(str(row.get("self_response") or ""), str(row.get("self_response") or "")),
            str(row.get("self_comment") or ""),
            "\n".join(str(item) for item in row.get("evidence_refs", []) if str(item).strip())
            if isinstance(row.get("evidence_refs"), list)
            else "",
        ]
        for col_idx, value in enumerate(values, 1):
            cell = worksheet.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = thin_border
        worksheet.row_dimensions[row_idx].height = 60

    worksheet.append(["合计", total_max, "", "", "", ""])
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _build_analysis_report_pdf(*, project_name: str, nodes: list[Any]) -> bytes:
    _register_report_font()
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        leftMargin=14 * mm,
        rightMargin=14 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
        title=f"{project_name} - 招标文件解析报告",
    )
    styles = _analysis_report_styles()
    story: list[Any] = [
        Paragraph(_escape_report_text(f"{project_name} - 招标文件解析报告"), styles["title"]),
        Paragraph(f"导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}", styles["meta"]),
        Spacer(1, 8),
    ]
    story.extend(_analysis_report_node_flowables(nodes, styles, depth=0))
    document.build(story)
    return output.getvalue()


def _register_report_font() -> None:
    for font_name in ("STSong-Light",):
        try:
            pdfmetrics.getFont(font_name)
            return
        except Exception:
            try:
                pdfmetrics.registerFont(UnicodeCIDFont(font_name))
                return
            except Exception:
                continue


def _analysis_report_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    font_name = "STSong-Light"
    return {
        "title": ParagraphStyle(
            "BidReportTitle",
            parent=base["Title"],
            fontName=font_name,
            fontSize=18,
            leading=24,
            textColor=colors.HexColor("#0c4a6e"),
            spaceAfter=4,
        ),
        "meta": ParagraphStyle(
            "BidReportMeta",
            parent=base["Normal"],
            fontName=font_name,
            fontSize=9,
            leading=13,
            textColor=colors.HexColor("#6b7280"),
            spaceAfter=12,
        ),
        "h2": ParagraphStyle(
            "BidReportH2",
            parent=base["Heading2"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            textColor=colors.HexColor("#0c4a6e"),
            spaceBefore=12,
            spaceAfter=6,
        ),
        "h3": ParagraphStyle(
            "BidReportH3",
            parent=base["Heading3"],
            fontName=font_name,
            fontSize=12,
            leading=16,
            textColor=colors.HexColor("#1e3a5f"),
            spaceBefore=8,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "BidReportBody",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=10,
            leading=15,
            textColor=colors.HexColor("#374151"),
            spaceAfter=6,
        ),
        "empty": ParagraphStyle(
            "BidReportEmpty",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=9,
            leading=13,
            textColor=colors.HexColor("#9ca3af"),
            spaceAfter=6,
        ),
    }


def _analysis_report_node_flowables(nodes: list[Any], styles: Mapping[str, ParagraphStyle], *, depth: int) -> list[Any]:
    flowables: list[Any] = []
    for raw_node in nodes:
        if not isinstance(raw_node, Mapping):
            continue
        label = str(raw_node.get("label") or raw_node.get("title") or raw_node.get("id") or "未命名节点")
        content = str(raw_node.get("content") or "")
        children = raw_node.get("children") if isinstance(raw_node.get("children"), list) else []
        heading_style = styles["h2"] if depth <= 0 else styles["h3"]
        flowables.append(Paragraph(_escape_report_text(label), heading_style))
        if children:
            flowables.extend(_analysis_report_node_flowables(children, styles, depth=depth + 1))
        elif content.strip():
            flowables.extend(_analysis_report_content_flowables(content, styles))
        else:
            flowables.append(Paragraph("（未提取）", styles["empty"]))
    return flowables


def _analysis_report_content_flowables(content: str, styles: Mapping[str, ParagraphStyle]) -> list[Any]:
    parsed_score_rows = _parse_analysis_score_rows(content)
    if parsed_score_rows:
        return [_build_analysis_score_table(parsed_score_rows, styles)]

    field_rows = _parse_analysis_xml_fields(content)
    if field_rows:
        rows = [
            [
                Paragraph(_escape_report_text(label), styles["body"]),
                Paragraph(_escape_report_text(value), styles["body"]),
            ]
            for label, value in field_rows
        ]
        table = Table(rows, colWidths=[35 * mm, 135 * mm])
        table.setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#e5e7eb")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f8fafc")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ]
            )
        )
        return [table, Spacer(1, 6)]

    paragraphs = [line.strip() for line in content.replace("\r\n", "\n").split("\n") if line.strip()]
    return [Paragraph(_escape_report_text(line), styles["body"]) for line in (paragraphs or [content])]


def _parse_analysis_score_rows(content: str) -> list[Mapping[str, Any]]:
    text_value = str(content or "").strip()
    if not text_value or text_value[0:1] not in {"{", "[", '"'}:
        return []
    try:
        parsed: Any = json.loads(text_value)
        if isinstance(parsed, str):
            parsed = json.loads(parsed)
    except Exception:
        return []
    if isinstance(parsed, Mapping):
        items = parsed.get("items")
    elif isinstance(parsed, list):
        items = parsed
    else:
        items = None
    if not isinstance(items, list):
        return []
    return [item for item in items if isinstance(item, Mapping)]


def _build_analysis_score_table(items: list[Mapping[str, Any]], styles: Mapping[str, ParagraphStyle]) -> Table:
    rows: list[list[Any]] = [
        [
            Paragraph("评分项", styles["body"]),
            Paragraph("评分规则", styles["body"]),
            Paragraph("满分", styles["body"]),
        ]
    ]
    total = 0.0
    for item in items:
        max_score = _number_value(item.get("max_score"))
        total += max_score
        rows.append(
            [
                Paragraph(_escape_report_text(str(item.get("name") or "")), styles["body"]),
                Paragraph(_escape_report_text(str(item.get("criteria") or "")), styles["body"]),
                Paragraph(_escape_report_text(f"{max_score:g}分"), styles["body"]),
            ]
        )
    rows.append([Paragraph("合计", styles["body"]), "", Paragraph(_escape_report_text(f"{total:g}分"), styles["body"])])
    table = Table(rows, colWidths=[35 * mm, 125 * mm, 20 * mm], repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#d1d5db")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e0f2fe")),
                ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#f3f4f6")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    return table


def _parse_analysis_xml_fields(content: str) -> list[tuple[str, str]]:
    rows = []
    for label, value in re.findall(r"<([^/>\s][^>]*)>([\s\S]*?)</\1>", str(content or "")):
        if label.startswith("要点"):
            continue
        text_value = re.sub(r"<[^>]+>", "", value).strip()
        if text_value:
            rows.append((label, text_value))
    return rows


def _escape_report_text(value: str) -> str:
    return html.escape(str(value or "")).replace("\n", "<br/>")


def _build_forge_document_docx(
    *,
    project_id: str,
    sections: list[Any],
    scoring_rows: list[Any],
    attachments: list[Any],
    mapping_table: dict[str, Any],
    bidder_info: dict[str, Any],
    image_map: dict[str, Any],
) -> bytes:
    normalized_sections = [dict(item) for item in sections if isinstance(item, Mapping)]
    normalized_scoring_rows = [dict(item) for item in scoring_rows if isinstance(item, Mapping)]
    normalized_attachments = [dict(item) for item in attachments if isinstance(item, Mapping)]
    all_content = " ".join(
        [str(section.get("content") or "") for section in normalized_sections]
        + [str(attachment.get("content") or "") for attachment in normalized_attachments]
    )
    pipt_mapping = _load_forge_pipt_mapping(all_content)
    dynamic_image_map = _load_forge_image_map(all_content)
    full_mapping = {**pipt_mapping, **{str(key): str(value) for key, value in mapping_table.items()}}
    merged_image_map = {**image_map, **dynamic_image_map}

    forge = create_document_forge(
        mapping_table=full_mapping,
        bidder_info=bidder_info,
        image_map=merged_image_map,
        project_id=project_id,
    )
    has_docx_slice = any(
        str(section.get("source_type") or "").strip().lower() == "docx_slice"
        for section in normalized_sections
    )
    heading_sanitized = True
    if has_docx_slice:
        docx_bytes, heading_sanitized = _build_hybrid_forge_docx(
            project_id=project_id,
            sections=normalized_sections,
            scoring_rows=normalized_scoring_rows,
            attachments=normalized_attachments,
            forge=forge,
        )
    else:
        docx_bytes = forge.build(
            sections=normalized_sections,
            scoring_rows=normalized_scoring_rows,
            attachments=normalized_attachments,
        )
    docx_bytes = _rebind_heading_numbering_bytes(docx_bytes)
    return _apply_toc_for_export(
        docx_bytes,
        normalized_sections,
        prefer_native=True,
        heading_sanitized=heading_sanitized,
    )


def _build_hybrid_forge_docx(
    *,
    project_id: str,
    sections: list[dict[str, Any]],
    scoring_rows: list[dict[str, Any]],
    attachments: list[dict[str, Any]],
    forge: Any,
) -> tuple[bytes, bool]:
    try:
        import docx as docx_module
        from docx.enum.text import WD_BREAK
        from docxcompose.composer import Composer
    except ImportError as exc:
        raise PlatformError(code="FORGE_FAILED", message=f"统一后端缺少 DOCX 拼装依赖: {exc}", status_code=500) from exc

    normalized_project_id = _ensure_safe_project_id(project_id)
    docx_segments: list[tuple[bytes, bool]] = []
    heading_sanitized = True
    pending_markdown_sections: list[dict[str, Any]] = []

    def flush_markdown_sections() -> None:
        nonlocal pending_markdown_sections
        if not pending_markdown_sections:
            return
        markdown_segment = forge.build(
            sections=pending_markdown_sections,
            scoring_rows=[],
            attachments=[],
        )
        docx_segments.append((markdown_segment, True))
        pending_markdown_sections = []

    def prepend_page_break(doc_obj: Any) -> Any:
        paragraph = doc_obj.paragraphs[0].insert_paragraph_before() if doc_obj.paragraphs else doc_obj.add_paragraph()
        paragraph.add_run().add_break(WD_BREAK.PAGE)
        return doc_obj

    for section in sections:
        source_type = str(section.get("source_type") or "markdown").strip().lower()
        if source_type == "docx_slice":
            flush_markdown_sections()
            if section.get("inject_title") and str(section.get("title") or "").strip():
                heading_segment = forge.build(
                    sections=[
                        {
                            "id": section.get("id", ""),
                            "title": section.get("title", ""),
                            "heading_number": section.get("heading_number", ""),
                            "heading_text": section.get("heading_text", ""),
                            "bookmark_id": section.get("bookmark_id", ""),
                            "content": "",
                            "heading_level": section.get("heading_level", 1),
                            "title_only": True,
                        }
                    ],
                    scoring_rows=[],
                    attachments=[],
                )
                docx_segments.append((heading_segment, True))
            slice_segment, sanitized_ok = _build_docx_slice_segment(section=section, project_id=normalized_project_id)
            heading_sanitized = heading_sanitized and bool(sanitized_ok)
            docx_segments.append((slice_segment, False))
            continue

        pending_markdown_sections.append({
            "id": section.get("id", ""),
            "title": section.get("title", ""),
            "heading_number": section.get("heading_number", ""),
            "heading_text": section.get("heading_text", ""),
            "bookmark_id": section.get("bookmark_id", ""),
            "content": section.get("content", ""),
            "heading_level": section.get("heading_level", 1),
            "title_only": section.get("title_only", False),
        })

    flush_markdown_sections()

    if not docx_segments:
        return forge.build(sections=sections, scoring_rows=scoring_rows, attachments=attachments), heading_sanitized

    master = docx_module.Document(io.BytesIO(docx_segments[0][0]))
    composer = Composer(master)
    for segment_bytes, with_page_break in docx_segments[1:]:
        segment_doc = docx_module.Document(io.BytesIO(segment_bytes))
        composer.append(prepend_page_break(segment_doc) if with_page_break else segment_doc)

    merged_buffer = io.BytesIO()
    composer.save(merged_buffer)
    merged_buffer.seek(0)

    final_doc = docx_module.Document(merged_buffer)
    add_scoring_table_and_attachments(final_doc, scoring_rows, attachments)
    _rebind_heading_numbering_for_export(final_doc)

    output = io.BytesIO()
    final_doc.save(output)
    output.seek(0)
    return output.read(), heading_sanitized


def _build_docx_slice_segment(*, section: Mapping[str, Any], project_id: str) -> tuple[bytes, bool]:
    start_block_id = str(section.get("start_block_id") or "").strip()
    end_block_id = str(section.get("end_block_id") or "").strip()
    if not start_block_id or not end_block_id:
        raise PlatformError(code="INVALID_REQUEST", message="docx_slice 段缺少 start_block_id/end_block_id", status_code=400)

    blocks = get_project_doc_blocks_payload(project_id)["blocks"]
    start_block = _find_doc_block_by_id(blocks, start_block_id)
    end_block = _find_doc_block_by_id(blocks, end_block_id)
    if start_block is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message=f"block_id {start_block_id} 未找到", status_code=404)
    if end_block is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message=f"block_id {end_block_id} 未找到", status_code=404)

    start_idx = _non_negative_int(start_block.get("body_idx"))
    end_idx = _non_negative_int(end_block.get("body_idx"))
    if start_idx > end_idx:
        start_idx, end_idx = end_idx, start_idx

    docx_path = _docx_cache_path(project_id)
    if not docx_path.exists():
        raise PlatformError(
            code="BUSINESS_DIRECT_ERROR",
            message="原始 DOCX 不可用，无法生成保格式切片；请上传原始 DOCX 重建定位缓存",
            status_code=409,
        )
    try:
        sliced = _slice_docx_bytes_by_body_range(docx_path.read_bytes(), start_idx, end_idx)
    except Exception as exc:
        raise PlatformError(code="FORGE_FAILED", message=f"DOCX 切片失败: {exc}", status_code=500) from exc
    return _sanitize_docx_slice_heading_semantics(sliced)


def _sanitize_docx_slice_heading_semantics(docx_bytes: bytes) -> tuple[bytes, bool]:
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": word_namespace}

    def qn(tag: str) -> str:
        return f"{{{word_namespace}}}{tag}"

    def is_heading_like_style_id(style_id: str) -> bool:
        normalized = str(style_id or "").strip().lower()
        return bool(re.match(r"heading[1-9]\d*$", normalized) or normalized.startswith("toc"))

    input_buffer = io.BytesIO(docx_bytes)
    output_buffer = io.BytesIO()
    with zipfile.ZipFile(input_buffer, "r") as zin:
        entries = {info.filename: zin.read(info.filename) for info in zin.infolist()}

    doc_xml = entries.get("word/document.xml")
    if not doc_xml:
        return docx_bytes, False

    heading_style_ids: set[str] = set()
    style_font_size_map: dict[str, str] = {}
    styles_xml = entries.get("word/styles.xml")
    if styles_xml:
        try:
            styles_root = ET.fromstring(styles_xml)
            for style in styles_root.findall(".//w:style", ns):
                if style.attrib.get(qn("type")) != "paragraph":
                    continue
                style_id = style.attrib.get(qn("styleId"), "")
                size_el = style.find("w:rPr/w:sz", ns)
                if size_el is not None and size_el.attrib.get(qn("val")):
                    style_font_size_map[style_id] = str(size_el.attrib.get(qn("val")))
                name_el = style.find("w:name", ns)
                style_name = name_el.attrib.get(qn("val"), "") if name_el is not None else ""
                if is_heading_like_style_id(style_id) or str(style_name).strip().lower().startswith(("heading", "toc")):
                    heading_style_ids.add(style_id)
        except Exception:
            heading_style_ids = set()
            style_font_size_map = {}

    changed = False
    try:
        root = ET.fromstring(doc_xml)
        for paragraph_props in root.findall(".//w:p/w:pPr", ns):
            p_style = paragraph_props.find("w:pStyle", ns)
            if p_style is not None:
                style_id = p_style.attrib.get(qn("val"), "")
                if style_id in heading_style_ids or is_heading_like_style_id(style_id):
                    inherited_size = style_font_size_map.get(style_id, "")
                    paragraph_props.remove(p_style)
                    changed = True
                    if inherited_size:
                        run_props = paragraph_props.find("w:rPr", ns)
                        if run_props is None:
                            run_props = ET.SubElement(paragraph_props, qn("rPr"))
                        size = run_props.find("w:sz", ns)
                        if size is None:
                            size = ET.SubElement(run_props, qn("sz"))
                        size.set(qn("val"), inherited_size)
                        size_cs = run_props.find("w:szCs", ns)
                        if size_cs is None:
                            size_cs = ET.SubElement(run_props, qn("szCs"))
                        size_cs.set(qn("val"), inherited_size)
            outline_level = paragraph_props.find("w:outlineLvl", ns)
            if outline_level is not None:
                paragraph_props.remove(outline_level)
                changed = True
    except Exception:
        return docx_bytes, False

    if not changed:
        return docx_bytes, True
    entries["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(output_buffer, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, content in entries.items():
            zout.writestr(name, content)
    return output_buffer.getvalue(), True


def _build_toc_entries(sections: list[dict[str, Any]]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    for section in sections:
        level = _non_negative_int(section.get("toc_level") or section.get("heading_level"))
        if level < 1 or level > 3:
            continue
        number = str(section.get("heading_number") or "").strip()
        raw_title = str(section.get("heading_text") or section.get("title") or "").strip()
        title = f"{number} {raw_title}".strip() if number else raw_title
        if title:
            entries.append({"level": level, "text": title})
    return entries


def _doc_has_unexpected_heading_semantics(doc: Any, allowed_headings: set[str]) -> bool:
    def normalized(value: str) -> str:
        return re.sub(r"\s+", "", str(value or "")).strip().lower()

    for paragraph in getattr(doc, "paragraphs", []) or []:
        try:
            style_name = (paragraph.style.name if paragraph.style else "") or ""
        except Exception:
            style_name = ""
        text_value = str(getattr(paragraph, "text", "") or "").strip()
        if not text_value:
            continue
        style_lower = style_name.lower()
        if style_lower.startswith("toc"):
            return True
        if style_lower.startswith("heading") and normalized(text_value) not in allowed_headings and text_value != "目录":
            return True
    return False


def _insert_toc_page(doc: Any, toc_entries: list[dict[str, Any]], *, use_native_toc: bool) -> None:
    if not toc_entries:
        return
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as docx_qn

    anchor = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    title_paragraph = anchor.insert_paragraph_before("目录")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_paragraph.runs:
        title_paragraph.runs[0].bold = True

    if use_native_toc:
        field_paragraph = anchor.insert_paragraph_before("")
        field = OxmlElement("w:fldSimple")
        field.set(docx_qn("w:instr"), ' TOC \\o "1-3" \\h \\z \\u ')
        run = OxmlElement("w:r")
        text_node = OxmlElement("w:t")
        text_node.text = "（在 Word 中右键目录并选择“更新域”）"
        run.append(text_node)
        field.append(run)
        field_paragraph._p.append(field)
    else:
        for row in toc_entries:
            indent = "    " * max(0, _non_negative_int(row.get("level")) - 1)
            anchor.insert_paragraph_before(f"{indent}{row.get('text', '')}")

    split_paragraph = anchor.insert_paragraph_before("")
    split_paragraph.add_run().add_break(WD_BREAK.PAGE)


def _rebind_heading_numbering_for_export(doc: Any) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn as docx_qn

    try:
        numbering = doc.part.numbering_part.numbering_definitions._numbering
    except Exception:
        return

    abstract_ids = [
        int(node.get(docx_qn("w:abstractNumId")))
        for node in numbering.findall(docx_qn("w:abstractNum"))
        if node.get(docx_qn("w:abstractNumId"))
    ]
    num_ids = [
        int(node.get(docx_qn("w:numId")))
        for node in numbering.findall(docx_qn("w:num"))
        if node.get(docx_qn("w:numId"))
    ]
    next_abstract_id = max(abstract_ids, default=1999) + 1
    next_num_id = max(num_ids, default=1999) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(docx_qn("w:abstractNumId"), str(next_abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(docx_qn("w:val"), "multilevel")
    abstract.append(multi)
    for level, number_format, text_value in ((0, "chineseCounting", "%1、"), (1, "decimal", "%1.%2"), (2, "decimal", "%1.%2.%3")):
        lvl = OxmlElement("w:lvl")
        lvl.set(docx_qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(docx_qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(docx_qn("w:val"), number_format)
        lvl.append(fmt)
        if level >= 1:
            lvl.append(OxmlElement("w:isLgl"))
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(docx_qn("w:val"), text_value)
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(docx_qn("w:val"), "left")
        lvl.append(lvl_jc)
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(docx_qn("w:numId"), str(next_num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(docx_qn("w:val"), str(next_abstract_id))
    num.append(abs_ref)
    numbering.append(num)

    def heading_level(style_name: str) -> int:
        normalized = str(style_name or "").strip().lower()
        if normalized.startswith(("heading 1", "heading1", "标题1", "标题 1")):
            return 1
        if normalized.startswith(("heading 2", "heading2", "标题2", "标题 2")):
            return 2
        if normalized.startswith(("heading 3", "heading3", "标题3", "标题 3")):
            return 3
        return 0

    for paragraph in getattr(doc, "paragraphs", []) or []:
        text_value = str(getattr(paragraph, "text", "") or "").strip()
        if not text_value or text_value == "目录":
            continue
        try:
            level = heading_level(paragraph.style.name if paragraph.style else "")
        except Exception:
            level = 0
        if level < 1:
            continue
        paragraph_props = paragraph._p.get_or_add_pPr()
        old = paragraph_props.find(docx_qn("w:numPr"))
        if old is not None:
            paragraph_props.remove(old)
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(docx_qn("w:val"), str(level - 1))
        num_id = OxmlElement("w:numId")
        num_id.set(docx_qn("w:val"), str(next_num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id)
        paragraph_props.append(num_pr)


def _rebind_heading_numbering_bytes(docx_bytes: bytes) -> bytes:
    try:
        import docx as docx_module

        doc = docx_module.Document(io.BytesIO(docx_bytes))
        _rebind_heading_numbering_for_export(doc)
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()
    except Exception:
        return docx_bytes


def _apply_toc_for_export(
    docx_bytes: bytes,
    sections: list[dict[str, Any]],
    *,
    prefer_native: bool,
    heading_sanitized: bool,
) -> bytes:
    try:
        import docx as docx_module
    except ImportError as exc:
        raise PlatformError(code="FORGE_FAILED", message=f"统一后端缺少 python-docx 依赖: {exc}", status_code=500) from exc

    def strip_heading_prefix(value: str) -> str:
        return re.sub(r"^(([一二三四五六七八九十百千万]+、)|(\d+(?:\.\d+){1,2}))\s*", "", str(value or "").strip()).strip()

    toc_entries = _build_toc_entries(sections)
    if not toc_entries:
        return docx_bytes

    doc = docx_module.Document(io.BytesIO(docx_bytes))
    allowed: set[str] = set()
    for row in toc_entries:
        text_value = str(row.get("text") or "").strip()
        if text_value:
            allowed.add(re.sub(r"\s+", "", text_value).strip().lower())
        stripped = strip_heading_prefix(text_value)
        if stripped:
            allowed.add(re.sub(r"\s+", "", stripped).strip().lower())
    use_native = bool(prefer_native and heading_sanitized and not _doc_has_unexpected_heading_semantics(doc, allowed))
    _insert_toc_page(doc, toc_entries, use_native_toc=use_native)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()


def _load_forge_pipt_mapping(content: str) -> dict[str, str]:
    placeholders = _find_forge_pipt_placeholders(content)
    if not placeholders:
        return {}
    legacy_placeholders = [item for item in placeholders if item.startswith("{{__PIPT_")]
    strong_placeholders = [item for item in placeholders if _is_strong_pipt_token(item)]
    mapping: dict[str, str] = {}
    engine = get_engine()
    with engine.connect() as conn:
        exists = conn.execute(text("SELECT to_regclass('bid_generator.entity_registry') IS NOT NULL")).scalar_one()
        if not exists:
            return {}
        stmt = text(
            """
            SELECT placeholder, strong_placeholder, original_text_enc
            FROM bid_generator.entity_registry
            WHERE placeholder IN :legacy_placeholders
               OR strong_placeholder IN :strong_placeholders
            """
        ).bindparams(
            bindparam("legacy_placeholders", expanding=True),
            bindparam("strong_placeholders", expanding=True),
        )
        rows = conn.execute(
            stmt,
            {
                "legacy_placeholders": legacy_placeholders or ["__none__"],
                "strong_placeholders": strong_placeholders or ["__none__"],
            },
        ).mappings().all()
    for row in rows:
        original = _decrypt_forge_original_text(str(row.get("original_text_enc") or ""))
        placeholder = str(row.get("placeholder") or "")
        strong_placeholder = str(row.get("strong_placeholder") or "")
        if placeholder in placeholders:
            mapping[placeholder] = original
        if strong_placeholder in placeholders:
            mapping[strong_placeholder] = original
    return mapping


def _find_forge_pipt_placeholders(content: str) -> list[str]:
    seen: set[str] = set()
    placeholders: list[str] = []
    for pattern in (r"\{\{__PIPT_[a-zA-Z0-9_]+__\}\}", r"@@PIPT:v1:e\d{6}:k[a-f0-9]{8}@@"):
        for match in re.finditer(pattern, str(content or "")):
            token = match.group(0)
            if token not in seen:
                seen.add(token)
                placeholders.append(token)
    return placeholders


def _is_strong_pipt_token(value: str) -> bool:
    return bool(re.fullmatch(r"@@PIPT:v1:e\d{6}:k[a-f0-9]{8}@@", str(value or "")))


def _decrypt_forge_original_text(value: str) -> str:
    raw_key = os.environ.get("PIPT_DB_KEY", "")
    if not raw_key:
        return value
    try:
        from cryptography.fernet import Fernet

        return Fernet(raw_key.encode() if isinstance(raw_key, str) else raw_key).decrypt(value.encode("ascii")).decode("utf-8")
    except Exception:
        return value


def _load_forge_image_map(content: str) -> dict[str, dict[str, str]]:
    placeholders = sorted(set(re.findall(r"__PRO_IMG_[a-f0-9]+__", str(content or ""))))
    if not placeholders:
        return {}
    image_hashes = [placeholder.replace("__PRO_IMG_", "").replace("__", "") for placeholder in placeholders]
    engine = get_engine()
    with engine.connect() as conn:
        exists = conn.execute(text("SELECT to_regclass('bid_generator.image_registry') IS NOT NULL")).scalar_one()
        if not exists:
            return {}
        stmt = text(
            """
            SELECT placeholder, abs_path, preview_url
            FROM bid_generator.image_registry
            WHERE image_hash IN :image_hashes
            """
        ).bindparams(bindparam("image_hashes", expanding=True))
        rows = conn.execute(stmt, {"image_hashes": image_hashes or ["__none__"]}).mappings().all()
    return {
        str(row.get("placeholder") or ""): {
            "abs_path": str(row.get("abs_path") or ""),
            "preview_url": str(row.get("preview_url") or ""),
        }
        for row in rows
        if str(row.get("placeholder") or "")
    }


def _analysis_report_mirror_path(project_id: str) -> Path:
    return _bid_generator_legacy_root() / "data" / "projects" / f"{project_id}_analysis.json"


def _pdf_cache_path(project_id: str) -> Path:
    return _bid_generator_legacy_root() / "data" / "pdf_cache" / f"{project_id}.pdf"


def _docx_cache_path(project_id: str) -> Path:
    return _bid_generator_legacy_root() / "data" / "docx_cache" / f"{project_id}.docx"


def _raw_doc_cache_path(project_id: str) -> Path:
    return _bid_generator_legacy_root() / "data" / "raw_doc_cache" / f"{project_id}.txt"


def _load_raw_document(project_id: str) -> str:
    normalized_id = str(project_id or "").strip()
    if not normalized_id:
        return ""
    raw_path = _raw_doc_cache_path(normalized_id)
    if not raw_path.exists():
        return ""
    try:
        return raw_path.read_text(encoding="utf-8")
    except OSError as exc:
        logger.warning("[%s] 读取 raw_document 缓存失败: %s", normalized_id, exc)
        return ""


def _extracted_image_path(filename: str) -> Path:
    return _bid_generator_legacy_root() / "data" / "extracted_images" / filename


def _kb_sync_status_dir() -> Path:
    return _bid_generator_legacy_root() / "data" / "kb_sync_status"


def _kb_sync_status_path(job_id: str) -> Path:
    return _kb_sync_status_dir() / f"{job_id}.json"


def _read_json_file(path: Path) -> Any:
    try:
        with path.open("r", encoding="utf-8") as file:
            return json.load(file)
    except json.JSONDecodeError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message=f"JSON 文件格式无效: {path.name}", status_code=500) from exc
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message=f"读取文件失败: {path.name}", status_code=500) from exc


def _diagram_artifact_dir() -> Path:
    return Path(os.environ.get("DIAGRAM_ARTIFACT_DIR", str(_bid_generator_legacy_root() / "data" / "diagram_artifacts")))


def _read_text_file(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message=f"读取文件失败: {path.name}", status_code=500) from exc


def _read_text_artifact_payload(path: Path, *, media_type: str) -> BidGeneratorFilePayload:
    return BidGeneratorFilePayload(
        content=_read_text_file(path).encode("utf-8"),
        media_type=media_type,
        filename=path.name,
        inline=True,
        cache_control="public, max-age=86400",
    )


def _find_diagram_artifact_path(root: Path, diagram_id: str, project: str, suffix: str) -> Path | None:
    path = root / project / f"{diagram_id}{suffix}"
    if path.exists():
        return path
    try:
        for candidate in root.glob(f"*/{diagram_id}{suffix}"):
            return candidate
    except OSError:
        return None
    return None


def _safe_diagram_project_dir(project_id: str) -> str:
    return re.sub(r"[^a-zA-Z0-9_.-]+", "_", str(project_id or "default"))


def _read_yaml_mapping(path: Path | None) -> dict[str, Any]:
    if path is None:
        return {}
    try:
        with path.open("r", encoding="utf-8") as file:
            loaded = yaml.safe_load(file) or {}
    except FileNotFoundError as exc:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message=f"文件不存在: {path.name}", status_code=404) from exc
    except yaml.YAMLError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message=f"YAML 文件格式无效: {path.name}", status_code=500) from exc
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message=f"读取 YAML 文件失败: {path.name}", status_code=500) from exc
    return loaded if isinstance(loaded, dict) else {}


def _write_yaml_mapping(path: Path, data: Mapping[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as file:
        yaml.safe_dump(dict(data), file, allow_unicode=True, sort_keys=False)


def _template_structures_dir() -> Path:
    return _bid_generator_resource_root() / "templates" / "structures"


_BUILTIN_TEMPLATE_NAMES: frozenset[str] = frozenset({"standard.yaml", "AI 专属大纲"})


def _list_template_structure_names(templates_dir: Path) -> list[str]:
    if not templates_dir.exists():
        return []
    names = []
    for path in templates_dir.iterdir():
        if not path.is_file():
            continue
        if path.suffix in {".yaml", ".yml"} or path.name in _BUILTIN_TEMPLATE_NAMES:
            names.append(path.name)
    return sorted(names)


def _ensure_safe_template_name(template_name: str, *, allow_standard: bool) -> str:
    normalized = str(template_name or "").strip()
    if not normalized:
        raise PlatformError(code="INVALID_REQUEST", message="Template name cannot be empty", status_code=400)
    if "/" in normalized or "\\" in normalized:
        raise PlatformError(code="INVALID_REQUEST", message="Invalid template name", status_code=400)
    if not (normalized.endswith(".yaml") or normalized.endswith(".yml") or normalized in _BUILTIN_TEMPLATE_NAMES):
        raise PlatformError(code="INVALID_REQUEST", message="Template name must end with .yaml", status_code=400)
    if not allow_standard and normalized in _BUILTIN_TEMPLATE_NAMES:
        raise PlatformError(
            code="INVALID_REQUEST",
            message=f"Cannot delete pre-configured template: {normalized}",
            status_code=400,
        )
    return normalized


def _read_image_file_payload(path: Path, *, fallback_name: str) -> BidGeneratorFilePayload:
    if not path.exists():
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="图片文件不存在", status_code=404)
    try:
        content = path.read_bytes()
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="读取图片文件失败。", status_code=500) from exc
    filename = path.name or fallback_name
    return BidGeneratorFilePayload(
        content=content,
        media_type=_image_media_type(path.suffix),
        filename=filename,
        inline=True,
        cache_control="public, max-age=86400",
    )


def _image_media_type(suffix: str) -> str:
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".bmp": "image/bmp",
    }.get(str(suffix or "").lower(), "application/octet-stream")


def _ensure_knowledge_image_tables(conn: Any) -> None:
    assets_exists = conn.execute(text("SELECT to_regclass('bid_generator.knowledge_image_assets') IS NOT NULL")).scalar_one()
    if not assets_exists:
        raise PlatformError(
            code="DATABASE_ERROR",
            message="知识库图片语义资产表不存在，请先执行数据库迁移。",
            status_code=500,
            details={"table": "bid_generator.knowledge_image_assets"},
        )


def _load_image_registry_preview_urls(conn: Any, image_hashes: list[str]) -> dict[str, str]:
    hashes = [str(item or "").lower() for item in image_hashes if str(item or "").strip()]
    if not hashes:
        return {}
    exists = conn.execute(text("SELECT to_regclass('bid_generator.image_registry') IS NOT NULL")).scalar_one()
    if not exists:
        return {}
    rows = conn.execute(
        text(
            """
            SELECT image_hash, preview_url
            FROM bid_generator.image_registry
            WHERE image_hash = ANY(:image_hashes)
            """
        ),
        {"image_hashes": hashes},
    ).mappings().all()
    return {str(row.get("image_hash") or "").lower(): str(row.get("preview_url") or "") for row in rows}


def _knowledge_image_asset_payload(row: Mapping[str, Any], preview_url: str = "") -> dict[str, Any]:
    return {
        "image_hash": row.get("image_hash"),
        "placeholder": row.get("placeholder"),
        "source_doc": row.get("source_doc"),
        "source_page": row.get("source_page"),
        "caption": row.get("caption"),
        "image_type": row.get("image_type"),
        "summary": row.get("summary"),
        "tags": _parse_json_tags(str(row.get("tags_json") or "[]")),
        "caption_status": row.get("caption_status"),
        "preview_url": preview_url,
        "created_at": _iso_value(row.get("created_at")),
    }


def _parse_json_object_text(value: str) -> dict[str, Any]:
    text_value = str(value or "").strip()
    if not text_value:
        return {}
    code_block_match = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", text_value, flags=re.IGNORECASE)
    raw = code_block_match.group(1).strip() if code_block_match else text_value
    object_match = re.search(r"\{[\s\S]*\}", raw)
    if object_match:
        raw = object_match.group(0)
    try:
        parsed = json.loads(raw)
    except json.JSONDecodeError:
        return {}
    return parsed if isinstance(parsed, dict) else {}


def _tag_image_with_vlm_native(image_path: Path) -> str:
    """调用远端多模态服务给图片打标；未启用时返回空字符串。"""
    try:
        if os.environ.get("REMOTE_VISION_ENABLED", "false").strip().lower() not in {"1", "true", "yes", "on"}:
            return ""
        vlm_api_url = os.environ.get("REMOTE_VISION_API_URL", "").strip()
        vlm_model = os.environ.get("REMOTE_VISION_MODEL", "").strip()
        if not vlm_api_url or not vlm_model:
            logger.warning("REMOTE_VISION_API_URL / REMOTE_VISION_MODEL 未配置，跳过图片打标")
            return ""
        timeout_seconds = int(os.environ.get("REMOTE_VISION_TIMEOUT", "120") or "120")
        max_tokens = int(os.environ.get("REMOTE_VISION_MAX_TOKENS", "400") or "400")
        api_key = os.environ.get("REMOTE_VISION_API_KEY", "").strip()
        encoded = base64.b64encode(image_path.read_bytes()).decode("utf-8")
        payload = {
            "model": vlm_model,
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "text",
                            "text": (
                                "你是知识库图片标注器。请识别图片内容并只输出 JSON，不要输出解释。\n"
                                "字段：caption(12-30字图注)、image_type(如系统架构图/流程图/截图/表格/其他)、"
                                "summary(120字以内说明)、key_elements(字符串数组)、tags(字符串数组)。"
                                "不得照抄图片中的真实人名、手机号、邮箱、IP、机构名称，需泛化为某人、某机构、[IP地址]。"
                            ),
                        },
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{encoded}"}},
                    ],
                }
            ],
            "max_tokens": max_tokens,
        }
        headers = {"Authorization": f"Bearer {api_key}"} if api_key else {}
        response = requests.post(vlm_api_url, json=payload, headers=headers, timeout=timeout_seconds)
        if response.status_code != 200:
            return ""
        content = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
        return str(content or "").strip()
    except Exception as exc:
        logger.warning("远端图片打标失败 (%s): %s", image_path, exc)
        return ""


def _normalize_image_caption_payload(raw_caption: str, fallback_caption: str = "知识库配图") -> dict[str, Any]:
    data = _parse_json_object_text(raw_caption)
    caption = str(data.get("caption") or "").strip() if data else ""
    image_type = str(data.get("image_type") or "").strip() if data else ""
    summary = str(data.get("summary") or "").strip() if data else ""
    tags = data.get("tags") if data else []
    if not isinstance(tags, list):
        tags = []
    if not caption and raw_caption and len(raw_caption.strip()) >= 4:
        caption = raw_caption.strip().splitlines()[0][:40]
        summary = raw_caption.strip()[:180]
    return {
        "caption": caption or fallback_caption,
        "image_type": image_type or "其他",
        "summary": summary or caption or fallback_caption,
        "tags": [str(item).strip() for item in tags if str(item).strip()][:8],
        "caption_status": "captioned" if caption and caption != fallback_caption else "weak",
    }


def _build_knowledge_image_block(asset: Mapping[str, Any]) -> str:
    tags = asset.get("tags") or []
    tag_text = "、".join(tags) if isinstance(tags, list) else str(tags or "")
    source = str(asset.get("source_doc") or "")
    page = asset.get("source_page")
    source_text = f"{source}，第 {page} 页" if source and page else source
    caption = str(asset.get("caption") or "知识库配图")
    placeholder = str(asset.get("placeholder") or "")
    return "\n".join(
        [
            "【知识库图片】",
            f"图片占位符：{placeholder}",
            f"图注：{caption}",
            f"类型：{asset.get('image_type', '其他')}",
            f"来源：{source_text}",
            f"说明：{asset.get('summary', '')}",
            f"标签：{tag_text}",
            f"使用规则：如正文需要引用该图，必须输出 ![图：{caption}]({placeholder})",
        ]
    ).strip()


def _ensure_image_asset_tables(conn: Any) -> None:
    registry_exists = conn.execute(text("SELECT to_regclass('bid_generator.image_registry') IS NOT NULL")).scalar_one()
    assets_exists = conn.execute(text("SELECT to_regclass('bid_generator.knowledge_image_assets') IS NOT NULL")).scalar_one()
    if not registry_exists or not assets_exists:
        raise PlatformError(
            code="DATABASE_ERROR",
            message="知识库图片资产表不存在，请先执行数据库迁移。",
            status_code=500,
            details={"tables": ["bid_generator.image_registry", "bid_generator.knowledge_image_assets"]},
        )


def _register_knowledge_image_asset_native(
    *,
    filename: str,
    image_bytes: bytes,
    original_name: str,
    fallback_caption: str = "知识库配图",
    source_page: int | None = None,
) -> tuple[str, str, dict[str, Any]]:
    """持久化图片资产；入参为图片字节和来源信息，出参为占位符、图注和 image_map 条目。"""
    image_hash = hashlib.md5(image_bytes).hexdigest()
    safe_original_name = Path(original_name or "image.bin").name or "image.bin"
    stored_name = f"{image_hash}_{safe_original_name}"
    image_dir = _bid_generator_legacy_root() / "data" / "extracted_images"
    image_path = image_dir / stored_name
    try:
        image_dir.mkdir(parents=True, exist_ok=True)
        if not image_path.exists():
            image_path.write_bytes(image_bytes)
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="保存提取图片失败。", status_code=500) from exc

    placeholder = f"__PRO_IMG_{image_hash}__"
    preview_url = f"/api/extracted-images/{stored_name}"
    caption_payload = _normalize_image_caption_payload(
        _tag_image_with_vlm_native(image_path),
        fallback_caption=fallback_caption,
    )
    caption = str(caption_payload["caption"])
    try:
        with get_engine().begin() as conn:
            _ensure_image_asset_tables(conn)
            registry_exists = conn.execute(
                text("SELECT 1 FROM bid_generator.image_registry WHERE image_hash = :image_hash"),
                {"image_hash": image_hash},
            ).first()
            if registry_exists is None:
                conn.execute(
                    text(
                        """
                        INSERT INTO bid_generator.image_registry
                          (image_hash, project_id, abs_path, preview_url, placeholder, vlm_caption, is_reference_only)
                        VALUES
                          (:image_hash, NULL, :abs_path, :preview_url, :placeholder, :vlm_caption, 1)
                        """
                    ),
                    {
                        "image_hash": image_hash,
                        "abs_path": str(image_path),
                        "preview_url": preview_url,
                        "placeholder": placeholder,
                        "vlm_caption": caption,
                    },
                )
            asset_exists = conn.execute(
                text("SELECT 1 FROM bid_generator.knowledge_image_assets WHERE image_hash = :image_hash"),
                {"image_hash": image_hash},
            ).first()
            if asset_exists is None:
                conn.execute(
                    text(
                        """
                        INSERT INTO bid_generator.knowledge_image_assets
                          (image_hash, placeholder, source_doc, source_page, nearby_text_sanitized,
                           caption, image_type, summary, tags_json, caption_status)
                        VALUES
                          (:image_hash, :placeholder, :source_doc, :source_page, '',
                           :caption, :image_type, :summary, :tags_json, :caption_status)
                        """
                    ),
                    {
                        "image_hash": image_hash,
                        "placeholder": placeholder,
                        "source_doc": filename or "",
                        "source_page": source_page,
                        "caption": caption,
                        "image_type": str(caption_payload["image_type"]),
                        "summary": str(caption_payload["summary"]),
                        "tags_json": json.dumps(caption_payload["tags"], ensure_ascii=False),
                        "caption_status": str(caption_payload["caption_status"]),
                    },
                )
    except PlatformError:
        raise
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc

    image_info = {
        "abs_path": str(image_path),
        "preview_url": preview_url,
        "description": caption,
        "knowledge_block": _build_knowledge_image_block(
            {
                "placeholder": placeholder,
                "caption": caption,
                "image_type": caption_payload["image_type"],
                "summary": caption_payload["summary"],
                "tags": caption_payload["tags"],
                "source_doc": filename or "",
                "source_page": source_page,
            }
        ),
        "caption_status": caption_payload["caption_status"],
    }
    return placeholder, caption, image_info


def _knowledge_document_payload(row: Mapping[str, Any]) -> dict[str, Any]:
    created_at = row.get("created_at")
    upload_time = "-"
    if created_at:
        try:
            upload_time = datetime.fromtimestamp(float(created_at)).strftime("%Y-%m-%d %H:%M")
        except (TypeError, ValueError, OSError):
            upload_time = "-"
    word_count = _non_negative_int(row.get("word_count"))
    tokens = _non_negative_int(row.get("tokens")) if "tokens" in row else 0
    chunks = tokens if tokens > 0 else word_count // 500
    return {
        "id": str(row.get("id") or ""),
        "name": str(row.get("name") or ""),
        "size": _estimated_size_from_word_count(word_count),
        "uploadTime": upload_time,
        "status": _knowledge_document_status(row.get("indexing_status")),
        "chunks": chunks,
    }


def _knowledge_document_status(value: Any) -> str:
    raw_status = str(value or "completed")
    if raw_status == "completed":
        return "success"
    if raw_status == "error":
        return "failed"
    return "indexing"


def _estimated_size_from_word_count(word_count: int) -> str:
    size_bytes = max(0, int(word_count or 0)) * 2
    if size_bytes > 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.1f} MB"
    return f"{size_bytes / 1024:.1f} KB"


def _non_negative_int(value: Any) -> int:
    try:
        return max(0, int(value or 0))
    except (TypeError, ValueError):
        return 0


def _number_value(value: Any) -> int | float:
    try:
        number = float(value or 0)
    except (TypeError, ValueError):
        return 0
    return int(number) if number.is_integer() else number


async def _call_dify_workflow(api_key: str, inputs: Mapping[str, Any], max_retries: int = 2) -> dict[str, Any]:
    dify_base = os.environ.get("DIFY_API_URL", "http://localhost/v1").rstrip("/")
    dify_url = f"{dify_base}/workflows/run"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "inputs": dict(inputs),
        "response_mode": "blocking",
        "user": "pro-engine-backend",
    }
    last_err: Exception = RuntimeError("未知错误")
    for attempt in range(max(1, int(max_retries or 1))):
        try:
            async with httpx.AsyncClient(timeout=1800) as client:
                response = await client.post(dify_url, headers=headers, json=payload)
                response.raise_for_status()
                return response.json()
        except Exception as exc:
            last_err = exc
            if attempt < max_retries - 1:
                wait_seconds = 5 * (attempt + 1)
                logger.warning("[Dify blocking] 第 %s 次调用失败，%ss 后重试: %s", attempt + 1, wait_seconds, exc)
                await asyncio.sleep(wait_seconds)
    raise last_err


async def _call_dify_workflow_stream(api_key: str, inputs: Mapping[str, Any]) -> Any:
    dify_base = os.environ.get("DIFY_API_URL", "http://localhost/v1").rstrip("/")
    dify_url = f"{dify_base}/workflows/run"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "inputs": dict(inputs),
        "response_mode": "streaming",
        "user": "pro-engine-backend",
    }
    loop_round = 0
    dify_task_id = ""
    workflow_run_id = ""
    async with httpx.AsyncClient(timeout=1800) as client:
        async with client.stream("POST", dify_url, headers=headers, json=payload) as response:
            if response.is_error:
                await response.aread()
            response.raise_for_status()
            buffer = ""
            async for chunk in response.aiter_text():
                buffer += chunk
                while "\n\n" in buffer:
                    event_str, buffer = buffer.split("\n\n", 1)
                    for line in event_str.strip().split("\n"):
                        if not line.startswith("data: "):
                            continue
                        data_str = line[6:]
                        try:
                            data = json.loads(data_str)
                        except json.JSONDecodeError:
                            continue
                        event_type = str(data.get("event") or "")
                        if not dify_task_id:
                            dify_task_id = str(data.get("task_id") or "")
                        if not workflow_run_id:
                            workflow_run_id = (
                                str(data.get("workflow_run_id") or "")
                                or str(data.get("data", {}).get("workflow_run_id") or "")
                                or str(data.get("data", {}).get("id") or "")
                            )
                        if event_type == "text_chunk":
                            text_value = str(data.get("data", {}).get("text") or "")
                            if text_value:
                                yield text_value
                            continue
                        if event_type == "node_started":
                            node_title = str(data.get("data", {}).get("title") or "")
                            if "Reviewer" in node_title or "审查" in node_title:
                                loop_round += 1
                            stage_label = _DIFY_NODE_STAGE_MAP.get(node_title)
                            if stage_label:
                                if loop_round > 0 and ("审查" in stage_label or "润色" in stage_label or "评分" in stage_label):
                                    stage_label = f"{stage_label} ({loop_round}/3)"
                                yield {
                                    "__stage__": stage_label,
                                    "node_title": node_title,
                                    "node_id": str(data.get("data", {}).get("node_id") or ""),
                                    "dify_task_id": dify_task_id,
                                    "workflow_run_id": workflow_run_id,
                                }
                            continue
                        if event_type == "workflow_finished":
                            run_id = (
                                str(data.get("workflow_run_id") or "")
                                or str(data.get("data", {}).get("workflow_run_id") or "")
                                or str(data.get("data", {}).get("id") or "")
                            )
                            if run_id:
                                workflow_run_id = run_id
                            yield {
                                "__finished__": True,
                                "outputs": data.get("data", {}).get("outputs", {}),
                                "workflow_run_id": workflow_run_id,
                                "dify_task_id": dify_task_id,
                            }


async def _get_dify_workflow_run_result(api_key: str, workflow_run_id: str) -> dict[str, Any]:
    normalized_run_id = str(workflow_run_id or "").strip()
    if not normalized_run_id:
        return {}
    dify_base = os.environ.get("DIFY_API_URL", "http://localhost/v1").rstrip("/")
    headers = {"Authorization": f"Bearer {api_key}"}
    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.get(f"{dify_base}/workflows/run/{normalized_run_id}", headers=headers)
        response.raise_for_status()
        payload = response.json()
    return payload if isinstance(payload, dict) else {}


def _resolve_outline_sections_from_outputs(
    outputs: Any,
    *,
    seed_headings: list[dict],
    max_diagrams: int,
) -> list[dict]:
    structured_data = parse_dify_outputs({"data": {"outputs": outputs}}) if outputs else {}
    sections_raw = extract_outline_sections_raw(structured_data)
    if not sections_raw and isinstance(outputs, dict):
        for value in outputs.values():
            candidate = value
            if isinstance(candidate, str):
                candidate = candidate.strip()
                if candidate.startswith("```"):
                    candidate = candidate.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
                try:
                    candidate = json.loads(candidate)
                except Exception:
                    pass
            if isinstance(candidate, list):
                sections_raw = candidate
                break
            if isinstance(candidate, dict):
                nested = candidate.get("outline") or candidate.get("sections")
                if nested:
                    sections_raw = nested
                    break
    return build_seeded_outline_sections(sections_raw, seed_headings, max_diagrams=max_diagrams)


async def _stream_native_outline_generation(
    *,
    dify_key: str,
    inputs: Mapping[str, Any],
    seed_headings: list[dict],
    max_diagrams: int,
    expected_total_words: int,
) -> Any:
    workflow_run_id = ""
    sections: list[dict] = []
    used_fallback = False
    async for chunk in _call_dify_workflow_stream(dify_key, inputs):
        if not isinstance(chunk, dict):
            continue
        if chunk.get("__stage__"):
            workflow_run_id = str(chunk.get("workflow_run_id") or workflow_run_id or "")
            yield {"stage": str(chunk.get("__stage__") or "")}
            continue
        if chunk.get("__finished__"):
            workflow_run_id = str(chunk.get("workflow_run_id") or workflow_run_id or "")
            sections = _resolve_outline_sections_from_outputs(
                chunk.get("outputs", {}),
                seed_headings=seed_headings,
                max_diagrams=max_diagrams,
            )
            break

    quality_report = evaluate_outline_quality(sections, seed_headings)
    if workflow_run_id and (not sections or not quality_report["pass"]):
        fallback_payload = await _get_dify_workflow_run_result(dify_key, workflow_run_id)
        sections = _resolve_outline_sections_from_outputs(
            fallback_payload.get("data", {}).get("outputs", {}) if isinstance(fallback_payload, dict) else {},
            seed_headings=seed_headings,
            max_diagrams=max_diagrams,
        )
        used_fallback = True

    quality_report = evaluate_outline_quality(sections, seed_headings)
    if not quality_report["pass"]:
        logger.error(
            "[generate_outline_stream] 结构校验失败: fallback_used=%s report=%s",
            used_fallback,
            quality_report,
        )
        raise PlatformError(
            code="OUTLINE_GENERATE_STREAM_FAILED",
            message="大纲生成结构不完整，请重试：" + "；".join(quality_report.get("issues") or []),
            status_code=502,
        )

    normalize_outline_word_budget_dict(sections, expected_total_words)
    yield {"done": True, "sections": sections}


async def _extract_docanalysis_group_results(
    *,
    system_prompt_base: str,
    dify_key: str,
    subset_nodes: list[dict],
    subset_label: str,
    document_text: str,
) -> list[dict[str, str]]:
    combined_system = build_docanalysis_system_prompt(system_prompt_base, subset_nodes, subset_label)
    raw_text = ""
    try:
        dify_res = await _call_dify_workflow(
            dify_key,
            {
                "system_prompt": combined_system,
                "raw_document": document_text,
                "node_label": subset_label,
            },
        )
        outputs = dify_res.get("data", {}).get("outputs", {}) if isinstance(dify_res, dict) else {}
        raw_text = extract_docanalysis_text_output(outputs)
        raw_text, _attachments_payload = split_bid_attachments_tag(raw_text)
        result_map = parse_docanalysis_result_map(raw_text)
        results: list[dict[str, str]] = []
        for node in subset_nodes:
            content = extract_docanalysis_node_content(result_map, str(node.get("id") or ""))
            if isinstance(content, (dict, list)):
                content = json.dumps(content, ensure_ascii=False, indent=2)
            results.append(
                {
                    "node_id": str(node.get("id") or ""),
                    "label": str(node.get("label") or ""),
                    "content": str(content),
                }
            )
        return results
    except Exception as exc:
        logger.warning("分组 [%s] 结果解析失败，降级逐节点提取: %s", subset_label, exc)
        if raw_text:
            return [{"node_id": str(subset_nodes[0].get("id") or ""), "label": str(subset_nodes[0].get("label") or ""), "content": raw_text}]
        return [
            {
                "node_id": str(node.get("id") or ""),
                "label": str(node.get("label") or ""),
                "content": "**提取失败，请重新生成**",
            }
            for node in subset_nodes
        ]


def _format_dify_runtime_error(exc: Exception) -> str:
    message = str(exc or "").strip()
    lower = message.lower()
    if "dashscope.aliyuncs.com" in lower and ("nameresolutionerror" in lower or "failed to resolve" in lower):
        return (
            "Dify 模型供应商 DashScope DNS 解析失败：dashscope.aliyuncs.com 无法解析。"
            "请在 Dify API/Worker 运行环境检查 DNS、代理或出网策略；标书后端已成功调用 Dify，但模型节点不可用。"
        )
    if "[models]" in lower and "server unavailable" in lower:
        return (
            "Dify 模型节点不可用（[models] Server Unavailable）。"
            "请检查 Dify 模型供应商配置、API Key、DNS/代理与出网策略。"
            + (f" 原始错误：{message}" if message else "")
        )
    if "name or service not known" in lower or "failed to resolve" in lower:
        return "Dify 或其上游服务域名解析失败，请检查运行环境 DNS / 代理 / 出网策略。" + (f" 原始错误：{message}" if message else "")
    return message or "Dify 工作流调用失败。"


def _parse_json_tags(tags_json: str) -> list[str]:
    try:
        loaded = json.loads(tags_json or "[]")
    except json.JSONDecodeError:
        return []
    if not isinstance(loaded, list):
        return []
    return [str(item).strip() for item in loaded if str(item).strip()]


def _ensure_safe_project_id(project_id: str) -> str:
    normalized = _required_string(project_id, field="project_id")
    if not re.fullmatch(r"[a-zA-Z0-9_-]+", normalized):
        raise PlatformError(code="INVALID_REQUEST", message="无效的 project_id。", status_code=400)
    return normalized


def _ensure_safe_image_hash(image_hash: str) -> str:
    normalized = str(image_hash or "").strip().lower()
    if not re.fullmatch(r"[a-f0-9]+", normalized):
        raise PlatformError(code="INVALID_REQUEST", message="无效的散列格式", status_code=400)
    return normalized


def _ensure_safe_diagram_artifact_id(diagram_id: str) -> str:
    normalized = str(diagram_id or "").strip()
    if not re.fullmatch(r"[a-fA-F0-9]{16,64}", normalized):
        raise PlatformError(code="INVALID_REQUEST", message="无效的图表 ID", status_code=400)
    return normalized.lower()


def _require_task_owner(task_id: str, project_id: str | None) -> Any:
    task_id_value = _required_string(task_id, field="task_id")
    task = _get_task(task_id_value)
    if not task:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="任务不存在或已过期", status_code=404)
    pid = str(project_id or "").strip()
    task_project_id = str(getattr(task, "project_id", "") or "").strip()
    if pid and task_project_id and task_project_id != pid:
        raise PlatformError(code="PERMISSION_DENIED", message="任务不属于当前项目", status_code=403)
    return task


def _get_task(task_id: str) -> Any | None:
    return _task_manager().get_task(task_id)


def _task_manager() -> Any:
    return native_task_manager


async def _ensure_project_slot_native(project_id: str, task_type: str) -> None:
    """检查项目任务并发槽位；入参为项目 ID/任务类型，不再调用 legacy task_routes 私有函数。"""
    project_id_value = str(project_id or "").strip()
    if not project_id_value:
        raise PlatformError(code="INVALID_REQUEST", message="project_id 不能为空", status_code=400)
    task_manager = _task_manager()
    try:
        task_manager.ensure_backend_ready()
    except RuntimeError as exc:
        raise PlatformError(
            code="TASK_BACKEND_UNAVAILABLE",
            message=str(exc),
            status_code=500,
            details={"code": "TASK_BACKEND_UNAVAILABLE", "message": str(exc)},
        ) from exc

    limits = task_manager.get_limits()
    content_project_limit = int(limits.get("max_project_content_running", 2) or 2)
    project_limit_override = content_project_limit if task_type == "content" else None
    type_limit_override = int(os.environ.get("MAX_DIAGRAM_RUNNING_TASKS", "1") or "1") if task_type == "diagram" else None
    allowed, details = await task_manager.try_acquire_task_slot(
        project_id_value,
        task_type,
        enforce_project_limit=(task_type != "diagram"),
        max_project_running=project_limit_override,
        max_type_running=type_limit_override,
    )
    if allowed:
        return

    reason = (details or {}).get("reason", "limit")
    if reason == "global_limit":
        message = "后台任务并发达到全局上限，请稍后重试"
    elif reason == "project_limit":
        limit_num = (details or {}).get("max_project_running")
        if task_type == "content" and limit_num:
            message = f"项目 {project_id_value} 正在运行 {limit_num} 个正文任务，请等待空闲后再发起"
        else:
            message = f"项目 {project_id_value} 正在运行任务，请等待当前任务完成后再发起"
    else:
        message = "任务并发受限，请稍后重试"
    detail = {
        "code": "TASK_LIMIT_REACHED",
        "message": message,
        "limit_reason": reason,
        "requested_project_id": project_id_value,
        "task_type": task_type,
        "limits": task_manager.get_limits(),
        "metrics": details,
    }
    raise PlatformError(code="TASK_LIMIT_REACHED", message=message, status_code=409, details=detail)


def _cache_pdf_file_native(project_id: str, content_bytes: bytes) -> str:
    """缓存上传 PDF；入参为项目 ID 和 PDF 字节，出参保持 legacy 预览 URL。"""
    normalized_id = _ensure_safe_project_id(project_id)
    pdf_path = _pdf_cache_path(normalized_id)
    try:
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        pdf_path.write_bytes(content_bytes)
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="缓存 PDF 文件失败。", status_code=500) from exc
    return f"/api/projects/pdf/{normalized_id}"


def _extract_pdf_pages_text_native(content_bytes: bytes) -> list[dict[str, Any]]:
    """按页提取 PDF 文本；入参为 PDF 字节，出参为页码和文本列表。"""
    pages_text: list[dict[str, Any]] = []
    try:
        import pymupdf

        doc = pymupdf.open(stream=content_bytes, filetype="pdf")
        try:
            for page_idx, page in enumerate(doc):
                pages_text.append({"page": page_idx, "text": page.get_text("text") or ""})
        finally:
            doc.close()
    except ImportError:
        logger.warning("PyMuPDF 未安装，无法进行分页文本索引")
    except Exception as exc:
        logger.warning("PDF 分页文本提取异常: %s", exc)
    return pages_text


def _preprocess_docx_alignment_native(src_path: str) -> str:
    """预处理 DOCX 对齐属性，降低 LibreOffice 转 PDF 时的字距异常概率。"""
    try:
        import docx as docx_module
    except ImportError:
        return src_path

    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    bad_alignments = {
        "distribute",
        "distributeLetter",
        "distributeAllLines",
        "thaiDistribute",
        "both",
        "justify",
        "lowKashida",
        "mediumKashida",
        "highKashida",
    }
    document = docx_module.Document(src_path)

    def fix_paragraphs(paragraphs: Any) -> None:
        for paragraph in paragraphs:
            paragraph_props = paragraph._element.find(f"{{{word_namespace}}}pPr")
            if paragraph_props is None:
                continue
            justification = paragraph_props.find(f"{{{word_namespace}}}jc")
            if justification is None:
                continue
            value = str(justification.get(f"{{{word_namespace}}}val", "") or "")
            if value in bad_alignments:
                justification.set(f"{{{word_namespace}}}val", "left")

    fix_paragraphs(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                fix_paragraphs(cell.paragraphs)
    document.save(src_path)
    return src_path


def _convert_to_pdf_and_cache_native(project_id: str, content_bytes: bytes, filename: str) -> str:
    """将 DOC/DOCX 转为缓存 PDF；入参为项目 ID、原文件字节和文件名，出参保持 legacy 预览 URL。"""
    normalized_id = _ensure_safe_project_id(project_id)
    pdf_path = _pdf_cache_path(normalized_id)
    try:
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="创建 PDF 缓存目录失败。", status_code=500) from exc

    ext = Path(filename or "").suffix.lower().lstrip(".") or "docx"
    with tempfile.TemporaryDirectory() as tmp_dir:
        src_path = Path(tmp_dir) / f"source.{ext}"
        src_path.write_bytes(content_bytes)
        if ext == "docx":
            try:
                _preprocess_docx_alignment_native(str(src_path))
            except Exception as exc:
                logger.warning("DOCX 排版预处理失败，使用原文件: %s", exc)

        try:
            import docx2pdf

            docx2pdf.convert(str(src_path), str(pdf_path))
            if pdf_path.exists() and pdf_path.stat().st_size > 0:
                return f"/api/projects/pdf/{normalized_id}"
        except Exception as exc:
            logger.debug("docx2pdf 不可用: %s", exc)

        try:
            env = os.environ.copy()
            env["SAL_USE_VCLPLUGIN"] = "svp"
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--norestore",
                    "--nofirststartwizard",
                    "--convert-to",
                    "pdf:writer_pdf_Export",
                    "--outdir",
                    tmp_dir,
                    str(src_path),
                ],
                capture_output=True,
                timeout=120,
                env=env,
                check=False,
            )
            converted_pdf = Path(tmp_dir) / "source.pdf"
            if result.returncode == 0 and converted_pdf.exists():
                shutil.copy(str(converted_pdf), str(pdf_path))
                return f"/api/projects/pdf/{normalized_id}"
            stderr_text = result.stderr.decode("utf-8", errors="replace")[:200]
            logger.warning("LibreOffice 转换失败: %s", stderr_text)
        except FileNotFoundError:
            logger.warning("LibreOffice 未安装，DOCX/DOC 转 PDF 不可用")
        except Exception as exc:
            logger.warning("LibreOffice 转换异常: %s", exc)

    logger.warning("所有 DOC/DOCX 转 PDF 方案均失败，project_id=%s", normalized_id)
    return ""


def _docx_blocks_to_locator_text(doc_blocks: list[dict[str, Any]]) -> tuple[str, dict[str, int]]:
    locator_map: dict[str, int] = {}
    lines: list[str] = []
    for block in doc_blocks:
        locator = str(block.get("locator") or "").strip().upper()
        text_value = str(block.get("text") or "").strip()
        if not locator or not text_value:
            continue
        try:
            locator_map[locator] = int(block.get("body_idx") or 0)
        except (TypeError, ValueError):
            locator_map[locator] = 0
        lines.append(f"[{locator}] {text_value}")
    return "\n".join(lines), locator_map


def _extract_docx_with_tables_native(content_bytes: bytes, *, filename: str = "", extract_images: bool = False) -> tuple[str, dict[str, Any]]:
    """按 DOCX body 顺序提取段落/表格/图片；入参为 DOCX 字节，出参为文本和图片映射。"""
    if not extract_images:
        blocks = _extract_docx_blocks(content_bytes)
        return "\n".join(str(block.get("text") or "") for block in blocks if str(block.get("text") or "").strip()), {}

    try:
        import docx as docx_module
    except ImportError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="统一后端缺少 python-docx 依赖。", status_code=500) from exc

    document = docx_module.Document(io.BytesIO(content_bytes))
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    paragraph_tag = f"{{{word_namespace}}}p"
    table_tag = f"{{{word_namespace}}}tbl"
    row_tag = f"{{{word_namespace}}}tr"
    cell_tag = f"{{{word_namespace}}}tc"
    drawing_tag = f"{{{word_namespace}}}drawing"
    drawing_namespace = "http://schemas.openxmlformats.org/drawingml/2006/main"
    blip_tag = f"{{{drawing_namespace}}}blip"
    relationship_namespace = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    embed_attr = f"{{{relationship_namespace}}}embed"
    image_map: dict[str, Any] = {}

    def cell_text(cell_elem: Any) -> str:
        texts: list[str] = []
        for paragraph in cell_elem.findall(f".//{paragraph_tag}"):
            text_value = "".join(node.text or "" for node in paragraph.iter(f"{{{word_namespace}}}t"))
            if text_value.strip():
                texts.append(text_value.strip())
        return " ".join(texts)

    def table_to_markdown(table_elem: Any) -> str:
        rows = table_elem.findall(f".//{row_tag}")
        if not rows:
            return ""
        markdown_rows: list[str] = []
        for row in rows:
            cells = row.findall(f".//{cell_tag}")
            markdown_rows.append("| " + " | ".join(cell_text(cell) for cell in cells) + " |")
        if len(markdown_rows) > 1:
            column_count = markdown_rows[0].count("|") - 1
            markdown_rows.insert(1, "| " + " | ".join(["---"] * max(column_count, 1)) + " |")
        return "\n".join(markdown_rows)

    def image_refs_from_element(element: Any) -> list[str]:
        refs: list[str] = []
        for drawing in element.findall(f".//{drawing_tag}"):
            for blip in drawing.findall(f".//{blip_tag}"):
                embed_id = blip.get(embed_attr)
                if not embed_id or embed_id not in document.part.rels:
                    continue
                rel = document.part.rels[embed_id]
                if "image" not in str(rel.reltype or ""):
                    continue
                image_bytes = rel.target_part.blob
                original_name = Path(str(rel.target_part.partname).split("/")[-1]).name
                placeholder, caption, image_info = _register_knowledge_image_asset_native(
                    filename=filename,
                    image_bytes=image_bytes,
                    original_name=original_name,
                )
                image_map[placeholder] = image_info
                refs.append(f"![{caption}]({placeholder})")
        return refs

    parts: list[str] = []
    for child in document.element.body:
        if child.tag == paragraph_tag:
            text_value = "".join(node.text or "" for node in child.iter(f"{{{word_namespace}}}t"))
            if text_value.strip():
                parts.append(text_value)
            parts.extend(image_refs_from_element(child))
        elif child.tag == table_tag:
            table_text = table_to_markdown(child)
            if table_text:
                parts.append(f"\n{table_text}\n")
            parts.extend(image_refs_from_element(child))
    return "\n".join(parts), image_map


def _extract_docx_with_locators_native(content_bytes: bytes) -> tuple[str, dict[str, Any], list[dict[str, Any]]]:
    """提取 DOCX 定位符文本；入参为 DOCX 字节，出参为定位文本、定位映射和块结构。"""
    blocks = _extract_docx_blocks(content_bytes)
    locator_text, locator_map = _docx_blocks_to_locator_text(blocks)
    return locator_text, locator_map, blocks


def _extract_pdf_text_native(content_bytes: bytes) -> str:
    try:
        import pymupdf

        doc = pymupdf.open(stream=content_bytes, filetype="pdf")
        try:
            return "\n".join(page.get_text("text") or "" for page in doc)
        finally:
            doc.close()
    except ImportError:
        logger.warning("PyMuPDF 未安装，PDF 文本提取降级为字节解码")
    except Exception as exc:
        logger.warning("PyMuPDF 解析异常，PDF 文本提取降级为字节解码: %s", exc)
    return content_bytes.decode("latin-1", errors="replace")


def _extract_pdf_text_with_images_native(filename: str, content_bytes: bytes) -> tuple[str, dict[str, Any]]:
    image_map: dict[str, Any] = {}
    try:
        import pymupdf4llm

        with tempfile.TemporaryDirectory() as temp_dir:
            pdf_path = Path(temp_dir) / "source.pdf"
            images_dir = Path(temp_dir) / "images"
            images_dir.mkdir(parents=True, exist_ok=True)
            pdf_path.write_bytes(content_bytes)
            markdown_text = pymupdf4llm.to_markdown(str(pdf_path), write_images=True, image_path=str(images_dir))
            min_size = int(os.environ.get("VLM_MIN_IMAGE_SIZE_KB", "5") or "5") * 1024

            def replace_image(match: re.Match[str]) -> str:
                source_img = Path(temp_dir) / match.group(2)
                try:
                    if not source_img.exists() or source_img.stat().st_size < min_size:
                        return ""
                    placeholder, caption, image_info = _register_knowledge_image_asset_native(
                        filename=filename,
                        image_bytes=source_img.read_bytes(),
                        original_name=source_img.name,
                    )
                    image_map[placeholder] = image_info
                    return f"![{caption}]({placeholder})"
                except Exception as exc:
                    logger.warning("PDF 图片资产注册失败: %s", exc)
                    return ""

            return re.sub(r"!\[(.*?)\]\((.*?)\)", replace_image, markdown_text), image_map
    except ImportError:
        logger.warning("服务器缺少 pymupdf4llm，PDF 视觉解析降级为基础文本解析")
    except Exception as exc:
        logger.warning("PyMuPDF4LLM 解析异常，PDF 视觉解析降级为基础文本解析: %s", exc)
    return _extract_pdf_text_native(content_bytes), image_map


def _extract_doc_text_native(content_bytes: bytes) -> str:
    if content_bytes[:4] == b"PK\x03\x04":
        text_value, _image_map = _extract_docx_with_tables_native(content_bytes)
        return text_value
    if content_bytes[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":
        try:
            import olefile

            with olefile.OleFileIO(io.BytesIO(content_bytes)) as ole:
                if ole.exists("WordDocument"):
                    raw = ole.openstream("WordDocument").read()
                    text_value = raw.decode("utf-16-le", errors="replace")
                    text_value = re.sub(r"[\x00-\x08\x0b-\x1f\x7f]", " ", text_value)
                    text_value = re.sub(r" {3,}", "\n", text_value).strip()
                    if len(text_value) > 50:
                        return text_value
        except ImportError:
            logger.warning("olefile 未安装，无法解析旧版 .doc 文件")
        except Exception as exc:
            logger.warning(".doc OLE2 解析失败: %s", exc)
    return "[无法解析旧版 .doc 文件，请将文件另存为 .docx 后重新上传]"


def _extract_raw_text_with_images_native(filename: str, content_bytes: bytes, *, use_vision_parsing: bool) -> tuple[str, dict[str, Any]]:
    """提取上传文档文本；入参为文件名/字节/视觉开关，出参为文本和图片映射。"""
    ext = Path(filename or "").suffix.lower().lstrip(".")
    if ext == "pdf":
        if use_vision_parsing:
            return _extract_pdf_text_with_images_native(filename, content_bytes)
        return _extract_pdf_text_native(content_bytes), {}
    if ext == "docx":
        return _extract_docx_with_tables_native(content_bytes, filename=filename, extract_images=use_vision_parsing)
    if ext == "doc":
        return _extract_doc_text_native(content_bytes), {}
    try:
        return content_bytes.decode("utf-8"), {}
    except UnicodeDecodeError:
        return content_bytes.decode("latin-1", errors="replace"), {}


def _load_desensitize_profile(profile_name: str) -> dict[str, Any]:
    config = _read_yaml_mapping(_bid_generator_config_path())
    pipt = config.get("pipt") if isinstance(config, dict) else {}
    profiles = pipt.get("profiles") if isinstance(pipt, dict) else {}
    if not isinstance(profiles, dict):
        return {}
    profile = profiles.get(profile_name)
    if isinstance(profile, dict):
        return profile
    default_profile = profiles.get("default")
    return default_profile if isinstance(default_profile, dict) else {}


def _run_bid_pipt_preprocess(*, text: str, project_id: str, task_id: str, profile_name: str) -> dict[str, Any]:
    profile = _load_desensitize_profile(profile_name)
    try:
        from app.services.pipt_config_service import get_module_pipt_runtime_config

        runtime_config = get_module_pipt_runtime_config("bid-generator")
        enabled = bool(runtime_config.get("enabled"))
        target_entities = runtime_config.get("target_entities")
    except Exception:
        logger.exception("Failed to load bid generator PIPT config, falling back to legacy profile.")
        enabled = True
        target_entities = profile.get("target_entities")
    if not isinstance(target_entities, list):
        target_entities = ["name", "phone", "email", "id_number"]
    method = str(profile.get("method") or "placeholder").strip().lower()
    if method != "placeholder":
        method = "placeholder"
    result = preprocess_internal_payload(
        {
            "text": text,
            "module_code": "bid-generator",
            "purpose": "document_preprocess",
            "mode": "strong",
            "enabled": enabled,
            "request_id": f"{project_id}:{task_id}",
            "target_entities": target_entities,
            "llm_mode": os.environ.get("PIPT_LLM_MODE_EXTRACT", "verify_only"),
        }
    )
    if method == "mask":
        mapping_table = _string_mapping(result.get("mapping_table"))
        masked_text = str(result.get("text") or "")
        for token, original in mapping_table.items():
            if token in masked_text:
                masked_text = masked_text.replace(token, "*" * len(original))
        result = dict(result)
        result["text"] = masked_text
    return result


def _compose_runtime_writing_hint(
    writing_hint: str,
    section_title: str,
    expected_words: int,
    keywords: str,
    *,
    section_outline_slice: str = "",
    analysis_context: str = "",
) -> str:
    parts: list[str] = []
    outline_block = _build_outline_slice_block(section_outline_slice)
    if outline_block:
        parts.append(outline_block)

    bridge_block = _build_section_bridge_block(section_title, section_outline_slice)
    if bridge_block:
        parts.append(bridge_block)

    parts.append(_build_formal_response_style_block())

    core = _extract_core_writing_intent(writing_hint)
    if core:
        parts.append(core)

    analysis_block = _build_analysis_context_block(analysis_context)
    if analysis_block:
        parts.append(analysis_block)

    parts.append(_build_content_expansion_constraints(section_title, int(expected_words or 0), keywords))
    return _join_hint_segments(parts)


_RUNTIME_HINT_BLOCK_TITLES: tuple[str, ...] = (
    "【本节目录层级定位（勿用 # 标题重复以下编号）】",
    "【本节目录层级定位（只用于理解，不得输出）】",
    "【章内承接与开篇导入要求】",
    "【正式应答文件行文范式】",
    "【招标文件解析参考（优先级最高，严格对应本章节要求）】",
    "【正文扩写与技术深度约束（必须遵守）】",
)

_IMPLICIT_HINT_TAIL_ANCHORS: tuple[str, ...] = (
    "正文应按“需求理解、方案机制、落地措施、验证与风险控制”展开",
    "不要重复目录编号",
    "不得编造缺乏依据",
)


def _normalize_hint_text(text_value: str) -> str:
    return str(text_value or "").replace("\r\n", "\n").strip()


def _join_hint_segments(segments: list[str]) -> str:
    return "\n\n".join(part.strip() for part in segments if part and part.strip()).strip()


def _normalize_response_writing_intent(text_value: str) -> str:
    """将旧大纲里的评分说明式提示归一为正式应答文件写作意图。"""
    normalized = _normalize_hint_text(text_value)
    if not normalized:
        return ""
    normalized = re.sub(r"[“\"]([^”\"]+?)（\s*\d+\s*分\s*）[”\"]", r"“\1”", normalized)
    normalized = re.sub(r"（\s*满分\s*\d+\s*分\s*）", "", normalized)
    normalized = re.sub(r"（\s*\d+\s*分\s*）", "", normalized)
    normalized = re.sub(
        r"对标评分标准“([^”]+)”[，,]?",
        r"围绕“\1”形成正式技术应答，",
        normalized,
    )
    normalized = re.sub(
        r"当前已识别的核心侧重点是：本章对应“([^”]+)”\s*\d*\s*分?评分项。?",
        r"当前写作侧重点：围绕“\1”形成正式应答。",
        normalized,
    )
    normalized = re.sub(
        r"当前已识别的核心侧重点是：本章综合对应[^。]*。",
        "当前写作侧重点：综合承接项目理解、重点难点、质量保障与技术能力要求。",
        normalized,
    )
    normalized = normalized.replace("当前已识别的核心侧重点是：", "当前写作侧重点：")
    normalized = re.sub(
        r"围绕“([^”]+)”撰写本节内容，([^。]*?)，先说明本节要解决的问题和响应目标，"
        r"再把招标文件或评分细则要求转化为可执行方案。",
        r"围绕“\1”形成正式应答，\2，开篇直接进入项目理解、方案机制或实施安排，"
        r"并将采购需求、技术条款与交付约束转化为可执行方案。",
        normalized,
    )
    normalized = re.sub(r"围绕“([^”]+)”撰写本节内容，", r"围绕“\1”形成正式应答，", normalized)
    normalized = normalized.replace(
        "先说明本节要解决的问题和响应目标，再把招标文件或评分细则要求转化为可执行方案",
        "开篇直接进入项目理解、方案机制或实施安排，并将采购需求、技术条款与交付约束转化为可执行方案",
    )
    normalized = normalized.replace("评分细则要求", "采购需求与评审关注点")
    normalized = normalized.replace("评分标准", "评审关注点")
    normalized = normalized.replace("评分项", "评审关注点")
    normalized = re.sub(r"对应评审关注点（约\s*\d+\s*分）", "对应评审关注点", normalized)
    normalized = re.sub(r"约\s*\d+\s*分", "对应权重", normalized)
    return re.sub(r"\s+", " ", normalized).strip()


def _find_hint_block_ranges(text_value: str) -> list[tuple[int, int]]:
    matches = [
        (title, text_value.find(title))
        for title in _RUNTIME_HINT_BLOCK_TITLES
        if text_value.find(title) >= 0
    ]
    matches.sort(key=lambda item: item[1])
    return [
        (start, matches[index + 1][1] if index + 1 < len(matches) else len(text_value))
        for index, (_, start) in enumerate(matches)
    ]


def _build_outline_slice_block(section_outline_slice: str) -> str:
    slice_text = _normalize_hint_text(section_outline_slice)
    if not slice_text:
        return ""
    return (
        "【本节目录层级定位（只用于理解，不得输出）】\n"
        + slice_text
        + "\n- 上述目录编号和标题由外部编辑器统一渲染，正文中不得复述、改写或另起同级/下级标题。"
    )


def _build_analysis_context_block(analysis_context: str) -> str:
    context_text = _normalize_hint_text(analysis_context)
    if not context_text:
        return ""
    return "【招标文件解析参考（优先级最高，严格对应本章节要求）】\n" + context_text


def _build_formal_response_style_block() -> str:
    return (
        "【正式应答文件行文范式】\n"
        "- 写作身份：以投标人/响应人技术方案正文口吻成文，内容应可直接进入应答文件，不写提示词说明、任务说明或评分说明；\n"
        "- 开篇方式：参考成熟应答文件，先进入项目事实、政策/业务背景、我方理解、技术路线或交付安排，再自然展开论证；\n"
        "- 样例范式：第一段可从采购项目事实、政策依据或建设目标切入，第二段转入我方理解和交付路径；如章节需要总结，可用自然的“核心总结”式内容承接；\n"
        "- 评审要求：评审细则只作为内部覆盖依据，不写成正文表述，不把评审条目作为首句主语；\n"
        "- 叙述层次：优先采用“项目理解 → 方案机制 → 实施路径 → 交付/验收/风险控制”的连续论证，让段落服务于投标响应，不把写作动作本身写进正文；\n"
        "- 语言风格：正式、克制、可验收，避免宣传稿、口号化、过度夸张和模板化自我解释。"
    )


def _parse_outline_lines(section_outline_slice: str) -> list[str]:
    text_value = _normalize_hint_text(section_outline_slice)
    if not text_value:
        return []
    return [line.rstrip() for line in text_value.splitlines() if line.strip()]


def _outline_line_depth(line: str) -> int:
    raw = re.sub(r"\[当前\]\s*", "", str(line or ""))
    return len(raw) - len(raw.lstrip(" "))


def _current_outline_item_is_first_sibling(lines: list[str]) -> bool:
    for index, line in enumerate(lines):
        if "[当前]" not in line:
            continue
        depth = _outline_line_depth(line)
        for previous in reversed(lines[:index]):
            if not previous.strip():
                continue
            previous_depth = _outline_line_depth(previous)
            if previous_depth < depth:
                break
            if previous_depth == depth:
                return False
        return True
    return False


def _looks_like_first_section(section_title: str, section_outline_slice: str) -> bool:
    title = _normalize_hint_text(section_title)
    lines = _parse_outline_lines(section_outline_slice)
    numbered_text = "\n".join(line.strip() for line in lines + [title])
    if re.search(r"(^|\n)\s*(?:第一节|第1节|1[.．、]1|[（(]一[）)]|一[、.．])", numbered_text):
        return True
    if _current_outline_item_is_first_sibling(lines):
        return True
    single_section_titles = ("响应情况", "响应程度", "符合性响应", "符合性偏离", "偏离情况")
    return len(lines) <= 1 and any(marker in title for marker in single_section_titles)


def _build_section_bridge_block(section_title: str, section_outline_slice: str) -> str:
    if not _looks_like_first_section(section_title, section_outline_slice):
        return ""
    return (
        "【章内承接与开篇导入要求】\n"
        "- 本节若是所在章节的第一个正文单元，开头写 1 个投标响应定位段：先交代项目事实、业务/政策背景和我方理解，再落到技术路线、交付边界或响应策略；\n"
        "- 政策背景必须服务于本项目的技术判断、合规路径或交付安排，避免写成采购人宣传稿、城市宣传稿或立项报告；\n"
        "- 导入段不得使用“本节主要介绍/该节将阐述/本节需系统回应”这类机械句式，不得直接罗列标题或评审条目；\n"
        "- 导入段之后立即进入具体响应内容，优先围绕需求理解、偏离控制、方案措施、交付物、验收与风险控制展开。"
    )


def _build_content_expansion_constraints(section_title: str, expected_words: int, keywords: str) -> str:
    target = max(int(expected_words or 0), 0)
    keyword_text = _normalize_hint_text(keywords)
    density = (
        "不少于 4 个技术要点段（每段需包含“结论 + 依据/机制 + 落地方式”）"
        if target < 1200
        else "不少于 6 个技术要点段（每段需包含“结论 + 依据/机制 + 落地方式”）"
    )
    return (
        "【正文扩写与技术深度约束（必须遵守）】\n"
        f"- 本节标题：{_normalize_hint_text(section_title) or '未命名章节'}\n"
        "- 输出边界：只输出本节标题下面应出现的正文内容，第一行必须直接进入正文句子；\n"
        "- 禁止在正文开头或段落独立行输出本节标题、章节编号、Markdown 标题、加粗标题；\n"
        "- 禁止输出“一、/二、/三、”“1.1/1.2/1.1.1”这类独立小标题行；如需分层，只能在自然段内承接或使用普通列表项说明措施；\n"
        f"- 目标篇幅：约 {target if target > 0 else 800} 字，建议控制在目标值的 90%-110%，不得明显短于用户设置字数；\n"
        "- 允许的组织形式仅限：常规正文段落、普通有序/无序列表；列表项必须是具体措施或论证内容，不能退化成目录标题清单；\n"
        "- 不得重复输出章节名或目录结构，不得把目录当正文写出；\n"
        f"- 内容密度：{density}；\n"
        "- 每个要点优先使用“技术方案 → 实施步骤 → 验证方式/度量指标”结构；\n"
        "- 如涉及架构设计，需明确组件职责、接口边界、数据流与异常处理；\n"
        "- 如涉及实施保障，需补充可执行细节（人员角色、里程碑、风险控制、验收标准）；\n"
        "- 避免泛泛表述与同义反复，不得仅停留在原则层面；\n"
        f"- 关键词覆盖：{keyword_text if keyword_text else '按章节主题提炼 3-5 个技术关键词并自然覆盖'}。"
    )


def _extract_core_writing_intent(writing_hint: str) -> str:
    normalized = _normalize_hint_text(writing_hint)
    if not normalized:
        return ""

    ranges = _find_hint_block_ranges(normalized)
    if ranges:
        segments: list[str] = []
        cursor = 0
        for start, end in ranges:
            between = normalized[cursor:start].strip()
            if between:
                segments.append(between)
            cursor = end
        tail = normalized[cursor:].strip()
        if tail:
            segments.append(tail)
        normalized = _join_hint_segments(segments)
        if not normalized:
            return ""

    anchor_indexes = [
        normalized.find(anchor)
        for anchor in _IMPLICIT_HINT_TAIL_ANCHORS
        if normalized.find(anchor) >= 0
    ]
    if anchor_indexes:
        normalized = normalized[: min(anchor_indexes)].strip()
    return _normalize_response_writing_intent(normalized)


def _validate_required_bidder_info(bidder_info: Mapping[str, Any] | None) -> None:
    validate_required_bidder_info(bidder_info)


def _merge_bidder_pipt_context(
    *,
    mapping_table: dict[str, Any],
    placeholder_hint: str,
    bidder_info: Mapping[str, Any] | None,
) -> tuple[dict[str, str], str, Any]:
    merged_mapping_table, merged_placeholder_hint, bidder_context = merge_bidder_pipt_context(
        mapping_table=mapping_table,
        placeholder_hint=placeholder_hint,
        bidder_info=bidder_info,
    )
    return _string_mapping(merged_mapping_table), str(merged_placeholder_hint or ""), bidder_context


def _resolve_body_placeholders(
    *,
    content: str,
    request_mapping_flat: dict[str, str],
    audit_source: str,
) -> tuple[str, list[dict[str, Any]]]:
    resolved, _replace_map, replace_report = resolve_body_placeholders_native(
        content,
        {},
        request_mapping_flat,
        audit_source=audit_source,
    )
    return str(resolved or ""), replace_report if isinstance(replace_report, list) else []


def _build_placeholder_warning(
    *,
    illegal_placeholders: list[str] | None = None,
    unresolved_placeholders: list[str] | None = None,
) -> dict[str, Any] | None:
    illegal = sorted({str(item) for item in (illegal_placeholders or []) if str(item).strip()})
    unresolved = sorted({str(item) for item in (unresolved_placeholders or []) if str(item).strip()})
    if not illegal and not unresolved:
        return None
    return {
        "code": "placeholder_restore_warning",
        "message": "模型生成发生错误，请手动修改异常文本或重新生成。",
        "illegal_count": len(illegal),
        "unresolved_count": len(unresolved),
        "has_illegal_placeholder": bool(illegal),
        "has_unresolved_placeholder": bool(unresolved),
    }


def _find_illegal_pipt_bidder_placeholders(content: str) -> list[str]:
    return find_illegal_pipt_bidder_placeholders_native(content)


def _finalize_generated_body(content: str, section_title: str, *, strip_structural_numbering: bool = False) -> str:
    """正文保存前只做确定性格式清理，不依赖 legacy task_routes 私有函数。"""
    body = _clean_markdown_artifacts(str(content or ""))
    body = _normalize_generated_markdown(body, str(section_title or ""))
    if strip_structural_numbering:
        body = _strip_response_section_numbering(body)
    return body.strip()


def _clean_markdown_artifacts(text_value: str) -> str:
    if not text_value:
        return ""
    cleaned = str(text_value).strip()
    cleaned = re.sub(r"^\s*```(?:markdown|md)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```\s*$", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"^\s*#+\s*$", "", cleaned, flags=re.MULTILINE)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def _normalize_generated_markdown(content: str, section_title: str) -> str:
    try:
        return str(normalize_generated_markdown(content, section_title) or "")
    except Exception:
        return str(content or "").strip()


def _strip_response_section_numbering(text_value: str) -> str:
    if not text_value:
        return ""
    cleaned_lines: list[str] = []
    pattern = re.compile(
        r"^(?:"
        r"[一二三四五六七八九十]+、"
        r"|（[一二三四五六七八九十]+）"
        r"|\([一二三四五六七八九十]+\)"
        r"|\d+(?:\.\d+){1,3}"
        r"|\d+\."
        r")\s*"
    )
    for raw_line in str(text_value).splitlines():
        line = raw_line.strip()
        if not line:
            cleaned_lines.append("")
            continue
        line = pattern.sub("", line, count=1).strip()
        if line:
            cleaned_lines.append(line)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(cleaned_lines)).strip()


def _resolve_content_workflow_name(generation_strategy: str = "") -> str:
    strategy = str(generation_strategy or "").strip().lower()
    if strategy == "response_special":
        return "response_content_writer"
    return "content_writer"


def _extract_content_diagram_specs(outputs: dict[str, Any]) -> Any:
    if not isinstance(outputs, dict):
        return None
    return outputs.get("diagram_specs") or outputs.get("diagram_spec") or outputs.get("diagram") or None


def _persist_diagram_artifact(project_id: str, section_id: str, svg: str) -> dict[str, Any]:
    raw_svg = str(svg or "").strip()
    if not raw_svg:
        raise ValueError("svg 不能为空")
    seed = f"{project_id}\n{section_id}\n{raw_svg}".encode("utf-8", errors="ignore")
    diagram_id = hashlib.sha256(seed).hexdigest()[:24]
    project_dir = _diagram_artifact_dir() / _safe_diagram_project_dir(project_id or "default")
    project_dir.mkdir(parents=True, exist_ok=True)
    path = project_dir / f"{diagram_id}.svg"
    if not path.exists():
        path.write_text(raw_svg, encoding="utf-8")
    return {
        "diagram_id": diagram_id,
        "svg_length": len(raw_svg),
        "svg_url": f"/api/diagram-artifacts/{diagram_id}.svg?project_id={project_id}",
    }


def _persist_mermaid_artifact(project_id: str, section_id: str, mermaid: str) -> dict[str, Any]:
    raw_mermaid = str(mermaid or "").strip()
    if not raw_mermaid:
        raise ValueError("mermaid 不能为空")
    seed = f"{project_id}\n{section_id}\n{raw_mermaid}".encode("utf-8", errors="ignore")
    diagram_id = hashlib.sha256(seed).hexdigest()[:24]
    project_dir = _diagram_artifact_dir() / _safe_diagram_project_dir(project_id or "default")
    project_dir.mkdir(parents=True, exist_ok=True)
    path = project_dir / f"{diagram_id}.mmd"
    if not path.exists():
        path.write_text(raw_mermaid, encoding="utf-8")
    return {
        "diagram_id": diagram_id,
        "mermaid_length": len(raw_mermaid),
        "mermaid_url": f"/api/diagram-artifacts/{diagram_id}.mmd?project_id={project_id}",
    }


def _build_diagram_reference_tag(diagram: dict[str, Any]) -> str:
    diagram_id = str(diagram.get("diagram_id") or "").strip()
    title = str(diagram.get("title") or "架构图").replace('"', "&quot;")
    diagram_type = str(diagram.get("type") or "architecture").replace('"', "&quot;")
    return f'<diagram data-diagram-id="{diagram_id}" type="{diagram_type}" title="{title}"></diagram>'


def _build_diagram_source_context(
    diagram_brief: str,
    writing_hint: str,
    keywords: str,
    section_title: str,
    global_outline: str,
    diagram_type_hint: str,
    content_context: str = "",
    max_len: int = 2800,
) -> str:
    def clip(text_value: str, limit: int) -> str:
        normalized = str(text_value or "").strip()
        if not normalized:
            return ""
        if len(normalized) <= limit:
            return normalized
        return normalized[: max(1, limit - 8)].rstrip() + " ...(截断)"

    def focus_hint(text_value: str, limit: int) -> str:
        normalized = str(text_value or "").strip()
        if not normalized:
            return ""
        lines = [line.strip() for line in re.split(r"[\n\r]+", normalized) if line.strip()]
        keep: list[str] = []
        hot_words = (
            "架构", "模块", "接口", "数据", "链路", "流程", "分层", "服务", "数据库", "缓存",
            "消息", "安全", "鉴权", "高可用", "容灾", "监控", "日志", "告警", "规则", "算法",
        )
        for line in lines:
            if any(word in line for word in hot_words):
                keep.append(line)
            if len("\n".join(keep)) >= limit:
                break
        if not keep:
            keep = lines[:8]
        return clip("\n".join(keep), limit)

    def outline_window(outline: str, title: str, limit: int) -> str:
        raw = str(outline or "").strip()
        if not raw:
            return ""
        lines = [line.rstrip() for line in raw.splitlines() if line.strip()]
        if not lines:
            return ""
        anchor = str(title or "").strip()
        index = -1
        if anchor:
            for idx, line in enumerate(lines):
                if anchor in line:
                    index = idx
                    break
        if index < 0:
            return clip("\n".join(lines[:20]), limit)
        return clip("\n".join(lines[max(0, index - 8): min(len(lines), index + 12)]), limit)

    diagram_type = str(diagram_type_hint or "architecture").strip()
    parts = [
        clip(
            "【图生成硬约束】\n"
            "1) 必须体现分层边界与主链路；\n"
            "2) 节点命名必须是技术实体（服务/模块/中间件/存储）；\n"
            "3) 必须包含关键连线语义（调用/数据/约束）；\n"
            f"4) 图类型优先按 {diagram_type} 组织结构。",
            260,
        )
    ]
    for label, value in (
        ("【diagramBrief】\n", clip(diagram_brief, 850)),
        ("【章节标题】", clip(section_title, 120)),
        ("【关键词】", clip(keywords, 220)),
        ("【写作引导-技术相关摘要】\n", focus_hint(writing_hint, 1200)),
        ("【已生成正文摘要】\n", clip(content_context, 1200)),
        ("【大纲邻域窗口】\n", outline_window(global_outline, section_title, 650)),
    ):
        if value:
            parts.append(label + value)
    out = "\n\n---\n\n".join(part for part in parts if part.strip())
    if len(out) > max_len:
        out = out[: max(1, max_len - 8)].rstrip() + " ...(截断)"
    return out


def _extract_svg_from_candidate(raw: Any) -> str:
    if raw is None:
        return ""
    if isinstance(raw, str):
        text_value = raw.strip()
        if not text_value:
            return ""
        svg_match = re.search(r"<svg\b[\s\S]*?</svg>", text_value, flags=re.IGNORECASE)
        if svg_match:
            return svg_match.group(0).strip()
        parsed = _try_parse_jsonish(text_value)
        if parsed is not None and parsed is not raw:
            return _extract_svg_from_candidate(parsed)
        return ""
    if isinstance(raw, Mapping):
        for key in ("svg", "svg_content", "content", "result", "text", "output", "structured_output"):
            svg = _extract_svg_from_candidate(raw.get(key))
            if svg:
                return svg
        for value in raw.values():
            svg = _extract_svg_from_candidate(value)
            if svg:
                return svg
    if isinstance(raw, list):
        for item in raw:
            svg = _extract_svg_from_candidate(item)
            if svg:
                return svg
    return ""


def _extract_mermaid_from_candidate(raw: Any) -> str:
    if raw is None:
        return ""
    if isinstance(raw, str):
        text_value = raw.strip()
        if not text_value:
            return ""
        fence = re.search(r"```mermaid\s*([\s\S]*?)```", text_value, flags=re.IGNORECASE)
        if fence:
            return fence.group(1).strip()
        parsed = _try_parse_jsonish(text_value)
        if parsed is not None and parsed is not raw:
            return _extract_mermaid_from_candidate(parsed)
        first_line = text_value.splitlines()[0].strip().lower() if text_value.splitlines() else ""
        if first_line.startswith(("graph ", "flowchart ", "sequencediagram", "classdiagram", "statediagram", "erdiagram", "journey", "gantt", "pie ")):
            return text_value
        return ""
    if isinstance(raw, Mapping):
        for key in ("mermaid", "mermaid_source", "mmd", "code", "content", "result", "text", "output", "structured_output"):
            mermaid = _extract_mermaid_from_candidate(raw.get(key))
            if mermaid:
                return mermaid
        for value in raw.values():
            mermaid = _extract_mermaid_from_candidate(value)
            if mermaid:
                return mermaid
    if isinstance(raw, list):
        for item in raw:
            mermaid = _extract_mermaid_from_candidate(item)
            if mermaid:
                return mermaid
    return ""


def _extract_diagram_svg_output(outputs: dict[str, Any]) -> str:
    if not isinstance(outputs, dict):
        return ""
    for key in ("svg", "svg_content", "result", "text", "output", "structured_output"):
        svg = _extract_svg_from_candidate(outputs.get(key))
        if svg:
            return svg
    for value in outputs.values():
        svg = _extract_svg_from_candidate(value)
        if svg:
            return svg
    return ""


def _extract_diagram_mermaid_output(outputs: dict[str, Any]) -> str:
    if not isinstance(outputs, dict):
        return ""
    for key in ("mermaid", "mermaid_source", "mmd", "code", "result", "text", "output", "structured_output"):
        mermaid = _extract_mermaid_from_candidate(outputs.get(key))
        if mermaid:
            return mermaid
    for value in outputs.values():
        mermaid = _extract_mermaid_from_candidate(value)
        if mermaid:
            return mermaid
    return ""


def _is_fallback_diagram_svg(svg: str) -> bool:
    text_value = re.sub(r"\s+", "", str(svg or ""))
    if not text_value:
        return False
    if "架构图" in text_value and "降级模板" in text_value:
        return True
    return ">上游模块<" in text_value and ">下游模块<" in text_value


def _normalize_diagram_spec_list(value: Any) -> list[Any]:
    if isinstance(value, list):
        return value
    if isinstance(value, Mapping):
        return [dict(value)]
    return []


def _parse_diagram_specs(raw: Any) -> dict[str, list[Any]]:
    parsed = _try_parse_jsonish(raw)
    if isinstance(parsed, list):
        parsed = {"elements": parsed}
    if not isinstance(parsed, Mapping):
        return {"elements": [], "flows": [], "emphasis": []}
    return {
        "elements": _normalize_diagram_spec_list(parsed.get("elements") or parsed.get("nodes") or parsed.get("components") or parsed.get("modules") or []),
        "flows": _normalize_diagram_spec_list(parsed.get("flows") or parsed.get("edges") or parsed.get("links") or parsed.get("relations") or []),
        "emphasis": _normalize_diagram_spec_list(parsed.get("emphasis") or parsed.get("highlights") or parsed.get("key_nodes") or []),
    }


def _summarize_workflow_outputs(outputs: dict[str, Any]) -> str:
    if not isinstance(outputs, dict):
        return f"outputs 类型异常: {type(outputs).__name__}"
    parts: list[str] = []
    for key, value in outputs.items():
        text = value if isinstance(value, str) else json.dumps(value, ensure_ascii=False, default=str)
        parts.append(f"{key}={str(text)[:180]}")
    return "; ".join(parts)[:700] if parts else "outputs 为空"


async def _collect_workflow_outputs(
    task_id: str,
    dify_key: str,
    inputs: dict[str, Any],
    *,
    _r: Any = None,
    initial_stage: str,
) -> dict[str, Any]:
    task_manager = _task_manager()
    task_manager.update_stage(task_id, initial_stage)
    outputs: dict[str, Any] = {}
    got_finished = False
    workflow_run_id = ""
    async for chunk in _call_dify_workflow_stream(dify_key, inputs):
        _ensure_task_running(task_id)
        if isinstance(chunk, dict):
            if chunk.get("dify_task_id"):
                task_manager.set_dify_task_id(task_id, chunk["dify_task_id"])
            if chunk.get("__error__"):
                raise RuntimeError(_format_dify_runtime_error(RuntimeError(str(chunk.get("error") or "Dify 工作流返回错误事件"))))
            if chunk.get("__stage__"):
                workflow_run_id = str(chunk.get("workflow_run_id") or workflow_run_id or "")
                task_manager.update_stage(task_id, chunk["__stage__"])
            elif chunk.get("__finished__"):
                outputs = chunk.get("outputs", {}) or {}
                workflow_run_id = str(chunk.get("workflow_run_id") or workflow_run_id or "")
                got_finished = True
                break
    if not got_finished and workflow_run_id:
        logger.warning("[Task %s] 内容工作流未收到 finished，尝试 fallback GET /workflows/run/%s", task_id, workflow_run_id)
        try:
            fb_data = await _get_dify_workflow_run_result(dify_key, workflow_run_id)
            outputs = (((fb_data or {}).get("data") or {}).get("outputs") or {}) if isinstance(fb_data, dict) else {}
            if outputs:
                task_manager.update_stage(task_id, "📥 正在回收远端完成结果")
                got_finished = True
        except Exception as fb_err:
            logger.warning("[Task %s] 内容工作流 fallback GET 失败: %s", task_id, _format_dify_runtime_error(fb_err))
    if not got_finished:
        raise RuntimeError("内容工作流异常中断（未收到 finished 事件）")
    return outputs


def _finalize_legacy_content_output(
    raw_content: Any,
    section_title: str,
    *,
    feedback: str = "",
    request_mapping_flat: dict[str, str] | None = None,
    strip_structural_numbering: bool = False,
    audit_source: str = "apps_api.content_result",
) -> tuple[str, list[dict[str, Any]], dict[str, Any] | None]:
    if isinstance(raw_content, list):
        raw_content = "\n\n".join(str(item) for item in raw_content)
    content = re.sub(r"<think>.*?</think>", "", str(raw_content or ""), flags=re.DOTALL).strip()
    content = _finalize_generated_body(
        content,
        section_title,
        strip_structural_numbering=strip_structural_numbering,
    )
    fb_clean = str(feedback or "").strip()
    if fb_clean and len(fb_clean) > 10 and content.startswith(fb_clean):
        content = content[len(fb_clean):].strip()
        content = _finalize_generated_body(
            content,
            section_title,
            strip_structural_numbering=strip_structural_numbering,
        )

    content, replace_report = _resolve_body_placeholders(
        content=content,
        request_mapping_flat=request_mapping_flat or {},
        audit_source=audit_source,
    )
    placeholder_issues = _find_illegal_pipt_bidder_placeholders(content)
    unresolved_placeholders = [
        str(item.get("placeholder") or "")
        for item in replace_report
        if isinstance(item, Mapping) and item.get("status") == "miss" and item.get("placeholder")
    ]
    placeholder_warning = _build_placeholder_warning(
        illegal_placeholders=placeholder_issues,
        unresolved_placeholders=unresolved_placeholders,
    )
    return content, replace_report, placeholder_warning


async def _run_inline_content_diagram(
    *,
    payload: dict[str, Any],
    workflow_name: str,
    content: str,
    writing_hint: str,
    outputs: dict[str, Any],
) -> tuple[str, int, dict[str, Any] | None, Any]:
    diagram_specs = _extract_content_diagram_specs(outputs)
    if workflow_name != "content_writer":
        return content, 0, None, diagram_specs

    enable_diagrams = bool(payload.get("enable_diagrams") and _diagram_generation_enabled())
    max_diagrams = _int_or_default(payload.get("max_diagrams"), default=0) if enable_diagrams else 0
    need_diagram = bool(payload.get("need_diagram") and enable_diagrams)
    diagram_brief = str(payload.get("diagram_brief") or "") if enable_diagrams else ""
    wants_diagram = enable_diagrams and need_diagram and bool(diagram_brief.strip()) and max_diagrams > 0
    if not wants_diagram:
        return content, 0, None, diagram_specs

    diagram_key = _get_workflow_key(_get_diagram_workflow_name())
    if not diagram_key:
        return content, 0, {
            "code": "diagram_key_missing",
            "message": f"{_get_diagram_workflow_name()} 工作流 API Key 未配置",
            "section_title": str(payload.get("section_title") or ""),
        }, diagram_specs

    task_manager = _task_manager()
    project_id = str(payload.get("project_id") or "").strip() or "legacy-content"
    task_id = task_manager.create_task("diagram", project_id, workflow_name=_get_diagram_workflow_name())

    try:
        diagrams_generated, diagram_slot_reserved, diagram_error = await _execute_diagram_for_section(
            task_id=task_id,
            project_id=project_id,
            diagram_key=diagram_key,
            enable_diagrams=enable_diagrams,
            need_diagram=need_diagram,
            diagram_brief=diagram_brief,
            max_diagrams=max_diagrams,
            diagram_type_hint=str(payload.get("diagram_type_hint") or "architecture"),
            section_title=str(payload.get("section_title") or ""),
            writing_hint=writing_hint,
            raw_keywords=str(payload.get("keywords") or ""),
            raw_global_outline=str(payload.get("global_outline") or ""),
            content_context=content,
            diagram_specs=diagram_specs,
        )
        if not diagrams_generated and diagram_slot_reserved:
            await task_manager.release_diagram_slot(project_id)
        if diagrams_generated:
            content = content + "\n" + "\n".join(_build_diagram_reference_tag(item) for item in diagrams_generated)
        task_manager.set_result(
            task_id,
            {
                "done": True,
                "section_id": str(payload.get("section_id") or ""),
                "diagrams_count": len(diagrams_generated),
                "diagram_error": diagram_error,
            },
        )
        return content, len(diagrams_generated), diagram_error, diagram_specs
    except Exception as exc:
        diagram_error = {
            "code": "diagram_inline_error",
            "message": _format_dify_runtime_error(exc),
            "section_title": str(payload.get("section_title") or ""),
        }
        task_manager.set_error(task_id, diagram_error["message"])
        return content, 0, diagram_error, diagram_specs


def _build_diagram_skip_payload(
    *,
    workflow_name: str,
    enable_diagrams: bool,
    need_diagram: bool,
    diagram_brief: str,
    max_diagrams: int,
    diagram_key: str,
) -> dict[str, Any] | None:
    reasons: list[str] = []
    if workflow_name != "content_writer":
        reasons.append(f"workflow_name={workflow_name}")
    if not _diagram_generation_enabled():
        reasons.append("ENABLE_DIAGRAM_GENERATION=false")
    if not enable_diagrams:
        reasons.append("enable_diagrams=false")
    if not need_diagram:
        reasons.append("need_diagram=false")
    if not str(diagram_brief or "").strip():
        reasons.append("diagram_brief=empty")
    if max_diagrams <= 0:
        reasons.append(f"max_diagrams={max_diagrams}")
    if enable_diagrams and need_diagram and str(diagram_brief or "").strip() and max_diagrams > 0 and not diagram_key:
        reasons.append(f"{_get_diagram_workflow_name()}_key_missing")
    if not reasons:
        return None
    return {
        "code": "diagram_skipped",
        "mode": "mermaid" if _get_diagram_workflow_name() == "diagram_generator_mermaid" else "svg",
        "workflow": _get_diagram_workflow_name(),
        "reasons": reasons,
    }


def _build_diagram_task_result(
    request: dict[str, Any],
    content: str,
    diagrams_generated: list[Any],
    diagram_error: dict[str, Any] | None = None,
) -> dict[str, Any]:
    section_id = str(request.get("section_id") or "")
    replace_report = request.get("replace_report", []) or []
    raw_score = request.get("quality_score")
    quality_score = None
    if raw_score is not None:
        try:
            quality_score = int(float(raw_score))
        except (ValueError, TypeError):
            pass
    result_payload = {
        "done": True,
        "section_id": section_id,
        "content": content,
        "word_count": _count_visible_chars(content),
        "quality_score": quality_score,
        "feedback": request.get("feedback"),
        "replace_report": replace_report,
        "diagrams_count": len(diagrams_generated or []),
    }
    if diagram_error:
        result_payload["diagram_error"] = diagram_error
    return result_payload


async def _execute_diagram_for_section(
    *,
    task_id: str,
    project_id: str,
    diagram_key: str | None,
    enable_diagrams: bool,
    need_diagram: bool,
    diagram_brief: str,
    max_diagrams: int,
    diagram_type_hint: str,
    section_title: str,
    writing_hint: str,
    raw_keywords: str,
    raw_global_outline: str,
    content_context: str = "",
    diagram_specs: Any = None,
) -> tuple[list[dict[str, Any]], bool, dict[str, Any] | None]:
    """执行图表工作流；入参为章节上下文，出参为生成图表、额度占用标记和错误载荷。"""
    skip_reasons: list[str] = []
    if not enable_diagrams:
        skip_reasons.append("enable_diagrams=false")
    if not need_diagram:
        skip_reasons.append("need_diagram=false")
    if not str(diagram_brief or "").strip():
        skip_reasons.append("diagram_brief=empty")
    if max_diagrams <= 0:
        skip_reasons.append(f"max_diagrams={max_diagrams}")
    if not diagram_key:
        skip_reasons.append("diagram_key_missing")
    if skip_reasons:
        logger.info(
            "[Task %s] 图表生成跳过: section=%s; reasons=%s",
            task_id,
            str(section_title or "").strip() or "<unknown>",
            ", ".join(skip_reasons),
        )
        return [], False, None

    task_manager = _task_manager()
    diagram_slot_reserved = await task_manager.reserve_diagram_slot(project_id, max_diagrams)
    if not diagram_slot_reserved:
        task_manager.update_stage(task_id, "⏭️ 图表额度已满，跳过图表")
        return [], False, None

    try:
        task_manager.update_stage(task_id, "🎨 图表生成中")
        source_context = _build_diagram_source_context(
            diagram_brief,
            writing_hint,
            raw_keywords,
            section_title,
            raw_global_outline,
            diagram_type_hint,
            content_context,
        )
        spec_payload = _parse_diagram_specs(diagram_specs)
        diagram_inputs = {
            "diagram_type": diagram_type_hint,
            "diagram_title": (section_title or "架构图")[:120],
            "source_excerpt": source_context,
            "elements": json.dumps(spec_payload["elements"], ensure_ascii=False),
            "flows": json.dumps(spec_payload["flows"], ensure_ascii=False),
            "emphasis": json.dumps(spec_payload["emphasis"], ensure_ascii=False),
            "style_preset": "consulting",
        }
        outputs: dict[str, Any] = {}
        svg_content = ""
        mermaid_content = ""
        workflow_run_id = ""
        async for chunk in _call_dify_workflow_stream(str(diagram_key), diagram_inputs):
            _ensure_task_running(task_id)
            if not isinstance(chunk, dict):
                continue
            workflow_run_id = str(chunk.get("workflow_run_id") or workflow_run_id or "")
            if chunk.get("dify_task_id"):
                task_manager.set_dify_task_id(task_id, chunk["dify_task_id"])
            if chunk.get("__error__"):
                raise RuntimeError(str(chunk.get("error") or "图表工作流返回错误事件"))
            if chunk.get("__finished__"):
                outputs = chunk.get("outputs", {}) or {}
                svg_content = _extract_diagram_svg_output(outputs)
                mermaid_content = "" if svg_content else _extract_diagram_mermaid_output(outputs)
                break

        if not svg_content and not mermaid_content and workflow_run_id:
            try:
                fallback_data = await _get_dify_workflow_run_result(str(diagram_key), workflow_run_id)
                fallback_outputs = (((fallback_data or {}).get("data") or {}).get("outputs") or {}) if isinstance(fallback_data, dict) else {}
                if isinstance(fallback_outputs, dict):
                    outputs = fallback_outputs
                    svg_content = _extract_diagram_svg_output(outputs)
                    mermaid_content = "" if svg_content else _extract_diagram_mermaid_output(outputs)
            except Exception as exc:
                logger.warning("[Task %s] 图表 workflow_run fallback 失败: %s", task_id, exc)

        if svg_content:
            output_keys = list(outputs.keys()) if isinstance(outputs, dict) else []
            if _is_fallback_diagram_svg(svg_content):
                error_payload: dict[str, Any] = {
                    "code": "diagram_fallback_svg",
                    "message": "图表工作流返回了降级模板，已拦截并保留正文。",
                    "output_keys": output_keys,
                    "svg_length": len(svg_content),
                    "section_title": str(section_title or "").strip() or "未命名章节",
                }
                if isinstance(outputs, dict):
                    for field in ("quality_report", "layout_plan", "semantic_plan"):
                        value = str(outputs.get(field) or "").strip()
                        if value:
                            error_payload[field] = _shrink_error_text(value, limit=800)
                return [], True, error_payload
            artifact = _persist_diagram_artifact(project_id, section_title, svg_content.strip())
            return (
                [{
                    "title": (section_title or "架构图")[:120],
                    "type": diagram_type_hint,
                    "diagram_id": artifact["diagram_id"],
                    "svg_url": artifact["svg_url"],
                    "svg_length": artifact["svg_length"],
                }],
                True,
                None,
            )

        if mermaid_content:
            artifact = _persist_mermaid_artifact(project_id, section_title, mermaid_content.strip())
            return (
                [{
                    "title": (section_title or "架构图")[:120],
                    "type": diagram_type_hint,
                    "diagram_id": artifact["diagram_id"],
                    "mermaid_url": artifact["mermaid_url"],
                    "mermaid_length": artifact["mermaid_length"],
                }],
                True,
                None,
            )

        return [], True, {
            "code": "diagram_output_missing",
            "message": "图表工作流已完成，但未返回可识别的 SVG 或 Mermaid 输出。",
            "output_keys": list(outputs.keys()) if isinstance(outputs, dict) else [],
            "section_title": str(section_title or "").strip() or "未命名章节",
        }
    except Exception as exc:
        error_payload = _build_diagram_error_payload(exc, section_title)
        logger.warning("[Task %s] 图表生成失败: %s", task_id, error_payload["message"])
        if diagram_slot_reserved:
            await task_manager.release_diagram_slot(project_id)
        return [], False, error_payload


def _emit_outline_stage_event_local(task_id: str, label: str, *, elapsed_sec: int = 0, heartbeat: bool = False) -> None:
    phase, percent = _outline_stage_meta_from_label(label)
    _push_task_event(
        task_id,
        "stage",
        {
            "label": label,
            "phase": phase,
            "percent": percent,
            "elapsed_sec": int(max(elapsed_sec, 0)),
            "heartbeat": bool(heartbeat),
        },
    )


def _make_h2_seed_sections_local(seed_headings: list[dict[str, Any]]) -> list[dict[str, Any]]:
    sections: list[dict[str, Any]] = []
    for idx, seed in enumerate(seed_headings or []):
        sections.append(
            {
                "id": str(seed.get("id") or f"seed_{idx + 1}"),
                "title": str(seed.get("title") or ""),
                "wordCount": int(seed.get("wordCount") or 0),
                "writingHint": str(seed.get("writingHint") or ""),
                "keywords": seed.get("keywords") or [],
                "relatedAnalysisIds": seed.get("relatedAnalysisIds") or [],
                "needDiagram": False,
                "diagramBrief": "",
                "diagramPlan": {},
                "headingLevel": 2,
                "children": [],
            }
        )
    return sections


def _outline_sections_window_batches_local(sections: list[dict[str, Any]], window_size: int = 2) -> list[list[dict[str, Any]]]:
    if not sections:
        return []
    size = max(1, int(window_size or 1))
    return [sections[i:i + size] for i in range(0, len(sections), size)]


def _shrink_error_text(text: str, limit: int = 220) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: max(1, limit - 8)].rstrip() + " ...(截断)"


def _build_diagram_error_payload(exc: Exception, section_title: str) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "code": "diagram_failed",
        "message": "图表工作流调用失败",
    }
    title = str(section_title or "").strip() or "未命名章节"
    if isinstance(exc, httpx.HTTPStatusError):
        response = exc.response
        status_code = int(response.status_code)
        payload["status_code"] = status_code
        detail = ""
        try:
            body = response.json()
            if isinstance(body, dict):
                detail = str(body.get("message") or body.get("detail") or body.get("error") or body.get("code") or "").strip()
            elif body is not None:
                detail = str(body).strip()
        except Exception:
            try:
                detail = response.text.strip()
            except Exception:
                detail = ""
        detail = _shrink_error_text(detail)
        if status_code == 401:
            payload["code"] = "diagram_auth_failed"
            payload["message"] = f"图表工作流鉴权失败（401）。请检查 {_get_diagram_workflow_name()} 对应的 Dify API Key 是否有效。"
        elif status_code == 404:
            payload["code"] = "diagram_endpoint_not_found"
            payload["message"] = "图表工作流接口不存在（404）。请检查 DIFY_API_URL 或目标 Dify 实例。"
        else:
            payload["code"] = f"diagram_http_{status_code}"
            payload["message"] = f"图表工作流调用失败（HTTP {status_code}）。"
        if detail:
            payload["message"] = f"{payload['message']} Dify 返回：{detail}"
        payload["detail"] = detail
        payload["section_title"] = title
        return payload
    if isinstance(exc, httpx.RequestError):
        payload["code"] = "diagram_request_error"
        payload["message"] = f"图表工作流请求失败：{_shrink_error_text(_format_dify_runtime_error(exc))}"
        payload["section_title"] = title
        return payload
    payload["message"] = f"图表工作流异常：{_shrink_error_text(_format_dify_runtime_error(exc))}"
    payload["section_title"] = title
    return payload


async def _run_diagram_request(
    task_id: str,
    request: dict[str, Any],
    diagram_key: str,
) -> dict[str, Any]:
    project_id = str(request.get("project_id") or "").strip()
    section_title = str(request.get("section_title") or "").strip()
    base_content = str(request.get("base_content") or "")
    writing_hint = str(request.get("writing_hint") or "")
    keywords = str(request.get("keywords") or "")
    expected_words = _int_or_default(request.get("expected_words"), default=900)
    analysis_context = str(request.get("analysis_context") or "")
    slice_text = str(request.get("section_outline_slice") or "")
    composed_hint = _compose_runtime_writing_hint(
        writing_hint,
        section_title,
        expected_words,
        keywords,
        section_outline_slice=slice_text,
        analysis_context=analysis_context,
    )

    enable_diagrams = bool(request.get("enable_diagrams", False) and _diagram_generation_enabled())
    max_diagrams = _int_or_default(request.get("max_diagrams"), default=0) if enable_diagrams else 0
    need_diagram = bool(request.get("need_diagram", False))
    diagram_brief = str(request.get("diagram_brief") or "")
    diagram_type_hint = str(request.get("diagram_type_hint") or "architecture")
    diagram_specs = request.get("diagram_specs") or request.get("diagram_spec")
    raw_global_outline = str(request.get("global_outline") or "")

    request_mapping_flat = _string_mapping(request.get("mapping_table"))
    replace_map_seed: dict[str, str] = {}
    for row in request.get("replace_report", []) or []:
        if isinstance(row, dict) and row.get("placeholder"):
            replace_map_seed[str(row["placeholder"])] = str(row.get("original", ""))

    diagrams_generated, diagram_slot_reserved, diagram_error = await _execute_diagram_for_section(
        task_id=task_id,
        project_id=project_id,
        diagram_key=diagram_key,
        enable_diagrams=enable_diagrams,
        need_diagram=need_diagram,
        diagram_brief=diagram_brief,
        max_diagrams=max_diagrams,
        diagram_type_hint=diagram_type_hint,
        section_title=section_title,
        writing_hint=composed_hint,
        raw_keywords=keywords,
        raw_global_outline=raw_global_outline,
        content_context=base_content,
        diagram_specs=diagram_specs,
    )
    if not diagrams_generated and diagram_slot_reserved:
        await _task_manager().release_diagram_slot(project_id)

    content = base_content
    if diagrams_generated:
        content = content + "\n" + "\n".join(_build_diagram_reference_tag(item) for item in diagrams_generated)
    content, replace_report = _resolve_body_placeholders(
        content=content,
        request_mapping_flat=request_mapping_flat,
        audit_source="apps_api.task.diagram_section",
    )
    req_for_result = {**request, "replace_report": replace_report}
    return _build_diagram_task_result(req_for_result, content, diagrams_generated, diagram_error)


def _dedupe_join(parts: list[Any], *, max_len: int) -> str:
    seen: set[str] = set()
    kept: list[str] = []
    total = 0
    for part in parts:
        text = re.sub(r"\s+", " ", str(part or "")).strip()
        if not text:
            continue
        key = text.lower()
        if key in seen:
            continue
        seen.add(key)
        append_len = len(text) + (1 if kept else 0)
        if kept and total + append_len > max_len:
            break
        if not kept and len(text) > max_len:
            kept.append(text[:max_len].strip())
            break
        kept.append(text)
        total += append_len
    return "\n".join(kept).strip()


def _normalize_group_title_key(value: str) -> str:
    text = re.sub(r"\s+", "", str(value or "").strip()).lower()
    text = re.sub(r"^[\d一二三四五六七八九十百]+[.、)）]?", "", text)
    return text


def _strip_code_fence(text: str) -> str:
    raw = str(text or "").strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[-1]
        raw = raw.rsplit("```", 1)[0]
    return raw.strip()


def _try_parse_jsonish(value: Any) -> Any:
    if isinstance(value, (dict, list)):
        return value
    text = _strip_code_fence(str(value or ""))
    if not text:
        return None
    try:
        return json.loads(text)
    except Exception:
        try:
            return ast.literal_eval(text)
        except Exception:
            return None


def _extract_group_sections_payload(outputs: dict[str, Any]) -> list[dict[str, Any]]:
    candidates = [
        outputs.get("sections"),
        outputs.get("result"),
        outputs.get("text"),
        outputs.get("structured_output"),
        outputs.get("sections_json"),
    ]
    for candidate in candidates:
        parsed = _try_parse_jsonish(candidate)
        if isinstance(parsed, list):
            return [row for row in parsed if isinstance(row, dict)]
        if isinstance(parsed, dict):
            sections = parsed.get("sections") or parsed.get("items") or parsed.get("data")
            if isinstance(sections, list):
                return [row for row in sections if isinstance(row, dict)]
    return []


def _build_group_writing_children(children: list[dict]) -> list[dict]:
    normalized: list[dict] = []
    for idx, child in enumerate(children):
        section_id = str(child.get("section_id") or child.get("id") or f"group_child_{idx + 1}").strip()
        section_title = str(child.get("section_title") or child.get("title") or "").strip()
        keywords = str(child.get("keywords") or section_title).strip()
        expected_words = _int_or_default(child.get("expected_words"), default=0)
        analysis_context = str(child.get("analysis_context") or "").strip()
        normalized.append(
            {
                "section_id": section_id,
                "section_title": section_title,
                "keywords": keywords,
                "expected_words": expected_words,
                "writing_hint": _compose_runtime_writing_hint(
                    str(child.get("writing_hint") or ""),
                    section_title,
                    expected_words,
                    keywords,
                    section_outline_slice=str(child.get("section_outline_slice") or ""),
                    analysis_context=analysis_context,
                ),
                "analysis_context": analysis_context,
                "section_outline_slice": str(child.get("section_outline_slice") or ""),
                "requires_search": bool(child.get("requires_search", False)),
                "generation_strategy": str(child.get("generation_strategy") or "general").strip() or "general",
                "need_diagram": bool(child.get("need_diagram", False)),
                "diagram_brief": str(child.get("diagram_brief") or "").strip(),
                "diagram_type_hint": str(child.get("diagram_type_hint") or "architecture").strip() or "architecture",
                "diagram_priority": _int_or_default(child.get("diagram_priority"), default=0),
            }
        )
    return normalized


def _build_group_search_query(group_title: str, children: list[dict], max_terms: int = 12, max_len: int = 160) -> str:
    terms: list[str] = []
    seen: set[str] = set()

    def push(raw: Any) -> None:
        text = re.sub(r"\s+", " ", str(raw or "").strip())
        if not text:
            return
        key = re.sub(r"\s+", "", text).lower()
        if key in seen:
            return
        seen.add(key)
        terms.append(text)

    def split_keywords(raw: Any) -> list[str]:
        text = str(raw or "").strip()
        if not text:
            return []
        normalized = re.sub(r"[，、；;/|]+", ",", text)
        return [item.strip() for item in normalized.split(",") if item.strip()]

    push(group_title)
    for child in children:
        push(child.get("section_title"))
    for child in children:
        for keyword in split_keywords(child.get("keywords")):
            push(keyword)

    compact: list[str] = []
    current_len = 0
    for term in terms:
        if len(compact) >= max_terms:
            break
        next_len = current_len + (1 if compact else 0) + len(term)
        if compact and next_len > max_len:
            break
        compact.append(term)
        current_len = next_len
    return " ".join(compact).strip() or str(group_title or "").strip() or "招标技术方案"


def _match_group_section_item(item: dict[str, Any], children: list[dict]) -> Optional[dict]:
    item_id = str(item.get("section_id") or item.get("id") or "").strip()
    item_title = _normalize_group_title_key(str(item.get("section_title") or item.get("title") or ""))
    for child in children:
        if item_id and item_id == child["section_id"]:
            return child
        if item_title and item_title == _normalize_group_title_key(child["section_title"]):
            return child
    return None


def _finalize_single_content_result(
    section_title: str,
    outputs: dict[str, Any],
    request_mapping_flat: dict[str, str],
    *,
    strip_structural_numbering: bool = False,
) -> dict[str, Any]:
    content, replace_report, placeholder_warning = _finalize_legacy_content_output(
        outputs.get("text") or outputs.get("result") or outputs.get("structured_output") or "",
        section_title,
        feedback=str(outputs.get("feedback") or ""),
        request_mapping_flat=request_mapping_flat,
        strip_structural_numbering=strip_structural_numbering,
        audit_source="apps_api.group_content_result",
    )
    content, referenced_images = _normalize_referenced_images(content)
    quality_score = None
    raw_score = outputs.get("quality_score")
    if raw_score is not None:
        try:
            quality_score = int(float(raw_score))
        except (TypeError, ValueError):
            quality_score = None
    return {
        "content": content,
        "word_count": _count_visible_chars(content),
        "quality_score": quality_score,
        "feedback": outputs.get("feedback") or None,
        "replace_report": replace_report,
        "placeholder_warning": placeholder_warning,
        "referenced_images": referenced_images,
        "placeholder_issues": [],
    }


def _parse_group_content_results(
    outputs: dict[str, Any],
    children: list[dict],
    request_mapping_flat: dict[str, str],
) -> dict[str, Any]:
    sections_raw = _extract_group_sections_payload(outputs)
    child_by_id = {child["section_id"]: child for child in children}
    rank = {child["section_id"]: idx for idx, child in enumerate(children)}
    failed_by_id: dict[str, str] = {}
    if not sections_raw:
        return {
            "sections": [],
            "failed_sections": [
                {
                    "section_id": child["section_id"],
                    "section_title": child["section_title"],
                    "error": "批量正文工作流未返回可解析的 sections",
                }
                for child in children
            ],
            "parse_error": "批量正文工作流未返回可解析的 sections",
        }

    ordered: list[dict[str, Any]] = []
    used_ids: set[str] = set()
    for item in sections_raw:
        child = _match_group_section_item(item, children)
        if not child:
            continue
        section_id = child["section_id"]
        if section_id in used_ids:
            continue
        raw_content = item.get("content") or item.get("text") or item.get("body") or ""
        if not str(raw_content or "").strip():
            failed_by_id[section_id] = "批量正文结果正文为空"
            continue
        try:
            payload = _finalize_single_content_result(child["section_title"], {"text": raw_content}, request_mapping_flat)
        except RuntimeError as exc:
            failed_by_id[section_id] = _format_dify_runtime_error(exc)
            continue
        payload.update({"section_id": section_id, "section_title": child["section_title"]})
        raw_score = item.get("quality_score")
        if raw_score is not None:
            try:
                payload["quality_score"] = int(float(raw_score))
            except (TypeError, ValueError):
                pass
        if item.get("feedback"):
            payload["feedback"] = str(item.get("feedback") or "")
        ordered.append(payload)
        used_ids.add(section_id)

    for child in children:
        section_id = child["section_id"]
        if section_id in used_ids or section_id in failed_by_id:
            continue
        failed_by_id[section_id] = "批量正文结果缺失子章节"

    ordered.sort(key=lambda row: rank.get(str(row.get("section_id") or ""), 9999))
    failed_sections = [
        {
            "section_id": section_id,
            "section_title": child_by_id.get(section_id, {}).get("section_title", section_id),
            "error": error,
        }
        for section_id, error in failed_by_id.items()
    ]
    failed_sections.sort(key=lambda row: rank.get(str(row.get("section_id") or ""), 9999))
    parse_error = ""
    if not ordered:
        parse_error = "批量正文工作流返回了 sections，但没有可用正文"
    elif failed_sections:
        parse_error = "批量正文结果存在缺失子章节"
    return {"sections": ordered, "failed_sections": failed_sections, "parse_error": parse_error}


async def _repair_group_failed_sections(
    *,
    task_id: str,
    children: list[dict],
    failed_sections: list[dict],
    request: dict,
    request_mapping_flat: dict[str, str],
    group_placeholder_hint: str,
    group_outline_slice: str,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    if not failed_sections:
        return [], []

    child_by_id = {child["section_id"]: child for child in children}
    repaired: list[dict[str, Any]] = []
    still_failed: list[dict[str, Any]] = []
    task_manager = _task_manager()
    for failed in failed_sections:
        section_id = str(failed.get("section_id") or "").strip()
        child = child_by_id.get(section_id)
        if not child:
            still_failed.append(failed)
            continue

        workflow_name = _resolve_content_workflow_name(str(child.get("generation_strategy") or "general"))
        dify_key = _get_workflow_key(workflow_name)
        if not dify_key:
            still_failed.append({**failed, "error": f"{workflow_name} 工作流 API Key 未配置，无法补生成"})
            continue

        task_manager.update_stage(task_id, f"🩹 子章节补生成中：{child['section_title']}")
        inputs: dict[str, Any] = {
            "section_title": child["section_title"],
            "writing_hint": _compose_runtime_writing_hint(
                str(child.get("writing_hint") or ""),
                child["section_title"],
                _int_or_default(child.get("expected_words"), default=0),
                str(child.get("keywords") or ""),
                section_outline_slice=str(child.get("section_outline_slice") or group_outline_slice),
                analysis_context=str(child.get("analysis_context") or ""),
            ),
            "keywords": child["keywords"] if str(child.get("keywords") or "").strip() else child["section_title"],
            "expected_words": child["expected_words"],
            "project_summary": request.get("project_summary", ""),
            "global_outline": group_outline_slice,
            "placeholder_hint": group_placeholder_hint,
        }
        if workflow_name == "content_writer":
            inputs["requires_search"] = "true" if bool(child.get("requires_search", False)) else "false"
            inputs["image_map_hint"] = request.get("image_map_hint", "")
        try:
            outputs = await _collect_workflow_outputs(
                task_id,
                dify_key,
                inputs,
                _r=None,
                initial_stage=f"🩹 子章节补生成中：{child['section_title']}",
            )
            payload = _finalize_single_content_result(
                child["section_title"],
                outputs,
                request_mapping_flat,
                strip_structural_numbering=workflow_name == "response_content_writer",
            )
            diagram_specs = outputs.get("diagram_specs") or outputs.get("diagram_spec") or outputs.get("diagram")
            if diagram_specs:
                payload["diagram_specs"] = diagram_specs
            payload.update(
                {
                    "section_id": section_id,
                    "section_title": child["section_title"],
                    "repaired": True,
                    "repair_source": "single_content_writer",
                }
            )
            repaired.append(payload)
            task_manager.update_stage(task_id, f"✅ 子章节补生成完成：{child['section_title']}")
        except Exception as exc:
            still_failed.append({**failed, "error": "批量正文缺失且补生成失败: " + _format_dify_runtime_error(exc)})
            logger.warning("[Task %s] H2 子章节补生成失败: section=%s; error=%s", task_id, child["section_title"], _format_dify_runtime_error(exc))
    return repaired, still_failed


def _parse_group_review_result(outputs: dict[str, Any]) -> dict[str, Any]:
    candidates = [
        outputs.get("group_feedback"),
        outputs.get("result"),
        outputs.get("text"),
        outputs.get("structured_output"),
    ]
    for candidate in candidates:
        parsed = _try_parse_jsonish(candidate)
        if isinstance(parsed, dict):
            group_feedback = str(parsed.get("group_feedback") or parsed.get("feedback") or "").strip()
            quality_score = parsed.get("quality_score")
            payload: dict[str, Any] = {"group_feedback": group_feedback}
            if quality_score is not None:
                try:
                    payload["quality_score"] = int(float(quality_score))
                except (TypeError, ValueError):
                    pass
            return payload
        text = _strip_code_fence(str(candidate or ""))
        if text:
            return {"group_feedback": text}
    return {"group_feedback": ""}


def _normalize_referenced_images(content: str) -> tuple[str, list[dict[str, Any]]]:
    """正文内图片引用校验；当前最小迁移版先透传正文并返回空引用列表。"""
    return str(content or ""), []


def _count_visible_chars(text: str) -> int:
    if not text:
        return 0
    plain = re.sub(r"<diagram\b[\s\S]*?</diagram>", "", str(text), flags=re.IGNORECASE)
    plain = re.sub(r"<svg\b[\s\S]*?</svg>", "", plain, flags=re.IGNORECASE)
    plain = re.sub(r"<[^>]+>", "", plain)
    return len(plain.replace(" ", "").replace("\n", ""))


def _split_diagram_blocks(text: str) -> tuple[str, str]:
    raw = str(text or "")
    blocks = re.findall(r"<diagram\b[\s\S]*?</diagram>", raw, flags=re.IGNORECASE)
    content = re.sub(r"\n?<diagram\b[\s\S]*?</diagram>\n?", "\n", raw, flags=re.IGNORECASE).strip()
    suffix = "\n".join(blocks).strip()
    return content, suffix


def _persist_content_result_to_project(
    project_id: str,
    section_id: str,
    payload: dict[str, Any],
    *,
    status: str = "done",
    error: str = "",
) -> None:
    """将正文结果回写项目；最小迁移版仅保证 generatedContent 结构可落库。"""
    if not project_id or not section_id:
        return
    try:
        project = get_project_payload(project_id)
    except Exception:
        return
    data = project.get("data") if isinstance(project, dict) else {}
    if not isinstance(data, dict):
        data = {}
    generated = data.get("generatedContent")
    if not isinstance(generated, dict):
        generated = {}
    existing = generated.get(section_id) if isinstance(generated.get(section_id), dict) else {}
    if status == "done":
        content = str(payload.get("content") or "")
        generated[section_id] = {
            **existing,
            "status": "done",
            "content": content,
            "wordCount": int(payload.get("word_count") or payload.get("wordCount") or _count_visible_chars(content)),
            "qualityScore": payload.get("quality_score"),
            "feedback": payload.get("feedback"),
            "diagramError": payload.get("diagram_error"),
            "placeholderWarning": payload.get("placeholder_warning"),
            "replaceReport": payload.get("replace_report", []) or [],
            "previousContent": None,
            "previousWordCount": None,
        }
        generated[section_id].pop("error", None)
        generated[section_id].pop("stage", None)
    else:
        generated[section_id] = {
            **existing,
            "status": "error",
            "content": str(existing.get("content") or ""),
            "wordCount": int(existing.get("wordCount") or existing.get("word_count") or 0),
            "error": error or "生成失败",
            "stage": None,
        }
    patch_project_payload(project_id, {"data_patch": {"generatedContent": generated}})


def _persist_group_content_result_to_project(
    project_id: str,
    sections: list[dict[str, Any]],
    failed_sections: list[dict[str, Any]],
) -> None:
    if not project_id:
        return
    try:
        project = get_project_payload(project_id)
    except Exception:
        return
    data = project.get("data") if isinstance(project, dict) else {}
    if not isinstance(data, dict):
        data = {}
    generated = data.get("generatedContent")
    if not isinstance(generated, dict):
        generated = {}

    failed_by_id = {
        str(item.get("section_id") or ""): str(item.get("error") or "分组生成失败")
        for item in failed_sections
        if str(item.get("section_id") or "").strip()
    }
    for row in sections:
        section_id = str(row.get("section_id") or "").strip()
        if not section_id:
            continue
        content = str(row.get("content") or "")
        existing = generated.get(section_id) if isinstance(generated.get(section_id), dict) else {}
        generated[section_id] = {
            **existing,
            "status": "done",
            "content": content,
            "wordCount": int(row.get("word_count") or row.get("wordCount") or _count_visible_chars(content)),
            "qualityScore": row.get("quality_score"),
            "feedback": row.get("feedback"),
            "diagramError": row.get("diagram_error"),
            "placeholderWarning": row.get("placeholder_warning"),
            "replaceReport": row.get("replace_report", []) or [],
            "previousContent": None,
            "previousWordCount": None,
        }
        generated[section_id].pop("error", None)
        generated[section_id].pop("stage", None)

    for section_id, error in failed_by_id.items():
        existing = generated.get(section_id) if isinstance(generated.get(section_id), dict) else {}
        generated[section_id] = {
            **existing,
            "status": "error",
            "content": str(existing.get("content") or ""),
            "wordCount": int(existing.get("wordCount") or existing.get("word_count") or 0),
            "error": error or "分组生成失败",
            "stage": None,
        }
    patch_project_payload(project_id, {"data_patch": {"generatedContent": generated}})


def _match_source_pages(excerpt: str, pages_text: list[dict[str, Any]], min_match_len: int = 8) -> list[dict[str, Any]]:
    if not excerpt or not pages_text:
        return []
    clean_excerpt = excerpt.replace("\n", "").replace("\r", "").replace(" ", "").strip()
    if len(clean_excerpt) < min_match_len:
        return []
    search_key = clean_excerpt[:50]
    matched_pages: list[dict[str, Any]] = []
    for page_info in pages_text:
        page_text_clean = str(page_info.get("text") or "").replace("\n", "").replace("\r", "").replace(" ", "")
        if search_key in page_text_clean:
            matched_pages.append({"page": _int_or_default(page_info.get("page"), default=0), "excerpt": excerpt[:200]})
    if matched_pages:
        return matched_pages
    short_key = clean_excerpt[:min_match_len]
    for page_info in pages_text:
        page_text_clean = str(page_info.get("text") or "").replace("\n", "").replace("\r", "").replace(" ", "")
        if short_key in page_text_clean:
            matched_pages.append({"page": _int_or_default(page_info.get("page"), default=0), "excerpt": excerpt[:200]})
    return matched_pages


def _optional_int_value(value: Any) -> int | None:
    if value is None or value == "":
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _build_extract_requirements_payload(
    *,
    structured_data: Mapping[str, Any],
    text_for_dify: str,
    raw_image_map: dict[str, Any],
    mapping_table: dict[str, str],
    entity_count: int,
    placeholder_manifest: dict[str, Any],
    placeholder_policy: dict[str, Any],
    pdf_url: str,
    pages_text: list[dict[str, Any]],
) -> dict[str, Any]:
    requirements_raw = structured_data.get("requirements", [])
    requirements: list[dict[str, Any]] = []
    if isinstance(requirements_raw, list):
        for item in requirements_raw:
            if not isinstance(item, Mapping):
                continue
            content = str(item.get("content") or "")
            if not content:
                continue
            source_excerpt = str(item.get("source_excerpt") or "")
            requirements.append(
                {
                    "type": str(item.get("type") or "tech"),
                    "content": content,
                    "points": _optional_int_value(item.get("points")),
                    "source_excerpt": source_excerpt,
                    "source_pages": _match_source_pages(source_excerpt, pages_text) if source_excerpt else [],
                }
            )
    analysis_report = structured_data.get("analysis_report", [])
    return {
        "bid_type": str(structured_data.get("bid_type") or "tech"),
        "project_summary": str(structured_data.get("project_summary") or ""),
        "requirements": requirements,
        "analysis_report": analysis_report if isinstance(analysis_report, list) else [],
        "analysis_v2": structured_data.get("analysis_v2") if isinstance(structured_data.get("analysis_v2"), dict) else {},
        "mapping_table": mapping_table,
        "entity_count": int(entity_count or 0),
        "placeholder_manifest": placeholder_manifest,
        "placeholder_policy": placeholder_policy,
        "image_map": raw_image_map,
        "required_attachments": structured_data.get("required_attachments", []) if isinstance(structured_data.get("required_attachments"), list) else [],
        "scoring_table_template": structured_data.get("scoring_table_template", []) if isinstance(structured_data.get("scoring_table_template"), list) else [],
        "raw_document": text_for_dify,
        "pdf_url": pdf_url,
        "expected_word_count": _optional_int_value(structured_data.get("expected_word_count")),
        "expected_chapter_count": _optional_int_value(structured_data.get("expected_chapter_count")),
    }


def _build_extract_preprocess_only_payload(
    *,
    text_for_dify: str,
    raw_image_map: dict[str, Any],
    mapping_table: dict[str, str],
    entity_count: int,
    placeholder_manifest: dict[str, Any],
    placeholder_policy: dict[str, Any],
    pdf_url: str,
) -> dict[str, Any]:
    return {
        "bid_type": "tech",
        "project_summary": "",
        "requirements": [],
        "analysis_report": [],
        "analysis_v2": {},
        "mapping_table": mapping_table,
        "entity_count": int(entity_count or 0),
        "placeholder_manifest": placeholder_manifest,
        "placeholder_policy": placeholder_policy,
        "image_map": raw_image_map,
        "required_attachments": [],
        "scoring_table_template": [],
        "raw_document": text_for_dify,
        "pdf_url": pdf_url,
        "expected_word_count": None,
        "expected_chapter_count": None,
    }


def _serialize_sse_event(event: str, payload: Mapping[str, Any]) -> str:
    return f"event: {event}\ndata: {json.dumps(dict(payload), ensure_ascii=False)}\n\n"


def _persist_extract_raw_document(project_id: str, text_for_dify: str) -> None:
    _persist_raw_document(project_id, text_for_dify[:300000])


def _prepare_extract_document(
    *,
    filename: str,
    content_bytes: bytes,
    project_id: str,
    enable_desensitize: bool,
    desensitize_profile: str,
    use_vision_parsing: bool,
    task_id: str = "",
) -> dict[str, Any]:
    suffix = Path(filename).suffix.lower()
    pdf_url = ""
    pages_text: list[dict[str, Any]] = []
    cache_id = project_id or uuid.uuid4().hex[:12]
    if suffix == ".pdf":
        pdf_url = _cache_pdf_file_native(cache_id, content_bytes)
        pages_text = _extract_pdf_pages_text_native(content_bytes)
    elif suffix in {".docx", ".doc"}:
        try:
            pdf_url = _convert_to_pdf_and_cache_native(cache_id, content_bytes, filename)
        except Exception as exc:
            logger.warning("DOC/DOCX 转 PDF 失败: %s", exc)

    raw_document, raw_image_map = _extract_raw_text_with_images_native(
        filename,
        content_bytes,
        use_vision_parsing=use_vision_parsing,
    )
    if raw_document.startswith("["):
        raise PlatformError(
            code="INVALID_REQUEST",
            message="旧版 .doc 文件无法自动解析，请将文件另存为 .docx 后重新上传。",
            status_code=400,
        )

    text_for_dify = raw_document
    if suffix in {".docx", ".doc"}:
        try:
            loc_text, _loc_map, doc_blocks = _extract_docx_with_locators_native(content_bytes)
            if doc_blocks:
                _persist_project_doc_blocks_snapshot(project_id=cache_id, doc_blocks=doc_blocks)
            if suffix == ".docx":
                _persist_docx_cache(cache_id, content_bytes)
            if loc_text:
                text_for_dify = loc_text
        except Exception as exc:
            logger.warning("[extract] 定位符缓存写入失败: %s", exc)

    mapping_table: dict[str, str] = {}
    entity_count = 0
    placeholder_manifest: dict[str, Any] = {}
    placeholder_policy: dict[str, Any] = {}
    if enable_desensitize:
        try:
            desensitized = _run_bid_pipt_preprocess(
                text=text_for_dify[:300000],
                project_id=project_id or cache_id,
                task_id=task_id or "extract",
                profile_name=desensitize_profile,
            )
            text_for_dify = str(desensitized.get("text") or text_for_dify[:300000])
            mapping_table = _string_mapping(desensitized.get("mapping_table"))
            entity_count = _int_or_default(desensitized.get("mapping_table_count"), default=len(mapping_table))
            placeholder_manifest = _dict_or_default(desensitized.get("placeholder_manifest"))
            placeholder_policy = _dict_or_default(desensitized.get("placeholder_policy"))
        except Exception as exc:
            logger.warning("脱敏处理失败，使用原文继续: %s", exc)
            text_for_dify = text_for_dify[:300000]

    return {
        "cache_id": cache_id,
        "pdf_url": pdf_url,
        "pages_text": pages_text,
        "raw_image_map": raw_image_map if isinstance(raw_image_map, dict) else {},
        "text_for_dify": text_for_dify,
        "mapping_table": mapping_table,
        "entity_count": entity_count,
        "placeholder_manifest": placeholder_manifest,
        "placeholder_policy": placeholder_policy,
    }


def _workflow_env_key(workflow_name: str) -> str:
    normalized = str(workflow_name or "").strip().upper()
    return f"DIFY_WORKFLOW_{normalized}" if normalized else ""


def _get_workflow_key(workflow_name: str) -> str:
    env_key = _workflow_env_key(workflow_name)
    if not env_key:
        return ""
    value = os.environ.get(env_key, "").strip()
    if value:
        return value
    return _read_root_env_value(env_key)


def _task_workflow_name(task: Any) -> str:
    explicit = str(getattr(task, "workflow_name", "") or "").strip()
    if explicit:
        return explicit
    return {
        "outline": "structure_generator",
        "content": "content_writer",
        "diagram": "diagram_generator",
        "analyze": "doc_analysis",
    }.get(str(getattr(task, "task_type", "") or "").strip(), "")


def _task_dify_task_ids(task: Any) -> list[str]:
    values = [getattr(task, "dify_task_id", None), *(getattr(task, "dify_task_ids", None) or [])]
    return list(dict.fromkeys(str(item).strip() for item in values if str(item or "").strip()))


async def _stop_dify_workflows_for_task(task: Any) -> tuple[bool, str]:
    workflow_name = _task_workflow_name(task)
    if not workflow_name:
        return False, "not_applicable"
    dify_key = _get_workflow_key(workflow_name)
    if not dify_key:
        return False, "missing_key"
    task_ids = _task_dify_task_ids(task)
    if not task_ids:
        return False, "not_bound"

    dify_base = os.environ.get("DIFY_API_URL", "http://localhost/v1").rstrip("/")
    stopped = 0
    failed = 0
    async with httpx.AsyncClient(timeout=10) as client:
        for dify_task_id in task_ids:
            try:
                response = await client.post(
                    f"{dify_base}/workflows/tasks/{dify_task_id}/stop",
                    headers={"Authorization": f"Bearer {dify_key}"},
                    json={"user": "pro-engine-backend"},
                )
                if response.status_code == 200:
                    stopped += 1
                else:
                    failed += 1
            except Exception as exc:
                failed += 1
                logger.warning("Dify stop failed for bid task %s: %s", dify_task_id, exc)

    if stopped == len(task_ids):
        return True, "stopped"
    if stopped > 0:
        return True, "partial"
    if failed > 0:
        return False, "failed"
    return False, "not_bound"


def _persist_project_task_runtime(
    task: Any,
    *,
    runtime_state: str,
    message: str = "",
    cancellable: bool = False,
) -> None:
    project_id = str(getattr(task, "project_id", "") or "").strip()
    if not project_id:
        return
    runtime = {
        "state": runtime_state,
        "taskId": str(getattr(task, "task_id", "") or ""),
        "taskType": str(getattr(task, "task_type", "") or ""),
        "message": message,
        "progress": 100 if runtime_state == "succeeded" else 0,
        "startedAt": _utc_iso_from_timestamp(getattr(task, "created_at", None)),
        "cancellable": bool(cancellable),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    try:
        patch_project_payload(project_id, {"data_patch": {"taskRuntime": runtime}})
    except Exception as exc:
        logger.warning("Bid task runtime persistence skipped: project=%s task=%s error=%s", project_id, runtime["taskId"], exc)


def _persist_project_runtime(
    project_id: str,
    *,
    task_id: str,
    task_type: str,
    runtime_state: str,
    message: str = "",
    cancellable: bool = False,
) -> None:
    if not project_id:
        return
    runtime = {
        "state": runtime_state,
        "taskId": task_id,
        "taskType": task_type,
        "message": message,
        "progress": 100 if runtime_state == "succeeded" else 0,
        "startedAt": datetime.now(timezone.utc).isoformat(),
        "cancellable": bool(cancellable),
        "updatedAt": datetime.now(timezone.utc).isoformat(),
    }
    try:
        patch_project_payload(project_id, {"data_patch": {"taskRuntime": runtime}})
    except Exception as exc:
        logger.warning("Bid runtime bootstrap skipped: project=%s task=%s error=%s", project_id, task_id, exc)


def _push_task_event(task_id: str, event: str, payload: dict[str, Any]) -> None:
    task_manager = _task_manager()
    encoded = json.dumps({"event": event, **(payload or {})}, ensure_ascii=False)
    task_manager.update_stage(task_id, f"{_TASK_EVENT_STAGE_PREFIX}{encoded}")


def _sync_project_runtime_from_task(task: Any | None) -> None:
    if task is None:
        return
    status = str(getattr(task, "status", "") or "")
    runtime_state = _task_status_to_api_state(status)
    _persist_project_task_runtime(
        task,
        runtime_state=runtime_state,
        message=str(getattr(task, "current_stage", "") or ""),
        cancellable=status == "running",
    )


def _ensure_task_running(task_id: str) -> None:
    task = _get_task(task_id)
    if task is None:
        raise RuntimeError("任务不存在或已过期")
    status = str(getattr(task, "status", "") or "")
    if status == "running":
        return
    if status == "cancelled":
        _sync_project_runtime_from_task(task)
        raise asyncio.CancelledError()
    _sync_project_runtime_from_task(task)
    raise RuntimeError(f"任务状态异常: {status or 'unknown'}")


def _persist_raw_document(project_id: str, raw_document: str) -> None:
    if not project_id:
        return
    path = _raw_doc_cache_path(project_id)
    try:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(str(raw_document or ""), encoding="utf-8")
    except OSError as exc:
        raise PlatformError(code="BUSINESS_DIRECT_ERROR", message="缓存原文失败。", status_code=500) from exc


async def update_template_config_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """更新大纲模板；入参为模板名和模板字典，出参兼容 legacy config/template。"""
    payload = _json_object_body(body)
    template_name = _ensure_safe_template_name(str(payload.get("template_name") or ""), allow_standard=True)
    template_dict = payload.get("template_dict")
    if not isinstance(template_dict, Mapping):
        raise PlatformError(code="INVALID_REQUEST", message="template_dict 必须是对象。", status_code=400)
    template_path = _template_structures_dir() / template_name
    try:
        template_path.parent.mkdir(parents=True, exist_ok=True)
        _write_yaml_mapping(template_path, dict(template_dict))
    except OSError as exc:
        raise PlatformError(code="TEMPLATE_UPDATE_FAILED", message=f"Template {template_name} updated failed", status_code=500) from exc
    return {"status": "success", "message": f"Template {template_name} updated successfully"}


async def delete_template_config_payload(template_name: str) -> dict[str, Any]:
    """删除大纲模板；入参为模板文件名，出参兼容 legacy config/template。"""
    normalized_name = _ensure_safe_template_name(template_name, allow_standard=False)
    template_path = _template_structures_dir() / normalized_name
    try:
        if template_path.exists():
            template_path.unlink()
    except OSError as exc:
        raise PlatformError(code="TEMPLATE_DELETE_FAILED", message=f"Template {normalized_name} delete failed", status_code=500) from exc
    return {"status": "success", "message": f"Template {normalized_name} deleted successfully"}


async def update_global_config_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """更新标书全局配置；入参为 config_dict，出参为更新结果。"""
    payload = _json_object_body(body)
    config_dict = payload.get("config_dict")
    if not isinstance(config_dict, Mapping):
        raise PlatformError(code="INVALID_REQUEST", message="config_dict 必须是对象。", status_code=400)
    try:
        _write_yaml_mapping(_bid_generator_config_path(), dict(config_dict))
    except OSError as exc:
        raise PlatformError(code="GLOBAL_CONFIG_UPDATE_FAILED", message="Config updated failed", status_code=500) from exc
    return {"status": "success", "message": "Config updated successfully"}


async def generate_template_architecture_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """生成项目专属模板结构；入参为项目/蓝图/结构化数据，出参兼容 legacy template/generate。"""
    return await bid_workflow_execution_adapter.generate_template_architecture_payload(body)


async def extract_requirements_payload(
    file: UploadFile,
    *,
    project_name: str = "",
    project_id: str = "",
    enable_desensitize: bool = True,
    desensitize_profile: str = "tender",
    use_vision_parsing: bool = False,
) -> dict[str, Any]:
    """同步解析招标文件；入参为上传文件和解析配置，出参兼容 legacy extract。"""
    normalized_project_id = str(project_id or "").strip()
    filename = str(getattr(file, "filename", "") or "")
    content_bytes = await file.read()

    prepared = _prepare_extract_document(
        filename=filename,
        content_bytes=content_bytes,
        project_id=normalized_project_id,
        enable_desensitize=enable_desensitize,
        desensitize_profile=desensitize_profile,
        use_vision_parsing=use_vision_parsing,
        task_id="projects_extract",
    )

    dify_key = _get_workflow_key("requirement_extractor")
    if not dify_key:
        raise PlatformError(
            code="REQUIREMENTS_EXTRACT_FAILED",
            message="需求提取工作流 API Key 未配置，请在 .env 中设置 DIFY_WORKFLOW_REQUIREMENT_EXTRACTOR",
            status_code=500,
        )
    try:
        dify_res = await _call_dify_workflow(
            dify_key,
            {
                "raw_document": prepared["text_for_dify"],
                "project_name": str(project_name or "").strip() or Path(filename).stem,
            },
        )
    except Exception as exc:
        raise PlatformError(code="REQUIREMENTS_EXTRACT_FAILED", message=_format_dify_runtime_error(exc), status_code=500) from exc

    structured_data = parse_dify_outputs(dify_res)
    _persist_extract_raw_document(str(prepared["cache_id"]), str(prepared["text_for_dify"]))
    return _build_extract_requirements_payload(
        structured_data=structured_data,
        text_for_dify=str(prepared["text_for_dify"]),
        raw_image_map=prepared["raw_image_map"] if isinstance(prepared["raw_image_map"], dict) else {},
        mapping_table=prepared["mapping_table"] if isinstance(prepared["mapping_table"], dict) else {},
        entity_count=_int_or_default(prepared["entity_count"], default=0),
        placeholder_manifest=prepared["placeholder_manifest"] if isinstance(prepared["placeholder_manifest"], dict) else {},
        placeholder_policy=prepared["placeholder_policy"] if isinstance(prepared["placeholder_policy"], dict) else {},
        pdf_url=str(prepared["pdf_url"] or ""),
        pages_text=prepared["pages_text"] if isinstance(prepared["pages_text"], list) else [],
    )


async def extract_requirements_stream_response(
    file: UploadFile,
    *,
    project_name: str = "",
    project_id: str = "",
    enable_desensitize: bool = True,
    desensitize_profile: str = "tender",
    use_vision_parsing: bool = False,
) -> Any:
    """流式解析招标文件；入参为上传文件和解析配置，出参保持 legacy SSE 协议。"""
    content_bytes = await file.read()
    filename = str(getattr(file, "filename", "") or "")
    normalized_project_id = str(project_id or "").strip()

    async def event_stream() -> Any:
        try:
            yield _serialize_sse_event("progress", {"step": 0, "label": "解析文档结构", "percent": 5})
            prepared = _prepare_extract_document(
                filename=filename,
                content_bytes=content_bytes,
                project_id=normalized_project_id,
                enable_desensitize=False,
                desensitize_profile=desensitize_profile,
                use_vision_parsing=use_vision_parsing,
                task_id="projects_extract_stream",
            )
            yield _serialize_sse_event("progress", {"step": 0, "label": "文档结构解析完成", "percent": 15})

            text_for_dify = str(prepared["text_for_dify"] or "")
            mapping_table: dict[str, str] = {}
            entity_count = 0
            placeholder_manifest: dict[str, Any] = {}
            placeholder_policy: dict[str, Any] = {}

            if enable_desensitize:
                yield _serialize_sse_event("progress", {"step": 1, "label": "隐私脱敏处理中", "percent": 20})
                try:
                    desensitized = _run_bid_pipt_preprocess(
                        text=text_for_dify[:300000],
                        project_id=normalized_project_id or str(prepared["cache_id"]),
                        task_id="projects_extract_stream",
                        profile_name=desensitize_profile,
                    )
                    text_for_dify = str(desensitized.get("text") or text_for_dify[:300000])
                    mapping_table = _string_mapping(desensitized.get("mapping_table"))
                    entity_count = _int_or_default(desensitized.get("mapping_table_count"), default=len(mapping_table))
                    placeholder_manifest = _dict_or_default(desensitized.get("placeholder_manifest"))
                    placeholder_policy = _dict_or_default(desensitized.get("placeholder_policy"))
                    yield _serialize_sse_event(
                        "progress",
                        {"step": 1, "label": f"脱敏完成，识别 {entity_count} 处实体", "percent": 50},
                    )
                except Exception as exc:
                    logger.warning("流式脱敏处理失败，使用原文继续: %s", exc)
                    text_for_dify = text_for_dify[:300000]
                    yield _serialize_sse_event("progress", {"step": 1, "label": "脱敏跳过（使用原文）", "percent": 50})
            else:
                yield _serialize_sse_event("progress", {"step": 1, "label": "跳过脱敏", "percent": 50})

            yield _serialize_sse_event("progress", {"step": 2, "label": "预处理完成", "percent": 100})
            _persist_extract_raw_document(str(prepared["cache_id"]), text_for_dify)
            yield _serialize_sse_event(
                "result",
                _build_extract_preprocess_only_payload(
                    text_for_dify=text_for_dify,
                    raw_image_map=prepared["raw_image_map"] if isinstance(prepared["raw_image_map"], dict) else {},
                    mapping_table=mapping_table,
                    entity_count=entity_count,
                    placeholder_manifest=placeholder_manifest,
                    placeholder_policy=placeholder_policy,
                    pdf_url=str(prepared["pdf_url"] or ""),
                ),
            )
        except PlatformError as exc:
            yield _serialize_sse_event("error", {"message": exc.message})
        except Exception as exc:
            logger.error("SSE extract 异常: %s", exc, exc_info=True)
            yield _serialize_sse_event("error", {"message": str(exc)})

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


async def re_extract_requirements_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """基于缓存原文重新提取需求；入参为项目和缓存原文，出参兼容 legacy re-extract。"""
    payload = _json_object_body(body)
    project_id = _required_string(payload.get("project_id"), field="project_id")
    project_name = str(payload.get("project_name") or "").strip()
    raw_document = str(payload.get("raw_document") or "").strip() or _load_raw_document(project_id)
    if not raw_document:
        raise PlatformError(
            code="RESOURCE_NOT_FOUND",
            message="未找到项目缓存原文，请先重新上传并解析文档",
            status_code=404,
        )

    dify_key = _get_workflow_key("requirement_extractor")
    if not dify_key:
        raise PlatformError(
            code="REQUIREMENTS_RE_EXTRACT_FAILED",
            message="需求提取工作流 API Key 未配置，请在 .env 中设置 DIFY_WORKFLOW_REQUIREMENT_EXTRACTOR",
            status_code=500,
        )
    try:
        dify_res = await _call_dify_workflow(
            dify_key,
            {
                "raw_document": raw_document,
                "project_name": project_name,
            },
        )
    except Exception as exc:
        raise PlatformError(code="REQUIREMENTS_RE_EXTRACT_FAILED", message=_format_dify_runtime_error(exc), status_code=500) from exc

    structured_data = parse_dify_outputs(dify_res)
    requirements_raw = structured_data.get("requirements", [])
    requirements: list[dict[str, Any]] = []
    if isinstance(requirements_raw, list):
        for item in requirements_raw:
            if not isinstance(item, Mapping):
                continue
            content = str(item.get("content") or "")
            if not content:
                continue
            requirements.append(
                {
                    "type": str(item.get("type") or "tech"),
                    "content": content,
                    "points": _optional_int_value(item.get("points")),
                    "source_excerpt": str(item.get("source_excerpt") or ""),
                    "source_pages": [],
                }
            )

    return {
        "bid_type": str(structured_data.get("bid_type") or "tech"),
        "project_summary": str(structured_data.get("project_summary") or ""),
        "requirements": requirements,
        "analysis_report": structured_data.get("analysis_report", []) if isinstance(structured_data.get("analysis_report"), list) else [],
        "analysis_v2": structured_data.get("analysis_v2") if isinstance(structured_data.get("analysis_v2"), dict) else {},
        "mapping_table": {},
        "entity_count": 0,
        "placeholder_manifest": {},
        "placeholder_policy": {},
        "image_map": {},
        "required_attachments": structured_data.get("required_attachments", []) if isinstance(structured_data.get("required_attachments"), list) else [],
        "scoring_table_template": structured_data.get("scoring_table_template", []) if isinstance(structured_data.get("scoring_table_template"), list) else [],
        "raw_document": raw_document,
        "pdf_url": "",
        "expected_word_count": _optional_int_value(structured_data.get("expected_word_count")),
        "expected_chapter_count": _optional_int_value(structured_data.get("expected_chapter_count")),
    }


async def generate_outline_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """同步生成标书大纲；入参为大纲生成 JSON，出参兼容 legacy generate-outline。"""
    payload = _json_object_body(body)
    dify_key = _get_workflow_key("structure_generator")
    if not dify_key:
        raise PlatformError(
            code="OUTLINE_GENERATE_FAILED",
            message="大纲生成工作流 API Key 未配置，请在 .env 中设置 DIFY_WORKFLOW_STRUCTURE_GENERATOR",
            status_code=500,
        )

    bundle = build_outline_generation_bundle(
        requirements=payload.get("requirements", []) if isinstance(payload.get("requirements"), list) else [],
        analysis_context=str(payload.get("analysis_context") or ""),
        expected_total_words=_int_or_default(payload.get("expected_total_words"), default=0),
        scoring_details_json=str(payload.get("scoring_details_json") or ""),
        structure_heading_seed_json=str(payload.get("structure_heading_seed_json") or ""),
        technical_h2_bindings_json=str(payload.get("technical_h2_bindings_json") or ""),
        technical_targets_json=str(payload.get("technical_targets_json") or ""),
    )
    inputs = dict(bundle["inputs"])
    inputs["bid_type"] = str(payload.get("bid_type") or "tech")
    inputs["use_knowledge"] = "true" if bool(payload.get("use_knowledge")) else "false"
    enable_diagrams = bool(payload.get("enable_diagrams") and _diagram_generation_enabled())
    max_diagrams = _int_or_default(payload.get("max_diagrams"), default=0) if enable_diagrams else 0
    inputs["enable_diagrams"] = "true" if enable_diagrams else "false"
    inputs["max_diagrams"] = max_diagrams

    try:
        dify_res = await _call_dify_workflow(dify_key, inputs)
    except Exception as exc:
        raise PlatformError(code="OUTLINE_GENERATE_FAILED", message=_format_dify_runtime_error(exc), status_code=500) from exc

    structured_data = parse_dify_outputs(dify_res)
    sections_raw = extract_outline_sections_raw(structured_data)
    sections_data = build_seeded_outline_sections(sections_raw, bundle["seed_headings"], max_diagrams=max_diagrams)
    quality_report = evaluate_outline_quality(sections_data, bundle["seed_headings"])
    if not quality_report["pass"]:
        logger.error("[generate_outline] 结构校验失败: %s", quality_report)
        raise PlatformError(
            code="OUTLINE_GENERATE_FAILED",
            message="大纲生成结构不完整，请重试：" + "；".join(quality_report.get("issues") or []),
            status_code=502,
        )

    if not enable_diagrams:
        for section in sections_data:
            if not isinstance(section, dict):
                continue
            section["needDiagram"] = False
            section["diagramBrief"] = ""
            section["diagramPlan"] = {"enabled": False, "brief": ""}
            for child in section.get("children") if isinstance(section.get("children"), list) else []:
                if isinstance(child, dict):
                    child["needDiagram"] = False
                    child["diagramBrief"] = ""
                    child["diagramPlan"] = {"enabled": False, "brief": ""}

    return {"sections": sections_data}


async def generate_outline_stream_response(body: Mapping[str, Any]) -> Any:
    """流式生成标书大纲；入参为大纲生成 JSON，出参保持 legacy SSE 协议。"""
    payload = _json_object_body(body)
    dify_key = _get_workflow_key("structure_generator")
    if not dify_key:
        raise PlatformError(
            code="OUTLINE_GENERATE_STREAM_FAILED",
            message="大纲生成工作流 API Key 未配置，请在 .env 中设置 DIFY_WORKFLOW_STRUCTURE_GENERATOR",
            status_code=500,
        )

    bundle = build_outline_generation_bundle(
        requirements=payload.get("requirements", []) if isinstance(payload.get("requirements"), list) else [],
        analysis_context=str(payload.get("analysis_context") or ""),
        expected_total_words=_int_or_default(payload.get("expected_total_words"), default=0),
        scoring_details_json=str(payload.get("scoring_details_json") or ""),
        structure_heading_seed_json=str(payload.get("structure_heading_seed_json") or ""),
        technical_h2_bindings_json=str(payload.get("technical_h2_bindings_json") or ""),
        technical_targets_json=str(payload.get("technical_targets_json") or ""),
    )
    inputs = dict(bundle["inputs"])
    inputs["bid_type"] = str(payload.get("bid_type") or "tech")
    inputs["use_knowledge"] = "true" if bool(payload.get("use_knowledge")) else "false"
    enable_diagrams = bool(payload.get("enable_diagrams") and _diagram_generation_enabled())
    max_diagrams = _int_or_default(payload.get("max_diagrams"), default=0) if enable_diagrams else 0
    inputs["enable_diagrams"] = "true" if enable_diagrams else "false"
    inputs["max_diagrams"] = max_diagrams
    expected_total_words = _int_or_default(payload.get("expected_total_words"), default=0)

    async def event_stream() -> Any:
        try:
            async for chunk in _stream_native_outline_generation(
                dify_key=dify_key,
                inputs=inputs,
                seed_headings=bundle["seed_headings"],
                max_diagrams=max_diagrams,
                expected_total_words=expected_total_words,
            ):
                yield f"data: {json.dumps(chunk, ensure_ascii=False)}\n\n"
        except Exception as exc:
            logger.error("SSE 大纲流式生成失败: %s", exc, exc_info=True)
            yield f"data: {json.dumps({'error': str(exc)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


async def generate_content_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """同步生成章节正文；入参为章节生成 JSON，出参兼容 legacy generate-content。"""
    payload = _json_object_body(body)
    try:
        _validate_required_bidder_info(payload.get("bidder_info", {}) or {})
    except Exception as exc:
        detail = str(exc)
        if detail:
            raise PlatformError(code="INVALID_REQUEST", message=detail, status_code=400) from exc
        raise

    workflow_name = _resolve_content_workflow_name(str(payload.get("generation_strategy") or "general"))
    dify_key = _get_workflow_key(workflow_name)
    if not dify_key:
        raise PlatformError(
            code="CONTENT_GENERATE_FAILED",
            message=f"{workflow_name} 工作流 API Key 未配置",
            status_code=500,
        )

    section_title = str(payload.get("section_title") or "")
    expected_words = _int_or_default(payload.get("expected_words"), default=500)
    keywords = str(payload.get("keywords") or "").strip() or section_title
    writing_hint_merged = _compose_runtime_writing_hint(
        str(payload.get("writing_hint") or ""),
        section_title,
        expected_words,
        keywords,
        section_outline_slice=str(payload.get("section_outline_slice") or ""),
        analysis_context=str(payload.get("analysis_context") or ""),
    )
    request_mapping_flat = _string_mapping(payload.get("mapping_table"))
    merged_placeholder_hint = str(payload.get("placeholder_hint") or "")
    try:
        request_mapping_flat, merged_placeholder_hint, _bidder_context = _merge_bidder_pipt_context(
            mapping_table=request_mapping_flat,
            placeholder_hint=merged_placeholder_hint,
            bidder_info=payload.get("bidder_info", {}) or {},
        )
    except Exception:
        logger.warning("投标人信息 PIPT 归一化失败，正文直连接口使用请求原始占位符上下文", exc_info=True)

    inputs = {
        "section_title": section_title,
        "writing_hint": writing_hint_merged,
        "keywords": keywords,
        "expected_words": expected_words,
        "project_summary": str(payload.get("project_summary") or ""),
        "global_outline": str(payload.get("global_outline") or ""),
        "placeholder_hint": merged_placeholder_hint,
    }
    if workflow_name == "content_writer":
        inputs["requires_search"] = "true" if bool(payload.get("requires_search")) else "false"
        inputs["decoupling_instruction"] = (
            "【重要越界防范】：你在编写本技术正文章节时，绝对不要自行捏造或生成任何诸如“法定代表人授权书”、“无违规记录承诺函”之类的独立格式化商务附件。"
            "任何商务附件都将由专门的商务审核工作流单独生成并在汇总时拼接，你只负责纯粹的技术方案与实施规划正文编写。"
        )
        inputs["format_guardrails"] = (
            "【输出格式硬约束】：禁止输出任何 Markdown 标题（# / ## / ###）或“一、/1.1/1.1.1”式自拟小节标题；"
            "允许形式仅限：常规正文段落、编号项（有序列表）、bullet point（无序列表）。"
        )

    try:
        dify_res = await _call_dify_workflow(dify_key, inputs)
    except Exception as exc:
        raise PlatformError(code="CONTENT_GENERATE_FAILED", message=_format_dify_runtime_error(exc), status_code=500) from exc

    outputs = dify_res.get("data", {}).get("outputs", {}) if isinstance(dify_res, dict) else {}
    content, replace_report, placeholder_warning = _finalize_legacy_content_output(
        outputs.get("text") or outputs.get("result") or outputs.get("structured_output") or outputs.get("content") or "",
        section_title,
        feedback=str(outputs.get("feedback") or ""),
        request_mapping_flat=request_mapping_flat,
        strip_structural_numbering=workflow_name == "response_content_writer",
        audit_source="apps_api.generate_content",
    )
    content, diagrams_count, diagram_error, diagram_specs = await _run_inline_content_diagram(
        payload=payload,
        workflow_name=workflow_name,
        content=content,
        writing_hint=writing_hint_merged,
        outputs=outputs if isinstance(outputs, dict) else {},
    )
    word_count = len(content.replace(" ", "").replace("\n", ""))
    raw_score = outputs.get("quality_score") if isinstance(outputs, dict) else None
    quality_score = None
    if raw_score is not None:
        try:
            quality_score = int(float(raw_score))
        except (TypeError, ValueError):
            quality_score = None

    response: dict[str, Any] = {
        "section_id": str(payload.get("section_id") or ""),
        "content": content,
        "word_count": word_count,
        "quality_score": quality_score,
        "feedback": outputs.get("feedback") if isinstance(outputs, dict) else None,
        "replace_report": replace_report,
        "placeholder_warning": placeholder_warning,
        "diagrams_count": diagrams_count,
    }
    if diagram_error:
        response["diagram_error"] = diagram_error
    if diagram_specs:
        response["diagram_specs"] = diagram_specs
    return response


async def generate_content_stream_response(body: Mapping[str, Any]) -> Any:
    """流式生成章节正文；入参为章节生成 JSON，出参保持 legacy SSE 协议。"""
    payload = _json_object_body(body)
    try:
        _validate_required_bidder_info(payload.get("bidder_info", {}) or {})
    except Exception as exc:
        detail = str(exc)
        if detail:
            raise PlatformError(code="INVALID_REQUEST", message=detail, status_code=400) from exc
        raise

    workflow_name = _resolve_content_workflow_name(str(payload.get("generation_strategy") or "general"))
    dify_key = _get_workflow_key(workflow_name)
    if not dify_key:
        raise PlatformError(
            code="CONTENT_GENERATE_STREAM_FAILED",
            message=f"{workflow_name} 工作流 API Key 未配置",
            status_code=500,
        )

    section_title = str(payload.get("section_title") or "")
    expected_words = _int_or_default(payload.get("expected_words"), default=500)
    keywords = str(payload.get("keywords") or "").strip() or section_title
    writing_hint_merged = _compose_runtime_writing_hint(
        str(payload.get("writing_hint") or ""),
        section_title,
        expected_words,
        keywords,
        section_outline_slice=str(payload.get("section_outline_slice") or ""),
        analysis_context=str(payload.get("analysis_context") or ""),
    )
    request_mapping_flat = _string_mapping(payload.get("mapping_table"))
    merged_placeholder_hint = str(payload.get("placeholder_hint") or "")
    try:
        request_mapping_flat, merged_placeholder_hint, _bidder_context = _merge_bidder_pipt_context(
            mapping_table=request_mapping_flat,
            placeholder_hint=merged_placeholder_hint,
            bidder_info=payload.get("bidder_info", {}) or {},
        )
    except Exception:
        logger.warning("投标人信息 PIPT 归一化失败，正文流式接口使用请求原始占位符上下文", exc_info=True)

    inputs = {
        "section_title": section_title,
        "writing_hint": writing_hint_merged,
        "keywords": keywords,
        "expected_words": expected_words,
        "project_summary": str(payload.get("project_summary") or ""),
        "global_outline": str(payload.get("global_outline") or ""),
        "placeholder_hint": merged_placeholder_hint,
    }
    if workflow_name == "content_writer":
        inputs["requires_search"] = "true" if bool(payload.get("requires_search")) else "false"
        inputs["decoupling_instruction"] = (
            "【重要越界防范】：你在编写本技术正文章节时，绝对不要自行捏造或生成任何诸如\"法定代表人授权书\"、\"无违规记录承诺函\"之类的独立格式化商务附件。"
            "任何商务附件都将由专门的商务审核工作流单独生成并在汇总时拼接，你只负责纯粹的技术方案与实施规划正文编写。"
        )
        inputs["format_guardrails"] = (
            "【输出格式硬约束】：禁止输出任何 Markdown 标题（# / ## / ###）或“一、/1.1/1.1.1”式自拟小节标题；"
            "允许形式仅限：常规正文段落、编号项（有序列表）、bullet point（无序列表）。"
        )

    async def event_stream() -> Any:
        full_content = ""
        buffer = ""
        in_think = False
        try:
            async for chunk in _call_dify_workflow_stream(dify_key, inputs):
                if isinstance(chunk, dict):
                    if chunk.get("__finished__"):
                        outputs = chunk.get("outputs", {}) if isinstance(chunk.get("outputs"), dict) else {}
                        raw_score = outputs.get("quality_score")
                        quality_score = None
                        if raw_score is not None:
                            try:
                                quality_score = int(float(raw_score))
                            except (TypeError, ValueError):
                                quality_score = None
                        final_raw_content = (
                            outputs.get("text")
                            or outputs.get("result")
                            or outputs.get("structured_output")
                            or outputs.get("content")
                            or full_content
                        )
                        final_content, replace_report, placeholder_warning = _finalize_legacy_content_output(
                            final_raw_content,
                            section_title,
                            feedback=str(outputs.get("feedback") or ""),
                            request_mapping_flat=request_mapping_flat,
                            strip_structural_numbering=workflow_name == "response_content_writer",
                            audit_source="apps_api.generate_content_stream",
                        )
                        final_content, diagrams_count, diagram_error, diagram_specs = await _run_inline_content_diagram(
                            payload=payload,
                            workflow_name=workflow_name,
                            content=final_content,
                            writing_hint=writing_hint_merged,
                            outputs=outputs,
                        )
                        if final_content != full_content:
                            full_content = final_content
                            yield f"data: {json.dumps({'text': full_content, 'replace': True}, ensure_ascii=False)}\n\n"
                        done_payload: dict[str, Any] = {
                            "done": True,
                            "section_id": str(payload.get("section_id") or ""),
                            "word_count": len(full_content.replace(' ', '').replace('\n', '')),
                            "quality_score": quality_score,
                            "feedback": outputs.get("feedback"),
                            "replace_report": replace_report,
                            "placeholder_warning": placeholder_warning,
                            "diagrams_count": diagrams_count,
                        }
                        if diagram_error:
                            done_payload["diagram_error"] = diagram_error
                        if diagram_specs:
                            done_payload["diagram_specs"] = diagram_specs
                        yield f"data: {json.dumps(done_payload, ensure_ascii=False)}\n\n"
                    elif chunk.get("__stage__"):
                        yield f"data: {json.dumps({'stage': chunk['__stage__']}, ensure_ascii=False)}\n\n"
                elif isinstance(chunk, str):
                    buffer += chunk
                    while True:
                        if not in_think:
                            think_start = buffer.find("<think>")
                            if think_start == -1:
                                safe_len = max(0, len(buffer) - 7)
                                safe = buffer[:safe_len]
                                buffer = buffer[safe_len:]
                                if safe:
                                    full_content += safe
                                    yield f"data: {json.dumps({'text': safe}, ensure_ascii=False)}\n\n"
                                break
                            before = buffer[:think_start]
                            if before:
                                full_content += before
                                yield f"data: {json.dumps({'text': before}, ensure_ascii=False)}\n\n"
                            buffer = buffer[think_start + 7:]
                            in_think = True
                        else:
                            think_end = buffer.find("</think>")
                            if think_end == -1:
                                buffer = ""
                                break
                            buffer = buffer[think_end + 8:]
                            in_think = False
            if buffer and not in_think:
                full_content += buffer
                yield f"data: {json.dumps({'text': buffer}, ensure_ascii=False)}\n\n"
        except Exception as exc:
            logger.error("SSE 流式生成失败 '%s': %s", section_title, exc, exc_info=True)
            yield f"data: {json.dumps({'error': _format_dify_runtime_error(exc)}, ensure_ascii=False)}\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


async def generate_attachment_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """生成商务附件正文；入参为附件生成 JSON，出参兼容 legacy generate-attachment。"""
    payload = _json_object_body(body)
    attachment_type = str(payload.get("attachment_type") or "").strip()
    if not attachment_type:
        raise PlatformError(code="INVALID_REQUEST", message="attachment_type 不能为空。", status_code=400)

    if attachment_type in ATTACHMENT_LABELS:
        context = {
            "org_name": str(payload.get("org_name") or "（投标单位）"),
            "legal_rep": str(payload.get("legal_rep") or "（法定代表人）"),
            "project_lead": str(payload.get("project_lead") or "（项目负责人）"),
            "phone": str(payload.get("phone") or "（电话）"),
            "doc_date": str(payload.get("doc_date") or "____年__月__日"),
            "project_name": str(payload.get("project_name") or "本项目"),
            "recipient": str(payload.get("recipient") or "采购人"),
            "bid_no": str(payload.get("bid_no") or ""),
            "agent_name": str(payload.get("agent_name") or "（被委托人）"),
            "agent_id": str(payload.get("agent_id") or ""),
        }
        try:
            content = render_attachment(attachment_type, context)
        except ValueError as exc:
            raise PlatformError(code="INVALID_REQUEST", message=str(exc), status_code=400) from exc
        return {
            "attachment_type": attachment_type,
            "label": ATTACHMENT_LABELS[attachment_type],
            "content": content,
        }

    dify_key = _get_workflow_key("attachment_generator")
    if not dify_key:
        raise PlatformError(
            code="ATTACHMENT_GENERATE_FAILED",
            message="附件动态生成工作流的 API Key (DIFY_WORKFLOW_ATTACHMENT_GENERATOR) 未配置",
            status_code=400,
        )

    try:
        dify_res = await _call_dify_workflow(
            dify_key,
            {
                "attachment_name": str(payload.get("attachment_name") or "未命名附件"),
                "attachment_desc": str(payload.get("attachment_desc") or ""),
                "project_name": str(payload.get("project_name") or ""),
                "org_name": str(payload.get("org_name") or ""),
                "legal_rep": str(payload.get("legal_rep") or ""),
            },
        )
    except Exception as exc:
        raise PlatformError(code="ATTACHMENT_GENERATE_FAILED", message=_format_dify_runtime_error(exc), status_code=500) from exc

    outputs = dify_res.get("data", {}).get("outputs", {})
    if not isinstance(outputs, Mapping):
        outputs = {}
    content = outputs.get("text") or outputs.get("result") or outputs.get("content") or "AI 工作流未返回内容，请检查 Dify 中的 `text` 输出变量配置。"
    if isinstance(content, list):
        content = "\n\n".join(str(item) for item in content)
    return {
        "attachment_type": attachment_type,
        "label": str(payload.get("attachment_name") or attachment_type),
        "content": str(content),
    }


async def build_scoring_table_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """构建自评评分表；入参为评分模板或需求，出参兼容 legacy build-scoring-table。"""
    payload = _json_object_body(body)
    scoring_table_template = payload.get("scoring_table_template", [])
    score_requirements = payload.get("score_requirements", [])

    if scoring_table_template is not None and not isinstance(scoring_table_template, list):
        raise PlatformError(code="INVALID_REQUEST", message="scoring_table_template 必须是数组。", status_code=400)
    if score_requirements is not None and not isinstance(score_requirements, list):
        raise PlatformError(code="INVALID_REQUEST", message="score_requirements 必须是数组。", status_code=400)

    rows: list[dict[str, Any]] = []
    if scoring_table_template:
        for index, template in enumerate(scoring_table_template):
            if not isinstance(template, Mapping):
                raise PlatformError(code="INVALID_REQUEST", message="scoring_table_template 元素必须是对象。", status_code=400)
            rows.append(
                {
                    "id": str(template.get("id") or f"scored_{index}"),
                    "indicator": str(template.get("indicator") or template.get("name") or f"评分项 {index + 1}"),
                    "max_score": _int_or_default(template.get("max_score"), default=_int_or_default(template.get("points"), default=10)),
                    "criteria": str(template.get("criteria") or template.get("description") or ""),
                }
            )
    else:
        for index, requirement in enumerate(score_requirements or []):
            if not isinstance(requirement, Mapping):
                raise PlatformError(code="INVALID_REQUEST", message="score_requirements 元素必须是对象。", status_code=400)
            rows.append(
                {
                    "id": str(requirement.get("id") or f"score_req_{index}"),
                    "indicator": str(requirement.get("content") or f"评分项 {index + 1}"),
                    "max_score": _int_or_default(requirement.get("points"), default=10),
                    "criteria": "",
                }
            )

    return {"rows": rows}


async def generate_blueprint_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """生成全局蓝图；入参为项目需求和大纲，出参兼容 legacy generate-blueprint。"""
    payload = _json_object_body(body)
    dify_key = _get_workflow_key("blueprint_generator")
    if not dify_key:
        return {
            "blueprint": {
                "positioning": "展示深厚行业经验与技术领先性，打造高性价比方案",
                "strategy": "充分响应招标需求，在附加分项目上寻求突破，提供超出预期的售后保障",
                "highlights": ["自研核心技术的安全可靠性", "行业首创的快速交付模式", "总包一站式闭环服务"],
                "writing_style": "正式、专业、数据驱动",
            }
        }

    requirements = payload.get("requirements", [])
    outline = payload.get("outline", [])
    if requirements is not None and not isinstance(requirements, list):
        raise PlatformError(code="INVALID_REQUEST", message="requirements 必须是数组。", status_code=400)
    if outline is not None and not isinstance(outline, list):
        raise PlatformError(code="INVALID_REQUEST", message="outline 必须是数组。", status_code=400)

    req_summary = "\n".join(
        f"- [{str(item.get('type') or '')}] {str(item.get('content') or '')}"
        for item in requirements[:30]
        if isinstance(item, Mapping)
    )
    outline_summary = "\n".join(
        f"- {str(item.get('title') or '')}"
        for item in outline
        if isinstance(item, Mapping)
    )

    try:
        dify_res = await _call_dify_workflow(
            dify_key,
            {
                "bid_type": str(payload.get("bid_type") or ""),
                "project_summary": str(payload.get("project_summary") or ""),
                "requirements_summary": req_summary,
                "outline_summary": outline_summary,
            },
        )
    except Exception as exc:
        raise PlatformError(code="BLUEPRINT_GENERATE_FAILED", message=_format_dify_runtime_error(exc), status_code=500) from exc

    outputs = dify_res.get("data", {}).get("outputs", {})
    if not isinstance(outputs, Mapping):
        outputs = {}

    if outputs.get("positioning"):
        highlights_list = _json_array_or_empty(outputs.get("highlights"))
        return {
            "blueprint": {
                "positioning": str(outputs.get("positioning") or ""),
                "strategy": str(outputs.get("strategy") or ""),
                "highlights": highlights_list,
                "writing_style": str(outputs.get("writing_style") or "正式、严谨庄重"),
            }
        }

    raw_text = str(outputs.get("text") or outputs.get("result") or outputs.get("content") or "")
    parsed = _extract_json_object(raw_text)
    return {
        "blueprint": {
            "positioning": str(parsed.get("positioning") or "展示高质量、高性价的专业方案"),
            "strategy": str(parsed.get("strategy") or "严格遵守所有要求，提供具有竞争力的优势方案"),
            "highlights": _list_of_strings(parsed.get("highlights"))
            or ["优秀的行业业绩案例", "专业的技术服务团队", "完善的售后保障"],
            "writing_style": str(parsed.get("writing_style") or "正式、严谨庄重"),
        }
    }


async def fill_scoring_row_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """AI 填写评分行；入参为单行评分上下文，出参兼容 legacy fill-scoring-row。"""
    payload = _json_object_body(body)
    dify_key = _get_workflow_key("scoring_assistant") or _get_workflow_key("requirement_extractor")
    if not dify_key:
        raise PlatformError(code="SCORING_ROW_FILL_FAILED", message="未配置 Dify 密钥，无法 AI 填写评分行", status_code=500)

    user_msg = (
        f"评分指标：{str(payload.get('indicator') or '')}\n"
        f"最高分：{_int_or_default(payload.get('max_score'), default=0)} 分\n"
        f"评分标准：{str(payload.get('criteria') or '（未提供具体标准）')}\n"
        f"项目概要：{str(payload.get('project_summary') or '（未提供）')}\n"
        f"其他需求上下文：{str(payload.get('requirements_context') or '（未提供）')}\n\n"
        "请按要求输出 JSON。"
    )

    try:
        dify_res = await _call_dify_workflow(
            dify_key,
            {
                "raw_document": user_msg,
                "_system_override": _SCORING_SYSTEM_PROMPT,
            },
        )
    except Exception as exc:
        raise PlatformError(code="SCORING_ROW_FILL_FAILED", message=_format_dify_runtime_error(exc), status_code=500) from exc

    outputs = dify_res.get("data", {}).get("outputs", {})
    if not isinstance(outputs, Mapping):
        outputs = {}

    row_id = str(payload.get("row_id") or "")
    default_comment = "我方具备相关能力，能够响应本评分项要求。"

    if outputs.get("self_response"):
        self_response = str(outputs.get("self_response") or "partial")
        if self_response not in {"full", "partial"}:
            self_response = "partial"
        return {
            "row_id": row_id,
            "self_response": self_response,
            "self_comment": str(outputs.get("self_comment") or default_comment),
            "evidence_refs": _list_of_strings(_json_array_or_empty(outputs.get("evidence_refs")))[:3],
        }

    raw_text = outputs.get("text") or outputs.get("result") or outputs.get("content") or ""
    if isinstance(raw_text, list):
        raw_text = "\n".join(str(item) for item in raw_text)
    clean_text = re.sub(r"<think>.*?</think>", "", str(raw_text), flags=re.DOTALL).strip()
    code_block_match = re.search(r"```json\s*(.*?)\s*```", clean_text, re.DOTALL)
    parsed = _extract_json_object(code_block_match.group(1) if code_block_match else clean_text)

    self_response = str(parsed.get("self_response") or "partial")
    if self_response not in {"full", "partial"}:
        self_response = "partial"
    return {
        "row_id": row_id,
        "self_response": self_response,
        "self_comment": str(parsed.get("self_comment") or default_comment),
        "evidence_refs": _list_of_strings(parsed.get("evidence_refs"))[:3],
    }


async def analyze_document_response(
    *,
    raw_document: str = "",
    project_id: str = "",
    selected_node_ids: str = "",
) -> Any:
    """流式生成解析报告；入参为原文/项目/节点选择，出参保持 legacy SSE 协议。"""
    normalized_project_id = _ensure_safe_project_id(project_id)
    config_path = _bid_generator_root() / "config" / "analysis_framework.json"
    if not config_path.exists():
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="analysis_framework.json 不存在", status_code=404)
    system_prompt_base, all_nodes = load_docanalysis_framework(config_path)
    if not all_nodes:
        raise PlatformError(code="INVALID_REQUEST", message="框架中无节点", status_code=400)

    dify_key = _get_workflow_key("doc_analysis") or _get_workflow_key("requirement_extractor")
    if not dify_key:
        raise PlatformError(code="ANALYSIS_STREAM_FAILED", message="需求提取工作流 API Key 未配置", status_code=500)

    selected_ids = {node_id.strip() for node_id in str(selected_node_ids or "").split(",") if node_id.strip()} or None
    document_source = str(raw_document or "").strip() or _load_raw_document(normalized_project_id)
    if not document_source:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message="未找到项目原文缓存，请先重新上传并解析文档", status_code=404)
    document_text = document_source[:300000]

    groups = build_docanalysis_groups(all_nodes, selected_ids)
    if not groups:
        raise PlatformError(code="INVALID_REQUEST", message="未找到可提取节点", status_code=400)
    total_nodes = sum(len(group.get("nodes") or []) for group in groups)

    async def sse_generator() -> Any:
        queue: asyncio.Queue = asyncio.Queue()
        success_count = 0
        yield _sse_event(
            "progress",
            {"phase": "analyzing", "message": f"并行解析 {len(groups)} 组 / 共 {total_nodes} 个节点", "total": total_nodes},
        )

        async def run_group(group: dict[str, Any], index: int) -> None:
            try:
                results = await _extract_docanalysis_group_results(
                    system_prompt_base=system_prompt_base,
                    dify_key=dify_key,
                    subset_nodes=group.get("nodes") if isinstance(group.get("nodes"), list) else [],
                    subset_label=str(group.get("group_label") or ""),
                    document_text=document_text,
                )
                await queue.put(("group_done", index, str(group.get("group_label") or ""), results, None))
            except Exception as exc:
                await queue.put(("group_error", index, str(group.get("group_label") or ""), [], str(exc)))

        tasks = [asyncio.create_task(run_group(group, index)) for index, group in enumerate(groups)]
        done_groups = 0
        while done_groups < len(groups):
            event_type, _index, group_label, results, error_message = await queue.get()
            done_groups += 1
            if event_type == "group_error":
                yield _sse_event("error", {"group": group_label, "error": error_message or "提取失败"})
                continue

            yield _sse_event("progress", {"phase": "group_done", "message": f"完成: {group_label} ({done_groups}/{len(groups)})"})
            for result in results:
                yield _sse_event(
                    "node_complete",
                    {
                        "node_id": result["node_id"],
                        "label": result["label"],
                        "content": result["content"],
                    },
                )
                success_count += 1

        await asyncio.gather(*tasks, return_exceptions=True)
        yield _sse_event("complete", {"total_nodes": total_nodes, "success_count": success_count})

    return StreamingResponse(
        sse_generator(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


async def analyze_node_response(project_id: str, body: Mapping[str, Any]) -> Any:
    """流式重提取单个解析节点；入参为项目 ID 和节点上下文，出参保持 legacy SSE 协议。"""
    normalized_project_id = _ensure_safe_project_id(project_id)
    payload = _json_object_body(body)
    node_id = str(payload.get("node_id") or "").strip()
    node_label = str(payload.get("node_label") or "").strip()
    extraction_prompt = str(payload.get("extraction_prompt") or "").strip()
    raw_document = str(payload.get("raw_document") or "").strip()
    if not raw_document:
        raw_document = _load_raw_document(normalized_project_id)
    raw_document = raw_document[:300000]
    if not node_id or not raw_document:
        raise PlatformError(code="INVALID_REQUEST", message="缺少 node_id 或项目原文缓存不存在", status_code=400)

    dify_key = _get_workflow_key("doc_analysis") or _get_workflow_key("requirement_extractor")
    if not dify_key:
        raise PlatformError(code="ANALYSIS_NODE_STREAM_FAILED", message="工作流 API Key 未配置", status_code=500)

    config_path = _bid_generator_root() / "config" / "analysis_framework.json"
    system_prompt_base = ""
    node_by_id: dict[str, dict] = {}
    if config_path.exists():
        system_prompt_base, all_nodes = load_docanalysis_framework(config_path)
        node_by_id = build_docanalysis_node_index(all_nodes)
    node_def = node_by_id.get(node_id) or {}
    node_label = str(node_def.get("label") or node_label or "").strip() or node_id
    extraction_prompt = str(node_def.get("extractionPrompt") or extraction_prompt or "").strip()
    if not extraction_prompt:
        raise PlatformError(code="INVALID_REQUEST", message="缺少 extraction_prompt，且未在解析框架中找到该节点定义", status_code=400)

    async def event_generator() -> Any:
        prompt_nodes = [{"id": node_id, "label": node_label, "extractionPrompt": extraction_prompt}]
        combined_system = build_docanalysis_system_prompt(system_prompt_base, prompt_nodes, f"单节点重提取：{node_label}")
        try:
            outputs: dict[str, Any] = {}
            got_finished = False
            async for chunk in _call_dify_workflow_stream(
                dify_key,
                {
                    "system_prompt": combined_system,
                    "raw_document": raw_document,
                    "node_label": node_label,
                },
            ):
                if isinstance(chunk, dict) and chunk.get("__finished__"):
                    got_finished = True
                    outputs = chunk.get("outputs", {}) or {}
                    break
            if not got_finished:
                raise RuntimeError("解析工作流异常中断（未收到 finished 事件）")

            raw_text = extract_docanalysis_text_output(outputs)
            content_text, attachments_payload = split_bid_attachments_tag(raw_text)
            result_map = parse_docanalysis_result_map(content_text)
            content = extract_docanalysis_node_content(result_map, node_id)
            if isinstance(content, (dict, list)):
                content = json.dumps(content, ensure_ascii=False, indent=2)
            bid_items = parse_bid_attachments_payload(attachments_payload)
            if bid_items:
                yield _sse_data({"type": "bid_attachments", "items": bid_items})
            yield _sse_data({"type": "done", "node_id": node_id, "content": str(content).strip()})
        except Exception as exc:
            logger.error("单节点提取 SSE 失败 [%s]: %s", node_id, exc)
            yield _sse_data({"type": "error", "message": str(exc)})

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


async def extract_bid_attachment_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """按定位符提取 DOCX 附件 HTML；入参为定位符范围，出参兼容 legacy attachment extract。"""
    payload = _json_object_body(body)
    project_id = _ensure_safe_project_id(str(payload.get("project_id") or ""))
    start_locator = _normalize_locator_token(str(payload.get("start_locator") or ""))
    end_locator = _normalize_locator_token(str(payload.get("end_locator") or ""))
    attachment_name = str(payload.get("attachment_name") or "").strip()
    if not start_locator or not end_locator:
        raise PlatformError(code="INVALID_REQUEST", message="start_locator 和 end_locator 不能为空", status_code=400)

    blocks = get_project_doc_blocks_payload(project_id)["blocks"]
    start_block = _find_doc_block_by_locator(blocks, start_locator)
    end_block = _find_doc_block_by_locator(blocks, end_locator)
    if start_block is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message=f"定位符 {start_locator} 未找到", status_code=404)
    if end_block is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message=f"定位符 {end_locator} 未找到", status_code=404)

    start_idx = _non_negative_int(start_block.get("body_idx"))
    end_idx = _non_negative_int(end_block.get("body_idx"))
    html_text = _doc_blocks_slice_to_html(blocks, start_idx, end_idx)
    if not html_text:
        raise PlatformError(
            code="BUSINESS_DIRECT_ERROR",
            message="当前仅恢复了文档块索引快照，且快照正文为空；请上传原始 DOCX 执行“重建定位”后重试",
            status_code=409,
        )
    if start_idx > end_idx:
        start_idx, end_idx = end_idx, start_idx

    return {
        "html": html_text,
        "attachment_name": attachment_name,
        "start_locator": start_locator,
        "end_locator": end_locator,
        "resolved_start_locator": start_locator,
        "resolved_end_locator": end_locator,
        "paragraph_count": end_idx - start_idx + 1,
        "snapshot_only": True,
    }


async def test_locators_payload(project_id: str) -> dict[str, Any]:
    """查看项目 DOCX 定位符映射；入参为项目 ID，出参兼容 legacy test-locators。"""
    blocks_payload = get_project_doc_blocks_payload(project_id)
    blocks = blocks_payload["blocks"]
    preview: list[dict[str, Any]] = []
    locator_count = 0
    for block in blocks:
        if not isinstance(block, Mapping):
            continue
        locator = str(block.get("locator") or "").strip().upper()
        if not locator:
            continue
        locator_count += 1
        if len(preview) >= 20:
            continue
        preview.append(
            {
                "locator": locator,
                "body_idx": _non_negative_int(block.get("body_idx")),
                "snippet": str(block.get("text") or "")[:60],
            }
        )

    return {
        "project_id": blocks_payload["project_id"],
        "total_locators": locator_count,
        "preview": preview,
        "snapshot_only": True,
    }


async def rebuild_locator_payload(project_id: str, file: UploadFile) -> dict[str, Any]:
    """重建 DOCX 定位缓存；入参为项目 ID 和 DOCX 文件，出参兼容 legacy rebuild-locator。"""
    normalized_id = _ensure_safe_project_id(project_id)
    filename = str(getattr(file, "filename", "") or "").lower()
    if not filename.endswith(".docx"):
        raise PlatformError(code="INVALID_REQUEST", message="仅支持上传 DOCX 文件", status_code=400)
    content = await file.read()
    if not content:
        raise PlatformError(code="INVALID_REQUEST", message="上传文件为空", status_code=400)

    try:
        doc_blocks = _extract_docx_blocks(content)
        _persist_docx_cache(normalized_id, content)
        _persist_project_doc_blocks_snapshot(project_id=normalized_id, doc_blocks=doc_blocks)
    except PlatformError:
        raise
    except Exception as exc:
        logger.error("[%s] 重建定位缓存失败: %s", normalized_id, exc, exc_info=True)
        raise PlatformError(code="BID_LOCATOR_REBUILD_FAILED", message=f"重建定位缓存失败: {exc}", status_code=500) from exc

    locator_count = sum(1 for block in doc_blocks if str(block.get("locator") or "").strip())
    logger.info("[%s] 重建定位缓存成功: %s 个文档块", normalized_id, len(doc_blocks))
    return {
        "project_id": normalized_id,
        "blocks": len(doc_blocks),
        "locators": locator_count,
        "snapshot_only": True,
    }


async def extract_bid_attachment_by_block_payload(body: Mapping[str, Any]) -> dict[str, Any]:
    """按 block_id 提取 DOCX 附件 HTML；入参为 block 范围，出参兼容 legacy block extract。"""
    payload = _json_object_body(body)
    project_id = _ensure_safe_project_id(str(payload.get("project_id") or ""))
    start_block_id = str(payload.get("start_block_id") or "").strip()
    end_block_id = str(payload.get("end_block_id") or "").strip()
    attachment_name = str(payload.get("attachment_name") or "").strip()
    if not start_block_id or not end_block_id:
        raise PlatformError(code="INVALID_REQUEST", message="project_id/start_block_id/end_block_id 不能为空", status_code=400)

    blocks = get_project_doc_blocks_payload(project_id)["blocks"]
    start_block = _find_doc_block_by_id(blocks, start_block_id)
    end_block = _find_doc_block_by_id(blocks, end_block_id)
    if start_block is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message=f"block_id {start_block_id} 未找到", status_code=404)
    if end_block is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message=f"block_id {end_block_id} 未找到", status_code=404)

    start_idx = _non_negative_int(start_block.get("body_idx"))
    end_idx = _non_negative_int(end_block.get("body_idx"))
    if start_idx > end_idx:
        start_idx, end_idx = end_idx, start_idx
        start_block_id, end_block_id = end_block_id, start_block_id
    html_text = _doc_blocks_slice_to_html(blocks, start_idx, end_idx)
    if not html_text:
        raise PlatformError(
            code="BUSINESS_DIRECT_ERROR",
            message="当前仅恢复了文档块索引快照，且快照正文为空；请上传原始 DOCX 执行“重建定位”后重试",
            status_code=409,
        )

    return {
        "html": html_text,
        "attachment_name": attachment_name,
        "start_block_id": start_block_id,
        "end_block_id": end_block_id,
        "paragraph_count": end_idx - start_idx + 1,
        "snapshot_only": True,
    }


async def extract_bid_attachment_by_block_docx_response(body: Mapping[str, Any]) -> Any:
    """按 block_id 返回 DOCX 切片；入参为 block 范围，出参保持 legacy 二进制响应。"""
    payload = _json_object_body(body)
    project_id = _ensure_safe_project_id(str(payload.get("project_id") or ""))
    start_block_id = str(payload.get("start_block_id") or "").strip()
    end_block_id = str(payload.get("end_block_id") or "").strip()
    attachment_name = str(payload.get("attachment_name") or "").strip() or "slice"
    if not start_block_id or not end_block_id:
        raise PlatformError(code="INVALID_REQUEST", message="project_id/start_block_id/end_block_id 不能为空", status_code=400)

    blocks = get_project_doc_blocks_payload(project_id)["blocks"]
    start_block = _find_doc_block_by_id(blocks, start_block_id)
    end_block = _find_doc_block_by_id(blocks, end_block_id)
    if start_block is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message=f"block_id {start_block_id} 未找到", status_code=404)
    if end_block is None:
        raise PlatformError(code="RESOURCE_NOT_FOUND", message=f"block_id {end_block_id} 未找到", status_code=404)

    start_idx = _non_negative_int(start_block.get("body_idx"))
    end_idx = _non_negative_int(end_block.get("body_idx"))
    if start_idx > end_idx:
        start_idx, end_idx = end_idx, start_idx
        start_block_id, end_block_id = end_block_id, start_block_id

    docx_path = _docx_cache_path(project_id)
    if not docx_path.exists():
        raise PlatformError(
            code="BUSINESS_DIRECT_ERROR",
            message="原始 DOCX 不可用，无法生成保格式切片；请上传原始 DOCX 重建定位缓存",
            status_code=409,
        )
    try:
        source_bytes = docx_path.read_bytes()
        sliced_bytes = _slice_docx_bytes_by_body_range(source_bytes, start_idx, end_idx)
    except PlatformError:
        raise
    except Exception as exc:
        logger.error("[%s] DOCX 切片失败: %s", project_id, exc, exc_info=True)
        raise PlatformError(code="BID_ATTACHMENT_BLOCK_DOCX_EXTRACT_FAILED", message=f"DOCX 切片失败: {exc}", status_code=500) from exc

    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", attachment_name).strip("_") or "slice"
    filename = f"{safe_name}_{start_block_id}_{end_block_id}.docx"
    return BidGeneratorFilePayload(
        content=sliced_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        inline=True,
        cache_control="public, max-age=3600",
        headers={
            "X-Start-Block-Id": start_block_id,
            "X-End-Block-Id": end_block_id,
        },
    )


def _task_status_to_api_state(status: str) -> str:
    return {
        "running": "running",
        "done": "succeeded",
        "error": "failed",
        "cancelled": "cancelled",
        "timeout": "timed_out",
    }.get(str(status or ""), "failed")


def _task_started_datetime(task: Any) -> datetime:
    value = getattr(task, "started_at", None)
    if isinstance(value, datetime):
        return value
    raw = getattr(task, "created_at", None)
    try:
        return datetime.fromtimestamp(float(raw), tz=timezone.utc)
    except (TypeError, ValueError, OSError):
        return datetime.now(timezone.utc)


def _elapsed_seconds(started_at: datetime) -> int:
    now = datetime.now(started_at.tzinfo or timezone.utc)
    return max(0, int((now - started_at).total_seconds()))


def _outline_stage_meta_from_label(stage: str) -> tuple[int, int]:
    text_value = str(stage or "")
    if "模型连接中" in text_value or "模型预热中" in text_value:
        return 0, 3
    if "生成大纲" in text_value:
        return 2, 12
    if "大纲润色" in text_value:
        return 3, 75
    if "数据校验" in text_value or "解析中" in text_value:
        return 4, 86
    if "归一化中" in text_value:
        return 5, 94
    if "结构已就绪" in text_value:
        return 6, 100
    return 0, 0


def _sse_event(event: str, payload: Any) -> str:
    return f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n"


def _sse_data(payload: Any) -> str:
    return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"


def _task_stage_sse_events(task_id: str, raw_stage: Any, index: int, started_at: datetime) -> list[str]:
    stage = str(raw_stage or "")
    if stage.startswith("__text__"):
        return [_sse_data({"text": stage[8:]})]
    if stage.startswith("__node__"):
        try:
            return [_sse_event("node_complete", json.loads(stage[8:]))]
        except json.JSONDecodeError:
            return []
    if stage.startswith(_BID_ATTACH_STAGE_PREFIX):
        try:
            return [_sse_event("bid_attachments", json.loads(stage[len(_BID_ATTACH_STAGE_PREFIX):]))]
        except json.JSONDecodeError as exc:
            logger.warning("[BidTask %s] bid_attachments SSE replay skipped: %s", task_id, exc)
            return []
    if stage.startswith(_ANALYSIS_V2_STAGE_PREFIX):
        try:
            return [_sse_event("analysis_v2", json.loads(stage[len(_ANALYSIS_V2_STAGE_PREFIX):]))]
        except json.JSONDecodeError as exc:
            logger.warning("[BidTask %s] analysis_v2 SSE replay skipped: %s", task_id, exc)
            return []
    if stage.startswith(_TASK_EVENT_STAGE_PREFIX):
        try:
            payload = json.loads(stage[len(_TASK_EVENT_STAGE_PREFIX):])
        except json.JSONDecodeError as exc:
            logger.warning("[BidTask %s] task_event SSE replay skipped: %s", task_id, exc)
            return []
        if isinstance(payload, dict):
            payload["event_id"] = f"{task_id}:{index}"
            event_name = str(payload.get("event") or "task_event")
            return [_sse_event(event_name, payload)]
        return []
    phase, percent = _outline_stage_meta_from_label(stage)
    stage_payload = {
        "event_id": f"{task_id}:{index}",
        "stage": stage,
        "phase": phase,
        "percent": percent,
        "elapsed_sec": _elapsed_seconds(started_at),
    }
    return [_sse_event("stage", stage_payload), _sse_data(stage_payload)]


def _task_terminal_sse_events(task: Any) -> list[str]:
    status = str(getattr(task, "status", "") or "")
    if status == "done":
        payload = getattr(task, "result", None)
        return [_sse_event("done", payload), _sse_data(payload)]
    if status == "error":
        payload = {"error": getattr(task, "error", None)}
        return [_sse_event("error", payload), _sse_data(payload)]
    if status == "timeout":
        payload = {"error": getattr(task, "error", None), "timed_out": True}
        return [_sse_event("error", payload), _sse_data(payload)]
    if status == "cancelled":
        payload = {"cancelled": True}
        return [_sse_event("cancelled", payload), _sse_data(payload)]
    return []


def _ensure_safe_kb_sync_job_id(job_id: str) -> str:
    normalized = str(job_id or "").strip()
    if not re.fullmatch(r"[a-f0-9]{12}", normalized):
        raise PlatformError(code="INVALID_REQUEST", message="无效的 job_id 格式", status_code=400)
    return normalized


def _escape_svg_text(text_value: str) -> str:
    return (
        str(text_value or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _mermaid_to_fallback_svg(mermaid: str, title: str = "数据流图") -> str:
    lines = [line.strip() for line in str(mermaid or "").splitlines() if line.strip()]
    body_lines = [line for line in lines if not re.match(r"^(?:flowchart|graph)\s+", line, flags=re.IGNORECASE)]
    if not body_lines:
        body_lines = ["Mermaid 图表源码已生成"]
    body_lines = body_lines[:18]
    width = 1120
    row_h = 30
    height = max(180, 92 + len(body_lines) * row_h)
    escaped_title = _escape_svg_text(title or "数据流图")
    rows = []
    for idx, line in enumerate(body_lines):
        y = 88 + idx * row_h
        rows.append(
            f'<text x="40" y="{y}" font-size="16" fill="#334155" font-family="monospace">{_escape_svg_text(line[:118])}</text>'
        )
    footer_y = height - 28
    return (
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" viewBox="0 0 {width} {height}">'
        '<rect width="100%" height="100%" rx="16" fill="#f8fafc"/>'
        '<rect x="24" y="22" width="1072" height="44" rx="10" fill="#e0f2fe" stroke="#bae6fd"/>'
        f'<text x="40" y="50" font-size="20" font-weight="700" fill="#0369a1" font-family="Arial, sans-serif">{escaped_title}</text>'
        f'{"".join(rows)}'
        f'<text x="40" y="{footer_y}" font-size="13" fill="#64748b" font-family="Arial, sans-serif">Mermaid 源码预览；导出 DOCX 时会渲染为正式图片。</text>'
        "</svg>"
    )


def _ensure_safe_image_filename(filename: str) -> str:
    normalized = str(filename or "").strip()
    if not re.fullmatch(r"[a-zA-Z0-9_.()-]+", normalized):
        raise PlatformError(code="INVALID_REQUEST", message="无效的文件名格式", status_code=400)
    return normalized


def _utc_now() -> datetime:
    return datetime.now(timezone.utc)


def _utc_iso_from_timestamp(value: Any) -> str:
    try:
        timestamp = float(value or datetime.now(timezone.utc).timestamp())
    except (TypeError, ValueError):
        timestamp = datetime.now(timezone.utc).timestamp()
    return datetime.fromtimestamp(timestamp, timezone.utc).replace(tzinfo=None).isoformat()


def _required_string(value: Any, *, field: str) -> str:
    normalized = str(value or "").strip()
    if not normalized:
        raise PlatformError(code="INVALID_REQUEST", message=f"{field} 不能为空。", status_code=400)
    return normalized


def _json_object_body(body: Mapping[str, Any]) -> dict[str, Any]:
    return dict(body) if isinstance(body, Mapping) else {}


def _json_array_or_empty(value: Any) -> list[Any]:
    if isinstance(value, list):
        return value
    if isinstance(value, str):
        try:
            loaded = json.loads(value)
        except json.JSONDecodeError:
            return []
        return loaded if isinstance(loaded, list) else []
    return []


def _extract_json_object(value: str) -> dict[str, Any]:
    raw = str(value or "").strip()
    if not raw:
        return {}
    match = re.search(r"\{.*\}", raw, re.DOTALL)
    if not match:
        return {}
    try:
        loaded = json.loads(match.group())
    except json.JSONDecodeError:
        return {}
    return dict(loaded) if isinstance(loaded, Mapping) else {}


def _model_or_mapping_to_dict(value: Any) -> dict[str, Any]:
    if isinstance(value, Mapping):
        return dict(value)
    model_dump = getattr(value, "model_dump", None)
    if callable(model_dump):
        dumped = model_dump()
        return dict(dumped) if isinstance(dumped, Mapping) else {"data": dumped}
    legacy_dict = getattr(value, "dict", None)
    if callable(legacy_dict):
        dumped = legacy_dict()
        return dict(dumped) if isinstance(dumped, Mapping) else {"data": dumped}
    return {"data": value}


def _optional_string(value: Any) -> str | None:
    if value is None:
        return None
    return str(value)


def _int_or_default(value: Any, *, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return int(default)


def _list_of_strings(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def _string_or_default(value: Any, *, default: str) -> str:
    normalized = _optional_string(value)
    return normalized if normalized is not None else default


def _dict_or_default(value: Any) -> dict[str, Any]:
    if value is None:
        return {}
    if not isinstance(value, dict):
        raise PlatformError(code="INVALID_REQUEST", message="data 必须是对象。", status_code=400)
    return value


def _string_mapping(value: Any) -> dict[str, str]:
    if not isinstance(value, dict):
        return {}
    return {str(key): str(item) for key, item in value.items() if str(key).strip()}


def _deep_merge_dict(base: dict[str, Any], patch: dict[str, Any]) -> dict[str, Any]:
    result = dict(base or {})
    for key, value in (patch or {}).items():
        if key in result and isinstance(result.get(key), dict) and isinstance(value, dict):
            result[key] = _deep_merge_dict(result[key], value)
        else:
            result[key] = value
    return result


def _json_detail_value(value: Any) -> Any:
    if value is None:
        return {}
    if isinstance(value, str):
        try:
            return json.loads(value)
        except json.JSONDecodeError:
            return {}
    return value


def list_pipt_audit_logs_payload(
    *,
    project_id: str | None = None,
    task_id: str | None = None,
    session_id: str | None = None,
    operation: str | None = None,
    status: str | None = None,
    placeholder: str | None = None,
    limit: int = 100,
) -> dict[str, Any]:
    """查询 PIPT 脱敏识别/回映射审计日志。返回值不包含敏感明文，仅包含 hash 与结构化上下文。"""
    bounded_limit = max(1, min(int(limit or 100), 500))
    filters = []
    params: dict[str, Any] = {"limit": bounded_limit}
    for key, value in (
        ("project_id", project_id),
        ("task_id", task_id),
        ("session_id", session_id),
        ("operation", operation),
        ("status", status),
        ("placeholder", placeholder),
    ):
        normalized = str(value or "").strip()
        if not normalized:
            continue
        filters.append(f"{key} = :{key}")
        params[key] = normalized
    where_sql = f"WHERE {' AND '.join(filters)}" if filters else ""
    try:
        with get_engine().begin() as conn:
            exists = conn.execute(text("SELECT to_regclass('bid_generator.pipt_audit_logs') IS NOT NULL")).scalar_one()
            if not exists:
                raise PlatformError(
                    code="DATABASE_ERROR",
                    message="PIPT 审计日志表不存在，请先执行数据库迁移。",
                    status_code=500,
                    details={"table": "bid_generator.pipt_audit_logs"},
                )
            rows = conn.execute(
                text(
                    f"""
                    SELECT id, operation, status, source, session_id, project_id, task_id,
                           placeholder, entity_type, original_hash, text_hash, details, created_at
                    FROM bid_generator.pipt_audit_logs
                    {where_sql}
                    ORDER BY created_at DESC
                    LIMIT :limit
                    """
                ),
                params,
            ).mappings().all()
    except PlatformError:
        raise
    except (SQLAlchemyError, RuntimeError) as exc:
        raise _database_error(exc) from exc

    items = [
        {
            "id": str(row["id"]),
            "operation": str(row["operation"]),
            "status": str(row["status"]),
            "source": str(row.get("source") or ""),
            "session_id": row.get("session_id"),
            "project_id": row.get("project_id"),
            "task_id": row.get("task_id"),
            "placeholder": row.get("placeholder"),
            "entity_type": row.get("entity_type"),
            "original_hash": row.get("original_hash"),
            "text_hash": row.get("text_hash"),
            "details": _json_detail_value(row.get("details")),
            "created_at": _iso_value(row.get("created_at")),
        }
        for row in rows
    ]
    return {"items": items, "count": len(items), "limit": bounded_limit}
